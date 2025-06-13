"""Service d'apprentissage du style de rédaction depuis des documents Word/PDF"""

import re
import statistics
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import docx
import PyPDF2
from collections import Counter
import streamlit as st
from dataclasses import dataclass, field

# ========================= STRUCTURES DE DONNÉES =========================

@dataclass
class StyleLearningResult:
    """Résultat de l'apprentissage du style"""
    style_name: str
    documents_analyzed: int
    average_sentence_length: int = 0
    average_paragraph_length: int = 0
    paragraph_numbering_style: Optional[str] = None
    paragraph_numbering_pattern: Optional[str] = None
    formality_score: float = 0.5
    technical_terms_frequency: float = 0.0
    transition_words: List[str] = field(default_factory=list)
    argument_patterns: List[str] = field(default_factory=list)
    citation_patterns: List[str] = field(default_factory=list)
    conclusion_patterns: List[str] = field(default_factory=list)
    common_phrases: List[str] = field(default_factory=list)
    use_bold: bool = False
    use_italic: bool = False
    heading_style: str = "Normal"
    confidence_score: float = 0.5

@dataclass
class Document:
    """Document pour l'analyse de style"""
    id: str
    title: str
    content: str
    source: str
    metadata: Dict = field(default_factory=dict)

# ========================= SERVICE PRINCIPAL =========================

class StyleLearningService:
    """Service pour apprendre le style de rédaction depuis des documents"""
    
    def __init__(self):
        self.learned_styles = {}  # Cache des styles appris
        
        # Patterns de numérotation
        self.numbering_patterns = {
            'numeric': r'^\d+\.',
            'numeric_paren': r'^\d+\)',
            'roman_upper': r'^[IVX]+\.',
            'roman_lower': r'^[ivx]+\.',
            'alpha_upper': r'^[A-Z]\.',
            'alpha_lower': r'^[a-z]\.',
            'hierarchical': r'^\d+\.\d+',
            'section': r'^§\s*\d+',
            'dash': r'^-\s',
            'bullet': r'^[•·▪▫]\s'
        }
        
        # Indicateurs de formalité
        self.formality_indicators = {
            'tres_formel': [
                'attendu que', 'considérant que', 'il appert', 'nonobstant',
                'aux termes de', 'en l\'espèce', 'il échet', 'partant',
                'au demeurant', 'quant à ce', 'force est de constater',
                'il convient de relever', 'à toutes fins utiles'
            ],
            'formel': [
                'en effet', 'par ailleurs', 'en outre', 'toutefois',
                'néanmoins', 'cependant', 'dès lors', 'ainsi',
                'par conséquent', 'il résulte', 'il s\'ensuit'
            ],
            'moderne': [
                'donc', 'mais', 'car', 'parce que', 'puisque',
                'c\'est pourquoi', 'en fait', 'notamment'
            ]
        }
        
        # Termes juridiques techniques
        self.technical_terms = [
            'assignation', 'citation', 'conclusions', 'dispositif',
            'moyens', 'prétentions', 'grief', 'cassation', 'appel',
            'intimé', 'appelant', 'demandeur', 'défendeur', 'magistrat',
            'juridiction', 'compétence', 'recevabilité', 'prescription',
            'forclusion', 'péremption', 'déchéance', 'nullité',
            'inopposabilité', 'chose jugée', 'autorité', 'exequatur'
        ]
        
        # Structures d'argumentation typiques
        self.argument_structures = {
            'classique': [
                'En fait', 'En droit', 'Par ces motifs'
            ],
            'moderne': [
                'Les faits', 'L\'analyse juridique', 'Les demandes'
            ],
            'detaille': [
                'Exposé des faits', 'Discussion', 'Sur la procédure',
                'Sur le fond', 'Sur les demandes', 'Par ces motifs'
            ]
        }
    
    async def learn_from_file(self, file_path: str, file_type: str) -> StyleLearningResult:
        """Apprend le style depuis un fichier"""
        content = ""
        
        if file_type == 'docx':
            content = self._extract_from_docx(file_path)
        elif file_type == 'pdf':
            content = self._extract_from_pdf(file_path)
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        if not content:
            raise ValueError("Impossible d'extraire le contenu du fichier")
        
        # Créer un document temporaire pour l'analyse
        doc = Document(
            id=f"temp_{datetime.now().timestamp()}",
            title=f"Document d'apprentissage",
            content=content,
            source=file_path
        )
        
        return await self.learn_from_documents([doc])
    
    async def learn_from_documents(self, documents: List[Document]) -> StyleLearningResult:
        """Apprend le style depuis plusieurs documents"""
        if not documents:
            raise ValueError("Aucun document fourni")
        
        result = StyleLearningResult(
            style_name=f"Style personnalisé ({len(documents)} documents)",
            documents_analyzed=len(documents)
        )
        
        # Analyser chaque aspect du style
        self._analyze_structure(documents, result)
        self._analyze_numbering(documents, result)
        self._analyze_vocabulary(documents, result)
        self._analyze_argumentation(documents, result)
        self._analyze_formatting(documents, result)
        self._extract_common_phrases(documents, result)
        
        # Calculer le score de confiance
        result.confidence_score = self._calculate_confidence(result)
        
        # Sauvegarder le style appris
        style_id = f"learned_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        self.learned_styles[style_id] = result
        
        return result
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extrait le texte d'un fichier Word"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                # Conserver la structure des paragraphes
                if paragraph.text.strip():
                    content.append(paragraph.text)
                else:
                    content.append("")  # Ligne vide pour séparer
            
            return "\n".join(content)
        except Exception as e:
            st.error(f"Erreur lors de la lecture du fichier Word: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extrait le texte d'un fichier PDF"""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
            
            return "\n".join(content)
        except Exception as e:
            st.error(f"Erreur lors de la lecture du fichier PDF: {e}")
            return ""
    
    def _analyze_structure(self, documents: List[Document], result: StyleLearningResult):
        """Analyse la structure générale des documents"""
        all_sentences = []
        all_paragraphs = []
        
        for doc in documents:
            # Séparer en phrases
            sentences = re.split(r'[.!?]+', doc.content)
            sentences = [s.strip() for s in sentences if s.strip()]
            all_sentences.extend(sentences)
            
            # Séparer en paragraphes
            paragraphs = doc.content.split('\n\n')
            paragraphs = [p.strip() for p in paragraphs if p.strip()]
            all_paragraphs.extend(paragraphs)
        
        # Calculer les moyennes
        if all_sentences:
            sentence_lengths = [len(s.split()) for s in all_sentences]
            result.average_sentence_length = int(statistics.mean(sentence_lengths))
        
        if all_paragraphs:
            paragraph_lengths = [len(p.split()) for p in all_paragraphs]
            result.average_paragraph_length = int(statistics.mean(paragraph_lengths))
    
    def _analyze_numbering(self, documents: List[Document], result: StyleLearningResult):
        """Analyse le style de numérotation des paragraphes"""
        numbering_counts = Counter()
        
        for doc in documents:
            lines = doc.content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Tester chaque pattern
                for style, pattern in self.numbering_patterns.items():
                    if re.match(pattern, line):
                        numbering_counts[style] += 1
                        break
        
        if numbering_counts:
            # Style dominant
            dominant_style = numbering_counts.most_common(1)[0][0]
            result.paragraph_numbering_style = dominant_style
            
            # Pattern correspondant
            patterns_map = {
                'numeric': '1.',
                'numeric_paren': '1)',
                'roman_upper': 'I.',
                'roman_lower': 'i.',
                'alpha_upper': 'A.',
                'alpha_lower': 'a.',
                'hierarchical': '1.1',
                'section': '§ 1',
                'dash': '-',
                'bullet': '•'
            }
            result.paragraph_numbering_pattern = patterns_map.get(dominant_style, '1.')
    
    def _analyze_vocabulary(self, documents: List[Document], result: StyleLearningResult):
        """Analyse le vocabulaire et le niveau de formalité"""
        text_combined = ' '.join(doc.content.lower() for doc in documents)
        total_words = len(text_combined.split())
        
        # Analyser la formalité
        formality_scores = {}
        for level, indicators in self.formality_indicators.items():
            count = sum(1 for indicator in indicators if indicator in text_combined)
            formality_scores[level] = count
        
        # Déterminer le niveau dominant
        if formality_scores['tres_formel'] > 5:
            result.formality_score = 0.9
        elif formality_scores['formel'] > 10:
            result.formality_score = 0.7
        else:
            result.formality_score = 0.5
        
        # Analyser les termes techniques
        technical_count = sum(1 for term in self.technical_terms if term in text_combined)
        result.technical_terms_frequency = technical_count / max(total_words, 1) * 100
        
        # Extraire les mots de transition utilisés
        all_transitions = []
        for indicators in self.formality_indicators.values():
            for indicator in indicators:
                if indicator in text_combined:
                    all_transitions.append(indicator)
        
        result.transition_words = list(set(all_transitions))
    
    def _analyze_argumentation(self, documents: List[Document], result: StyleLearningResult):
        """Analyse la structure d'argumentation"""
        text_combined = ' '.join(doc.content for doc in documents)
        
        # Identifier les patterns d'argumentation
        argument_patterns = []
        
        # Patterns d'introduction
        intro_patterns = [
            r'[Ii]l convient de (?:rappeler|préciser|relever|souligner)',
            r'[Ii]l résulte de (?:ce qui précède|l\'instruction|la procédure)',
            r'[Ii]l apparaît que',
            r'[Ff]orce est de constater',
            r'[Ii]l est constant que',
            r'[Àà] titre (?:principal|subsidiaire|liminaire)',
            r'[Ee]n premier lieu',
            r'[Dd]\'une part',
            r'[Tt]out d\'abord'
        ]
        
        for pattern in intro_patterns:
            if re.search(pattern, text_combined):
                argument_patterns.append(pattern.replace('(?:', '').replace(')', '').replace('|', '/'))
        
        result.argument_patterns = argument_patterns
        
        # Identifier les patterns de citation
        citation_patterns = []
        
        # Patterns de citation de jurisprudence
        if re.search(r'Cass\.?\s+\w+\.?\s+\d+', text_combined):
            citation_patterns.append("Cass. [chambre] [date], n°[numero]")
        if re.search(r'CA\s+\w+\s+\d+', text_combined):
            citation_patterns.append("CA [ville] [date], n°[numero]")
        if re.search(r'aux termes de l\'article', text_combined):
            citation_patterns.append("aux termes de l'article [numero] du [code]")
        
        result.citation_patterns = citation_patterns
        
        # Patterns de conclusion
        conclusion_patterns = []
        concl_markers = [
            'Par ces motifs',
            'En conséquence',
            'Au regard de ce qui précède',
            'Il en résulte que',
            'C\'est pourquoi',
            'Ainsi'
        ]
        
        for marker in concl_markers:
            if marker in text_combined:
                conclusion_patterns.append(marker)
        
        result.conclusion_patterns = conclusion_patterns
    
    def _analyze_formatting(self, documents: List[Document], result: StyleLearningResult):
        """Analyse la mise en forme (gras, italique, etc.)"""
        # Pour les fichiers Word, on pourrait extraire ces infos
        # Pour l'instant, on se base sur des conventions
        
        for doc in documents:
            # Détecter l'usage de majuscules pour les titres
            lines = doc.content.split('\n')
            for line in lines:
                if line.isupper() and len(line.split()) > 2:
                    result.use_bold = True
                    result.heading_style = "MAJUSCULES"
                    break
            
            # Détecter l'usage de tirets ou astérisques pour l'emphase
            if '*' in doc.content or '_' in doc.content:
                result.use_italic = True
    
    def _extract_common_phrases(self, documents: List[Document], result: StyleLearningResult):
        """Extrait les phrases et formulations récurrentes"""
        # Extraire les débuts de phrases
        sentence_starters = []
        
        for doc in documents:
            sentences = re.split(r'[.!?]+', doc.content)
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence.split()) > 5:
                    # Prendre les 3-5 premiers mots
                    words = sentence.split()[:5]
                    starter = ' '.join(words)
                    sentence_starters.append(starter)
        
        # Compter les occurrences
        starter_counts = Counter(sentence_starters)
        
        # Garder les plus fréquents (au moins 2 occurrences)
        common_starters = [
            starter for starter, count in starter_counts.most_common(20)
            if count >= 2
        ]
        
        result.common_phrases = common_starters[:10]
    
    def _calculate_confidence(self, result: StyleLearningResult) -> float:
        """Calcule un score de confiance pour l'apprentissage"""
        confidence = 0.5
        
        # Plus de documents = plus de confiance
        if result.documents_analyzed >= 5:
            confidence += 0.2
        elif result.documents_analyzed >= 3:
            confidence += 0.1
        
        # Patterns cohérents = plus de confiance
        if result.paragraph_numbering_style:
            confidence += 0.1
        
        if len(result.argument_patterns) >= 3:
            confidence += 0.1
        
        if len(result.common_phrases) >= 5:
            confidence += 0.1
        
        return min(confidence, 1.0)
    
    def apply_style_to_text(self, text: str, style: StyleLearningResult) -> str:
        """Applique un style appris à un texte"""
        # Appliquer la numérotation des paragraphes
        if style.paragraph_numbering_pattern:
            text = self._apply_paragraph_numbering(text, style)
        
        # Appliquer les phrases types au début
        if style.common_phrases and not any(phrase in text[:200] for phrase in style.common_phrases):
            # Suggérer d'utiliser une phrase type
            suggestion = f"Suggestion de début: {style.common_phrases[0]}..."
            text = f"[{suggestion}]\n\n{text}"
        
        # Appliquer la mise en forme
        if style.use_bold and style.heading_style == "MAJUSCULES":
            # Mettre les titres en majuscules
            lines = text.split('\n')
            formatted_lines = []
            for line in lines:
                if line.strip() and len(line.split()) <= 5 and not any(char in line for char in '.!?'):
                    # Probablement un titre
                    formatted_lines.append(line.upper())
                else:
                    formatted_lines.append(line)
            text = '\n'.join(formatted_lines)
        
        return text
    
    def _apply_paragraph_numbering(self, text: str, style: StyleLearningResult) -> str:
        """Applique le style de numérotation aux paragraphes"""
        paragraphs = text.split('\n\n')
        numbered_paragraphs = []
        
        counter = 1
        pattern = style.paragraph_numbering_pattern
        
        for para in paragraphs:
            if para.strip() and not self._is_title(para):
                # Générer le numéro selon le pattern
                if pattern == '1.':
                    number = f"{counter}."
                elif pattern == '1)':
                    number = f"{counter})"
                elif pattern == 'I.':
                    number = f"{self._to_roman(counter)}."
                elif pattern == 'A.':
                    number = f"{chr(64 + counter)}."
                elif pattern == 'a.':
                    number = f"{chr(96 + counter)}."
                elif pattern == '§ 1':
                    number = f"§ {counter}"
                else:
                    number = f"{counter}."
                
                numbered_paragraphs.append(f"{number} {para}")
                counter += 1
            else:
                numbered_paragraphs.append(para)
        
        return '\n\n'.join(numbered_paragraphs)
    
    def _is_title(self, text: str) -> bool:
        """Détermine si un texte est probablement un titre"""
        text = text.strip()
        
        # Critères pour identifier un titre
        if text.isupper():
            return True
        if len(text.split()) <= 5 and not any(char in text for char in '.!?,'):
            return True
        if text.startswith(('I.', 'II.', 'III.', 'A.', 'B.', 'C.')):
            return True
        
        return False
    
    def _to_roman(self, num: int) -> str:
        """Convertit un nombre en chiffres romains"""
        values = [
            (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
        ]
        result = ''
        for value, letter in values:
            count = num // value
            if count:
                result += letter * count
                num -= value * count
        return result
    
    def get_style_summary(self, style: StyleLearningResult) -> str:
        """Génère un résumé du style appris"""
        summary = []
        
        summary.append(f"**Style appris depuis {style.documents_analyzed} document(s)**\n")
        
        if style.paragraph_numbering_pattern:
            summary.append(f"- Numérotation : {style.paragraph_numbering_pattern}")
        
        summary.append(f"- Longueur moyenne des phrases : {style.average_sentence_length} mots")
        summary.append(f"- Longueur moyenne des paragraphes : {style.average_paragraph_length} mots")
        
        # Niveau de formalité
        if style.formality_score > 0.8:
            summary.append("- Niveau de formalité : Très formel")
        elif style.formality_score > 0.6:
            summary.append("- Niveau de formalité : Formel")
        else:
            summary.append("- Niveau de formalité : Moderne")
        
        if style.technical_terms_frequency > 1:
            summary.append(f"- Densité technique : {style.technical_terms_frequency:.1f}%")
        
        if style.common_phrases:
            summary.append(f"\n**Phrases types identifiées :**")
            for phrase in style.common_phrases[:5]:
                summary.append(f"  • {phrase}...")
        
        if style.transition_words:
            summary.append(f"\n**Mots de liaison privilégiés :**")
            summary.append(f"  {', '.join(style.transition_words[:10])}")
        
        summary.append(f"\n**Confiance dans l'analyse : {style.confidence_score:.0%}**")
        
        return "\n".join(summary)

# Singleton
_style_learning_service = None

def get_style_learning_service() -> StyleLearningService:
    """Retourne l'instance singleton du service"""
    global _style_learning_service
    if _style_learning_service is None:
        _style_learning_service = StyleLearningService()
    return _style_learning_service