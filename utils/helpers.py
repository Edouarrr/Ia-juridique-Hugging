# utils/helpers.py
"""Fonctions utilitaires pour l'application"""

import re
import os
import io
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any, Tuple
import hashlib
import unicodedata
from collections import Counter, defaultdict

logger = logging.getLogger(__name__)

# Import conditionnel pour docx
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("Module python-docx non disponible")


def format_date(date: datetime, format_str: str = "%d/%m/%Y") -> str:
    """Formate une date selon le format spécifié"""
    if not date:
        return ""
    
    if isinstance(date, str):
        try:
            date = datetime.fromisoformat(date)
        except:
            return date
    
    return date.strftime(format_str)


def format_currency(amount: float, symbol: str = "€") -> str:
    """Formate un montant en devise"""
    if amount is None:
        return f"0,00 {symbol}"
    
    # Formater avec séparateur de milliers et 2 décimales
    formatted = f"{amount:,.2f}".replace(",", " ").replace(".", ",")
    return f"{formatted} {symbol}"


def sanitize_text(text: str) -> str:
    """Nettoie et sécurise un texte"""
    if not text:
        return ""
    
    # Supprimer les caractères de contrôle
    text = ''.join(char for char in text if unicodedata.category(char)[0] != 'C')
    
    # Normaliser les espaces
    text = ' '.join(text.split())
    
    # Limiter la longueur
    max_length = 50000
    if len(text) > max_length:
        text = text[:max_length] + "..."
    
    return text.strip()


def extract_numbers(text: str) -> List[float]:
    """Extrait tous les nombres d'un texte"""
    if not text:
        return []
    
    # Pattern pour les nombres avec virgules françaises
    pattern = r'-?\d+(?:[.,]\d+)?'
    matches = re.findall(pattern, text)
    
    numbers = []
    for match in matches:
        try:
            # Remplacer la virgule par un point
            number = float(match.replace(',', '.'))
            numbers.append(number)
        except ValueError:
            continue
    
    return numbers


def calculate_risk_score(factors: Dict[str, float]) -> float:
    """Calcule un score de risque basé sur plusieurs facteurs"""
    if not factors:
        return 0.0
    
    # Pondérations par défaut
    weights = {
        'montant': 0.3,
        'complexite': 0.2,
        'antecedents': 0.2,
        'cooperation': 0.15,
        'premeditation': 0.15
    }
    
    total_score = 0.0
    total_weight = 0.0
    
    for factor, value in factors.items():
        weight = weights.get(factor, 0.1)
        total_score += value * weight
        total_weight += weight
    
    # Normaliser entre 0 et 10
    if total_weight > 0:
        return min(10.0, max(0.0, (total_score / total_weight) * 10))
    
    return 5.0  # Score neutre par défaut


def generate_unique_id(prefix: str = "doc") -> str:
    """Génère un identifiant unique"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    random_part = hashlib.md5(os.urandom(16)).hexdigest()[:8]
    return f"{prefix}_{timestamp}_{random_part}"


def validate_email(email: str) -> bool:
    """Valide une adresse email"""
    if not email:
        return False
    
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return bool(re.match(pattern, email))


def validate_phone(phone: str) -> bool:
    """Valide un numéro de téléphone français"""
    if not phone:
        return False
    
    # Nettoyer le numéro
    phone = re.sub(r'[^\d+]', '', phone)
    
    # Patterns français
    patterns = [
        r'^0[1-9]\d{8}$',  # 0123456789
        r'^\+33[1-9]\d{8}$',  # +33123456789
        r'^33[1-9]\d{8}$'  # 33123456789
    ]
    
    return any(re.match(pattern, phone) for pattern in patterns)


def clean_html(text: str) -> str:
    """Supprime les balises HTML d'un texte"""
    if not text:
        return ""
    
    # Supprimer les scripts et styles
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Supprimer toutes les balises
    text = re.sub(r'<[^>]+>', '', text)
    
    # Décoder les entités HTML
    html_entities = {
        '&nbsp;': ' ',
        '&amp;': '&',
        '&lt;': '<',
        '&gt;': '>',
        '&quot;': '"',
        '&#39;': "'",
        '&euro;': '€'
    }
    
    for entity, char in html_entities.items():
        text = text.replace(entity, char)
    
    return text.strip()


def truncate_text(text: str, max_length: int = 100, suffix: str = "...") -> str:
    """Tronque un texte à une longueur maximale"""
    if not text or len(text) <= max_length:
        return text
    
    # Trouver le dernier espace avant la limite
    truncated = text[:max_length]
    last_space = truncated.rfind(' ')
    
    if last_space > 0:
        truncated = truncated[:last_space]
    
    return truncated + suffix


def highlight_text(text: str, terms: List[str], tag: str = "mark") -> str:
    """Surligne des termes dans un texte"""
    if not text or not terms:
        return text
    
    # Créer un pattern pour tous les termes
    escaped_terms = [re.escape(term) for term in terms if term]
    if not escaped_terms:
        return text
    
    pattern = '|'.join(escaped_terms)
    
    # Remplacer en conservant la casse
    def replace_match(match):
        return f"<{tag}>{match.group()}</{tag}>"
    
    return re.sub(pattern, replace_match, text, flags=re.IGNORECASE)


def create_summary(text: str, max_sentences: int = 3) -> str:
    """Crée un résumé du texte"""
    if not text:
        return ""
    
    # Diviser en phrases
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    if len(sentences) <= max_sentences:
        return text
    
    # Prendre les premières phrases
    summary = '. '.join(sentences[:max_sentences])
    
    # Ajouter un point si nécessaire
    if not summary.endswith('.'):
        summary += '.'
    
    return summary


def parse_legal_reference(reference: str) -> Dict[str, str]:
    """Parse une référence juridique"""
    result = {
        'type': '',
        'juridiction': '',
        'date': '',
        'numero': '',
        'original': reference
    }
    
    if not reference:
        return result
    
    # Patterns pour différents types de références
    patterns = {
        'arret': r'(Cass\.|CA|CE|CC|CJUE|CEDH)\s+(.+?),?\s+(\d{1,2}[\s\-./]\w+[\s\-./]\d{2,4}),?\s+n[°o]?\s*(.+)',
        'article': r'(Art\.?|Article)\s+(\d+(?:-\d+)?)\s+(?:du\s+)?(.+)',
        'loi': r'(Loi|Décret|Ordonnance)\s+n[°o]?\s*(.+?)\s+du\s+(.+)'
    }
    
    for type_ref, pattern in patterns.items():
        match = re.search(pattern, reference, re.IGNORECASE)
        if match:
            result['type'] = type_ref
            if type_ref == 'arret':
                result['juridiction'] = match.group(1)
                result['date'] = match.group(3)
                result['numero'] = match.group(4)
            elif type_ref == 'article':
                result['numero'] = match.group(2)
                result['source'] = match.group(3)
            elif type_ref == 'loi':
                result['type_texte'] = match.group(1)
                result['numero'] = match.group(2)
                result['date'] = match.group(3)
            break
    
    return result


def extract_legal_articles(text: str) -> List[str]:
    """Extrait les références d'articles de loi d'un texte"""
    if not text:
        return []
    
    articles = []
    
    # Patterns pour les articles
    patterns = [
        r'articles?\s+(\d+(?:-\d+)?(?:\s+et\s+\d+)*)\s+(?:du\s+)?([^,\.;]+?)(?:[,\.;]|$)',
        r'art\.?\s+(\d+(?:-\d+)?)\s+(?:du\s+)?([^,\.;]+?)(?:[,\.;]|$)'
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            article = f"Art. {match.group(1)} {match.group(2).strip()}"
            if article not in articles:
                articles.append(article)
    
    return articles


def calculate_prescription_date(date_faits: datetime, type_infraction: str) -> Tuple[datetime, str]:
    """Calcule la date de prescription selon le type d'infraction"""
    from config.app_config import PRESCRIPTION_CONFIG
    
    if not date_faits:
        return None, "Date des faits non spécifiée"
    
    # Déterminer le délai selon le type
    delai_annees = PRESCRIPTION_CONFIG.get(type_infraction.lower(), 6)
    
    # Calculer la date de prescription
    try:
        from dateutil.relativedelta import relativedelta
        date_prescription = date_faits + relativedelta(years=delai_annees)
        
        # Vérifier si prescrit
        if datetime.now() > date_prescription:
            statut = f"Prescrit depuis le {format_date(date_prescription)}"
        else:
            jours_restants = (date_prescription - datetime.now()).days
            statut = f"Prescription le {format_date(date_prescription)} ({jours_restants} jours)"
        
        return date_prescription, statut
        
    except ImportError:
        # Fallback sans dateutil
        import datetime as dt
        date_prescription = date_faits.replace(year=date_faits.year + delai_annees)
        
        if datetime.now() > date_prescription:
            statut = f"Prescrit depuis le {format_date(date_prescription)}"
        else:
            statut = f"Prescription le {format_date(date_prescription)}"
        
        return date_prescription, statut


def clean_key(text: str) -> str:
    """Nettoie une chaîne pour en faire une clé Streamlit valide"""
    if not text:
        return "key"
    
    replacements = {
        'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e',
        'à': 'a', 'â': 'a', 'ä': 'a',
        'ù': 'u', 'û': 'u', 'ü': 'u',
        'ô': 'o', 'ö': 'o',
        'î': 'i', 'ï': 'i',
        'ç': 'c',
        ' ': '_', '-': '_', "'": '_', '"': '_',
        '.': '_', ',': '_', '(': '_', ')': '_',
        '[': '_', ']': '_', '/': '_', '\\': '_',
        ':': '_', ';': '_', '!': '_', '?': '_',
        '@': '_', '#': '_', '$': '_', '%': '_',
        '^': '_', '&': '_', '*': '_', '+': '_',
        '=': '_', '{': '_', '}': '_', '|': '_',
        '<': '_', '>': '_', '~': '_', '`': '_'
    }
    
    cleaned = text.lower()
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    
    # Garder uniquement les caractères alphanumériques et underscore
    cleaned = ''.join(c if c.isalnum() or c == '_' else '_' for c in cleaned)
    
    # Supprimer les underscores multiples
    cleaned = re.sub(r'_+', '_', cleaned)
    
    # Supprimer les underscores au début et à la fin
    cleaned = cleaned.strip('_')
    
    # S'assurer que la clé n'est pas vide
    if not cleaned:
        cleaned = "key"
    
    # Limiter la longueur
    if len(cleaned) > 50:
        cleaned = cleaned[:50]
    
    return cleaned


def merge_structures(structures: List[Dict]) -> Dict[str, Any]:
    """Fusionne plusieurs structures de documents"""
    if not structures:
        return {}
    
    merged = {
        'sections_communes': [],
        'longueur_moyenne': 0
    }
    
    # Trouver les sections communes
    all_sections = []
    for struct in structures:
        all_sections.extend([s['titre'] for s in struct.get('sections', [])])
    
    # Compter les occurrences
    section_counts = Counter(all_sections)
    
    # Garder les sections présentes dans au moins 50% des documents
    threshold = len(structures) / 2
    merged['sections_communes'] = [
        section for section, count in section_counts.items()
        if count >= threshold
    ]
    
    # Ajouter les informations Word si présentes
    word_styles = []
    for struct in structures:
        if 'word_styles' in struct:
            word_styles.extend(struct['word_styles'])
    
    if word_styles:
        merged['word_styles'] = list(set(word_styles))
    
    return merged


def merge_formules(formules_list: List[List[str]]) -> List[str]:
    """Fusionne les formules types"""
    all_formules = []
    for formules in formules_list:
        all_formules.extend(formules)
    
    # Compter et garder les plus fréquentes
    formule_counts = Counter(all_formules)
    
    return [formule for formule, count in formule_counts.most_common(20)]


def merge_formatting(formats: List[Dict]) -> Dict[str, Any]:
    """Fusionne les paramètres de mise en forme"""
    if not formats:
        return {}
    
    merged = {}
    
    # Moyennes et valeurs communes
    for key in formats[0].keys():
        values = [f.get(key) for f in formats if key in f]
        
        if all(isinstance(v, bool) for v in values):
            # Pour les booléens, prendre la majorité
            merged[key] = sum(values) > len(values) / 2
        elif all(isinstance(v, (int, float)) for v in values):
            # Pour les nombres, prendre la moyenne
            merged[key] = sum(values) / len(values)
        else:
            # Pour le reste, prendre la valeur la plus fréquente
            merged[key] = Counter(values).most_common(1)[0][0] if values else None
    
    return merged


def merge_vocabulary(vocab_list: List[Dict[str, int]]) -> Dict[str, int]:
    """Fusionne les vocabulaires"""
    merged = defaultdict(int)
    
    for vocab in vocab_list:
        for word, count in vocab.items():
            merged[word] += count
    
    # Garder les 100 mots les plus fréquents
    return dict(sorted(merged.items(), key=lambda x: x[1], reverse=True)[:100])


def create_letterhead_from_template(template, content: str):
    """Crée un document avec papier en-tête à partir d'un template"""
    if not DOCX_AVAILABLE:
        logger.error("Module python-docx non disponible pour créer le document")
        return None
        
    try:
        doc = DocxDocument()
        
        # Définir les marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(template.margins['top'])
            section.bottom_margin = Inches(template.margins['bottom'])
            section.left_margin = Inches(template.margins['left'])
            section.right_margin = Inches(template.margins['right'])
        
        # Ajouter l'en-tête
        if template.header_content:
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = template.header_content
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter le contenu principal
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            p = doc.add_paragraph(para_text)
            p.style.font.name = template.font_family
            p.style.font.size = Pt(template.font_size)
            p.paragraph_format.line_spacing = template.line_spacing
        
        # Ajouter le pied de page
        if template.footer_content:
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = template.footer_content
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sauvegarder en mémoire
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer
        
    except Exception as e:
        logger.error(f"Erreur création document avec papier en-tête: {e}")
        return None


def format_legal_date(date: datetime) -> str:
    """Formate une date au format juridique français"""
    if not date:
        return ""
    
    months = {
        1: "janvier", 2: "février", 3: "mars", 4: "avril",
        5: "mai", 6: "juin", 7: "juillet", 8: "août",
        9: "septembre", 10: "octobre", 11: "novembre", 12: "décembre"
    }
    
    return f"{date.day} {months.get(date.month, '')} {date.year}"


def extract_amount_from_text(text: str) -> Optional[float]:
    """Extrait un montant d'un texte"""
    if not text:
        return None
    
    # Patterns pour les montants
    patterns = [
        r'(\d+(?:\s*\d{3})*(?:,\d+)?)\s*(?:€|EUR|euros?)',
        r'(\d+(?:\.\d{3})*(?:,\d+)?)\s*(?:€|EUR|euros?)',
        r'(\d+(?:,\d+)?)\s*(?:€|EUR|euros?)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            amount_str = match.group(1)
            # Normaliser le format
            amount_str = amount_str.replace(' ', '').replace('.', '').replace(',', '.')
            try:
                return float(amount_str)
            except ValueError:
                continue
    
    return None


def anonymize_text(text: str, level: str = "medium") -> str:
    """Anonymise un texte selon le niveau demandé"""
    if not text:
        return ""
    
    # Patterns à anonymiser
    patterns = {
        'email': (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL]'),
        'phone': (r'(?:(?:\+|00)33|0)\s*[1-9](?:[\s.-]*\d{2}){4}', '[TELEPHONE]'),
        'iban': (r'[A-Z]{2}\d{2}\s?(?:\w{4}\s?){2,7}\w{1,4}', '[IBAN]'),
        'ssn': (r'[12]\s?\d{2}\s?\d{2}\s?\d{2}\s?\d{3}\s?\d{3}(?:\s?\d{2})?', '[NUM_SECU]')
    }
    
    result = text
    
    # Niveau basic : emails et téléphones
    if level in ["basic", "medium", "high"]:
        for key in ['email', 'phone']:
            pattern, replacement = patterns[key]
            result = re.sub(pattern, replacement, result)
    
    # Niveau medium : + IBAN et numéros de sécurité sociale
    if level in ["medium", "high"]:
        for key in ['iban', 'ssn']:
            pattern, replacement = patterns[key]
            result = re.sub(pattern, replacement, result)
    
    # Niveau high : + noms propres (approximatif)
    if level == "high":
        # Remplacer les mots commençant par une majuscule (sauf début de phrase)
        result = re.sub(r'(?<!^)(?<!\. )\b[A-Z][a-z]+\b', '[NOM]', result)
    
    return result


def generate_file_hash(file_content: bytes) -> str:
    """Génère un hash SHA-256 d'un fichier"""
    return hashlib.sha256(file_content).hexdigest()


def format_file_size(size_bytes: int) -> str:
    """Formate une taille de fichier de manière lisible"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} TB"


def is_valid_french_company_number(siret: str) -> bool:
    """Vérifie la validité d'un numéro SIRET français"""
    if not siret:
        return False
    
    # Nettoyer le numéro
    siret = re.sub(r'[^\d]', '', siret)
    
    if len(siret) != 14:
        return False
    
    # Algorithme de Luhn pour vérifier la validité
    total = 0
    for i, digit in enumerate(siret):
        n = int(digit)
        if i % 2 == 0:
            n *= 2
            if n > 9:
                n -= 9
        total += n
    
    return total % 10 == 0