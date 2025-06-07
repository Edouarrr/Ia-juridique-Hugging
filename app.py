import streamlit as st

# PREMIÈRE commande Streamlit OBLIGATOIREMENT
st.set_page_config(
    page_title="Assistant Pénal des Affaires IA", 
    page_icon="⚖️", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Configuration de l'encodage pour les emojis
import sys
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# ================== IMPORTS ET CONFIGURATION ==================

# Charger les variables d'environnement
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Vérifier la disponibilité d'Azure
try:
    from azure.search.documents import SearchClient
    from azure.search.documents.indexes import SearchIndexClient
    from azure.search.documents.models import VectorizedQuery
    from azure.core.credentials import AzureKeyCredential
    from azure.storage.blob import BlobServiceClient
    AZURE_AVAILABLE = True
except ImportError:
    AZURE_AVAILABLE = False
    st.warning("⚠️ Modules Azure non disponibles - Fonctionnalités limitées")

import os
import io
from typing import Dict, List, Set, Tuple, Optional, Any
import json
from datetime import datetime, timedelta
import re
import pandas as pd
import base64
import hashlib
from dataclasses import dataclass, field, asdict
from enum import Enum
import logging
from functools import lru_cache, wraps
import asyncio
from concurrent.futures import ThreadPoolExecutor
import difflib
from collections import defaultdict, Counter

# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Imports conditionnels pour les IA
try:
    from openai import OpenAI, AzureOpenAI
except ImportError:
    OpenAI = None
    AzureOpenAI = None

try:
    import anthropic
except ImportError:
    anthropic = None

try:
    import google.generativeai as genai
except ImportError:
    genai = None

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ================== FIX POUR HUGGING FACE SPACES ==================
def clean_env_for_azure():
    """Nettoie l'environnement pour Azure sur Hugging Face Spaces"""
    # Supprimer les variables de proxy qui interfèrent avec Azure
    proxy_vars = ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy', 
                  'NO_PROXY', 'no_proxy', 'REQUESTS_CA_BUNDLE', 'CURL_CA_BUNDLE']
    
    for var in proxy_vars:
        if var in os.environ:
            del os.environ[var]
    
    # Désactiver les proxies dans requests
    os.environ['CURL_CA_BUNDLE'] = ""
    os.environ['REQUESTS_CA_BUNDLE'] = ""
    
# Appeler cette fonction AVANT toute initialisation Azure
clean_env_for_azure()

# ================== CONFIGURATION CENTRALISÉE ==================
class AppConfig:
    """Configuration centralisée de l'application"""
    
    # UI
    PAGE_SIZE = 10
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Container Azure par défaut
    DEFAULT_CONTAINER = "sharepoint-documents"
    
    # Prescription
    PRESCRIPTION_CONTRAVENTION = 1  # an
    PRESCRIPTION_DELIT = 6  # ans
    PRESCRIPTION_CRIME = 20  # ans
    
    # Délais spécifiques pénal des affaires
    PRESCRIPTION_FRAUDE_FISCALE = 6  # ans (sauf dissimulation)
    PRESCRIPTION_TRAVAIL_DISSIMULE = 6  # ans
    
    # Export
    EXPORT_FORMAT = "%Y%m%d_%H%M%S"
    
    # Azure Search
    SEARCH_INDEX_NAME = "juridique-index"
    VECTOR_DIMENSION = 1536  # OpenAI embeddings
    
    # Formats de citation
    CITATION_FORMATS = {
        'jurisprudence': "{juridiction}, {date}, n° {numero}",
        'article_code': "Art. {numero} {code}",
        'doctrine': "{auteur}, « {titre} », {revue} {annee}, n° {numero}, p. {page}",
        'circulaire': "Circ. {reference} du {date}",
        'reponse_ministerielle': "Rép. min. n° {numero}, {date}"
    }

# ================== INITIALISATION SESSION STATE ==================

def initialize_session_state():
    """Initialise toutes les variables de session"""
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        st.session_state.dossier_actif = None
        st.session_state.dossiers = {}
        st.session_state.pieces_selectionnees = {}
        st.session_state.current_page = 1
        st.session_state.search_query = ""
        st.session_state.docs_for_analysis = []
        st.session_state.document_content = ""
        st.session_state.custom_templates = {}
        st.session_state.letterhead = None
        st.session_state.citation_manager = None
        st.session_state.victimes_adapter = []
        st.session_state.plainte_originale = None
        st.session_state.plaintes_adaptees = {}
        st.session_state.azure_documents = {}
        st.session_state.selected_container = None
        st.session_state.current_folder_path = ""
        st.session_state.style_models = {}  # Documents modèles pour le style
        st.session_state.learned_styles = {}  # Styles appris
        st.session_state.vector_store = None  # Store de vecteurs
        st.session_state.azure_search_client = None
        st.session_state.azure_blob_manager = None
        st.session_state.azure_search_manager = None
        st.session_state.dynamic_search_prompts = {}  # Cache pour les prompts générés
        st.session_state.dynamic_templates = {}  # Cache pour les modèles générés
        st.session_state.selected_folders = set()  # Dossiers sélectionnés
        st.session_state.letterhead_template = None  # Template de papier en-tête
        st.session_state.letterhead_image = None  # Image du papier en-tête

# ================== ENUMERATIONS ==================

class InfractionAffaires(Enum):
    """Types d'infractions en droit pénal des affaires - Liste non exhaustive"""
    ABS = "Abus de biens sociaux"
    ABUS_CONFIANCE = "Abus de confiance"
    CORRUPTION = "Corruption"
    TRAFIC_INFLUENCE = "Trafic d'influence"
    PRISE_ILLEGALE = "Prise illégale d'intérêts"
    FAVORITISME = "Favoritisme"
    BLANCHIMENT = "Blanchiment"
    FRAUDE_FISCALE = "Fraude fiscale"
    ESCROQUERIE = "Escroquerie"
    FAUX_USAGE_FAUX = "Faux et usage de faux"
    BANQUEROUTE = "Banqueroute"
    DELIT_INITIE = "Délit d'initié"
    MANIPULATION_COURS = "Manipulation de cours"
    ENTRAVE = "Entrave"
    TRAVAIL_DISSIMULE = "Travail dissimulé"
    HARCELEMENT = "Harcèlement moral/sexuel"
    MISE_DANGER = "Mise en danger d'autrui"
    BLESSURES_INVOLONTAIRES = "Blessures involontaires"
    POLLUTION = "Atteinte à l'environnement"
    AUTRE = "Autre infraction"

class SearchMode(Enum):
    """Modes de recherche disponibles"""
    HYBRID = "Recherche hybride (textuelle + sémantique)"
    TEXT_ONLY = "Recherche textuelle uniquement"
    VECTOR_ONLY = "Recherche vectorielle uniquement"
    LOCAL = "Recherche locale uniquement"

class LLMProvider(Enum):
    """Providers LLM disponibles"""
    AZURE_OPENAI = "Azure OpenAI (GPT-4)"
    CLAUDE_OPUS = "Claude Opus 4"
    CHATGPT_4O = "ChatGPT 4o"
    GEMINI = "Google Gemini"
    PERPLEXITY = "Perplexity AI"

# ================== DATACLASSES ==================

@dataclass
class Document:
    """Représente un document juridique"""
    id: str
    title: str
    content: str
    source: str = 'local'
    page_number: int = 1
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: datetime = field(default_factory=datetime.now)
    folder_path: Optional[str] = None
    selected: bool = False
    embedding: Optional[List[float]] = None  # Vecteur pour la recherche

@dataclass
class PieceSelectionnee:
    """Pièce sélectionnée pour un dossier"""
    document_id: str
    titre: str
    categorie: str
    date_selection: datetime = field(default_factory=datetime.now)
    notes: str = ""
    pertinence: int = 5  # Score de 1 à 10

@dataclass
class StylePattern:
    """Modèle de style extrait d'un document"""
    document_id: str
    type_acte: str
    structure: Dict[str, Any]  # Structure du document
    formules: List[str]  # Formules types extraites
    mise_en_forme: Dict[str, Any]  # Paramètres de mise en forme
    vocabulaire: Dict[str, int]  # Fréquence des mots
    paragraphes_types: List[str]  # Exemples de paragraphes

@dataclass
class LetterheadTemplate:
    """Template de papier en-tête"""
    name: str
    header_content: str  # Contenu HTML/Markdown de l'en-tête
    footer_content: str  # Contenu HTML/Markdown du pied de page
    logo_path: Optional[str] = None
    margins: Dict[str, float] = field(default_factory=lambda: {
        'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5
    })
    font_family: str = "Arial"
    font_size: int = 11
    line_spacing: float = 1.5
    metadata: Dict[str, Any] = field(default_factory=dict)

# ================== CSS PERSONNALISÉ ==================

st.markdown("""
<style>
    :root {
        --primary-color: #1a237e;
        --secondary-color: #283593;
        --success-color: #2e7d32;
        --warning-color: #f57c00;
        --error-color: #c62828;
    }
    
    .main { 
        background-color: #f5f5f5; 
    }
    
    .stButton > button {
        background-color: var(--primary-color);
        color: white;
        border-radius: 5px;
        border: none;
        padding: 0.5rem 1rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        background-color: var(--secondary-color);
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(26, 35, 126, 0.3);
    }
    
    .document-card {
        background: white;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        border-left: 5px solid var(--primary-color);
        transition: all 0.3s ease;
    }
    
    .document-card:hover {
        box-shadow: 0 4px 16px rgba(0,0,0,0.15);
    }
    
    .folder-card {
        background: #e3f2fd;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        border-left: 5px solid var(--secondary-color);
        transition: all 0.3s ease;
    }
    
    .folder-card:hover {
        box-shadow: 0 4px 16px rgba(0,0,0,0.15);
    }
    
    .folder-nav {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        margin-bottom: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .search-section {
        background: white;
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    .piece-selectionnee {
        background: #e8f5e9;
        border-left: 3px solid var(--success-color);
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 0 8px 8px 0;
    }
    
    .azure-browser {
        background: #f0f7ff;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    
    .style-model {
        background: #fff3e0;
        border-left: 3px solid var(--warning-color);
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 0 8px 8px 0;
    }
    
    .letterhead-preview {
        background: white;
        border: 1px solid #ddd;
        padding: 2rem;
        margin: 1rem 0;
        min-height: 400px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
</style>
""", unsafe_allow_html=True)

# ================== UTILITAIRES ==================

def clean_key(text: str) -> str:
    """Nettoie une chaîne pour en faire une clé Streamlit valide"""
    replacements = {
        'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e',
        'à': 'a', 'â': 'a', 'ä': 'a',
        'ù': 'u', 'û': 'u', 'ü': 'u',
        'ô': 'o', 'ö': 'o',
        'î': 'i', 'ï': 'i',
        'ç': 'c',
        ' ': '_', '-': '_', "'": '_', '"': '_',
        '.': '_', ',': '_', '(': '_', ')': '_',
        '[': '_', ']': '_', '/': '_', '\\': '_'
    }
    
    cleaned = text.lower()
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    
    cleaned = ''.join(c if c.isalnum() or c == '_' else '_' for c in cleaned)
    cleaned = re.sub(r'_+', '_', cleaned)
    cleaned = cleaned.strip('_')
    
    return cleaned

# ================== AZURE SEARCH MANAGER ==================

class AzureSearchManager:
    """Gestionnaire pour Azure Cognitive Search avec vectorisation"""
    
    def __init__(self):
        self.search_client = None
        self.index_client = None
        self.openai_client = None
        try:
            self._init_clients()
        except Exception as e:
            logger.error(f"Erreur lors de l'initialisation d'AzureSearchManager : {e}")
            self.search_client = None
            self.index_client = None
            self.openai_client = None
    
    def _init_clients(self):
        """Initialise les clients Azure Search et OpenAI"""
        try:
            # Nettoyer l'environnement pour Hugging Face
            clean_env_for_azure()
            
            # Azure Search
            search_endpoint = os.getenv('AZURE_SEARCH_ENDPOINT')
            search_key = os.getenv('AZURE_SEARCH_KEY')
            
            if search_endpoint and search_key and AZURE_AVAILABLE:
                self.index_client = SearchIndexClient(
                    endpoint=search_endpoint,
                    credential=AzureKeyCredential(search_key)
                )
                
                self.search_client = SearchClient(
                    endpoint=search_endpoint,
                    index_name=AppConfig.SEARCH_INDEX_NAME,
                    credential=AzureKeyCredential(search_key)
                )
                
                logger.info("Client Azure Search initialisé")
            
            # OpenAI pour les embeddings
            openai_key = os.getenv('AZURE_OPENAI_KEY')
            openai_endpoint = os.getenv('AZURE_OPENAI_ENDPOINT')
            
            if openai_key and openai_endpoint and AzureOpenAI:
                # Créer le client sans paramètres proxy
                self.openai_client = AzureOpenAI(
                    azure_endpoint=openai_endpoint,
                    api_key=openai_key,
                    api_version="2024-02-01"
                )
                logger.info("Client OpenAI pour embeddings initialisé")
                
        except Exception as e:
            logger.error(f"Erreur initialisation Azure Search: {e}")
            # Ne pas propager l'erreur pour permettre à l'app de continuer
            self.search_client = None
            self.index_client = None
            self.openai_client = None
    
    def generate_embedding(self, text: str) -> Optional[List[float]]:
        """Génère un embedding pour un texte"""
        if not self.openai_client:
            return None
        
        try:
            response = self.openai_client.embeddings.create(
                input=text[:8000],  # Limite pour l'API
                model="text-embedding-ada-002"
            )
            return response.data[0].embedding
            
        except Exception as e:
            logger.error(f"Erreur génération embedding: {e}")
            return None
    
    def index_document(self, document: Document):
        """Indexe un document avec son vecteur"""
        if not self.search_client:
            return False
        
        try:
            # Générer l'embedding
            embedding = self.generate_embedding(f"{document.title} {document.content}")
            
            if not embedding:
                logger.warning(f"Pas d'embedding pour {document.id}")
                embedding = [0.0] * AppConfig.VECTOR_DIMENSION
            
            # Préparer le document pour l'index
            doc_to_index = {
                "id": document.id,
                "title": document.title,
                "content": document.content,
                "source": document.source,
                "embedding": embedding
            }
            
            # Indexer
            self.search_client.upload_documents([doc_to_index])
            logger.info(f"Document {document.id} indexé")
            return True
            
        except Exception as e:
            logger.error(f"Erreur indexation document: {e}")
            return False
    
    def search_hybrid(self, query: str, mode: SearchMode = SearchMode.HYBRID, 
                     top: int = 50, filters: Optional[str] = None) -> List[Dict[str, Any]]:
        """Recherche hybride (textuelle + vectorielle)"""
        if not self.search_client:
            return []
        
        try:
            results = []
            
            if mode in [SearchMode.HYBRID, SearchMode.VECTOR_ONLY]:
                # Générer l'embedding de la requête
                query_embedding = self.generate_embedding(query)
                
                if query_embedding:
                    # Recherche vectorielle
                    vector_query = VectorizedQuery(
                        vector=query_embedding,
                        k_nearest_neighbors=top,
                        fields="embedding"
                    )
                    
                    if mode == SearchMode.VECTOR_ONLY:
                        # Recherche vectorielle uniquement
                        response = self.search_client.search(
                            search_text=None,
                            vector_queries=[vector_query],
                            filter=filters,
                            top=top
                        )
                    else:
                        # Recherche hybride
                        response = self.search_client.search(
                            search_text=query,
                            vector_queries=[vector_query],
                            filter=filters,
                            top=top
                        )
                    
                    for result in response:
                        results.append({
                            'id': result['id'],
                            'title': result['title'],
                            'content': result['content'],
                            'score': result['@search.score'],
                            'source': result['source']
                        })
            
            elif mode == SearchMode.TEXT_ONLY:
                # Recherche textuelle uniquement
                response = self.search_client.search(
                    search_text=query,
                    filter=filters,
                    top=top
                )
                
                for result in response:
                    results.append({
                        'id': result['id'],
                        'title': result['title'],
                        'content': result['content'],
                        'score': result['@search.score'],
                        'source': result['source']
                    })
            
            return results
            
        except Exception as e:
            logger.error(f"Erreur recherche hybride: {e}")
            return []

# ================== GESTIONNAIRE AZURE BLOB ==================

class AzureBlobManager:
    """Gestionnaire pour Azure Blob Storage avec navigation dans les dossiers"""
    
    def __init__(self):
        self.blob_service_client = None
        self.container_client = None
        try:
            self._init_blob_client()
        except Exception as e:
            logger.error(f"Erreur lors de l'initialisation d'AzureBlobManager : {e}")
            self.blob_service_client = None
    
    def _init_blob_client(self):
        """Initialise le client Azure Blob"""
        try:
            connection_string = os.getenv('AZURE_STORAGE_CONNECTION_STRING')
            if connection_string and AZURE_AVAILABLE:
                self.blob_service_client = BlobServiceClient.from_connection_string(connection_string)
                logger.info("Client Azure Blob Storage initialisé avec succès")
            else:
                if not connection_string:
                    logger.warning("AZURE_STORAGE_CONNECTION_STRING non définie")
                if not AZURE_AVAILABLE:
                    logger.warning("Modules Azure non disponibles")
        except Exception as e:
            logger.error(f"Erreur initialisation Azure Blob : {e}")
            self.blob_service_client = None
    
    def is_connected(self) -> bool:
        """Vérifie si la connexion est établie"""
        return self.blob_service_client is not None
    
    def list_containers(self) -> List[str]:
        """Liste tous les containers disponibles"""
        if not self.blob_service_client:
            return []
        
        try:
            containers = self.blob_service_client.list_containers()
            container_list = [container.name for container in containers]
            logger.info(f"Containers trouvés: {container_list}")
            return container_list
        except Exception as e:
            logger.error(f"Erreur listing containers : {e}")
            return []
    
    def list_folders(self, container_name: str, prefix: str = "") -> List[Dict[str, Any]]:
        """Liste les dossiers et fichiers dans un chemin donné"""
        if not self.blob_service_client:
            return []
        
        try:
            container_client = self.blob_service_client.get_container_client(container_name)
            
            # Structure pour stocker l'arborescence
            folders = set()
            files = []
            
            # Lister tous les blobs avec le préfixe
            blobs = container_client.list_blobs(name_starts_with=prefix)
            
            for blob in blobs:
                # Extraire le chemin relatif
                relative_path = blob.name[len(prefix):] if prefix else blob.name
                parts = relative_path.split('/')
                
                # Si c'est un dossier (a des sous-éléments)
                if len(parts) > 1 and parts[0]:
                    folders.add(parts[0])
                # Si c'est un fichier direct
                elif len(parts) == 1 and parts[0]:
                    files.append({
                        'name': parts[0],
                        'size': blob.size,
                        'last_modified': blob.last_modified,
                        'content_type': blob.content_settings.content_type if blob.content_settings else None,
                        'full_path': blob.name
                    })
            
            # Combiner dossiers et fichiers
            result = []
            
            # Ajouter les dossiers
            for folder in sorted(folders):
                result.append({
                    'name': folder,
                    'type': 'folder',
                    'path': f"{prefix}{folder}/" if prefix else f"{folder}/"
                })
            
            # Ajouter les fichiers
            for file in sorted(files, key=lambda x: x['name']):
                result.append({
                    'name': file['name'],
                    'type': 'file',
                    'size': file['size'],
                    'last_modified': file['last_modified'],
                    'content_type': file['content_type'],
                    'full_path': file['full_path']
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Erreur listing dossiers : {e}")
            return []
    
    def download_file(self, container_name: str, blob_name: str) -> Optional[bytes]:
        """Télécharge un fichier depuis Azure Blob"""
        if not self.blob_service_client:
            return None
        
        try:
            blob_client = self.blob_service_client.get_blob_client(
                container=container_name,
                blob=blob_name
            )
            
            return blob_client.download_blob().readall()
            
        except Exception as e:
            logger.error(f"Erreur téléchargement blob : {e}")
            return None
    
    def extract_text_from_blob(self, container_name: str, blob_name: str) -> Optional[str]:
        """Extrait le texte d'un blob"""
        content_bytes = self.download_file(container_name, blob_name)
        
        if not content_bytes:
            return None
        
        file_ext = os.path.splitext(blob_name)[1].lower()
        
        try:
            if file_ext == '.txt':
                return content_bytes.decode('utf-8', errors='ignore')
            elif file_ext in ['.docx', '.doc'] and DOCX_AVAILABLE:
                doc = DocxDocument(io.BytesIO(content_bytes))
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            else:
                return f"[Format {file_ext} non supporté pour l'extraction de texte]"
        except Exception as e:
            logger.error(f"Erreur extraction texte : {e}")
            return None
    
    def get_all_files_in_folder(self, container_name: str, folder_path: str) -> List[Dict[str, Any]]:
        """Récupère récursivement tous les fichiers d'un dossier"""
        if not self.blob_service_client:
            return []
        
        try:
            container_client = self.blob_service_client.get_container_client(container_name)
            all_files = []
            
            # Lister tous les blobs avec le préfixe du dossier
            blobs = container_client.list_blobs(name_starts_with=folder_path)
            
            for blob in blobs:
                # Ignorer les dossiers vides
                if not blob.name.endswith('/'):
                    all_files.append({
                        'name': os.path.basename(blob.name),
                        'full_path': blob.name,
                        'size': blob.size,
                        'last_modified': blob.last_modified,
                        'folder': os.path.dirname(blob.name)
                    })
            
            return all_files
            
        except Exception as e:
            logger.error(f"Erreur récupération fichiers du dossier : {e}")
            return []

# ================== GESTIONNAIRE DE SÉLECTION DE PIÈCES ==================

class GestionnairePiecesSelectionnees:
    """Gère la sélection et l'organisation des pièces pour un dossier"""
    
    def __init__(self):
        self.pieces: Dict[str, PieceSelectionnee] = {}
        self.categories = [
            "📄 Procédure",
            "💰 Comptabilité",
            "📊 Expertise",
            "📧 Correspondances",
            "📋 Contrats",
            "🏢 Documents sociaux",
            "🔍 Preuves",
            "📑 Autres"
        ]
    
    def ajouter_piece(self, document: Document, categorie: str, notes: str = "", pertinence: int = 5):
        """Ajoute une pièce à la sélection"""
        piece = PieceSelectionnee(
            document_id=document.id,
            titre=document.title,
            categorie=categorie,
            notes=notes,
            pertinence=pertinence
        )
        
        self.pieces[document.id] = piece
        
        # Sauvegarder dans session state
        if 'pieces_selectionnees' not in st.session_state:
            st.session_state.pieces_selectionnees = {}
        st.session_state.pieces_selectionnees[document.id] = piece
    
    def retirer_piece(self, document_id: str):
        """Retire une pièce de la sélection"""
        if document_id in self.pieces:
            del self.pieces[document_id]
        
        if document_id in st.session_state.pieces_selectionnees:
            del st.session_state.pieces_selectionnees[document_id]
    
    def get_pieces_par_categorie(self) -> Dict[str, List[PieceSelectionnee]]:
        """Retourne les pièces organisées par catégorie"""
        pieces_par_cat = {cat: [] for cat in self.categories}
        
        for piece in self.pieces.values():
            if piece.categorie in pieces_par_cat:
                pieces_par_cat[piece.categorie].append(piece)
        
        return pieces_par_cat
    
    def generer_bordereau(self) -> str:
        """Génère un bordereau des pièces sélectionnées"""
        bordereau = "BORDEREAU DES PIÈCES SÉLECTIONNÉES\n"
        bordereau += "=" * 50 + "\n\n"
        bordereau += f"Date : {datetime.now().strftime('%d/%m/%Y')}\n\n"
        
        pieces_par_cat = self.get_pieces_par_categorie()
        numero = 1
        
        for categorie, pieces in pieces_par_cat.items():
            if pieces:
                bordereau += f"\n{categorie}\n{'-' * len(categorie)}\n"
                for piece in sorted(pieces, key=lambda x: x.pertinence, reverse=True):
                    bordereau += f"{numero}. {piece.titre}"
                    if piece.notes:
                        bordereau += f" - {piece.notes}"
                    bordereau += f" (Pertinence: {piece.pertinence}/10)\n"
                    numero += 1
        
        bordereau += f"\n\nTOTAL : {len(self.pieces)} pièces\n"
        
        return bordereau

# ================== ANALYSEUR DE STYLE ==================

class StyleAnalyzer:
    """Analyse et apprend le style de rédaction des documents"""
    
    def __init__(self):
        self.patterns = defaultdict(list)
        self.formules_types = defaultdict(set)
        self.structures = defaultdict(list)
    
    def analyze_document(self, document: Document, type_acte: str) -> StylePattern:
        """Analyse un document pour en extraire le style"""
        content = document.content
        
        # Analyser la structure
        structure = self._extract_structure(content)
        
        # Extraire les formules types
        formules = self._extract_formules(content)
        
        # Analyser la mise en forme
        mise_en_forme = self._analyze_formatting(content)
        
        # Analyser le vocabulaire
        vocabulaire = self._analyze_vocabulary(content)
        
        # Extraire des paragraphes types
        paragraphes_types = self._extract_sample_paragraphs(content)
        
        pattern = StylePattern(
            document_id=document.id,
            type_acte=type_acte,
            structure=structure,
            formules=list(formules),
            mise_en_forme=mise_en_forme,
            vocabulaire=vocabulaire,
            paragraphes_types=paragraphes_types
        )
        
        # Stocker le pattern
        self.patterns[type_acte].append(pattern)
        
        return pattern
    
    def analyze_word_document(self, doc_bytes: bytes, type_acte: str) -> Optional[StylePattern]:
        """Analyse un document Word pour en extraire le style"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            doc = DocxDocument(io.BytesIO(doc_bytes))
            
            # Extraire le contenu et la structure
            content = []
            structure = {
                'sections': [],
                'styles_utilises': set(),
                'mise_en_forme_paragraphes': []
            }
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
                    
                    # Analyser le style du paragraphe
                    para_style = {
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': str(paragraph.alignment) if paragraph.alignment else 'LEFT',
                        'first_line_indent': paragraph.paragraph_format.first_line_indent,
                        'left_indent': paragraph.paragraph_format.left_indent,
                        'right_indent': paragraph.paragraph_format.right_indent,
                        'space_before': paragraph.paragraph_format.space_before,
                        'space_after': paragraph.paragraph_format.space_after,
                        'line_spacing': paragraph.paragraph_format.line_spacing
                    }
                    
                    structure['mise_en_forme_paragraphes'].append(para_style)
                    
                    if paragraph.style:
                        structure['styles_utilises'].add(paragraph.style.name)
                    
                    # Détecter les sections
                    if paragraph.style and 'Heading' in paragraph.style.name:
                        structure['sections'].append({
                            'titre': text,
                            'niveau': paragraph.style.name
                        })
            
            # Créer un document temporaire pour l'analyse
            temp_doc = Document(
                id=f"word_doc_{datetime.now().timestamp()}",
                title=type_acte,
                content='\n'.join(content),
                source='word'
            )
            
            # Analyser avec la méthode standard
            pattern = self.analyze_document(temp_doc, type_acte)
            
            # Enrichir avec les informations Word spécifiques
            pattern.structure.update({
                'word_styles': list(structure['styles_utilises']),
                'word_formatting': structure['mise_en_forme_paragraphes'][:10]  # Garder un échantillon
            })
            
            return pattern
            
        except Exception as e:
            logger.error(f"Erreur analyse document Word: {e}")
            return None
    
    def _extract_structure(self, content: str) -> Dict[str, Any]:
        """Extrait la structure du document"""
        lines = content.split('\n')
        structure = {
            'sections': [],
            'niveau_hierarchie': 0,
            'longueur_sections': []
        }
        
        current_section = None
        section_content = []
        
        for line in lines:
            # Détecter les titres (lignes en majuscules, numérotées, etc.)
            if self._is_title(line):
                if current_section:
                    structure['sections'].append({
                        'titre': current_section,
                        'longueur': len(section_content)
                    })
                    structure['longueur_sections'].append(len(section_content))
                
                current_section = line.strip()
                section_content = []
            else:
                section_content.append(line)
        
        # Ajouter la dernière section
        if current_section:
            structure['sections'].append({
                'titre': current_section,
                'longueur': len(section_content)
            })
        
        return structure
    
    def _extract_formules(self, content: str) -> Set[str]:
        """Extrait les formules types du document"""
        formules = set()
        
        # Patterns de formules juridiques courantes
        patterns = [
            r"J'ai l'honneur de.*?[.!]",
            r"Il résulte de.*?[.!]",
            r"Aux termes de.*?[.!]",
            r"En l'espèce.*?[.!]",
            r"Par ces motifs.*?[.!]",
            r"Il convient de.*?[.!]",
            r"Force est de constater.*?[.!]",
            r"Il apparaît que.*?[.!]",
            r"Attendu que.*?[.!]",
            r"Considérant que.*?[.!]",
            r"Je vous prie d'agréer.*?[.!]",
            r"Veuillez agréer.*?[.!]",
            r"Dans l'attente de.*?[.!]",
            r"Je reste à votre disposition.*?[.!]"
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                formule = match.group(0).strip()
                if len(formule) < 200:  # Éviter les formules trop longues
                    formules.add(formule)
        
        return formules
    
    def _analyze_formatting(self, content: str) -> Dict[str, Any]:
        """Analyse la mise en forme du document"""
        return {
            'longueur_moyenne_paragraphe': self._avg_paragraph_length(content),
            'utilise_tirets': '-' in content,
            'utilise_numerotation': bool(re.search(r'\d+\.', content)),
            'utilise_majuscules_titres': bool(re.search(r'^[A-Z\s]+$', content, re.MULTILINE)),
            'espacement_sections': content.count('\n\n')
        }
    
    def _analyze_vocabulary(self, content: str) -> Dict[str, int]:
        """Analyse le vocabulaire utilisé"""
        # Nettoyer le texte
        words = re.findall(r'\b[a-zA-ZÀ-ÿ]+\b', content.lower())
        
        # Compter les fréquences
        word_freq = defaultdict(int)
        for word in words:
            if len(word) > 3:  # Ignorer les mots courts
                word_freq[word] += 1
        
        # Garder les mots les plus fréquents
        return dict(sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:100])
    
    def _extract_sample_paragraphs(self, content: str) -> List[str]:
        """Extrait des paragraphes types"""
        paragraphs = content.split('\n\n')
        
        # Filtrer les paragraphes intéressants
        samples = []
        for para in paragraphs:
            if 50 < len(para) < 500:  # Longueur raisonnable
                samples.append(para.strip())
        
        return samples[:10]  # Garder max 10 exemples
    
    def _is_title(self, line: str) -> bool:
        """Détermine si une ligne est un titre"""
        line = line.strip()
        
        if not line:
            return False
        
        # Critères pour identifier un titre
        if line.isupper() and len(line) > 3:
            return True
        
        if re.match(r'^[IVX]+\.?\s+', line):  # Numérotation romaine
            return True
        
        if re.match(r'^\d+\.?\s+', line):  # Numérotation arabe
            return True
        
        if re.match(r'^[A-Z]\.\s+', line):  # Lettre majuscule
            return True
        
        return False
    
    def _avg_paragraph_length(self, content: str) -> int:
        """Calcule la longueur moyenne des paragraphes"""
        paragraphs = content.split('\n\n')
        lengths = [len(p) for p in paragraphs if p.strip()]
        
        return sum(lengths) // len(lengths) if lengths else 0
    
    def generate_with_style(self, type_acte: str, contenu_base: str) -> str:
        """Génère du contenu en appliquant le style appris"""
        if type_acte not in self.patterns:
            return contenu_base
        
        patterns = self.patterns[type_acte]
        if not patterns:
            return contenu_base
        
        # Utiliser le premier pattern comme référence
        pattern = patterns[0]
        
        # Appliquer la structure
        styled_content = self._apply_structure(contenu_base, pattern.structure)
        
        # Insérer des formules types
        styled_content = self._insert_formules(styled_content, pattern.formules)
        
        # Appliquer la mise en forme
        styled_content = self._apply_formatting(styled_content, pattern.mise_en_forme)
        
        return styled_content
    
    def _apply_structure(self, content: str, structure: Dict[str, Any]) -> str:
        """Applique une structure au contenu"""
        # Diviser le contenu en sections basées sur la structure modèle
        sections = structure.get('sections', [])
        
        if not sections:
            return content
        
        # Restructurer le contenu
        lines = content.split('\n')
        structured = []
        
        section_size = len(lines) // len(sections) if sections else len(lines)
        
        for i, section in enumerate(sections):
            # Ajouter le titre de section
            structured.append(f"\n{section['titre']}\n")
            
            # Ajouter le contenu de la section
            start = i * section_size
            end = start + section_size if i < len(sections) - 1 else len(lines)
            
            structured.extend(lines[start:end])
        
        return '\n'.join(structured)
    
    def _insert_formules(self, content: str, formules: List[str]) -> str:
        """Insère des formules types dans le contenu"""
        if not formules:
            return content
        
        # Insérer quelques formules au début des paragraphes
        paragraphs = content.split('\n\n')
        
        for i in range(0, len(paragraphs), 3):  # Tous les 3 paragraphes
            if i < len(paragraphs) and formules:
                formule = formules[i % len(formules)]
                # Adapter la formule au contexte
                paragraphs[i] = f"{formule} {paragraphs[i]}"
        
        return '\n\n'.join(paragraphs)
    
    def _apply_formatting(self, content: str, formatting: Dict[str, Any]) -> str:
        """Applique la mise en forme au contenu"""
        # Ajuster l'espacement
        if formatting.get('espacement_sections', 0) > 1:
            content = re.sub(r'\n{2,}', '\n\n\n', content)
        
        # Ajouter la numérotation si nécessaire
        if formatting.get('utilise_numerotation'):
            lines = content.split('\n')
            numbered_lines = []
            counter = 1
            
            for line in lines:
                if line.strip() and not self._is_title(line):
                    numbered_lines.append(f"{counter}. {line}")
                    counter += 1
                else:
                    numbered_lines.append(line)
            
            content = '\n'.join(numbered_lines)
        
        return content

# ================== CONFIGURATION LLM ==================

class LLMConfig:
    """Configuration des LLMs"""
    
    @staticmethod
    def get_configs() -> Dict[LLMProvider, Dict[str, Any]]:
        return {
            LLMProvider.AZURE_OPENAI: {
                'endpoint': os.getenv('AZURE_OPENAI_ENDPOINT'),
                'key': os.getenv('AZURE_OPENAI_KEY'),
                'deployment': os.getenv('AZURE_OPENAI_DEPLOYMENT', 'gpt-4'),
                'api_version': '2024-02-01'
            },
            LLMProvider.CLAUDE_OPUS: {
                'api_key': os.getenv('ANTHROPIC_API_KEY'),
                'model': 'claude-3-opus-20240229'
            },
            LLMProvider.CHATGPT_4O: {
                'api_key': os.getenv('OPENAI_API_KEY'),
                'model': 'gpt-4-turbo-preview'
            },
            LLMProvider.GEMINI: {
                'api_key': os.getenv('GOOGLE_API_KEY'),
                'model': 'gemini-pro'
            },
            LLMProvider.PERPLEXITY: {
                'api_key': os.getenv('PERPLEXITY_API_KEY'),
                'model': 'pplx-70b-online'
            }
        }

# ================== GESTIONNAIRE MULTI-LLM ==================

class MultiLLMManager:
    """Gestionnaire pour interroger plusieurs LLMs"""
    
    def __init__(self):
        self.configs = LLMConfig.get_configs()
        self.clients = self._initialize_clients()
        self.executor = ThreadPoolExecutor(max_workers=5)
    
    def _initialize_clients(self) -> Dict[LLMProvider, Any]:
        """Initialise les clients LLM"""
        # Nettoyer l'environnement pour Hugging Face
        clean_env_for_azure()
        
        clients = {}
        
        # Azure OpenAI
        if self.configs[LLMProvider.AZURE_OPENAI]['key']:
            try:
                if AzureOpenAI:
                    clients[LLMProvider.AZURE_OPENAI] = AzureOpenAI(
                        azure_endpoint=self.configs[LLMProvider.AZURE_OPENAI]['endpoint'],
                        api_key=self.configs[LLMProvider.AZURE_OPENAI]['key'],
                        api_version=self.configs[LLMProvider.AZURE_OPENAI]['api_version']
                    )
            except Exception as e:
                logger.warning(f"Azure OpenAI non disponible: {e}")
        
        # Claude
        if self.configs[LLMProvider.CLAUDE_OPUS]['api_key']:
            try:
                if anthropic:
                    clients[LLMProvider.CLAUDE_OPUS] = anthropic.Anthropic(
                        api_key=self.configs[LLMProvider.CLAUDE_OPUS]['api_key']
                    )
            except Exception as e:
                logger.warning(f"Claude non disponible: {e}")
        
        # ChatGPT
        if self.configs[LLMProvider.CHATGPT_4O]['api_key']:
            try:
                if OpenAI:
                    clients[LLMProvider.CHATGPT_4O] = OpenAI(
                        api_key=self.configs[LLMProvider.CHATGPT_4O]['api_key']
                    )
            except Exception as e:
                logger.warning(f"ChatGPT non disponible: {e}")
        
        # Gemini
        if self.configs[LLMProvider.GEMINI]['api_key']:
            try:
                if genai:
                    genai.configure(api_key=self.configs[LLMProvider.GEMINI]['api_key'])
                    clients[LLMProvider.GEMINI] = genai.GenerativeModel(
                        self.configs[LLMProvider.GEMINI]['model']
                    )
            except Exception as e:
                logger.warning(f"Gemini non disponible: {e}")
        
        # Perplexity
        if self.configs[LLMProvider.PERPLEXITY]['api_key']:
            try:
                if OpenAI:
                    clients[LLMProvider.PERPLEXITY] = OpenAI(
                        api_key=self.configs[LLMProvider.PERPLEXITY]['api_key'],
                        base_url="https://api.perplexity.ai"
                    )
            except Exception as e:
                logger.warning(f"Perplexity non disponible: {e}")
        
        return clients
    
    async def query_single_llm(self, 
                              provider: LLMProvider, 
                              prompt: str,
                              system_prompt: str = None) -> Dict[str, Any]:
        """Interroge un seul LLM"""
        
        if provider not in self.clients:
            return {
                'provider': provider.value,
                'success': False,
                'error': 'Provider non configuré',
                'response': None
            }
        
        try:
            client = self.clients[provider]
            
            # Azure OpenAI
            if provider == LLMProvider.AZURE_OPENAI:
                messages = []
                if system_prompt:
                    messages.append({"role": "system", "content": system_prompt})
                messages.append({"role": "user", "content": prompt})
                
                response = client.chat.completions.create(
                    model=self.configs[provider]['deployment'],
                    messages=messages,
                    temperature=0.7,
                    max_tokens=4000
                )
                
                return {
                    'provider': provider.value,
                    'success': True,
                    'response': response.choices[0].message.content,
                    'usage': response.usage.model_dump() if response.usage else None
                }
            
            # Claude
            elif provider == LLMProvider.CLAUDE_OPUS:
                messages = []
                messages.append({"role": "user", "content": prompt})
                
                response = client.messages.create(
                    model=self.configs[provider]['model'],
                    messages=messages,
                    system=system_prompt if system_prompt else "Tu es un assistant juridique expert en droit pénal des affaires.",
                    max_tokens=4000
                )
                
                return {
                    'provider': provider.value,
                    'success': True,
                    'response': response.content[0].text,
                    'usage': {'total_tokens': response.usage.input_tokens + response.usage.output_tokens}
                }
            
            # ChatGPT 4o
            elif provider == LLMProvider.CHATGPT_4O:
                messages = []
                if system_prompt:
                    messages.append({"role": "system", "content": system_prompt})
                messages.append({"role": "user", "content": prompt})
                
                response = client.chat.completions.create(
                    model=self.configs[provider]['model'],
                    messages=messages,
                    temperature=0.7,
                    max_tokens=4000
                )
                
                return {
                    'provider': provider.value,
                    'success': True,
                    'response': response.choices[0].message.content,
                    'usage': response.usage.model_dump() if response.usage else None
                }
            
            # Gemini
            elif provider == LLMProvider.GEMINI:
                full_prompt = f"{system_prompt}\n\n{prompt}" if system_prompt else prompt
                response = client.generate_content(full_prompt)
                
                return {
                    'provider': provider.value,
                    'success': True,
                    'response': response.text,
                    'usage': None
                }
            
            # Perplexity
            elif provider == LLMProvider.PERPLEXITY:
                messages = []
                if system_prompt:
                    messages.append({"role": "system", "content": system_prompt})
                messages.append({"role": "user", "content": prompt})
                
                response = client.chat.completions.create(
                    model=self.configs[provider]['model'],
                    messages=messages,
                    temperature=0.7,
                    max_tokens=4000
                )
                
                return {
                    'provider': provider.value,
                    'success': True,
                    'response': response.choices[0].message.content,
                    'usage': response.usage.model_dump() if response.usage else None,
                    'citations': getattr(response, 'citations', [])
                }
            
        except Exception as e:
            logger.error(f"Erreur {provider.value}: {str(e)}")
            return {
                'provider': provider.value,
                'success': False,
                'error': str(e),
                'response': None
            }
    
    async def query_multiple_llms(self, providers: List[LLMProvider], prompt: str, 
                                 system_prompt: str = None) -> List[Dict[str, Any]]:
        """Interroge plusieurs LLMs en parallèle"""
        tasks = []
        
        for provider in providers:
            task = self.query_single_llm(provider, prompt, system_prompt)
            tasks.append(task)
        
        results = await asyncio.gather(*tasks)
        return results
    
    def fusion_responses(self, responses: List[Dict[str, Any]]) -> str:
        """Fusionne intelligemment plusieurs réponses"""
        valid_responses = [r for r in responses if r['success']]
        
        if not valid_responses:
            return "Aucune réponse valide obtenue."
        
        if len(valid_responses) == 1:
            return valid_responses[0]['response']
        
        # Construire un prompt de fusion
        fusion_prompt = "Voici plusieurs analyses d'experts. Synthétise-les en gardant les points essentiels:\n\n"
        
        for resp in valid_responses:
            fusion_prompt += f"### {resp['provider']}\n{resp['response']}\n\n"
        
        # Utiliser le premier LLM disponible pour la fusion
        if self.clients:
            provider = list(self.clients.keys())[0]
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            fusion_result = loop.run_until_complete(
                self.query_single_llm(
                    provider,
                    fusion_prompt,
                    "Tu es un expert en synthèse. Fusionne ces analyses en gardant le meilleur de chaque."
                )
            )
            
            if fusion_result['success']:
                return fusion_result['response']
        
        # Fallback: concatenation simple
        return "\n\n".join([f"### {r['provider']}\n{r['response']}" for r in valid_responses])

# ================== GÉNÉRATEURS DYNAMIQUES ==================

async def generate_dynamic_search_prompts(search_query: str, context: str = "") -> Dict[str, Dict[str, List[str]]]:
    """Génère dynamiquement des prompts de recherche basés sur la requête"""
    llm_manager = MultiLLMManager()
    
    # Utiliser Claude Opus 4 et ChatGPT 4o si disponibles
    preferred_providers = []
    if LLMProvider.CLAUDE_OPUS in llm_manager.clients:
        preferred_providers.append(LLMProvider.CLAUDE_OPUS)
    if LLMProvider.CHATGPT_4O in llm_manager.clients:
        preferred_providers.append(LLMProvider.CHATGPT_4O)
    
    if not preferred_providers and llm_manager.clients:
        preferred_providers = [list(llm_manager.clients.keys())[0]]
    
    if not preferred_providers:
        # Retour aux prompts statiques si aucun LLM disponible
        return {
            "🔍 Recherches suggérées": {
                "Générique": [
                    f"{search_query} jurisprudence récente",
                    f"{search_query} éléments constitutifs",
                    f"{search_query} moyens de défense",
                    f"{search_query} sanctions encourues"
                ]
            }
        }
    
    prompt = f"""En tant qu'expert en droit pénal des affaires, génère des prompts de recherche juridique pertinents basés sur cette requête : "{search_query}"
{f"Contexte supplémentaire : {context}" if context else ""}
Crée une structure JSON avec des catégories et sous-catégories de prompts de recherche.
Chaque prompt doit être concis (max 80 caractères) et cibler un aspect juridique précis.
Format attendu :
{{
    "🔍 Éléments constitutifs": {{
        "Élément matériel": ["prompt1", "prompt2", ...],
        "Élément intentionnel": ["prompt1", "prompt2", ...]
    }},
    "⚖️ Jurisprudence": {{
        "Décisions récentes": ["prompt1", "prompt2", ...],
        "Arrêts de principe": ["prompt1", "prompt2", ...]
    }},
    "🛡️ Moyens de défense": {{
        "Exceptions": ["prompt1", "prompt2", ...],
        "Stratégies": ["prompt1", "prompt2", ...]
    }}
}}
Génère au moins 3 catégories avec 2 sous-catégories chacune, et 4 prompts par sous-catégorie."""
    
    system_prompt = """Tu es un avocat spécialisé en droit pénal des affaires avec 20 ans d'expérience.
Tu maîtrises parfaitement la recherche juridique et sais formuler des requêtes précises pour trouver
la jurisprudence, la doctrine et les textes pertinents. Tes prompts sont toujours en français,
techniquement précis et adaptés au contexte du droit pénal économique français."""
    
    try:
        response = await llm_manager.query_single_llm(
            preferred_providers[0],
            prompt,
            system_prompt
        )
        
        if response['success']:
            # Extraire le JSON de la réponse
            import re
            json_match = re.search(r'\{[\s\S]*\}', response['response'])
            if json_match:
                try:
                    return json.loads(json_match.group(0))
                except json.JSONDecodeError:
                    pass
        
    except Exception as e:
        logger.error(f"Erreur génération prompts dynamiques: {e}")
    
    # Fallback
    return {
        "🔍 Recherches suggérées": {
            "Générique": [
                f"{search_query} jurisprudence",
                f"{search_query} éléments constitutifs",
                f"{search_query} défense",
                f"{search_query} sanctions"
            ]
        }
    }

async def generate_dynamic_templates(type_acte: str, context: Dict[str, Any] = None) -> Dict[str, str]:
    """Génère dynamiquement des modèles d'actes juridiques"""
    llm_manager = MultiLLMManager()
    
    # Utiliser Claude Opus 4 et ChatGPT 4o si disponibles
    preferred_providers = []
    if LLMProvider.CLAUDE_OPUS in llm_manager.clients:
        preferred_providers.append(LLMProvider.CLAUDE_OPUS)
    if LLMProvider.CHATGPT_4O in llm_manager.clients:
        preferred_providers.append(LLMProvider.CHATGPT_4O)
    
    if not preferred_providers and llm_manager.clients:
        preferred_providers = [list(llm_manager.clients.keys())[0]]
    
    if not preferred_providers:
        return {}
    
    context_str = ""
    if context:
        context_str = f"""
Contexte spécifique :
- Client : {context.get('client', 'Non spécifié')}
- Infraction : {context.get('infraction', 'Non spécifiée')}
- Juridiction : {context.get('juridiction', 'Non spécifiée')}
"""
    
    prompt = f"""Génère 3 modèles d'actes juridiques pour : "{type_acte}"
{context_str}
Pour chaque modèle, fournis :
1. Un titre descriptif avec emoji (ex: "📨 Demande d'audition libre")
2. Le contenu complet du modèle avec les balises [CHAMP] pour les éléments à personnaliser
Utilise un style juridique professionnel, formel et conforme aux usages du barreau français.
Les modèles doivent être immédiatement utilisables par un avocat.
Format de réponse attendu (JSON) :
{{
    "📄 Modèle standard de {type_acte}": "Contenu du modèle...",
    "⚖️ Modèle approfondi de {type_acte}": "Contenu du modèle...",
    "🔍 Modèle détaillé de {type_acte}": "Contenu du modèle..."
}}"""
    
    system_prompt = """Tu es un avocat au barreau de Paris, spécialisé en droit pénal des affaires.
Tu rédiges des actes juridiques depuis 20 ans et maîtrises parfaitement les formules consacrées,
la structure des actes et les mentions obligatoires. Tes modèles sont toujours conformes
aux exigences procédurales et aux usages de la profession."""
    
    try:
        # Interroger les LLMs
        responses = await llm_manager.query_multiple_llms(
            preferred_providers,
            prompt,
            system_prompt
        )
        
        # Fusionner les réponses si plusieurs LLMs
        if len(responses) > 1:
            fusion_prompt = f"""Voici plusieurs propositions de modèles pour "{type_acte}".
Fusionne-les intelligemment pour créer les 3 meilleurs modèles en gardant le meilleur de chaque proposition.
{chr(10).join([f"Proposition {i+1}: {r['response']}" for i, r in enumerate(responses) if r['success']])}
Retourne un JSON avec 3 modèles fusionnés."""
            
            fusion_response = await llm_manager.query_single_llm(
                preferred_providers[0],
                fusion_prompt,
                "Tu es un expert en fusion de contenus juridiques."
            )
            
            if fusion_response['success']:
                response_text = fusion_response['response']
            else:
                response_text = responses[0]['response'] if responses[0]['success'] else ""
        else:
            response_text = responses[0]['response'] if responses and responses[0]['success'] else ""
        
        # Extraire le JSON
        if response_text:
            import re
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                try:
                    return json.loads(json_match.group(0))
                except json.JSONDecodeError:
                    pass
        
    except Exception as e:
        logger.error(f"Erreur génération modèles dynamiques: {e}")
    
    # Fallback avec un modèle basique
    return {
        f"📄 Modèle standard de {type_acte}": f"""[EN-TÊTE AVOCAT]

À l'attention de [DESTINATAIRE]

Objet : {type_acte}
Référence : [RÉFÉRENCE]

[FORMULE D'APPEL],

J'ai l'honneur de [OBJET DE LA DEMANDE].

[DÉVELOPPEMENT]

[CONCLUSION]

Je vous prie d'agréer, [FORMULE DE POLITESSE].

[SIGNATURE]"""
    }

# ================== GESTIONNAIRE DE PAPIER EN-TÊTE ==================

def create_letterhead_from_template(template: LetterheadTemplate, content: str) -> str:
    """Crée un document avec papier en-tête à partir d'un template"""
    if not DOCX_AVAILABLE:
        return content
    
    try:
        doc = DocxDocument()
        
        # Définir les marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(template.margins['top'])
            section.bottom_margin = Inches(template.margins['bottom'])
            section.left_margin = Inches(template.margins['left'])
            section.right_margin = Inches(template.margins['right'])
        
        # Ajouter l'en-tête
        if template.header_content:
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = template.header_content
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter le contenu principal
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            p = doc.add_paragraph(para_text)
            p.style.font.name = template.font_family
            p.style.font.size = Pt(template.font_size)
            p.paragraph_format.line_spacing = template.line_spacing
        
        # Ajouter le pied de page
        if template.footer_content:
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = template.footer_content
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sauvegarder en mémoire
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer
        
    except Exception as e:
        logger.error(f"Erreur création document avec papier en-tête: {e}")
        return None

# ================== ANALYSE PROMPTS PÉNAL DES AFFAIRES ==================

ANALYSIS_PROMPTS_AFFAIRES = {
    "🎯 Analyse infractions économiques": [
        "Analysez précisément les éléments constitutifs de l'infraction reprochée",
        "Identifiez l'élément intentionnel et les moyens de le contester",
        "Recherchez les causes d'exonération ou d'atténuation",
        "Proposez une stratégie axée sur la bonne foi et l'intérêt social"
    ],
    "🏢 Responsabilité personne morale": [
        "Vérifiez les conditions d'imputation à la personne morale",
        "Analysez si les faits ont été commis pour le compte de la PM",
        "Examinez le rôle des organes et représentants",
        "Évaluez l'impact d'une éventuelle délégation de pouvoirs"
    ],
    "🛡️ Moyens de défense affaires": [
        "Valorisez le programme de conformité existant",
        "Démontrez les mesures correctives prises",
        "Argumentez sur l'absence d'enrichissement personnel",
        "Mettez en avant la transparence et la bonne gouvernance"
    ],
    "💰 Enjeux financiers": [
        "Calculez précisément le préjudice allégué",
        "Contestez les méthodes de calcul du préjudice",
        "Évaluez l'impact financier des sanctions encourues",
        "Proposez des modalités de réparation adaptées"
    ]
}

# ================== INTERFACE PRINCIPALE ==================

def main():
    """Interface principale de l'application"""
    
    # Initialisation
    initialize_session_state()
    
    # Initialiser les gestionnaires dans session state
    if 'azure_blob_manager' not in st.session_state or st.session_state.azure_blob_manager is None:
        try:
            st.session_state.azure_blob_manager = AzureBlobManager()
        except Exception as e:
            logger.error(f"Erreur création Blob Manager: {e}")
            st.session_state.azure_blob_manager = None
    
    if 'azure_search_manager' not in st.session_state or st.session_state.azure_search_manager is None:
        try:
            st.session_state.azure_search_manager = AzureSearchManager()
        except Exception as e:
            logger.error(f"Erreur création Search Manager: {e}")
            st.session_state.azure_search_manager = None
    
    # Titre principal avec style
    st.markdown("""
    <div style='text-align: center; padding: 2rem 0;'>
        <h1 style='color: #1a237e; font-size: 3rem;'>⚖️ Assistant Pénal des Affaires IA</h1>
        <p style='color: #666; font-size: 1.2rem;'>Intelligence artificielle au service du droit pénal économique</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.markdown("### 🎯 Navigation")
        
        # Menu principal
        page = st.selectbox(
            "Choisir une fonctionnalité",
            [
                "🔍 Recherche de documents",
                "📁 Sélection de pièces",
                "🤖 Analyse IA",
                "📝 Rédaction assistée",
                "✉️ Rédaction de courrier",
                "⚙️ Configuration"
            ],
            key="main_navigation"
        )
        
        st.markdown("---")
        
        # Infos système
        st.markdown("### 📊 État du système")
        
        # État Azure
        if AZURE_AVAILABLE:
            # Vérifier Azure Blob
            if st.session_state.azure_blob_manager and st.session_state.azure_blob_manager.is_connected():
                st.success("✅ Azure Blob connecté")
            else:
                st.error("❌ Azure Blob non connecté")
            
            # Vérifier Azure Search
            if st.session_state.azure_search_manager and st.session_state.azure_search_manager.search_client:
                st.success("✅ Azure Search connecté")
            else:
                st.warning("⚠️ Azure Search non disponible")
        else:
            st.warning("⚠️ Modules Azure non installés")
        
        # Documents chargés
        nb_docs = len(st.session_state.azure_documents)
        nb_pieces = len(st.session_state.pieces_selectionnees)
        
        st.metric("Documents disponibles", nb_docs)
        st.metric("Pièces sélectionnées", nb_pieces)
    
    # Pages
    if page == "🔍 Recherche de documents":
        page_recherche_documents()
    elif page == "📁 Sélection de pièces":
        page_selection_pieces()
    elif page == "🤖 Analyse IA":
        page_analyse_ia()
    elif page == "📝 Rédaction assistée":
        page_redaction_assistee()
    elif page == "✉️ Rédaction de courrier":
        page_redaction_courrier()
    elif page == "⚙️ Configuration":
        page_configuration()

def page_recherche_documents():
    """Page de recherche dans les documents Azure Blob"""
    
    st.header("🔍 Recherche de documents")
    
    # Vérifier la connexion Azure
    if not st.session_state.azure_blob_manager or not st.session_state.azure_blob_manager.is_connected():
        st.error("❌ Connexion Azure Blob non configurée. Veuillez vérifier vos variables d'environnement (AZURE_STORAGE_CONNECTION_STRING).")
        return
    
    # Section de recherche avec Azure Search
    if st.session_state.azure_search_manager and st.session_state.azure_search_manager.search_client:
        with st.container():
            st.markdown('<div class="search-section">', unsafe_allow_html=True)
            
            # Barre de recherche principale simplifiée
            with st.form(key="search_form"):
                col1, col2 = st.columns([4, 1])
                with col1:
                    search_query = st.text_input(
                        "Rechercher dans tous les documents",
                        value=st.session_state.search_query,
                        placeholder="Ex: abus de biens sociaux, délégation de pouvoirs, fraude fiscale...",
                        key="search_input_main"
                    )
                
                with col2:
                    search_clicked = st.form_submit_button("🔍 Rechercher", type="primary")
            
            # Options de recherche avancée (masquées par défaut)
            with st.expander("Options avancées"):
                search_mode = st.selectbox(
                    "Mode de recherche",
                    [mode.value for mode in SearchMode],
                    key="search_mode_select"
                )
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Effectuer la recherche
            if search_clicked and search_query:
                st.session_state.search_query = search_query
                
                with st.spinner("Recherche en cours..."):
                    # Mode de recherche
                    mode_enum = SearchMode.HYBRID
                    for mode in SearchMode:
                        if mode.value == search_mode:
                            mode_enum = mode
                            break
                    
                    results = st.session_state.azure_search_manager.search_hybrid(
                        search_query,
                        mode=mode_enum
                    )
                    
                    # Organiser les résultats : dossiers d'abord
                    folders_results = []
                    files_results = []
                    
                    # Simuler des dossiers basés sur les chemins
                    folder_paths = set()
                    for result in results:
                        if 'folder_path' in result and result['folder_path']:
                            # Extraire tous les niveaux de dossiers
                            parts = result['folder_path'].split('/')
                            for i in range(1, len(parts)):
                                folder_paths.add('/'.join(parts[:i]))
                    
                    # Créer des résultats pour les dossiers
                    for folder_path in sorted(folder_paths):
                        folders_results.append({
                            'type': 'folder',
                            'path': folder_path,
                            'name': folder_path.split('/')[-1],
                            'count': len([r for r in results if r.get('folder_path', '').startswith(folder_path)])
                        })
                    
                    # Ajouter les fichiers
                    for result in results:
                        result['type'] = 'file'
                        files_results.append(result)
                    
                    # Afficher les résultats
                    total_results = len(folders_results) + len(files_results)
                    if total_results > 0:
                        st.success(f"✅ {total_results} résultats trouvés ({len(folders_results)} dossiers, {len(files_results)} fichiers)")
                        
                        # Afficher les dossiers en premier
                        if folders_results:
                            st.markdown("### 📁 Dossiers")
                            for folder in folders_results:
                                with st.container():
                                    st.markdown('<div class="folder-card">', unsafe_allow_html=True)
                                    
                                    col1, col2 = st.columns([4, 1])
                                    
                                    with col1:
                                        st.markdown(f"📁 **{folder['name']}**")
                                        st.caption(f"Chemin: {folder['path']} | {folder['count']} documents")
                                    
                                    with col2:
                                        if st.button("➕ Ajouter tout", key=f"add_folder_{clean_key(folder['path'])}"):
                                            # Ajouter tous les fichiers du dossier
                                            added_count = 0
                                            with st.spinner(f"Ajout des documents du dossier {folder['name']}..."):
                                                for result in files_results:
                                                    if result.get('folder_path', '').startswith(folder['path']):
                                                        doc = Document(
                                                            id=result['id'],
                                                            title=result['title'],
                                                            content=result['content'],
                                                            source=result['source'],
                                                            folder_path=result.get('folder_path')
                                                        )
                                                        st.session_state.azure_documents[doc.id] = doc
                                                        added_count += 1
                                            
                                            if added_count > 0:
                                                st.success(f"✅ {added_count} documents ajoutés du dossier {folder['name']}")
                                    
                                    st.markdown('</div>', unsafe_allow_html=True)
                        
                        # Afficher les fichiers
                        if files_results:
                            st.markdown("### 📄 Documents")
                            for result in files_results:
                                with st.container():
                                    st.markdown('<div class="document-card">', unsafe_allow_html=True)
                                    
                                    col1, col2 = st.columns([4, 1])
                                    
                                    with col1:
                                        st.markdown(f"**{result['title']}**")
                                        st.caption(f"Score: {result['score']:.2f} | Source: {result['source']}")
                                        
                                        # Extrait du contenu
                                        excerpt = result['content'][:300] + "..." if len(result['content']) > 300 else result['content']
                                        st.text(excerpt)
                                    
                                    with col2:
                                        if result['id'] in st.session_state.azure_documents:
                                            st.success("✅ Déjà ajouté")
                                        else:
                                            if st.button("➕ Ajouter", key=f"add_search_{result['id']}"):
                                                # Créer un document
                                                doc = Document(
                                                    id=result['id'],
                                                    title=result['title'],
                                                    content=result['content'],
                                                    source=result['source'],
                                                    folder_path=result.get('folder_path')
                                                )
                                                
                                                st.session_state.azure_documents[doc.id] = doc
                                                st.success(f"✅ {doc.title} ajouté")
                                                st.rerun()
                                    
                                    st.markdown('</div>', unsafe_allow_html=True)
                    else:
                        st.info("Aucun résultat trouvé")
    
    # Navigation Azure Blob
    st.markdown("### 📂 Explorer les documents SharePoint")
    
    azure_manager = st.session_state.azure_blob_manager
    
    # Sélection du container
    containers = azure_manager.list_containers()
    
    if containers:
        col1, col2 = st.columns([3, 1])
        
        with col1:
            # Utiliser le container par défaut depuis la config
            default_index = 0
            if AppConfig.DEFAULT_CONTAINER in containers:
                default_index = containers.index(AppConfig.DEFAULT_CONTAINER)
            
            selected_container = st.selectbox(
                "Sélectionner un espace de stockage",
                containers,
                index=default_index,
                key="container_select"
            )
            
            st.session_state.selected_container = selected_container
        
        with col2:
            if st.button("🔄 Actualiser", key="refresh_containers"):
                st.rerun()
        
        # Navigation dans les dossiers
        if selected_container:
            st.markdown('<div class="folder-nav">', unsafe_allow_html=True)
            
            # Fil d'Ariane
            current_path = st.session_state.get('current_folder_path', '')
            if current_path:
                path_parts = current_path.split('/')
                path_parts = [p for p in path_parts if p]  # Enlever les parties vides
                
                breadcrumb = "📁 "
                if st.button("Racine", key="breadcrumb_root"):
                    st.session_state.current_folder_path = ""
                    st.rerun()
                
                for i, part in enumerate(path_parts):
                    breadcrumb += f" > "
                    partial_path = '/'.join(path_parts[:i+1]) + '/'
                    if st.button(part, key=f"breadcrumb_{clean_key(part)}_{i}"):
                        st.session_state.current_folder_path = partial_path
                        st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Lister le contenu du dossier actuel
            items = azure_manager.list_folders(selected_container, current_path)
            
            if items:
                # Séparer dossiers et fichiers
                folders = [item for item in items if item['type'] == 'folder']
                files = [item for item in items if item['type'] == 'file']
                
                # Afficher les dossiers en premier
                for item in folders + files:
                    with st.container():
                        if item['type'] == 'folder':
                            st.markdown('<div class="folder-card">', unsafe_allow_html=True)
                        else:
                            st.markdown('<div class="document-card">', unsafe_allow_html=True)
                        
                        col1, col2, col3 = st.columns([3, 1, 1])
                        
                        with col1:
                            if item['type'] == 'folder':
                                st.markdown(f"📁 **{item['name']}**")
                            else:
                                # Icône selon le type de fichier
                                file_ext = os.path.splitext(item['name'])[1].lower()
                                icon = {
                                    '.pdf': '📄',
                                    '.docx': '📝',
                                    '.doc': '📝',
                                    '.txt': '📃',
                                    '.xlsx': '📊',
                                    '.xls': '📊'
                                }.get(file_ext, '📎')
                                
                                st.markdown(f"{icon} **{item['name']}**")
                                
                                if item.get('size'):
                                    size_mb = item['size'] / (1024 * 1024)
                                    st.caption(f"Taille: {size_mb:.2f} MB")
                        
                        with col2:
                            if item.get('last_modified'):
                                st.caption(f"Modifié: {item['last_modified'].strftime('%d/%m/%Y')}")
                        
                        with col3:
                            if item['type'] == 'folder':
                                col_open, col_add = st.columns(2)
                                
                                with col_open:
                                    if st.button("📂", key=f"open_folder_{clean_key(item['name'])}", help="Ouvrir"):
                                        st.session_state.current_folder_path = item['path']
                                        st.rerun()
                                
                                with col_add:
                                    if st.button("➕", key=f"add_folder_all_{clean_key(item['path'])}", help="Ajouter tout le dossier"):
                                        # Ajouter tous les fichiers du dossier
                                        with st.spinner(f"Ajout du dossier {item['name']}..."):
                                            all_files = azure_manager.get_all_files_in_folder(selected_container, item['path'])
                                            
                                            added_count = 0
                                            for file_info in all_files:
                                                # Extraire le texte
                                                content = azure_manager.extract_text_from_blob(
                                                    selected_container,
                                                    file_info['full_path']
                                                )
                                                
                                                if content:
                                                    doc_id = f"azure_{clean_key(file_info['full_path'])}"
                                                    doc = Document(
                                                        id=doc_id,
                                                        title=file_info['name'],
                                                        content=content,
                                                        source='azure',
                                                        metadata={
                                                            'container': selected_container,
                                                            'path': file_info['full_path'],
                                                            'size': file_info.get('size'),
                                                            'last_modified': file_info.get('last_modified')
                                                        },
                                                        folder_path=file_info['folder']
                                                    )
                                                    
                                                    st.session_state.azure_documents[doc_id] = doc
                                                    
                                                    # Indexer dans Azure Search
                                                    if st.session_state.azure_search_manager and st.session_state.azure_search_manager.search_client:
                                                        st.session_state.azure_search_manager.index_document(doc)
                                                    
                                                    added_count += 1
                                            
                                            if added_count > 0:
                                                st.success(f"✅ {added_count} documents ajoutés du dossier {item['name']}")
                            else:
                                col_view, col_select = st.columns(2)
                                
                                with col_view:
                                    if st.button("👁️", key=f"view_file_{clean_key(item['full_path'])}"):
                                        # Extraire et afficher le contenu
                                        with st.spinner("Chargement..."):
                                            content = azure_manager.extract_text_from_blob(
                                                selected_container,
                                                item['full_path']
                                            )
                                            
                                            if content:
                                                st.text_area(
                                                    f"Contenu de {item['name']}",
                                                    content[:2000] + "..." if len(content) > 2000 else content,
                                                    height=300,
                                                    key=f"content_view_{clean_key(item['full_path'])}"
                                                )
                                
                                with col_select:
                                    # Créer un document et l'ajouter à la session
                                    doc_id = f"azure_{clean_key(item['full_path'])}"
                                    
                                    if doc_id in st.session_state.azure_documents:
                                        st.success("✅")
                                    else:
                                        if st.button("➕", key=f"add_doc_{doc_id}", help="Ajouter à la sélection"):
                                            # Télécharger le contenu si pas déjà fait
                                            with st.spinner("Ajout..."):
                                                content = azure_manager.extract_text_from_blob(
                                                    selected_container,
                                                    item['full_path']
                                                )
                                                
                                                if content:
                                                    doc = Document(
                                                        id=doc_id,
                                                        title=item['name'],
                                                        content=content,
                                                        source='azure',
                                                        metadata={
                                                            'container': selected_container,
                                                            'path': item['full_path'],
                                                            'size': item.get('size'),
                                                            'last_modified': item.get('last_modified')
                                                        },
                                                        folder_path=current_path
                                                    )
                                                    
                                                    st.session_state.azure_documents[doc_id] = doc
                                                    
                                                    # Indexer dans Azure Search si disponible
                                                    if st.session_state.azure_search_manager and st.session_state.azure_search_manager.search_client:
                                                        st.session_state.azure_search_manager.index_document(doc)
                                                    
                                                    st.success(f"✅ {item['name']} ajouté")
                                                    st.rerun()
                        
                        st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info("📭 Aucun document dans ce dossier")
    else:
        st.warning("Aucun container disponible")
    
    # Prompts de recherche suggérés dynamiques
    if st.session_state.search_query:
        st.markdown("### 💡 Recherches suggérées")
        
        # Générer des prompts dynamiques si une recherche est active
        search_query = st.session_state.search_query
        
        # Vérifier le cache
        cache_key = f"prompts_{clean_key(search_query)}"
        
        if cache_key not in st.session_state.dynamic_search_prompts:
            with st.spinner("Génération de suggestions intelligentes..."):
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                
                dynamic_prompts = loop.run_until_complete(
                    generate_dynamic_search_prompts(search_query)
                )
                
                st.session_state.dynamic_search_prompts[cache_key] = dynamic_prompts
        else:
            dynamic_prompts = st.session_state.dynamic_search_prompts[cache_key]
        
        # Afficher les prompts dynamiques
        for categorie, sous_categories in dynamic_prompts.items():
            with st.expander(categorie):
                for sous_cat, prompts in sous_categories.items():
                    st.markdown(f"**{sous_cat}**")
                    
                    # Créer une grille de boutons
                    cols = st.columns(2)
                    for idx, prompt in enumerate(prompts):
                        col = cols[idx % 2]
                        
                        with col:
                            # Clé unique pour chaque bouton
                            button_key = f"dyn_prompt_{clean_key(categorie)}_{clean_key(sous_cat)}_{idx}"
                            
                            if st.button(
                                prompt[:50] + "..." if len(prompt) > 50 else prompt,
                                key=button_key,
                                help=prompt
                            ):
                                st.session_state.search_query = prompt
                                st.rerun()

def page_selection_pieces():
    """Page de sélection et organisation des pièces"""
    
    st.header("📁 Sélection et organisation des pièces")
    
    # Gestionnaire de pièces
    gestionnaire = GestionnairePiecesSelectionnees()
    
    # Charger les pièces depuis la session
    if 'pieces_selectionnees' in st.session_state:
        gestionnaire.pieces = st.session_state.pieces_selectionnees
    
    # Vue d'ensemble
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Documents disponibles", len(st.session_state.azure_documents))
    
    with col2:
        st.metric("Pièces sélectionnées", len(gestionnaire.pieces))
    
    with col3:
        if st.button("📋 Générer bordereau", key="generer_bordereau"):
            bordereau = gestionnaire.generer_bordereau()
            st.text_area("Bordereau", bordereau, height=400, key="bordereau_display")
            
            # Option de téléchargement
            st.download_button(
                "💾 Télécharger le bordereau",
                bordereau,
                "bordereau_pieces.txt",
                "text/plain",
                key="download_bordereau"
            )
    
    # Deux colonnes : documents disponibles et pièces sélectionnées
    col_docs, col_pieces = st.columns([1, 1])
    
    # Colonne des documents disponibles
    with col_docs:
        st.markdown("### 📄 Documents disponibles")
        
        # Filtre
        filtre = st.text_input(
            "Filtrer les documents",
            placeholder="Rechercher par nom...",
            key="filtre_docs_pieces"
        )
        
        # Liste des documents
        for doc_id, doc in st.session_state.azure_documents.items():
            if filtre.lower() in doc.title.lower():
                with st.container():
                    st.markdown('<div class="document-card">', unsafe_allow_html=True)
                    
                    # Titre et métadonnées
                    st.markdown(f"**{doc.title}**")
                    
                    if doc.metadata.get('last_modified'):
                        st.caption(f"Modifié: {doc.metadata['last_modified']}")
                    
                    # Si déjà sélectionné, afficher différemment
                    if doc_id in gestionnaire.pieces:
                        st.success("✅ Déjà sélectionné")
                        
                        if st.button(f"❌ Retirer", key=f"remove_{doc_id}"):
                            gestionnaire.retirer_piece(doc_id)
                            st.rerun()
                    else:
                        # Sélection avec catégorie
                        col_cat, col_btn = st.columns([3, 1])
                        
                        with col_cat:
                            categorie = st.selectbox(
                                "Catégorie",
                                gestionnaire.categories,
                                key=f"cat_select_{doc_id}"
                            )
                        
                        with col_btn:
                            if st.button("➕", key=f"add_piece_{doc_id}"):
                                gestionnaire.ajouter_piece(doc, categorie)
                                st.rerun()
                    
                    st.markdown('</div>', unsafe_allow_html=True)
    
    # Colonne des pièces sélectionnées
    with col_pieces:
        st.markdown("### ✅ Pièces sélectionnées")
        
        pieces_par_cat = gestionnaire.get_pieces_par_categorie()
        
        for categorie, pieces in pieces_par_cat.items():
            if pieces:
                with st.expander(f"{categorie} ({len(pieces)})", expanded=True):
                    for piece in sorted(pieces, key=lambda x: x.pertinence, reverse=True):
                        st.markdown('<div class="piece-selectionnee">', unsafe_allow_html=True)
                        
                        col1, col2 = st.columns([3, 1])
                        
                        with col1:
                            st.markdown(f"**{piece.titre}**")
                            
                            # Notes
                            notes = st.text_input(
                                "Notes",
                                value=piece.notes,
                                key=f"notes_{piece.document_id}",
                                placeholder="Ajouter des notes..."
                            )
                            
                            # Pertinence
                            pertinence = st.slider(
                                "Pertinence",
                                1, 10,
                                value=piece.pertinence,
                                key=f"pertinence_{piece.document_id}"
                            )
                            
                            # Mettre à jour la pièce
                            if notes != piece.notes or pertinence != piece.pertinence:
                                piece.notes = notes
                                piece.pertinence = pertinence
                                st.session_state.pieces_selectionnees[piece.document_id] = piece
                        
                        with col2:
                            if st.button("🗑️", key=f"delete_piece_{piece.document_id}"):
                                gestionnaire.retirer_piece(piece.document_id)
                                st.rerun()
                        
                        st.markdown('</div>', unsafe_allow_html=True)

def page_analyse_ia():
    """Page d'analyse par IA"""
    
    st.header("🤖 Analyse IA des documents")
    
    # Vérifier qu'il y a des pièces sélectionnées
    if not st.session_state.pieces_selectionnees:
        st.warning("⚠️ Veuillez d'abord sélectionner des pièces dans l'onglet 'Sélection de pièces'")
        return
    
    # Configuration de l'analyse
    st.markdown("### ⚙️ Configuration de l'analyse")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Type d'infraction - Saisie libre avec suggestions
        infractions_list = [inf.value for inf in InfractionAffaires]
        
        infraction = st.text_input(
            "Type d'infraction",
            placeholder="Ex: Abus de biens sociaux, corruption, fraude fiscale...",
            key="infraction_input",
            help="Saisissez librement l'infraction ou choisissez dans les suggestions"
        )
        
        # Suggestions
        if not infraction:
            st.info("💡 Suggestions : " + ", ".join(infractions_list[:5]) + "...")
        
        # Client
        client_nom = st.text_input(
            "Nom du client",
            placeholder="Personne physique ou morale",
            key="client_nom_analyse"
        )
        
        client_type = st.radio(
            "Type de client",
            ["Personne physique", "Personne morale"],
            key="client_type_analyse"
        )
    
    with col2:
        # Providers IA à utiliser
        llm_manager = MultiLLMManager()
        available_providers = list(llm_manager.clients.keys())
        
        if available_providers:
            selected_providers = st.multiselect(
                "IA à utiliser",
                [p.value for p in available_providers],
                default=[available_providers[0].value] if available_providers else [],
                key="selected_providers_analyse"
            )
            
            # Mode de fusion
            fusion_mode = st.radio(
                "Mode de fusion des réponses",
                ["Synthèse IA", "Concatenation simple"],
                key="fusion_mode_analyse"
            )
        else:
            st.error("❌ Aucune IA configurée")
            return
    
    # Sélection des pièces à analyser
    st.markdown("### 📄 Pièces à analyser")
    
    pieces_a_analyser = []
    for piece_id, piece in st.session_state.pieces_selectionnees.items():
        if st.checkbox(
            f"{piece.titre} ({piece.categorie})",
            value=True,
            key=f"analyse_piece_{piece_id}"
        ):
            pieces_a_analyser.append(piece_id)
    
    # Type d'analyse
    st.markdown("### 🎯 Type d'analyse")
    
    analyse_types = st.multiselect(
        "Sélectionner les analyses à effectuer",
        list(ANALYSIS_PROMPTS_AFFAIRES.keys()),
        default=list(ANALYSIS_PROMPTS_AFFAIRES.keys())[:2],
        key="analyse_types_select"
    )
    
    # Bouton d'analyse
    if st.button("🚀 Lancer l'analyse", type="primary", key="lancer_analyse"):
        if not infraction or not client_nom or not pieces_a_analyser:
            st.error("❌ Veuillez remplir tous les champs")
            return
        
        # Préparer le contenu pour l'analyse
        contenu_pieces = []
        for piece_id in pieces_a_analyser:
            if piece_id in st.session_state.azure_documents:
                doc = st.session_state.azure_documents[piece_id]
                piece = st.session_state.pieces_selectionnees[piece_id]
                
                contenu_pieces.append(f"""
=== {doc.title} ({piece.categorie}) ===
Pertinence: {piece.pertinence}/10
Notes: {piece.notes}
Contenu:
{doc.content[:3000]}...
""")
        
        # Construire le prompt
        prompt_base = f"""Tu es un avocat expert en droit pénal des affaires.
Client: {client_nom} ({client_type})
Infraction reprochée: {infraction}

Documents analysés:
{chr(10).join(contenu_pieces)}

Analyses demandées:
"""
        
        # Lancer les analyses
        with st.spinner("🔄 Analyse en cours..."):
            resultats = {}
            
            # Pour chaque type d'analyse
            for analyse_type in analyse_types:
                prompts = ANALYSIS_PROMPTS_AFFAIRES[analyse_type]
                
                # Construire le prompt complet
                prompt_complet = prompt_base + f"\n{analyse_type}:\n"
                for p in prompts:
                    prompt_complet += f"- {p}\n"
                
                # Convertir les providers
                providers_enum = [
                    p for p in LLMProvider 
                    if p.value in selected_providers
                ]
                
                # Interroger les IA
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                
                responses = loop.run_until_complete(
                    llm_manager.query_multiple_llms(
                        providers_enum,
                        prompt_complet
                    )
                )
                
                # Fusionner les réponses
                if fusion_mode == "Synthèse IA" and len(responses) > 1:
                    fusion = llm_manager.fusion_responses(responses)
                else:
                    fusion = "\n\n".join([
                        f"### {r['provider']}\n{r['response']}"
                        for r in responses
                        if r['success']
                    ])
                
                resultats[analyse_type] = fusion
            
            # Afficher les résultats
            st.markdown("## 📊 Résultats de l'analyse")
            
            for analyse_type, resultat in resultats.items():
                with st.expander(analyse_type, expanded=True):
                    st.markdown(resultat)
                    
                    # Options d'export
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            "💾 Télécharger",
                            resultat,
                            f"analyse_{clean_key(analyse_type)}.txt",
                            "text/plain",
                            key=f"download_analyse_{clean_key(analyse_type)}"
                        )
                    
                    with col2:
                        if st.button(
                            "📋 Copier",
                            key=f"copy_analyse_{clean_key(analyse_type)}"
                        ):
                            st.write("Contenu copié!")

def page_redaction_assistee():
    """Page de rédaction assistée avec apprentissage de style"""
    
    st.header("📝 Rédaction assistée par IA")
    
    # Onglets pour organiser la fonctionnalité
    tabs = st.tabs(["✍️ Rédaction", "🎨 Apprentissage de style", "📚 Modèles"])
    
    # Onglet Rédaction
    with tabs[0]:
        st.markdown("### 📄 Créer un nouvel acte")
        
        # Type d'acte flexible
        col1, col2 = st.columns([2, 1])
        
        with col1:
            type_acte = st.text_input(
                "Type d'acte à rédiger",
                placeholder="Ex: Plainte avec constitution de partie civile, Conclusions, Demande d'audition...",
                key="type_acte_input"
            )
        
        with col2:
            # Utiliser un style appris ?
            if 'learned_styles' in st.session_state and st.session_state.learned_styles:
                use_style = st.checkbox("Utiliser un style appris", key="use_learned_style")
            else:
                use_style = False
                st.info("Aucun style appris")
        
        # Sélection du style si activé
        if use_style and st.session_state.learned_styles:
            selected_style = st.selectbox(
                "Choisir un style",
                list(st.session_state.learned_styles.keys()),
                key="select_style_redaction"
            )
        else:
            selected_style = None
        
        # Informations de base
        st.markdown("### 📋 Informations essentielles")
        
        col1, col2 = st.columns(2)
        
        with col1:
            destinataire = st.text_input(
                "Destinataire",
                placeholder="Ex: Monsieur le Procureur de la République",
                key="destinataire_acte"
            )
            
            client_nom = st.text_input(
                "Client",
                placeholder="Nom du client",
                key="client_nom_acte"
            )
            
            avocat_nom = st.text_input(
                "Avocat",
                placeholder="Maître...",
                key="avocat_nom_acte"
            )
        
        with col2:
            reference = st.text_input(
                "Référence",
                placeholder="N° de procédure, dossier...",
                key="reference_acte"
            )
            
            infraction = st.text_input(
                "Infraction(s)",
                placeholder="Ex: Abus de biens sociaux",
                key="infraction_acte"
            )
            
            date_faits = st.date_input(
                "Date des faits",
                key="date_faits_acte"
            )
        
        # Contenu principal
        st.markdown("### 📝 Contenu de l'acte")
        
        # Points clés à inclure
        points_cles = st.text_area(
            "Points clés à développer",
            placeholder="""Ex:
- Absence d'élément intentionnel
- Actions réalisées dans l'intérêt de la société
- Bonne foi du dirigeant
- Préjudice non caractérisé""",
            height=150,
            key="points_cles_acte"
        )
        
        # Pièces à mentionner
        if st.session_state.pieces_selectionnees:
            st.markdown("#### 📎 Pièces à citer")
            
            pieces_a_citer = []
            for piece_id, piece in st.session_state.pieces_selectionnees.items():
                if st.checkbox(
                    f"Pièce n°{len(pieces_a_citer)+1} : {piece.titre}",
                    key=f"cite_piece_{piece_id}"
                ):
                    pieces_a_citer.append(piece)
        
        # Options de génération
        st.markdown("### ⚙️ Options de génération")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            ton = st.select_slider(
                "Ton",
                options=["Très formel", "Formel", "Neutre", "Direct", "Combatif"],
                value="Formel",
                key="ton_generation"
            )
        
        with col2:
            longueur = st.select_slider(
                "Longueur",
                options=["Concis", "Standard", "Détaillé", "Très détaillé"],
                value="Standard",
                key="longueur_generation"
            )
        
        with col3:
            inclure_jurisprudence = st.checkbox(
                "Inclure des références jurisprudentielles",
                value=True,
                key="inclure_juris"
            )
        
        # Bouton de génération
        if st.button("🚀 Générer l'acte", type="primary", key="generer_acte"):
            if not type_acte:
                st.error("❌ Veuillez spécifier le type d'acte")
                return
            
            # Construire le prompt
            prompt = f"""Tu es un avocat expert en droit pénal des affaires.
Rédige un(e) {type_acte} avec les informations suivantes :

Destinataire : {destinataire or 'Non spécifié'}
Client : {client_nom or 'Non spécifié'}
Avocat : {avocat_nom or 'Non spécifié'}
Référence : {reference or 'Non spécifiée'}
Infraction(s) : {infraction or 'Non spécifiée'}
Date des faits : {date_faits.strftime('%d/%m/%Y') if date_faits else 'Non spécifiée'}

Points clés à développer :
{points_cles}

Ton souhaité : {ton}
Longueur : {longueur}
{"Inclure des références jurisprudentielles pertinentes" if inclure_jurisprudence else ""}

Structure l'acte de manière professionnelle avec :
- Un en-tête approprié
- Une introduction claire
- Un développement structuré des arguments
- Une conclusion percutante
- Les formules de politesse adaptées"""
            
            # Si un style est sélectionné, l'ajouter au prompt
            if selected_style and selected_style in st.session_state.learned_styles:
                style_info = st.session_state.learned_styles[selected_style]
                prompt += f"\n\nApplique le style suivant :\n{json.dumps(style_info, ensure_ascii=False, indent=2)}"
            
            # Générer avec l'IA
            llm_manager = MultiLLMManager()
            
            with st.spinner("🔄 Génération en cours..."):
                if llm_manager.clients:
                    provider = list(llm_manager.clients.keys())[0]
                    
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    
                    response = loop.run_until_complete(
                        llm_manager.query_single_llm(
                            provider,
                            prompt,
                            "Tu es un avocat spécialisé en droit pénal des affaires, expert en rédaction d'actes juridiques."
                        )
                    )
                    
                    if response['success']:
                        # Appliquer le style si nécessaire
                        contenu_genere = response['response']
                        
                        if selected_style and 'style_analyzer' not in st.session_state:
                            st.session_state.style_analyzer = StyleAnalyzer()
                        
                        if selected_style and hasattr(st.session_state, 'style_analyzer'):
                            contenu_genere = st.session_state.style_analyzer.generate_with_style(
                                selected_style,
                                contenu_genere
                            )
                        
                        # Afficher le résultat
                        st.markdown("### 📄 Acte généré")
                        
                        st.text_area(
                            "Contenu",
                            value=contenu_genere,
                            height=600,
                            key="acte_genere_content"
                        )
                        
                        # Options d'export
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.download_button(
                                "💾 Télécharger (.txt)",
                                contenu_genere,
                                f"{clean_key(type_acte)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                "text/plain",
                                key="download_txt_acte"
                            )
                        
                        with col2:
                            if DOCX_AVAILABLE:
                                # Créer un document Word
                                doc = DocxDocument()
                                
                                # Ajouter le contenu
                                for paragraph in contenu_genere.split('\n'):
                                    doc.add_paragraph(paragraph)
                                
                                # Sauvegarder en mémoire
                                docx_buffer = io.BytesIO()
                                doc.save(docx_buffer)
                                docx_buffer.seek(0)
                                
                                st.download_button(
                                    "💾 Télécharger (.docx)",
                                    docx_buffer.getvalue(),
                                    f"{clean_key(type_acte)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_docx_acte"
                                )
                        
                        with col3:
                            if st.button("📧 Préparer l'envoi", key="prepare_send"):
                                st.info("Fonctionnalité d'envoi à implémenter")
                    else:
                        st.error(f"❌ Erreur : {response['error']}")
                else:
                    st.error("❌ Aucune IA disponible")
    
    # Onglet Apprentissage de style
    with tabs[1]:
        st.markdown("### 🎨 Apprendre un style de rédaction")
        
        st.info("""
        Cette fonctionnalité analyse vos documents modèles pour apprendre votre style de rédaction :
        - Structure des documents
        - Formules types utilisées
        - Mise en forme préférée
        - Vocabulaire spécifique
        
        **Nouveauté** : Vous pouvez maintenant analyser des documents Word (.docx) directement !
        """)
        
        # Type d'acte pour le style
        type_style = st.text_input(
            "Nom du style à apprendre",
            placeholder="Ex: Plainte pénale, Conclusions de relaxe...",
            key="type_style_learn"
        )
        
        # Sélection des documents modèles
        if st.session_state.azure_documents:
            st.markdown("#### 📚 Sélectionner les documents modèles depuis Azure")
            
            # Sélection des documents
            docs_modeles = []
            for doc_id, doc in st.session_state.azure_documents.items():
                if st.checkbox(
                    f"{doc.title}",
                    key=f"model_{doc_id}",
                    help=f"Source: {doc.source}"
                ):
                    docs_modeles.append(doc)
        
        # Upload de documents Word
        st.markdown("#### 📤 Ou télécharger des documents Word")
        
        uploaded_files = st.file_uploader(
            "Choisir des fichiers Word (.docx)",
            type=['docx'],
            accept_multiple_files=True,
            key="upload_word_models"
        )
        
        if st.button("🧠 Apprendre le style", key="learn_style") and type_style:
            if not docs_modeles and not uploaded_files:
                st.error("❌ Veuillez sélectionner au moins un document modèle")
                return
            
            with st.spinner("Analyse en cours..."):
                # Initialiser l'analyseur
                if 'style_analyzer' not in st.session_state:
                    st.session_state.style_analyzer = StyleAnalyzer()
                
                patterns = []
                
                # Analyser les documents Azure
                for doc in docs_modeles:
                    pattern = st.session_state.style_analyzer.analyze_document(doc, type_style)
                    patterns.append(pattern)
                
                # Analyser les documents Word uploadés
                if uploaded_files and DOCX_AVAILABLE:
                    for uploaded_file in uploaded_files:
                        # Lire le contenu du fichier
                        doc_bytes = uploaded_file.read()
                        
                        # Analyser le document Word
                        pattern = st.session_state.style_analyzer.analyze_word_document(doc_bytes, type_style)
                        if pattern:
                            patterns.append(pattern)
                            st.success(f"✅ {uploaded_file.name} analysé")
                
                if patterns:
                    # Fusionner les patterns
                    merged_pattern = {
                        'nombre_documents': len(patterns),
                        'structure_commune': merge_structures([p.structure for p in patterns]),
                        'formules_frequentes': merge_formules([p.formules for p in patterns]),
                        'mise_en_forme_type': merge_formatting([p.mise_en_forme for p in patterns]),
                        'vocabulaire_cle': merge_vocabulary([p.vocabulaire for p in patterns])
                    }
                    
                    # Sauvegarder le style
                    if 'learned_styles' not in st.session_state:
                        st.session_state.learned_styles = {}
                    
                    st.session_state.learned_styles[type_style] = merged_pattern
                    
                    st.success(f"✅ Style '{type_style}' appris avec succès à partir de {len(patterns)} documents!")
                    
                    # Afficher un résumé
                    with st.expander("Voir le résumé du style appris"):
                        st.json(merged_pattern)
        else:
            st.warning("⚠️ Aucun document disponible. Chargez d'abord des documents.")
    
    # Onglet Modèles
    with tabs[2]:
        st.markdown("### 📚 Bibliothèque de modèles")
        
        # Options pour générer des modèles dynamiques
        st.markdown("#### 🤖 Générer des modèles personnalisés")
        
        col1, col2 = st.columns([3, 1])
        
        with col1:
            type_modele_generer = st.text_input(
                "Type d'acte pour lequel générer des modèles",
                placeholder="Ex: Plainte avec constitution de partie civile, Mémoire en défense...",
                key="type_modele_generer"
            )
        
        with col2:
            if st.button("🎯 Générer", key="generer_modeles_button"):
                if type_modele_generer:
                    with st.spinner("Génération de modèles intelligents..."):
                        # Contexte optionnel
                        context = {
                            'client': st.session_state.get('client_nom_acte', ''),
                            'infraction': st.session_state.get('infraction_acte', ''),
                            'juridiction': ''
                        }
                        
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        
                        modeles = loop.run_until_complete(
                            generate_dynamic_templates(type_modele_generer, context)
                        )
                        
                        # Stocker dans le cache
                        cache_key = f"templates_{clean_key(type_modele_generer)}"
                        st.session_state.dynamic_templates[cache_key] = modeles
                        
                        st.success("✅ Modèles générés avec succès!")
        
        # Afficher les modèles générés dynamiquement
        if st.session_state.dynamic_templates:
            st.markdown("#### 🎨 Modèles générés par IA")
            
            for cache_key, modeles in st.session_state.dynamic_templates.items():
                type_clean = cache_key.replace("templates_", "").replace("_", " ").title()
                
                with st.expander(f"📁 Modèles pour : {type_clean}"):
                    for titre, contenu in modeles.items():
                        st.markdown(f"**{titre}**")
                        
                        st.text_area(
                            "Modèle",
                            value=contenu,
                            height=300,
                            key=f"dyn_template_view_{clean_key(titre)}"
                        )
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if st.button(f"📋 Utiliser", key=f"use_dyn_template_{clean_key(titre)}"):
                                st.session_state.template_to_use = contenu
                                st.info("Modèle copié. Retournez à l'onglet Rédaction.")
                        
                        with col2:
                            st.download_button(
                                "💾 Télécharger",
                                contenu,
                                f"{clean_key(titre)}.txt",
                                "text/plain",
                                key=f"download_dyn_template_{clean_key(titre)}"
                            )
                        
                        st.markdown("---")

def page_redaction_courrier():
    """Page de rédaction de courrier avec papier en-tête"""
    
    st.header("✉️ Rédaction de courrier")
    
    # Vérifier si un papier en-tête est configuré
    if 'letterhead_template' not in st.session_state or st.session_state.letterhead_template is None:
        st.warning("⚠️ Aucun papier en-tête configuré. Veuillez d'abord configurer votre papier en-tête dans l'onglet Configuration.")
        
        # Option rapide pour créer un papier en-tête
        if st.button("➕ Créer un papier en-tête rapide"):
            with st.form("quick_letterhead"):
                st.markdown("### Configuration rapide du papier en-tête")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    cabinet_nom = st.text_input("Nom du cabinet", placeholder="Cabinet d'avocats...")
                    avocat_nom = st.text_input("Nom de l'avocat", placeholder="Maître...")
                    barreau = st.text_input("Barreau", placeholder="Barreau de Paris")
                
                with col2:
                    adresse = st.text_area("Adresse", placeholder="123 rue de la Justice\n75001 Paris")
                    telephone = st.text_input("Téléphone", placeholder="01 23 45 67 89")
                    email = st.text_input("Email", placeholder="contact@cabinet.fr")
                
                if st.form_submit_button("Créer le papier en-tête"):
                    # Créer l'en-tête
                    header_content = f"""{cabinet_nom}
{avocat_nom}
Avocat au {barreau}
{adresse}
Tél : {telephone}
Email : {email}"""
                    
                    # Créer le pied de page
                    footer_content = f"{cabinet_nom} - {adresse} - Tél : {telephone}"
                    
                    # Créer le template
                    st.session_state.letterhead_template = LetterheadTemplate(
                        name="Papier en-tête principal",
                        header_content=header_content,
                        footer_content=footer_content
                    )
                    
                    st.success("✅ Papier en-tête créé avec succès!")
                    st.rerun()
        return
    
    # Afficher le papier en-tête actuel
    with st.expander("📄 Papier en-tête actuel", expanded=False):
        template = st.session_state.letterhead_template
        st.text(template.header_content)
        st.caption(f"Pied de page : {template.footer_content}")
    
    # Formulaire de rédaction du courrier
    st.markdown("### 📝 Nouveau courrier")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Destinataire
        destinataire_nom = st.text_input(
            "Destinataire",
            placeholder="Nom du destinataire",
            key="courrier_destinataire_nom"
        )
        
        destinataire_qualite = st.text_input(
            "Qualité",
            placeholder="Ex: Procureur de la République, Directeur...",
            key="courrier_destinataire_qualite"
        )
        
        destinataire_adresse = st.text_area(
            "Adresse",
            placeholder="Adresse complète du destinataire",
            height=80,
            key="courrier_destinataire_adresse"
        )
    
    with col2:
        # Informations courrier
        objet = st.text_input(
            "Objet",
            placeholder="Objet du courrier",
            key="courrier_objet"
        )
        
        reference = st.text_input(
            "Référence",
            placeholder="Votre référence / Notre référence",
            key="courrier_reference"
        )
        
        date_courrier = st.date_input(
            "Date",
            value=datetime.now(),
            key="courrier_date"
        )
        
        pieces_jointes = st.text_area(
            "Pièces jointes",
            placeholder="Liste des pièces jointes (optionnel)",
            height=60,
            key="courrier_pj"
        )
    
    # Type de courrier
    st.markdown("### 📨 Type de courrier")
    
    type_courrier = st.selectbox(
        "Sélectionner le type",
        [
            "Courrier simple",
            "Demande d'information",
            "Demande d'audition",
            "Transmission de pièces",
            "Réponse à courrier",
            "Mise en demeure",
            "Notification",
            "Autre"
        ],
        key="courrier_type_select"
    )
    
    # Contenu principal
    st.markdown("### 📄 Contenu du courrier")
    
    # Option : rédaction manuelle ou assistée par IA
    mode_redaction = st.radio(
        "Mode de rédaction",
        ["Rédaction assistée par IA", "Rédaction manuelle"],
        key="courrier_mode"
    )
    
    if mode_redaction == "Rédaction assistée par IA":
        # Points clés pour l'IA
        points_courrier = st.text_area(
            "Points clés à développer",
            placeholder="""Ex:
- Demander un rendez-vous
- Rappeler la procédure en cours
- Solliciter la communication de pièces
- Proposer des dates de disponibilité""",
            height=150,
            key="courrier_points"
        )
        
        # Options de génération
        col1, col2 = st.columns(2)
        
        with col1:
            ton_courrier = st.select_slider(
                "Ton",
                options=["Très formel", "Formel", "Courtois", "Direct"],
                value="Formel",
                key="courrier_ton"
            )
        
        with col2:
            style_courrier = st.select_slider(
                "Style",
                options=["Concis", "Standard", "Détaillé"],
                value="Standard",
                key="courrier_style"
            )
        
        # Bouton de génération
        if st.button("🤖 Générer le courrier", type="primary", key="generer_courrier_ia"):
            if not destinataire_nom or not objet or not points_courrier:
                st.error("❌ Veuillez remplir tous les champs obligatoires")
                return
            
            # Construire le prompt
            prompt = f"""Tu es un avocat rédigeant un courrier professionnel.

Type de courrier : {type_courrier}
Destinataire : {destinataire_nom} ({destinataire_qualite})
Objet : {objet}
Référence : {reference or 'Aucune'}
Points à développer : {points_courrier}
Ton : {ton_courrier}
Style : {style_courrier}
{"Pièces jointes : " + pieces_jointes if pieces_jointes else ""}

Rédige un courrier professionnel complet avec :
- Une formule d'appel appropriée
- Une introduction claire
- Le développement des points demandés
- Une conclusion adaptée
- Une formule de politesse conforme aux usages

Ne pas inclure l'en-tête ni l'adresse (sera ajouté automatiquement)."""
            
            # Générer avec l'IA
            llm_manager = MultiLLMManager()
            
            with st.spinner("🔄 Génération en cours..."):
                if llm_manager.clients:
                    provider = list(llm_manager.clients.keys())[0]
                    
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    
                    response = loop.run_until_complete(
                        llm_manager.query_single_llm(
                            provider,
                            prompt,
                            "Tu es un avocat expert en rédaction de courriers professionnels. Tes courriers sont toujours clairs, courtois et efficaces."
                        )
                    )
                    
                    if response['success']:
                        st.session_state.courrier_content = response['response']
                        st.success("✅ Courrier généré avec succès!")
                    else:
                        st.error(f"❌ Erreur : {response['error']}")
                else:
                    st.error("❌ Aucune IA disponible")
    
    else:  # Rédaction manuelle
        contenu_manuel = st.text_area(
            "Contenu du courrier",
            placeholder="Rédigez votre courrier ici...",
            height=400,
            key="courrier_contenu_manuel"
        )
        
        if contenu_manuel:
            st.session_state.courrier_content = contenu_manuel
    
    # Affichage et édition du courrier
    if 'courrier_content' in st.session_state and st.session_state.courrier_content:
        st.markdown("### ✏️ Courrier généré")
        
        # Édition possible
        courrier_final = st.text_area(
            "Vous pouvez modifier le courrier",
            value=st.session_state.courrier_content,
            height=400,
            key="courrier_final_edit"
        )
        
        # Prévisualisation avec papier en-tête
        if st.button("👁️ Prévisualiser avec papier en-tête", key="preview_courrier"):
            st.markdown("### 📄 Prévisualisation")
            
            # Construire le courrier complet
            template = st.session_state.letterhead_template
            
            # Créer la structure complète
            courrier_complet = f"""{template.header_content}

{destinataire_nom}
{destinataire_qualite}
{destinataire_adresse}

{date_courrier.strftime('%d %B %Y')}

Objet : {objet}
{"Réf : " + reference if reference else ""}

{courrier_final}

{template.footer_content}"""
            
            # Afficher dans un conteneur stylé
            st.markdown('<div class="letterhead-preview">', unsafe_allow_html=True)
            st.text(courrier_complet)
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Options d'export
        st.markdown("### 💾 Export")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Export texte simple
            courrier_complet_export = f"""{template.header_content}

{destinataire_nom}
{destinataire_qualite}
{destinataire_adresse}

{date_courrier.strftime('%d %B %Y')}

Objet : {objet}
{"Réf : " + reference if reference else ""}

{courrier_final}

{template.footer_content}"""
            
            st.download_button(
                "💾 Télécharger (.txt)",
                courrier_complet_export,
                f"courrier_{clean_key(objet)}_{date_courrier.strftime('%Y%m%d')}.txt",
                "text/plain",
                key="download_courrier_txt"
            )
        
        with col2:
            if DOCX_AVAILABLE:
                # Créer le document Word avec papier en-tête
                docx_buffer = create_letterhead_from_template(
                    template,
                    f"""{destinataire_nom}
{destinataire_qualite}
{destinataire_adresse}

{date_courrier.strftime('%d %B %Y')}

Objet : {objet}
{"Réf : " + reference if reference else ""}

{courrier_final}"""
                )
                
                if docx_buffer:
                    st.download_button(
                        "💾 Télécharger (.docx)",
                        docx_buffer.getvalue(),
                        f"courrier_{clean_key(objet)}_{date_courrier.strftime('%Y%m%d')}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_courrier_docx"
                    )
        
        with col3:
            if st.button("📧 Préparer l'envoi par email", key="prepare_email"):
                st.info("📧 Copier le contenu ci-dessous dans votre client email")
                st.code(courrier_final)

def page_configuration():
    """Page de configuration"""
    
    st.header("⚙️ Configuration")
    
    tabs = st.tabs(["🔑 Clés API", "📄 Papier en-tête", "📊 État du système", "💾 Export/Import"])
    
    # Onglet Clés API
    with tabs[0]:
        st.markdown("### Configuration des clés API")
        
        st.info("""
        ℹ️ Les clés API doivent être configurées dans les variables d'environnement ou dans un fichier .env
        
        Variables nécessaires:
        - AZURE_OPENAI_KEY
        - AZURE_OPENAI_ENDPOINT
        - AZURE_SEARCH_KEY
        - AZURE_SEARCH_ENDPOINT
        - AZURE_STORAGE_CONNECTION_STRING
        - ANTHROPIC_API_KEY
        - OPENAI_API_KEY
        - GOOGLE_API_KEY
        - PERPLEXITY_API_KEY
        """)
        
        # Créer un fichier .env exemple
        env_example = """# Azure Search
AZURE_SEARCH_ENDPOINT=https://search-rag-juridique.search.windows.net
AZURE_SEARCH_KEY=Votre_Clé_Azure_Search

# Azure OpenAI pour les embeddings
AZURE_OPENAI_ENDPOINT=https://openai-juridique-rag2.openai.azure.com
AZURE_OPENAI_KEY=Votre_Clé_Azure_OpenAI
AZURE_OPENAI_DEPLOYMENT=text-embedding-ada-002

# Azure Blob Storage
AZURE_STORAGE_CONNECTION_STRING=DefaultEndpointsProtocol=https;AccountName=VOTRE_COMPTE;AccountKey=VOTRE_CLE;EndpointSuffix=core.windows.net

# Autres IA (optionnel)
ANTHROPIC_API_KEY=Votre_Clé_Anthropic
OPENAI_API_KEY=Votre_Clé_OpenAI
GOOGLE_API_KEY=Votre_Clé_Google
PERPLEXITY_API_KEY=Votre_Clé_Perplexity"""
        
        st.download_button(
            "📥 Télécharger un fichier .env exemple",
            env_example,
            ".env.example",
            "text/plain",
            key="download_env_example"
        )
        
        # Vérifier l'état des clés
        configs = LLMConfig.get_configs()
        
        st.markdown("#### 🤖 Providers IA")
        for provider, config in configs.items():
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.text(provider.value)
            
            with col2:
                if config.get('key') or config.get('api_key'):
                    st.success("✅")
                else:
                    st.error("❌")
        
        # Services Azure
        st.markdown("#### 🔷 Services Azure")
        
        services = {
            "Azure Blob Storage": bool(os.getenv('AZURE_STORAGE_CONNECTION_STRING')),
            "Azure Search": bool(os.getenv('AZURE_SEARCH_ENDPOINT')),
            "Azure OpenAI": bool(os.getenv('AZURE_OPENAI_ENDPOINT'))
        }
        
        for service, configured in services.items():
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.text(service)
            
            with col2:
                if configured:
                    st.success("✅")
                else:
                    st.error("❌")
    
    # Onglet Papier en-tête
    with tabs[1]:
        st.markdown("### 📄 Configuration du papier en-tête")
        
        # Papier en-tête actuel
        if 'letterhead_template' in st.session_state and st.session_state.letterhead_template:
            current_template = st.session_state.letterhead_template
            
            with st.expander("Papier en-tête actuel", expanded=True):
                st.text("En-tête :")
                st.code(current_template.header_content)
                st.text("Pied de page :")
                st.code(current_template.footer_content)
                
                if st.button("🗑️ Supprimer le papier en-tête actuel"):
                    st.session_state.letterhead_template = None
                    st.success("✅ Papier en-tête supprimé")
                    st.rerun()
        
        # Créer/Modifier papier en-tête
        st.markdown("#### ➕ Créer/Modifier le papier en-tête")
        
        with st.form("letterhead_form"):
            col1, col2 = st.columns(2)
            
            with col1:
                nom_template = st.text_input(
                    "Nom du template",
                    value=st.session_state.letterhead_template.name if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else "Papier en-tête principal"
                )
                
                header_content = st.text_area(
                    "En-tête",
                    value=st.session_state.letterhead_template.header_content if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else "",
                    height=150,
                    placeholder="""Cabinet d'avocats XYZ
Maître Jean DUPONT
Avocat au Barreau de Paris
123 rue de la Justice
75001 PARIS
Tél : 01 23 45 67 89
Email : contact@cabinet-xyz.fr"""
                )
                
                footer_content = st.text_area(
                    "Pied de page",
                    value=st.session_state.letterhead_template.footer_content if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else "",
                    height=80,
                    placeholder="Cabinet XYZ - 123 rue de la Justice, 75001 PARIS - Tél : 01 23 45 67 89"
                )
            
            with col2:
                # Paramètres de mise en forme
                st.markdown("**Mise en forme**")
                
                font_family = st.selectbox(
                    "Police",
                    ["Arial", "Times New Roman", "Calibri", "Garamond", "Helvetica"],
                    index=0 if not ('letterhead_template' in st.session_state and st.session_state.letterhead_template) else ["Arial", "Times New Roman", "Calibri", "Garamond", "Helvetica"].index(st.session_state.letterhead_template.font_family)
                )
                
                font_size = st.number_input(
                    "Taille de police",
                    min_value=8,
                    max_value=16,
                    value=11 if not ('letterhead_template' in st.session_state and st.session_state.letterhead_template) else st.session_state.letterhead_template.font_size
                )
                
                line_spacing = st.number_input(
                    "Interligne",
                    min_value=1.0,
                    max_value=2.0,
                    step=0.1,
                    value=1.5 if not ('letterhead_template' in st.session_state and st.session_state.letterhead_template) else st.session_state.letterhead_template.line_spacing
                )
                
                st.markdown("**Marges (cm)**")
                
                margin_top = st.number_input("Haut", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
                margin_bottom = st.number_input("Bas", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
                margin_left = st.number_input("Gauche", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
                margin_right = st.number_input("Droite", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
            
            # Upload logo (optionnel)
            logo_file = st.file_uploader(
                "Logo (optionnel)",
                type=['png', 'jpg', 'jpeg'],
                key="letterhead_logo"
            )
            
            if st.form_submit_button("💾 Sauvegarder le papier en-tête", type="primary"):
                # Créer le template
                new_template = LetterheadTemplate(
                    name=nom_template,
                    header_content=header_content,
                    footer_content=footer_content,
                    font_family=font_family,
                    font_size=font_size,
                    line_spacing=line_spacing,
                    margins={
                        'top': margin_top,
                        'bottom': margin_bottom,
                        'left': margin_left,
                        'right': margin_right
                    }
                )
                
                # Sauvegarder le logo si uploadé
                if logo_file:
                    st.session_state.letterhead_image = logo_file.read()
                    new_template.logo_path = "logo_uploaded"
                
                st.session_state.letterhead_template = new_template
                st.success("✅ Papier en-tête sauvegardé avec succès!")
                st.rerun()
        
        # Import de papier en-tête depuis Word
        st.markdown("#### 📤 Importer depuis un document Word")
        
        uploaded_letterhead = st.file_uploader(
            "Charger un document Word avec papier en-tête",
            type=['docx'],
            key="upload_letterhead_word"
        )
        
        if uploaded_letterhead and st.button("📥 Extraire le papier en-tête"):
            if DOCX_AVAILABLE:
                try:
                    doc = DocxDocument(uploaded_letterhead)
                    
                    # Extraire l'en-tête
                    header_text = ""
                    if doc.sections:
                        header = doc.sections[0].header
                        for paragraph in header.paragraphs:
                            if paragraph.text.strip():
                                header_text += paragraph.text + "\n"
                    
                    # Extraire le pied de page
                    footer_text = ""
                    if doc.sections:
                        footer = doc.sections[0].footer
                        for paragraph in footer.paragraphs:
                            if paragraph.text.strip():
                                footer_text += paragraph.text + "\n"
                    
                    if header_text or footer_text:
                        st.success("✅ Papier en-tête extrait avec succès!")
                        st.text("En-tête extrait :")
                        st.code(header_text)
                        st.text("Pied de page extrait :")
                        st.code(footer_text)
                        
                        if st.button("Utiliser ce papier en-tête"):
                            st.session_state.letterhead_template = LetterheadTemplate(
                                name="Papier en-tête importé",
                                header_content=header_text.strip(),
                                footer_content=footer_text.strip()
                            )
                            st.success("✅ Papier en-tête importé!")
                            st.rerun()
                    else:
                        st.warning("⚠️ Aucun en-tête ou pied de page trouvé dans le document")
                        
                except Exception as e:
                    st.error(f"❌ Erreur lors de l'extraction : {str(e)}")
    
    # Onglet État du système
    with tabs[2]:
        st.markdown("### État du système")
        
        # Métriques
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Documents Azure", len(st.session_state.azure_documents))
        
        with col2:
            st.metric("Pièces sélectionnées", len(st.session_state.pieces_selectionnees))
        
        with col3:
            llm_manager = MultiLLMManager()
            st.metric("IA disponibles", len(llm_manager.clients))
        
        with col4:
            st.metric("Styles appris", len(st.session_state.get('learned_styles', {})))
        
        # État des connexions
        st.markdown("### 🔗 État des connexions")
        
        # Azure Blob
        if st.session_state.azure_blob_manager and st.session_state.azure_blob_manager.is_connected():
            st.success("✅ Azure Blob Storage : Connecté")
            
            # Test de connexion
            if st.button("🧪 Tester la connexion Blob", key="test_blob"):
                try:
                    containers = st.session_state.azure_blob_manager.list_containers()
                    st.success(f"✅ {len(containers)} containers trouvés : {', '.join(containers)}")
                except Exception as e:
                    st.error(f"❌ Erreur : {str(e)}")
        else:
            st.error("❌ Azure Blob Storage : Non connecté")
        
        # Azure Search
        if st.session_state.azure_search_manager and st.session_state.azure_search_manager.search_client:
            st.success("✅ Azure Search : Connecté")
            
            # Test de connexion
            if st.button("🧪 Tester la connexion Search", key="test_search"):
                try:
                    # Tester avec une recherche simple
                    results = st.session_state.azure_search_manager.search_hybrid("test", top=1)
                    st.success("✅ Connexion fonctionnelle")
                except Exception as e:
                    st.error(f"❌ Erreur : {str(e)}")
        else:
            st.warning("⚠️ Azure Search : Non configuré")
        
        # Papier en-tête
        if 'letterhead_template' in st.session_state and st.session_state.letterhead_template:
            st.success("✅ Papier en-tête : Configuré")
        else:
            st.warning("⚠️ Papier en-tête : Non configuré")
    
    # Onglet Export/Import
    with tabs[3]:
        st.markdown("### 💾 Export/Import de configuration")
        
        # Export
        st.markdown("#### 📥 Export")
        
        export_data = {
            'pieces_selectionnees': {
                k: {
                    'document_id': v.document_id,
                    'titre': v.titre,
                    'categorie': v.categorie,
                    'notes': v.notes,
                    'pertinence': v.pertinence
                }
                for k, v in st.session_state.pieces_selectionnees.items()
            },
            'learned_styles': st.session_state.get('learned_styles', {}),
            'letterhead_template': {
                'name': st.session_state.letterhead_template.name,
                'header_content': st.session_state.letterhead_template.header_content,
                'footer_content': st.session_state.letterhead_template.footer_content,
                'font_family': st.session_state.letterhead_template.font_family,
                'font_size': st.session_state.letterhead_template.font_size,
                'line_spacing': st.session_state.letterhead_template.line_spacing,
                'margins': st.session_state.letterhead_template.margins
            } if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else None,
            'timestamp': datetime.now().isoformat()
        }
        
        export_json = json.dumps(export_data, indent=2, ensure_ascii=False)
        
        st.download_button(
            "💾 Exporter la configuration complète",
            export_json,
            f"config_penal_affaires_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "application/json",
            key="export_config"
        )
        
        # Import
        st.markdown("#### 📤 Import")
        
        uploaded_file = st.file_uploader(
            "Charger une configuration",
            type=['json'],
            key="import_config_file"
        )
        
        if uploaded_file:
            try:
                config_data = json.load(uploaded_file)
                
                # Prévisualisation
                with st.expander("Voir le contenu"):
                    st.json(config_data)
                
                if st.button("⬆️ Importer", key="import_config_button"):
                    # Importer les pièces sélectionnées
                    if 'pieces_selectionnees' in config_data:
                        for piece_id, piece_data in config_data['pieces_selectionnees'].items():
                            piece = PieceSelectionnee(
                                document_id=piece_data['document_id'],
                                titre=piece_data['titre'],
                                categorie=piece_data['categorie'],
                                notes=piece_data.get('notes', ''),
                                pertinence=piece_data.get('pertinence', 5)
                            )
                            st.session_state.pieces_selectionnees[piece_id] = piece
                    
                    # Importer les styles appris
                    if 'learned_styles' in config_data:
                        st.session_state.learned_styles = config_data['learned_styles']
                    
                    # Importer le papier en-tête
                    if 'letterhead_template' in config_data and config_data['letterhead_template']:
                        lt = config_data['letterhead_template']
                        st.session_state.letterhead_template = LetterheadTemplate(
                            name=lt['name'],
                            header_content=lt['header_content'],
                            footer_content=lt['footer_content'],
                            font_family=lt.get('font_family', 'Arial'),
                            font_size=lt.get('font_size', 11),
                            line_spacing=lt.get('line_spacing', 1.5),
                            margins=lt.get('margins', {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5})
                        )
                    
                    st.success("✅ Configuration importée avec succès")
                    st.rerun()
                    
            except Exception as e:
                st.error(f"❌ Erreur lors de l'import: {str(e)}")

# Fonctions auxiliaires pour la fusion des styles
def merge_structures(structures: List[Dict]) -> Dict[str, Any]:
    """Fusionne plusieurs structures de documents"""
    if not structures:
        return {}
    
    merged = {
        'sections_communes': [],
        'longueur_moyenne': 0
    }
    
    # Trouver les sections communes
    all_sections = []
    for struct in structures:
        all_sections.extend([s['titre'] for s in struct.get('sections', [])])
    
    # Compter les occurrences
    section_counts = Counter(all_sections)
    
    # Garder les sections présentes dans au moins 50% des documents
    threshold = len(structures) / 2
    merged['sections_communes'] = [
        section for section, count in section_counts.items()
        if count >= threshold
    ]
    
    # Ajouter les informations Word si présentes
    word_styles = []
    for struct in structures:
        if 'word_styles' in struct:
            word_styles.extend(struct['word_styles'])
    
    if word_styles:
        merged['word_styles'] = list(set(word_styles))
    
    return merged

def merge_formules(formules_list: List[List[str]]) -> List[str]:
    """Fusionne les formules types"""
    all_formules = []
    for formules in formules_list:
        all_formules.extend(formules)
    
    # Compter et garder les plus fréquentes
    formule_counts = Counter(all_formules)
    
    return [formule for formule, count in formule_counts.most_common(20)]

def merge_formatting(formats: List[Dict]) -> Dict[str, Any]:
    """Fusionne les paramètres de mise en forme"""
    if not formats:
        return {}
    
    merged = {}
    
    # Moyennes et valeurs communes
    for key in formats[0].keys():
        values = [f.get(key) for f in formats if key in f]
        
        if all(isinstance(v, bool) for v in values):
            # Pour les booléens, prendre la majorité
            merged[key] = sum(values) > len(values) / 2
        elif all(isinstance(v, (int, float)) for v in values):
            # Pour les nombres, prendre la moyenne
            merged[key] = sum(values) / len(values)
        else:
            # Pour le reste, prendre la valeur la plus fréquente
            merged[key] = Counter(values).most_common(1)[0][0] if values else None
    
    return merged

def merge_vocabulary(vocab_list: List[Dict[str, int]]) -> Dict[str, int]:
    """Fusionne les vocabulaires"""
    merged = defaultdict(int)
    
    for vocab in vocab_list:
        for word, count in vocab.items():
            merged[word] += count
    
    # Garder les 100 mots les plus fréquents
    return dict(sorted(merged.items(), key=lambda x: x[1], reverse=True)[:100])

# ================== LANCEMENT DE L'APPLICATION ==================

if __name__ == "__main__":
    main()