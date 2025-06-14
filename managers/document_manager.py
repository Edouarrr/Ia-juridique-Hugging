"""
Gestionnaire de documents juridiques - Import, export, création et analyse
"""
import asyncio
import base64
import io
import json
import logging
import os
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import openpyxl
import pandas as pd
import PyPDF2
import streamlit as st
# Import des bibliothèques de traitement de documents
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image

from managers.azure_search_manager import AzureSearchManager
from managers.export_manager import ExportManager
from managers.jurisprudence_verifier import JurisprudenceVerifier
# Import des gestionnaires
from managers.llm_manager import LLMManager
from managers.template_manager import TemplateManager
# CORRECTION : Import depuis modules au lieu de models
from modules.dataclasses import (AnalyseJuridique, CasJuridique,
                                 DocumentJuridique)

# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentManager:
    """Gestionnaire centralisé pour tous les documents juridiques"""
    
    SUPPORTED_FORMATS = {
        'import': ['.pdf', '.docx', '.txt', '.json', '.xlsx', '.csv'],
        'export': ['pdf', 'docx', 'xlsx', 'json', 'txt']
    }
    
    def __init__(self):
        # Gestionnaires externes
        self.llm_manager = LLMManager()
        self.template_manager = TemplateManager()
        self.export_manager = ExportManager()
        self.jurisprudence_verifier = JurisprudenceVerifier()
        self.azure_search_manager = AzureSearchManager()
        
        # Stockage local
        self.imported_documents = []
        self.processed_texts = []
        
        # Mapping des types de documents
        self.document_types = {
            "plainte": "Plainte pénale",
            "constitution": "Constitution de partie civile",
            "conclusions": "Conclusions",
            "assignation": "Assignation",
            "requete": "Requête",
            "memoire": "Mémoire",
            "courrier": "Courrier juridique"
        }
    
    # ========== MÉTHODES D'IMPORT ==========
    
    def import_document(self, file) -> Tuple[bool, str, str]:
        """
        Importe un document et extrait son contenu
        Retourne: (succès, contenu, message)
        """
        try:
            file_extension = Path(file.name).suffix.lower()
            
            if file_extension not in self.SUPPORTED_FORMATS['import']:
                return False, "", f"Format non supporté: {file_extension}"
            
            # Traitement selon le type
            if file_extension == '.pdf':
                content = self._extract_pdf_content(file)
            elif file_extension == '.docx':
                content = self._extract_docx_content(file)
            elif file_extension == '.txt':
                content = str(file.read(), 'utf-8')
            elif file_extension == '.json':
                content = self._extract_json_content(file)
            elif file_extension in ['.xlsx', '.csv']:
                content = self._extract_table_content(file, file_extension)
            else:
                return False, "", "Format non reconnu"
            
            # Stocker le document importé
            self.imported_documents.append({
                'filename': file.name,
                'content': content,
                'import_date': datetime.now(),
                'size': file.size
            })
            
            return True, content, f"Document '{file.name}' importé avec succès"
            
        except Exception as e:
            logger.error(f"Erreur import document: {e}")
            return False, "", f"Erreur lors de l'import: {str(e)}"
    
    def batch_import(self, files: List) -> Dict[str, Any]:
        """Import en lot de plusieurs documents"""
        results = {
            'success': [],
            'failed': [],
            'total_content': ""
        }
        
        for file in files:
            success, content, message = self.import_document(file)
            
            if success:
                results['success'].append({
                    'filename': file.name,
                    'message': message,
                    'content_length': len(content)
                })
                results['total_content'] += f"\n\n--- {file.name} ---\n{content}"
            else:
                results['failed'].append({
                    'filename': file.name,
                    'error': message
                })
        
        return results
    
    # ========== MÉTHODES DE CRÉATION DE DOCUMENTS JURIDIQUES ==========
    
    def creer_document(self, 
                      type_doc: str,
                      contexte: Dict[str, Any],
                      style: str = "formel",
                      options: Optional[Dict[str, Any]] = None) -> DocumentJuridique:
        """
        Crée un nouveau document juridique
        
        Args:
            type_doc: Type de document à créer
            contexte: Informations contextuelles (parties, faits, etc.)
            style: Style rédactionnel
            options: Options supplémentaires
            
        Returns:
            DocumentJuridique créé
        """
        try:
            # Enrichir le contexte avec des informations supplémentaires
            contexte_enrichi = self._enrichir_contexte(contexte, type_doc)
            
            # Obtenir le template approprié
            template = self.template_manager.get_template(type_doc)
            if not template:
                st.error(f"Template non trouvé pour le type: {type_doc}")
                return None
            
            # Générer le contenu via LLM
            prompt = self._construire_prompt(type_doc, contexte_enrichi, template, style)
            contenu = self.llm_manager.generate(prompt, temperature=0.3)
            
            # Vérifier et enrichir avec la jurisprudence
            if options and options.get("verifier_jurisprudence", True):
                jurisprudences = self.jurisprudence_verifier.verifier_arguments(contenu)
                if jurisprudences:
                    contenu = self._integrer_jurisprudence(contenu, jurisprudences)
            
            # Créer le document
            doc = DocumentJuridique(
                id=str(uuid.uuid4()),
                titre=f"{self.document_types.get(type_doc, 'Document')} - {contexte.get('affaire', 'Sans titre')}",
                type_document=type_doc,
                numero_reference=self._generer_reference(type_doc),
                date_document=datetime.now(),
                auteur=contexte.get('avocat', 'Me. X'),
                destinataires=self._extraire_destinataires(type_doc, contexte),
                contenu=contenu,
                juridiction=contexte.get('juridiction', ''),
                procedure_liee=contexte.get('numero_procedure', ''),
                statut_legal="projet"
            )
            
            # Sauvegarder dans la session
            if 'documents_generes' not in st.session_state:
                st.session_state.documents_generes = {}
            st.session_state.documents_generes[doc.id] = doc
            
            # Indexer dans Azure Search si disponible
            if self.azure_search_manager.is_connected():
                self._indexer_document(doc)
            
            return doc
            
        except Exception as e:
            st.error(f"Erreur lors de la création du document: {str(e)}")
            return None
    
    def analyser_document(self, document: Any, type_analyse: str = "complete") -> AnalyseJuridique:
        """
        Analyse un document juridique existant
        
        Args:
            document: Document à analyser (texte ou DocumentJuridique)
            type_analyse: Type d'analyse à effectuer
            
        Returns:
            AnalyseJuridique avec les résultats
        """
        try:
            # Extraire le contenu
            if isinstance(document, str):
                contenu = document
                titre = "Document sans titre"
            elif hasattr(document, 'contenu'):
                contenu = document.contenu
                titre = document.titre
            else:
                contenu = str(document)
                titre = "Document converti"
            
            # Construire le prompt d'analyse
            prompt = f"""
            Analysez ce document juridique de manière approfondie:
            
            {contenu[:3000]}...
            
            Fournissez:
            1. Résumé des points clés
            2. Forces et faiblesses juridiques
            3. Risques identifiés
            4. Recommandations
            5. Références juridiques pertinentes
            """
            
            # Générer l'analyse
            analyse_brute = self.llm_manager.generate(prompt, temperature=0.2)
            
            # Parser l'analyse
            points_cles = self._extraire_section(analyse_brute, "points clés")
            risques = self._extraire_section(analyse_brute, "risques")
            recommandations = self._extraire_section(analyse_brute, "recommandations")
            
            # Vérifier la jurisprudence citée
            jurisprudences = self.jurisprudence_verifier.extraire_citations(analyse_brute)
            
            # Créer l'analyse
            analyse = AnalyseJuridique(
                id=str(uuid.uuid4()),
                titre=f"Analyse de {titre}",
                date_analyse=datetime.now(),
                type_analyse=type_analyse,
                objet_analyse=titre,
                contenu_analyse=analyse_brute,
                points_cles=points_cles,
                risques_identifies=risques,
                recommandations=recommandations,
                references_juridiques=[j['reference'] for j in jurisprudences],
                niveau_confiance=0.85,
                auteur="Assistant IA"
            )
            
            return analyse
            
        except Exception as e:
            st.error(f"Erreur lors de l'analyse: {str(e)}")
            return None
    
    def generer_cas_juridique(self, faits: str, contexte: Optional[Dict] = None) -> CasJuridique:
        """
        Génère un cas juridique structuré à partir de faits
        
        Args:
            faits: Description des faits
            contexte: Contexte additionnel
            
        Returns:
            CasJuridique structuré
        """
        try:
            prompt = f"""
            Analysez ces faits et créez un cas juridique structuré:
            
            {faits}
            
            Identifiez:
            1. Les questions juridiques principales
            2. Les infractions potentielles
            3. Les articles de loi applicables
            4. La stratégie juridique recommandée
            5. L'évaluation des risques
            """
            
            analyse = self.llm_manager.generate(prompt, temperature=0.2)
            
            # Parser les éléments
            questions = self._extraire_liste(analyse, "questions juridiques")
            infractions = self._extraire_liste(analyse, "infractions")
            articles = self._extraire_liste(analyse, "articles")
            
            cas = CasJuridique(
                id=str(uuid.uuid4()),
                titre=self._generer_titre_cas(faits),
                description=faits,
                type_cas="penal" if "pénal" in faits.lower() else "civil",
                date_fait=datetime.now(),
                faits=[faits],
                questions_juridiques=questions,
                infractions_potentielles=infractions,
                articles_applicables=articles,
                strategie_proposee=self._extraire_section(analyse, "stratégie")[0] if self._extraire_section(analyse, "stratégie") else "",
                risque_evaluation=self._extraire_section(analyse, "risques")[0] if self._extraire_section(analyse, "risques") else "",
                statut="en_analyse"
            )
            
            return cas
            
        except Exception as e:
            st.error(f"Erreur lors de la génération du cas: {str(e)}")
            return None
    
    # ========== MÉTHODES D'EXTRACTION DE CONTENU ==========
    
    def _extract_pdf_content(self, file) -> str:
        """Extrait le texte d'un PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.extract_text()
                
            return text.strip()
        except Exception as e:
            logger.error(f"Erreur extraction PDF: {e}")
            raise
    
    def _extract_docx_content(self, file) -> str:
        """Extrait le texte d'un DOCX"""
        try:
            doc = Document(file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extraire aussi les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
                    
            return text.strip()
        except Exception as e:
            logger.error(f"Erreur extraction DOCX: {e}")
            raise
    
    def _extract_json_content(self, file) -> str:
        """Extrait et formate le contenu JSON"""
        try:
            data = json.load(file)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Erreur extraction JSON: {e}")
            raise
    
    def _extract_table_content(self, file, extension: str) -> str:
        """Extrait le contenu d'un fichier tabulaire"""
        try:
            if extension == '.csv':
                df = pd.read_csv(file)
            else:  # xlsx
                df = pd.read_excel(file)
                
            # Convertir en texte structuré
            text = f"Tableau avec {len(df)} lignes et {len(df.columns)} colonnes\n\n"
            text += "Colonnes: " + ", ".join(df.columns) + "\n\n"
            
            # Résumé des données
            text += "Aperçu des données:\n"
            text += df.head(10).to_string()
            
            return text
        except Exception as e:
            logger.error(f"Erreur extraction tableau: {e}")
            raise
    
    # ========== MÉTHODES D'EXPORT ==========
    
    def export_analysis(
        self,
        analysis: AnalyseJuridique,
        format: str,
        include_metadata: bool = True
    ) -> Tuple[bytes, str]:
        """
        Exporte une analyse dans le format demandé
        Retourne: (contenu_bytes, nom_fichier)
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"analyse_juridique_{timestamp}"
        
        if format == 'pdf':
            # Note: Pour un vrai PDF, utiliser reportlab ou weasyprint
            # Ici on fait une version simplifiée
            content = self._create_pdf_content(analysis, include_metadata)
            filename = f"{base_filename}.pdf"
        elif format == 'docx':
            content = self._create_docx_content(analysis, include_metadata)
            filename = f"{base_filename}.docx"
        elif format == 'xlsx':
            content = self._create_xlsx_content(analysis, include_metadata)
            filename = f"{base_filename}.xlsx"
        elif format == 'json':
            content = self._create_json_content(analysis, include_metadata)
            filename = f"{base_filename}.json"
        else:  # txt
            content = self._create_txt_content(analysis, include_metadata)
            filename = f"{base_filename}.txt"
            
        return content, filename
    
    def _create_docx_content(self, analysis: AnalyseJuridique, include_metadata: bool) -> bytes:
        """Crée un document DOCX"""
        doc = Document()
        
        # Titre
        title = doc.add_heading('Analyse Juridique', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Métadonnées
        if include_metadata:
            doc.add_paragraph(f"Date: {analysis.date_analyse.strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph(f"Type d'analyse: {analysis.type_analyse}")
            doc.add_paragraph(f"ID: {analysis.id}")
            doc.add_paragraph()
        
        # Titre de l'analyse
        doc.add_heading(analysis.titre, 1)
        
        # Contenu de l'analyse
        doc.add_heading('Analyse détaillée', 2)
        doc.add_paragraph(analysis.contenu_analyse)
        
        # Points clés
        if analysis.points_cles:
            doc.add_heading('Points clés', 2)
            for point in analysis.points_cles:
                doc.add_paragraph(f"• {point}", style='List Bullet')
        
        # Risques identifiés
        if analysis.risques_identifies:
            doc.add_heading('Risques identifiés', 2)
            for risque in analysis.risques_identifies:
                doc.add_paragraph(f"• {risque}", style='List Bullet')
        
        # Recommandations
        if analysis.recommandations:
            doc.add_heading('Recommandations', 2)
            for i, reco in enumerate(analysis.recommandations, 1):
                doc.add_paragraph(f"{i}. {reco}")
        
        # Références juridiques
        if analysis.references_juridiques:
            doc.add_heading('Références juridiques', 2)
            for ref in analysis.references_juridiques:
                doc.add_paragraph(f"• {ref}", style='List Bullet')
        
        # Niveau de confiance
        doc.add_heading('Évaluation', 2)
        doc.add_paragraph(f"Niveau de confiance: {analysis.niveau_confiance:.0%}")
        doc.add_paragraph(f"Statut: {analysis.statut}")
        
        # Sauvegarder en mémoire
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io.getvalue()
    
    def _create_xlsx_content(self, analysis: AnalyseJuridique, include_metadata: bool) -> bytes:
        """Crée un fichier Excel"""
        wb = openpyxl.Workbook()
        
        # Feuille de synthèse
        ws_synthese = wb.active
        ws_synthese.title = "Synthèse"
        
        # En-têtes avec style
        headers = ['Élément', 'Valeur']
        for col, header in enumerate(headers, 1):
            cell = ws_synthese.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")
        
        # Données de synthèse
        row = 2
        if include_metadata:
            ws_synthese.cell(row=row, column=1, value="Date d'analyse")
            ws_synthese.cell(row=row, column=2, value=analysis.date_analyse.strftime('%d/%m/%Y %H:%M'))
            row += 1
            
            ws_synthese.cell(row=row, column=1, value="Type d'analyse")
            ws_synthese.cell(row=row, column=2, value=analysis.type_analyse)
            row += 1
        
        ws_synthese.cell(row=row, column=1, value="Titre")
        ws_synthese.cell(row=row, column=2, value=analysis.titre)
        row += 1
        
        ws_synthese.cell(row=row, column=1, value="Niveau de confiance")
        ws_synthese.cell(row=row, column=2, value=f"{analysis.niveau_confiance:.0%}")
        row += 1
        
        ws_synthese.cell(row=row, column=1, value="Nombre de points clés")
        ws_synthese.cell(row=row, column=2, value=len(analysis.points_cles))
        row += 1
        
        ws_synthese.cell(row=row, column=1, value="Nombre de risques")
        ws_synthese.cell(row=row, column=2, value=len(analysis.risques_identifies))
        
        # Feuille détaillée
        ws_detail = wb.create_sheet("Détails")
        
        # Points clés
        ws_detail.cell(row=1, column=1, value="Points clés").font = Font(bold=True)
        for idx, point in enumerate(analysis.points_cles, 2):
            ws_detail.cell(row=idx, column=1, value=point)
        
        # Risques
        start_row = len(analysis.points_cles) + 3
        ws_detail.cell(row=start_row, column=1, value="Risques identifiés").font = Font(bold=True)
        for idx, risque in enumerate(analysis.risques_identifies):
            ws_detail.cell(row=start_row + idx + 1, column=1, value=risque)
        
        # Recommandations
        start_row = start_row + len(analysis.risques_identifies) + 2
        ws_detail.cell(row=start_row, column=1, value="Recommandations").font = Font(bold=True)
        for idx, reco in enumerate(analysis.recommandations):
            ws_detail.cell(row=start_row + idx + 1, column=1, value=reco)
        
        # Ajuster les largeurs de colonnes
        for ws in [ws_synthese, ws_detail]:
            for column_cells in ws.columns:
                length = max(len(str(cell.value)) for cell in column_cells if cell.value)
                ws.column_dimensions[column_cells[0].column_letter].width = min(length + 2, 50)
        
        # Sauvegarder
        excel_io = io.BytesIO()
        wb.save(excel_io)
        excel_io.seek(0)
        
        return excel_io.getvalue()
    
    def _create_json_content(self, analysis: AnalyseJuridique, include_metadata: bool) -> bytes:
        """Crée un export JSON"""
        data = {
            "titre": analysis.titre,
            "type_analyse": analysis.type_analyse,
            "objet_analyse": analysis.objet_analyse,
            "contenu_analyse": analysis.contenu_analyse,
            "points_cles": analysis.points_cles,
            "risques_identifies": analysis.risques_identifies,
            "recommandations": analysis.recommandations,
            "references_juridiques": analysis.references_juridiques,
            "niveau_confiance": analysis.niveau_confiance,
            "auteur": analysis.auteur,
            "statut": analysis.statut
        }
        
        if include_metadata:
            data["metadata"] = {
                "id": analysis.id,
                "date_analyse": analysis.date_analyse.isoformat(),
                "validee_par": analysis.validee_par,
                "metadata_additionnelle": analysis.metadata
            }
        
        # Ajouter les opportunités si présentes
        if analysis.opportunites:
            data["opportunites"] = analysis.opportunites
        
        return json.dumps(data, indent=2, ensure_ascii=False).encode('utf-8')
    
    def _create_txt_content(self, analysis: AnalyseJuridique, include_metadata: bool) -> bytes:
        """Crée un export texte"""
        lines = []
        
        lines.append("ANALYSE JURIDIQUE")
        lines.append("=" * 50)
        lines.append("")
        
        if include_metadata:
            lines.append(f"Date: {analysis.date_analyse.strftime('%d/%m/%Y %H:%M')}")
            lines.append(f"Type: {analysis.type_analyse}")
            lines.append(f"ID: {analysis.id}")
            lines.append("")
        
        lines.append(f"TITRE: {analysis.titre}")
        lines.append("-" * 30)
        lines.append("")
        
        lines.append("ANALYSE DÉTAILLÉE")
        lines.append("-" * 30)
        lines.append(analysis.contenu_analyse)
        lines.append("")
        
        if analysis.points_cles:
            lines.append("POINTS CLÉS")
            lines.append("-" * 30)
            for point in analysis.points_cles:
                lines.append(f"• {point}")
            lines.append("")
        
        if analysis.risques_identifies:
            lines.append("RISQUES IDENTIFIÉS")
            lines.append("-" * 30)
            for risque in analysis.risques_identifies:
                lines.append(f"• {risque}")
            lines.append("")
        
        if analysis.recommandations:
            lines.append("RECOMMANDATIONS")
            lines.append("-" * 30)
            for i, reco in enumerate(analysis.recommandations, 1):
                lines.append(f"{i}. {reco}")
            lines.append("")
        
        if analysis.references_juridiques:
            lines.append("RÉFÉRENCES JURIDIQUES")
            lines.append("-" * 30)
            for ref in analysis.references_juridiques:
                lines.append(f"• {ref}")
            lines.append("")
        
        lines.append("ÉVALUATION")
        lines.append("-" * 30)
        lines.append(f"Niveau de confiance: {analysis.niveau_confiance:.0%}")
        lines.append(f"Statut: {analysis.statut}")
        lines.append(f"Auteur: {analysis.auteur}")
        
        return '\n'.join(lines).encode('utf-8')
    
    def _create_pdf_content(self, analysis: AnalyseJuridique, include_metadata: bool) -> bytes:
        """
        Crée un PDF (version simplifiée - pour un vrai PDF, utiliser reportlab)
        Pour l'instant, on retourne le contenu texte qui peut être converti en PDF
        """
        # Dans une vraie implémentation, utiliser reportlab ou weasyprint
        # Pour l'instant, on retourne le contenu texte
        return self._create_txt_content(analysis, include_metadata)
    
    # ========== MÉTHODES UTILITAIRES ==========
    
    def _enrichir_contexte(self, contexte: Dict, type_doc: str) -> Dict:
        """Enrichit le contexte avec des informations supplémentaires"""
        contexte_enrichi = contexte.copy()
        
        # Ajouter la date du jour
        contexte_enrichi['date_jour'] = datetime.now().strftime("%d/%m/%Y")
        
        # Ajouter des formules de politesse selon le type
        if type_doc in ["plainte", "requete"]:
            contexte_enrichi['formule_appel'] = "Monsieur le Procureur de la République"
        elif type_doc == "assignation":
            contexte_enrichi['formule_appel'] = "Le Tribunal"
        
        return contexte_enrichi
    
    def _construire_prompt(self, type_doc: str, contexte: Dict, template: str, style: str) -> str:
        """Construit le prompt pour la génération LLM"""
        prompt = f"""
        Rédigez un(e) {self.document_types.get(type_doc, 'document juridique')} 
        en respectant le style {style} et les informations suivantes:
        
        Contexte:
        {json.dumps(contexte, indent=2, ensure_ascii=False)}
        
        Template de base:
        {template}
        
        Instructions:
        - Respectez le formalisme juridique français
        - Utilisez un ton {style}
        - Structurez clairement le document
        - Citez les textes de loi pertinents
        - Assurez la cohérence juridique
        """
        
        return prompt
    
    def _generer_reference(self, type_doc: str) -> str:
        """Génère une référence unique pour le document"""
        prefix = type_doc[:3].upper()
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        return f"{prefix}-{timestamp}"
    
    def _extraire_destinataires(self, type_doc: str, contexte: Dict) -> List[str]:
        """Extrait les destinataires selon le type de document"""
        destinataires = []
        
        if type_doc == "plainte":
            destinataires.append("Procureur de la République")
        elif type_doc == "assignation":
            destinataires.extend(contexte.get('parties_adverses', []))
        elif type_doc == "conclusions":
            destinataires.append("Tribunal")
            destinataires.extend(contexte.get('autres_parties', []))
        
        return destinataires
    
    def _integrer_jurisprudence(self, contenu: str, jurisprudences: List[Dict]) -> str:
        """Intègre les références de jurisprudence dans le contenu"""
        # Ajouter une section jurisprudence si elle n'existe pas
        if "JURISPRUDENCE" not in contenu:
            section_juris = "\n\nIII. JURISPRUDENCE APPLICABLE\n\n"
            for juris in jurisprudences:
                section_juris += f"- {juris['reference']} : {juris['principe']}\n"
            
            # Insérer avant la conclusion
            if "CONCLUSION" in contenu:
                contenu = contenu.replace("CONCLUSION", section_juris + "\nCONCLUSION")
            else:
                contenu += section_juris
        
        return contenu
    
    def _indexer_document(self, doc: DocumentJuridique):
        """Indexe le document dans Azure Search"""
        try:
            doc_index = {
                "id": doc.id,
                "titre": doc.titre,
                "type": doc.type_document,
                "contenu": doc.contenu,
                "date": doc.date_document.isoformat(),
                "auteur": doc.auteur,
                "juridiction": doc.juridiction,
                "tags": [doc.type_document, doc.statut_legal]
            }
            
            self.azure_search_manager.index_documents([doc_index])
            
        except Exception as e:
            print(f"Erreur indexation: {e}")
    
    def _extraire_section(self, texte: str, section: str) -> List[str]:
        """Extrait une section spécifique du texte"""
        lignes = []
        in_section = False
        
        for ligne in texte.split('\n'):
            if section.lower() in ligne.lower():
                in_section = True
                continue
            elif in_section and ligne.strip() and not ligne.startswith(' '):
                # Nouvelle section
                break
            elif in_section and ligne.strip():
                lignes.append(ligne.strip('- •').strip())
        
        return lignes
    
    def _extraire_liste(self, texte: str, marqueur: str) -> List[str]:
        """Extrait une liste à partir d'un marqueur"""
        items = []
        lignes = texte.split('\n')
        
        for i, ligne in enumerate(lignes):
            if marqueur.lower() in ligne.lower():
                # Chercher les éléments de la liste après le marqueur
                for j in range(i + 1, len(lignes)):
                    if lignes[j].strip().startswith(('-', '•', '*', '1', '2', '3')):
                        items.append(lignes[j].strip('- •*123.').strip())
                    elif lignes[j].strip() and not lignes[j].startswith(' '):
                        break
                break
        
        return items
    
    def _generer_titre_cas(self, faits: str) -> str:
        """Génère un titre pour le cas juridique"""
        # Extraire les mots clés importants
        mots_cles = []
        
        if "abus de biens sociaux" in faits.lower():
            mots_cles.append("Abus de biens sociaux")
        if "corruption" in faits.lower():
            mots_cles.append("Corruption")
        if "escroquerie" in faits.lower():
            mots_cles.append("Escroquerie")
        
        if mots_cles:
            return f"Cas de {' et '.join(mots_cles)}"
        else:
            # Prendre les 50 premiers caractères
            return f"Cas juridique: {faits[:50]}..."
    
    def create_download_link(self, content: bytes, filename: str, text: str) -> str:
        """Crée un lien de téléchargement pour Streamlit"""
        b64 = base64.b64encode(content).decode()
        return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    
    def get_import_stats(self) -> Dict[str, Any]:
        """Retourne des statistiques sur les imports"""
        if not self.imported_documents:
            return {'count': 0}
            
        return {
            'count': len(self.imported_documents),
            'total_size': sum(doc['size'] for doc in self.imported_documents),
            'types': pd.Series([Path(doc['filename']).suffix for doc in self.imported_documents]).value_counts().to_dict(),
            'latest': self.imported_documents[-1]['filename'] if self.imported_documents else None
        }


# ========== FONCTIONS D'INTERFACE STREAMLIT ==========

def display_import_interface():
    """Interface d'import de documents"""
    st.markdown("### 📤 Import de documents")
    
    # Options d'import
    col1, col2 = st.columns([3, 1])
    
    with col1:
        uploaded_files = st.file_uploader(
            "Sélectionnez un ou plusieurs fichiers",
            type=['pdf', 'docx', 'txt', 'json', 'xlsx', 'csv'],
            accept_multiple_files=True,
            help="Formats supportés: PDF, DOCX, TXT, JSON, XLSX, CSV"
        )
    
    with col2:
        process_option = st.radio(
            "Traitement",
            ["Individuel", "Fusionné"],
            help="Individuel: chaque fichier séparément\nFusionné: combiner tous les fichiers"
        )
    
    if uploaded_files:
        # Créer ou récupérer le gestionnaire
        if 'doc_manager' not in st.session_state:
            st.session_state.doc_manager = DocumentManager()
        doc_manager = st.session_state.doc_manager
        
        if st.button("🚀 Importer", type="primary"):
            with st.spinner("Import en cours..."):
                if process_option == "Fusionné":
                    results = doc_manager.batch_import(uploaded_files)
                    
                    # Afficher les résultats
                    if results['success']:
                        st.success(f"✅ {len(results['success'])} fichiers importés avec succès")
                        
                        # Afficher le contenu fusionné
                        with st.expander("Contenu importé", expanded=True):
                            st.text_area(
                                "Texte extrait",
                                value=results['total_content'],
                                height=400
                            )
                            
                        # Stocker dans session_state
                        st.session_state['imported_content'] = results['total_content']
                    
                    if results['failed']:
                        st.error(f"❌ {len(results['failed'])} échecs")
                        for fail in results['failed']:
                            st.caption(f"- {fail['filename']}: {fail['error']}")
                else:
                    # Import individuel
                    for file in uploaded_files:
                        success, content, message = doc_manager.import_document(file)
                        
                        if success:
                            st.success(f"✅ {message}")
                            with st.expander(f"Contenu de {file.name}"):
                                st.text_area(
                                    "Texte extrait",
                                    value=content[:1000] + "..." if len(content) > 1000 else content,
                                    height=200,
                                    key=f"content_{file.name}"
                                )
                        else:
                            st.error(f"❌ {message}")
        
        # Statistiques
        if 'doc_manager' in st.session_state:
            stats = doc_manager.get_import_stats()
            if stats['count'] > 0:
                st.info(f"📊 {stats['count']} documents importés au total")

def display_export_interface(analysis: Optional[AnalyseJuridique] = None):
    """Interface d'export de documents"""
    if not analysis:
        st.warning("Aucune analyse à exporter")
        return
        
    st.markdown("### 💾 Export de l'analyse")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        format = st.selectbox(
            "Format d'export",
            options=['docx', 'xlsx', 'json', 'txt', 'pdf'],
            help="Sélectionnez le format de sortie"
        )
    
    with col2:
        include_metadata = st.checkbox(
            "Inclure métadonnées",
            value=True,
            help="Date, modèle utilisé, etc."
        )
    
    with col3:
        if st.button("📥 Générer export", type="primary"):
            if 'doc_manager' not in st.session_state:
                st.session_state.doc_manager = DocumentManager()
            doc_manager = st.session_state.doc_manager
            
            with st.spinner("Génération en cours..."):
                content, filename = doc_manager.export_analysis(
                    analysis,
                    format,
                    include_metadata
                )
                
                # Créer le bouton de téléchargement
                st.download_button(
                    label=f"⬇️ Télécharger {filename}",
                    data=content,
                    file_name=filename,
                    mime="application/octet-stream"
                )
                
                st.success(f"Export '{filename}' prêt !")

# Export
__all__ = [
    'DocumentManager',
    'display_import_interface',
    'display_export_interface'
]