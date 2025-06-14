# managers/export_manager.py
"""
Gestionnaire d'export de documents dans différents formats
"""

import base64
import io
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

logger = logging.getLogger(__name__)

class ExportManager:
    """Gestionnaire centralisé pour l'export de documents"""
    
    SUPPORTED_FORMATS = ['txt', 'json', 'docx', 'pdf', 'html', 'md']
    
    def __init__(self):
        self.export_count = 0
        self.last_export = None
    
    def export_document(
        self,
        content: Any,
        format: str,
        filename: Optional[str] = None,
        metadata: Optional[Dict] = None
    ) -> Tuple[bytes, str, str]:
        """
        Exporte un document dans le format spécifié
        
        Args:
            content: Le contenu à exporter (str, dict, ou objet)
            format: Format d'export (txt, json, docx, pdf, html, md)
            filename: Nom du fichier (généré automatiquement si None)
            metadata: Métadonnées additionnelles
            
        Returns:
            Tuple (contenu_bytes, nom_fichier, mime_type)
        """
        format = format.lower()
        if format not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Format non supporté: {format}. Formats disponibles: {self.SUPPORTED_FORMATS}")
        
        # Générer le nom de fichier si non fourni
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.{format}"
        elif not filename.endswith(f".{format}"):
            filename = f"{filename}.{format}"
        
        # Exporter selon le format
        if format == 'txt':
            content_bytes, mime_type = self._export_txt(content, metadata)
        elif format == 'json':
            content_bytes, mime_type = self._export_json(content, metadata)
        elif format == 'docx':
            content_bytes, mime_type = self._export_docx(content, metadata)
        elif format == 'pdf':
            content_bytes, mime_type = self._export_pdf(content, metadata)
        elif format == 'html':
            content_bytes, mime_type = self._export_html(content, metadata)
        elif format == 'md':
            content_bytes, mime_type = self._export_markdown(content, metadata)
        else:
            raise ValueError(f"Format {format} non implémenté")
        
        # Mettre à jour les statistiques
        self.export_count += 1
        self.last_export = {
            'filename': filename,
            'format': format,
            'size': len(content_bytes),
            'timestamp': datetime.now()
        }
        
        return content_bytes, filename, mime_type
    
    def _export_txt(self, content: Any, metadata: Optional[Dict]) -> Tuple[bytes, str]:
        """Export au format texte"""
        if isinstance(content, str):
            text = content
        elif isinstance(content, dict):
            text = self._dict_to_text(content)
        elif hasattr(content, '__dict__'):
            text = self._object_to_text(content)
        else:
            text = str(content)
        
        # Ajouter les métadonnées si présentes
        if metadata:
            header = "=== MÉTADONNÉES ===\n"
            for key, value in metadata.items():
                header += f"{key}: {value}\n"
            header += "\n=== CONTENU ===\n"
            text = header + text
        
        return text.encode('utf-8'), 'text/plain'
    
    def _export_json(self, content: Any, metadata: Optional[Dict]) -> Tuple[bytes, str]:
        """Export au format JSON"""
        data = {}
        
        if isinstance(content, dict):
            data['content'] = content
        elif isinstance(content, str):
            data['content'] = {'text': content}
        elif hasattr(content, '__dict__'):
            data['content'] = content.__dict__
        else:
            data['content'] = str(content)
        
        if metadata:
            data['metadata'] = metadata
        
        data['export_info'] = {
            'timestamp': datetime.now().isoformat(),
            'format': 'json',
            'version': '1.0'
        }
        
        json_str = json.dumps(data, ensure_ascii=False, indent=2, default=str)
        return json_str.encode('utf-8'), 'application/json'
    
    def _export_docx(self, content: Any, metadata: Optional[Dict]) -> Tuple[bytes, str]:
        """Export au format DOCX"""
        doc = Document()
        
        # Ajouter le titre
        title = metadata.get('title', 'Document exporté') if metadata else 'Document exporté'
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter les métadonnées
        if metadata:
            doc.add_heading('Informations', 1)
            for key, value in metadata.items():
                if key != 'title':
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
        
        doc.add_paragraph()  # Espace
        
        # Ajouter le contenu
        doc.add_heading('Contenu', 1)
        
        if isinstance(content, str):
            # Diviser par paragraphes
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        elif isinstance(content, dict):
            for key, value in content.items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
        elif isinstance(content, list):
            for item in content:
                doc.add_paragraph(f"• {str(item)}", style='List Bullet')
        else:
            doc.add_paragraph(str(content))
        
        # Ajouter le pied de page
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}").italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Sauvegarder en mémoire
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    def _export_pdf(self, content: Any, metadata: Optional[Dict]) -> Tuple[bytes, str]:
        """Export au format PDF"""
        # Créer le PDF en mémoire
        pdf_io = io.BytesIO()
        doc = SimpleDocTemplate(pdf_io, pagesize=A4)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Centre
        )
        
        # Contenu du PDF
        story = []
        
        # Titre
        title = metadata.get('title', 'Document exporté') if metadata else 'Document exporté'
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Métadonnées
        if metadata:
            story.append(Paragraph("Informations", styles['Heading2']))
            for key, value in metadata.items():
                if key != 'title':
                    story.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
            story.append(Spacer(1, 0.3 * inch))
        
        # Contenu principal
        story.append(Paragraph("Contenu", styles['Heading2']))
        story.append(Spacer(1, 0.1 * inch))
        
        if isinstance(content, str):
            # Diviser le contenu en paragraphes
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    story.append(Paragraph(paragraph.strip(), styles['Normal']))
                    story.append(Spacer(1, 0.1 * inch))
        elif isinstance(content, dict):
            data = [[key, str(value)] for key, value in content.items()]
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
        else:
            story.append(Paragraph(str(content), styles['Normal']))
        
        # Pied de page
        story.append(Spacer(1, 0.5 * inch))
        footer_text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        story.append(Paragraph(footer_text, styles['Italic']))
        
        # Construire le PDF
        doc.build(story)
        pdf_io.seek(0)
        
        return pdf_io.getvalue(), 'application/pdf'
    
    def _export_html(self, content: Any, metadata: Optional[Dict]) -> Tuple[bytes, str]:
        """Export au format HTML"""
        title = metadata.get('title', 'Document exporté') if metadata else 'Document exporté'
        
        html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f4f4f4;
        }}
        .container {{
            background-color: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #333;
            text-align: center;
            border-bottom: 2px solid #333;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #666;
            margin-top: 30px;
        }}
        .metadata {{
            background-color: #f9f9f9;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        .metadata dt {{
            font-weight: bold;
            float: left;
            width: 150px;
        }}
        .metadata dd {{
            margin-left: 160px;
            margin-bottom: 5px;
        }}
        .content {{
            margin-top: 30px;
        }}
        .footer {{
            text-align: right;
            font-style: italic;
            color: #666;
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{title}</h1>
"""
        
        # Ajouter les métadonnées
        if metadata:
            html += """
        <div class="metadata">
            <h2>Informations</h2>
            <dl>
"""
            for key, value in metadata.items():
                if key != 'title':
                    html += f"                <dt>{key}:</dt><dd>{value}</dd>\n"
            html += """
            </dl>
        </div>
"""
        
        # Ajouter le contenu
        html += """
        <div class="content">
            <h2>Contenu</h2>
"""
        
        if isinstance(content, str):
            # Convertir les sauts de ligne en paragraphes HTML
            paragraphs = content.split('\n\n')
            for p in paragraphs:
                if p.strip():
                    html += f"            <p>{p.strip()}</p>\n"
        elif isinstance(content, dict):
            html += "            <dl>\n"
            for key, value in content.items():
                html += f"                <dt>{key}:</dt><dd>{value}</dd>\n"
            html += "            </dl>\n"
        elif isinstance(content, list):
            html += "            <ul>\n"
            for item in content:
                html += f"                <li>{item}</li>\n"
            html += "            </ul>\n"
        else:
            html += f"            <p>{str(content)}</p>\n"
        
        html += f"""
        </div>
        
        <div class="footer">
            Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}
        </div>
    </div>
</body>
</html>"""
        
        return html.encode('utf-8'), 'text/html'
    
    def _export_markdown(self, content: Any, metadata: Optional[Dict]) -> Tuple[bytes, str]:
        """Export au format Markdown"""
        title = metadata.get('title', 'Document exporté') if metadata else 'Document exporté'
        
        md = f"# {title}\n\n"
        
        # Métadonnées
        if metadata:
            md += "## Informations\n\n"
            for key, value in metadata.items():
                if key != 'title':
                    md += f"**{key}:** {value}  \n"
            md += "\n"
        
        # Contenu
        md += "## Contenu\n\n"
        
        if isinstance(content, str):
            md += content
        elif isinstance(content, dict):
            for key, value in content.items():
                md += f"**{key}:** {value}\n\n"
        elif isinstance(content, list):
            for item in content:
                md += f"- {item}\n"
        else:
            md += str(content)
        
        # Pied de page
        md += f"\n\n---\n\n*Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}*"
        
        return md.encode('utf-8'), 'text/markdown'
    
    def _dict_to_text(self, data: Dict) -> str:
        """Convertit un dictionnaire en texte formaté"""
        lines = []
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{key}:")
                for sub_key, sub_value in value.items():
                    lines.append(f"  {sub_key}: {sub_value}")
            elif isinstance(value, list):
                lines.append(f"{key}:")
                for item in value:
                    lines.append(f"  - {item}")
            else:
                lines.append(f"{key}: {value}")
        return '\n'.join(lines)
    
    def _object_to_text(self, obj: Any) -> str:
        """Convertit un objet en texte formaté"""
        if hasattr(obj, 'to_dict'):
            return self._dict_to_text(obj.to_dict())
        elif hasattr(obj, '__dict__'):
            return self._dict_to_text(obj.__dict__)
        else:
            return str(obj)
    
    def create_download_button(self, content_bytes: bytes, filename: str, mime_type: str) -> str:
        """Crée un bouton de téléchargement HTML (pour Streamlit ou autre)"""
        b64 = base64.b64encode(content_bytes).decode()
        return f'<a href="data:{mime_type};base64,{b64}" download="{filename}" class="download-button">Télécharger {filename}</a>'
    
    def get_export_stats(self) -> Dict[str, Any]:
        """Retourne les statistiques d'export"""
        return {
            'total_exports': self.export_count,
            'last_export': self.last_export,
            'supported_formats': self.SUPPORTED_FORMATS,
        }