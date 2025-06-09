# managers/style_analyzer.py
"""Analyseur de style pour documents juridiques - Version complète"""

import re
import io
import logging
import statistics
from typing import Dict, List, Set, Optional, Any, Tuple
from collections import defaultdict, Counter
from datetime import datetime
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("Module python-docx non disponible")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("Module PyPDF2 non disponible")

from modules.dataclasses import Document, StylePattern, StyleLearningResult, StyleConfig

class StyleAnalyzer:
    """Analyse et apprend le style de rédaction des documents"""
    
    def __init__(self):
        self.patterns = defaultdict(list)
        self.formules_types = defaultdict(set)
        self.structures = defaultdict(list)
        self.learned_styles = {}  # Cache des styles appris
        
        # Patterns de numérotation étendus
        self.numbering_patterns = {
            'numeric': r'^\d+\.',
            'numeric_paren': r'^\d+\)',
            'roman_upper': r'^[IVX]+\.',
            'roman_lower': r'^[ivx]+\.',
            'alpha_upper': r'^[A-Z]\.',
            'alpha_lower': r'^[a-z]\.',
            'hierarchical': r'^\d+\.\d+',
            'section': r'^§\s*\d+',
            'dash': r'^-\s',
            'bullet': r'^[•·▪▫]\s',
            'arrow': r'^[→➤]\s'
        }
        
        # Indicateurs de formalité
        self.formality_indicators = {
            'tres_formel': [
                'attendu que', 'considérant que', 'il appert', 'nonobstant',
                'aux termes de', 'en l\'espèce', 'il échet', 'partant',
                'au demeurant', 'quant à ce', 'force est de constater',
                'il convient de relever', 'à toutes fins utiles', 'subséquemment'
            ],
            'formel': [
                'en effet', 'par ailleurs', 'en outre', 'toutefois',
                'néanmoins', 'cependant', 'dès lors', 'ainsi',
                'par conséquent', 'il résulte', 'il s\'ensuit', 'au surplus'
            ],
            'moderne': [
                'donc', 'mais', 'car', 'parce que', 'puisque',
                'c\'est pourquoi', 'en fait', 'notamment', 'aussi'
            ]
        }
        
        # Termes juridiques techniques
        self.technical_terms = [
            'assignation', 'citation', 'conclusions', 'dispositif',
            'moyens', 'prétentions', 'grief', 'cassation', 'appel',
            'intimé', 'appelant', 'demandeur', 'défendeur', 'magistrat',
            'juridiction', 'compétence', 'recevabilité', 'prescription',
            'forclusion', 'péremption', 'déchéance', 'nullité',
            'inopposabilité', 'chose jugée', 'autorité', 'exequatur',
            'contradictoire', 'délibéré', 'référé', 'ordonnance',
            'jugement', 'arrêt', 'pourvoi', 'mémoire', 'dire'
        ]
    
    def analyze_document(self, document: Document, type_acte: str) -> StylePattern:
        """Analyse un document pour en extraire le style"""
        content = document.content
        
        # Analyser la structure
        structure = self._extract_structure(content)
        
        # Extraire les formules types
        formules = self._extract_formules(content)
        
        # Analyser la mise en forme
        mise_en_forme = self._analyze_formatting(content)
        
        # Analyser le vocabulaire
        vocabulaire = self._analyze_vocabulary(content)
        
        # Extraire des paragraphes types
        paragraphes_types = self._extract_sample_paragraphs(content)
        
        # Analyser la numérotation
        numerotation = self._analyze_numbering(content)
        
        # Analyser le niveau de formalité
        formalite = self._analyze_formality(content)
        
        # Extraire les patterns d'argumentation
        argumentation = self._extract_argumentation_patterns(content)
        
        pattern = StylePattern(
            document_id=document.id,
            type_acte=type_acte,
            structure=structure,
            formules=list(formules),
            mise_en_forme=mise_en_forme,
            vocabulaire=vocabulaire,
            paragraphes_types=paragraphes_types,
            numerotation=numerotation,
            formalite=formalite,
            argumentation=argumentation
        )
        
        # Stocker le pattern
        self.patterns[type_acte].append(pattern)
        
        return pattern
    
    def analyze_word_document(self, doc_bytes: bytes, type_acte: str) -> Optional[StylePattern]:
        """Analyse un document Word pour en extraire le style"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx non disponible")
            return None
        
        try:
            doc = DocxDocument(io.BytesIO(doc_bytes))
            
            # Extraire le contenu et la structure
            content = []
            structure = {
                'sections': [],
                'styles_utilises': set(),
                'mise_en_forme_paragraphes': [],
                'fonts_utilises': set(),
                'tailles_police': set()
            }
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
                    
                    # Analyser le style du paragraphe
                    para_style = {
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': str(paragraph.alignment) if paragraph.alignment else 'LEFT',
                        'first_line_indent': paragraph.paragraph_format.first_line_indent,
                        'left_indent': paragraph.paragraph_format.left_indent,
                        'right_indent': paragraph.paragraph_format.right_indent,
                        'space_before': paragraph.paragraph_format.space_before,
                        'space_after': paragraph.paragraph_format.space_after,
                        'line_spacing': paragraph.paragraph_format.line_spacing
                    }
                    
                    # Analyser les runs pour la mise en forme du texte
                    for run in paragraph.runs:
                        if run.font.name:
                            structure['fonts_utilises'].add(run.font.name)
                        if run.font.size:
                            structure['tailles_police'].add(run.font.size)
                    
                    structure['mise_en_forme_paragraphes'].append(para_style)
                    
                    if paragraph.style:
                        structure['styles_utilises'].add(paragraph.style.name)
                    
                    # Détecter les sections
                    if paragraph.style and 'Heading' in paragraph.style.name:
                        structure['sections'].append({
                            'titre': text,
                            'niveau': paragraph.style.name,
                            'numerotation': self._detect_section_numbering(text)
                        })
            
            # Créer un document temporaire pour l'analyse
            temp_doc = Document(
                id=f"word_doc_{datetime.now().timestamp()}",
                title=type_acte,
                content='\n'.join(content),
                source='word'
            )
            
            # Analyser avec la méthode standard
            pattern = self.analyze_document(temp_doc, type_acte)
            
            # Enrichir avec les informations Word spécifiques
            pattern.structure.update({
                'word_styles': list(structure['styles_utilises']),
                'word_formatting': structure['mise_en_forme_paragraphes'][:10],
                'fonts': list(structure['fonts_utilises']),
                'font_sizes': list(structure['tailles_police'])
            })
            
            return pattern
            
        except Exception as e:
            logger.error(f"Erreur analyse document Word: {e}")
            return None
    
    def analyze_pdf_document(self, pdf_bytes: bytes, type_acte: str) -> Optional[StylePattern]:
        """Analyse un document PDF pour en extraire le style"""
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 non disponible")
            return None
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            content = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    content.append(text)
            
            # Créer un document temporaire
            temp_doc = Document(
                id=f"pdf_doc_{datetime.now().timestamp()}",
                title=type_acte,
                content='\n'.join(content),
                source='pdf'
            )
            
            # Analyser avec la méthode standard
            return self.analyze_document(temp_doc, type_acte)
            
        except Exception as e:
            logger.error(f"Erreur analyse document PDF: {e}")
            return None
    
    def learn_from_documents(self, documents: List[Document], style_name: str = None) -> StyleLearningResult:
        """Apprend le style depuis plusieurs documents"""
        if not documents:
            raise ValueError("Aucun document fourni")
        
        if not style_name:
            style_name = f"Style personnalisé ({len(documents)} documents)"
        
        result = StyleLearningResult(
            style_name=style_name,
            documents_analyzed=len(documents)
        )
        
        # Analyser chaque document
        all_patterns = []
        for doc in documents:
            pattern = self.analyze_document(doc, "general")
            all_patterns.append(pattern)
        
        # Combiner les analyses
        self._combine_patterns(all_patterns, result)
        
        # Calculer le score de confiance
        result.confidence_score = self._calculate_confidence(result)
        
        # Créer la configuration de style
        result.style_config = result.to_style_config()
        
        # Sauvegarder
        self.learned_styles[style_name] = result
        
        return result
    
    def _extract_structure(self, content: str) -> Dict[str, Any]:
        """Extrait la structure du document"""
        lines = content.split('\n')
        structure = {
            'sections': [],
            'niveau_hierarchie': 0,
            'longueur_sections': [],
            'plan_type': None,  # classique, moderne, personnalisé
            'sous_sections': defaultdict(list)
        }
        
        current_section = None
        current_level = 0
        section_content = []
        section_stack = []  # Pour gérer la hiérarchie
        
        for line in lines:
            title_info = self._analyze_title_level(line)
            
            if title_info:
                # Sauvegarder la section précédente
                if current_section:
                    structure['sections'].append({
                        'titre': current_section,
                        'niveau': current_level,
                        'longueur': len(section_content),
                        'sous_sections': len([s for s in structure['sous_sections'][current_section]])
                    })
                    structure['longueur_sections'].append(len(section_content))
                
                current_section = line.strip()
                current_level = title_info['level']
                section_content = []
                
                # Gérer la hiérarchie
                if current_level > structure['niveau_hierarchie']:
                    structure['niveau_hierarchie'] = current_level
                
                # Déterminer le type de plan
                if title_info['type'] == 'roman_upper' and current_level == 1:
                    structure['plan_type'] = 'classique'
                elif title_info['type'] == 'numeric' and current_level == 1:
                    structure['plan_type'] = 'moderne'
                
                # Gérer les sous-sections
                if section_stack and current_level > section_stack[-1][1]:
                    parent_section = section_stack[-1][0]
                    structure['sous_sections'][parent_section].append(current_section)
                
                section_stack.append((current_section, current_level))
            else:
                section_content.append(line)
        
        # Sauvegarder la dernière section
        if current_section:
            structure['sections'].append({
                'titre': current_section,
                'niveau': current_level,
                'longueur': len(section_content)
            })
        
        return structure
    
    def _analyze_title_level(self, line: str) -> Optional[Dict[str, Any]]:
        """Analyse le niveau et le type d'un titre"""
        line = line.strip()
        
        if not line:
            return None
        
        # Titres en majuscules
        if line.isupper() and len(line) > 3:
            return {'level': 1, 'type': 'uppercase'}
        
        # Numérotation romaine
        if re.match(r'^[IVX]+\.?\s+', line):
            return {'level': 1, 'type': 'roman_upper'}
        
        # Numérotation numérique
        if re.match(r'^\d+\.?\s+', line):
            return {'level': 1, 'type': 'numeric'}
        
        # Numérotation alphabétique
        if re.match(r'^[A-Z]\.\s+', line):
            return {'level': 2, 'type': 'alpha_upper'}
        
        # Numérotation hiérarchique
        if re.match(r'^\d+\.\d+\.?\s+', line):
            return {'level': 2, 'type': 'hierarchical'}
        
        # Sous-numérotation
        if re.match(r'^[a-z]\)\s+', line):
            return {'level': 3, 'type': 'alpha_lower_paren'}
        
        return None
    
    def _analyze_numbering(self, content: str) -> Dict[str, Any]:
        """Analyse détaillée de la numérotation"""
        numbering_info = {
            'primary_style': None,
            'secondary_style': None,
            'uses_hierarchy': False,
            'consistency': 1.0,
            'patterns_found': {}
        }
        
        lines = content.split('\n')
        pattern_counts = Counter()
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            for style, pattern in self.numbering_patterns.items():
                if re.match(pattern, line):
                    pattern_counts[style] += 1
                    break
        
        if pattern_counts:
            # Style principal
            most_common = pattern_counts.most_common()
            numbering_info['primary_style'] = most_common[0][0]
            numbering_info['patterns_found'] = dict(pattern_counts)
            
            # Style secondaire
            if len(most_common) > 1:
                numbering_info['secondary_style'] = most_common[1][0]
            
            # Vérifier la hiérarchie
            if 'hierarchical' in pattern_counts or (
                'numeric' in pattern_counts and 'alpha_upper' in pattern_counts
            ):
                numbering_info['uses_hierarchy'] = True
            
            # Calculer la cohérence
            total_numbered = sum(pattern_counts.values())
            main_style_count = most_common[0][1]
            numbering_info['consistency'] = main_style_count / total_numbered
        
        return numbering_info
    
    def _analyze_formality(self, content: str) -> Dict[str, Any]:
        """Analyse le niveau de formalité du document"""
        formality_info = {
            'level': 'moderne',  # moderne, formel, tres_formel
            'score': 0.5,
            'indicators_found': {},
            'density': 0.0
        }
        
        content_lower = content.lower()
        total_words = len(content_lower.split())
        
        # Compter les indicateurs
        indicator_counts = {}
        total_indicators = 0
        
        for level, indicators in self.formality_indicators.items():
            count = 0
            found = []
            for indicator in indicators:
                occurrences = content_lower.count(indicator)
                if occurrences > 0:
                    count += occurrences
                    found.append(indicator)
            
            indicator_counts[level] = count
            total_indicators += count
            if found:
                formality_info['indicators_found'][level] = found
        
        # Déterminer le niveau
        if indicator_counts['tres_formel'] >= 5:
            formality_info['level'] = 'tres_formel'
            formality_info['score'] = 0.9
        elif indicator_counts['formel'] >= 10:
            formality_info['level'] = 'formel'
            formality_info['score'] = 0.7
        else:
            formality_info['level'] = 'moderne'
            formality_info['score'] = 0.4
        
        # Calculer la densité
        formality_info['density'] = total_indicators / max(total_words, 1) * 100
        
        # Ajuster le score selon la densité technique
        technical_count = sum(1 for term in self.technical_terms if term in content_lower)
        technical_density = technical_count / max(total_words, 1) * 100
        
        if technical_density > 2:
            formality_info['score'] = min(1.0, formality_info['score'] + 0.1)
        
        return formality_info
    
    def _extract_argumentation_patterns(self, content: str) -> Dict[str, List[str]]:
        """Extrait les patterns d'argumentation"""
        patterns = {
            'introduction': [],
            'developpement': [],
            'transition': [],
            'conclusion': [],
            'citation': []
        }
        
        # Patterns d'introduction
        intro_patterns = [
            r'[Ii]l convient (?:tout d\'abord |préalablement |)de (?:rappeler|préciser|relever|souligner)',
            r'[Ii]l résulte de (?:ce qui précède|l\'instruction|la procédure|ces éléments)',
            r'[Ii]l apparaît (?:clairement |manifestement |)que',
            r'[Ff]orce est de constater que',
            r'[Ii]l est (?:constant|établi|démontré|avéré) que',
            r'[Àà] titre (?:principal|subsidiaire|liminaire|préliminaire)',
            r'[Ee]n (?:premier|second|dernier) lieu',
            r'[Dd]\'une part',
            r'[Tt]out d\'abord',
            r'[Pp]remièrement'
        ]
        
        # Patterns de développement
        dev_patterns = [
            r'[Ee]n l\'(?:espèce|occurrence)',
            r'[Ss]\'agissant de',
            r'[Cc]oncernant',
            r'[Qq]uant à',
            r'[Pp]our ce qui (?:est|concerne)',
            r'[Dd]ans (?:ces conditions|ce contexte|cette hypothèse)',
            r'[Cc]ompte tenu de',
            r'[Ee]u égard à',
            r'[Aa]u (?:regard|vu) de'
        ]
        
        # Patterns de transition
        trans_patterns = [
            r'[Pp]ar ailleurs',
            r'[Ee]n outre',
            r'[Dd]e (?:plus|surcroît)',
            r'[Aa]u (?:surplus|demeurant)',
            r'[Cc]ependant',
            r'[Tt]outefois',
            r'[Nn]éanmoins',
            r'[Pp]our autant',
            r'[Dd]\'autre part',
            r'[Ee]nfin',
            r'[Aa]u final'
        ]
        
        # Patterns de conclusion
        concl_patterns = [
            r'[Pp]ar (?:conséquent|suite)',
            r'[Dd]ès lors',
            r'[Aa]insi',
            r'[Dd]onc',
            r'[Ii]l s\'ensuit que',
            r'[Ee]n (?:conséquence|définitive)',
            r'[Aa]u (?:total|final)',
            r'[Ee]n (?:conclusion|résumé|somme)',
            r'[Pp]artant',
            r'[Cc]\'est (?:pourquoi|la raison pour laquelle)'
        ]
        
        # Patterns de citation
        citation_patterns = [
            r'[Aa]ux termes de (?:l\'article|la jurisprudence)',
            r'[Ss]elon (?:l\'article|la Cour|le Tribunal)',
            r'[Cc]onformément à',
            r'[Ee]n application de',
            r'[Aa]u sens de',
            r'[Tt]el que (?:défini|prévu|stipulé) (?:par|à)',
            r'[Vv]isé(?:e|s)? (?:à|aux) article'
        ]
        
        # Rechercher les patterns dans le contenu
        all_patterns = {
            'introduction': intro_patterns,
            'developpement': dev_patterns,
            'transition': trans_patterns,
            'conclusion': concl_patterns,
            'citation': citation_patterns
        }
        
        for category, pattern_list in all_patterns.items():
            for pattern in pattern_list:
                matches = re.findall(pattern, content, re.IGNORECASE)
                if matches:
                    # Nettoyer et ajouter les patterns trouvés
                    for match in matches[:3]:  # Limiter à 3 exemples
                        clean_match = re.sub(r'\s+', ' ', match).strip()
                        if clean_match not in patterns[category]:
                            patterns[category].append(clean_match)
        
        return patterns
    
    def _extract_formules(self, content: str) -> Set[str]:
        """Extrait les formules types du document"""
        formules = set()
        
        # Patterns étendus
        patterns = [
            # Formules d'introduction
            r"J'ai l'honneur de.*?[.!]",
            r"Je soussigné.*?[.!]",
            r"Par la présente.*?[.!]",
            r"Je me permets de.*?[.!]",
            
            # Formules juridiques
            r"Il résulte de.*?[.!]",
            r"Aux termes de.*?[.!]",
            r"En l'espèce.*?[.!]",
            r"Par ces motifs.*?[.!]",
            r"Il convient de.*?[.!]",
            r"Force est de constater.*?[.!]",
            r"Il apparaît que.*?[.!]",
            r"Attendu que.*?[.!]",
            r"Considérant que.*?[.!]",
            r"Il est constant que.*?[.!]",
            r"Il n'est pas contesté que.*?[.!]",
            
            # Formules de demande
            r"Je vous (?:prie|saurais gré) de.*?[.!]",
            r"Je sollicite.*?[.!]",
            r"Il est demandé.*?[.!]",
            r"Plaise au Tribunal.*?[.!]",
            
            # Formules de conclusion
            r"Je vous prie d'agréer.*?[.!]",
            r"Veuillez agréer.*?[.!]",
            r"Dans l'attente de.*?[.!]",
            r"Je reste à votre disposition.*?[.!]",
            r"Vous en souhaitant.*?[.!]",
            r"En vous remerciant.*?[.!]",
            r"Avec mes remerciements.*?[.!]",
            
            # Formules de référence
            r"Comme suite à.*?[.!]",
            r"En référence à.*?[.!]",
            r"Pour faire suite à.*?[.!]",
            r"Faisant suite à.*?[.!]"
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                formule = match.group(0).strip()
                # Nettoyer la formule
                formule = re.sub(r'\s+', ' ', formule)
                # Limiter la longueur
                if 10 < len(formule) < 200:
                    formules.add(formule)
        
        return formules
    
    def _combine_patterns(self, patterns: List[StylePattern], result: StyleLearningResult):
        """Combine plusieurs patterns pour créer un style unifié"""
        # Combiner les structures
        all_sections = []
        for pattern in patterns:
            all_sections.extend(pattern.structure.get('sections', []))
        
        # Analyser la longueur des phrases et paragraphes
        all_sentences = []
        all_paragraphs = []
        
        for pattern in patterns:
            content = pattern.vocabulaire  # Utiliser le contenu original
            if isinstance(content, dict):
                # Si c'est le vocabulaire, reconstituer depuis les paragraphes types
                content = '\n'.join(pattern.paragraphes_types)
            
            sentences = re.split(r'[.!?]+', str(content))
            all_sentences.extend([s.strip() for s in sentences if s.strip()])
            
            paragraphs = str(content).split('\n\n')
            all_paragraphs.extend([p.strip() for p in paragraphs if p.strip()])
        
        # Calculer les moyennes
        if all_sentences:
            sentence_lengths = [len(s.split()) for s in all_sentences]
            result.average_sentence_length = int(statistics.mean(sentence_lengths))
        
        if all_paragraphs:
            paragraph_lengths = [len(p.split()) for p in all_paragraphs]
            result.average_paragraph_length = int(statistics.mean(paragraph_lengths))
        
        # Combiner la numérotation
        numbering_styles = Counter()
        for pattern in patterns:
            if hasattr(pattern, 'numerotation') and pattern.numerotation:
                primary = pattern.numerotation.get('primary_style')
                if primary:
                    numbering_styles[primary] += 1
        
        if numbering_styles:
            result.paragraph_numbering_style = numbering_styles.most_common(1)[0][0]
            result.paragraph_numbering_pattern = self._get_pattern_for_style(
                result.paragraph_numbering_style
            )
        
        # Combiner la formalité
        formality_scores = []
        for pattern in patterns:
            if hasattr(pattern, 'formalite') and pattern.formalite:
                formality_scores.append(pattern.formalite.get('score', 0.5))
        
        if formality_scores:
            result.formality_score = statistics.mean(formality_scores)
        
        # Combiner les formules et phrases types
        all_formules = set()
        for pattern in patterns:
            all_formules.update(pattern.formules)
        
        # Sélectionner les plus représentatives
        result.common_phrases = list(all_formules)[:20]
        
        # Combiner les patterns d'argumentation
        all_arg_patterns = defaultdict(list)
        for pattern in patterns:
            if hasattr(pattern, 'argumentation'):
                for category, items in pattern.argumentation.items():
                    all_arg_patterns[category].extend(items)
        
        # Dédupliquer et limiter
        for category, items in all_arg_patterns.items():
            unique_items = list(set(items))
            if category == 'transition':
                result.transition_words = unique_items[:10]
            elif category in ['introduction', 'developpement', 'conclusion']:
                result.argument_patterns.extend(unique_items[:5])
        
        # Extraire les patterns de citation
        result.citation_patterns = list(set(all_arg_patterns.get('citation', [])))[:5]
    
    def _get_pattern_for_style(self, style: str) -> str:
        """Retourne le pattern de numérotation pour un style"""
        patterns_map = {
            'numeric': '1.',
            'numeric_paren': '1)',
            'roman_upper': 'I.',
            'roman_lower': 'i.',
            'alpha_upper': 'A.',
            'alpha_lower': 'a.',
            'hierarchical': '1.1',
            'section': '§ 1',
            'dash': '-',
            'bullet': '•'
        }
        return patterns_map.get(style, '1.')
    
    def _calculate_confidence(self, result: StyleLearningResult) -> float:
        """Calcule un score de confiance pour l'apprentissage"""
        confidence = 0.5
        
        # Plus de documents = plus de confiance
        if result.documents_analyzed >= 5:
            confidence += 0.2
        elif result.documents_analyzed >= 3:
            confidence += 0.1
        
        # Patterns cohérents = plus de confiance
        if result.paragraph_numbering_style:
            confidence += 0.1
        
        if len(result.argument_patterns) >= 3:
            confidence += 0.1
        
        if len(result.common_phrases) >= 5:
            confidence += 0.1
        
        # Formalité cohérente
        if 0.3 < result.formality_score < 0.9:
            confidence += 0.05
        
        return min(confidence, 1.0)
    
    def generate_with_style(self, type_acte: str, contenu_base: str, style_name: str = None) -> str:
        """Génère du contenu en appliquant le style appris"""
        if style_name and style_name in self.learned_styles:
            # Utiliser un style appris spécifique
            style = self.learned_styles[style_name]
            return self.apply_learned_style(contenu_base, style)
        elif type_acte in self.patterns:
            # Utiliser les patterns du type d'acte
            patterns = self.patterns[type_acte]
            if not patterns:
                return contenu_base
            
            pattern = patterns[0]  # Prendre le premier ou faire une moyenne
            
            styled_content = self._apply_structure(contenu_base, pattern.structure)
            styled_content = self._insert_formules(styled_content, pattern.formules)
            styled_content = self._apply_formatting(styled_content, pattern.mise_en_forme)
            styled_content = self._apply_numbering_style(styled_content, pattern.numerotation)
            
            return styled_content
        else:
            return contenu_base
    
    def apply_learned_style(self, text: str, style: StyleLearningResult) -> str:
        """Applique un style appris à un texte"""
        # Appliquer la numérotation
        if style.paragraph_numbering_pattern:
            text = self._apply_paragraph_numbering(text, style)
        
        # Insérer des phrases types
        if style.common_phrases:
            text = self._enhance_with_common_phrases(text, style.common_phrases)
        
        # Appliquer les patterns d'argumentation
        if style.argument_patterns:
            text = self._apply_argumentation_patterns(text, style)
        
        # Ajuster la formalité
        if style.formality_score > 0.7:
            text = self._increase_formality(text, style)
        
        return text
    
    def _apply_paragraph_numbering(self, text: str, style: StyleLearningResult) -> str:
        """Applique le style de numérotation aux paragraphes"""
        paragraphs = text.split('\n\n')
        numbered_paragraphs = []
        
        counter = 1
        pattern = style.paragraph_numbering_pattern
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                numbered_paragraphs.append('')
                continue
            
            # Ne pas numéroter les titres
            if self._is_title(para):
                numbered_paragraphs.append(para)
                continue
            
            # Générer le numéro selon le pattern
            if pattern == '1.':
                number = f"{counter}."
            elif pattern == '1)':
                number = f"{counter})"
            elif pattern == 'I.':
                number = f"{self._to_roman(counter)}."
            elif pattern == 'A.':
                number = f"{chr(64 + counter)}."
            elif pattern == 'a.':
                number = f"{chr(96 + counter)}."
            elif pattern == '§ 1':
                number = f"§ {counter}"
            elif pattern == '-':
                number = "-"
            elif pattern == '•':
                number = "•"
            else:
                number = f"{counter}."
            
            numbered_paragraphs.append(f"{number} {para}")
            
            # Incrémenter seulement pour les styles numérotés
            if pattern not in ['-', '•']:
                counter += 1
        
        return '\n\n'.join(numbered_paragraphs)
    
    def _enhance_with_common_phrases(self, text: str, phrases: List[str]) -> str:
        """Enrichit le texte avec des phrases types"""
        # Insérer des phrases d'introduction si le texte n'en a pas
        intro_phrases = [p for p in phrases if any(
            start in p.lower() for start in ["j'ai l'honneur", "je soussigné", "par la présente"]
        )]
        
        if intro_phrases and not any(phrase in text[:200] for phrase in intro_phrases):
            # Suggérer une phrase d'introduction
            text = f"[Suggestion: {intro_phrases[0]}]\n\n{text}"
        
        # Ajouter des formules de conclusion si absentes
        conclusion_phrases = [p for p in phrases if any(
            end in p.lower() for end in ["veuillez agréer", "je vous prie d'agréer", "dans l'attente"]
        )]
        
        if conclusion_phrases and not any(phrase in text[-300:] for phrase in conclusion_phrases):
            text = f"{text}\n\n[Suggestion de conclusion: {conclusion_phrases[0]}]"
        
        return text
    
    def _apply_argumentation_patterns(self, text: str, style: StyleLearningResult) -> str:
        """Applique les patterns d'argumentation au texte"""
        paragraphs = text.split('\n\n')
        enhanced_paragraphs = []
        
        # Insérer des transitions entre les paragraphes
        for i, para in enumerate(paragraphs):
            if i > 0 and style.transition_words and i % 3 == 0:
                # Ajouter une transition tous les 3 paragraphes
                transition = style.transition_words[i % len(style.transition_words)]
                if not para.lower().startswith(transition.lower()):
                    para = f"{transition}, {para[0].lower()}{para[1:]}"
            
            enhanced_paragraphs.append(para)
        
        return '\n\n'.join(enhanced_paragraphs)
    
    def _increase_formality(self, text: str, style: StyleLearningResult) -> str:
        """Augmente le niveau de formalité du texte"""
        # Remplacements pour augmenter la formalité
        replacements = {
            r'\bmais\b': 'toutefois',
            r'\bdonc\b': 'par conséquent',
            r'\bcar\b': 'en effet',
            r'\baussi\b': 'également',
            r'\bpour\b': 'aux fins de',
            r'\bavec\b': 'au moyen de',
            r'\bsans\b': 'en l\'absence de',
            r'\bil faut\b': 'il convient de',
            r'\bon doit\b': 'il y a lieu de',
            r'\bje pense\b': 'j\'estime',
            r'\bje crois\b': 'je considère'
        }
        
        for pattern, replacement in replacements.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _to_roman(self, num: int) -> str:
        """Convertit un nombre en chiffres romains"""
        values = [
            (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
        ]
        result = ''
        for value, letter in values:
            count = num // value
            if count:
                result += letter * count
                num -= value * count
        return result
    
    def _detect_section_numbering(self, title: str) -> Optional[str]:
        """Détecte le type de numérotation d'une section"""
        for style, pattern in self.numbering_patterns.items():
            if re.match(pattern, title.strip()):
                return style
        return None
    
    def _apply_numbering_style(self, content: str, numbering_info: Dict[str, Any]) -> str:
        """Applique un style de numérotation au contenu"""
        if not numbering_info or not numbering_info.get('primary_style'):
            return content
        
        style = numbering_info['primary_style']
        pattern = self._get_pattern_for_style(style)
        
        # Appliquer selon le style
        lines = content.split('\n')
        numbered_lines = []
        counter = 1
        
        for line in lines:
            if line.strip() and not self._is_title(line):
                if style in ['dash', 'bullet']:
                    numbered_lines.append(f"{pattern} {line}")
                else:
                    numbered_lines.append(f"{counter}{pattern[1:]} {line}")
                    counter += 1
            else:
                numbered_lines.append(line)
        
        return '\n'.join(numbered_lines)
    
    def get_style_summary(self, style: StyleLearningResult) -> str:
        """Génère un résumé du style appris"""
        summary = []
        
        summary.append(f"**Style : {style.style_name}**\n")
        summary.append(f"Documents analysés : {style.documents_analyzed}")
        summary.append(f"Confiance : {style.confidence_score:.0%}\n")
        
        # Structure
        summary.append("**Structure :**")
        if style.paragraph_numbering_pattern:
            summary.append(f"- Numérotation : {style.paragraph_numbering_pattern}")
        summary.append(f"- Longueur moyenne des phrases : {style.average_sentence_length} mots")
        summary.append(f"- Longueur moyenne des paragraphes : {style.average_paragraph_length} mots")
        
        # Formalité
        summary.append("\n**Niveau de formalité :**")
        if style.formality_score > 0.8:
            summary.append("- Très formel (style traditionnel)")
        elif style.formality_score > 0.6:
            summary.append("- Formel (style classique)")
        else:
            summary.append("- Moderne (style contemporain)")
        
        # Vocabulaire
        if style.transition_words:
            summary.append(f"\n**Mots de liaison privilégiés :**")
            summary.append(f"- {', '.join(style.transition_words[:8])}")
        
        # Phrases types
        if style.common_phrases:
            summary.append(f"\n**Phrases types identifiées :**")
            for i, phrase in enumerate(style.common_phrases[:5], 1):
                # Tronquer si trop long
                if len(phrase) > 80:
                    phrase = phrase[:77] + "..."
                summary.append(f"{i}. {phrase}")
        
        # Argumentation
        if style.argument_patterns:
            summary.append(f"\n**Patterns d'argumentation :**")
            for pattern in style.argument_patterns[:5]:
                summary.append(f"- {pattern}")
        
        return "\n".join(summary)