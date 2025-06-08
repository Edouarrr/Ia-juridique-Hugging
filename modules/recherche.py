# modules/recherche.py
"""Module de recherche unifiée avec 100% des fonctionnalités intégrées"""

import streamlit as st
import asyncio
from datetime import datetime, timedelta
import re
import json
from collections import defaultdict, Counter
from typing import List, Dict, Optional, Any, Tuple, Union
import io
import base64
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

# Imports optionnels avec gestion d'erreur
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("⚠️ pandas non installé - Certaines fonctionnalités seront limitées")

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False
    print("⚠️ plotly non installé - Visualisations simplifiées")

try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False
    print("⚠️ networkx non installé - Analyse réseau limitée")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx non installé - Export Word non disponible")

# Imports des modules internes
from config.app_config import SearchMode, ANALYSIS_PROMPTS_AFFAIRES, LLMProvider
from managers.azure_blob_manager import AzureBlobManager
from managers.azure_search_manager import AzureSearchManager
from managers.multi_llm_manager import MultiLLMManager
from managers.document_manager import display_import_interface
from managers.dynamic_generators import generate_dynamic_search_prompts
from managers.legal_search import LegalSearchManager, display_legal_search_interface
from managers.jurisprudence_verifier import JurisprudenceVerifier
from models.dataclasses import Document, PieceSelectionnee
from utils.helpers import clean_key

# === FONCTION PRINCIPALE D'AFFICHAGE ===

def show_page():
    """Fonction principale de la page recherche universelle"""
    
    st.markdown("## 🔍 Recherche Universelle")
    
    # Barre de recherche principale
    col1, col2 = st.columns([5, 1])
    
    with col1:
        query = st.text_input(
            "Entrez votre commande ou recherche",
            placeholder="Ex: rédiger conclusions @affaire_martin, analyser risques, importer documents...",
            key="universal_query",
            help="Utilisez @ pour référencer une affaire spécifique"
        )
    
    with col2:
        search_button = st.button("🔍 Rechercher", key="search_button", use_container_width=True)
    
    # Suggestions de commandes
    with st.expander("💡 Exemples de commandes", expanded=False):
        st.markdown("""
        **Recherche :**
        - `contrats société XYZ`
        - `@affaire_martin documents comptables`
        
        **Analyse :**
        - `analyser les risques @dossier_pénal`
        - `identifier les infractions @affaire_corruption`
        
        **Rédaction :**
        - `rédiger conclusions défense @affaire_martin abus biens sociaux`
        - `créer plainte avec constitution partie civile escroquerie`
        
        **Visualisations :**
        - `chronologie des faits @affaire_martin`
        - `cartographie des sociétés @groupe_abc`
        
        **Gestion :**
        - `importer documents PDF`
        - `exporter analyse format word`
        - `envoyer par email @destinataire`
        """)
    
    # Traiter la requête
    if query and (search_button or st.session_state.get('process_query', False)):
        with st.spinner("🔄 Traitement en cours..."):
            process_universal_query(query)
    
    # Afficher les résultats
    show_unified_results_tab()
    
    # Réinitialiser le flag de traitement
    if 'process_query' in st.session_state:
        st.session_state.process_query = False

def process_universal_query(query: str):
    """Traite une requête universelle et route vers la bonne fonction"""
    
    # Sauvegarder la requête
    st.session_state.last_universal_query = query
    
    # Analyser la requête
    analysis = analyze_query(query)
    st.session_state.current_analysis = analysis
    
    # Router selon le type détecté
    if analysis['type'] == 'import':
        process_import_request(query, analysis)
    
    elif analysis['type'] == 'export':
        process_export_request(query, analysis)
    
    elif analysis['type'] == 'email':
        process_email_request(query, analysis)
    
    elif analysis['type'] == 'piece_selection':
        process_piece_selection_request(query, analysis)
    
    elif analysis['type'] == 'bordereau':
        process_bordereau_request(query, analysis)
    
    elif analysis['type'] == 'synthesis':
        process_synthesis_request(query, analysis)
    
    elif analysis['type'] == 'template':
        process_template_request(query, analysis)
    
    elif analysis['type'] == 'jurisprudence':
        process_jurisprudence_request(query, analysis)
    
    elif analysis['type'] == 'redaction':
        process_redaction_request(query, analysis)
    
    elif analysis['type'] == 'plaidoirie':
        process_plaidoirie_request(query, analysis)
    
    elif analysis['type'] == 'preparation_client':
        process_preparation_client_request(query, analysis)
    
    elif analysis['type'] == 'timeline':
        process_timeline_request(query, analysis)
    
    elif analysis['type'] == 'mapping':
        process_mapping_request(query, analysis)
    
    elif analysis['type'] == 'comparison':
        process_comparison_request(query, analysis)
    
    elif analysis['type'] == 'analysis':
        process_analysis_request(query, analysis)
    
    else:  # Recherche simple par défaut
        process_search_request(query, analysis)

def analyze_query(query: str) -> dict:
    """Analyse une requête pour déterminer le type d'action"""
    
    query_lower = query.lower()
    
    # Patterns de détection
    patterns = {
        'import': ['import', 'importer', 'charger', 'upload'],
        'export': ['export', 'exporter', 'télécharger', 'download'],
        'email': ['email', 'envoyer', 'mail', 'courrier électronique'],
        'piece_selection': ['sélectionner pièces', 'choisir pièces', 'pièces'],
        'bordereau': ['bordereau', 'liste pièces', 'inventaire'],
        'synthesis': ['synthèse', 'synthétiser', 'résumer'],
        'template': ['template', 'modèle', 'gabarit'],
        'jurisprudence': ['jurisprudence', 'juris', 'décision', 'arrêt'],
        'redaction': ['rédiger', 'écrire', 'créer', 'conclusions', 'plainte'],
        'plaidoirie': ['plaidoirie', 'plaider', 'audience'],
        'preparation_client': ['préparer client', 'préparation', 'coaching'],
        'timeline': ['chronologie', 'timeline', 'frise'],
        'mapping': ['cartographie', 'mapping', 'carte', 'réseau'],
        'comparison': ['comparer', 'comparaison', 'différences'],
        'analysis': ['analyser', 'analyse', 'étudier', 'examiner']
    }
    
    # Détecter le type
    detected_type = 'search'  # Par défaut
    
    for pattern_type, keywords in patterns.items():
        if any(kw in query_lower for kw in keywords):
            detected_type = pattern_type
            break
    
    # Extraire la référence @ si présente
    reference = None
    ref_match = re.search(r'@(\w+)', query)
    if ref_match:
        reference = ref_match.group(1)
    
    # Extraire d'autres détails selon le type
    details = extract_query_details(query, detected_type)
    
    return {
        'type': detected_type,
        'query': query,
        'reference': reference,
        'details': details
    }

def extract_query_details(query: str, query_type: str) -> dict:
    """Extrait les détails spécifiques selon le type de requête"""
    
    details = {}
    
    if query_type == 'export':
        # Détecter le format
        if 'word' in query.lower() or 'docx' in query.lower():
            details['format'] = 'docx'
        elif 'pdf' in query.lower():
            details['format'] = 'pdf'
        elif 'excel' in query.lower() or 'xlsx' in query.lower():
            details['format'] = 'xlsx'
        else:
            details['format'] = 'docx'  # Par défaut
    
    elif query_type == 'redaction':
        # Détecter le type de document
        if 'conclusions' in query.lower():
            details['document_type'] = 'conclusions'
        elif 'plainte' in query.lower():
            details['document_type'] = 'plainte'
        elif 'courrier' in query.lower():
            details['document_type'] = 'courrier'
        else:
            details['document_type'] = 'general'
    
    # Ajouter d'autres extractions selon les besoins
    
    return details

# === FONCTIONS STUB MANQUANTES ===

def process_redaction_request(query: str, analysis: dict):
    """Traite une demande de rédaction"""
    st.info("🚧 Fonction de rédaction en cours d'implémentation")
    # TODO: Implémenter la rédaction

def process_plaidoirie_request(query: str, analysis: dict):
    """Traite une demande de plaidoirie"""
    st.info("🚧 Fonction de plaidoirie en cours d'implémentation")
    # TODO: Implémenter la plaidoirie

def process_preparation_client_request(query: str, analysis: dict):
    """Traite une demande de préparation client"""
    st.info("🚧 Fonction de préparation client en cours d'implémentation")
    # TODO: Implémenter la préparation

def process_timeline_request(query: str, analysis: dict):
    """Traite une demande de chronologie"""
    st.info("🚧 Fonction de chronologie en cours d'implémentation")
    # TODO: Implémenter la timeline

def process_mapping_request(query: str, analysis: dict):
    """Traite une demande de cartographie"""
    st.info("🚧 Fonction de cartographie en cours d'implémentation")
    # TODO: Implémenter le mapping

def process_comparison_request(query: str, analysis: dict):
    """Traite une demande de comparaison"""
    st.info("🚧 Fonction de comparaison en cours d'implémentation")
    # TODO: Implémenter la comparaison

# Configuration des styles de rédaction
REDACTION_STYLES = {
    'formel': {
        'name': 'Formel',
        'description': 'Style juridique classique et solennel',
        'tone': 'respectueux et distant',
        'vocabulary': 'technique et précis'
    },
    'persuasif': {
        'name': 'Persuasif',
        'description': 'Style argumentatif et convaincant',
        'tone': 'assertif et engagé',
        'vocabulary': 'percutant et imagé'
    },
    'technique': {
        'name': 'Technique',
        'description': 'Style factuel et détaillé',
        'tone': 'neutre et objectif',
        'vocabulary': 'spécialisé et exhaustif'
    },
    'synthétique': {
        'name': 'Synthétique',
        'description': 'Style concis et efficace',
        'tone': 'direct et clair',
        'vocabulary': 'simple et précis'
    },
    'pédagogique': {
        'name': 'Pédagogique',
        'description': 'Style explicatif et accessible',
        'tone': 'bienveillant et didactique',
        'vocabulary': 'vulgarisé et illustré'
    }
}

# Templates prédéfinis
DOCUMENT_TEMPLATES = {
    'conclusions_defense': {
        'name': 'Conclusions en défense',
        'structure': [
            'I. FAITS ET PROCÉDURE',
            'II. DISCUSSION',
            '   A. Sur la recevabilité',
            '   B. Sur le fond',
            '      1. Sur l\'élément matériel',
            '      2. Sur l\'élément intentionnel',
            '      3. Sur le préjudice',
            'III. PAR CES MOTIFS'
        ],
        'style': 'formel'
    },
    'plainte_simple': {
        'name': 'Plainte simple',
        'structure': [
            'Objet : Plainte',
            'EXPOSÉ DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PRÉJUDICES SUBIS',
            'DEMANDES',
            'PIÈCES JOINTES'
        ],
        'style': 'formel'
    },
    'plainte_avec_cpc': {
        'name': 'Plainte avec constitution de partie civile',
        'structure': [
            'Objet : Plainte avec constitution de partie civile',
            'EXPOSÉ DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PRÉJUDICES SUBIS',
            'CONSTITUTION DE PARTIE CIVILE',
            'ÉVALUATION DU PRÉJUDICE',
            'DEMANDES',
            'PIÈCES JOINTES'
        ],
        'style': 'formel'
    },
    'mise_en_demeure': {
        'name': 'Mise en demeure',
        'structure': [
            'MISE EN DEMEURE',
            'Rappel des faits',
            'Obligations non respectées',
            'Délai accordé',
            'Conséquences du défaut',
            'Réserves'
        ],
        'style': 'persuasif'
    },
    'note_synthese': {
        'name': 'Note de synthèse',
        'structure': [
            'SYNTHÈSE EXÉCUTIVE',
            'I. CONTEXTE',
            'II. ANALYSE',
            'III. RECOMMANDATIONS',
            'IV. PLAN D\'ACTION'
        ],
        'style': 'synthétique'
    }
}

# === FONCTIONS DE TRAITEMENT POUR IMPORT ===

def process_import_request(query: str, analysis: dict):
    """Traite une demande d'import"""
    
    st.info("📥 Interface d'import activée")
    
    # Configuration de l'import depuis la config
    file_types = analysis['details'].get('file_types', ['pdf', 'docx', 'txt'])
    destination = st.session_state.get('import_destination', 'Documents locaux')
    
    # Si des fichiers sont déjà uploadés
    if st.session_state.get('import_files'):
        with st.spinner("📥 Import des fichiers en cours..."):
            imported_count = 0
            
            for file in st.session_state.import_files:
                try:
                    # Lire le contenu
                    content = file.read()
                    
                    # Traiter selon le type
                    if file.name.endswith('.pdf'):
                        text_content = extract_text_from_pdf(content)
                    elif file.name.endswith('.docx'):
                        text_content = extract_text_from_docx(content)
                    else:
                        text_content = content.decode('utf-8', errors='ignore')
                    
                    # Stocker selon la destination
                    if destination == 'Azure Blob':
                        store_in_azure_blob(file.name, content, analysis.get('reference'))
                    else:
                        store_locally(file.name, text_content, analysis.get('reference'))
                    
                    imported_count += 1
                    
                except Exception as e:
                    st.error(f"❌ Erreur import {file.name}: {str(e)}")
            
            st.success(f"✅ {imported_count} fichiers importés avec succès")
            
            # Analyse automatique si demandée
            if st.session_state.get('auto_analyze_import', True):
                st.info("🤖 Analyse automatique des documents importés...")
                analyze_imported_documents()
    
    else:
        st.warning("⚠️ Aucun fichier sélectionné. Utilisez l'interface d'import ci-dessous.")

def extract_text_from_pdf(content: bytes) -> str:
    """Extrait le texte d'un PDF"""
    try:
        import PyPDF2
        import io
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text
        
    except ImportError:
        return "PDF - Installation de PyPDF2 requise pour l'extraction"
    except Exception as e:
        return f"Erreur extraction PDF: {str(e)}"

def extract_text_from_docx(content: bytes) -> str:
    """Extrait le texte d'un DOCX"""
    if DOCX_AVAILABLE:
        try:
            import io
            doc = docx.Document(io.BytesIO(content))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"Erreur extraction DOCX: {str(e)}"
    else:
        return "DOCX - Installation de python-docx requise"

def store_locally(filename: str, content: str, reference: str = None):
    """Stocke un document localement"""
    doc_id = f"imported_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{clean_key(filename)}"
    
    if 'azure_documents' not in st.session_state:
        st.session_state.azure_documents = {}
    
    st.session_state.azure_documents[doc_id] = Document(
        id=doc_id,
        title=filename,
        content=content,
        source=f"Import {reference if reference else 'direct'}",
        metadata={
            'imported_at': datetime.now().isoformat(),
            'reference': reference,
            'original_filename': filename
        }
    )

def store_in_azure_blob(filename: str, content: bytes, reference: str = None):
    """Stocke un document dans Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        container = reference if reference else 'imports'
        blob_name = f"{datetime.now().strftime('%Y%m%d')}_{filename}"
        
        try:
            blob_manager.upload_blob(container, blob_name, content)
            st.success(f"✅ {filename} uploadé dans Azure")
        except Exception as e:
            st.error(f"❌ Erreur Azure: {str(e)}")
    else:
        st.warning("⚠️ Azure Blob non connecté - Stockage local")
        store_locally(filename, content.decode('utf-8', errors='ignore'), reference)

def analyze_imported_documents():
    """Analyse automatique des documents importés"""
    # Récupérer les derniers documents importés
    recent_docs = []
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if 'imported_at' in doc.metadata:
            import_time = datetime.fromisoformat(doc.metadata['imported_at'])
            if (datetime.now() - import_time).seconds < 300:  # Moins de 5 minutes
                recent_docs.append(doc)
    
    if recent_docs:
        # Lancer une analyse automatique
        analysis_prompt = f"""Analyse ces {len(recent_docs)} documents récemment importés.
        
Identifie :
1. Type de document
2. Parties impliquées
3. Dates importantes
4. Points juridiques clés
5. Risques identifiés

Documents :
{chr(10).join([f"- {doc.title}: {doc.content[:500]}..." for doc in recent_docs[:5]])}"""
        
        llm_manager = MultiLLMManager()
        if llm_manager.clients:
            provider = list(llm_manager.clients.keys())[0]
            response = llm_manager.query_single_llm(
                provider,
                analysis_prompt,
                "Tu es un expert en analyse documentaire juridique."
            )
            
            if response['success']:
                st.session_state.import_analysis = response['response']
                st.success("✅ Analyse des imports terminée")

# === FONCTIONS DE TRAITEMENT POUR EXPORT ===

def process_export_request(query: str, analysis: dict):
    """Traite une demande d'export"""
    
    format = analysis['details'].get('format', 'docx')
    
    # Déterminer ce qu'il faut exporter
    content_to_export = None
    filename_base = "export"
    
    # Priorité : rédaction > analyse > recherche > sélection
    if st.session_state.get('redaction_result'):
        content_to_export = st.session_state.redaction_result['document']
        filename_base = f"{st.session_state.redaction_result['type']}"
    
    elif st.session_state.get('timeline_result'):
        content_to_export = export_timeline_content(st.session_state.timeline_result)
        filename_base = "chronologie"
    
    elif st.session_state.get('mapping_result'):
        content_to_export = export_mapping_content(st.session_state.mapping_result)
        filename_base = "cartographie"
    
    elif st.session_state.get('comparison_result'):
        content_to_export = export_comparison_content(st.session_state.comparison_result)
        filename_base = "comparaison"
    
    elif st.session_state.get('ai_analysis_results'):
        content_to_export = format_ai_analysis_for_export(st.session_state.ai_analysis_results)
        filename_base = "analyse"
    
    elif st.session_state.get('search_results'):
        content_to_export = format_search_results_for_export(st.session_state.search_results)
        filename_base = "recherche"
    
    else:
        st.warning("⚠️ Aucun contenu à exporter")
        return
    
    # Exporter selon le format
    export_functions = {
        'docx': export_to_docx,
        'pdf': export_to_pdf,
        'txt': export_to_txt,
        'html': export_to_html,
        'xlsx': export_to_xlsx
    }
    
    export_func = export_functions.get(format, export_to_txt)
    
    try:
        exported_data = export_func(content_to_export, analysis)
        
        # Créer le bouton de téléchargement
        st.download_button(
            f"💾 Télécharger {format.upper()}",
            exported_data,
            f"{filename_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}",
            get_mime_type(format),
            key=f"download_export_{format}"
        )
        
        st.success(f"✅ Export {format.upper()} prêt")
        
    except Exception as e:
        st.error(f"❌ Erreur export: {str(e)}")

def export_timeline_content(timeline_result: dict) -> str:
    """Exporte le contenu d'une chronologie"""
    content = f"CHRONOLOGIE - {timeline_result['type'].upper()}\n"
    content += f"Générée le {timeline_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Basée sur {timeline_result['document_count']} documents\n\n"
    
    for event in timeline_result['events']:
        content += f"{event.get('date', 'N/A')} - {event.get('description', '')}\n"
        if 'actors' in event:
            content += f"   Acteurs: {', '.join(event['actors'])}\n"
        if 'source' in event:
            content += f"   Source: {event['source']}\n"
        content += "\n"
    
    return content

def export_mapping_content(mapping_result: dict) -> str:
    """Exporte le contenu d'une cartographie"""
    content = f"CARTOGRAPHIE DES RELATIONS - {mapping_result['type'].upper()}\n"
    content += f"Générée le {mapping_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n\n"
    
    content += f"STATISTIQUES:\n"
    content += f"- Entités: {mapping_result['analysis']['node_count']}\n"
    content += f"- Relations: {mapping_result['analysis']['edge_count']}\n"
    content += f"- Densité: {mapping_result['analysis']['density']:.2%}\n\n"
    
    content += "ACTEURS PRINCIPAUX:\n"
    for i, player in enumerate(mapping_result['analysis']['key_players'], 1):
        content += f"{i}. {player}\n"
    
    content += "\nENTITÉS:\n"
    for entity in mapping_result['entities']:
        content += f"- {entity['name']} ({entity.get('type', 'N/A')})\n"
    
    content += "\nRELATIONS:\n"
    for relation in mapping_result['relations']:
        content += f"- {relation['source']} → {relation['target']}: {relation.get('type', 'N/A')}\n"
    
    return content

def export_comparison_content(comparison_result: dict) -> str:
    """Exporte le contenu d'une comparaison"""
    content = f"COMPARAISON - {comparison_result['type'].upper()}\n"
    content += f"Générée le {comparison_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Documents comparés: {comparison_result['document_count']}\n\n"
    
    comparison = comparison_result['comparison']
    
    content += "CONVERGENCES:\n"
    for conv in comparison.get('convergences', []):
        content += f"- {conv['point']}\n  {conv['details']}\n\n"
    
    content += "DIVERGENCES:\n"
    for div in comparison.get('divergences', []):
        content += f"- {div['point']}\n  {div['details']}\n\n"
    
    if comparison.get('analysis'):
        content += f"ANALYSE:\n{comparison['analysis']}\n"
    
    return content

def export_to_docx(content: str, analysis: dict) -> bytes:
    """Exporte vers Word avec mise en forme"""
    if not DOCX_AVAILABLE:
        # Fallback en texte brut
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Styles personnalisés
        styles = doc.styles
        
        # En-tête
        if st.session_state.get('export_metadata', True):
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Titre principal
        title = doc.add_heading(analysis.get('document_type', 'Document').upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenu avec mise en forme
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            # Détecter les niveaux de titre
            if line.strip().isupper() and len(line.strip()) > 5:
                doc.add_heading(line.strip(), 1)
            elif line.strip().startswith('###'):
                doc.add_heading(line.strip().replace('#', '').strip(), 3)
            elif line.strip().startswith('##'):
                doc.add_heading(line.strip().replace('#', '').strip(), 2)
            elif line.strip().startswith('#'):
                doc.add_heading(line.strip().replace('#', '').strip(), 1)
            elif line.strip().startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
                p = doc.add_paragraph()
                p.add_run(line.strip()).bold = True
                p.style.font.size = Pt(14)
            elif line.strip().startswith(('A.', 'B.', 'C.', 'D.')):
                p = doc.add_paragraph(line.strip(), style='List Number')
                p.style.font.size = Pt(12)
            elif line.strip().startswith('-'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(line.strip())
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
        
        # Table des matières si demandée
        if st.session_state.get('docx_toc', True):
            # Insérer au début
            doc.paragraphs[0].insert_paragraph_before("TABLE DES MATIÈRES", style='Heading 1')
            # Note: La vraie TOC nécessite des manipulations XML complexes
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création DOCX: {e}")
        return content.encode('utf-8')

def export_to_pdf(content: str, analysis: dict) -> bytes:
    """Exporte vers PDF"""
    # Nécessiterait reportlab ou weasyprint
    st.warning("⚠️ Export PDF nécessite l'installation de bibliothèques supplémentaires")
    
    # Pour l'instant, retourner le contenu texte
    return content.encode('utf-8')

def export_to_txt(content: str, analysis: dict) -> bytes:
    """Exporte en texte brut"""
    return content.encode('utf-8')

def export_to_html(content: str, analysis: dict) -> bytes:
    """Exporte vers HTML avec style"""
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <title>{analysis.get('document_type', 'Document').title()}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; margin: 40px; }}
        h1 {{ text-align: center; color: #2c3e50; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        p {{ text-align: justify; line-height: 1.6; }}
        .metadata {{ text-align: right; color: #95a5a6; font-size: 0.9em; }}
        .section {{ margin: 20px 0; }}
        ul, ol {{ margin-left: 20px; }}
    </style>
</head>
<body>
    <div class="metadata">Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</div>
    <h1>{analysis.get('document_type', 'Document').upper()}</h1>
    <div class="content">
"""
    
    # Convertir le contenu en HTML
    lines = content.split('\n')
    in_list = False
    
    for line in lines:
        if not line.strip():
            if in_list:
                html += "</ul>"
                in_list = False
            continue
        
        if line.strip().isupper() and len(line.strip()) > 5:
            html += f"<h2>{line.strip()}</h2>\n"
        elif line.strip().startswith('-'):
            if not in_list:
                html += "<ul>\n"
                in_list = True
            html += f"<li>{line.strip()[1:].strip()}</li>\n"
        else:
            if in_list:
                html += "</ul>\n"
                in_list = False
            html += f"<p>{line.strip()}</p>\n"
    
    html += """
    </div>
</body>
</html>"""
    
    return html.encode('utf-8')

def export_to_xlsx(content: str, analysis: dict) -> bytes:
    """Exporte vers Excel"""
    if not PANDAS_AVAILABLE:
        st.error("❌ pandas requis pour l'export Excel")
        return content.encode('utf-8')
    
    try:
        import io
        
        # Créer un DataFrame selon le type de contenu
        if 'timeline_result' in st.session_state:
            df = pd.DataFrame(st.session_state.timeline_result['events'])
        elif 'mapping_result' in st.session_state:
            df_entities = pd.DataFrame(st.session_state.mapping_result['entities'])
            df_relations = pd.DataFrame(st.session_state.mapping_result['relations'])
        else:
            # Contenu générique
            lines = content.split('\n')
            df = pd.DataFrame({'Contenu': lines})
        
        # Exporter vers Excel
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            if 'mapping_result' in st.session_state:
                df_entities.to_excel(writer, sheet_name='Entités', index=False)
                df_relations.to_excel(writer, sheet_name='Relations', index=False)
            else:
                df.to_excel(writer, sheet_name='Données', index=False)
            
            # Ajouter une feuille de métadonnées
            metadata = pd.DataFrame({
                'Propriété': ['Type', 'Date génération', 'Source'],
                'Valeur': [
                    analysis.get('document_type', 'Export'),
                    datetime.now().strftime('%d/%m/%Y %H:%M'),
                    analysis.get('reference', 'N/A')
                ]
            })
            metadata.to_excel(writer, sheet_name='Métadonnées', index=False)
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur export Excel: {e}")
        return content.encode('utf-8')

def get_mime_type(format: str) -> str:
    """Retourne le type MIME pour un format"""
    mime_types = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'html': 'text/html',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    }
    return mime_types.get(format, 'application/octet-stream')

# === FONCTIONS DE TRAITEMENT POUR EMAIL ===

def process_email_request(query: str, analysis: dict):
    """Traite une demande d'envoi par email"""
    
    recipients = analysis.get('recipients', [])
    
    if not recipients:
        st.error("❌ Aucun destinataire spécifié")
        return
    
    # Préparer le contenu
    content = prepare_email_content(analysis)
    
    if not content:
        st.warning("⚠️ Aucun contenu à envoyer")
        return
    
    # Configuration email
    email_config = {
        'to': st.session_state.get('email_to', ', '.join(recipients)),
        'cc': st.session_state.get('email_cc', ''),
        'subject': st.session_state.get('email_subject', 'Document juridique'),
        'body': content['body'],
        'attachments': []
    }
    
    # Préparer les pièces jointes
    if st.session_state.get('email_attach_current', True):
        attachment_format = st.session_state.get('email_attachment_format', 'pdf')
        attachment_data = export_current_content(attachment_format)
        
        if attachment_data:
            email_config['attachments'].append({
                'filename': f"document.{attachment_format}",
                'data': attachment_data,
                'mime_type': get_mime_type(attachment_format)
            })
    
    # Afficher l'aperçu
    show_email_preview(email_config)
    
    # Bouton d'envoi
    if st.button("📧 Envoyer l'email", key="send_email_button"):
        if send_email(email_config):
            st.success("✅ Email envoyé avec succès")
            st.session_state.email_sent = True
        else:
            st.error("❌ Erreur lors de l'envoi")

def prepare_email_content(analysis: dict) -> dict:
    """Prépare le contenu de l'email"""
    content = {'body': '', 'attachments': []}
    
    # Corps de l'email selon le contexte
    if st.session_state.get('redaction_result'):
        result = st.session_state.redaction_result
        content['body'] = f"""Bonjour,

Veuillez trouver ci-joint le document {result['type']} demandé.

Ce document a été généré automatiquement à partir de l'analyse de votre dossier.

Cordialement,
[Votre nom]"""
        
    else:
        content['body'] = """Bonjour,

Veuillez trouver ci-joint le document demandé.

Cordialement,
[Votre nom]"""
    
    return content

def show_email_preview(email_config: dict):
    """Affiche un aperçu de l'email"""
    with st.expander("📧 Aperçu de l'email", expanded=True):
        st.text_input("À:", value=email_config['to'], disabled=True)
        st.text_input("Cc:", value=email_config['cc'], disabled=True)
        st.text_input("Objet:", value=email_config['subject'], disabled=True)
        st.text_area("Corps:", value=email_config['body'], height=200, disabled=True)
        
        if email_config['attachments']:
            st.write("📎 Pièces jointes:")
            for att in email_config['attachments']:
                st.write(f"- {att['filename']}")

def send_email(email_config: dict) -> bool:
    """Envoie l'email (fonction simplifiée)"""
    try:
        # Configuration SMTP (à adapter)
        smtp_config = {
            'server': st.secrets.get('smtp_server', 'smtp.gmail.com'),
            'port': st.secrets.get('smtp_port', 587),
            'username': st.secrets.get('smtp_username', ''),
            'password': st.secrets.get('smtp_password', '')
        }
        
        if not smtp_config['username']:
            st.error("❌ Configuration email manquante")
            return False
        
        # Créer le message
        msg = MIMEMultipart()
        msg['From'] = smtp_config['username']
        msg['To'] = email_config['to']
        msg['Cc'] = email_config['cc']
        msg['Subject'] = email_config['subject']
        
        # Corps
        msg.attach(MIMEText(email_config['body'], 'plain'))
        
        # Pièces jointes
        for att in email_config['attachments']:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(att['data'])
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {att["filename"]}'
            )
            msg.attach(part)
        
        # Envoyer
        with smtplib.SMTP(smtp_config['server'], smtp_config['port']) as server:
            server.starttls()
            server.login(smtp_config['username'], smtp_config['password'])
            
            recipients = [email_config['to']]
            if email_config['cc']:
                recipients.extend(email_config['cc'].split(','))
            
            server.send_message(msg)
        
        return True
        
    except Exception as e:
        st.error(f"Erreur envoi email: {e}")
        return False

def export_current_content(format: str) -> bytes:
    """Exporte le contenu actuel dans le format demandé"""
    # Réutiliser les fonctions d'export
    analysis = st.session_state.get('current_analysis', {})
    
    if st.session_state.get('redaction_result'):
        content = st.session_state.redaction_result['document']
    else:
        content = "Aucun contenu disponible"
    
    export_functions = {
        'pdf': export_to_pdf,
        'docx': export_to_docx,
        'txt': export_to_txt
    }
    
    export_func = export_functions.get(format, export_to_txt)
    return export_func(content, analysis)

# === FONCTIONS DE TRAITEMENT POUR SÉLECTION DE PIÈCES ===

def process_piece_selection_request(query: str, analysis: dict):
    """Traite une demande de sélection de pièces"""
    
    # Interface de sélection
    st.markdown("### 📋 Sélection de pièces")
    
    # Collecter les documents disponibles
    available_docs = collect_available_documents(analysis)
    
    if not available_docs:
        st.warning("⚠️ Aucun document disponible")
        return
    
    # Grouper par catégorie
    categories = group_documents_by_category(available_docs)
    
    # Interface de sélection par catégorie
    selected_pieces = []
    
    for category, docs in categories.items():
        with st.expander(f"📁 {category} ({len(docs)} documents)", expanded=True):
            select_all = st.checkbox(f"Tout sélectionner - {category}", key=f"select_all_{category}")
            
            for doc in docs:
                is_selected = st.checkbox(
                    f"📄 {doc['title']}",
                    value=select_all,
                    key=f"select_doc_{doc['id']}",
                    help=f"Source: {doc.get('source', 'N/A')}"
                )
                
                if is_selected:
                    selected_pieces.append(PieceSelectionnee(
                        numero=len(selected_pieces) + 1,
                        titre=doc['title'],
                        description=doc.get('description', ''),
                        categorie=category,
                        date=doc.get('date'),
                        source=doc.get('source', ''),
                        pertinence=calculate_piece_relevance(doc, analysis)
                    ))
    
    # Sauvegarder la sélection
    st.session_state.selected_pieces = selected_pieces
    
    # Actions sur la sélection
    if selected_pieces:
        st.success(f"✅ {len(selected_pieces)} pièces sélectionnées")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📊 Créer bordereau", key="create_bordereau_from_selection"):
                process_bordereau_request(query, analysis)
        
        with col2:
            if st.button("📄 Synthétiser", key="synthesize_selection"):
                synthesize_selected_pieces(selected_pieces)
        
        with col3:
            if st.button("📤 Exporter liste", key="export_piece_list"):
                export_piece_list(selected_pieces)

def collect_available_documents(analysis: dict) -> list:
    """Collecte tous les documents disponibles"""
    documents = []
    
    # Documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        documents.append({
            'id': doc_id,
            'title': doc.title,
            'content': doc.content,
            'source': doc.source,
            'metadata': doc.metadata
        })
    
    # Si référence spécifique
    if analysis.get('reference'):
        ref_docs = search_by_reference(f"@{analysis['reference']}")
        documents.extend(ref_docs)
    
    return documents

def group_documents_by_category(documents: list) -> dict:
    """Groupe les documents par catégorie"""
    categories = defaultdict(list)
    
    for doc in documents:
        # Déterminer la catégorie
        category = determine_document_category(doc)
        categories[category].append(doc)
    
    return dict(categories)

def determine_document_category(doc: dict) -> str:
    """Détermine la catégorie d'un document"""
    title_lower = doc.get('title', '').lower()
    content_lower = doc.get('content', '')[:500].lower()
    
    # Patterns de catégories
    category_patterns = {
        'Procédure': ['plainte', 'procès-verbal', 'audition', 'perquisition', 'garde à vue'],
        'Expertise': ['expertise', 'expert', 'rapport technique', 'analyse'],
        'Comptabilité': ['bilan', 'compte', 'comptable', 'facture', 'devis'],
        'Contrats': ['contrat', 'convention', 'accord', 'avenant'],
        'Correspondance': ['courrier', 'email', 'lettre', 'mail'],
        'Pièces d\'identité': ['carte identité', 'passeport', 'kbis', 'statuts'],
        'Bancaire': ['relevé', 'virement', 'compte bancaire', 'rib']
    }
    
    for category, keywords in category_patterns.items():
        if any(kw in title_lower or kw in content_lower for kw in keywords):
            return category
    
    return 'Autres'

def calculate_piece_relevance(doc: dict, analysis: dict) -> float:
    """Calcule la pertinence d'une pièce"""
    score = 0.5
    
    # Si le document contient des mots-clés de l'analyse
    if analysis.get('subject_matter'):
        if analysis['subject_matter'] in doc.get('content', '').lower():
            score += 0.3
    
    # Si référence dans le titre
    if analysis.get('reference'):
        if analysis['reference'] in doc.get('title', '').lower():
            score += 0.2
    
    return min(score, 1.0)

# === FONCTIONS DE TRAITEMENT POUR BORDEREAU ===

def process_bordereau_request(query: str, analysis: dict):
    """Traite une demande de création de bordereau"""
    
    pieces = st.session_state.get('selected_pieces', [])
    
    if not pieces:
        st.warning("⚠️ Aucune pièce sélectionnée pour le bordereau")
        return
    
    # Créer le bordereau
    bordereau = create_bordereau(pieces, analysis)
    
    # Afficher le bordereau
    st.markdown("### 📊 Bordereau de communication de pièces")
    
    # En-tête
    st.text_area(
        "En-tête du bordereau",
        value=bordereau['header'],
        height=150,
        key="bordereau_header"
    )
    
    # Table des pièces
    if PANDAS_AVAILABLE:
        df = pd.DataFrame([
            {
                'N°': p.numero,
                'Titre': p.titre,
                'Description': p.description,
                'Catégorie': p.categorie,
                'Date': p.date.strftime('%d/%m/%Y') if p.date else 'N/A'
            }
            for p in pieces
        ])
        
        st.dataframe(df, use_container_width=True)
    else:
        # Affichage sans pandas
        for piece in pieces:
            st.write(f"**{piece.numero}.** {piece.titre}")
            if piece.description:
                st.caption(piece.description)
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.download_button(
            "💾 Télécharger bordereau",
            create_bordereau_document(bordereau),
            f"bordereau_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_bordereau"
        ):
            st.success("✅ Bordereau téléchargé")
    
    with col2:
        if st.button("✏️ Modifier sélection", key="modify_bordereau_selection"):
            st.session_state.show_piece_selection = True
            st.rerun()
    
    with col3:
        if st.button("📧 Envoyer bordereau", key="send_bordereau"):
            st.session_state.universal_query = f"envoyer bordereau @{analysis.get('reference', 'dossier')}"
            st.rerun()
    
    # Stocker le bordereau
    st.session_state.current_bordereau = bordereau

def create_bordereau(pieces: list, analysis: dict) -> dict:
    """Crée un bordereau structuré"""
    
    bordereau = {
        'header': f"""BORDEREAU DE COMMUNICATION DE PIÈCES

AFFAIRE : {analysis.get('reference', 'N/A').upper()}
DATE : {datetime.now().strftime('%d/%m/%Y')}
NOMBRE DE PIÈCES : {len(pieces)}

POUR : [À compléter]
CONTRE : [À compléter]

PIÈCES COMMUNIQUÉES :""",
        'pieces': pieces,
        'footer': """Je certifie que les pièces communiquées sont conformes aux originaux en ma possession.

Fait à [Ville], le {datetime.now().strftime('%d/%m/%Y')}

[Signature]""",
        'metadata': {
            'created_at': datetime.now(),
            'piece_count': len(pieces),
            'categories': list(set(p.categorie for p in pieces)),
            'reference': analysis.get('reference')
        }
    }
    
    return bordereau

def create_bordereau_document(bordereau: dict) -> bytes:
    """Crée le document Word du bordereau"""
    if not DOCX_AVAILABLE:
        # Fallback texte
        content = bordereau['header'] + '\n\n'
        
        for piece in bordereau['pieces']:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            content += f"   Catégorie: {piece.categorie}\n"
            if piece.date:
                content += f"   Date: {piece.date.strftime('%d/%m/%Y')}\n"
            content += "\n"
        
        content += bordereau['footer']
        
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # En-tête
        for line in bordereau['header'].split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                if 'BORDEREAU' in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(16)
        
        # Table des pièces
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # En-têtes de colonnes
        headers = ['N°', 'Titre', 'Description', 'Catégorie', 'Date']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Lignes de pièces
        for piece in bordereau['pieces']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(piece.numero)
            row_cells[1].text = piece.titre
            row_cells[2].text = piece.description or ''
            row_cells[3].text = piece.categorie
            row_cells[4].text = piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A'
        
        # Pied de page
        doc.add_paragraph()
        for line in bordereau['footer'].split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création bordereau: {e}")
        return create_bordereau_document.__wrapped__(bordereau)  # Fallback

# === FONCTIONS DE TRAITEMENT POUR SYNTHÈSE ===

def process_synthesis_request(query: str, analysis: dict):
    """Traite une demande de synthèse"""
    
    # Déterminer la source de la synthèse
    if st.session_state.get('selected_pieces'):
        content_to_synthesize = synthesize_selected_pieces(st.session_state.selected_pieces)
    
    elif analysis.get('reference'):
        docs = search_by_reference(f"@{analysis['reference']}")
        content_to_synthesize = synthesize_documents(docs)
    
    else:
        st.warning("⚠️ Aucun contenu à synthétiser")
        return
    
    # Stocker le résultat
    st.session_state.synthesis_result = content_to_synthesize

def synthesize_selected_pieces(pieces: list) -> dict:
    """Synthétise les pièces sélectionnées"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le contexte
    context = "PIÈCES À SYNTHÉTISER:\n\n"
    
    for piece in pieces[:20]:  # Limiter
        context += f"Pièce {piece.numero}: {piece.titre}\n"
        context += f"Catégorie: {piece.categorie}\n"
        if piece.description:
            context += f"Description: {piece.description}\n"
        context += "\n"
    
    # Prompt de synthèse
    synthesis_prompt = f"""{context}

Crée une synthèse structurée de ces pièces.

La synthèse doit inclure:
1. Vue d'ensemble des pièces
2. Points clés par catégorie
3. Chronologie si applicable
4. Liens et relations entre pièces
5. Points d'attention juridiques
6. Recommandations

Format professionnel avec titres et sous-sections."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            synthesis_prompt,
            "Tu es un expert en analyse et synthèse de documents juridiques."
        )
        
        if response['success']:
            return {
                'content': response['response'],
                'piece_count': len(pieces),
                'categories': list(set(p.categorie for p in pieces)),
                'timestamp': datetime.now()
            }
        else:
            return {'error': 'Échec de la synthèse'}
            
    except Exception as e:
        return {'error': f'Erreur synthèse: {str(e)}'}

def synthesize_documents(documents: list) -> dict:
    """Synthétise une liste de documents"""
    # Convertir en pièces pour réutiliser la fonction
    pieces = []
    
    for i, doc in enumerate(documents):
        pieces.append(PieceSelectionnee(
            numero=i + 1,
            titre=doc.get('title', 'Sans titre'),
            description=doc.get('content', '')[:200] + '...' if doc.get('content') else '',
            categorie=determine_document_category(doc),
            source=doc.get('source', '')
        ))
    
    return synthesize_selected_pieces(pieces)

# === FONCTIONS DE TRAITEMENT POUR TEMPLATES ===

def process_template_request(query: str, analysis: dict):
    """Traite une demande liée aux templates"""
    
    action = analysis['details'].get('action', 'apply')
    
    if action == 'create':
        create_new_template()
    
    elif action == 'edit':
        edit_existing_template()
    
    else:  # apply
        apply_template()

def create_new_template():
    """Crée un nouveau template"""
    
    template_name = st.session_state.get('new_template_name', '')
    
    if not template_name:
        st.warning("⚠️ Nom du template requis")
        return
    
    # Base du template
    base_template = st.session_state.get('base_template', 'Vide')
    
    if base_template == 'Vide':
        template_content = {
            'name': template_name,
            'structure': [],
            'style': 'formel',
            'category': st.session_state.get('template_category', 'Autre')
        }
    else:
        # Copier depuis un template existant
        template_content = DOCUMENT_TEMPLATES.get(base_template, {}).copy()
        template_content['name'] = template_name
    
    # Éditeur de structure
    st.markdown("### 📝 Structure du template")
    
    structure = st.text_area(
        "Sections (une par ligne)",
        value='\n'.join(template_content.get('structure', [])),
        height=300,
        key="template_structure_editor"
    )
    
    template_content['structure'] = [s.strip() for s in structure.split('\n') if s.strip()]
    
    # Sauvegarder
    if st.button("💾 Sauvegarder le template", key="save_new_template"):
        if 'saved_templates' not in st.session_state:
            st.session_state.saved_templates = {}
        
        st.session_state.saved_templates[clean_key(template_name)] = template_content
        st.success(f"✅ Template '{template_name}' sauvegardé")
        
        # Optionnel : sauvegarder dans un fichier
        save_templates_to_file()

def edit_existing_template():
    """Édite un template existant"""
    
    template_to_edit = st.session_state.get('template_to_edit')
    
    if not template_to_edit:
        st.warning("⚠️ Sélectionnez un template à modifier")
        return
    
    # Charger le template
    if template_to_edit in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[template_to_edit].copy()
        is_builtin = True
    else:
        template = st.session_state.saved_templates.get(template_to_edit, {})
        is_builtin = False
    
    if not template:
        st.error("❌ Template introuvable")
        return
    
    # Éditeur
    st.markdown(f"### ✏️ Édition du template '{template.get('name', template_to_edit)}'")
    
    if is_builtin:
        st.info("ℹ️ Template intégré - Les modifications seront sauvegardées comme nouveau template")
    
    # Nom
    new_name = st.text_input(
        "Nom du template",
        value=template.get('name', ''),
        key="edit_template_name"
    )
    
    # Structure
    structure = st.text_area(
        "Structure",
        value='\n'.join(template.get('structure', [])),
        height=300,
        key="edit_template_structure"
    )
    
    # Style
    style = st.selectbox(
        "Style par défaut",
        list(REDACTION_STYLES.keys()),
        index=list(REDACTION_STYLES.keys()).index(template.get('style', 'formel')),
        format_func=lambda x: REDACTION_STYLES[x]['name'],
        key="edit_template_style"
    )
    
    # Sauvegarder les modifications
    if st.button("💾 Sauvegarder les modifications", key="save_template_edits"):
        updated_template = {
            'name': new_name,
            'structure': [s.strip() for s in structure.split('\n') if s.strip()],
            'style': style,
            'category': template.get('category', 'Autre')
        }
        
        if is_builtin:
            # Sauvegarder comme nouveau
            st.session_state.saved_templates[clean_key(new_name)] = updated_template
            st.success(f"✅ Nouveau template '{new_name}' créé")
        else:
            # Mettre à jour l'existant
            st.session_state.saved_templates[template_to_edit] = updated_template
            st.success(f"✅ Template '{new_name}' mis à jour")

def apply_template():
    """Applique un template sélectionné"""
    
    selected_template = st.session_state.get('selected_template')
    
    if not selected_template:
        st.warning("⚠️ Sélectionnez un template à appliquer")
        return
    
    # Charger le template
    if selected_template in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[selected_template]
    else:
        template = st.session_state.saved_templates.get(selected_template, {})
    
    if not template:
        st.error("❌ Template introuvable")
        return
    
    # Créer une requête de rédaction avec le template
    st.session_state.universal_query = f"rédiger {template.get('name', 'document')} avec template {selected_template}"
    
    # Définir le style
    st.session_state.redaction_style = template.get('style', 'formel')
    
    # Déclencher la rédaction
    st.info(f"✅ Template '{template.get('name')}' appliqué - Lancez la rédaction")

def save_templates_to_file():
    """Sauvegarde les templates dans un fichier"""
    try:
        import json
        
        templates_data = {
            'builtin': DOCUMENT_TEMPLATES,
            'custom': st.session_state.get('saved_templates', {})
        }
        
        # Créer un fichier téléchargeable
        json_str = json.dumps(templates_data, indent=2, ensure_ascii=False)
        
        st.download_button(
            "💾 Exporter tous les templates",
            json_str,
            f"templates_{datetime.now().strftime('%Y%m%d')}.json",
            "application/json",
            key="export_templates"
        )
        
    except Exception as e:
        st.error(f"Erreur sauvegarde templates: {e}")

# === FONCTIONS DE TRAITEMENT POUR JURISPRUDENCE ===

def process_jurisprudence_request(query: str, analysis: dict):
    """Traite une demande de recherche de jurisprudence"""
    
    # Utiliser le gestionnaire de recherche juridique
    st.session_state.jurisprudence_search_active = True
    
    # Rediriger vers l'onglet jurisprudence
    st.info("⚖️ Recherche de jurisprudence activée - Voir l'onglet Jurisprudence")

# === FONCTIONS DE TRAITEMENT POUR ANALYSE ===

def process_analysis_request(query: str, analysis: dict):
    """Traite une demande d'analyse IA"""
    
    # Collecter les documents pertinents
    documents = []
    
    if analysis.get('reference'):
        documents = search_by_reference(f"@{analysis['reference']}")
    else:
        # Recherche générale
        documents = perform_search(query)
    
    if not documents:
        st.warning("⚠️ Aucun document trouvé pour l'analyse")
        return
    
    # Déterminer le type d'analyse
    analysis_type = detect_analysis_type(query)
    
    # Lancer l'analyse appropriée
    with st.spinner("🤖 Analyse en cours..."):
        if analysis_type == 'risks':
            results = analyze_legal_risks(documents, query)
        elif analysis_type == 'compliance':
            results = analyze_compliance(documents, query)
        elif analysis_type == 'strategy':
            results = analyze_strategy(documents, query)
        else:
            results = perform_general_analysis(documents, query)
    
    # Stocker les résultats
    st.session_state.ai_analysis_results = results

def detect_analysis_type(query: str) -> str:
    """Détecte le type d'analyse demandé"""
    query_lower = query.lower()
    
    if any(word in query_lower for word in ['risque', 'danger', 'exposition', 'vulnérabilité']):
        return 'risks'
    elif any(word in query_lower for word in ['conformité', 'compliance', 'réglementation']):
        return 'compliance'
    elif any(word in query_lower for word in ['stratégie', 'défense', 'approche', 'tactique']):
        return 'strategy'
    else:
        return 'general'

def analyze_legal_risks(documents: list, query: str) -> dict:
    """Analyse les risques juridiques"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le prompt d'analyse
    risk_prompt = f"""Analyse les risques juridiques dans ces documents.

DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}

QUESTION: {query}

Identifie et évalue:
1. RISQUES PÉNAUX
   - Infractions potentielles
   - Éléments constitutifs
   - Niveau de risque (faible/moyen/élevé)
   - Sanctions encourues

2. RISQUES CIVILS
   - Responsabilités contractuelles
   - Préjudices potentiels
   - Montants estimés

3. RISQUES RÉPUTATIONNELS
   - Impact médiatique
   - Conséquences business

4. RECOMMANDATIONS
   - Actions préventives
   - Stratégies de mitigation
   - Priorités

Format structuré avec évaluation précise."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            risk_prompt,
            "Tu es un expert en analyse de risques juridiques et compliance."
        )
        
        if response['success']:
            return {
                'type': 'risk_analysis',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

def perform_general_analysis(documents: list, query: str) -> dict:
    """Analyse générale des documents"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Prompt général
    general_prompt = f"""Analyse ces documents pour répondre à la question.

DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}

QUESTION: {query}

Fournis une analyse complète et structurée."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            general_prompt,
            "Tu es un expert en analyse juridique."
        )
        
        if response['success']:
            return {
                'type': 'general_analysis',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

# === FONCTIONS DE TRAITEMENT POUR RECHERCHE ===

def process_search_request(query: str, analysis: dict):
    """Traite une demande de recherche simple"""
    
    # Recherche selon le contexte
    if analysis.get('reference'):
        results = search_by_reference(f"@{analysis['reference']}")
    else:
        results = perform_search(query)
    
    # Stocker les résultats
    st.session_state.search_results = results
    
    if not results:
        st.warning("⚠️ Aucun résultat trouvé")

def search_by_reference(reference: str) -> list:
    """Recherche par référence @"""
    results = []
    
    # Nettoyer la référence
    ref_clean = reference.replace('@', '').strip().lower()
    
    # Recherche dans les documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if ref_clean in doc.title.lower() or ref_clean in doc.source.lower():
            results.append({
                'id': doc_id,
                'title': doc.title,
                'content': doc.content,
                'source': doc.source,
                'score': 1.0 if ref_clean == doc.title.lower() else 0.8
            })
    
    # Recherche Azure si disponible
    search_manager = st.session_state.get('azure_search_manager')
    if search_manager and search_manager.is_connected():
        try:
            azure_results = search_manager.search(reference)
            results.extend(azure_results)
        except:
            pass
    
    # Trier par score
    results.sort(key=lambda x: x.get('score', 0), reverse=True)
    
    return results

def perform_search(query: str) -> list:
    """Effectue une recherche générale"""
    results = []
    
    # Recherche locale basique
    query_lower = query.lower()
    query_words = query_lower.split()
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        # Score simple basé sur les mots trouvés
        score = 0
        content_lower = doc.content.lower()
        title_lower = doc.title.lower()
        
        for word in query_words:
            if word in title_lower:
                score += 2
            if word in content_lower:
                score += 1
        
        if score > 0:
            results.append({
                'id': doc_id,
                'title': doc.title,
                'content': doc.content,
                'source': doc.source,
                'score': score / len(query_words)
            })
    
    # Recherche Azure si disponible
    search_manager = st.session_state.get('azure_search_manager')
    if search_manager and search_manager.is_connected():
        try:
            azure_results = search_manager.search(query)
            results.extend(azure_results)
        except:
            pass
    
    # Trier par score
    results.sort(key=lambda x: x.get('score', 0), reverse=True)
    
    return results[:50]  # Limiter à 50 résultats

# === FONCTIONS D'AFFICHAGE DES RÉSULTATS ===

def show_unified_results_tab():
    """Affiche tous les types de résultats dans un onglet unifié"""
    
    # Vérifier quel type de résultat afficher
    has_results = False
    
    # RÉSULTATS DE RÉDACTION (Priorité 1)
    if st.session_state.get('redaction_result'):
        show_redaction_results()
        has_results = True
    
    # RÉSULTATS DE TIMELINE (Priorité 2)
    elif st.session_state.get('timeline_result'):
        show_timeline_results()
        has_results = True
    
    # RÉSULTATS DE MAPPING (Priorité 3)
    elif st.session_state.get('mapping_result'):
        show_mapping_results()
        has_results = True
    
    # RÉSULTATS DE COMPARAISON (Priorité 4)
    elif st.session_state.get('comparison_result'):
        show_comparison_results()
        has_results = True
    
    # RÉSULTATS DE SYNTHÈSE (Priorité 5)
    elif st.session_state.get('synthesis_result'):
        show_synthesis_results()
        has_results = True
    
    # RÉSULTATS D'ANALYSE IA (Priorité 6)
    elif st.session_state.get('ai_analysis_results'):
        show_ai_analysis_results()
        has_results = True
    
    # RÉSULTATS DE RECHERCHE (Priorité 7)
    elif st.session_state.get('search_results'):
        show_search_results()
        has_results = True
    
    # Message si aucun résultat
    if not has_results:
        st.info("💡 Utilisez la barre de recherche universelle pour commencer")
        
        # Suggestions d'utilisation
        st.markdown("""
        ### 🚀 Exemples de commandes
        
        **Recherche :**
        - `contrats société XYZ`
        - `@affaire_martin documents comptables`
        
        **Analyse :**
        - `analyser les risques @dossier_pénal`
        - `identifier les infractions @affaire_corruption`
        
        **Rédaction :**
        - `rédiger conclusions défense @affaire_martin abus biens sociaux`
        - `créer plainte avec constitution partie civile escroquerie`
        
        **Visualisations :**
        - `chronologie des faits @affaire_martin`
        - `cartographie des sociétés @groupe_abc`
        - `comparer les auditions @témoins`
        
        **Gestion :**
        - `sélectionner pièces @dossier catégorie procédure`
        - `créer bordereau @pièces_sélectionnées`
        - `exporter analyse format word`
        """)

def show_redaction_results():
    """Affiche les résultats de rédaction avec toutes les fonctionnalités"""
    
    result = st.session_state.redaction_result
    
    st.markdown("### 📝 Document juridique généré")
    
    # Métadonnées
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        doc_icons = {
            'conclusions': '⚖️ Conclusions',
            'plainte': '📋 Plainte',
            'constitution_pc': '🛡️ Constitution PC',
            'courrier': '✉️ Courrier',
            'assignation': '📜 Assignation',
            'mémoire': '📚 Mémoire',
            'requête': '📄 Requête'
        }
        st.metric("Type", doc_icons.get(result['type'], '📄 Document'))
    
    with col2:
        providers_count = len([r for r in result.get('responses', []) if r.get('success')])
        st.metric("IA utilisées", providers_count)
    
    with col3:
        word_count = len(result['document'].split())
        st.metric("Mots", f"{word_count:,}")
    
    with col4:
        char_count = len(result['document'])
        st.metric("Caractères", f"{char_count:,}")
    
    # Zone d'édition principale
    st.markdown("#### 📄 Contenu du document")
    
    # Ajuster la hauteur selon le type de document
    # Les conclusions et plaintes sont généralement très longues
    height_by_type = {
        'conclusions': 800,  # Très long pour dossiers complexes
        'plainte': 700,      # Long aussi
        'constitution_pc': 700,
        'mémoire': 800,
        'assignation': 600,
        'requête': 600,
        'courrier': 400
    }
    
    text_height = height_by_type.get(result['type'], 600)
    
    edited_content = st.text_area(
        "Vous pouvez éditer le document",
        value=result['document'],
        height=text_height,
        key="edit_redaction_main",
        help="Le document a été généré pour être complet et détaillé, adapté aux dossiers complexes"
    )
    
    # Barre d'outils d'édition
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        if st.button("🔄 Régénérer", key="regenerate_main"):
            if st.session_state.get('last_universal_query'):
                process_universal_query(st.session_state.last_universal_query)
                st.rerun()
    
    with col2:
        if st.button("➕ Enrichir", key="enrich_document"):
            enrich_current_document(edited_content)
    
    with col3:
        if st.button("✂️ Synthétiser", key="summarize_document"):
            create_document_summary(edited_content)
    
    with col4:
        if st.button("🔍 Vérifier", key="verify_document"):
            verify_document_content(edited_content)
    
    with col5:
        if st.button("📊 Statistiques", key="document_stats"):
            show_document_statistics(edited_content)
    
    # Options d'export
    st.markdown("#### 💾 Export et partage")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # Export Word avec mise en forme
        docx_data = create_formatted_docx(edited_content, result['type'])
        if st.download_button(
            "📄 Word (.docx)",
            docx_data,
            f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx_main"
        ):
            st.success("✅ Document Word téléchargé")
    
    with col2:
        # Export PDF (si disponible)
        try:
            pdf_data = create_pdf_from_content(edited_content, result['type'])
            if st.download_button(
                "📑 PDF",
                pdf_data,
                f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "application/pdf",
                key="download_pdf_main"
            ):
                st.success("✅ PDF téléchargé")
        except:
            st.info("PDF nécessite des libs supplémentaires")
    
    with col3:
        # Export texte brut
        if st.download_button(
            "📝 Texte (.txt)",
            edited_content.encode('utf-8'),
            f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_txt_main"
        ):
            st.success("✅ Fichier texte téléchargé")
    
    with col4:
        # Préparer email
        if st.button("📧 Envoyer", key="prepare_email_main"):
            prepare_document_email(edited_content, result['type'])
    
    # Versions individuelles des IA (si plusieurs)
    if len(result.get('responses', [])) > 1:
        with st.expander("🤖 Versions par IA", expanded=False):
            for i, response in enumerate(result['responses']):
                if response.get('success'):
                    st.markdown(f"#### {response['provider']}")
                    
                    # Métriques par version
                    col1, col2 = st.columns([4, 1])
                    
                    with col1:
                        words = len(response['response'].split())
                        st.caption(f"📊 {words:,} mots")
                    
                    with col2:
                        if st.button(f"Utiliser", key=f"use_version_{i}"):
                            st.session_state.edit_redaction_main = response['response']
                            st.rerun()
                    
                    # Contenu
                    st.text_area(
                        f"Version {response['provider']}",
                        value=response['response'],
                        height=400,
                        key=f"version_{response['provider']}_{i}",
                        disabled=True
                    )
    
    # Jurisprudence utilisée
    if result.get('jurisprudence_used') and result.get('jurisprudence_references'):
        with st.expander("⚖️ Jurisprudence citée", expanded=False):
            for ref in result['jurisprudence_references']:
                st.markdown(f"- [{ref['title']}]({ref.get('url', '#')})")
                if ref.get('summary'):
                    st.caption(ref['summary'])

def create_formatted_docx(content: str, doc_type: str) -> bytes:
    """Crée un document Word avec mise en forme professionnelle complète"""
    
    if not DOCX_AVAILABLE:
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = docx.Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
        
        # Styles personnalisés pour documents juridiques
        styles = doc.styles
        
        # Style pour le titre principal
        title_style = styles.add_style('JuridicalTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # Style pour les sections principales (I., II., III.)
        heading1_style = styles.add_style('JuridicalHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.name = 'Times New Roman'
        heading1_style.font.size = Pt(14)
        heading1_style.font.bold = True
        heading1_style.paragraph_format.space_before = Pt(18)
        heading1_style.paragraph_format.space_after = Pt(12)
        
        # Style pour les sous-sections (A., B., C.)
        heading2_style = styles.add_style('JuridicalHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.name = 'Times New Roman'
        heading2_style.font.size = Pt(13)
        heading2_style.font.bold = True
        heading2_style.font.underline = True
        heading2_style.paragraph_format.space_before = Pt(12)
        heading2_style.paragraph_format.space_after = Pt(6)
        
        # Style pour le corps du texte
        body_style = styles['Normal']
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.first_line_indent = Inches(0.5)
        
        # Analyser et formater le contenu
        lines = content.split('\n')
        
        # Variables pour tracker l'état
        in_header = True
        current_section = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Ligne vide - ajouter un saut seulement si pas dans l'en-tête
                if not in_header:
                    doc.add_paragraph()
                continue
            
            # Détecter le type de ligne
            # Titre principal (tout en majuscules, court)
            if line.isupper() and len(line.split()) < 10 and in_header:
                p = doc.add_paragraph(line, style='JuridicalTitle')
                in_header = False
            
            # En-tête (POUR:, CONTRE:, etc.)
            elif line.startswith(('POUR', 'CONTRE', 'TRIBUNAL', 'AFFAIRE')) and ':' in line:
                p = doc.add_paragraph()
                parts = line.split(':', 1)
                run1 = p.add_run(parts[0] + ':')
                run1.bold = True
                if len(parts) > 1:
                    p.add_run(' ' + parts[1])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Sections principales (I., II., III., etc.)
            elif re.match(r'^[IVX]+\.\s+', line):
                p = doc.add_paragraph(line, style='JuridicalHeading1')
                current_section = line
                in_header = False
            
            # Sous-sections (A., B., C., etc.)
            elif re.match(r'^[A-Z]\.\s+', line):
                p = doc.add_paragraph(line, style='JuridicalHeading2')
            
            # Sous-sous-sections (1., 2., 3., etc.)
            elif re.match(r'^\d+\.\s+', line):
                p = doc.add_paragraph(line)
                p.style.font.bold = True
                p.style.paragraph_format.left_indent = Inches(0.5)
            
            # Points avec tirets
            elif line.startswith('-'):
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                p.style.paragraph_format.left_indent = Inches(0.5)
            
            # PAR CES MOTIFS
            elif line.startswith('PAR CES MOTIFS'):
                # Saut de page avant PAR CES MOTIFS
                doc.add_page_break()
                p = doc.add_paragraph(line, style='JuridicalHeading1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Formules de conclusion
            elif any(line.startswith(phrase) for phrase in [
                'PLAISE AU TRIBUNAL',
                'PLAISE À LA COUR',
                'IL EST DEMANDÉ',
                'SOUS TOUTES RÉSERVES'
            ]):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style.font.bold = True
            
            # Texte normal
            else:
                p = doc.add_paragraph(line, style='Normal')
                
                # Mettre en gras les références juridiques
                for run in p.runs:
                    text = run.text
                    # Articles de loi
                    if re.search(r'article\s+[LR]?\s*\d+', text, re.IGNORECASE):
                        run.font.bold = True
                    # Jurisprudence
                    elif re.search(r'(Cass\.|CA|CE|Cons\.\s*const\.)', text):
                        run.font.italic = True
        
        # Ajouter les métadonnées dans le pied de page
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} - {doc_type.title()}"
        footer_para.style.font.size = Pt(9)
        footer_para.style.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Numérotation des pages
        # Note: La numérotation automatique nécessite des manipulations XML complexes
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création document Word: {e}")
        # Fallback : retourner le texte brut
        return content.encode('utf-8')

def create_pdf_from_content(content: str, doc_type: str) -> bytes:
    """Crée un PDF à partir du contenu (nécessite reportlab ou weasyprint)"""
    # Pour l'instant, retourner une version texte
    # L'implémentation complète nécessiterait reportlab
    return content.encode('utf-8')

def enrich_current_document(content: str):
    """Enrichit le document actuel avec plus de détails"""
    st.info("🔄 Enrichissement en cours...")
    # Implémenter l'enrichissement via IA

def show_timeline_results():
    """Affiche les résultats de chronologie"""
    result = st.session_state.timeline_result
    
    st.markdown(f"### ⏱️ Chronologie des {result['type']}")
    
    # Métadonnées
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Événements", len(result['events']))
    with col2:
        st.metric("Documents analysés", result['document_count'])
    with col3:
        st.metric("Type", result['type'].title())
    
    # Visualisation
    if result.get('visualization'):
        st.plotly_chart(result['visualization'], use_container_width=True)
    else:
        # Affichage alternatif
        for event in result['events']:
            with st.container():
                col1, col2 = st.columns([1, 4])
                with col1:
                    st.write(f"**{event.get('date', 'N/A')}**")
                with col2:
                    st.write(event.get('description', ''))
                    if event.get('actors'):
                        st.caption(f"👥 {', '.join(event['actors'])}")

def show_mapping_results():
    """Affiche les résultats de cartographie"""
    result = st.session_state.mapping_result
    
    st.markdown(f"### 🗺️ Cartographie des relations - {result['type']}")
    
    # Statistiques
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Entités", result['analysis']['node_count'])
    with col2:
        st.metric("Relations", result['analysis']['edge_count'])
    with col3:
        st.metric("Densité", f"{result['analysis']['density']:.2%}")
    with col4:
        st.metric("Composantes", len(result['analysis']['components']))
    
    # Acteurs clés
    if result['analysis']['key_players']:
        st.markdown("#### 🎯 Acteurs principaux")
        for i, player in enumerate(result['analysis']['key_players'], 1):
            st.write(f"{i}. **{player}**")
    
    # Visualisation
    if result.get('visualization'):
        st.plotly_chart(result['visualization'], use_container_width=True)

def show_comparison_results():
    """Affiche les résultats de comparaison"""
    result = st.session_state.comparison_result
    
    st.markdown(f"### 🔄 Comparaison - {result['type']}")
    
    # Statistiques
    stats = result['comparison'].get('statistics', {})
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Documents", result['document_count'])
    with col2:
        st.metric("Convergences", len(result['comparison'].get('convergences', [])))
    with col3:
        st.metric("Divergences", len(result['comparison'].get('divergences', [])))
    with col4:
        reliability = stats.get('reliability_score', 0)
        color = "🟢" if reliability > 0.7 else "🟡" if reliability > 0.4 else "🔴"
        st.metric("Fiabilité", f"{color} {reliability:.0%}")
    
    # Visualisations
    if result.get('visualizations'):
        for viz_name, viz in result['visualizations'].items():
            if viz:
                st.plotly_chart(viz, use_container_width=True)
    
    # Détails textuels
    tabs = st.tabs(["Convergences", "Divergences", "Évolutions", "Analyse"])
    
    with tabs[0]:
        for conv in result['comparison'].get('convergences', []):
            st.write(f"**{conv['point']}**")
            st.write(conv['details'])
            st.divider()
    
    with tabs[1]:
        for div in result['comparison'].get('divergences', []):
            st.write(f"**{div['point']}**")
            st.write(div['details'])
            st.divider()
    
    with tabs[2]:
        for evo in result['comparison'].get('evolutions', []):
            st.write(f"**{evo.get('point', 'Évolution')}**")
            st.write(evo.get('details', ''))
            st.divider()
    
    with tabs[3]:
        st.write(result['comparison'].get('analysis', 'Pas d\'analyse disponible'))

def show_synthesis_results():
    """Affiche les résultats de synthèse"""
    result = st.session_state.synthesis_result
    
    if 'error' in result:
        st.error(f"❌ {result['error']}")
        return
    
    st.markdown("### 📝 Synthèse des documents")
    
    # Métadonnées
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Pièces analysées", result.get('piece_count', 0))
    with col2:
        st.metric("Catégories", len(result.get('categories', [])))
    with col3:
        st.metric("Généré le", result.get('timestamp', datetime.now()).strftime('%H:%M'))
    
    # Contenu de la synthèse
    st.markdown("#### 📄 Synthèse")
    
    synthesis_content = st.text_area(
        "Contenu de la synthèse",
        value=result.get('content', ''),
        height=600,
        key="synthesis_content_display"
    )
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.download_button(
            "💾 Télécharger",
            synthesis_content.encode('utf-8'),
            f"synthese_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_synthesis"
        ):
            st.success("✅ Synthèse téléchargée")
    
    with col2:
        if st.button("📊 Voir les pièces", key="view_synthesis_pieces"):
            st.session_state.show_piece_selection = True
    
    with col3:
        if st.button("🔄 Régénérer", key="regenerate_synthesis"):
            if st.session_state.get('selected_pieces'):
                synthesize_selected_pieces(st.session_state.selected_pieces)

def show_ai_analysis_results():
    """Affiche les résultats d'analyse IA"""
    results = st.session_state.ai_analysis_results
    
    if 'error' in results:
        st.error(f"❌ {results['error']}")
        return
    
    analysis_titles = {
        'risk_analysis': '⚠️ Analyse des risques',
        'compliance': '✅ Analyse de conformité',
        'strategy': '🎯 Analyse stratégique',
        'general_analysis': '🤖 Analyse générale'
    }
    
    st.markdown(f"### {analysis_titles.get(results.get('type'), '🤖 Analyse IA')}")
    
    # Métadonnées
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Documents analysés", results.get('document_count', 0))
    with col2:
        st.metric("Type", results.get('type', 'general').replace('_', ' ').title())
    with col3:
        st.metric("Généré le", results.get('timestamp', datetime.now()).strftime('%H:%M'))
    
    # Question originale
    if results.get('query'):
        st.info(f"**Question :** {results['query']}")
    
    # Contenu de l'analyse
    st.markdown("#### 📊 Résultats de l'analyse")
    
    analysis_content = st.text_area(
        "Analyse détaillée",
        value=results.get('content', ''),
        height=600,
        key="ai_analysis_content"
    )
    
    # Actions
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.download_button(
            "💾 Télécharger",
            analysis_content.encode('utf-8'),
            f"analyse_{results.get('type', 'general')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_analysis"
        ):
            st.success("✅ Analyse téléchargée")
    
    with col2:
        if st.button("📝 Convertir en rapport", key="convert_to_report"):
            convert_analysis_to_report(analysis_content, results)
    
    with col3:
        if st.button("📊 Visualiser", key="visualize_analysis"):
            create_analysis_visualization(results)
    
    with col4:
        if st.button("🔄 Approfondir", key="deepen_analysis"):
            deepen_current_analysis(results)

def show_search_results():
    """Affiche les résultats de recherche"""
    results = st.session_state.search_results
    
    st.markdown(f"### 🔍 Résultats de recherche ({len(results)} documents)")
    
    if not results:
        st.info("Aucun résultat trouvé")
        return
    
    # Options d'affichage
    col1, col2, col3 = st.columns(3)
    
    with col1:
        sort_by = st.selectbox(
            "Trier par",
            ["Pertinence", "Titre", "Date", "Source"],
            key="sort_search_results"
        )
    
    with col2:
        view_mode = st.radio(
            "Affichage",
            ["Compact", "Détaillé"],
            key="view_mode_search",
            horizontal=True
        )
    
    with col3:
        results_per_page = st.selectbox(
            "Résultats par page",
            [10, 20, 50, 100],
            key="results_per_page"
        )
    
    # Trier les résultats
    sorted_results = sort_search_results(results, sort_by)
    
    # Pagination
    total_pages = (len(sorted_results) - 1) // results_per_page + 1
    current_page = st.number_input(
        "Page",
        min_value=1,
        max_value=total_pages,
        value=1,
        key="search_page"
    )
    
    start_idx = (current_page - 1) * results_per_page
    end_idx = min(start_idx + results_per_page, len(sorted_results))
    
    # Afficher les résultats
    for i, result in enumerate(sorted_results[start_idx:end_idx], start=start_idx + 1):
        with st.container():
            if view_mode == "Compact":
                col1, col2, col3 = st.columns([4, 1, 1])
                
                with col1:
                    st.markdown(f"**{i}. {result.get('title', 'Sans titre')}**")
                    st.caption(f"Source: {result.get('source', 'N/A')}")
                
                with col2:
                    if result.get('score'):
                        st.metric("Score", f"{result['score']:.0%}")
                
                with col3:
                    if st.button("Voir", key=f"view_result_{i}"):
                        show_document_detail(result)
            
            else:  # Détaillé
                st.markdown(f"**{i}. {result.get('title', 'Sans titre')}**")
                st.caption(f"Source: {result.get('source', 'N/A')} | Score: {result.get('score', 0):.0%}")
                
                # Extrait du contenu
                content = result.get('content', '')
                if content:
                    st.text_area(
                        "Extrait",
                        value=content[:500] + '...' if len(content) > 500 else content,
                        height=150,
                        key=f"extract_{i}",
                        disabled=True
                    )
                
                # Actions
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if st.button("📖 Lire", key=f"read_{i}"):
                        show_document_detail(result)
                
                with col2:
                    if st.button("🤖 Analyser", key=f"analyze_{i}"):
                        analyze_single_document(result)
                
                with col3:
                    if st.button("📋 Sélectionner", key=f"select_{i}"):
                        add_to_selection(result)
            
            st.divider()
    
    # Navigation pages
    if total_pages > 1:
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col1:
            if current_page > 1:
                if st.button("⬅️ Précédent"):
                    st.session_state.search_page = current_page - 1
                    st.rerun()
        
        with col2:
            st.write(f"Page {current_page} / {total_pages}")
        
        with col3:
            if current_page < total_pages:
                if st.button("Suivant ➡️"):
                    st.session_state.search_page = current_page + 1
                    st.rerun()

def sort_search_results(results: list, sort_by: str) -> list:
    """Trie les résultats de recherche"""
    if sort_by == "Pertinence":
        return sorted(results, key=lambda x: x.get('score', 0), reverse=True)
    elif sort_by == "Titre":
        return sorted(results, key=lambda x: x.get('title', ''))
    elif sort_by == "Source":
        return sorted(results, key=lambda x: x.get('source', ''))
    else:  # Date
        # Essayer de parser les dates si disponibles
        return results

def show_document_detail(document: dict):
    """Affiche le détail d'un document"""
    st.session_state.current_document = document
    st.session_state.show_document_detail = True

def analyze_single_document(document: dict):
    """Lance l'analyse d'un seul document"""
    st.session_state.universal_query = f"analyser @{document.get('id', document.get('title', ''))}"
    st.rerun()

def add_to_selection(document: dict):
    """Ajoute un document à la sélection"""
    if 'selected_documents' not in st.session_state:
        st.session_state.selected_documents = []
    
    st.session_state.selected_documents.append(document)
    st.success("✅ Document ajouté à la sélection")

# === AUTRES FONCTIONS HELPER ===

def show_pieces_management_tab():
    """Affiche l'onglet de gestion des pièces"""
    st.markdown("### 📋 Gestion des pièces")
    
    if st.session_state.get('selected_pieces'):
        pieces = st.session_state.selected_pieces
        
        # Statistiques
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total pièces", len(pieces))
        with col2:
            categories = list(set(p.categorie for p in pieces))
            st.metric("Catégories", len(categories))
        with col3:
            st.metric("Pertinence moyenne", f"{sum(p.pertinence for p in pieces) / len(pieces):.0%}")
        
        # Actions groupées
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("📊 Créer bordereau", key="create_bordereau_tab"):
                process_bordereau_request("", {})
        
        with col2:
            if st.button("📝 Synthétiser", key="synthesize_pieces_tab"):
                synthesize_selected_pieces(pieces)
        
        with col3:
            if st.button("📤 Exporter liste", key="export_pieces_tab"):
                export_piece_list(pieces)
        
        with col4:
            if st.button("🗑️ Vider sélection", key="clear_selection_tab"):
                st.session_state.selected_pieces = []
                st.rerun()
        
        # Affichage des pièces
        for piece in pieces:
            with st.container():
                col1, col2, col3, col4 = st.columns([1, 3, 1, 1])
                
                with col1:
                    st.write(f"**{piece.numero}**")
                
                with col2:
                    st.write(piece.titre)
                    if piece.description:
                        st.caption(piece.description)
                
                with col3:
                    st.caption(piece.categorie)
                
                with col4:
                    if st.button("❌", key=f"remove_piece_{piece.numero}"):
                        pieces.remove(piece)
                        st.rerun()
            
            st.divider()
    
    else:
        st.info("Aucune pièce sélectionnée. Utilisez la commande `sélectionner pièces` dans la barre universelle.")

def export_piece_list(pieces: list):
    """Exporte la liste des pièces"""
    content = "LISTE DES PIÈCES SÉLECTIONNÉES\n"
    content += f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Nombre de pièces : {len(pieces)}\n\n"
    
    # Grouper par catégorie
    by_category = defaultdict(list)
    for piece in pieces:
        by_category[piece.categorie].append(piece)
    
    for category, cat_pieces in by_category.items():
        content += f"\n{category.upper()} ({len(cat_pieces)} pièces)\n"
        content += "-" * 50 + "\n"
        
        for piece in cat_pieces:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            if piece.date:
                content += f"   Date : {piece.date.strftime('%d/%m/%Y')}\n"
            content += "\n"
    
    # Télécharger
    st.download_button(
        "💾 Télécharger la liste",
        content.encode('utf-8'),
        f"liste_pieces_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
        "text/plain",
        key="download_piece_list_file"
    )

def show_jurisprudence_tab():
    """Affiche l'onglet jurisprudence"""
    # Utiliser le gestionnaire existant
    display_legal_search_interface()

def show_explorer_tab():
    """Affiche l'explorateur de fichiers"""
    st.markdown("### 📁 Explorateur de documents")
    
    # Sources disponibles
    sources = []
    
    if st.session_state.get('azure_documents'):
        sources.append("Documents locaux")
    
    if st.session_state.get('azure_blob_manager'):
        sources.append("Azure Blob Storage")
    
    if not sources:
        st.info("Aucune source de documents disponible")
        return
    
    selected_source = st.selectbox(
        "Source",
        sources,
        key="explorer_source"
    )
    
    if selected_source == "Documents locaux":
        show_local_documents_explorer()
    
    elif selected_source == "Azure Blob Storage":
        show_azure_blob_explorer()

def show_local_documents_explorer():
    """Affiche l'explorateur de documents locaux"""
    documents = st.session_state.get('azure_documents', {})
    
    if not documents:
        st.info("Aucun document local")
        return
    
    # Filtres
    col1, col2 = st.columns(2)
    
    with col1:
        search_filter = st.text_input(
            "Filtrer par nom",
            key="explorer_filter_name"
        )
    
    with col2:
        source_filter = st.selectbox(
            "Filtrer par source",
            ["Toutes"] + list(set(doc.source for doc in documents.values())),
            key="explorer_filter_source"
        )
    
    # Appliquer les filtres
    filtered_docs = {}
    for doc_id, doc in documents.items():
        if search_filter and search_filter.lower() not in doc.title.lower():
            continue
        if source_filter != "Toutes" and doc.source != source_filter:
            continue
        filtered_docs[doc_id] = doc
    
    # Afficher les documents
    st.write(f"**{len(filtered_docs)} documents**")
    
    for doc_id, doc in filtered_docs.items():
        with st.expander(f"📄 {doc.title}"):
            st.write(f"**ID :** {doc_id}")
            st.write(f"**Source :** {doc.source}")
            
            if doc.metadata:
                st.write("**Métadonnées :**")
                for key, value in doc.metadata.items():
                    st.write(f"- {key}: {value}")
            
            # Actions
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📖 Lire", key=f"read_doc_{doc_id}"):
                    st.session_state.current_document = {
                        'id': doc_id,
                        'title': doc.title,
                        'content': doc.content,
                        'source': doc.source
                    }
            
            with col2:
                if st.button("🤖 Analyser", key=f"analyze_doc_{doc_id}"):
                    st.session_state.universal_query = f"analyser @{doc_id}"
                    st.rerun()
            
            with col3:
                if st.button("🗑️ Supprimer", key=f"delete_doc_{doc_id}"):
                    del st.session_state.azure_documents[doc_id]
                    st.rerun()

def show_azure_blob_explorer():
    """Affiche l'explorateur Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if not blob_manager or not blob_manager.is_connected():
        st.warning("Azure Blob Storage non connecté")
        return
    
    try:
        containers = blob_manager.list_containers()
        
        if not containers:
            st.info("Aucun conteneur disponible")
            return
        
        selected_container = st.selectbox(
            "Conteneur",
            containers,
            key="blob_container"
        )
        
        # Explorer le conteneur
        current_path = st.session_state.get('blob_current_path', '')
        
        # Breadcrumb
        if current_path:
            path_parts = current_path.split('/')
            breadcrumb = "📁 " + " > ".join(path_parts)
            st.write(breadcrumb)
            
            if st.button("⬆️ Dossier parent"):
                st.session_state.blob_current_path = '/'.join(path_parts[:-1])
                st.rerun()
        
        # Lister le contenu
        items = blob_manager.list_folder_contents(selected_container, current_path)
        
        # Séparer dossiers et fichiers
        folders = [item for item in items if item['type'] == 'folder']
        files = [item for item in items if item['type'] == 'file']
        
        # Afficher les dossiers
        if folders:
            st.write("**📁 Dossiers**")
            for folder in folders:
                col1, col2 = st.columns([4, 1])
                
                with col1:
                    st.write(f"📁 {folder['name']}")
                
                with col2:
                    if st.button("Ouvrir", key=f"open_folder_{folder['name']}"):
                        st.session_state.blob_current_path = folder['path']
                        st.rerun()
        
        # Afficher les fichiers
        if files:
            st.write("**📄 Fichiers**")
            for file in files:
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    st.write(f"📄 {file['name']}")
                    st.caption(f"Taille: {file['size']} octets")
                
                with col2:
                    if st.button("💾", key=f"download_blob_{file['name']}"):
                        content = blob_manager.download_blob(selected_container, file['path'])
                        st.download_button(
                            "Télécharger",
                            content,
                            file['name'],
                            key=f"download_actual_{file['name']}"
                        )
                
                with col3:
                    if st.button("🤖", key=f"analyze_blob_{file['name']}"):
                        # Télécharger et analyser
                        content = blob_manager.download_blob(selected_container, file['path'])
                        # Stocker temporairement
                        store_locally(
                            file['name'],
                            content.decode('utf-8', errors='ignore'),
                            f"azure:{selected_container}/{file['path']}"
                        )
                        st.session_state.universal_query = f"analyser @{file['name']}"
                        st.rerun()
        
    except Exception as e:
        st.error(f"Erreur exploration Azure: {e}")

def show_configuration_tab():
    """Affiche l'onglet de configuration"""
    st.markdown("### ⚙️ Configuration")
    
    tabs = st.tabs(["IA", "Connexions", "Templates", "Préférences"])
    
    with tabs[0]:
        show_ai_configuration()
    
    with tabs[1]:
        show_connections_configuration()
    
    with tabs[2]:
        show_templates_configuration()
    
    with tabs[3]:
        show_preferences_configuration()

def show_ai_configuration():
    """Configuration des IA"""
    st.markdown("#### 🤖 Configuration des IA")
    
    llm_manager = MultiLLMManager()
    
    # État des connexions
    st.write("**État des connexions :**")
    
    for provider in LLMProvider:
        col1, col2 = st.columns([3, 1])
        
        with col1:
            is_connected = provider in llm_manager.clients
            status = "✅ Connecté" if is_connected else "❌ Non connecté"
            st.write(f"{provider.value}: {status}")
        
        with col2:
            if is_connected:
                if st.button(f"Tester", key=f"test_{provider.value}"):
                    test_llm_connection(provider)
    
    # Préférences par défaut
    st.write("**Préférences par défaut :**")
    
    default_providers = st.multiselect(
        "IA par défaut pour la rédaction",
        [p.value for p in LLMProvider if p in llm_manager.clients],
        default=[p.value for p in list(llm_manager.clients.keys())[:2]] if llm_manager.clients else [],
        key="default_redaction_providers"
    )
    
    fusion_mode = st.selectbox(
        "Mode de fusion par défaut",
        ["🎯 Fusion intelligente", "📋 Comparaison côte à côte", "🔗 Synthèse enrichie", "⚡ Meilleure version"],
        key="default_fusion_mode"
    )

def test_llm_connection(provider: LLMProvider):
    """Teste la connexion à une IA"""
    llm_manager = MultiLLMManager()
    
    with st.spinner(f"Test de {provider.value}..."):
        try:
            response = llm_manager.query_single_llm(
                provider,
                "Réponds simplement 'OK' si tu me reçois.",
                "Assistant de test"
            )
            
            if response['success']:
                st.success(f"✅ {provider.value} fonctionne correctement")
            else:
                st.error(f"❌ Erreur: {response.get('error', 'Inconnue')}")
                
        except Exception as e:
            st.error(f"❌ Erreur connexion: {str(e)}")

def show_connections_configuration():
    """Configuration des connexions externes"""
    st.markdown("#### 🔌 Connexions externes")
    
    # Azure Blob Storage
    st.write("**Azure Blob Storage**")
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        st.success("✅ Connecté")
        if st.button("🔄 Rafraîchir la connexion"):
            blob_manager.connect()
    else:
        st.warning("❌ Non connecté")
        st.info("Configurez les variables d'environnement Azure")
    
    # Azure Search
    st.write("**Azure Cognitive Search**")
    search_manager = st.session_state.get('azure_search_manager')
    
    if search_manager and search_manager.is_connected():
        st.success("✅ Connecté")
    else:
        st.warning("❌ Non connecté")
    
    # Email SMTP
    st.write("**Configuration Email**")
    
    smtp_configured = bool(st.secrets.get('smtp_username'))
    if smtp_configured:
        st.success("✅ Email configuré")
    else:
        st.warning("❌ Email non configuré")
        st.info("Ajoutez les paramètres SMTP dans les secrets")

def show_templates_configuration():
    """Configuration des templates"""
    st.markdown("#### 📄 Gestion des templates")
    
    # Templates intégrés
    st.write("**Templates intégrés :**")
    for key, template in DOCUMENT_TEMPLATES.items():
        st.write(f"- {template['name']} ({template.get('style', 'N/A')})")
    
    # Templates personnalisés
    custom_templates = st.session_state.get('saved_templates', {})
    
    if custom_templates:
        st.write(f"**Templates personnalisés ({len(custom_templates)}) :**")
        
        for key, template in custom_templates.items():
            col1, col2, col3 = st.columns([3, 1, 1])
            
            with col1:
                st.write(f"- {template.get('name', key)}")
            
            with col2:
                if st.button("✏️", key=f"edit_template_{key}"):
                    st.session_state.template_to_edit = key
                    st.session_state.universal_query = "modifier template"
            
            with col3:
                if st.button("🗑️", key=f"delete_template_{key}"):
                    del st.session_state.saved_templates[key]
                    st.rerun()
    
    # Import/Export
    st.divider()
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📤 Exporter tous les templates"):
            save_templates_to_file()
    
    with col2:
        uploaded_file = st.file_uploader(
            "📥 Importer des templates",
            type=['json'],
            key="import_templates"
        )
        
        if uploaded_file:
            try:
                import json
                templates_data = json.load(uploaded_file)
                
                if 'custom' in templates_data:
                    st.session_state.saved_templates.update(templates_data['custom'])
                    st.success(f"✅ {len(templates_data['custom'])} templates importés")
                    st.rerun()
                    
            except Exception as e:
                st.error(f"Erreur import: {e}")

def show_preferences_configuration():
    """Configuration des préférences utilisateur"""
    st.markdown("#### 🎨 Préférences")
    
    # Préférences d'affichage
    st.write("**Affichage :**")
    
    results_per_page = st.selectbox(
        "Résultats par page par défaut",
        [10, 20, 50, 100],
        index=1,
        key="pref_results_per_page"
    )
    
    default_view = st.radio(
        "Vue par défaut",
        ["Compact", "Détaillé"],
        key="pref_default_view",
        horizontal=True
    )
    
    # Préférences de rédaction
    st.write("**Rédaction :**")
    
    auto_jurisprudence = st.checkbox(
        "Recherche automatique de jurisprudence",
        value=True,
        key="pref_auto_juris"
    )
    
    create_hyperlinks = st.checkbox(
        "Créer des liens hypertextes automatiques",
        value=True,
        key="pref_hyperlinks"
    )
    
    default_doc_length = st.select_slider(
        "Longueur par défaut des documents",
        options=["Concis", "Standard", "Détaillé", "Très détaillé", "Exhaustif"],
        value="Très détaillé",
        key="pref_doc_length"
    )
    
    # Sauvegarder les préférences
    if st.button("💾 Sauvegarder les préférences"):
        preferences = {
            'results_per_page': results_per_page,
            'default_view': default_view,
            'auto_jurisprudence': auto_jurisprudence,
            'create_hyperlinks': create_hyperlinks,
            'default_doc_length': default_doc_length
        }
        
        st.session_state.user_preferences = preferences
        st.success("✅ Préférences sauvegardées")

# Fonctions helper additionnelles

def clear_universal_state():
    """Efface l'état de l'interface universelle"""
    keys_to_clear = [
        'universal_query', 'last_universal_query', 'current_analysis',
        'redaction_result', 'timeline_result', 'mapping_result',
        'comparison_result', 'synthesis_result', 'ai_analysis_results',
        'search_results', 'selected_pieces', 'import_files'
    ]
    
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    
    st.success("✅ Interface réinitialisée")
    st.rerun()

def save_current_work():
    """Sauvegarde le travail actuel"""
    work_data = {
        'timestamp': datetime.now().isoformat(),
        'query': st.session_state.get('universal_query', ''),
        'analysis': st.session_state.get('current_analysis', {}),
        'results': {}
    }
    
    # Collecter tous les résultats
    result_keys = [
        'redaction_result', 'timeline_result', 'mapping_result',
        'comparison_result', 'synthesis_result', 'ai_analysis_results',
        'search_results'
    ]
    
    for key in result_keys:
        if key in st.session_state:
            work_data['results'][key] = st.session_state[key]
    
    # Sauvegarder
    import json
    
    json_str = json.dumps(work_data, indent=2, ensure_ascii=False, default=str)
    
    st.download_button(
        "💾 Télécharger la sauvegarde",
        json_str,
        f"sauvegarde_travail_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        "application/json",
        key="download_work_save"
    )

def share_current_work():
    """Partage le travail actuel"""
    # Créer un lien de partage ou export
    st.info("🔗 Fonctionnalité de partage à implémenter")
    
    # Pour l'instant, proposer l'export
    save_current_work()

def prepare_document_email(content: str, doc_type: str):
    """Prépare l'envoi d'un document par email"""
    st.session_state.email_document = {
        'content': content,
        'type': doc_type
    }
    
    st.session_state.universal_query = f"envoyer {doc_type} par email"
    st.rerun()

def verify_document_content(content: str):
    """Vérifie le contenu d'un document"""
    # Vérifications basiques
    issues = []
    
    # Longueur
    if len(content) < 1000:
        issues.append("⚠️ Document très court (< 1000 caractères)")
    
    # Références juridiques
    if not re.search(r'article\s+\d+', content, re.IGNORECASE):
        issues.append("⚠️ Aucune référence d'article de loi")
    
    # Structure
    if not any(marker in content for marker in ['I.', 'II.', 'A.', 'B.']):
        issues.append("⚠️ Structure peu claire (pas de sections numérotées)")
    
    # Afficher le résultat
    if issues:
        st.warning("Points d'attention :")
        for issue in issues:
            st.write(issue)
    else:
        st.success("✅ Document bien structuré")

def show_document_statistics(content: str):
    """Affiche les statistiques d'un document"""
    
    # Calculs
    words = content.split()
    sentences = content.split('.')
    paragraphs = content.split('\n\n')
    
    # Références
    law_refs = len(re.findall(r'article\s+[LR]?\s*\d+', content, re.IGNORECASE))
    juris_refs = len(re.findall(r'(Cass\.|CA|CE|Cons\.\s*const\.)', content))
    
    # Affichage
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Mots", f"{len(words):,}")
        st.metric("Phrases", f"{len(sentences):,}")
    
    with col2:
        st.metric("Paragraphes", len(paragraphs))
        st.metric("Mots/phrase", f"{len(words) / max(len(sentences), 1):.1f}")
    
    with col3:
        st.metric("Articles cités", law_refs)
        st.metric("Jurisprudences", juris_refs)
    
    with col4:
        # Complexité (basique)
        avg_word_length = sum(len(w) for w in words) / max(len(words), 1)
        complexity = "Simple" if avg_word_length < 5 else "Moyen" if avg_word_length < 7 else "Complexe"
        st.metric("Complexité", complexity)
        st.metric("Longueur moy.", f"{avg_word_length:.1f} car/mot")

def create_document_summary(content: str):
    """Crée un résumé du document"""
    llm_manager = MultiLLMManager()
    
    if not llm_manager.clients:
        st.error("Aucune IA disponible pour le résumé")
        return
    
    with st.spinner("Création du résumé..."):
        summary_prompt = f"""Crée un résumé exécutif de ce document juridique.

Le résumé doit contenir :
1. Type et objet du document
2. Points clés (3-5 points)
3. Demandes principales
4. Enjeux juridiques

Document :
{content[:5000]}...

Résumé en 200-300 mots maximum."""
        
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            summary_prompt,
            "Tu es un expert en synthèse juridique."
        )
        
        if response['success']:
            st.markdown("### 📋 Résumé exécutif")
            st.write(response['response'])
            
            # Proposer de sauvegarder
            if st.button("💾 Sauvegarder le résumé"):
                st.session_state.document_summary = response['response']

def convert_analysis_to_report(analysis_content: str, analysis_metadata: dict):
    """Convertit une analyse en rapport formel"""
    st.info("📝 Conversion en rapport formel...")
    
    # Créer une requête de rédaction de rapport
    report_query = f"rédiger rapport formel basé sur analyse {analysis_metadata.get('type', 'générale')}"
    
    # Injecter le contenu de l'analyse
    st.session_state.analysis_to_report = analysis_content

# modules/recherche.py
"""Module de recherche unifiée avec 100% des fonctionnalités intégrées"""

import streamlit as st
import asyncio
from datetime import datetime, timedelta
import re
import json
from collections import defaultdict, Counter
from typing import List, Dict, Optional, Any, Tuple, Union
import io
import base64
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

# Imports optionnels avec gestion d'erreur
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("⚠️ pandas non installé - Certaines fonctionnalités seront limitées")

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False
    print("⚠️ plotly non installé - Visualisations simplifiées")

try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False
    print("⚠️ networkx non installé - Analyse réseau limitée")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx non installé - Export Word non disponible")

# Imports des modules internes
from config.app_config import SearchMode, ANALYSIS_PROMPTS_AFFAIRES, LLMProvider
from managers.azure_blob_manager import AzureBlobManager
from managers.azure_search_manager import AzureSearchManager
from managers.multi_llm_manager import MultiLLMManager
from managers.document_manager import display_import_interface
from managers.dynamic_generators import generate_dynamic_search_prompts
from managers.legal_search import LegalSearchManager, display_legal_search_interface
from managers.jurisprudence_verifier import JurisprudenceVerifier
from models.dataclasses import Document, PieceSelectionnee
from utils.helpers import clean_key

# Configuration des styles de rédaction
REDACTION_STYLES = {
    'formel': {
        'name': 'Formel',
        'description': 'Style juridique classique et solennel',
        'tone': 'respectueux et distant',
        'vocabulary': 'technique et précis'
    },
    'persuasif': {
        'name': 'Persuasif',
        'description': 'Style argumentatif et convaincant',
        'tone': 'assertif et engagé',
        'vocabulary': 'percutant et imagé'
    },
    'technique': {
        'name': 'Technique',
        'description': 'Style factuel et détaillé',
        'tone': 'neutre et objectif',
        'vocabulary': 'spécialisé et exhaustif'
    },
    'synthétique': {
        'name': 'Synthétique',
        'description': 'Style concis et efficace',
        'tone': 'direct et clair',
        'vocabulary': 'simple et précis'
    },
    'pédagogique': {
        'name': 'Pédagogique',
        'description': 'Style explicatif et accessible',
        'tone': 'bienveillant et didactique',
        'vocabulary': 'vulgarisé et illustré'
    }
}

# Templates prédéfinis
DOCUMENT_TEMPLATES = {
    'conclusions_defense': {
        'name': 'Conclusions en défense',
        'structure': [
            'I. FAITS ET PROCÉDURE',
            'II. DISCUSSION',
            '   A. Sur la recevabilité',
            '   B. Sur le fond',
            '      1. Sur l\'élément matériel',
            '      2. Sur l\'élément intentionnel',
            '      3. Sur le préjudice',
            'III. PAR CES MOTIFS'
        ],
        'style': 'formel'
    },
    'plainte_simple': {
        'name': 'Plainte simple',
        'structure': [
            'Objet : Plainte',
            'EXPOSÉ DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PRÉJUDICES SUBIS',
            'DEMANDES',
            'PIÈCES JOINTES'
        ],
        'style': 'formel'
    },
    'plainte_avec_cpc': {
        'name': 'Plainte avec constitution de partie civile',
        'structure': [
            'Objet : Plainte avec constitution de partie civile',
            'EXPOSÉ DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PRÉJUDICES SUBIS',
            'CONSTITUTION DE PARTIE CIVILE',
            'ÉVALUATION DU PRÉJUDICE',
            'DEMANDES',
            'PIÈCES JOINTES'
        ],
        'style': 'formel'
    },
    'mise_en_demeure': {
        'name': 'Mise en demeure',
        'structure': [
            'MISE EN DEMEURE',
            'Rappel des faits',
            'Obligations non respectées',
            'Délai accordé',
            'Conséquences du défaut',
            'Réserves'
        ],
        'style': 'persuasif'
    },
    'note_synthese': {
        'name': 'Note de synthèse',
        'structure': [
            'SYNTHÈSE EXÉCUTIVE',
            'I. CONTEXTE',
            'II. ANALYSE',
            'III. RECOMMANDATIONS',
            'IV. PLAN D\'ACTION'
        ],
        'style': 'synthétique'
    }
}

# === FONCTIONS DE TRAITEMENT POUR IMPORT ===

def process_import_request(query: str, analysis: dict):
    """Traite une demande d'import"""
    
    st.info("📥 Interface d'import activée")
    
    # Configuration de l'import depuis la config
    file_types = analysis['details'].get('file_types', ['pdf', 'docx', 'txt'])
    destination = st.session_state.get('import_destination', 'Documents locaux')
    
    # Si des fichiers sont déjà uploadés
    if st.session_state.get('import_files'):
        with st.spinner("📥 Import des fichiers en cours..."):
            imported_count = 0
            
            for file in st.session_state.import_files:
                try:
                    # Lire le contenu
                    content = file.read()
                    
                    # Traiter selon le type
                    if file.name.endswith('.pdf'):
                        text_content = extract_text_from_pdf(content)
                    elif file.name.endswith('.docx'):
                        text_content = extract_text_from_docx(content)
                    else:
                        text_content = content.decode('utf-8', errors='ignore')
                    
                    # Stocker selon la destination
                    if destination == 'Azure Blob':
                        store_in_azure_blob(file.name, content, analysis.get('reference'))
                    else:
                        store_locally(file.name, text_content, analysis.get('reference'))
                    
                    imported_count += 1
                    
                except Exception as e:
                    st.error(f"❌ Erreur import {file.name}: {str(e)}")
            
            st.success(f"✅ {imported_count} fichiers importés avec succès")
            
            # Analyse automatique si demandée
            if st.session_state.get('auto_analyze_import', True):
                st.info("🤖 Analyse automatique des documents importés...")
                analyze_imported_documents()
    
    else:
        st.warning("⚠️ Aucun fichier sélectionné. Utilisez l'interface d'import ci-dessous.")

def extract_text_from_pdf(content: bytes) -> str:
    """Extrait le texte d'un PDF"""
    try:
        import PyPDF2
        import io
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text
        
    except ImportError:
        return "PDF - Installation de PyPDF2 requise pour l'extraction"
    except Exception as e:
        return f"Erreur extraction PDF: {str(e)}"

def extract_text_from_docx(content: bytes) -> str:
    """Extrait le texte d'un DOCX"""
    if DOCX_AVAILABLE:
        try:
            import io
            doc = docx.Document(io.BytesIO(content))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"Erreur extraction DOCX: {str(e)}"
    else:
        return "DOCX - Installation de python-docx requise"

def store_locally(filename: str, content: str, reference: str = None):
    """Stocke un document localement"""
    doc_id = f"imported_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{clean_key(filename)}"
    
    if 'azure_documents' not in st.session_state:
        st.session_state.azure_documents = {}
    
    st.session_state.azure_documents[doc_id] = Document(
        id=doc_id,
        title=filename,
        content=content,
        source=f"Import {reference if reference else 'direct'}",
        metadata={
            'imported_at': datetime.now().isoformat(),
            'reference': reference,
            'original_filename': filename
        }
    )

def store_in_azure_blob(filename: str, content: bytes, reference: str = None):
    """Stocke un document dans Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        container = reference if reference else 'imports'
        blob_name = f"{datetime.now().strftime('%Y%m%d')}_{filename}"
        
        try:
            blob_manager.upload_blob(container, blob_name, content)
            st.success(f"✅ {filename} uploadé dans Azure")
        except Exception as e:
            st.error(f"❌ Erreur Azure: {str(e)}")
    else:
        st.warning("⚠️ Azure Blob non connecté - Stockage local")
        store_locally(filename, content.decode('utf-8', errors='ignore'), reference)

def analyze_imported_documents():
    """Analyse automatique des documents importés"""
    # Récupérer les derniers documents importés
    recent_docs = []
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if 'imported_at' in doc.metadata:
            import_time = datetime.fromisoformat(doc.metadata['imported_at'])
            if (datetime.now() - import_time).seconds < 300:  # Moins de 5 minutes
                recent_docs.append(doc)
    
    if recent_docs:
        # Lancer une analyse automatique
        analysis_prompt = f"""Analyse ces {len(recent_docs)} documents récemment importés.
        
Identifie :
1. Type de document
2. Parties impliquées
3. Dates importantes
4. Points juridiques clés
5. Risques identifiés

Documents :
{chr(10).join([f"- {doc.title}: {doc.content[:500]}..." for doc in recent_docs[:5]])}"""
        
        llm_manager = MultiLLMManager()
        if llm_manager.clients:
            provider = list(llm_manager.clients.keys())[0]
            response = llm_manager.query_single_llm(
                provider,
                analysis_prompt,
                "Tu es un expert en analyse documentaire juridique."
            )
            
            if response['success']:
                st.session_state.import_analysis = response['response']
                st.success("✅ Analyse des imports terminée")

# === FONCTIONS DE TRAITEMENT POUR EXPORT ===

def process_export_request(query: str, analysis: dict):
    """Traite une demande d'export"""
    
    format = analysis['details'].get('format', 'docx')
    
    # Déterminer ce qu'il faut exporter
    content_to_export = None
    filename_base = "export"
    
    # Priorité : rédaction > analyse > recherche > sélection
    if st.session_state.get('redaction_result'):
        content_to_export = st.session_state.redaction_result['document']
        filename_base = f"{st.session_state.redaction_result['type']}"
    
    elif st.session_state.get('timeline_result'):
        content_to_export = export_timeline_content(st.session_state.timeline_result)
        filename_base = "chronologie"
    
    elif st.session_state.get('mapping_result'):
        content_to_export = export_mapping_content(st.session_state.mapping_result)
        filename_base = "cartographie"
    
    elif st.session_state.get('comparison_result'):
        content_to_export = export_comparison_content(st.session_state.comparison_result)
        filename_base = "comparaison"
    
    elif st.session_state.get('ai_analysis_results'):
        content_to_export = format_ai_analysis_for_export(st.session_state.ai_analysis_results)
        filename_base = "analyse"
    
    elif st.session_state.get('search_results'):
        content_to_export = format_search_results_for_export(st.session_state.search_results)
        filename_base = "recherche"
    
    else:
        st.warning("⚠️ Aucun contenu à exporter")
        return
    
    # Exporter selon le format
    export_functions = {
        'docx': export_to_docx,
        'pdf': export_to_pdf,
        'txt': export_to_txt,
        'html': export_to_html,
        'xlsx': export_to_xlsx
    }
    
    export_func = export_functions.get(format, export_to_txt)
    
    try:
        exported_data = export_func(content_to_export, analysis)
        
        # Créer le bouton de téléchargement
        st.download_button(
            f"💾 Télécharger {format.upper()}",
            exported_data,
            f"{filename_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}",
            get_mime_type(format),
            key=f"download_export_{format}"
        )
        
        st.success(f"✅ Export {format.upper()} prêt")
        
    except Exception as e:
        st.error(f"❌ Erreur export: {str(e)}")

def export_timeline_content(timeline_result: dict) -> str:
    """Exporte le contenu d'une chronologie"""
    content = f"CHRONOLOGIE - {timeline_result['type'].upper()}\n"
    content += f"Générée le {timeline_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Basée sur {timeline_result['document_count']} documents\n\n"
    
    for event in timeline_result['events']:
        content += f"{event.get('date', 'N/A')} - {event.get('description', '')}\n"
        if 'actors' in event:
            content += f"   Acteurs: {', '.join(event['actors'])}\n"
        if 'source' in event:
            content += f"   Source: {event['source']}\n"
        content += "\n"
    
    return content

def export_mapping_content(mapping_result: dict) -> str:
    """Exporte le contenu d'une cartographie"""
    content = f"CARTOGRAPHIE DES RELATIONS - {mapping_result['type'].upper()}\n"
    content += f"Générée le {mapping_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n\n"
    
    content += f"STATISTIQUES:\n"
    content += f"- Entités: {mapping_result['analysis']['node_count']}\n"
    content += f"- Relations: {mapping_result['analysis']['edge_count']}\n"
    content += f"- Densité: {mapping_result['analysis']['density']:.2%}\n\n"
    
    content += "ACTEURS PRINCIPAUX:\n"
    for i, player in enumerate(mapping_result['analysis']['key_players'], 1):
        content += f"{i}. {player}\n"
    
    content += "\nENTITÉS:\n"
    for entity in mapping_result['entities']:
        content += f"- {entity['name']} ({entity.get('type', 'N/A')})\n"
    
    content += "\nRELATIONS:\n"
    for relation in mapping_result['relations']:
        content += f"- {relation['source']} → {relation['target']}: {relation.get('type', 'N/A')}\n"
    
    return content

def export_comparison_content(comparison_result: dict) -> str:
    """Exporte le contenu d'une comparaison"""
    content = f"COMPARAISON - {comparison_result['type'].upper()}\n"
    content += f"Générée le {comparison_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Documents comparés: {comparison_result['document_count']}\n\n"
    
    comparison = comparison_result['comparison']
    
    content += "CONVERGENCES:\n"
    for conv in comparison.get('convergences', []):
        content += f"- {conv['point']}\n  {conv['details']}\n\n"
    
    content += "DIVERGENCES:\n"
    for div in comparison.get('divergences', []):
        content += f"- {div['point']}\n  {div['details']}\n\n"
    
    if comparison.get('analysis'):
        content += f"ANALYSE:\n{comparison['analysis']}\n"
    
    return content

def export_to_docx(content: str, analysis: dict) -> bytes:
    """Exporte vers Word avec mise en forme"""
    if not DOCX_AVAILABLE:
        # Fallback en texte brut
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Styles personnalisés
        styles = doc.styles
        
        # En-tête
        if st.session_state.get('export_metadata', True):
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Titre principal
        title = doc.add_heading(analysis.get('document_type', 'Document').upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenu avec mise en forme
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            # Détecter les niveaux de titre
            if line.strip().isupper() and len(line.strip()) > 5:
                doc.add_heading(line.strip(), 1)
            elif line.strip().startswith('###'):
                doc.add_heading(line.strip().replace('#', '').strip(), 3)
            elif line.strip().startswith('##'):
                doc.add_heading(line.strip().replace('#', '').strip(), 2)
            elif line.strip().startswith('#'):
                doc.add_heading(line.strip().replace('#', '').strip(), 1)
            elif line.strip().startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
                p = doc.add_paragraph()
                p.add_run(line.strip()).bold = True
                p.style.font.size = Pt(14)
            elif line.strip().startswith(('A.', 'B.', 'C.', 'D.')):
                p = doc.add_paragraph(line.strip(), style='List Number')
                p.style.font.size = Pt(12)
            elif line.strip().startswith('-'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(line.strip())
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
        
        # Table des matières si demandée
        if st.session_state.get('docx_toc', True):
            # Insérer au début
            doc.paragraphs[0].insert_paragraph_before("TABLE DES MATIÈRES", style='Heading 1')
            # Note: La vraie TOC nécessite des manipulations XML complexes
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création DOCX: {e}")
        return content.encode('utf-8')

def export_to_pdf(content: str, analysis: dict) -> bytes:
    """Exporte vers PDF"""
    # Nécessiterait reportlab ou weasyprint
    st.warning("⚠️ Export PDF nécessite l'installation de bibliothèques supplémentaires")
    
    # Pour l'instant, retourner le contenu texte
    return content.encode('utf-8')

def export_to_txt(content: str, analysis: dict) -> bytes:
    """Exporte en texte brut"""
    return content.encode('utf-8')

def export_to_html(content: str, analysis: dict) -> bytes:
    """Exporte vers HTML avec style"""
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <title>{analysis.get('document_type', 'Document').title()}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; margin: 40px; }}
        h1 {{ text-align: center; color: #2c3e50; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        p {{ text-align: justify; line-height: 1.6; }}
        .metadata {{ text-align: right; color: #95a5a6; font-size: 0.9em; }}
        .section {{ margin: 20px 0; }}
        ul, ol {{ margin-left: 20px; }}
    </style>
</head>
<body>
    <div class="metadata">Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</div>
    <h1>{analysis.get('document_type', 'Document').upper()}</h1>
    <div class="content">
"""
    
    # Convertir le contenu en HTML
    lines = content.split('\n')
    in_list = False
    
    for line in lines:
        if not line.strip():
            if in_list:
                html += "</ul>"
                in_list = False
            continue
        
        if line.strip().isupper() and len(line.strip()) > 5:
            html += f"<h2>{line.strip()}</h2>\n"
        elif line.strip().startswith('-'):
            if not in_list:
                html += "<ul>\n"
                in_list = True
            html += f"<li>{line.strip()[1:].strip()}</li>\n"
        else:
            if in_list:
                html += "</ul>\n"
                in_list = False
            html += f"<p>{line.strip()}</p>\n"
    
    html += """
    </div>
</body>
</html>"""
    
    return html.encode('utf-8')

def export_to_xlsx(content: str, analysis: dict) -> bytes:
    """Exporte vers Excel"""
    if not PANDAS_AVAILABLE:
        st.error("❌ pandas requis pour l'export Excel")
        return content.encode('utf-8')
    
    try:
        import io
        
        # Créer un DataFrame selon le type de contenu
        if 'timeline_result' in st.session_state:
            df = pd.DataFrame(st.session_state.timeline_result['events'])
        elif 'mapping_result' in st.session_state:
            df_entities = pd.DataFrame(st.session_state.mapping_result['entities'])
            df_relations = pd.DataFrame(st.session_state.mapping_result['relations'])
        else:
            # Contenu générique
            lines = content.split('\n')
            df = pd.DataFrame({'Contenu': lines})
        
        # Exporter vers Excel
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            if 'mapping_result' in st.session_state:
                df_entities.to_excel(writer, sheet_name='Entités', index=False)
                df_relations.to_excel(writer, sheet_name='Relations', index=False)
            else:
                df.to_excel(writer, sheet_name='Données', index=False)
            
            # Ajouter une feuille de métadonnées
            metadata = pd.DataFrame({
                'Propriété': ['Type', 'Date génération', 'Source'],
                'Valeur': [
                    analysis.get('document_type', 'Export'),
                    datetime.now().strftime('%d/%m/%Y %H:%M'),
                    analysis.get('reference', 'N/A')
                ]
            })
            metadata.to_excel(writer, sheet_name='Métadonnées', index=False)
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur export Excel: {e}")
        return content.encode('utf-8')

def get_mime_type(format: str) -> str:
    """Retourne le type MIME pour un format"""
    mime_types = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'html': 'text/html',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    }
    return mime_types.get(format, 'application/octet-stream')

# === FONCTIONS DE TRAITEMENT POUR EMAIL ===

def process_email_request(query: str, analysis: dict):
    """Traite une demande d'envoi par email"""
    
    recipients = analysis.get('recipients', [])
    
    if not recipients:
        st.error("❌ Aucun destinataire spécifié")
        return
    
    # Préparer le contenu
    content = prepare_email_content(analysis)
    
    if not content:
        st.warning("⚠️ Aucun contenu à envoyer")
        return
    
    # Configuration email
    email_config = {
        'to': st.session_state.get('email_to', ', '.join(recipients)),
        'cc': st.session_state.get('email_cc', ''),
        'subject': st.session_state.get('email_subject', 'Document juridique'),
        'body': content['body'],
        'attachments': []
    }
    
    # Préparer les pièces jointes
    if st.session_state.get('email_attach_current', True):
        attachment_format = st.session_state.get('email_attachment_format', 'pdf')
        attachment_data = export_current_content(attachment_format)
        
        if attachment_data:
            email_config['attachments'].append({
                'filename': f"document.{attachment_format}",
                'data': attachment_data,
                'mime_type': get_mime_type(attachment_format)
            })
    
    # Afficher l'aperçu
    show_email_preview(email_config)
    
    # Bouton d'envoi
    if st.button("📧 Envoyer l'email", key="send_email_button"):
        if send_email(email_config):
            st.success("✅ Email envoyé avec succès")
            st.session_state.email_sent = True
        else:
            st.error("❌ Erreur lors de l'envoi")

def prepare_email_content(analysis: dict) -> dict:
    """Prépare le contenu de l'email"""
    content = {'body': '', 'attachments': []}
    
    # Corps de l'email selon le contexte
    if st.session_state.get('redaction_result'):
        result = st.session_state.redaction_result
        content['body'] = f"""Bonjour,

Veuillez trouver ci-joint le document {result['type']} demandé.

Ce document a été généré automatiquement à partir de l'analyse de votre dossier.

Cordialement,
[Votre nom]"""
        
    else:
        content['body'] = """Bonjour,

Veuillez trouver ci-joint le document demandé.

Cordialement,
[Votre nom]"""
    
    return content

def show_email_preview(email_config: dict):
    """Affiche un aperçu de l'email"""
    with st.expander("📧 Aperçu de l'email", expanded=True):
        st.text_input("À:", value=email_config['to'], disabled=True)
        st.text_input("Cc:", value=email_config['cc'], disabled=True)
        st.text_input("Objet:", value=email_config['subject'], disabled=True)
        st.text_area("Corps:", value=email_config['body'], height=200, disabled=True)
        
        if email_config['attachments']:
            st.write("📎 Pièces jointes:")
            for att in email_config['attachments']:
                st.write(f"- {att['filename']}")

def send_email(email_config: dict) -> bool:
    """Envoie l'email (fonction simplifiée)"""
    try:
        # Configuration SMTP (à adapter)
        smtp_config = {
            'server': st.secrets.get('smtp_server', 'smtp.gmail.com'),
            'port': st.secrets.get('smtp_port', 587),
            'username': st.secrets.get('smtp_username', ''),
            'password': st.secrets.get('smtp_password', '')
        }
        
        if not smtp_config['username']:
            st.error("❌ Configuration email manquante")
            return False
        
        # Créer le message
        msg = MIMEMultipart()
        msg['From'] = smtp_config['username']
        msg['To'] = email_config['to']
        msg['Cc'] = email_config['cc']
        msg['Subject'] = email_config['subject']
        
        # Corps
        msg.attach(MIMEText(email_config['body'], 'plain'))
        
        # Pièces jointes
        for att in email_config['attachments']:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(att['data'])
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {att["filename"]}'
            )
            msg.attach(part)
        
        # Envoyer
        with smtplib.SMTP(smtp_config['server'], smtp_config['port']) as server:
            server.starttls()
            server.login(smtp_config['username'], smtp_config['password'])
            
            recipients = [email_config['to']]
            if email_config['cc']:
                recipients.extend(email_config['cc'].split(','))
            
            server.send_message(msg)
        
        return True
        
    except Exception as e:
        st.error(f"Erreur envoi email: {e}")
        return False

def export_current_content(format: str) -> bytes:
    """Exporte le contenu actuel dans le format demandé"""
    # Réutiliser les fonctions d'export
    analysis = st.session_state.get('current_analysis', {})
    
    if st.session_state.get('redaction_result'):
        content = st.session_state.redaction_result['document']
    else:
        content = "Aucun contenu disponible"
    
    export_functions = {
        'pdf': export_to_pdf,
        'docx': export_to_docx,
        'txt': export_to_txt
    }
    
    export_func = export_functions.get(format, export_to_txt)
    return export_func(content, analysis)

# === FONCTIONS DE TRAITEMENT POUR SÉLECTION DE PIÈCES ===

def process_piece_selection_request(query: str, analysis: dict):
    """Traite une demande de sélection de pièces"""
    
    # Interface de sélection
    st.markdown("### 📋 Sélection de pièces")
    
    # Collecter les documents disponibles
    available_docs = collect_available_documents(analysis)
    
    if not available_docs:
        st.warning("⚠️ Aucun document disponible")
        return
    
    # Grouper par catégorie
    categories = group_documents_by_category(available_docs)
    
    # Interface de sélection par catégorie
    selected_pieces = []
    
    for category, docs in categories.items():
        with st.expander(f"📁 {category} ({len(docs)} documents)", expanded=True):
            select_all = st.checkbox(f"Tout sélectionner - {category}", key=f"select_all_{category}")
            
            for doc in docs:
                is_selected = st.checkbox(
                    f"📄 {doc['title']}",
                    value=select_all,
                    key=f"select_doc_{doc['id']}",
                    help=f"Source: {doc.get('source', 'N/A')}"
                )
                
                if is_selected:
                    selected_pieces.append(PieceSelectionnee(
                        numero=len(selected_pieces) + 1,
                        titre=doc['title'],
                        description=doc.get('description', ''),
                        categorie=category,
                        date=doc.get('date'),
                        source=doc.get('source', ''),
                        pertinence=calculate_piece_relevance(doc, analysis)
                    ))
    
    # Sauvegarder la sélection
    st.session_state.selected_pieces = selected_pieces
    
    # Actions sur la sélection
    if selected_pieces:
        st.success(f"✅ {len(selected_pieces)} pièces sélectionnées")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📊 Créer bordereau", key="create_bordereau_from_selection"):
                process_bordereau_request(query, analysis)
        
        with col2:
            if st.button("📄 Synthétiser", key="synthesize_selection"):
                synthesize_selected_pieces(selected_pieces)
        
        with col3:
            if st.button("📤 Exporter liste", key="export_piece_list"):
                export_piece_list(selected_pieces)

def collect_available_documents(analysis: dict) -> list:
    """Collecte tous les documents disponibles"""
    documents = []
    
    # Documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        documents.append({
            'id': doc_id,
            'title': doc.title,
            'content': doc.content,
            'source': doc.source,
            'metadata': doc.metadata
        })
    
    # Si référence spécifique
    if analysis.get('reference'):
        ref_docs = search_by_reference(f"@{analysis['reference']}")
        documents.extend(ref_docs)
    
    return documents

def group_documents_by_category(documents: list) -> dict:
    """Groupe les documents par catégorie"""
    categories = defaultdict(list)
    
    for doc in documents:
        # Déterminer la catégorie
        category = determine_document_category(doc)
        categories[category].append(doc)
    
    return dict(categories)

def determine_document_category(doc: dict) -> str:
    """Détermine la catégorie d'un document"""
    title_lower = doc.get('title', '').lower()
    content_lower = doc.get('content', '')[:500].lower()
    
    # Patterns de catégories
    category_patterns = {
        'Procédure': ['plainte', 'procès-verbal', 'audition', 'perquisition', 'garde à vue'],
        'Expertise': ['expertise', 'expert', 'rapport technique', 'analyse'],
        'Comptabilité': ['bilan', 'compte', 'comptable', 'facture', 'devis'],
        'Contrats': ['contrat', 'convention', 'accord', 'avenant'],
        'Correspondance': ['courrier', 'email', 'lettre', 'mail'],
        'Pièces d\'identité': ['carte identité', 'passeport', 'kbis', 'statuts'],
        'Bancaire': ['relevé', 'virement', 'compte bancaire', 'rib']
    }
    
    for category, keywords in category_patterns.items():
        if any(kw in title_lower or kw in content_lower for kw in keywords):
            return category
    
    return 'Autres'

def calculate_piece_relevance(doc: dict, analysis: dict) -> float:
    """Calcule la pertinence d'une pièce"""
    score = 0.5
    
    # Si le document contient des mots-clés de l'analyse
    if analysis.get('subject_matter'):
        if analysis['subject_matter'] in doc.get('content', '').lower():
            score += 0.3
    
    # Si référence dans le titre
    if analysis.get('reference'):
        if analysis['reference'] in doc.get('title', '').lower():
            score += 0.2
    
    return min(score, 1.0)

# === FONCTIONS DE TRAITEMENT POUR BORDEREAU ===

def process_bordereau_request(query: str, analysis: dict):
    """Traite une demande de création de bordereau"""
    
    pieces = st.session_state.get('selected_pieces', [])
    
    if not pieces:
        st.warning("⚠️ Aucune pièce sélectionnée pour le bordereau")
        return
    
    # Créer le bordereau
    bordereau = create_bordereau(pieces, analysis)
    
    # Afficher le bordereau
    st.markdown("### 📊 Bordereau de communication de pièces")
    
    # En-tête
    st.text_area(
        "En-tête du bordereau",
        value=bordereau['header'],
        height=150,
        key="bordereau_header"
    )
    
    # Table des pièces
    if PANDAS_AVAILABLE:
        df = pd.DataFrame([
            {
                'N°': p.numero,
                'Titre': p.titre,
                'Description': p.description,
                'Catégorie': p.categorie,
                'Date': p.date.strftime('%d/%m/%Y') if p.date else 'N/A'
            }
            for p in pieces
        ])
        
        st.dataframe(df, use_container_width=True)
    else:
        # Affichage sans pandas
        for piece in pieces:
            st.write(f"**{piece.numero}.** {piece.titre}")
            if piece.description:
                st.caption(piece.description)
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.download_button(
            "💾 Télécharger bordereau",
            create_bordereau_document(bordereau),
            f"bordereau_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_bordereau"
        ):
            st.success("✅ Bordereau téléchargé")
    
    with col2:
        if st.button("✏️ Modifier sélection", key="modify_bordereau_selection"):
            st.session_state.show_piece_selection = True
            st.rerun()
    
    with col3:
        if st.button("📧 Envoyer bordereau", key="send_bordereau"):
            st.session_state.universal_query = f"envoyer bordereau @{analysis.get('reference', 'dossier')}"
            st.rerun()
    
    # Stocker le bordereau
    st.session_state.current_bordereau = bordereau

def create_bordereau(pieces: list, analysis: dict) -> dict:
    """Crée un bordereau structuré"""
    
    bordereau = {
        'header': f"""BORDEREAU DE COMMUNICATION DE PIÈCES

AFFAIRE : {analysis.get('reference', 'N/A').upper()}
DATE : {datetime.now().strftime('%d/%m/%Y')}
NOMBRE DE PIÈCES : {len(pieces)}

POUR : [À compléter]
CONTRE : [À compléter]

PIÈCES COMMUNIQUÉES :""",
        'pieces': pieces,
        'footer': """Je certifie que les pièces communiquées sont conformes aux originaux en ma possession.

Fait à [Ville], le {datetime.now().strftime('%d/%m/%Y')}

[Signature]""",
        'metadata': {
            'created_at': datetime.now(),
            'piece_count': len(pieces),
            'categories': list(set(p.categorie for p in pieces)),
            'reference': analysis.get('reference')
        }
    }
    
    return bordereau

def create_bordereau_document(bordereau: dict) -> bytes:
    """Crée le document Word du bordereau"""
    if not DOCX_AVAILABLE:
        # Fallback texte
        content = bordereau['header'] + '\n\n'
        
        for piece in bordereau['pieces']:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            content += f"   Catégorie: {piece.categorie}\n"
            if piece.date:
                content += f"   Date: {piece.date.strftime('%d/%m/%Y')}\n"
            content += "\n"
        
        content += bordereau['footer']
        
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # En-tête
        for line in bordereau['header'].split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                if 'BORDEREAU' in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(16)
        
        # Table des pièces
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # En-têtes de colonnes
        headers = ['N°', 'Titre', 'Description', 'Catégorie', 'Date']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Lignes de pièces
        for piece in bordereau['pieces']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(piece.numero)
            row_cells[1].text = piece.titre
            row_cells[2].text = piece.description or ''
            row_cells[3].text = piece.categorie
            row_cells[4].text = piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A'
        
        # Pied de page
        doc.add_paragraph()
        for line in bordereau['footer'].split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création bordereau: {e}")
        return create_bordereau_document.__wrapped__(bordereau)  # Fallback

# === FONCTIONS DE TRAITEMENT POUR SYNTHÈSE ===

def process_synthesis_request(query: str, analysis: dict):
    """Traite une demande de synthèse"""
    
    # Déterminer la source de la synthèse
    if st.session_state.get('selected_pieces'):
        content_to_synthesize = synthesize_selected_pieces(st.session_state.selected_pieces)
    
    elif analysis.get('reference'):
        docs = search_by_reference(f"@{analysis['reference']}")
        content_to_synthesize = synthesize_documents(docs)
    
    else:
        st.warning("⚠️ Aucun contenu à synthétiser")
        return
    
    # Stocker le résultat
    st.session_state.synthesis_result = content_to_synthesize

def synthesize_selected_pieces(pieces: list) -> dict:
    """Synthétise les pièces sélectionnées"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le contexte
    context = "PIÈCES À SYNTHÉTISER:\n\n"
    
    for piece in pieces[:20]:  # Limiter
        context += f"Pièce {piece.numero}: {piece.titre}\n"
        context += f"Catégorie: {piece.categorie}\n"
        if piece.description:
            context += f"Description: {piece.description}\n"
        context += "\n"
    
    # Prompt de synthèse
    synthesis_prompt = f"""{context}

Crée une synthèse structurée de ces pièces.

La synthèse doit inclure:
1. Vue d'ensemble des pièces
2. Points clés par catégorie
3. Chronologie si applicable
4. Liens et relations entre pièces
5. Points d'attention juridiques
6. Recommandations

Format professionnel avec titres et sous-sections."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            synthesis_prompt,
            "Tu es un expert en analyse et synthèse de documents juridiques."
        )
        
        if response['success']:
            return {
                'content': response['response'],
                'piece_count': len(pieces),
                'categories': list(set(p.categorie for p in pieces)),
                'timestamp': datetime.now()
            }
        else:
            return {'error': 'Échec de la synthèse'}
            
    except Exception as e:
        return {'error': f'Erreur synthèse: {str(e)}'}

def synthesize_documents(documents: list) -> dict:
    """Synthétise une liste de documents"""
    # Convertir en pièces pour réutiliser la fonction
    pieces = []
    
    for i, doc in enumerate(documents):
        pieces.append(PieceSelectionnee(
            numero=i + 1,
            titre=doc.get('title', 'Sans titre'),
            description=doc.get('content', '')[:200] + '...' if doc.get('content') else '',
            categorie=determine_document_category(doc),
            source=doc.get('source', '')
        ))
    
    return synthesize_selected_pieces(pieces)

# === FONCTIONS DE TRAITEMENT POUR TEMPLATES ===

def process_template_request(query: str, analysis: dict):
    """Traite une demande liée aux templates"""
    
    action = analysis['details'].get('action', 'apply')
    
    if action == 'create':
        create_new_template()
    
    elif action == 'edit':
        edit_existing_template()
    
    else:  # apply
        apply_template()

def create_new_template():
    """Crée un nouveau template"""
    
    template_name = st.session_state.get('new_template_name', '')
    
    if not template_name:
        st.warning("⚠️ Nom du template requis")
        return
    
    # Base du template
    base_template = st.session_state.get('base_template', 'Vide')
    
    if base_template == 'Vide':
        template_content = {
            'name': template_name,
            'structure': [],
            'style': 'formel',
            'category': st.session_state.get('template_category', 'Autre')
        }
    else:
        # Copier depuis un template existant
        template_content = DOCUMENT_TEMPLATES.get(base_template, {}).copy()
        template_content['name'] = template_name
    
    # Éditeur de structure
    st.markdown("### 📝 Structure du template")
    
    structure = st.text_area(
        "Sections (une par ligne)",
        value='\n'.join(template_content.get('structure', [])),
        height=300,
        key="template_structure_editor"
    )
    
    template_content['structure'] = [s.strip() for s in structure.split('\n') if s.strip()]
    
    # Sauvegarder
    if st.button("💾 Sauvegarder le template", key="save_new_template"):
        if 'saved_templates' not in st.session_state:
            st.session_state.saved_templates = {}
        
        st.session_state.saved_templates[clean_key(template_name)] = template_content
        st.success(f"✅ Template '{template_name}' sauvegardé")
        
        # Optionnel : sauvegarder dans un fichier
        save_templates_to_file()

def edit_existing_template():
    """Édite un template existant"""
    
    template_to_edit = st.session_state.get('template_to_edit')
    
    if not template_to_edit:
        st.warning("⚠️ Sélectionnez un template à modifier")
        return
    
    # Charger le template
    if template_to_edit in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[template_to_edit].copy()
        is_builtin = True
    else:
        template = st.session_state.saved_templates.get(template_to_edit, {})
        is_builtin = False
    
    if not template:
        st.error("❌ Template introuvable")
        return
    
    # Éditeur
    st.markdown(f"### ✏️ Édition du template '{template.get('name', template_to_edit)}'")
    
    if is_builtin:
        st.info("ℹ️ Template intégré - Les modifications seront sauvegardées comme nouveau template")
    
    # Nom
    new_name = st.text_input(
        "Nom du template",
        value=template.get('name', ''),
        key="edit_template_name"
    )
    
    # Structure
    structure = st.text_area(
        "Structure",
        value='\n'.join(template.get('structure', [])),
        height=300,
        key="edit_template_structure"
    )
    
    # Style
    style = st.selectbox(
        "Style par défaut",
        list(REDACTION_STYLES.keys()),
        index=list(REDACTION_STYLES.keys()).index(template.get('style', 'formel')),
        format_func=lambda x: REDACTION_STYLES[x]['name'],
        key="edit_template_style"
    )
    
    # Sauvegarder les modifications
    if st.button("💾 Sauvegarder les modifications", key="save_template_edits"):
        updated_template = {
            'name': new_name,
            'structure': [s.strip() for s in structure.split('\n') if s.strip()],
            'style': style,
            'category': template.get('category', 'Autre')
        }
        
        if is_builtin:
            # Sauvegarder comme nouveau
            st.session_state.saved_templates[clean_key(new_name)] = updated_template
            st.success(f"✅ Nouveau template '{new_name}' créé")
        else:
            # Mettre à jour l'existant
            st.session_state.saved_templates[template_to_edit] = updated_template
            st.success(f"✅ Template '{new_name}' mis à jour")

def apply_template():
    """Applique un template sélectionné"""
    
    selected_template = st.session_state.get('selected_template')
    
    if not selected_template:
        st.warning("⚠️ Sélectionnez un template à appliquer")
        return
    
    # Charger le template
    if selected_template in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[selected_template]
    else:
        template = st.session_state.saved_templates.get(selected_template, {})
    
    if not template:
        st.error("❌ Template introuvable")
        return
    
    # Créer une requête de rédaction avec le template
    st.session_state.universal_query = f"rédiger {template.get('name', 'document')} avec template {selected_template}"
    
    # Définir le style
    st.session_state.redaction_style = template.get('style', 'formel')
    
    # Déclencher la rédaction
    st.info(f"✅ Template '{template.get('name')}' appliqué - Lancez la rédaction")

def save_templates_to_file():
    """Sauvegarde les templates dans un fichier"""
    try:
        import json
        
        templates_data = {
            'builtin': DOCUMENT_TEMPLATES,
            'custom': st.session_state.get('saved_templates', {})
        }
        
        # Créer un fichier téléchargeable
        json_str = json.dumps(templates_data, indent=2, ensure_ascii=False)
        
        st.download_button(
            "💾 Exporter tous les templates",
            json_str,
            f"templates_{datetime.now().strftime('%Y%m%d')}.json",
            "application/json",
            key="export_templates"
        )
        
    except Exception as e:
        st.error(f"Erreur sauvegarde templates: {e}")

# === FONCTIONS DE TRAITEMENT POUR JURISPRUDENCE ===

def process_jurisprudence_request(query: str, analysis: dict):
    """Traite une demande de recherche de jurisprudence"""
    
    # Utiliser le gestionnaire de recherche juridique
    st.session_state.jurisprudence_search_active = True
    
    # Rediriger vers l'onglet jurisprudence
    st.info("⚖️ Recherche de jurisprudence activée - Voir l'onglet Jurisprudence")

# === FONCTIONS DE TRAITEMENT POUR ANALYSE ===

def process_analysis_request(query: str, analysis: dict):
    """Traite une demande d'analyse IA"""
    
    # Collecter les documents pertinents
    documents = []
    
    if analysis.get('reference'):
        documents = search_by_reference(f"@{analysis['reference']}")
    else:
        # Recherche générale
        documents = perform_search(query)
    
    if not documents:
        st.warning("⚠️ Aucun document trouvé pour l'analyse")
        return
    
    # Déterminer le type d'analyse
    analysis_type = detect_analysis_type(query)
    
    # Lancer l'analyse appropriée
    with st.spinner("🤖 Analyse en cours..."):
        if analysis_type == 'risks':
            results = analyze_legal_risks(documents, query)
        elif analysis_type == 'compliance':
            results = analyze_compliance(documents, query)
        elif analysis_type == 'strategy':
            results = analyze_strategy(documents, query)
        else:
            results = perform_general_analysis(documents, query)
    
    # Stocker les résultats
    st.session_state.ai_analysis_results = results

def detect_analysis_type(query: str) -> str:
    """Détecte le type d'analyse demandé"""
    query_lower = query.lower()
    
    if any(word in query_lower for word in ['risque', 'danger', 'exposition', 'vulnérabilité']):
        return 'risks'
    elif any(word in query_lower for word in ['conformité', 'compliance', 'réglementation']):
        return 'compliance'
    elif any(word in query_lower for word in ['stratégie', 'défense', 'approche', 'tactique']):
        return 'strategy'
    else:
        return 'general'

def analyze_legal_risks(documents: list, query: str) -> dict:
    """Analyse les risques juridiques"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le prompt d'analyse
    risk_prompt = f"""Analyse les risques juridiques dans ces documents.

DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}

QUESTION: {query}

Identifie et évalue:
1. RISQUES PÉNAUX
   - Infractions potentielles
   - Éléments constitutifs
   - Niveau de risque (faible/moyen/élevé)
   - Sanctions encourues

2. RISQUES CIVILS
   - Responsabilités contractuelles
   - Préjudices potentiels
   - Montants estimés

3. RISQUES RÉPUTATIONNELS
   - Impact médiatique
   - Conséquences business

4. RECOMMANDATIONS
   - Actions préventives
   - Stratégies de mitigation
   - Priorités

Format structuré avec évaluation précise."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            risk_prompt,
            "Tu es un expert en analyse de risques juridiques et compliance."
        )
        
        if response['success']:
            return {
                'type': 'risk_analysis',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

def perform_general_analysis(documents: list, query: str) -> dict:
    """Analyse générale des documents"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Prompt général
    general_prompt = f"""Analyse ces documents pour répondre à la question.

DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}

QUESTION: {query}

Fournis une analyse complète et structurée."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            general_prompt,
            "Tu es un expert en analyse juridique."
        )
        
        if response['success']:
            return {
                'type': 'general_analysis',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

# === FONCTIONS DE TRAITEMENT POUR RECHERCHE ===

def process_search_request(query: str, analysis: dict):
    """Traite une demande de recherche simple"""
    
    # Recherche selon le contexte
    if analysis.get('reference'):
        results = search_by_reference(f"@{analysis['reference']}")
    else:
        results = perform_search(query)
    
    # Stocker les résultats
    st.session_state.search_results = results
    
    if not results:
        st.warning("⚠️ Aucun résultat trouvé")

def search_by_reference(reference: str) -> list:
    """Recherche par référence @"""
    results = []
    
    # Nettoyer la référence
    ref_clean = reference.replace('@', '').strip().lower()
    
    # Recherche dans les documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if ref_clean in doc.title.lower() or ref_clean in doc.source.lower():
            results.append({
                'id': doc_id,
                'title': doc.title,
                'content': doc.content,
                'source': doc.source,
                'score': 1.0 if ref_clean == doc.title.lower() else 0.8
            })
    
    # Recherche Azure si disponible
    search_manager = st.session_state.get('azure_search_manager')
    if search_manager and search_manager.is_connected():
        try:
            azure_results = search_manager.search(reference)
            results.extend(azure_results)
        except:
            pass
    
    # Trier par score
    results.sort(key=lambda x: x.get('score', 0), reverse=True)
    
    return results

def perform_search(query: str) -> list:
    """Effectue une recherche générale"""
    results = []
    
    # Recherche locale basique
    query_lower = query.lower()
    query_words = query_lower.split()
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        # Score simple basé sur les mots trouvés
        score = 0
        content_lower = doc.content.lower()
        title_lower = doc.title.lower()
        
        for word in query_words:
            if word in title_lower:
                score += 2
            if word in content_lower:
                score += 1
        
        if score > 0:
            results.append({
                'id': doc_id,
                'title': doc.title,
                'content': doc.content,
                'source': doc.source,
                'score': score / len(query_words)
            })
    
    # Recherche Azure si disponible
    search_manager = st.session_state.get('azure_search_manager')
    if search_manager and search_manager.is_connected():
        try:
            azure_results = search_manager.search(query)
            results.extend(azure_results)
        except:
            pass
    
    # Trier par score
    results.sort(key=lambda x: x.get('score', 0), reverse=True)
    
    return results[:50]  # Limiter à 50 résultats

# === FONCTIONS D'AFFICHAGE DES RÉSULTATS ===

def show_unified_results_tab():
    """Affiche tous les types de résultats dans un onglet unifié"""
    
    # Vérifier quel type de résultat afficher
    has_results = False
    
    # RÉSULTATS DE RÉDACTION (Priorité 1)
    if st.session_state.get('redaction_result'):
        show_redaction_results()
        has_results = True
    
    # RÉSULTATS DE PLAIDOIRIE (Priorité 2)
    elif st.session_state.get('plaidoirie_result'):
        show_plaidoirie_results()
        has_results = True
    
    # RÉSULTATS DE PRÉPARATION CLIENT (Priorité 3)
    elif st.session_state.get('preparation_client_result'):
        show_preparation_client_results()
        has_results = True
    
    # RÉSULTATS DE TIMELINE (Priorité 4)
    elif st.session_state.get('timeline_result'):
        show_timeline_results()
        has_results = True
    
    # RÉSULTATS DE MAPPING (Priorité 5)
    elif st.session_state.get('mapping_result'):
        show_mapping_results()
        has_results = True
    
    # RÉSULTATS DE COMPARAISON (Priorité 6)
    elif st.session_state.get('comparison_result'):
        show_comparison_results()
        has_results = True
    
    # RÉSULTATS DE SYNTHÈSE (Priorité 7)
    elif st.session_state.get('synthesis_result'):
        show_synthesis_results()
        has_results = True
    
    # RÉSULTATS D'ANALYSE IA (Priorité 8)
    elif st.session_state.get('ai_analysis_results'):
        show_ai_analysis_results()
        has_results = True
    
    # RÉSULTATS DE RECHERCHE (Priorité 9)
    elif st.session_state.get('search_results'):
        show_search_results()
        has_results = True
    
    # Message si aucun résultat
    if not has_results:
        st.info("💡 Utilisez la barre de recherche universelle pour commencer")
        
        # Suggestions d'utilisation
        st.markdown("""
        ### 🚀 Exemples de commandes
        
        **Recherche :**
        - `contrats société XYZ`
        - `@affaire_martin documents comptables`
        
        **Analyse :**
        - `analyser les risques @dossier_pénal`
        - `identifier les infractions @affaire_corruption`
        
        **Rédaction :**
        - `rédiger conclusions défense @affaire_martin abus biens sociaux`
        - `créer plainte avec constitution partie civile escroquerie`
        
        **Plaidoirie & Préparation :**
        - `préparer plaidoirie @affaire_martin audience correctionnelle`
        - `préparer client interrogatoire @dossier_fraude`
        - `questions réponses audience @affaire_escroquerie`
        
        **Visualisations :**
        - `chronologie des faits @affaire_martin`
        - `cartographie des sociétés @groupe_abc`
        - `comparer les auditions @témoins`
        
        **Gestion :**
        - `sélectionner pièces @dossier catégorie procédure`
        - `créer bordereau @pièces_sélectionnées`
        - `exporter analyse format word`
        """)

def show_redaction_results():
    """Affiche les résultats de rédaction avec toutes les fonctionnalités"""
    
    result = st.session_state.redaction_result
    
    st.markdown("### 📝 Document juridique généré")
    
    # Métadonnées
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        doc_icons = {
            'conclusions': '⚖️ Conclusions',
            'plainte': '📋 Plainte',
            'constitution_pc': '🛡️ Constitution PC',
            'courrier': '✉️ Courrier',
            'assignation': '📜 Assignation',
            'mémoire': '📚 Mémoire',
            'requête': '📄 Requête'
        }
        st.metric("Type", doc_icons.get(result['type'], '📄 Document'))
    
    with col2:
        providers_count = len([r for r in result.get('responses', []) if r.get('success')])
        st.metric("IA utilisées", providers_count)
    
    with col3:
        word_count = len(result['document'].split())
        st.metric("Mots", f"{word_count:,}")
    
    with col4:
        char_count = len(result['document'])
        st.metric("Caractères", f"{char_count:,}")
    
    # Zone d'édition principale
    st.markdown("#### 📄 Contenu du document")
    
    # Ajuster la hauteur selon le type de document
    # Les conclusions et plaintes sont généralement très longues
    height_by_type = {
        'conclusions': 800,  # Très long pour dossiers complexes
        'plainte': 700,      # Long aussi
        'constitution_pc': 700,
        'mémoire': 800,
        'assignation': 600,
        'requête': 600,
        'courrier': 400
    }
    
    text_height = height_by_type.get(result['type'], 600)
    
    edited_content = st.text_area(
        "Vous pouvez éditer le document",
        value=result['document'],
        height=text_height,
        key="edit_redaction_main",
        help="Le document a été généré pour être complet et détaillé, adapté aux dossiers complexes"
    )
    
    # Barre d'outils d'édition
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        if st.button("🔄 Régénérer", key="regenerate_main"):
            if st.session_state.get('last_universal_query'):
                process_universal_query(st.session_state.last_universal_query)
                st.rerun()
    
    with col2:
        if st.button("➕ Enrichir", key="enrich_document"):
            enrich_current_document(edited_content)
    
    with col3:
        if st.button("✂️ Synthétiser", key="summarize_document"):
            create_document_summary(edited_content)
    
    with col4:
        if st.button("🔍 Vérifier", key="verify_document"):
            verify_document_content(edited_content)
    
    with col5:
        if st.button("📊 Statistiques", key="document_stats"):
            show_document_statistics(edited_content)
    
    # Options d'export
    st.markdown("#### 💾 Export et partage")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # Export Word avec mise en forme
        docx_data = create_formatted_docx(edited_content, result['type'])
        if st.download_button(
            "📄 Word (.docx)",
            docx_data,
            f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx_main"
        ):
            st.success("✅ Document Word téléchargé")
    
    with col2:
        # Export PDF (si disponible)
        try:
            pdf_data = create_pdf_from_content(edited_content, result['type'])
            if st.download_button(
                "📑 PDF",
                pdf_data,
                f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "application/pdf",
                key="download_pdf_main"
            ):
                st.success("✅ PDF téléchargé")
        except:
            st.info("PDF nécessite des libs supplémentaires")
    
    with col3:
        # Export texte brut
        if st.download_button(
            "📝 Texte (.txt)",
            edited_content.encode('utf-8'),
            f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_txt_main"
        ):
            st.success("✅ Fichier texte téléchargé")
    
    with col4:
        # Préparer email
        if st.button("📧 Envoyer", key="prepare_email_main"):
            prepare_document_email(edited_content, result['type'])
    
    # Versions individuelles des IA (si plusieurs)
    if len(result.get('responses', [])) > 1:
        with st.expander("🤖 Versions par IA", expanded=False):
            for i, response in enumerate(result['responses']):
                if response.get('success'):
                    st.markdown(f"#### {response['provider']}")
                    
                    # Métriques par version
                    col1, col2 = st.columns([4, 1])
                    
                    with col1:
                        words = len(response['response'].split())
                        st.caption(f"📊 {words:,} mots")
                    
                    with col2:
                        if st.button(f"Utiliser", key=f"use_version_{i}"):
                            st.session_state.edit_redaction_main = response['response']
                            st.rerun()
                    
                    # Contenu
                    st.text_area(
                        f"Version {response['provider']}",
                        value=response['response'],
                        height=400,
                        key=f"version_{response['provider']}_{i}",
                        disabled=True
                    )
    
    # Jurisprudence utilisée
    if result.get('jurisprudence_used') and result.get('jurisprudence_references'):
        with st.expander("⚖️ Jurisprudence citée", expanded=False):
            for ref in result['jurisprudence_references']:
                st.markdown(f"- [{ref['title']}]({ref.get('url', '#')})")
                if ref.get('summary'):
                    st.caption(ref['summary'])

def create_formatted_docx(content: str, doc_type: str) -> bytes:
    """Crée un document Word avec mise en forme professionnelle complète"""
    
    if not DOCX_AVAILABLE:
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = docx.Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
        
        # Styles personnalisés pour documents juridiques
        styles = doc.styles
        
        # Style pour le titre principal
        title_style = styles.add_style('JuridicalTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # Style pour les sections principales (I., II., III.)
        heading1_style = styles.add_style('JuridicalHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.name = 'Times New Roman'
        heading1_style.font.size = Pt(14)
        heading1_style.font.bold = True
        heading1_style.paragraph_format.space_before = Pt(18)
        heading1_style.paragraph_format.space_after = Pt(12)
        
        # Style pour les sous-sections (A., B., C.)
        heading2_style = styles.add_style('JuridicalHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.name = 'Times New Roman'
        heading2_style.font.size = Pt(13)
        heading2_style.font.bold = True
        heading2_style.font.underline = True
        heading2_style.paragraph_format.space_before = Pt(12)
        heading2_style.paragraph_format.space_after = Pt(6)
        
        # Style pour le corps du texte
        body_style = styles['Normal']
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.first_line_indent = Inches(0.5)
        
        # Analyser et formater le contenu
        lines = content.split('\n')
        
        # Variables pour tracker l'état
        in_header = True
        current_section = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Ligne vide - ajouter un saut seulement si pas dans l'en-tête
                if not in_header:
                    doc.add_paragraph()
                continue
            
            # Détecter le type de ligne
            # Titre principal (tout en majuscules, court)
            if line.isupper() and len(line.split()) < 10 and in_header:
                p = doc.add_paragraph(line, style='JuridicalTitle')
                in_header = False
            
            # En-tête (POUR:, CONTRE:, etc.)
            elif line.startswith(('POUR', 'CONTRE', 'TRIBUNAL', 'AFFAIRE')) and ':' in line:
                p = doc.add_paragraph()
                parts = line.split(':', 1)
                run1 = p.add_run(parts[0] + ':')
                run1.bold = True
                if len(parts) > 1:
                    p.add_run(' ' + parts[1])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Sections principales (I., II., III., etc.)
            elif re.match(r'^[IVX]+\.\s+', line):
                p = doc.add_paragraph(line, style='JuridicalHeading1')
                current_section = line
                in_header = False
            
            # Sous-sections (A., B., C., etc.)
            elif re.match(r'^[A-Z]\.\s+', line):
                p = doc.add_paragraph(line, style='JuridicalHeading2')
            
            # Sous-sous-sections (1., 2., 3., etc.)
            elif re.match(r'^\d+\.\s+', line):
                p = doc.add_paragraph(line)
                p.style.font.bold = True
                p.style.paragraph_format.left_indent = Inches(0.5)
            
            # Points avec tirets
            elif line.startswith('-'):
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                p.style.paragraph_format.left_indent = Inches(0.5)
            
            # PAR CES MOTIFS
            elif line.startswith('PAR CES MOTIFS'):
                # Saut de page avant PAR CES MOTIFS
                doc.add_page_break()
                p = doc.add_paragraph(line, style='JuridicalHeading1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Formules de conclusion
            elif any(line.startswith(phrase) for phrase in [
                'PLAISE AU TRIBUNAL',
                'PLAISE À LA COUR',
                'IL EST DEMANDÉ',
                'SOUS TOUTES RÉSERVES'
            ]):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style.font.bold = True
            
            # Texte normal
            else:
                p = doc.add_paragraph(line, style='Normal')
                
                # Mettre en gras les références juridiques
                for run in p.runs:
                    text = run.text
                    # Articles de loi
                    if re.search(r'article\s+[LR]?\s*\d+', text, re.IGNORECASE):
                        run.font.bold = True
                    # Jurisprudence
                    elif re.search(r'(Cass\.|CA|CE|Cons\.\s*const\.)', text):
                        run.font.italic = True
        
        # Ajouter les métadonnées dans le pied de page
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} - {doc_type.title()}"
        footer_para.style.font.size = Pt(9)
        footer_para.style.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Numérotation des pages
        # Note: La numérotation automatique nécessite des manipulations XML complexes
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création document Word: {e}")
        # Fallback : retourner le texte brut
        return content.encode('utf-8')

def create_pdf_from_content(content: str, doc_type: str) -> bytes:
    """Crée un PDF à partir du contenu (nécessite reportlab ou weasyprint)"""
    # Pour l'instant, retourner une version texte
    # L'implémentation complète nécessiterait reportlab
    return content.encode('utf-8')

def enrich_current_document(content: str):
    """Enrichit le document actuel avec plus de détails"""
    st.info("🔄 Enrichissement en cours...")
    # Implémenter l'enrichissement via IA

def show_timeline_results():
    """Affiche les résultats de chronologie"""
    result = st.session_state.timeline_result
    
    st.markdown(f"### ⏱️ Chronologie des {result['type']}")
    
    # Métadonnées
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Événements", len(result['events']))
    with col2:
        st.metric("Documents analysés", result['document_count'])
    with col3:
        st.metric("Type", result['type'].title())
    
    # Visualisation
    if result.get('visualization'):
        st.plotly_chart(result['visualization'], use_container_width=True)
    else:
        # Affichage alternatif
        for event in result['events']:
            with st.container():
                col1, col2 = st.columns([1, 4])
                with col1:
                    st.write(f"**{event.get('date', 'N/A')}**")
                with col2:
                    st.write(event.get('description', ''))
                    if event.get('actors'):
                        st.caption(f"👥 {', '.join(event['actors'])}")

def show_mapping_results():
    """Affiche les résultats de cartographie"""
    result = st.session_state.mapping_result
    
    st.markdown(f"### 🗺️ Cartographie des relations - {result['type']}")
    
    # Statistiques
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Entités", result['analysis']['node_count'])
    with col2:
        st.metric("Relations", result['analysis']['edge_count'])
    with col3:
        st.metric("Densité", f"{result['analysis']['density']:.2%}")
    with col4:
        st.metric("Composantes", len(result['analysis']['components']))
    
    # Acteurs clés
    if result['analysis']['key_players']:
        st.markdown("#### 🎯 Acteurs principaux")
        for i, player in enumerate(result['analysis']['key_players'], 1):
            st.write(f"{i}. **{player}**")
    
    # Visualisation
    if result.get('visualization'):
        st.plotly_chart(result['visualization'], use_container_width=True)

def show_comparison_results():
    """Affiche les résultats de comparaison"""
    result = st.session_state.comparison_result
    
    st.markdown(f"### 🔄 Comparaison - {result['type']}")
    
    # Statistiques
    stats = result['comparison'].get('statistics', {})
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Documents", result['document_count'])
    with col2:
        st.metric("Convergences", len(result['comparison'].get('convergences', [])))
    with col3:
        st.metric("Divergences", len(result['comparison'].get('divergences', [])))
    with col4:
        reliability = stats.get('reliability_score', 0)
        color = "🟢" if reliability > 0.7 else "🟡" if reliability > 0.4 else "🔴"
        st.metric("Fiabilité", f"{color} {reliability:.0%}")
    
    # Visualisations
    if result.get('visualizations'):
        for viz_name, viz in result['visualizations'].items():
            if viz:
                st.plotly_chart(viz, use_container_width=True)
    
    # Détails textuels
    tabs = st.tabs(["Convergences", "Divergences", "Évolutions", "Analyse"])
    
    with tabs[0]:
        for conv in result['comparison'].get('convergences', []):
            st.write(f"**{conv['point']}**")
            st.write(conv['details'])
            st.divider()
    
    with tabs[1]:
        for div in result['comparison'].get('divergences', []):
            st.write(f"**{div['point']}**")
            st.write(div['details'])
            st.divider()
    
    with tabs[2]:
        for evo in result['comparison'].get('evolutions', []):
            st.write(f"**{evo.get('point', 'Évolution')}**")
            st.write(evo.get('details', ''))
            st.divider()
    
    with tabs[3]:
        st.write(result['comparison'].get('analysis', 'Pas d\'analyse disponible'))

def show_synthesis_results():
    """Affiche les résultats de synthèse"""
    result = st.session_state.synthesis_result
    
    if 'error' in result:
        st.error(f"❌ {result['error']}")
        return
    
    st.markdown("### 📝 Synthèse des documents")
    
    # Métadonnées
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Pièces analysées", result.get('piece_count', 0))
    with col2:
        st.metric("Catégories", len(result.get('categories', [])))
    with col3:
        st.metric("Généré le", result.get('timestamp', datetime.now()).strftime('%H:%M'))
    
    # Contenu de la synthèse
    st.markdown("#### 📄 Synthèse")
    
    synthesis_content = st.text_area(
        "Contenu de la synthèse",
        value=result.get('content', ''),
        height=600,
        key="synthesis_content_display"
    )
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.download_button(
            "💾 Télécharger",
            synthesis_content.encode('utf-8'),
            f"synthese_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_synthesis"
        ):
            st.success("✅ Synthèse téléchargée")
    
    with col2:
        if st.button("📊 Voir les pièces", key="view_synthesis_pieces"):
            st.session_state.show_piece_selection = True
    
    with col3:
        if st.button("🔄 Régénérer", key="regenerate_synthesis"):
            if st.session_state.get('selected_pieces'):
                synthesize_selected_pieces(st.session_state.selected_pieces)

def show_ai_analysis_results():
    """Affiche les résultats d'analyse IA"""
    results = st.session_state.ai_analysis_results
    
    if 'error' in results:
        st.error(f"❌ {results['error']}")
        return
    
    analysis_titles = {
        'risk_analysis': '⚠️ Analyse des risques',
        'compliance': '✅ Analyse de conformité',
        'strategy': '🎯 Analyse stratégique',
        'general_analysis': '🤖 Analyse générale'
    }
    
    st.markdown(f"### {analysis_titles.get(results.get('type'), '🤖 Analyse IA')}")
    
    # Métadonnées
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Documents analysés", results.get('document_count', 0))
    with col2:
        st.metric("Type", results.get('type', 'general').replace('_', ' ').title())
    with col3:
        st.metric("Généré le", results.get('timestamp', datetime.now()).strftime('%H:%M'))
    
    # Question originale
    if results.get('query'):
        st.info(f"**Question :** {results['query']}")
    
    # Contenu de l'analyse
    st.markdown("#### 📊 Résultats de l'analyse")
    
    analysis_content = st.text_area(
        "Analyse détaillée",
        value=results.get('content', ''),
        height=600,
        key="ai_analysis_content"
    )
    
    # Actions
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.download_button(
            "💾 Télécharger",
            analysis_content.encode('utf-8'),
            f"analyse_{results.get('type', 'general')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_analysis"
        ):
            st.success("✅ Analyse téléchargée")
    
    with col2:
        if st.button("📝 Convertir en rapport", key="convert_to_report"):
            convert_analysis_to_report(analysis_content, results)
    
    with col3:
        if st.button("📊 Visualiser", key="visualize_analysis"):
            create_analysis_visualization(results)
    
    with col4:
        if st.button("🔄 Approfondir", key="deepen_analysis"):
            deepen_current_analysis(results)

def show_search_results():
    """Affiche les résultats de recherche"""
    results = st.session_state.search_results
    
    st.markdown(f"### 🔍 Résultats de recherche ({len(results)} documents)")
    
    if not results:
        st.info("Aucun résultat trouvé")
        return
    
    # Options d'affichage
    col1, col2, col3 = st.columns(3)
    
    with col1:
        sort_by = st.selectbox(
            "Trier par",
            ["Pertinence", "Titre", "Date", "Source"],
            key="sort_search_results"
        )
    
    with col2:
        view_mode = st.radio(
            "Affichage",
            ["Compact", "Détaillé"],
            key="view_mode_search",
            horizontal=True
        )
    
    with col3:
        results_per_page = st.selectbox(
            "Résultats par page",
            [10, 20, 50, 100],
            key="results_per_page"
        )
    
    # Trier les résultats
    sorted_results = sort_search_results(results, sort_by)
    
    # Pagination
    total_pages = (len(sorted_results) - 1) // results_per_page + 1
    current_page = st.number_input(
        "Page",
        min_value=1,
        max_value=total_pages,
        value=1,
        key="search_page"
    )
    
    start_idx = (current_page - 1) * results_per_page
    end_idx = min(start_idx + results_per_page, len(sorted_results))
    
    # Afficher les résultats
    for i, result in enumerate(sorted_results[start_idx:end_idx], start=start_idx + 1):
        with st.container():
            if view_mode == "Compact":
                col1, col2, col3 = st.columns([4, 1, 1])
                
                with col1:
                    st.markdown(f"**{i}. {result.get('title', 'Sans titre')}**")
                    st.caption(f"Source: {result.get('source', 'N/A')}")
                
                with col2:
                    if result.get('score'):
                        st.metric("Score", f"{result['score']:.0%}")
                
                with col3:
                    if st.button("Voir", key=f"view_result_{i}"):
                        show_document_detail(result)
            
            else:  # Détaillé
                st.markdown(f"**{i}. {result.get('title', 'Sans titre')}**")
                st.caption(f"Source: {result.get('source', 'N/A')} | Score: {result.get('score', 0):.0%}")
                
                # Extrait du contenu
                content = result.get('content', '')
                if content:
                    st.text_area(
                        "Extrait",
                        value=content[:500] + '...' if len(content) > 500 else content,
                        height=150,
                        key=f"extract_{i}",
                        disabled=True
                    )
                
                # Actions
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if st.button("📖 Lire", key=f"read_{i}"):
                        show_document_detail(result)
                
                with col2:
                    if st.button("🤖 Analyser", key=f"analyze_{i}"):
                        analyze_single_document(result)
                
                with col3:
                    if st.button("📋 Sélectionner", key=f"select_{i}"):
                        add_to_selection(result)
            
            st.divider()
    
    # Navigation pages
    if total_pages > 1:
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col1:
            if current_page > 1:
                if st.button("⬅️ Précédent"):
                    st.session_state.search_page = current_page - 1
                    st.rerun()
        
        with col2:
            st.write(f"Page {current_page} / {total_pages}")
        
        with col3:
            if current_page < total_pages:
                if st.button("Suivant ➡️"):
                    st.session_state.search_page = current_page + 1
                    st.rerun()

def sort_search_results(results: list, sort_by: str) -> list:
    """Trie les résultats de recherche"""
    if sort_by == "Pertinence":
        return sorted(results, key=lambda x: x.get('score', 0), reverse=True)
    elif sort_by == "Titre":
        return sorted(results, key=lambda x: x.get('title', ''))
    elif sort_by == "Source":
        return sorted(results, key=lambda x: x.get('source', ''))
    else:  # Date
        # Essayer de parser les dates si disponibles
        return results

def show_document_detail(document: dict):
    """Affiche le détail d'un document"""
    st.session_state.current_document = document
    st.session_state.show_document_detail = True

def analyze_single_document(document: dict):
    """Lance l'analyse d'un seul document"""
    st.session_state.universal_query = f"analyser @{document.get('id', document.get('title', ''))}"
    st.rerun()

def add_to_selection(document: dict):
    """Ajoute un document à la sélection"""
    if 'selected_documents' not in st.session_state:
        st.session_state.selected_documents = []
    
    st.session_state.selected_documents.append(document)
    st.success("✅ Document ajouté à la sélection")

# === AUTRES FONCTIONS HELPER ===

def show_pieces_management_tab():
    """Affiche l'onglet de gestion des pièces"""
    st.markdown("### 📋 Gestion des pièces")
    
    if st.session_state.get('selected_pieces'):
        pieces = st.session_state.selected_pieces
        
        # Statistiques
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total pièces", len(pieces))
        with col2:
            categories = list(set(p.categorie for p in pieces))
            st.metric("Catégories", len(categories))
        with col3:
            st.metric("Pertinence moyenne", f"{sum(p.pertinence for p in pieces) / len(pieces):.0%}")
        
        # Actions groupées
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("📊 Créer bordereau", key="create_bordereau_tab"):
                process_bordereau_request("", {})
        
        with col2:
            if st.button("📝 Synthétiser", key="synthesize_pieces_tab"):
                synthesize_selected_pieces(pieces)
        
        with col3:
            if st.button("📤 Exporter liste", key="export_pieces_tab"):
                export_piece_list(pieces)
        
        with col4:
            if st.button("🗑️ Vider sélection", key="clear_selection_tab"):
                st.session_state.selected_pieces = []
                st.rerun()
        
        # Affichage des pièces
        for piece in pieces:
            with st.container():
                col1, col2, col3, col4 = st.columns([1, 3, 1, 1])
                
                with col1:
                    st.write(f"**{piece.numero}**")
                
                with col2:
                    st.write(piece.titre)
                    if piece.description:
                        st.caption(piece.description)
                
                with col3:
                    st.caption(piece.categorie)
                
                with col4:
                    if st.button("❌", key=f"remove_piece_{piece.numero}"):
                        pieces.remove(piece)
                        st.rerun()
            
            st.divider()
    
    else:
        st.info("Aucune pièce sélectionnée. Utilisez la commande `sélectionner pièces` dans la barre universelle.")

def export_piece_list(pieces: list):
    """Exporte la liste des pièces"""
    content = "LISTE DES PIÈCES SÉLECTIONNÉES\n"
    content += f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Nombre de pièces : {len(pieces)}\n\n"
    
    # Grouper par catégorie
    by_category = defaultdict(list)
    for piece in pieces:
        by_category[piece.categorie].append(piece)
    
    for category, cat_pieces in by_category.items():
        content += f"\n{category.upper()} ({len(cat_pieces)} pièces)\n"
        content += "-" * 50 + "\n"
        
        for piece in cat_pieces:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            if piece.date:
                content += f"   Date : {piece.date.strftime('%d/%m/%Y')}\n"
            content += "\n"
    
    # Télécharger
    st.download_button(
        "💾 Télécharger la liste",
        content.encode('utf-8'),
        f"liste_pieces_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
        "text/plain",
        key="download_piece_list_file"
    )

def show_jurisprudence_tab():
    """Affiche l'onglet jurisprudence"""
    # Utiliser le gestionnaire existant
    display_legal_search_interface()

def show_explorer_tab():
    """Affiche l'explorateur de fichiers"""
    st.markdown("### 📁 Explorateur de documents")
    
    # Sources disponibles
    sources = []
    
    if st.session_state.get('azure_documents'):
        sources.append("Documents locaux")
    
    if st.session_state.get('azure_blob_manager'):
        sources.append("Azure Blob Storage")
    
    if not sources:
        st.info("Aucune source de documents disponible")
        return
    
    selected_source = st.selectbox(
        "Source",
        sources,
        key="explorer_source"
    )
    
    if selected_source == "Documents locaux":
        show_local_documents_explorer()
    
    elif selected_source == "Azure Blob Storage":
        show_azure_blob_explorer()

def show_local_documents_explorer():
    """Affiche l'explorateur de documents locaux"""
    documents = st.session_state.get('azure_documents', {})
    
    if not documents:
        st.info("Aucun document local")
        return
    
    # Filtres
    col1, col2 = st.columns(2)
    
    with col1:
        search_filter = st.text_input(
            "Filtrer par nom",
            key="explorer_filter_name"
        )
    
    with col2:
        source_filter = st.selectbox(
            "Filtrer par source",
            ["Toutes"] + list(set(doc.source for doc in documents.values())),
            key="explorer_filter_source"
        )
    
    # Appliquer les filtres
    filtered_docs = {}
    for doc_id, doc in documents.items():
        if search_filter and search_filter.lower() not in doc.title.lower():
            continue
        if source_filter != "Toutes" and doc.source != source_filter:
            continue
        filtered_docs[doc_id] = doc
    
    # Afficher les documents
    st.write(f"**{len(filtered_docs)} documents**")
    
    for doc_id, doc in filtered_docs.items():
        with st.expander(f"📄 {doc.title}"):
            st.write(f"**ID :** {doc_id}")
            st.write(f"**Source :** {doc.source}")
            
            if doc.metadata:
                st.write("**Métadonnées :**")
                for key, value in doc.metadata.items():
                    st.write(f"- {key}: {value}")
            
            # Actions
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📖 Lire", key=f"read_doc_{doc_id}"):
                    st.session_state.current_document = {
                        'id': doc_id,
                        'title': doc.title,
                        'content': doc.content,
                        'source': doc.source
                    }
            
            with col2:
                if st.button("🤖 Analyser", key=f"analyze_doc_{doc_id}"):
                    st.session_state.universal_query = f"analyser @{doc_id}"
                    st.rerun()
            
            with col3:
                if st.button("🗑️ Supprimer", key=f"delete_doc_{doc_id}"):
                    del st.session_state.azure_documents[doc_id]
                    st.rerun()

def show_azure_blob_explorer():
    """Affiche l'explorateur Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if not blob_manager or not blob_manager.is_connected():
        st.warning("Azure Blob Storage non connecté")
        return
    
    try:
        containers = blob_manager.list_containers()
        
        if not containers:
            st.info("Aucun conteneur disponible")
            return
        
        selected_container = st.selectbox(
            "Conteneur",
            containers,
            key="blob_container"
        )
        
        # Explorer le conteneur
        current_path = st.session_state.get('blob_current_path', '')
        
        # Breadcrumb
        if current_path:
            path_parts = current_path.split('/')
            breadcrumb = "📁 " + " > ".join(path_parts)
            st.write(breadcrumb)
            
            if st.button("⬆️ Dossier parent"):
                st.session_state.blob_current_path = '/'.join(path_parts[:-1])
                st.rerun()
        
        # Lister le contenu
        items = blob_manager.list_folder_contents(selected_container, current_path)
        
        # Séparer dossiers et fichiers
        folders = [item for item in items if item['type'] == 'folder']
        files = [item for item in items if item['type'] == 'file']
        
        # Afficher les dossiers
        if folders:
            st.write("**📁 Dossiers**")
            for folder in folders:
                col1, col2 = st.columns([4, 1])
                
                with col1:
                    st.write(f"📁 {folder['name']}")
                
                with col2:
                    if st.button("Ouvrir", key=f"open_folder_{folder['name']}"):
                        st.session_state.blob_current_path = folder['path']
                        st.rerun()
        
        # Afficher les fichiers
        if files:
            st.write("**📄 Fichiers**")
            for file in files:
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    st.write(f"📄 {file['name']}")
                    st.caption(f"Taille: {file['size']} octets")
                
                with col2:
                    if st.button("💾", key=f"download_blob_{file['name']}"):
                        content = blob_manager.download_blob(selected_container, file['path'])
                        st.download_button(
                            "Télécharger",
                            content,
                            file['name'],
                            key=f"download_actual_{file['name']}"
                        )
                
                with col3:
                    if st.button("🤖", key=f"analyze_blob_{file['name']}"):
                        # Télécharger et analyser
                        content = blob_manager.download_blob(selected_container, file['path'])
                        # Stocker temporairement
                        store_locally(
                            file['name'],
                            content.decode('utf-8', errors='ignore'),
                            f"azure:{selected_container}/{file['path']}"
                        )
                        st.session_state.universal_query = f"analyser @{file['name']}"
                        st.rerun()
        
    except Exception as e:
        st.error(f"Erreur exploration Azure: {e}")

def show_configuration_tab():
    """Affiche l'onglet de configuration"""
    st.markdown("### ⚙️ Configuration")
    
    tabs = st.tabs(["IA", "Connexions", "Templates", "Préférences"])
    
    with tabs[0]:
        show_ai_configuration()
    
    with tabs[1]:
        show_connections_configuration()
    
    with tabs[2]:
        show_templates_configuration()
    
    with tabs[3]:
        show_preferences_configuration()

def show_ai_configuration():
    """Configuration des IA"""
    st.markdown("#### 🤖 Configuration des IA")
    
    llm_manager = MultiLLMManager()
    
    # État des connexions
    st.write("**État des connexions :**")
    
    for provider in LLMProvider:
        col1, col2 = st.columns([3, 1])
        
        with col1:
            is_connected = provider in llm_manager.clients
            status = "✅ Connecté" if is_connected else "❌ Non connecté"
            st.write(f"{provider.value}: {status}")
        
        with col2:
            if is_connected:
                if st.button(f"Tester", key=f"test_{provider.value}"):
                    test_llm_connection(provider)
    
    # Préférences par défaut
    st.write("**Préférences par défaut :**")
    
    default_providers = st.multiselect(
        "IA par défaut pour la rédaction",
        [p.value for p in LLMProvider if p in llm_manager.clients],
        default=[p.value for p in list(llm_manager.clients.keys())[:2]] if llm_manager.clients else [],
        key="default_redaction_providers"
    )
    
    fusion_mode = st.selectbox(
        "Mode de fusion par défaut",
        ["🎯 Fusion intelligente", "📋 Comparaison côte à côte", "🔗 Synthèse enrichie", "⚡ Meilleure version"],
        key="default_fusion_mode"
    )

def test_llm_connection(provider: LLMProvider):
    """Teste la connexion à une IA"""
    llm_manager = MultiLLMManager()
    
    with st.spinner(f"Test de {provider.value}..."):
        try:
            response = llm_manager.query_single_llm(
                provider,
                "Réponds simplement 'OK' si tu me reçois.",
                "Assistant de test"
            )
            
            if response['success']:
                st.success(f"✅ {provider.value} fonctionne correctement")
            else:
                st.error(f"❌ Erreur: {response.get('error', 'Inconnue')}")
                
        except Exception as e:
            st.error(f"❌ Erreur connexion: {str(e)}")

def show_connections_configuration():
    """Configuration des connexions externes"""
    st.markdown("#### 🔌 Connexions externes")
    
    # Azure Blob Storage
    st.write("**Azure Blob Storage**")
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        st.success("✅ Connecté")
        if st.button("🔄 Rafraîchir la connexion"):
            blob_manager.connect()
    else:
        st.warning("❌ Non connecté")
        st.info("Configurez les variables d'environnement Azure")
    
    # Azure Search
    st.write("**Azure Cognitive Search**")
    search_manager = st.session_state.get('azure_search_manager')
    
    if search_manager and search_manager.is_connected():
        st.success("✅ Connecté")
    else:
        st.warning("❌ Non connecté")
    
    # Email SMTP
    st.write("**Configuration Email**")
    
    smtp_configured = bool(st.secrets.get('smtp_username'))
    if smtp_configured:
        st.success("✅ Email configuré")
    else:
        st.warning("❌ Email non configuré")
        st.info("Ajoutez les paramètres SMTP dans les secrets")

def show_templates_configuration():
    """Configuration des templates"""
    st.markdown("#### 📄 Gestion des templates")
    
    # Templates intégrés
    st.write("**Templates intégrés :**")
    for key, template in DOCUMENT_TEMPLATES.items():
        st.write(f"- {template['name']} ({template.get('style', 'N/A')})")
    
    # Templates personnalisés
    custom_templates = st.session_state.get('saved_templates', {})
    
    if custom_templates:
        st.write(f"**Templates personnalisés ({len(custom_templates)}) :**")
        
        for key, template in custom_templates.items():
            col1, col2, col3 = st.columns([3, 1, 1])
            
            with col1:
                st.write(f"- {template.get('name', key)}")
            
            with col2:
                if st.button("✏️", key=f"edit_template_{key}"):
                    st.session_state.template_to_edit = key
                    st.session_state.universal_query = "modifier template"
            
            with col3:
                if st.button("🗑️", key=f"delete_template_{key}"):
                    del st.session_state.saved_templates[key]
                    st.rerun()
    
    # Import/Export
    st.divider()
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📤 Exporter tous les templates"):
            save_templates_to_file()
    
    with col2:
        uploaded_file = st.file_uploader(
            "📥 Importer des templates",
            type=['json'],
            key="import_templates"
        )
        
        if uploaded_file:
            try:
                import json
                templates_data = json.load(uploaded_file)
                
                if 'custom' in templates_data:
                    st.session_state.saved_templates.update(templates_data['custom'])
                    st.success(f"✅ {len(templates_data['custom'])} templates importés")
                    st.rerun()
                    
            except Exception as e:
                st.error(f"Erreur import: {e}")

def show_preferences_configuration():
    """Configuration des préférences utilisateur"""
    st.markdown("#### 🎨 Préférences")
    
    # Préférences d'affichage
    st.write("**Affichage :**")
    
    results_per_page = st.selectbox(
        "Résultats par page par défaut",
        [10, 20, 50, 100],
        index=1,
        key="pref_results_per_page"
    )
    
    default_view = st.radio(
        "Vue par défaut",
        ["Compact", "Détaillé"],
        key="pref_default_view",
        horizontal=True
    )
    
    # Préférences de rédaction
    st.write("**Rédaction :**")
    
    auto_jurisprudence = st.checkbox(
        "Recherche automatique de jurisprudence",
        value=True,
        key="pref_auto_juris"
    )
    
    create_hyperlinks = st.checkbox(
        "Créer des liens hypertextes automatiques",
        value=True,
        key="pref_hyperlinks"
    )
    
    default_doc_length = st.select_slider(
        "Longueur par défaut des documents",
        options=["Concis", "Standard", "Détaillé", "Très détaillé", "Exhaustif"],
        value="Très détaillé",
        key="pref_doc_length"
    )
    
    # Sauvegarder les préférences
    if st.button("💾 Sauvegarder les préférences"):
        preferences = {
            'results_per_page': results_per_page,
            'default_view': default_view,
            'auto_jurisprudence': auto_jurisprudence,
            'create_hyperlinks': create_hyperlinks,
            'default_doc_length': default_doc_length
        }
        
        st.session_state.user_preferences = preferences
        st.success("✅ Préférences sauvegardées")

# Fonctions helper additionnelles

def clear_universal_state():
    """Efface l'état de l'interface universelle"""
    keys_to_clear = [
        'universal_query', 'last_universal_query', 'current_analysis',
        'redaction_result', 'timeline_result', 'mapping_result',
        'comparison_result', 'synthesis_result', 'ai_analysis_results',
        'search_results', 'selected_pieces', 'import_files',
        'plaidoirie_result', 'preparation_client_result'
    ]
    
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    
    st.success("✅ Interface réinitialisée")
    st.rerun()

def save_current_work():
    """Sauvegarde le travail actuel"""
    work_data = {
        'timestamp': datetime.now().isoformat(),
        'query': st.session_state.get('universal_query', ''),
        'analysis': st.session_state.get('current_analysis', {}),
        'results': {}
    }
    
    # Collecter tous les résultats
    result_keys = [
        'redaction_result', 'timeline_result', 'mapping_result',
        'comparison_result', 'synthesis_result', 'ai_analysis_results',
        'search_results'
    ]
    
    for key in result_keys:
        if key in st.session_state:
            work_data['results'][key] = st.session_state[key]
    
    # Sauvegarder
    import json
    
    json_str = json.dumps(work_data, indent=2, ensure_ascii=False, default=str)
    
    st.download_button(
        "💾 Télécharger la sauvegarde",
        json_str,
        f"sauvegarde_travail_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        "application/json",
        key="download_work_save"
    )

def share_current_work():
    """Partage le travail actuel"""
    # Créer un lien de partage ou export
    st.info("🔗 Fonctionnalité de partage à implémenter")
    
    # Pour l'instant, proposer l'export
    save_current_work()

def prepare_document_email(content: str, doc_type: str):
    """Prépare l'envoi d'un document par email"""
    st.session_state.email_document = {
        'content': content,
        'type': doc_type
    }
    
    st.session_state.universal_query = f"envoyer {doc_type} par email"
    st.rerun()

def verify_document_content(content: str):
    """Vérifie le contenu d'un document"""
    # Vérifications basiques
    issues = []
    
    # Longueur
    if len(content) < 1000:
        issues.append("⚠️ Document très court (< 1000 caractères)")
    
    # Références juridiques
    if not re.search(r'article\s+\d+', content, re.IGNORECASE):
        issues.append("⚠️ Aucune référence d'article de loi")
    
    # Structure
    if not any(marker in content for marker in ['I.', 'II.', 'A.', 'B.']):
        issues.append("⚠️ Structure peu claire (pas de sections numérotées)")
    
    # Afficher le résultat
    if issues:
        st.warning("Points d'attention :")
        for issue in issues:
            st.write(issue)
    else:
        st.success("✅ Document bien structuré")

def show_document_statistics(content: str):
    """Affiche les statistiques d'un document"""
    
    # Calculs
    words = content.split()
    sentences = content.split('.')
    paragraphs = content.split('\n\n')
    
    # Références
    law_refs = len(re.findall(r'article\s+[LR]?\s*\d+', content, re.IGNORECASE))
    juris_refs = len(re.findall(r'(Cass\.|CA|CE|Cons\.\s*const\.)', content))
    
    # Affichage
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Mots", f"{len(words):,}")
        st.metric("Phrases", f"{len(sentences):,}")
    
    with col2:
        st.metric("Paragraphes", len(paragraphs))
        st.metric("Mots/phrase", f"{len(words) / max(len(sentences), 1):.1f}")
    
    with col3:
        st.metric("Articles cités", law_refs)
        st.metric("Jurisprudences", juris_refs)
    
    with col4:
        # Complexité (basique)
        avg_word_length = sum(len(w) for w in words) / max(len(words), 1)
        complexity = "Simple" if avg_word_length < 5 else "Moyen" if avg_word_length < 7 else "Complexe"
        st.metric("Complexité", complexity)
        st.metric("Longueur moy.", f"{avg_word_length:.1f} car/mot")

def create_document_summary(content: str):
    """Crée un résumé du document"""
    llm_manager = MultiLLMManager()
    
    if not llm_manager.clients:
        st.error("Aucune IA disponible pour le résumé")
        return
    
    with st.spinner("Création du résumé..."):
        summary_prompt = f"""Crée un résumé exécutif de ce document juridique.

Le résumé doit contenir :
1. Type et objet du document
2. Points clés (3-5 points)
3. Demandes principales
4. Enjeux juridiques

Document :
{content[:5000]}...

Résumé en 200-300 mots maximum."""
        
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            summary_prompt,
            "Tu es un expert en synthèse juridique."
        )
        
        if response['success']:
            st.markdown("### 📋 Résumé exécutif")
            st.write(response['response'])
            
            # Proposer de sauvegarder
            if st.button("💾 Sauvegarder le résumé"):
                st.session_state.document_summary = response['response']

def convert_analysis_to_report(analysis_content: str, analysis_metadata: dict):
    """Convertit une analyse en rapport formel"""
    st.info("📝 Conversion en rapport formel...")
    
    # Créer une requête de rédaction de rapport
    report_query = f"rédiger rapport formel basé sur analyse {analysis_metadata.get('type', 'générale')}"
    
    # Injecter le contenu de l'analyse
    st.session_state.analysis_to_report = analysis_content
    
# Fonctions helper pour plaidoirie

def detect_document_type(title: str, content: str) -> str:
    """Détecte le type d'un document"""
    title_lower = title.lower()
    content_preview = content[:1000].lower()
    
    if any(kw in title_lower for kw in ['audition', 'interrogatoire', 'garde à vue']):
        return 'temoignage'
    elif any(kw in title_lower for kw in ['expertise', 'expert', 'rapport']):
        return 'expertise'
    elif any(kw in title_lower for kw in ['pv', 'procès-verbal', 'constat']):
        return 'procedure'
    else:
        return 'piece'

def show_rehearsal_mode(content: str):
    """Mode répétition pour la plaidoirie"""
    st.markdown("### 🎬 Mode répétition")
    
    # Diviser en sections
    sections = re.split(r'\n\n+', content)
    
    # Navigation entre sections
    if 'rehearsal_section' not in st.session_state:
        st.session_state.rehearsal_section = 0
    
    current_section = st.session_state.rehearsal_section
    
    # Afficher la section courante
    st.info(f"Section {current_section + 1} / {len(sections)}")
    
    st.text_area(
        "Section actuelle",
        value=sections[current_section] if current_section < len(sections) else "",
        height=300,
        key="rehearsal_display"
    )
    
    # Navigation
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("⬅️ Précédent", disabled=current_section == 0):
            st.session_state.rehearsal_section -= 1
            st.rerun()
    
    with col2:
        st.write(f"Progression : {((current_section + 1) / len(sections) * 100):.0f}%")
    
    with col3:
        if st.button("Suivant ➡️", disabled=current_section >= len(sections) - 1):
            st.session_state.rehearsal_section += 1
            st.rerun()
    
    # Conseils
    st.markdown("""
    **💡 Conseils répétition :**
    - Répétez debout, comme en audience
    - Enregistrez-vous pour écouter le rythme
    - Marquez physiquement les [pauses]
    - Pratiquez les transitions entre sections
    """)

def show_plaidoirie_statistics(content: str):
    """Affiche les statistiques de la plaidoirie"""
    st.markdown("### 📊 Statistiques de la plaidoirie")
    
    # Calculs
    clean_content = re.sub(r'\[.*?\]', '', content)
    words = clean_content.split()
    sentences = clean_content.split('.')
    
    # Sections
    sections = re.findall(r'^[IVX]+\.\s+.*', content, re.MULTILINE)
    subsections = re.findall(r'^[A-Z]\.\s+.*', content, re.MULTILINE)
    
    # Affichage
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Mots totaux", f"{len(words):,}")
        st.metric("Phrases", len(sentences))
        st.metric("Mots/minute", f"{len(words) / 150:.0f} min")
    
    with col2:
        st.metric("Sections principales", len(sections))
        st.metric("Sous-sections", len(subsections))
        st.metric("Longueur moyenne phrase", f"{len(words) / max(len(sentences), 1):.1f} mots")
    
    with col3:
        annotations = len(re.findall(r'\[.*?\]', content))
        st.metric("Annotations orales", annotations)
        st.metric("Questions rhétoriques", content.count('?'))
        st.metric("Points d'exclamation", content.count('!'))
    
    # Graphique de répartition
    if sections:
        st.markdown("### 📈 Répartition du contenu")
        section_lengths = []
        
        for i, section in enumerate(sections):
            # Trouver le contenu de cette section
            start = content.find(section)
            if i < len(sections) - 1:
                end = content.find(sections[i + 1])
            else:
                end = len(content)
            
            section_content = content[start:end]
            section_lengths.append(len(section_content.split()))
        
        # Afficher la répartition
        for i, (section, length) in enumerate(zip(sections, section_lengths)):
            percentage = (length / len(words)) * 100
            st.write(f"{section[:30]}... : {percentage:.1f}% ({length} mots)")
            st.progress(percentage / 100)

def create_plaidoirie_mindmap(content: str):
    """Crée une carte mentale de la plaidoirie"""
    st.markdown("### 🗺️ Structure en carte mentale")
    
    # Extraire la structure
    structure = extract_plaidoirie_structure(content)
    
    # Afficher sous forme textuelle (sans bibliothèque de mindmap)
    mindmap_text = "PLAIDOIRIE\n"
    
    for section, subsections in structure.items():
        mindmap_text += f"├── {section}\n"
        for i, subsection in enumerate(subsections):
            is_last = i == len(subsections) - 1
            mindmap_text += f"│   {'└' if is_last else '├'}── {subsection[:50]}...\n"
        if subsections:
            mindmap_text += "│\n"
    
    st.text(mindmap_text)
    
    # Export
    st.download_button(
        "💾 Télécharger structure",
        mindmap_text.encode('utf-8'),
        "structure_plaidoirie.txt",
        "text/plain",
        key="download_mindmap"
    )

def extract_plaidoirie_structure(content: str) -> dict:
    """Extrait la structure hiérarchique de la plaidoirie"""
    structure = {}
    
    # Pattern pour les sections principales
    sections = re.findall(r'^[IVX]+\.\s+.*', content, re.MULTILINE)
    
    for section in sections:
        # Trouver le contenu de cette section
        start = content.find(section)
        
        # Trouver la prochaine section ou la fin
        next_section_match = re.search(
            r'^[IVX]+\.\s+',
            content[start + len(section):],
            re.MULTILINE
        )
        
        if next_section_match:
            end = start + len(section) + next_section_match.start()
        else:
            end = len(content)
        
        section_content = content[start:end]
        
        # Extraire les sous-sections
        subsections = re.findall(r'^([A-Z]\.\s+.*)', section_content, re.MULTILINE)
        
        structure[section] = subsections
    
    return structure

def create_plaidoirie_flashcards(result: dict):
    """Crée des fiches aide-mémoire pour la plaidoirie"""
    st.markdown("### 📇 Fiches mémo plaidoirie")
    
    # Créer des fiches par section
    flashcards = []
    
    # Points clés
    for i, point in enumerate(result.get('key_points', []), 1):
        flashcards.append({
            'title': f"Point clé {i}",
            'content': point,
            'type': 'key_point'
        })
    
    # Transitions
    transitions = re.findall(r'\[transition\](.*?)(?=\[|$)', result['content'])
    for i, trans in enumerate(transitions, 1):
        flashcards.append({
            'title': f"Transition {i}",
            'content': trans.strip(),
            'type': 'transition'
        })
    
    # Afficher les fiches
    for card in flashcards:
        with st.expander(f"{card['title']} ({card['type']})"):
            st.write(card['content'])

# Fonctions helper pour préparation client

def show_exercise_timer():
    """Timer pour les exercices de préparation"""
    st.markdown("### ⏱️ Chronomètre exercices")
    
    exercise_durations = {
        "Présentation personnelle": 60,
        "Récit des faits": 180,
        "Questions rapides": 120,
        "Gestion du silence": 30,
        "Reformulation": 60
    }
    
    selected_exercise = st.selectbox(
        "Choisir l'exercice",
        list(exercise_durations.keys()),
        key="exercise_select"
    )
    
    duration = exercise_durations[selected_exercise]
    
    st.info(f"Durée recommandée : {duration} secondes")
    
    # Instructions par exercice
    instructions = {
        "Présentation personnelle": "Présentez-vous en 1 minute : identité, profession, situation familiale",
        "Récit des faits": "Racontez les faits de manière chronologique en 3 minutes",
        "Questions rapides": "Répondez rapidement à des questions simples pendant 2 minutes",
        "Gestion du silence": "Restez calme pendant 30 secondes de silence complet",
        "Reformulation": "Reformulez une question complexe en vos propres mots"
    }
    
    st.write(f"**Instructions :** {instructions[selected_exercise]}")

def create_preparation_summary(result: dict):
    """Crée une fiche résumé de la préparation"""
    st.markdown("### 📋 Fiche résumé")
    
    summary = f"""FICHE RÉSUMÉ - PRÉPARATION {result['config']['prep_type'].upper()}
{'=' * 60}

PROFIL CLIENT : {result['config']['profil_client']}
STRATÉGIE : {result['config']['strategie']}
DURÉE PRÉPARATION : {result['stats']['duree_estimee']}

POINTS CLÉS À RETENIR :
"""
    
    # Extraire les points essentiels
    key_qa = extract_key_qa(result['content'])
    
    for i, qa in enumerate(key_qa[:15], 1):  # Top 15
        summary += f"\n{i}. Q: {qa['question']}"
        summary += f"\n   R: {qa['answer']}\n"
    
    # Phrases à mémoriser absolument
    summary += "\nPHRASES CLÉS :\n"
    key_phrases = re.findall(r'"([^"]{20,100})"', result['content'])[:5]
    
    for phrase in key_phrases:
        summary += f'• "{phrase}"\n'
    
    # Ne jamais dire
    summary += "\n⚠️ NE JAMAIS DIRE :\n"
    never_say = extract_never_say(result['content'])
    
    for item in never_say[:5]:
        summary += f"• {item}\n"
    
    # Afficher et télécharger
    st.text_area(
        "Résumé à imprimer",
        value=summary,
        height=600,
        key="prep_summary_display"
    )
    
    st.download_button(
        "💾 Télécharger résumé",
        summary.encode('utf-8'),
        f"resume_preparation_{datetime.now().strftime('%Y%m%d')}.txt",
        "text/plain",
        key="download_prep_summary"
    )

def extract_key_qa(content: str) -> list:
    """Extrait les questions-réponses clés"""
    qa_list = []
    
    # Pattern pour Q/R
    qa_pattern = r'(?:Question|Q)\s*:\s*([^\n]+).*?(?:Réponse|R)\s*:\s*([^\n]+)'
    
    matches = re.findall(qa_pattern, content, re.DOTALL | re.IGNORECASE)
    
    for q, a in matches:
        qa_list.append({
            'question': q.strip(),
            'answer': a.strip()
        })
    
    return qa_list

def extract_never_say(content: str) -> list:
    """Extrait les choses à ne jamais dire"""
    never_list = []
    
    # Chercher la section
    section = extract_section(content, "NE JAMAIS")
    
    if section:
        # Extraire les éléments
        items = re.findall(r'[-•]\s*([^\n]+)', section)
        never_list.extend(items)
    
    # Chercher aussi les patterns "ne pas dire"
    patterns = re.findall(r'[Nn]e (?:pas|jamais) dire[:\s]*([^\n]+)', content)
    never_list.extend(patterns)
    
    return list(set(never_list))[:10]

def show_interrogation_simulation(result: dict):
    """Mode simulation d'interrogatoire"""
    st.markdown("### 🎮 Simulation d'interrogatoire")
    
    # Extraire les questions
    questions = extract_key_qa(result['content'])
    
    if not questions:
        st.warning("Aucune question trouvée pour la simulation")
        return
    
    # État de la simulation
    if 'simulation_index' not in st.session_state:
        st.session_state.simulation_index = 0
        st.session_state.simulation_score = []
    
    current_q = st.session_state.simulation_index
    
    if current_q < len(questions):
        # Question actuelle
        st.info(f"Question {current_q + 1}/{len(questions)}")
        st.subheader(questions[current_q]['question'])
        
        # Réponse de l'utilisateur
        user_answer = st.text_area(
            "Votre réponse :",
            height=150,
            key=f"sim_answer_{current_q}"
        )
        
        # Afficher la réponse suggérée
        with st.expander("Voir la réponse suggérée"):
            st.success(questions[current_q]['answer'])
        
        # Évaluation
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("✅ Bonne réponse", key=f"good_{current_q}"):
                st.session_state.simulation_score.append(1)
                st.session_state.simulation_index += 1
                st.rerun()
        
        with col2:
            if st.button("😐 Moyenne", key=f"medium_{current_q}"):
                st.session_state.simulation_score.append(0.5)
                st.session_state.simulation_index += 1
                st.rerun()
        
        with col3:
            if st.button("❌ À retravailler", key=f"bad_{current_q}"):
                st.session_state.simulation_score.append(0)
                st.session_state.simulation_index += 1
                st.rerun()
    
    else:
        # Fin de simulation
        st.success("🎉 Simulation terminée !")
        
        # Score
        total_score = sum(st.session_state.simulation_score)
        max_score = len(st.session_state.simulation_score)
        percentage = (total_score / max_score * 100) if max_score > 0 else 0
        
        st.metric("Score global", f"{percentage:.0f}%")
        
        # Recommandations
        if percentage >= 80:
            st.success("Excellente préparation ! Vous êtes prêt.")
        elif percentage >= 60:
            st.warning("Bonne base, mais revoyez les questions ratées.")
        else:
            st.error("Préparation insuffisante. Refaites les exercices.")
        
        # Recommencer
        if st.button("🔄 Recommencer la simulation"):
            st.session_state.simulation_index = 0
            st.session_state.simulation_score = []
            st.rerun()

def extract_best_introduction(responses: list) -> str:
    """Extrait la meilleure introduction parmi les réponses"""
    best_intro = ""
    best_score = 0
    
    for response in responses:
        if response['success']:
            # Extraire l'introduction (premiers paragraphes)
            lines = response['response'].split('\n')
            intro_lines = []
            
            for line in lines:
                if line.strip():
                    intro_lines.append(line)
                if len(intro_lines) > 5:  # Environ 5 lignes d'intro
                    break
            
            intro = '\n'.join(intro_lines)
            
            # Scorer l'introduction
            score = 0
            if any(word in intro.lower() for word in ['messieurs', 'monsieur le président', 'mesdames']):
                score += 1
            if len(intro) > 200:
                score += 1
            if '?' in intro:  # Question rhétorique
                score += 1
            
            if score > best_score:
                best_score = score
                best_intro = intro
    
    return best_intro

def extract_best_development(responses: list) -> str:
    """Extrait le meilleur développement"""
    # Logique similaire pour le corps principal
    developments = []
    
    for response in responses:
        if response['success']:
            # Extraire la partie centrale
            content = response['response']
            # Enlever intro et conclusion approximativement
            lines = content.split('\n')
            
            # Chercher le début du développement
            dev_start = 0
            for i, line in enumerate(lines):
                if re.match(r'^[IVX]+\.', line.strip()):
                    dev_start = i
                    break
            
            # Chercher la fin (PAR CES MOTIFS ou équivalent)
            dev_end = len(lines)
            for i, line in enumerate(lines[dev_start:], dev_start):
                if 'PAR CES MOTIFS' in line or 'En conclusion' in line:
                    dev_end = i
                    break
            
            development = '\n'.join(lines[dev_start:dev_end])
            developments.append(development)
    
    # Retourner le plus complet
    return max(developments, key=len) if developments else ""

def extract_best_conclusion(responses: list) -> str:
    """Extrait la meilleure conclusion"""
    conclusions = []
    
    for response in responses:
        if response['success']:
            # Chercher la conclusion
            content = response['response']
            
            # Patterns de début de conclusion
            conclusion_start = None
            for pattern in ['PAR CES MOTIFS', 'En conclusion', 'Pour conclure', 'En définitive']:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    conclusion_start = match.start()
                    break
            
            if conclusion_start:
                conclusion = content[conclusion_start:]
                conclusions.append(conclusion)
    
    # Retourner la plus percutante (souvent la plus courte et directe)
    if conclusions:
        return min(conclusions, key=lambda x: len(x) if len(x) > 100 else float('inf'))
    return ""

def extract_key_phrases(responses: list) -> list:
    """Extrait les phrases clés de toutes les versions"""
    key_phrases = []
    
    for response in responses:
        if response['success']:
            # Chercher les phrases percutantes
            # Phrases avec points d'exclamation
            exclamations = re.findall(r'[^.!?]*!', response['response'])
            key_phrases.extend(exclamations)
            
            # Questions rhétoriques
            questions = re.findall(r'[^.!?]*\?', response['response'])
            key_phrases.extend(questions)
            
            # Phrases avec "jamais", "toujours", etc.
            strong_words = ['jamais', 'toujours', 'innocent', 'justice', 'vérité', 'preuve']
            for word in strong_words:
                pattern = rf'[^.!?]*\b{word}\b[^.!?]*[.!?]'
                matches = re.findall(pattern, response['response'], re.IGNORECASE)
                key_phrases.extend(matches)
    
    # Dédupliquer et nettoyer
    unique_phrases = []
    for phrase in key_phrases:
        cleaned = phrase.strip()
        if cleaned and len(cleaned) > 20 and cleaned not in unique_phrases:
            unique_phrases.append(cleaned)
    
    return unique_phrases[:20]  # Top 20
# === FONCTIONS MANQUANTES À AJOUTER ===

def format_ai_analysis_for_export(analysis_results: dict) -> str:
    """Formate les résultats d'analyse IA pour l'export"""
    content = f"ANALYSE IA - {analysis_results.get('type', 'Générale').upper()}\n"
    content += f"Date : {analysis_results.get('timestamp', datetime.now()).strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Documents analysés : {analysis_results.get('document_count', 0)}\n\n"
    
    if analysis_results.get('query'):
        content += f"QUESTION : {analysis_results['query']}\n\n"
    
    content += "ANALYSE :\n"
    content += analysis_results.get('content', 'Aucun contenu')
    
    return content

def format_search_results_for_export(search_results: list) -> str:
    """Formate les résultats de recherche pour l'export"""
    content = f"RÉSULTATS DE RECHERCHE\n"
    content += f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Nombre de résultats : {len(search_results)}\n\n"
    
    for i, result in enumerate(search_results, 1):
        content += f"{i}. {result.get('title', 'Sans titre')}\n"
        content += f"   Source : {result.get('source', 'N/A')}\n"
        content += f"   Score : {result.get('score', 0):.2%}\n"
        if result.get('content'):
            content += f"   Extrait : {result['content'][:200]}...\n"
        content += "\n"
    
    return content

def deepen_current_analysis(results: dict):
    """Approfondit l'analyse actuelle"""
    st.info("🔄 Approfondissement de l'analyse en cours...")
    query = results.get('query', '')
    deeper_query = f"{query} - analyse approfondie avec plus de détails"
    st.session_state.universal_query = deeper_query
    st.rerun()

def create_analysis_visualization(results: dict):
    """Crée une visualisation de l'analyse"""
    st.info("📊 Création de visualisations...")
    if results.get('type') == 'risk_analysis':
        st.warning("Visualisation des risques à implémenter")
    else:
        st.info("Visualisation générale à implémenter")

def analyze_compliance(documents: list, query: str) -> dict:
    """Analyse de conformité"""
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    compliance_prompt = f"""Analyse la conformité réglementaire dans ces documents.

DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}

QUESTION: {query}

Vérifie:
1. CONFORMITÉ LÉGALE
2. OBLIGATIONS RÉGLEMENTAIRES
3. ÉCARTS IDENTIFIÉS
4. RECOMMANDATIONS

Format structuré."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            compliance_prompt,
            "Tu es un expert en conformité juridique."
        )
        
        if response['success']:
            return {
                'type': 'compliance',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

def analyze_strategy(documents: list, query: str) -> dict:
    """Analyse stratégique"""
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    strategy_prompt = f"""Analyse stratégique de ces documents.

DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}

QUESTION: {query}

Développe:
1. ANALYSE DE SITUATION
2. OPTIONS STRATÉGIQUES
3. RECOMMANDATIONS
4. PLAN D'ACTION

Format professionnel."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            strategy_prompt,
            "Tu es un stratège juridique expert."
        )
        
        if response['success']:
            return {
                'type': 'strategy',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

def show_plaidoirie_results():
    """Affiche les résultats de plaidoirie"""
    result = st.session_state.plaidoirie_result
    
    st.markdown("### 🎤 Plaidoirie générée")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Type", result.get('type', 'Plaidoirie').title())
    with col2:
        st.metric("Durée estimée", result.get('duree_estimee', 'N/A'))
    with col3:
        st.metric("Points clés", len(result.get('key_points', [])))
    
    st.text_area(
        "Plaidoirie",
        value=result.get('content', ''),
        height=600,
        key="plaidoirie_display"
    )

def show_preparation_client_results():
    """Affiche les résultats de préparation client"""
    result = st.session_state.preparation_client_result
    
    st.markdown("### 👥 Préparation client")
    
    config = result.get('config', {})
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Type", config.get('prep_type', 'N/A'))
    with col2:
        st.metric("Profil client", config.get('profil_client', 'N/A'))
    with col3:
        st.metric("Durée", result.get('stats', {}).get('duree_estimee', 'N/A'))
    
    st.text_area(
        "Guide de préparation",
        value=result.get('content', ''),
        height=600,
        key="preparation_display"
    )

def extract_section(content: str, section_name: str) -> str:
    """Extrait une section spécifique du contenu"""
    pattern = rf'{section_name}[:\s]*\n(.*?)(?=\n[A-Z][A-Z\s]+:|$)'
    match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
    
    if match:
        return match.group(1).strip()
    return ""

def process_universal_query(query: str):
    """Traite une requête universelle - fonction temporaire"""
    st.info(f"Traitement de la requête : {query}")
    # Cette fonction sera remplacée par l'import correct plus tard