"""Module pour les fonctions d'import de documents"""

import streamlit as st
from datetime import datetime
from utils.helpers import clean_key
from models.dataclasses import Document

def process_import_request(query: str, analysis: dict):
    """Traite une demande d'import"""
    
    st.info("📥 Interface d'import activée")
    
    # Configuration de l'import depuis la config
    file_types = analysis.get('details', {}).get('file_types', ['pdf', 'docx', 'txt'])
    destination = st.session_state.get('import_destination', 'Documents locaux')
    
    # Interface d'upload
    uploaded_files = st.file_uploader(
        "Sélectionnez les fichiers à importer",
        type=['pdf', 'docx', 'txt', 'doc', 'rtf'],
        accept_multiple_files=True,
        key="import_files"
    )
    
    # Options d'import
    col1, col2 = st.columns(2)
    
    with col1:
        destination = st.selectbox(
            "Destination",
            ["Documents locaux", "Azure Blob Storage"],
            key="import_destination"
        )
    
    with col2:
        auto_analyze = st.checkbox(
            "Analyser automatiquement",
            value=True,
            key="auto_analyze_import"
        )
    
    # Si des fichiers sont uploadés
    if uploaded_files:
        if st.button("📥 Importer", key="process_import"):
            with st.spinner("📥 Import des fichiers en cours..."):
                imported_count = 0
                
                for file in uploaded_files:
                    try:
                        # Lire le contenu
                        content = file.read()
                        
                        # Traiter selon le type
                        if file.name.endswith('.pdf'):
                            text_content = extract_text_from_pdf(content)
                        elif file.name.endswith('.docx'):
                            text_content = extract_text_from_docx(content)
                        else:
                            text_content = content.decode('utf-8', errors='ignore')
                        
                        # Stocker selon la destination
                        if destination == 'Azure Blob Storage':
                            store_in_azure_blob(file.name, content, analysis.get('reference'))
                        else:
                            store_locally(file.name, text_content, analysis.get('reference'))
                        
                        imported_count += 1
                        
                    except Exception as e:
                        st.error(f"❌ Erreur import {file.name}: {str(e)}")
                
                st.success(f"✅ {imported_count} fichiers importés avec succès")
                
                # Analyse automatique si demandée
                if auto_analyze:
                    st.info("🤖 Analyse automatique des documents importés...")
                    analyze_imported_documents()

def extract_text_from_pdf(content: bytes) -> str:
    """Extrait le texte d'un PDF"""
    try:
        import PyPDF2
        import io
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text
        
    except ImportError:
        return "PDF - Installation de PyPDF2 requise pour l'extraction"
    except Exception as e:
        return f"Erreur extraction PDF: {str(e)}"

def extract_text_from_docx(content: bytes) -> str:
    """Extrait le texte d'un DOCX"""
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])
    except ImportError:
        return "DOCX - Installation de python-docx requise"
    except Exception as e:
        return f"Erreur extraction DOCX: {str(e)}"

def store_locally(filename: str, content: str, reference: str = None):
    """Stocke un document localement"""
    doc_id = f"imported_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{clean_key(filename)}"
    
    if 'azure_documents' not in st.session_state:
        st.session_state.azure_documents = {}
    
    st.session_state.azure_documents[doc_id] = Document(
        id=doc_id,
        title=filename,
        content=content,
        source=f"Import {reference if reference else 'direct'}",
        metadata={
            'imported_at': datetime.now().isoformat(),
            'reference': reference,
            'original_filename': filename
        }
    )

def store_in_azure_blob(filename: str, content: bytes, reference: str = None):
    """Stocke un document dans Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        container = reference if reference else 'imports'
        blob_name = f"{datetime.now().strftime('%Y%m%d')}_{filename}"
        
        try:
            blob_manager.upload_blob(container, blob_name, content)
            st.success(f"✅ {filename} uploadé dans Azure")
        except Exception as e:
            st.error(f"❌ Erreur Azure: {str(e)}")
    else:
        st.warning("⚠️ Azure Blob non connecté - Stockage local")
        store_locally(filename, content.decode('utf-8', errors='ignore'), reference)

def analyze_imported_documents():
    """Analyse automatique des documents importés"""
    # Récupérer les derniers documents importés
    recent_docs = []
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if 'imported_at' in doc.metadata:
            import_time = datetime.fromisoformat(doc.metadata['imported_at'])
            if (datetime.now() - import_time).seconds < 300:  # Moins de 5 minutes
                recent_docs.append(doc)
    
    if recent_docs:
        # Lancer une analyse automatique
        from managers.multi_llm_manager import MultiLLMManager
        
        analysis_prompt = f"""Analyse ces {len(recent_docs)} documents récemment importés.
        
Identifie :
1. Type de document
2. Parties impliquées
3. Dates importantes
4. Points juridiques clés
5. Risques identifiés

Documents :
{chr(10).join([f"- {doc.title}: {doc.content[:500]}..." for doc in recent_docs[:5]])}"""
        
        llm_manager = MultiLLMManager()
        if llm_manager.clients:
            provider = list(llm_manager.clients.keys())[0]
            response = llm_manager.query_single_llm(
                provider,
                analysis_prompt,
                "Tu es un expert en analyse documentaire juridique."
            )
            
            if response['success']:
                st.session_state.import_analysis = response['response']
                st.success("✅ Analyse des imports terminée")
                
                with st.expander("📊 Résultats de l'analyse", expanded=True):
                    st.write(response['response'])