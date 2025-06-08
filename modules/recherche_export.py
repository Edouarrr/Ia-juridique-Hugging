"""Module pour les fonctions d'export"""

import streamlit as st
from datetime import datetime
import re

# Vérifier la disponibilité des bibliothèques
try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

def process_export_request(query: str, analysis: dict):
    """Traite une demande d'export"""
    
    format = analysis.get('details', {}).get('format', 'docx')
    
    # Déterminer ce qu'il faut exporter
    content_to_export = None
    filename_base = "export"
    
    # Priorité : rédaction > analyse > recherche > sélection
    if st.session_state.get('redaction_result'):
        content_to_export = st.session_state.redaction_result['document']
        filename_base = f"{st.session_state.redaction_result['type']}"
    
    elif st.session_state.get('timeline_result'):
        content_to_export = export_timeline_content(st.session_state.timeline_result)
        filename_base = "chronologie"
    
    elif st.session_state.get('mapping_result'):
        content_to_export = export_mapping_content(st.session_state.mapping_result)
        filename_base = "cartographie"
    
    elif st.session_state.get('comparison_result'):
        content_to_export = export_comparison_content(st.session_state.comparison_result)
        filename_base = "comparaison"
    
    elif st.session_state.get('ai_analysis_results'):
        content_to_export = format_ai_analysis_for_export(st.session_state.ai_analysis_results)
        filename_base = "analyse"
    
    elif st.session_state.get('search_results'):
        content_to_export = format_search_results_for_export(st.session_state.search_results)
        filename_base = "recherche"
    
    else:
        st.warning("⚠️ Aucun contenu à exporter")
        return
    
    # Interface d'export
    st.markdown("### 💾 Export de documents")
    
    # Sélection du format
    col1, col2 = st.columns(2)
    
    with col1:
        export_format = st.selectbox(
            "Format d'export",
            ["docx", "pdf", "txt", "html", "xlsx"],
            key="export_format"
        )
    
    with col2:
        include_metadata = st.checkbox(
            "Inclure les métadonnées",
            value=True,
            key="export_metadata"
        )
    
    # Options spécifiques DOCX
    if export_format == "docx" and DOCX_AVAILABLE:
        with st.expander("Options Word", expanded=False):
            st.checkbox("Inclure table des matières", value=True, key="docx_toc")
            st.checkbox("Numérotation automatique", value=True, key="docx_numbering")
            st.checkbox("En-têtes et pieds de page", value=True, key="docx_headers")
    
    # Exporter selon le format
    export_functions = {
        'docx': export_to_docx,
        'pdf': export_to_pdf,
        'txt': export_to_txt,
        'html': export_to_html,
        'xlsx': export_to_xlsx
    }
    
    export_func = export_functions.get(export_format, export_to_txt)
    
    try:
        exported_data = export_func(content_to_export, analysis)
        
        # Créer le bouton de téléchargement
        st.download_button(
            f"💾 Télécharger {export_format.upper()}",
            exported_data,
            f"{filename_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{export_format}",
            get_mime_type(export_format),
            key=f"download_export_{export_format}"
        )
        
        st.success(f"✅ Export {export_format.upper()} prêt")
        
    except Exception as e:
        st.error(f"❌ Erreur export: {str(e)}")

def export_timeline_content(timeline_result: dict) -> str:
    """Exporte le contenu d'une chronologie"""
    content = f"CHRONOLOGIE - {timeline_result['type'].upper()}\n"
    content += f"Générée le {timeline_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Basée sur {timeline_result['document_count']} documents\n\n"
    
    for event in timeline_result['events']:
        content += f"{event.get('date', 'N/A')} - {event.get('description', '')}\n"
        if 'actors' in event:
            content += f"   Acteurs: {', '.join(event['actors'])}\n"
        if 'source' in event:
            content += f"   Source: {event['source']}\n"
        content += "\n"
    
    return content

def export_mapping_content(mapping_result: dict) -> str:
    """Exporte le contenu d'une cartographie"""
    content = f"CARTOGRAPHIE DES RELATIONS - {mapping_result['type'].upper()}\n"
    content += f"Générée le {mapping_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n\n"
    
    content += f"STATISTIQUES:\n"
    content += f"- Entités: {mapping_result['analysis']['node_count']}\n"
    content += f"- Relations: {mapping_result['analysis']['edge_count']}\n"
    content += f"- Densité: {mapping_result['analysis']['density']:.2%}\n\n"
    
    content += "ACTEURS PRINCIPAUX:\n"
    for i, player in enumerate(mapping_result['analysis']['key_players'], 1):
        content += f"{i}. {player}\n"
    
    return content

def export_comparison_content(comparison_result: dict) -> str:
    """Exporte le contenu d'une comparaison"""
    content = f"COMPARAISON - {comparison_result['type'].upper()}\n"
    content += f"Générée le {comparison_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Documents comparés: {comparison_result['document_count']}\n\n"
    
    comparison = comparison_result['comparison']
    
    content += "CONVERGENCES:\n"
    for conv in comparison.get('convergences', []):
        content += f"- {conv['point']}\n  {conv['details']}\n\n"
    
    content += "DIVERGENCES:\n"
    for div in comparison.get('divergences', []):
        content += f"- {div['point']}\n  {div['details']}\n\n"
    
    return content

def format_ai_analysis_for_export(analysis_results: dict) -> str:
    """Formate les résultats d'analyse IA pour l'export"""
    content = f"ANALYSE IA - {analysis_results.get('type', 'Générale').upper()}\n"
    content += f"Date : {analysis_results.get('timestamp', datetime.now()).strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Documents analysés : {analysis_results.get('document_count', 0)}\n\n"
    
    if analysis_results.get('query'):
        content += f"QUESTION : {analysis_results['query']}\n\n"
    
    content += "ANALYSE :\n"
    content += analysis_results.get('content', 'Aucun contenu')
    
    return content

def format_search_results_for_export(search_results: list) -> str:
    """Formate les résultats de recherche pour l'export"""
    content = f"RÉSULTATS DE RECHERCHE\n"
    content += f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Nombre de résultats : {len(search_results)}\n\n"
    
    for i, result in enumerate(search_results, 1):
        content += f"{i}. {result.get('title', 'Sans titre')}\n"
        content += f"   Source : {result.get('source', 'N/A')}\n"
        content += f"   Score : {result.get('score', 0):.2%}\n"
        if result.get('content'):
            content += f"   Extrait : {result['content'][:200]}...\n"
        content += "\n"
    
    return content

def export_to_docx(content: str, analysis: dict) -> bytes:
    """Exporte vers Word avec mise en forme"""
    if not DOCX_AVAILABLE:
        # Fallback en texte brut
        return content.encode('utf-8')
    
    try:
        doc = docx.Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
        
        # Styles personnalisés
        styles = doc.styles
        
        # Style pour le titre principal
        title_style = styles.add_style('JuridicalTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # En-tête si demandé
        if st.session_state.get('export_metadata', True):
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Contenu avec mise en forme
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            # Détecter les niveaux de titre
            if line.strip().isupper() and len(line.strip()) > 5:
                doc.add_heading(line.strip(), 1)
            elif line.strip().startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
                p = doc.add_paragraph()
                p.add_run(line.strip()).bold = True
                p.style.font.size = Pt(14)
            elif line.strip().startswith('-'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(line.strip())
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création DOCX: {e}")
        return content.encode('utf-8')

def create_formatted_docx(content: str, doc_type: str) -> bytes:
    """Crée un document Word avec mise en forme professionnelle complète"""
    
    if not DOCX_AVAILABLE:
        return content.encode('utf-8')
    
    try:
        doc = docx.Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
        
        # Styles personnalisés pour documents juridiques
        styles = doc.styles
        
        # Style pour le titre principal
        if 'JuridicalTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('JuridicalTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        
        # Style pour les sections principales
        if 'JuridicalHeading1' not in [s.name for s in styles]:
            heading1_style = styles.add_style('JuridicalHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = 'Times New Roman'
            heading1_style.font.size = Pt(14)
            heading1_style.font.bold = True
            heading1_style.paragraph_format.space_before = Pt(18)
            heading1_style.paragraph_format.space_after = Pt(12)
        
        # Style pour le corps du texte
        body_style = styles['Normal']
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.first_line_indent = Inches(0.5)
        
        # Analyser et formater le contenu
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            # Titre principal
            if line.isupper() and len(line.split()) < 10:
                p = doc.add_paragraph(line, style='JuridicalTitle')
            
            # Sections principales
            elif re.match(r'^[IVX]+\.\s+', line):
                p = doc.add_paragraph(line, style='JuridicalHeading1')
            
            # Points avec tirets
            elif line.startswith('-'):
                p