# modules/synthesis.py
"""Module de synthèse de documents et pièces juridiques"""

import streamlit as st
from datetime import datetime
from typing import List, Dict, Any, Optional
from collections import defaultdict

from models.dataclasses import PieceSelectionnee, Document
from managers.multi_llm_manager import MultiLLMManager
from utils.helpers import truncate_text, extract_key_phrases

def process_synthesis_request(query: str, analysis: dict):
    """Traite une demande de synthèse"""
    
    # Déterminer la source de la synthèse
    synthesis_content = None
    
    if st.session_state.get('selected_pieces'):
        synthesis_content = synthesize_selected_pieces(st.session_state.selected_pieces)
    
    elif analysis.get('reference'):
        docs = search_documents_by_reference(f"@{analysis['reference']}")
        synthesis_content = synthesize_documents(docs)
    
    elif st.session_state.get('search_results'):
        synthesis_content = synthesize_search_results(st.session_state.search_results)
    
    else:
        st.warning("⚠️ Aucun contenu à synthétiser")
        return
    
    # Afficher la synthèse
    if synthesis_content:
        display_synthesis_interface(synthesis_content)
        st.session_state.synthesis_result = synthesis_content

def synthesize_selected_pieces(pieces: List[PieceSelectionnee]) -> dict:
    """Synthétise les pièces sélectionnées"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le contexte
    context = construct_synthesis_context(pieces)
    
    # Prompt de synthèse
    synthesis_prompt = f"""{context}

Crée une synthèse structurée et détaillée de ces pièces.

La synthèse doit inclure:

1. VUE D'ENSEMBLE
   - Nombre et types de documents
   - Période couverte
   - Parties impliquées
   - Contexte général

2. POINTS CLÉS PAR CATÉGORIE
   - Pour chaque catégorie de documents
   - Informations essentielles
   - Éléments de preuve

3. CHRONOLOGIE DES ÉVÉNEMENTS
   - Timeline des faits importants
   - Enchaînement logique
   - Points de rupture

4. ANALYSE CROISÉE
   - Liens entre les pièces
   - Cohérences et incohérences
   - Éléments manquants

5. POINTS D'ATTENTION JURIDIQUES
   - Qualifications possibles
   - Risques identifiés
   - Forces et faiblesses

6. RECOMMANDATIONS
   - Actions prioritaires
   - Pièces complémentaires nécessaires
   - Stratégie suggérée

Format professionnel avec titres, sous-sections et mise en évidence des éléments importants."""
    
    try:
        # Utiliser plusieurs IA si disponibles
        providers = list(llm_manager.clients.keys())[:2]  # Max 2 pour la synthèse
        
        if len(providers) > 1:
            # Synthèse multi-IA
            responses = llm_manager.query_multiple_llms(
                providers,
                synthesis_prompt,
                "Tu es un expert en analyse et synthèse de documents juridiques complexes.",
                temperature=0.7,
                max_tokens=4000
            )
            
            # Fusionner les réponses
            synthesis = llm_manager.fusion_responses(responses)
        else:
            # Synthèse simple
            provider = providers[0]
            response = llm_manager.query_single_llm(
                provider,
                synthesis_prompt,
                "Tu es un expert en analyse et synthèse de documents juridiques complexes."
            )
            synthesis = response['response'] if response['success'] else "Erreur de synthèse"
        
        return {
            'content': synthesis,
            'piece_count': len(pieces),
            'categories': list(set(p.categorie for p in pieces)),
            'timestamp': datetime.now(),
            'method': 'multi_llm' if len(providers) > 1 else 'single_llm'
        }
        
    except Exception as e:
        return {'error': f'Erreur synthèse: {str(e)}'}

def synthesize_documents(documents: List[Dict[str, Any]]) -> dict:
    """Synthétise une liste de documents"""
    # Convertir en pièces pour réutiliser la fonction
    pieces = []
    
    for i, doc in enumerate(documents):
        pieces.append(PieceSelectionnee(
            numero=i + 1,
            titre=doc.get('title', 'Sans titre'),
            description=truncate_text(doc.get('content', ''), 200),
            categorie=determine_document_category(doc),
            source=doc.get('source', ''),
            pertinence=doc.get('score', 0.5)
        ))
    
    return synthesize_selected_pieces(pieces)

def synthesize_search_results(results: List[Dict[str, Any]]) -> dict:
    """Synthétise des résultats de recherche"""
    # Limiter aux 20 premiers résultats
    top_results = results[:20]
    
    # Convertir en format pour synthèse
    pieces = []
    for i, result in enumerate(top_results):
        pieces.append(PieceSelectionnee(
            numero=i + 1,
            titre=result.get('title', 'Sans titre'),
            description=truncate_text(result.get('content', ''), 200),
            categorie='Résultat de recherche',
            source=result.get('source', ''),
            pertinence=result.get('score', 0.5)
        ))
    
    return synthesize_selected_pieces(pieces)

def construct_synthesis_context(pieces: List[PieceSelectionnee]) -> str:
    """Construit le contexte pour la synthèse"""
    context = "PIÈCES À SYNTHÉTISER:\n\n"
    
    # Grouper par catégorie
    by_category = defaultdict(list)
    for piece in pieces:
        by_category[piece.categorie].append(piece)
    
    # Limiter à 30 pièces pour le contexte
    total_pieces = 0
    
    for category, cat_pieces in by_category.items():
        context += f"\n--- {category.upper()} ({len(cat_pieces)} pièces) ---\n"
        
        for piece in cat_pieces:
            if total_pieces >= 30:
                context += "\n[...Pièces supplémentaires tronquées pour la synthèse...]\n"
                break
            
            context += f"\nPièce {piece.numero}: {piece.titre}\n"
            
            if piece.cote:
                context += f"Cote: {piece.cote}\n"
            
            if piece.date:
                date_str = piece.date.strftime('%d/%m/%Y') if hasattr(piece.date, 'strftime') else str(piece.date)
                context += f"Date: {date_str}\n"
            
            if piece.description:
                context += f"Description: {piece.description}\n"
            
            context += f"Pertinence: {piece.pertinence:.0%}\n"
            
            total_pieces += 1
    
    context += f"\nTOTAL: {len(pieces)} pièces analysées\n"
    
    return context

def display_synthesis_interface(synthesis_result: dict):
    """Affiche l'interface de synthèse"""
    
    if 'error' in synthesis_result:
        st.error(f"❌ {synthesis_result['error']}")
        return
    
    st.markdown("### 📝 Synthèse des documents")
    
    # Métadonnées
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Pièces analysées", synthesis_result.get('piece_count', 0))
    
    with col2:
        st.metric("Catégories", len(synthesis_result.get('categories', [])))
    
    with col3:
        st.metric("Méthode", synthesis_result.get('method', 'simple').replace('_', ' ').title())
    
    with col4:
        timestamp = synthesis_result.get('timestamp', datetime.now())
        st.metric("Généré à", timestamp.strftime('%H:%M'))
    
    # Options d'affichage
    col1, col2 = st.columns([3, 1])
    
    with col1:
        view_mode = st.radio(
            "Mode d'affichage",
            ["📄 Document", "📊 Sections", "🔍 Points clés"],
            horizontal=True,
            key="synthesis_view_mode"
        )
    
    with col2:
        edit_mode = st.checkbox("✏️ Mode édition", key="edit_synthesis")
    
    # Contenu de la synthèse
    synthesis_content = synthesis_result.get('content', '')
    
    if view_mode == "📄 Document":
        # Vue document complète
        if edit_mode:
            edited_content = st.text_area(
                "Synthèse complète",
                value=synthesis_content,
                height=600,
                key="synthesis_content_editor"
            )
            synthesis_content = edited_content
        else:
            st.markdown(synthesis_content)
    
    elif view_mode == "📊 Sections":
        # Vue par sections
        sections = extract_sections_from_synthesis(synthesis_content)
        
        for section_title, section_content in sections.items():
            with st.expander(f"📌 {section_title}", expanded=True):
                if edit_mode:
                    edited = st.text_area(
                        f"Éditer {section_title}",
                        value=section_content,
                        height=200,
                        key=f"edit_section_{section_title}"
                    )
                    sections[section_title] = edited
                else:
                    st.markdown(section_content)
    
    else:  # Points clés
        # Extraction des points clés
        key_points = extract_key_points_from_synthesis(synthesis_content)
        
        if key_points:
            for i, point in enumerate(key_points, 1):
                st.write(f"**{i}.** {point}")
        else:
            st.info("Aucun point clé extrait")
    
    # Actions
    st.markdown("### 💾 Actions")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("📄 Exporter Word", key="export_synthesis_docx"):
            docx_content = export_synthesis_to_docx(synthesis_result)
            st.download_button(
                "💾 Télécharger DOCX",
                docx_content,
                f"synthese_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_synthesis_docx"
            )
    
    with col2:
        if st.button("📧 Envoyer", key="email_synthesis"):
            st.session_state.pending_email = {
                'type': 'synthesis',
                'content': synthesis_content,
                'subject': "Synthèse des documents"
            }
            st.info("📧 Prêt pour envoi")
    
    with col3:
        if st.button("🔄 Régénérer", key="regenerate_synthesis"):
            if st.session_state.get('selected_pieces'):
                new_synthesis = synthesize_selected_pieces(st.session_state.selected_pieces)
                st.session_state.synthesis_result = new_synthesis
                st.rerun()
    
    with col4:
        if st.button("📊 Statistiques", key="synthesis_stats"):
            show_synthesis_statistics(synthesis_result)

def extract_sections_from_synthesis(content: str) -> Dict[str, str]:
    """Extrait les sections d'une synthèse"""
    import re
    
    sections = {}
    current_section = "Introduction"
    current_content = []
    
    lines = content.split('\n')
    
    for line in lines:
        # Détecter les titres de section (en majuscules ou numérotés)
        if re.match(r'^[0-9]+\.\s+[A-Z]', line) or (line.isupper() and len(line) > 5 and len(line) < 50):
            # Sauvegarder la section précédente
            if current_content:
                sections[current_section] = '\n'.join(current_content).strip()
            
            # Nouvelle section
            current_section = line.strip()
            current_content = []
        else:
            current_content.append(line)
    
    # Dernière section
    if current_content:
        sections[current_section] = '\n'.join(current_content).strip()
    
    return sections

def extract_key_points_from_synthesis(content: str) -> List[str]:
    """Extrait les points clés d'une synthèse"""
    key_points = []
    
    # Patterns pour identifier les points importants
    patterns = [
        r'^\s*[-•]\s*(.+)$',  # Listes à puces
        r'^\s*\d+\.\s*(.+)$',  # Listes numérotées
        r'^(?:Il ressort|Il apparaît|On constate|Force est de constater)\s+(.+)$',  # Formulations juridiques
        r'^(?:Point clé|Important|À noter|Attention)\s*:\s*(.+)$',  # Marqueurs explicites
    ]
    
    lines = content.split('\n')
    
    for line in lines:
        for pattern in patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                point = match.group(1).strip()
                if len(point) > 20:  # Éviter les points trop courts
                    key_points.append(point)
                break
    
    # Dédupliquer
    seen = set()
    unique_points = []
    for point in key_points:
        if point.lower() not in seen:
            seen.add(point.lower())
            unique_points.append(point)
    
    return unique_points[:20]  # Limiter à 20 points

def export_synthesis_to_docx(synthesis_result: dict) -> bytes:
    """Exporte la synthèse en format Word"""
    try:
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Titre
        title = doc.add_heading('SYNTHÈSE DES DOCUMENTS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Métadonnées
        doc.add_paragraph(f"Date : {synthesis_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}")
        doc.add_paragraph(f"Nombre de pièces analysées : {synthesis_result['piece_count']}")
        doc.add_paragraph(f"Catégories : {', '.join(synthesis_result['categories'])}")
        doc.add_paragraph()
        
        # Contenu
        content_lines = synthesis_result['content'].split('\n')
        
        for line in content_lines:
            if not line.strip():
                doc.add_paragraph()
                continue
            
            # Détecter les niveaux de titre
            if line.strip().isupper() and len(line.strip()) > 5:
                doc.add_heading(line.strip(), 1)
            elif re.match(r'^\d+\.\s+[A-Z]', line):
                doc.add_heading(line.strip(), 2)
            elif line.strip().startswith(('- ', '• ', '* ')):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError:
        # Fallback en texte
        return synthesis_result['content'].encode('utf-8')

def show_synthesis_statistics(synthesis_result: dict):
    """Affiche les statistiques de la synthèse"""
    content = synthesis_result.get('content', '')
    
    with st.expander("📊 Statistiques de la synthèse", expanded=True):
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Mots", f"{len(content.split()):,}")
            st.metric("Caractères", f"{len(content):,}")
        
        with col2:
            st.metric("Paragraphes", len([p for p in content.split('\n\n') if p.strip()]))
            st.metric("Lignes", len(content.split('\n')))
        
        with col3:
            sections = extract_sections_from_synthesis(content)
            st.metric("Sections", len(sections))
            key_points = extract_key_points_from_synthesis(content)
            st.metric("Points clés", len(key_points))
        
        # Nuage de mots-clés si disponible
        if 'categories' in synthesis_result:
            st.write("**Catégories analysées:**")
            for cat in synthesis_result['categories']:
                st.caption(f"• {cat}")

def determine_document_category(doc: dict) -> str:
    """Détermine la catégorie d'un document (fonction helper)"""
    # Réutiliser depuis piece_selection ou importer
    title_lower = doc.get('title', '').lower()
    
    if 'procédure' in title_lower or 'pv' in title_lower:
        return 'Procédure'
    elif 'expertise' in title_lower:
        return 'Expertise'
    elif 'contrat' in title_lower:
        return 'Contrats'
    else:
        return 'Autres'

def search_documents_by_reference(reference: str) -> List[Dict[str, Any]]:
    """Recherche des documents par référence"""
    # Fonction helper - peut être importée depuis piece_selection
    results = []
    ref_clean = reference.replace('@', '').strip().lower()
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if ref_clean in doc.title.lower() or ref_clean in doc.source.lower():
            results.append({
                'id': doc_id,
                'title': doc.title,
                'content': doc.content,
                'source': doc.source
            })
    
    return results