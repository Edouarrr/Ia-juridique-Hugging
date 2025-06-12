"""Module d'export des actes juridiques en Word et PDF avec mise en forme Garamond"""

import os
import re
from datetime import datetime
from typing import Dict, Any, Optional
from io import BytesIO

# Import des bibliothèques d'export
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx non disponible - Export Word désactivé")

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ reportlab non disponible - Export PDF désactivé")

# Configuration du cahier des charges
from config.cahier_des_charges import STYLE_DOCUMENT, HIERARCHIE_NUMEROTATION

# ========================= EXPORT WORD =========================

class ExportWord:
    """Classe pour exporter les actes juridiques en format Word"""
    
    def __init__(self):
        self.doc = None
        self.styles_configured = False
    
    def export_acte(self, acte: 'ActeJuridique') -> bytes:
        """Exporte un acte juridique en format Word avec style Garamond"""
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx n'est pas installé")
        
        # Créer un nouveau document
        self.doc = Document()
        
        # Configurer les styles
        self._configure_styles()
        
        # Définir les marges
        self._set_margins()
        
        # Ajouter le contenu
        self._add_content(acte)
        
        # Ajouter la pagination
        self._add_page_numbers()
        
        # Sauvegarder dans un buffer
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _configure_styles(self):
        """Configure les styles selon le cahier des charges"""
        
        # Style pour le texte normal
        style_normal = self.doc.styles['Normal']
        style_normal.font.name = 'Garamond'
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Créer les styles hiérarchiques
        for niveau, config in HIERARCHIE_NUMEROTATION.items():
            style_name = f'Heading {niveau}'
            
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
            else:
                style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            
            style.font.name = 'Garamond'
            style.font.size = Pt(config['police_taille'])
            
            # Appliquer le style (gras, souligné, italique)
            if 'gras' in config['style']:
                style.font.bold = True
            if 'italique' in config['style']:
                style.font.italic = True
            
            # Pour le soulignement, on devra l'appliquer directement
            
        # Style pour l'en-tête centré
        style_header = self.doc.styles.add_style('HeaderCenter', WD_STYLE_TYPE.PARAGRAPH)
        style_header.font.name = 'Garamond'
        style_header.font.size = Pt(14)
        style_header.font.bold = True
        style_header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_header.font.all_caps = True
        
        # Style pour les formules de politesse
        style_politesse = self.doc.styles.add_style('FormulePolitesse', WD_STYLE_TYPE.PARAGRAPH)
        style_politesse.font.name = 'Garamond'
        style_politesse.font.size = Pt(12)
        style_politesse.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self.styles_configured = True
    
    def _set_margins(self):
        """Définit les marges selon le cahier des charges"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(STYLE_DOCUMENT['marges']['haut'] / 2.54)
            section.bottom_margin = Inches(STYLE_DOCUMENT['marges']['bas'] / 2.54)
            section.left_margin = Inches(STYLE_DOCUMENT['marges']['gauche'] / 2.54)
            section.right_margin = Inches(STYLE_DOCUMENT['marges']['droite'] / 2.54)
    
    def _add_content(self, acte: 'ActeJuridique'):
        """Ajoute le contenu de l'acte au document"""
        
        # Parser le contenu HTML-like et le convertir
        content_lines = acte.contenu.split('\n')
        
        for line in content_lines:
            line = line.strip()
            
            if not line:
                # Ligne vide = nouveau paragraphe
                self.doc.add_paragraph()
                continue
            
            # Détecter le type de ligne
            if line.isupper() and len(line) < 100:
                # En-tête en majuscules
                p = self.doc.add_paragraph(line, style='HeaderCenter')
            
            elif self._is_heading(line):
                # Titre hiérarchique
                niveau = self._get_heading_level(line)
                p = self.doc.add_paragraph(line, style=f'Heading {niveau}')
                
                # Ajouter le soulignement si nécessaire
                if niveau <= 2:
                    self._add_underline(p)
            
            elif line.startswith('Pièce n°'):
                # Pièce en gras souligné
                p = self.doc.add_paragraph()
                run = p.add_run(line)
                run.font.bold = True
                run.font.underline = True
                run.font.name = 'Garamond'
                run.font.size = Pt(12)
            
            elif line.startswith('(PIÈCE N°'):
                # Référence de pièce dans le texte
                p = self.doc.add_paragraph()
                # Extraire les parties
                parts = re.split(r'(\(PIÈCE N° \d+\))', line)
                for part in parts:
                    if part.startswith('(PIÈCE'):
                        run = p.add_run(part)
                        run.font.bold = True
                        run.font.underline = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Garamond'
                    run.font.size = Pt(12)
            
            else:
                # Paragraphe normal
                p = self.doc.add_paragraph(line, style='Normal')
    
    def _is_heading(self, line: str) -> bool:
        """Vérifie si une ligne est un titre"""
        patterns = [
            r'^[IVX]+\.\s+',  # I., II., III.
            r'^[A-Z]\.\s+',   # A., B., C.
            r'^\d+\.\s+',     # 1., 2., 3.
            r'^[a-z]\.\s+',   # a., b., c.
            r'^\([ivx]+\)\s+' # (i), (ii), (iii)
        ]
        
        for pattern in patterns:
            if re.match(pattern, line):
                return True
        return False
    
    def _get_heading_level(self, line: str) -> int:
        """Détermine le niveau hiérarchique d'un titre"""
        if re.match(r'^[IVX]+\.\s+', line):
            return 1
        elif re.match(r'^[A-Z]\.\s+', line):
            return 2
        elif re.match(r'^\d+\.\s+', line):
            return 3
        elif re.match(r'^[a-z]\.\s+', line):
            return 4
        elif re.match(r'^\([ivx]+\)\s+', line):
            return 5
        return 3  # Par défaut
    
    def _add_underline(self, paragraph):
        """Ajoute un soulignement à un paragraphe"""
        for run in paragraph.runs:
            run.font.underline = True
    
    def _add_page_numbers(self):
        """Ajoute la pagination (commence à la page 2)"""
        # Cette fonction est complexe en python-docx
        # Pour une implémentation complète, voir la documentation
        pass

# ========================= EXPORT PDF =========================

class ExportPDF:
    """Classe pour exporter les actes juridiques en format PDF"""
    
    def __init__(self):
        self.styles = None
        self._setup_fonts()
    
    def _setup_fonts(self):
        """Configure les polices pour le PDF"""
        if not PDF_AVAILABLE:
            return
        
        # Essayer d'enregistrer Garamond si disponible
        try:
            # Chercher Garamond dans les polices système
            garamond_path = self._find_garamond_font()
            if garamond_path:
                pdfmetrics.registerFont(TTFont('Garamond', garamond_path))
                self.font_name = 'Garamond'
            else:
                # Fallback sur Times-Roman (similaire à Garamond)
                self.font_name = 'Times-Roman'
        except:
            self.font_name = 'Times-Roman'
    
    def _find_garamond_font(self) -> Optional[str]:
        """Cherche la police Garamond dans le système"""
        possible_paths = [
            '/usr/share/fonts/truetype/garamond/Garamond.ttf',
            '/System/Library/Fonts/Garamond.ttf',
            'C:\\Windows\\Fonts\\GARAMOND.TTF',
            'C:\\Windows\\Fonts\\garamond.ttf'
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        return None
    
    def export_acte(self, acte: 'ActeJuridique') -> bytes:
        """Exporte un acte juridique en format PDF"""
        
        if not PDF_AVAILABLE:
            raise ImportError("reportlab n'est pas installé")
        
        # Créer un buffer
        buffer = BytesIO()
        
        # Créer le document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=STYLE_DOCUMENT['marges']['droite'] * cm,
            leftMargin=STYLE_DOCUMENT['marges']['gauche'] * cm,
            topMargin=STYLE_DOCUMENT['marges']['haut'] * cm,
            bottomMargin=STYLE_DOCUMENT['marges']['bas'] * cm
        )
        
        # Créer les styles
        self._create_styles()
        
        # Construire le contenu
        story = self._build_story(acte)
        
        # Générer le PDF
        doc.build(story, onFirstPage=self._first_page, onLaterPages=self._later_pages)
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def _create_styles(self):
        """Crée les styles pour le PDF"""
        self.styles = getSampleStyleSheet()
        
        # Style normal avec Garamond
        self.styles.add(ParagraphStyle(
            name='GaramondNormal',
            parent=self.styles['Normal'],
            fontName=self.font_name,
            fontSize=12,
            leading=18,  # Interligne 1.5
            alignment=4,  # Justifié
            spaceAfter=6
        ))
        
        # Styles pour les titres
        for i in range(1, 6):
            size = 14 if i == 1 else 12
            self.styles.add(ParagraphStyle(
                name=f'GaramondHeading{i}',
                parent=self.styles['Heading1'],
                fontName=self.font_name,
                fontSize=size,
                leading=size * 1.5,
                textColor=colors.black,
                spaceAfter=12,
                spaceBefore=12
            ))
        
        # Style pour l'en-tête centré
        self.styles.add(ParagraphStyle(
            name='HeaderCenter',
            parent=self.styles['Normal'],
            fontName=self.font_name,
            fontSize=14,
            alignment=1,  # Centré
            textTransform='uppercase'
        ))
    
    def _build_story(self, acte: 'ActeJuridique') -> list:
        """Construit le contenu du PDF"""
        story = []
        
        # Parser le contenu
        lines = acte.contenu.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                story.append(Spacer(1, 12))
                continue
            
            # Nettoyer le HTML
            line = self._clean_html(line)
            
            # Déterminer le style
            if line.isupper() and len(line) < 100:
                para = Paragraph(line, self.styles['HeaderCenter'])
            elif self._is_heading(line):
                level = self._get_heading_level(line)
                # Ajouter le formatage gras/souligné
                if level <= 2:
                    line = f'<b><u>{line}</u></b>'
                elif level == 3:
                    line = f'<b>{line}</b>'
                else:
                    line = f'<i>{line}</i>'
                para = Paragraph(line, self.styles[f'GaramondHeading{level}'])
            else:
                # Mettre en évidence les pièces
                line = re.sub(
                    r'\(PIÈCE N° (\d+)\)',
                    r'<b><u>(PIÈCE N° \1)</u></b>',
                    line
                )
                para = Paragraph(line, self.styles['GaramondNormal'])
            
            story.append(para)
        
        return story
    
    def _clean_html(self, text: str) -> str:
        """Nettoie et convertit le HTML pour reportlab"""
        # Remplacer les tags HTML custom par des tags reportlab
        text = re.sub(r'<h\d[^>]*>', '<b>', text)
        text = re.sub(r'</h\d>', '</b>', text)
        text = re.sub(r'<span[^>]*>', '', text)
        text = re.sub(r'</span>', '', text)
        
        # Échapper les caractères spéciaux
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;').replace('>', '&gt;')
        
        # Restaurer les tags autorisés
        text = text.replace('&lt;b&gt;', '<b>').replace('&lt;/b&gt;', '</b>')
        text = text.replace('&lt;i&gt;', '<i>').replace('&lt;/i&gt;', '</i>')
        text = text.replace('&lt;u&gt;', '<u>').replace('&lt;/u&gt;', '</u>')
        
        return text
    
    def _is_heading(self, line: str) -> bool:
        """Vérifie si une ligne est un titre"""
        patterns = [
            r'^[IVX]+\.\s+',
            r'^[A-Z]\.\s+',
            r'^\d+\.\s+',
            r'^[a-z]\.\s+',
            r'^\([ivx]+\)\s+'
        ]
        
        for pattern in patterns:
            if re.match(pattern, line):
                return True
        return False
    
    def _get_heading_level(self, line: str) -> int:
        """Détermine le niveau hiérarchique"""
        if re.match(r'^[IVX]+\.\s+', line):
            return 1
        elif re.match(r'^[A-Z]\.\s+', line):
            return 2
        elif re.match(r'^\d+\.\s+', line):
            return 3
        elif re.match(r'^[a-z]\.\s+', line):
            return 4
        elif re.match(r'^\([ivx]+\)\s+', line):
            return 5
        return 3
    
    def _first_page(self, canvas, doc):
        """Callback pour la première page (pas de numéro)"""
        canvas.saveState()
        # Pas de numéro sur la première page
        canvas.restoreState()
    
    def _later_pages(self, canvas, doc):
        """Callback pour les pages suivantes (avec numéro)"""
        canvas.saveState()
        canvas.setFont(self.font_name, 10)
        page_num = canvas.getPageNumber() - 1  # Commence à 2
        if page_num > 0:
            text = str(page_num)
            canvas.drawRightString(
                doc.pagesize[0] - doc.rightMargin,
                doc.bottomMargin / 2,
                text
            )
        canvas.restoreState()

# ========================= GESTIONNAIRE D'EXPORT =========================

class GestionnaireExport:
    """Gestionnaire principal pour l'export des actes juridiques"""
    
    def __init__(self):
        self.word_exporter = ExportWord() if DOCX_AVAILABLE else None
        self.pdf_exporter = ExportPDF() if PDF_AVAILABLE else None
    
    def export(self, acte: 'ActeJuridique', format: str = 'docx') -> tuple[bytes, str]:
        """
        Exporte un acte dans le format demandé
        
        Returns:
            tuple: (contenu en bytes, extension de fichier)
        """
        
        format = format.lower()
        
        if format == 'docx' or format == 'word':
            if not self.word_exporter:
                raise ValueError("Export Word non disponible - installez python-docx")
            return self.word_exporter.export_acte(acte), 'docx'
        
        elif format == 'pdf':
            if not self.pdf_exporter:
                raise ValueError("Export PDF non disponible - installez reportlab")
            return self.pdf_exporter.export_acte(acte), 'pdf'
        
        elif format == 'txt' or format == 'text':
            # Export texte simple
            return acte.contenu.encode('utf-8'), 'txt'
        
        elif format == 'rtf':
            # Export RTF basique
            return self._export_rtf(acte), 'rtf'
        
        else:
            raise ValueError(f"Format non supporté : {format}")
    
    def _export_rtf(self, acte: 'ActeJuridique') -> bytes:
        """Export basique en RTF"""
        rtf_content = r'{\rtf1\ansi\deff0 {\fonttbl{\f0 Times New Roman;}}' + '\n'
        rtf_content += r'\f0\fs24 '  # Police 12pt
        
        # Convertir le contenu
        lines = acte.contenu.split('\n')
        for line in lines:
            if line.strip():
                # Échapper les caractères spéciaux RTF
                line = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                rtf_content += line + r'\par '
            else:
                rtf_content += r'\par '
        
        rtf_content += '}'
        
        return rtf_content.encode('utf-8')
    
    def get_available_formats(self) -> list[str]:
        """Retourne la liste des formats disponibles"""
        formats = ['txt', 'rtf']
        
        if DOCX_AVAILABLE:
            formats.append('docx')
        
        if PDF_AVAILABLE:
            formats.append('pdf')
        
        return formats

# ========================= FONCTION UTILITAIRE =========================

def export_acte_juridique(acte: 'ActeJuridique', format: str = 'docx', 
                         filename: Optional[str] = None) -> str:
    """
    Fonction utilitaire pour exporter un acte juridique
    
    Args:
        acte: L'acte juridique à exporter
        format: Le format d'export (docx, pdf, txt, rtf)
        filename: Le nom du fichier (optionnel)
    
    Returns:
        str: Le chemin du fichier créé
    """
    
    gestionnaire = GestionnaireExport()
    
    # Générer le nom de fichier si non fourni
    if not filename:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{acte.type_acte}_{timestamp}"
    
    # Exporter
    content, extension = gestionnaire.export(acte, format)
    
    # Sauvegarder
    full_filename = f"{filename}.{extension}"
    with open(full_filename, 'wb') as f:
        f.write(content)
    
    return full_filename

# ========================= INTÉGRATION STREAMLIT =========================

def show_export_options(acte: 'ActeJuridique'):
    """Affiche les options d'export dans Streamlit"""
    import streamlit as st
    
    gestionnaire = GestionnaireExport()
    formats_disponibles = gestionnaire.get_available_formats()
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        format_choisi = st.selectbox(
            "Format d'export",
            formats_disponibles,
            format_func=lambda x: {
                'docx': 'Word (.docx)',
                'pdf': 'PDF',
                'txt': 'Texte (.txt)',
                'rtf': 'RTF'
            }.get(x, x)
        )
    
    with col2:
        if st.button("📥 Exporter", use_container_width=True):
            try:
                content, extension = gestionnaire.export(acte, format_choisi)
                
                st.download_button(
                    label="💾 Télécharger",
                    data=content,
                    file_name=f"{acte.type_acte}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{extension}",
                    mime={
                        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'pdf': 'application/pdf',
                        'txt': 'text/plain',
                        'rtf': 'application/rtf'
                    }.get(extension, 'application/octet-stream')
                )
                
                st.success(f"✅ Export {format_choisi.upper()} prêt !")
                
            except Exception as e:
                st.error(f"❌ Erreur lors de l'export : {str(e)}")

if __name__ == "__main__":
    # Test du module
    print("Module d'export juridique chargé")
    print(f"Export Word disponible : {DOCX_AVAILABLE}")
    print(f"Export PDF disponible : {PDF_AVAILABLE}")