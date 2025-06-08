# modules/bordereau.py
"""Module de création et gestion des bordereaux de communication de pièces"""

import streamlit as st
from datetime import datetime
from typing import List, Dict, Any, Optional
import io

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from models.dataclasses import PieceSelectionnee

def process_bordereau_request(query: str, analysis: dict):
    """Traite une demande de création de bordereau"""
    
    pieces = st.session_state.get('selected_pieces', [])
    
    if not pieces:
        st.warning("⚠️ Aucune pièce sélectionnée pour le bordereau")
        return
    
    # Créer le bordereau
    bordereau = create_bordereau(pieces, analysis)
    
    # Afficher le bordereau
    display_bordereau_interface(bordereau, pieces)

def create_bordereau(pieces: List[PieceSelectionnee], analysis: dict) -> dict:
    """Crée un bordereau structuré"""
    
    # Déterminer les parties si disponibles
    pour = analysis.get('client', '[À compléter]')
    contre = analysis.get('adversaire', '[À compléter]')
    juridiction = analysis.get('juridiction', '[Tribunal]')
    
    bordereau = {
        'header': f"""BORDEREAU DE COMMUNICATION DE PIÈCES

AFFAIRE : {analysis.get('reference', 'N/A').upper()}
DATE : {datetime.now().strftime('%d/%m/%Y')}
NOMBRE DE PIÈCES : {len(pieces)}

POUR : {pour}
CONTRE : {contre}
DEVANT : {juridiction}

PIÈCES COMMUNIQUÉES :""",
        'pieces': pieces,
        'footer': f"""
Je certifie que les pièces communiquées sont conformes aux originaux en ma possession.

Fait à [Ville], le {datetime.now().strftime('%d/%m/%Y')}

[Signature]""",
        'metadata': {
            'created_at': datetime.now(),
            'piece_count': len(pieces),
            'categories': list(set(p.categorie for p in pieces)),
            'reference': analysis.get('reference')
        }
    }
    
    return bordereau

def display_bordereau_interface(bordereau: dict, pieces: List[PieceSelectionnee]):
    """Affiche l'interface du bordereau"""
    st.markdown("### 📊 Bordereau de communication de pièces")
    
    # Options d'édition
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.info(f"📋 {bordereau['metadata']['piece_count']} pièces - {len(bordereau['metadata']['categories'])} catégories")
    
    with col2:
        edit_mode = st.checkbox("✏️ Mode édition", key="edit_bordereau")
    
    # En-tête
    if edit_mode:
        bordereau['header'] = st.text_area(
            "En-tête du bordereau",
            value=bordereau['header'],
            height=200,
            key="bordereau_header_edit"
        )
    else:
        st.text(bordereau['header'])
    
    st.divider()
    
    # Table des pièces
    display_pieces_table(pieces, edit_mode)
    
    st.divider()
    
    # Pied de page
    if edit_mode:
        bordereau['footer'] = st.text_area(
            "Pied de page",
            value=bordereau['footer'],
            height=100,
            key="bordereau_footer_edit"
        )
    else:
        st.text(bordereau['footer'])
    
    # Actions
    st.markdown("### 💾 Export du bordereau")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("📄 Word (.docx)", key="export_bordereau_docx"):
            docx_content = create_bordereau_docx(bordereau)
            st.download_button(
                "💾 Télécharger DOCX",
                docx_content,
                f"bordereau_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_bordereau_docx"
            )
    
    with col2:
        if st.button("📊 Excel (.xlsx)", key="export_bordereau_xlsx"):
            xlsx_content = create_bordereau_xlsx(bordereau)
            st.download_button(
                "💾 Télécharger XLSX",
                xlsx_content,
                f"bordereau_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_bordereau_xlsx"
            )
    
    with col3:
        if st.button("📝 Texte (.txt)", key="export_bordereau_txt"):
            txt_content = create_bordereau_txt(bordereau)
            st.download_button(
                "💾 Télécharger TXT",
                txt_content,
                f"bordereau_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                "text/plain",
                key="download_bordereau_txt"
            )
    
    with col4:
        if st.button("📧 Envoyer", key="send_bordereau"):
            st.session_state.pending_email = {
                'type': 'bordereau',
                'content': bordereau,
                'subject': f"Bordereau de pièces - {bordereau['metadata']['reference']}"
            }
            st.info("📧 Prêt pour envoi - Configurez l'email")
    
    # Stocker le bordereau
    st.session_state.current_bordereau = bordereau

def display_pieces_table(pieces: List[PieceSelectionnee], edit_mode: bool):
    """Affiche le tableau des pièces"""
    
    if PANDAS_AVAILABLE:
        # Créer un DataFrame
        df_data = []
        for piece in pieces:
            df_data.append({
                'N°': piece.numero,
                'Cote': piece.cote or f"P-{piece.numero:03d}",
                'Titre': piece.titre,
                'Description': piece.description,
                'Catégorie': piece.categorie,
                'Date': piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A',
                'Pages': piece.pages or '-',
                'Format': piece.format or '-'
            })
        
        df = pd.DataFrame(df_data)
        
        if edit_mode:
            # Mode édition avec ag-grid si disponible
            st.data_editor(
                df,
                use_container_width=True,
                num_rows="fixed",
                key="bordereau_pieces_editor"
            )
        else:
            # Affichage simple
            st.dataframe(
                df,
                use_container_width=True,
                hide_index=True
            )
        
        # Statistiques par catégorie
        with st.expander("📊 Statistiques par catégorie"):
            category_stats = df.groupby('Catégorie').agg({
                'N°': 'count',
                'Pages': lambda x: sum(p for p in x if isinstance(p, int))
            }).rename(columns={'N°': 'Nombre'})
            
            st.dataframe(category_stats)
    
    else:
        # Affichage sans pandas
        for piece in pieces:
            with st.container():
                col1, col2, col3, col4 = st.columns([1, 3, 1, 1])
                
                with col1:
                    st.write(f"**{piece.numero}**")
                
                with col2:
                    st.write(piece.titre)
                    if piece.description:
                        st.caption(piece.description)
                
                with col3:
                    st.caption(piece.categorie)
                
                with col4:
                    if piece.date:
                        st.caption(piece.date.strftime('%d/%m/%Y'))

def create_bordereau_docx(bordereau: dict) -> bytes:
    """Crée le document Word du bordereau"""
    if not DOCX_AVAILABLE:
        # Fallback texte
        return create_bordereau_txt(bordereau).encode('utf-8')
    
    try:
        doc = docx.Document()
        
        # Styles
        styles = doc.styles
        
        # En-tête
        for line in bordereau['header'].split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                if 'BORDEREAU' in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(16)
                elif any(x in line for x in ['AFFAIRE', 'DATE', 'NOMBRE']):
                    p.runs[0].bold = True
        
        doc.add_paragraph()
        
        # Table des pièces
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # En-têtes de colonnes
        headers = ['N°', 'Cote', 'Titre', 'Description', 'Catégorie', 'Date']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lignes de pièces
        for piece in bordereau['pieces']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(piece.numero)
            row_cells[1].text = piece.cote or f"P-{piece.numero:03d}"
            row_cells[2].text = piece.titre
            row_cells[3].text = piece.description or ''
            row_cells[4].text = piece.categorie
            row_cells[5].text = piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A'
            
            # Centrer le numéro et la cote
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Pied de page
        doc.add_paragraph()
        for line in bordereau['footer'].split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Sauvegarder
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création bordereau DOCX: {e}")
        return create_bordereau_txt(bordereau).encode('utf-8')

def create_bordereau_xlsx(bordereau: dict) -> bytes:
    """Crée le fichier Excel du bordereau"""
    if not PANDAS_AVAILABLE:
        return create_bordereau_txt(bordereau).encode('utf-8')
    
    try:
        # Créer un writer Excel
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            # Feuille principale
            df_data = []
            for piece in bordereau['pieces']:
                df_data.append({
                    'N°': piece.numero,
                    'Cote': piece.cote or f"P-{piece.numero:03d}",
                    'Titre': piece.titre,
                    'Description': piece.description or '',
                    'Catégorie': piece.categorie,
                    'Date': piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A',
                    'Pages': piece.pages or '',
                    'Format': piece.format or ''
                })
            
            df = pd.DataFrame(df_data)
            df.to_excel(writer, sheet_name='Bordereau', index=False)
            
            # Feuille de métadonnées
            metadata_df = pd.DataFrame([
                ['Référence', bordereau['metadata'].get('reference', 'N/A')],
                ['Date création', bordereau['metadata']['created_at'].strftime('%d/%m/%Y %H:%M')],
                ['Nombre de pièces', bordereau['metadata']['piece_count']],
                ['Catégories', ', '.join(bordereau['metadata']['categories'])]
            ], columns=['Propriété', 'Valeur'])
            
            metadata_df.to_excel(writer, sheet_name='Informations', index=False)
            
            # Ajuster les largeurs
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création bordereau Excel: {e}")
        return create_bordereau_txt(bordereau).encode('utf-8')

def create_bordereau_txt(bordereau: dict) -> str:
    """Crée le bordereau en format texte"""
    content = bordereau['header'] + '\n\n'
    
    # Largeurs de colonnes pour l'alignement
    col_widths = {
        'numero': 5,
        'cote': 10,
        'titre': 40,
        'categorie': 20,
        'date': 12
    }
    
    # En-tête du tableau
    header = (
        f"{'N°'.ljust(col_widths['numero'])}"
        f"{'COTE'.ljust(col_widths['cote'])}"
        f"{'TITRE'.ljust(col_widths['titre'])}"
        f"{'CATÉGORIE'.ljust(col_widths['categorie'])}"
        f"{'DATE'.ljust(col_widths['date'])}"
    )
    
    content += header + '\n'
    content += '-' * len(header) + '\n'
    
    # Lignes du tableau
    for piece in bordereau['pieces']:
        cote = piece.cote or f"P-{piece.numero:03d}"
        titre = piece.titre[:col_widths['titre']-3] + '...' if len(piece.titre) > col_widths['titre'] else piece.titre
        categorie = piece.categorie[:col_widths['categorie']-3] + '...' if len(piece.categorie) > col_widths['categorie'] else piece.categorie
        date_str = piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A'
        
        line = (
            f"{str(piece.numero).ljust(col_widths['numero'])}"
            f"{cote.ljust(col_widths['cote'])}"
            f"{titre.ljust(col_widths['titre'])}"
            f"{categorie.ljust(col_widths['categorie'])}"
            f"{date_str.ljust(col_widths['date'])}"
        )
        
        content += line + '\n'
        
        # Ajouter la description sur une ligne séparée si elle existe
        if piece.description:
            content += f"     → {piece.description}\n"
    
    content += '\n' + bordereau['footer']
    
    return content

def validate_bordereau(bordereau: dict) -> List[str]:
    """Valide un bordereau et retourne les erreurs"""
    errors = []
    
    if not bordereau.get('pieces'):
        errors.append("Aucune pièce dans le bordereau")
    
    # Vérifier les numéros de pièces
    numeros = [p.numero for p in bordereau.get('pieces', [])]
    if len(numeros) != len(set(numeros)):
        errors.append("Numéros de pièces en double détectés")
    
    # Vérifier les cotes
    cotes = [p.cote for p in bordereau.get('pieces', []) if p.cote]
    if len(cotes) != len(set(cotes)):
        errors.append("Cotes en double détectées")
    
    return errors

def generate_bordereau_summary(bordereau: dict) -> str:
    """Génère un résumé du bordereau"""
    pieces = bordereau.get('pieces', [])
    
    summary = f"""RÉSUMÉ DU BORDEREAU
    
Référence : {bordereau['metadata'].get('reference', 'N/A')}
Date : {bordereau['metadata']['created_at'].strftime('%d/%m/%Y')}
Nombre total de pièces : {len(pieces)}

Répartition par catégorie :
"""
    
    # Compter par catégorie
    from collections import Counter
    categories = Counter(p.categorie for p in pieces)
    
    for cat, count in categories.most_common():
        summary += f"- {cat} : {count} pièce(s)\n"
    
    return summary