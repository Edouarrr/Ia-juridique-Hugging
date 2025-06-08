# modules/timeline.py
"""Module de génération de chronologies (timelines) juridiques"""

import streamlit as st
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple
import re
from collections import defaultdict, OrderedDict

try:
    import plotly.graph_objects as go
    import plotly.express as px
    from plotly.subplots import make_subplots
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

from models.dataclasses import TimelineEvent, Document
from utils.helpers import extract_dates, extract_entities, format_legal_date
from managers.multi_llm_manager import MultiLLMManager

def process_timeline_request(query: str, analysis: dict):
    """Traite une demande de création de chronologie"""
    
    st.markdown("### ⏱️ Génération de chronologie")
    
    # Collecter les documents disponibles
    documents = collect_documents_for_timeline(analysis)
    
    if not documents:
        st.warning("⚠️ Aucun document disponible pour créer une chronologie")
        return
    
    # Configuration
    config = display_timeline_config_interface(documents, analysis)
    
    if st.button("🚀 Générer la chronologie", key="generate_timeline", type="primary"):
        with st.spinner("🔄 Extraction des événements en cours..."):
            timeline_result = generate_timeline(documents, config, analysis)
            
            if timeline_result:
                st.session_state.timeline_result = timeline_result
                display_timeline_results(timeline_result)

def collect_documents_for_timeline(analysis: dict) -> List[Dict[str, Any]]:
    """Collecte les documents pertinents pour la chronologie"""
    documents = []
    
    # Documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        documents.append({
            'id': doc_id,
            'title': doc.title,
            'content': doc.content,
            'source': doc.source,
            'metadata': doc.metadata
        })
    
    # Si référence spécifique
    if analysis.get('reference'):
        # Filtrer par référence
        ref_lower = analysis['reference'].lower()
        documents = [d for d in documents if ref_lower in d['title'].lower() or ref_lower in d.get('source', '').lower()]
    
    return documents

def display_timeline_config_interface(documents: List[Dict[str, Any]], analysis: dict) -> dict:
    """Interface de configuration pour la chronologie"""
    
    config = {}
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Type de chronologie
        config['timeline_type'] = st.selectbox(
            "📊 Type de chronologie",
            ["faits", "procedure", "financiere", "relationnelle", "complete"],
            format_func=lambda x: {
                "faits": "📅 Chronologie des faits",
                "procedure": "⚖️ Chronologie procédurale",
                "financiere": "💰 Chronologie financière",
                "relationnelle": "🤝 Chronologie relationnelle",
                "complete": "🌐 Chronologie complète"
            }.get(x, x.title()),
            key="timeline_type_select"
        )
        
        # Période
        config['date_range'] = st.radio(
            "📆 Période",
            ["auto", "custom"],
            format_func=lambda x: "🔍 Détection automatique" if x == "auto" else "📅 Personnalisée",
            key="date_range_select"
        )
    
    with col2:
        if config['date_range'] == "custom":
            config['start_date'] = st.date_input(
                "Date de début",
                value=datetime.now() - timedelta(days=365),
                key="start_date_input"
            )
            config['end_date'] = st.date_input(
                "Date de fin",
                value=datetime.now(),
                key="end_date_input"
            )
        
        # Granularité
        config['granularity'] = st.selectbox(
            "🔍 Granularité",
            ["jour", "semaine", "mois", "trimestre", "annee"],
            format_func=lambda x: {
                "jour": "📅 Par jour",
                "semaine": "📅 Par semaine",
                "mois": "📅 Par mois",
                "trimestre": "📅 Par trimestre",
                "annee": "📅 Par année"
            }.get(x, x.title()),
            key="granularity_select"
        )
    
    with col3:
        # Options d'affichage
        config['show_actors'] = st.checkbox(
            "👥 Afficher les acteurs",
            value=True,
            key="show_actors_check"
        )
        
        config['show_sources'] = st.checkbox(
            "📄 Afficher les sources",
            value=True,
            key="show_sources_check"
        )
        
        config['group_by_category'] = st.checkbox(
            "📁 Grouper par catégorie",
            value=False,
            key="group_by_category_check"
        )
    
    # Sélection des documents
    with st.expander("📄 Documents à analyser", expanded=True):
        # Option pour tout sélectionner
        select_all = st.checkbox("Tout sélectionner", value=True, key="select_all_docs_timeline")
        
        selected_docs = []
        for i, doc in enumerate(documents):
            is_selected = st.checkbox(
                f"📄 {doc['title']}",
                value=select_all,
                key=f"select_doc_timeline_{i}"
            )
            if is_selected:
                selected_docs.append(doc)
        
        config['selected_documents'] = selected_docs
        
        st.info(f"📊 {len(selected_docs)} documents sélectionnés")
    
    # Filtres additionnels
    with st.expander("🔍 Filtres avancés", expanded=False):
        config['keywords'] = st.text_area(
            "Mots-clés (un par ligne)",
            placeholder="virement\ncontrat\nréunion",
            height=100,
            key="keywords_filter"
        )
        
        config['actors_filter'] = st.text_input(
            "Acteurs spécifiques",
            placeholder="Ex: Jean Dupont, Société XYZ",
            key="actors_filter"
        )
        
        config['min_importance'] = st.slider(
            "Importance minimale",
            1, 10, 5,
            key="min_importance_slider"
        )
    
    return config

def generate_timeline(documents: List[Dict[str, Any]], config: dict, analysis: dict) -> Dict[str, Any]:
    """Génère une chronologie à partir des documents"""
    
    # Extraire les événements
    all_events = extract_timeline_events(documents, config)
    
    if not all_events:
        st.warning("⚠️ Aucun événement extrait des documents")
        return None
    
    # Filtrer selon la configuration
    filtered_events = filter_timeline_events(all_events, config)
    
    # Enrichir avec l'IA si nécessaire
    if len(filtered_events) < 10 and len(documents) > 5:
        # Utiliser l'IA pour extraire plus d'événements
        filtered_events = enrich_timeline_with_ai(documents, filtered_events, config)
    
    # Trier par date
    filtered_events.sort(key=lambda e: e.date if isinstance(e.date, datetime) else datetime.min)
    
    # Analyser la timeline
    analysis_result = analyze_timeline(filtered_events, config)
    
    # Créer la visualisation
    visualization = None
    if PLOTLY_AVAILABLE:
        visualization = create_timeline_visualization(filtered_events, config, analysis_result)
    
    return {
        'type': config['timeline_type'],
        'events': filtered_events,
        'analysis': analysis_result,
        'visualization': visualization,
        'document_count': len(config['selected_documents']),
        'event_count': len(filtered_events),
        'config': config,
        'timestamp': datetime.now()
    }

def extract_timeline_events(documents: List[Dict[str, Any]], config: dict) -> List[TimelineEvent]:
    """Extrait les événements des documents"""
    
    events = []
    
    for doc in documents:
        # Extraire les dates et leur contexte
        doc_dates = extract_dates(doc['content'])
        
        for date, context in doc_dates:
            # Créer un événement
            event = create_event_from_context(date, context, doc, config)
            if event:
                events.append(event)
        
        # Extraire aussi les événements sans date explicite mais avec marqueurs temporels
        temporal_events = extract_temporal_events(doc['content'], doc)
        events.extend(temporal_events)
    
    return events

def create_event_from_context(date: datetime, context: str, doc: Dict[str, Any], config: dict) -> Optional[TimelineEvent]:
    """Crée un événement à partir d'un contexte daté"""
    
    # Déterminer la catégorie selon le type de timeline
    category = determine_event_category(context, config['timeline_type'])
    
    if not category and config['timeline_type'] != 'complete':
        return None
    
    # Extraire les acteurs
    entities = extract_entities(context)
    actors = entities.get('persons', []) + entities.get('organizations', [])
    
    # Déterminer l'importance
    importance = calculate_event_importance(context, config)
    
    # Créer l'événement
    return TimelineEvent(
        date=date,
        description=clean_event_description(context),
        actors=actors[:5],  # Limiter à 5 acteurs
        location=entities.get('locations', [None])[0],
        category=category,
        importance=importance,
        source=doc.get('title', 'Document'),
        evidence=[doc.get('id', doc.get('title'))]
    )

def extract_temporal_events(content: str, doc: Dict[str, Any]) -> List[TimelineEvent]:
    """Extrait les événements avec marqueurs temporels relatifs"""
    
    events = []
    
    # Patterns temporels relatifs
    temporal_patterns = [
        (r'([Aa]u début|[Aa]u commencement)', -365),  # Un an avant
        (r'([Pp]ar la suite|[Ee]nsuite|[Pp]uis)', 30),  # Un mois après
        (r'([Qq]uelques jours plus tard)', 5),
        (r'([Ll]a veille)', -1),
        (r'([Ll]e lendemain)', 1),
        (r'([Uu]ne semaine plus tard)', 7),
        (r'([Uu]n mois plus tard)', 30)
    ]
    
    # Date de référence (chercher la première date absolue)
    ref_dates = extract_dates(content)
    if not ref_dates:
        return events
    
    ref_date = ref_dates[0][0]
    
    # Chercher les marqueurs temporels
    for pattern, days_offset in temporal_patterns:
        matches = re.finditer(pattern + r'[^.]*\.', content)
        
        for match in matches:
            event_date = ref_date + timedelta(days=days_offset)
            context = match.group(0)
            
            event = TimelineEvent(
                date=event_date,
                description=clean_event_description(context),
                category="temporal",
                importance=5,
                source=doc.get('title', 'Document')
            )
            events.append(event)
    
    return events

def determine_event_category(context: str, timeline_type: str) -> Optional[str]:
    """Détermine la catégorie d'un événement selon le type de timeline"""
    
    context_lower = context.lower()
    
    categories = {
        'faits': {
            'keywords': ['fait', 'événement', 'incident', 'action', 'acte'],
            'default': 'fait'
        },
        'procedure': {
            'keywords': ['plainte', 'audition', 'perquisition', 'jugement', 'appel', 'pourvoi'],
            'default': None
        },
        'financiere': {
            'keywords': ['virement', 'paiement', 'facture', 'compte', 'euros', '€', 'montant'],
            'default': None
        },
        'relationnelle': {
            'keywords': ['rencontre', 'réunion', 'contrat', 'accord', 'partenariat', 'relation'],
            'default': None
        },
        'complete': {
            'keywords': [],
            'default': 'general'
        }
    }
    
    type_config = categories.get(timeline_type, categories['complete'])
    
    # Vérifier les mots-clés
    for keyword in type_config['keywords']:
        if keyword in context_lower:
            return timeline_type
    
    return type_config['default']

def calculate_event_importance(context: str, config: dict) -> int:
    """Calcule l'importance d'un événement (1-10)"""
    
    importance = 5  # Base
    
    context_lower = context.lower()
    
    # Mots augmentant l'importance
    high_importance_words = [
        'crucial', 'décisif', 'majeur', 'important', 'signature', 'jugement',
        'condamnation', 'millions', 'principal', 'escroquerie', 'fraude'
    ]
    
    # Mots diminuant l'importance
    low_importance_words = [
        'mineur', 'secondaire', 'accessoire', 'simple', 'courrier', 'email'
    ]
    
    # Ajuster selon les mots trouvés
    for word in high_importance_words:
        if word in context_lower:
            importance += 2
    
    for word in low_importance_words:
        if word in context_lower:
            importance -= 1
    
    # Montants élevés
    amounts = re.findall(r'(\d+(?:\.\d{3})*(?:,\d{2})?)\s*(?:€|EUR|euros?)', context)
    if amounts:
        for amount_str in amounts:
            try:
                amount = float(amount_str.replace('.', '').replace(',', '.'))
                if amount > 100000:
                    importance += 2
                elif amount > 10000:
                    importance += 1
            except:
                pass
    
    # Limiter entre 1 et 10
    return max(1, min(10, importance))

def clean_event_description(context: str) -> str:
    """Nettoie la description d'un événement"""
    
    # Supprimer les références de dates redondantes
    context = re.sub(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', '', context)
    
    # Nettoyer les espaces
    context = ' '.join(context.split())
    
    # Limiter la longueur
    if len(context) > 200:
        context = context[:197] + "..."
    
    return context.strip()

def filter_timeline_events(events: List[TimelineEvent], config: dict) -> List[TimelineEvent]:
    """Filtre les événements selon la configuration"""
    
    filtered = events
    
    # Filtre par période
    if config['date_range'] == 'custom':
        start = datetime.combine(config['start_date'], datetime.min.time())
        end = datetime.combine(config['end_date'], datetime.max.time())
        filtered = [e for e in filtered if start <= e.date <= end]
    
    # Filtre par importance
    min_importance = config.get('min_importance', 1)
    filtered = [e for e in filtered if e.importance >= min_importance]
    
    # Filtre par mots-clés
    if config.get('keywords'):
        keywords = [k.strip().lower() for k in config['keywords'].split('\n') if k.strip()]
        if keywords:
            filtered = [
                e for e in filtered 
                if any(kw in e.description.lower() for kw in keywords)
            ]
    
    # Filtre par acteurs
    if config.get('actors_filter'):
        actors = [a.strip().lower() for a in config['actors_filter'].split(',') if a.strip()]
        if actors:
            filtered = [
                e for e in filtered
                if any(any(actor in a.lower() for a in e.actors) for actor in actors)
            ]
    
    return filtered

def enrich_timeline_with_ai(documents: List[Dict[str, Any]], existing_events: List[TimelineEvent], config: dict) -> List[TimelineEvent]:
    """Enrichit la timeline avec l'IA pour extraire plus d'événements"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return existing_events
    
    # Construire le prompt
    prompt = f"""Analyse ces documents et extrais TOUS les événements datés ou datables pour créer une chronologie {config['timeline_type']}.

DOCUMENTS:
"""
    
    for doc in documents[:10]:  # Limiter à 10 documents
        prompt += f"\n--- {doc['title']} ---\n{doc['content'][:1000]}...\n"
    
    prompt += f"""
Extrais les événements au format suivant:
- Date (format JJ/MM/AAAA ou approximation)
- Description courte et précise
- Acteurs impliqués
- Importance (1-10)

Focus sur les événements de type : {config['timeline_type']}
"""
    
    # Interroger l'IA
    provider = list(llm_manager.clients.keys())[0]
    response = llm_manager.query_single_llm(
        provider,
        prompt,
        "Tu es un expert en analyse chronologique de documents juridiques.",
        temperature=0.3
    )
    
    if response['success']:
        # Parser la réponse pour extraire les événements
        new_events = parse_ai_timeline_response(response['response'], documents[0])
        
        # Fusionner avec les événements existants
        all_events = existing_events + new_events
        
        # Dédupliquer
        unique_events = []
        seen = set()
        
        for event in all_events:
            key = (event.date, event.description[:50])
            if key not in seen:
                seen.add(key)
                unique_events.append(event)
        
        return unique_events
    
    return existing_events

def parse_ai_timeline_response(response: str, source_doc: Dict[str, Any]) -> List[TimelineEvent]:
    """Parse la réponse de l'IA pour extraire les événements"""
    
    events = []
    
    # Chercher les lignes qui ressemblent à des événements
    lines = response.split('\n')
    
    for line in lines:
        # Pattern : date - description
        match = re.match(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\s*[-:]\s*(.+)', line)
        if match:
            date_str, description = match.groups()
            
            # Parser la date
            date = extract_dates(date_str)
            if date:
                event = TimelineEvent(
                    date=date[0][0],
                    description=description.strip(),
                    importance=5,
                    source=source_doc.get('title', 'IA Analysis'),
                    category='ai_extracted'
                )
                events.append(event)
    
    return events

def analyze_timeline(events: List[TimelineEvent], config: dict) -> Dict[str, Any]:
    """Analyse la chronologie pour en extraire des insights"""
    
    analysis = {
        'total_events': len(events),
        'date_range': None,
        'busiest_periods': [],
        'key_actors': [],
        'event_distribution': {},
        'patterns': []
    }
    
    if not events:
        return analysis
    
    # Plage de dates
    dates = [e.date for e in events if isinstance(e.date, datetime)]
    if dates:
        analysis['date_range'] = {
            'start': min(dates),
            'end': max(dates),
            'duration_days': (max(dates) - min(dates)).days
        }
    
    # Périodes les plus actives
    date_counts = defaultdict(int)
    for event in events:
        if isinstance(event.date, datetime):
            # Grouper selon la granularité
            if config['granularity'] == 'jour':
                key = event.date.date()
            elif config['granularity'] == 'semaine':
                key = event.date.isocalendar()[1]  # Numéro de semaine
            elif config['granularity'] == 'mois':
                key = f"{event.date.year}-{event.date.month:02d}"
            elif config['granularity'] == 'trimestre':
                key = f"{event.date.year}-Q{(event.date.month - 1) // 3 + 1}"
            else:  # année
                key = event.date.year
            
            date_counts[key] += 1
    
    # Top 5 périodes
    analysis['busiest_periods'] = sorted(
        date_counts.items(),
        key=lambda x: x[1],
        reverse=True
    )[:5]
    
    # Acteurs clés
    actor_counts = defaultdict(int)
    for event in events:
        for actor in event.actors:
            actor_counts[actor] += 1
    
    analysis['key_actors'] = sorted(
        actor_counts.items(),
        key=lambda x: x[1],
        reverse=True
    )[:10]
    
    # Distribution par catégorie
    category_counts = defaultdict(int)
    for event in events:
        category_counts[event.category or 'autre'] += 1
    
    analysis['event_distribution'] = dict(category_counts)
    
    # Patterns détectés
    analysis['patterns'] = detect_timeline_patterns(events)
    
    return analysis

def detect_timeline_patterns(events: List[TimelineEvent]) -> List[Dict[str, Any]]:
    """Détecte des patterns dans la chronologie"""
    
    patterns = []
    
    # Pattern 1: Accélération des événements
    if len(events) >= 10:
        # Diviser en quartiles
        quarter = len(events) // 4
        first_quarter_density = calculate_event_density(events[:quarter])
        last_quarter_density = calculate_event_density(events[-quarter:])
        
        if last_quarter_density > first_quarter_density * 2:
            patterns.append({
                'type': 'acceleration',
                'description': 'Forte accélération des événements dans la période récente',
                'factor': last_quarter_density / max(first_quarter_density, 0.1)
            })
    
    # Pattern 2: Événements récurrents
    recurring = find_recurring_events(events)
    if recurring:
        patterns.append({
            'type': 'recurrence',
            'description': f'{len(recurring)} types d\'événements récurrents détectés',
            'events': recurring
        })
    
    # Pattern 3: Pics d'activité
    peaks = find_activity_peaks(events)
    if peaks:
        patterns.append({
            'type': 'peaks',
            'description': f'{len(peaks)} pics d\'activité identifiés',
            'peaks': peaks
        })
    
    return patterns

def calculate_event_density(events: List[TimelineEvent]) -> float:
    """Calcule la densité d'événements (événements/jour)"""
    
    if len(events) < 2:
        return 0.0
    
    dates = [e.date for e in events if isinstance(e.date, datetime)]
    if len(dates) < 2:
        return 0.0
    
    duration = (max(dates) - min(dates)).days
    if duration == 0:
        return len(dates)
    
    return len(dates) / duration

def find_recurring_events(events: List[TimelineEvent]) -> List[Dict[str, Any]]:
    """Trouve les événements récurrents"""
    
    # Grouper par similarité de description
    event_groups = defaultdict(list)
    
    for event in events:
        # Extraire les mots-clés principaux
        keywords = extract_event_keywords(event.description)
        key = tuple(sorted(keywords))
        
        if key:
            event_groups[key].append(event)
    
    # Identifier les récurrences
    recurring = []
    for keywords, group_events in event_groups.items():
        if len(group_events) >= 3:  # Au moins 3 occurrences
            recurring.append({
                'keywords': list(keywords),
                'count': len(group_events),
                'first': min(e.date for e in group_events if isinstance(e.date, datetime)),
                'last': max(e.date for e in group_events if isinstance(e.date, datetime))
            })
    
    return recurring

def extract_event_keywords(description: str) -> List[str]:
    """Extrait les mots-clés d'une description d'événement"""
    
    # Mots à ignorer
    stopwords = {'le', 'la', 'les', 'de', 'du', 'des', 'un', 'une', 'et', 'ou', 'à', 'pour'}
    
    # Extraire les mots significatifs
    words = re.findall(r'\b\w{4,}\b', description.lower())
    keywords = [w for w in words if w not in stopwords]
    
    return keywords[:3]  # Top 3

def find_activity_peaks(events: List[TimelineEvent]) -> List[Dict[str, Any]]:
    """Trouve les pics d'activité"""
    
    # Grouper par jour
    daily_counts = defaultdict(list)
    
    for event in events:
        if isinstance(event.date, datetime):
            day = event.date.date()
            daily_counts[day].append(event)
    
    # Calculer la moyenne
    avg_events_per_day = len(events) / len(daily_counts) if daily_counts else 0
    
    # Identifier les pics (jours avec 2x plus d'événements que la moyenne)
    peaks = []
    for day, day_events in daily_counts.items():
        if len(day_events) >= avg_events_per_day * 2:
            peaks.append({
                'date': day,
                'event_count': len(day_events),
                'events': day_events,
                'factor': len(day_events) / max(avg_events_per_day, 1)
            })
    
    return sorted(peaks, key=lambda x: x['event_count'], reverse=True)[:5]

def create_timeline_visualization(events: List[TimelineEvent], config: dict, analysis: Dict[str, Any]) -> Any:
    """Crée la visualisation Plotly de la timeline"""
    
    if not PLOTLY_AVAILABLE:
        return None
    
    # Préparer les données
    data = prepare_timeline_data(events, config)
    
    # Type de visualisation selon la configuration
    if config.get('group_by_category'):
        fig = create_grouped_timeline(data, config, analysis)
    else:
        fig = create_linear_timeline(data, config, analysis)
    
    # Mise en forme générale
    fig.update_layout(
        title=f"Chronologie {config['timeline_type'].title()}",
        height=600,
        showlegend=True,
        hovermode='closest',
        template='plotly_white'
    )
    
    return fig

def prepare_timeline_data(events: List[TimelineEvent], config: dict) -> pd.DataFrame:
    """Prépare les données pour la visualisation"""
    
    if not PANDAS_AVAILABLE:
        # Fallback sans pandas
        return []
    
    data = []
    
    for event in events:
        if isinstance(event.date, datetime):
            row = {
                'date': event.date,
                'description': event.description[:100] + '...' if len(event.description) > 100 else event.description,
                'full_description': event.description,
                'actors': ', '.join(event.actors) if event.actors else 'N/A',
                'category': event.category or 'autre',
                'importance': event.importance,
                'source': event.source
            }
            data.append(row)
    
    return pd.DataFrame(data)

def create_linear_timeline(data: pd.DataFrame, config: dict, analysis: Dict[str, Any]) -> go.Figure:
    """Crée une timeline linéaire"""
    
    fig = go.Figure()
    
    # Couleurs par catégorie
    category_colors = {
        'faits': '#1f77b4',
        'procedure': '#ff7f0e',
        'financiere': '#2ca02c',
        'relationnelle': '#d62728',
        'autre': '#9467bd'
    }
    
    # Ajouter les événements
    for category in data['category'].unique():
        cat_data = data[data['category'] == category]
        
        fig.add_trace(go.Scatter(
            x=cat_data['date'],
            y=cat_data['importance'],
            mode='markers+text',
            name=category.title(),
            marker=dict(
                size=cat_data['importance'] * 3,
                color=category_colors.get(category, '#666'),
                symbol='circle'
            ),
            text=cat_data['description'],
            textposition='top center',
            hovertemplate='<b>%{text}</b><br>' +
                         'Date: %{x}<br>' +
                         'Importance: %{y}<br>' +
                         'Acteurs: %{customdata[0]}<br>' +
                         'Source: %{customdata[1]}<extra></extra>',
            customdata=cat_data[['actors', 'source']]
        ))
    
    # Mise en forme
    fig.update_xaxis(
        title="Date",
        rangeslider=dict(visible=True),
        type="date"
    )
    
    fig.update_yaxis(
        title="Importance",
        range=[0, 11]
    )
    
    return fig

def create_grouped_timeline(data: pd.DataFrame, config: dict, analysis: Dict[str, Any]) -> go.Figure:
    """Crée une timeline groupée par catégorie"""
    
    # Créer des subplots pour chaque catégorie
    categories = data['category'].unique()
    n_categories = len(categories)
    
    fig = make_subplots(
        rows=n_categories,
        cols=1,
        shared_xaxes=True,
        subplot_titles=[cat.title() for cat in categories],
        vertical_spacing=0.05
    )
    
    # Couleurs
    colors = px.colors.qualitative.Set3
    
    # Ajouter les traces
    for i, category in enumerate(categories):
        cat_data = data[data['category'] == category]
        
        fig.add_trace(
            go.Scatter(
                x=cat_data['date'],
                y=[1] * len(cat_data),  # Tous au même niveau
                mode='markers+text',
                name=category.title(),
                marker=dict(
                    size=cat_data['importance'] * 2,
                    color=colors[i % len(colors)],
                    symbol='circle'
                ),
                text=cat_data['description'],
                textposition='top center',
                showlegend=False,
                hovertemplate='<b>%{text}</b><br>' +
                             'Date: %{x}<br>' +
                             'Importance: %{customdata[0]}<br>' +
                             'Acteurs: %{customdata[1]}<extra></extra>',
                customdata=cat_data[['importance', 'actors']]
            ),
            row=i+1,
            col=1
        )
        
        # Masquer l'axe Y
        fig.update_yaxes(showticklabels=False, row=i+1, col=1)
    
    # Mise en forme
    fig.update_xaxes(title="Date", row=n_categories, col=1)
    
    return fig

def display_timeline_results(timeline_result: Dict[str, Any]):
    """Affiche les résultats de la chronologie"""
    
    st.success(f"✅ Chronologie générée : {timeline_result['event_count']} événements")
    
    # Métadonnées
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Type", timeline_result['type'].title())
    
    with col2:
        st.metric("Événements", timeline_result['event_count'])
    
    with col3:
        st.metric("Documents analysés", timeline_result['document_count'])
    
    with col4:
        if timeline_result['analysis']['date_range']:
            duration = timeline_result['analysis']['date_range']['duration_days']
            st.metric("Période", f"{duration} jours")
    
    # Visualisation
    if timeline_result.get('visualization'):
        st.plotly_chart(timeline_result['visualization'], use_container_width=True)
    else:
        st.info("Visualisation non disponible - Installez plotly pour les graphiques")
    
    # Analyse
    with st.expander("📊 Analyse de la chronologie", expanded=True):
        display_timeline_analysis(timeline_result['analysis'])
    
    # Liste des événements
    with st.expander("📅 Liste détaillée des événements", expanded=False):
        display_events_list(timeline_result['events'], timeline_result['config'])
    
    # Actions
    st.markdown("### 💾 Actions")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("📄 Exporter Word", key="export_timeline_docx"):
            docx_content = export_timeline_to_docx(timeline_result)
            st.download_button(
                "💾 Télécharger DOCX",
                docx_content,
                f"chronologie_{timeline_result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_timeline_docx"
            )
    
    with col2:
        if st.button("📊 Exporter Excel", key="export_timeline_xlsx"):
            xlsx_content = export_timeline_to_excel(timeline_result)
            st.download_button(
                "💾 Télécharger XLSX",
                xlsx_content,
                f"chronologie_{timeline_result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_timeline_xlsx"
            )
    
    with col3:
        if st.button("📈 Statistiques", key="timeline_stats"):
            show_timeline_statistics(timeline_result)
    
    with col4:
        if st.button("🔍 Filtrer", key="filter_timeline"):
            st.session_state.show_timeline_filters = True

def display_timeline_analysis(analysis: Dict[str, Any]):
    """Affiche l'analyse de la chronologie"""
    
    # Période couverte
    if analysis['date_range']:
        st.markdown("#### 📅 Période analysée")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.write(f"**Début :** {analysis['date_range']['start'].strftime('%d/%m/%Y')}")
        
        with col2:
            st.write(f"**Fin :** {analysis['date_range']['end'].strftime('%d/%m/%Y')}")
        
        with col3:
            st.write(f"**Durée :** {analysis['date_range']['duration_days']} jours")
    
    # Périodes actives
    if analysis['busiest_periods']:
        st.markdown("#### 🔥 Périodes les plus actives")
        for period, count in analysis['busiest_periods']:
            st.write(f"- **{period}** : {count} événements")
    
    # Acteurs clés
    if analysis['key_actors']:
        st.markdown("#### 👥 Acteurs principaux")
        for actor, count in analysis['key_actors'][:5]:
            st.write(f"- **{actor}** : impliqué dans {count} événements")
    
    # Patterns
    if analysis['patterns']:
        st.markdown("#### 🔍 Patterns détectés")
        for pattern in analysis['patterns']:
            if pattern['type'] == 'acceleration':
                st.warning(f"⚡ {pattern['description']} (facteur x{pattern['factor']:.1f})")
            elif pattern['type'] == 'recurrence':
                st.info(f"🔄 {pattern['description']}")
            elif pattern['type'] == 'peaks':
                st.error(f"📊 {pattern['description']}")

def display_events_list(events: List[TimelineEvent], config: dict):
    """Affiche la liste détaillée des événements"""
    
    # Options de tri
    sort_by = st.selectbox(
        "Trier par",
        ["Date", "Importance", "Catégorie"],
        key="sort_timeline_events"
    )
    
    # Trier les événements
    if sort_by == "Date":
        sorted_events = sorted(events, key=lambda e: e.date if isinstance(e.date, datetime) else datetime.min)
    elif sort_by == "Importance":
        sorted_events = sorted(events, key=lambda e: e.importance, reverse=True)
    else:  # Catégorie
        sorted_events = sorted(events, key=lambda e: e.category or 'z')
    
    # Afficher
    for i, event in enumerate(sorted_events, 1):
        with st.container():
            col1, col2, col3, col4 = st.columns([1, 3, 1, 1])
            
            with col1:
                date_str = event.date.strftime('%d/%m/%Y') if isinstance(event.date, datetime) else str(event.date)
                st.write(f"**{date_str}**")
            
            with col2:
                st.write(event.description)
                if event.actors and config.get('show_actors'):
                    st.caption(f"👥 {', '.join(event.actors)}")
            
            with col3:
                importance_color = "🔴" if event.importance >= 8 else "🟡" if event.importance >= 5 else "🟢"
                st.write(f"{importance_color} {event.importance}/10")
            
            with col4:
                if config.get('show_sources'):
                    st.caption(event.source)
            
            if i < len(sorted_events):
                st.divider()

def export_timeline_to_docx(timeline_result: Dict[str, Any]) -> bytes:
    """Exporte la chronologie vers Word"""
    
    try:
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Titre
        title = doc.add_heading(f'Chronologie {timeline_result["type"].title()}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Métadonnées
        doc.add_paragraph(f"Générée le : {timeline_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}")
        doc.add_paragraph(f"Nombre d'événements : {timeline_result['event_count']}")
        doc.add_paragraph(f"Documents analysés : {timeline_result['document_count']}")
        
        if timeline_result['analysis']['date_range']:
            doc.add_paragraph(
                f"Période : du {timeline_result['analysis']['date_range']['start'].strftime('%d/%m/%Y')} "
                f"au {timeline_result['analysis']['date_range']['end'].strftime('%d/%m/%Y')}"
            )
        
        doc.add_paragraph()
        
        # Événements par date
        doc.add_heading('Chronologie détaillée', 1)
        
        # Grouper par mois
        events_by_month = defaultdict(list)
        for event in timeline_result['events']:
            if isinstance(event.date, datetime):
                month_key = event.date.strftime('%B %Y')
                events_by_month[month_key].append(event)
        
        # Afficher par mois
        for month, month_events in sorted(events_by_month.items()):
            doc.add_heading(month, 2)
            
            for event in sorted(month_events, key=lambda e: e.date):
                # Date et description
                p = doc.add_paragraph()
                p.add_run(event.date.strftime('%d/%m/%Y')).bold = True
                p.add_run(f" - {event.description}")
                
                # Détails
                if event.actors:
                    doc.add_paragraph(f"   Acteurs : {', '.join(event.actors)}", style='List Bullet')
                if event.source:
                    doc.add_paragraph(f"   Source : {event.source}", style='List Bullet')
                doc.add_paragraph()
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError:
        # Fallback texte
        return export_timeline_to_text(timeline_result).encode('utf-8')

def export_timeline_to_excel(timeline_result: Dict[str, Any]) -> bytes:
    """Exporte la chronologie vers Excel"""
    
    if not PANDAS_AVAILABLE:
        return export_timeline_to_text(timeline_result).encode('utf-8')
    
    try:
        import io
        
        # Créer un writer Excel
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            # Feuille des événements
            events_data = []
            for event in timeline_result['events']:
                events_data.append({
                    'Date': event.date.strftime('%d/%m/%Y') if isinstance(event.date, datetime) else str(event.date),
                    'Heure': event.date.strftime('%H:%M') if isinstance(event.date, datetime) else '',
                    'Description': event.description,
                    'Acteurs': ', '.join(event.actors) if event.actors else '',
                    'Lieu': event.location or '',
                    'Catégorie': event.category or '',
                    'Importance': event.importance,
                    'Source': event.source
                })
            
            df_events = pd.DataFrame(events_data)
            df_events.to_excel(writer, sheet_name='Événements', index=False)
            
            # Feuille d'analyse
            analysis_data = []
            
            # Acteurs clés
            for actor, count in timeline_result['analysis']['key_actors']:
                analysis_data.append({
                    'Type': 'Acteur clé',
                    'Valeur': actor,
                    'Occurrences': count
                })
            
            # Périodes actives
            for period, count in timeline_result['analysis']['busiest_periods']:
                analysis_data.append({
                    'Type': 'Période active',
                    'Valeur': str(period),
                    'Occurrences': count
                })
            
            if analysis_data:
                df_analysis = pd.DataFrame(analysis_data)
                df_analysis.to_excel(writer, sheet_name='Analyse', index=False)
            
            # Ajuster les largeurs de colonnes
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur export Excel: {e}")
        return export_timeline_to_text(timeline_result).encode('utf-8')

def export_timeline_to_text(timeline_result: Dict[str, Any]) -> str:
    """Exporte la chronologie en format texte"""
    
    content = f"CHRONOLOGIE {timeline_result['type'].upper()}\n"
    content += "=" * 60 + "\n\n"
    
    content += f"Générée le : {timeline_result['timestamp'].strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Nombre d'événements : {timeline_result['event_count']}\n"
    content += f"Documents analysés : {timeline_result['document_count']}\n"
    
    if timeline_result['analysis']['date_range']:
        content += (
            f"Période : du {timeline_result['analysis']['date_range']['start'].strftime('%d/%m/%Y')} "
            f"au {timeline_result['analysis']['date_range']['end'].strftime('%d/%m/%Y')}\n"
        )
    
    content += "\n" + "=" * 60 + "\n"
    content += "ÉVÉNEMENTS CHRONOLOGIQUES\n"
    content += "=" * 60 + "\n\n"
    
    # Événements par date
    for event in sorted(timeline_result['events'], key=lambda e: e.date if isinstance(e.date, datetime) else datetime.min):
        date_str = event.date.strftime('%d/%m/%Y') if isinstance(event.date, datetime) else str(event.date)
        content += f"{date_str} - {event.description}\n"
        
        if event.actors:
            content += f"   Acteurs : {', '.join(event.actors)}\n"
        if event.location:
            content += f"   Lieu : {event.location}\n"
        if event.source:
            content += f"   Source : {event.source}\n"
        content += f"   Importance : {event.importance}/10\n"
        content += "\n"
    
    return content

def show_timeline_statistics(timeline_result: Dict[str, Any]):
    """Affiche des statistiques détaillées de la timeline"""
    
    st.markdown("### 📊 Statistiques détaillées")
    
    events = timeline_result['events']
    
    # Statistiques temporelles
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Total événements", len(events))
        
        # Événements par mois
        events_by_month = defaultdict(int)
        for event in events:
            if isinstance(event.date, datetime):
                month_key = event.date.strftime('%Y-%m')
                events_by_month[month_key] += 1
        
        avg_per_month = len(events) / len(events_by_month) if events_by_month else 0
        st.metric("Moyenne/mois", f"{avg_per_month:.1f}")
    
    with col2:
        # Importance moyenne
        avg_importance = sum(e.importance for e in events) / len(events) if events else 0
        st.metric("Importance moyenne", f"{avg_importance:.1f}/10")
        
        # Événements critiques
        critical_events = [e for e in events if e.importance >= 8]
        st.metric("Événements critiques", len(critical_events))
    
    with col3:
        # Acteurs uniques
        all_actors = set()
        for event in events:
            all_actors.update(event.actors)
        st.metric("Acteurs uniques", len(all_actors))
        
        # Sources uniques
        sources = set(e.source for e in events if e.source)
        st.metric("Sources", len(sources))
    
    # Graphiques si plotly disponible
    if PLOTLY_AVAILABLE and PANDAS_AVAILABLE:
        # Distribution temporelle
        st.markdown("#### 📈 Distribution temporelle")
        
        # Créer un histogramme des événements par mois
        dates = [e.date for e in events if isinstance(e.date, datetime)]
        if dates:
            df_dates = pd.DataFrame({'date': dates})
            df_dates['month'] = df_dates['date'].dt.to_period('M')
            
            monthly_counts = df_dates.groupby('month').size().reset_index(name='count')
            
            fig = px.bar(
                monthly_counts,
                x='month',
                y='count',
                title="Nombre d'événements par mois"
            )
            
            st.plotly_chart(fig, use_container_width=True)
        
        # Distribution par catégorie
        st.markdown("#### 📊 Répartition par catégorie")
        
        categories = [e.category or 'autre' for e in events]
        if categories:
            df_cat = pd.DataFrame({'category': categories})
            cat_counts = df_cat['category'].value_counts()
            
            fig = px.pie(
                values=cat_counts.values,
                names=cat_counts.index,
                title="Distribution des événements par catégorie"
            )
            
            st.plotly_chart(fig, use_container_width=True)