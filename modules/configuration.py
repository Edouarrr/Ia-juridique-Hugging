#modules/configuration.py
"""Page de configuration de l'application"""

import streamlit as st
import json
import io
from datetime import datetime

from config.app_config import LLMProvider, get_llm_configs
from models.dataclasses import LetterheadTemplate, PieceSelectionnee
from managers.multi_llm_manager import MultiLLMManager
from utils.helpers import create_env_example

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def show_page():
    """Affiche la page de configuration"""
    st.header("⚙️ Configuration")
    
    tabs = st.tabs(["🔑 Clés API", "📄 Papier en-tête", "📊 État du système", "💾 Export/Import"])
    
    with tabs[0]:
        show_api_keys_tab()
    
    with tabs[1]:
        show_letterhead_tab()
    
    with tabs[2]:
        show_system_status_tab()
    
    with tabs[3]:
        show_export_import_tab()

def show_api_keys_tab():
    """Affiche l'onglet de configuration des clés API"""
    st.markdown("### Configuration des clés API")
    
    st.info("""
    ℹ️ Les clés API doivent être configurées dans les variables d'environnement ou dans un fichier .env
    
    Variables nécessaires:
    - AZURE_OPENAI_KEY
    - AZURE_OPENAI_ENDPOINT
    - AZURE_SEARCH_KEY
    - AZURE_SEARCH_ENDPOINT
    - AZURE_STORAGE_CONNECTION_STRING
    - ANTHROPIC_API_KEY
    - OPENAI_API_KEY
    - GOOGLE_API_KEY
    - PERPLEXITY_API_KEY
    """)
    
    # Créer un fichier .env exemple
    env_example = create_env_example()
    
    st.download_button(
        "📥 Télécharger un fichier .env exemple",
        env_example,
        ".env.example",
        "text/plain",
        key="download_env_example"
    )
    
    # Vérifier l'état des clés
    configs = get_llm_configs()
    
    st.markdown("#### 🤖 Providers IA")
    for provider in LLMProvider:
        col1, col2 = st.columns([3, 1])
        
        config = configs.get(provider, {})
        
        with col1:
            st.text(provider.value)
        
        with col2:
            if config.get('key') or config.get('api_key'):
                st.success("✅")
            else:
                st.error("❌")
    
    # Services Azure
    st.markdown("#### 🔷 Services Azure")
    
    services = {
        "Azure Blob Storage": bool(os.getenv('AZURE_STORAGE_CONNECTION_STRING')),
        "Azure Search": bool(os.getenv('AZURE_SEARCH_ENDPOINT')),
        "Azure OpenAI": bool(os.getenv('AZURE_OPENAI_ENDPOINT'))
    }
    
    for service, configured in services.items():
        col1, col2 = st.columns([3, 1])
        
        with col1:
            st.text(service)
        
        with col2:
            if configured:
                st.success("✅")
            else:
                st.error("❌")

def show_letterhead_tab():
    """Affiche l'onglet de configuration du papier en-tête"""
    st.markdown("### 📄 Configuration du papier en-tête")
    
    # Papier en-tête actuel
    if 'letterhead_template' in st.session_state and st.session_state.letterhead_template:
        show_current_letterhead()
    
    # Créer/Modifier papier en-tête
    show_letterhead_form()
    
    # Import de papier en-tête depuis Word
    show_letterhead_import()

def show_current_letterhead():
    """Affiche le papier en-tête actuel"""
    current_template = st.session_state.letterhead_template
    
    with st.expander("Papier en-tête actuel", expanded=True):
        st.text("En-tête :")
        st.code(current_template.header_content)
        st.text("Pied de page :")
        st.code(current_template.footer_content)
        
        if st.button("🗑️ Supprimer le papier en-tête actuel"):
            st.session_state.letterhead_template = None
            st.success("✅ Papier en-tête supprimé")
            st.rerun()

def show_letterhead_form():
    """Affiche le formulaire de création/modification du papier en-tête"""
    st.markdown("#### ➕ Créer/Modifier le papier en-tête")
    
    with st.form("letterhead_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            nom_template = st.text_input(
                "Nom du template",
                value=st.session_state.letterhead_template.name if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else "Papier en-tête principal"
            )
            
            header_content = st.text_area(
                "En-tête",
                value=st.session_state.letterhead_template.header_content if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else "",
                height=150,
                placeholder="""Cabinet d'avocats XYZ
Maître Jean DUPONT
Avocat au Barreau de Paris
123 rue de la Justice
75001 PARIS
Tél : 01 23 45 67 89
Email : contact@cabinet-xyz.fr"""
            )
            
            footer_content = st.text_area(
                "Pied de page",
                value=st.session_state.letterhead_template.footer_content if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else "",
                height=80,
                placeholder="Cabinet XYZ - 123 rue de la Justice, 75001 PARIS - Tél : 01 23 45 67 89"
            )
        
        with col2:
            # Paramètres de mise en forme
            st.markdown("**Mise en forme**")
            
            font_family = st.selectbox(
                "Police",
                ["Arial", "Times New Roman", "Calibri", "Garamond", "Helvetica"],
                index=0 if not ('letterhead_template' in st.session_state and st.session_state.letterhead_template) else ["Arial", "Times New Roman", "Calibri", "Garamond", "Helvetica"].index(st.session_state.letterhead_template.font_family)
            )
            
            font_size = st.number_input(
                "Taille de police",
                min_value=8,
                max_value=16,
                value=11 if not ('letterhead_template' in st.session_state and st.session_state.letterhead_template) else st.session_state.letterhead_template.font_size
            )
            
            line_spacing = st.number_input(
                "Interligne",
                min_value=1.0,
                max_value=2.0,
                step=0.1,
                value=1.5 if not ('letterhead_template' in st.session_state and st.session_state.letterhead_template) else st.session_state.letterhead_template.line_spacing
            )
            
            st.markdown("**Marges (cm)**")
            
            margin_top = st.number_input("Haut", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
            margin_bottom = st.number_input("Bas", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
            margin_left = st.number_input("Gauche", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
            margin_right = st.number_input("Droite", min_value=1.0, max_value=5.0, value=2.5, step=0.5)
        
        # Upload logo (optionnel)
        logo_file = st.file_uploader(
            "Logo (optionnel)",
            type=['png', 'jpg', 'jpeg'],
            key="letterhead_logo"
        )
        
        if st.form_submit_button("💾 Sauvegarder le papier en-tête", type="primary"):
            # Créer le template
            new_template = LetterheadTemplate(
                name=nom_template,
                header_content=header_content,
                footer_content=footer_content,
                font_family=font_family,
                font_size=font_size,
                line_spacing=line_spacing,
                margins={
                    'top': margin_top,
                    'bottom': margin_bottom,
                    'left': margin_left,
                    'right': margin_right
                }
            )
            
            # Sauvegarder le logo si uploadé
            if logo_file:
                st.session_state.letterhead_image = logo_file.read()
                new_template.logo_path = "logo_uploaded"
            
            st.session_state.letterhead_template = new_template
            st.success("✅ Papier en-tête sauvegardé avec succès!")
            st.rerun()

def show_letterhead_import():
    """Affiche l'import de papier en-tête depuis Word"""
    st.markdown("#### 📤 Importer depuis un document Word")
    
    uploaded_letterhead = st.file_uploader(
        "Charger un document Word avec papier en-tête",
        type=['docx'],
        key="upload_letterhead_word"
    )
    
    if uploaded_letterhead and st.button("📥 Extraire le papier en-tête"):
        if DOCX_AVAILABLE:
            try:
                doc = DocxDocument(uploaded_letterhead)
                
                # Extraire l'en-tête
                header_text = ""
                if doc.sections:
                    header = doc.sections[0].header
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            header_text += paragraph.text + "\n"
                
                # Extraire le pied de page
                footer_text = ""
                if doc.sections:
                    footer = doc.sections[0].footer
                    for paragraph in footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text += paragraph.text + "\n"
                
                if header_text or footer_text:
                    st.success("✅ Papier en-tête extrait avec succès!")
                    st.text("En-tête extrait :")
                    st.code(header_text)
                    st.text("Pied de page extrait :")
                    st.code(footer_text)
                    
                    if st.button("Utiliser ce papier en-tête"):
                        st.session_state.letterhead_template = LetterheadTemplate(
                            name="Papier en-tête importé",
                            header_content=header_text.strip(),
                            footer_content=footer_text.strip()
                        )
                        st.success("✅ Papier en-tête importé!")
                        st.rerun()
                else:
                    st.warning("⚠️ Aucun en-tête ou pied de page trouvé dans le document")
                    
            except Exception as e:
                st.error(f"❌ Erreur lors de l'extraction : {str(e)}")

def show_system_status_tab():
    """Affiche l'onglet d'état du système"""
    st.markdown("### État du système")
    
    # Métriques
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Documents Azure", len(st.session_state.azure_documents))
    
    with col2:
        st.metric("Pièces sélectionnées", len(st.session_state.pieces_selectionnees))
    
    with col3:
        llm_manager = MultiLLMManager()
        st.metric("IA disponibles", len(llm_manager.clients))
    
    with col4:
        st.metric("Styles appris", len(st.session_state.get('learned_styles', {})))
    
    # État des connexions
    st.markdown("### 🔗 État des connexions")
    
    # Azure Blob
    if st.session_state.azure_blob_manager and st.session_state.azure_blob_manager.is_connected():
        st.success("✅ Azure Blob Storage : Connecté")
        
        if st.button("🧪 Tester la connexion Blob", key="test_blob"):
            try:
                containers = st.session_state.azure_blob_manager.list_containers()
                st.success(f"✅ {len(containers)} containers trouvés : {', '.join(containers)}")
            except Exception as e:
                st.error(f"❌ Erreur : {str(e)}")
    else:
        st.error("❌ Azure Blob Storage : Non connecté")
    
    # Azure Search
    if st.session_state.azure_search_manager and st.session_state.azure_search_manager.search_client:
        st.success("✅ Azure Search : Connecté")
        
        if st.button("🧪 Tester la connexion Search", key="test_search"):
            try:
                # Tester avec une recherche simple
                results = st.session_state.azure_search_manager.search_hybrid("test", top=1)
                st.success("✅ Connexion fonctionnelle")
            except Exception as e:
                st.error(f"❌ Erreur : {str(e)}")
    else:
        st.warning("⚠️ Azure Search : Non configuré")
    
    # Papier en-tête
    if 'letterhead_template' in st.session_state and st.session_state.letterhead_template:
        st.success("✅ Papier en-tête : Configuré")
    else:
        st.warning("⚠️ Papier en-tête : Non configuré")

def show_export_import_tab():
    """Affiche l'onglet d'export/import"""
    st.markdown("### 💾 Export/Import de configuration")
    
    # Export
    show_export_section()
    
    # Import
    show_import_section()

def show_export_section():
    """Affiche la section d'export"""
    st.markdown("#### 📥 Export")
    
    export_data = {
        'pieces_selectionnees': {
            k: {
                'document_id': v.document_id,
                'titre': v.titre,
                'categorie': v.categorie,
                'notes': v.notes,
                'pertinence': v.pertinence
            }
            for k, v in st.session_state.pieces_selectionnees.items()
        },
        'learned_styles': st.session_state.get('learned_styles', {}),
        'letterhead_template': {
            'name': st.session_state.letterhead_template.name,
            'header_content': st.session_state.letterhead_template.header_content,
            'footer_content': st.session_state.letterhead_template.footer_content,
            'font_family': st.session_state.letterhead_template.font_family,
            'font_size': st.session_state.letterhead_template.font_size,
            'line_spacing': st.session_state.letterhead_template.line_spacing,
            'margins': st.session_state.letterhead_template.margins
        } if 'letterhead_template' in st.session_state and st.session_state.letterhead_template else None,
        'timestamp': datetime.now().isoformat()
    }
    
    export_json = json.dumps(export_data, indent=2, ensure_ascii=False)
    
    st.download_button(
        "💾 Exporter la configuration complète",
        export_json,
        f"config_penal_affaires_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        "application/json",
        key="export_config"
    )

def show_import_section():
    """Affiche la section d'import"""
    st.markdown("#### 📤 Import")
    
    uploaded_file = st.file_uploader(
        "Charger une configuration",
        type=['json'],
        key="import_config_file"
    )
    
    if uploaded_file:
        try:
            config_data = json.load(uploaded_file)
            
            # Prévisualisation
            with st.expander("Voir le contenu"):
                st.json(config_data)
            
            if st.button("⬆️ Importer", key="import_config_button"):
                import_configuration(config_data)
                st.success("✅ Configuration importée avec succès")
                st.rerun()
                
        except Exception as e:
            st.error(f"❌ Erreur lors de l'import: {str(e)}")

def import_configuration(config_data):
    """Importe la configuration depuis les données"""
    # Importer les pièces sélectionnées
    if 'pieces_selectionnees' in config_data:
        for piece_id, piece_data in config_data['pieces_selectionnees'].items():
            piece = PieceSelectionnee(
                document_id=piece_data['document_id'],
                titre=piece_data['titre'],
                categorie=piece_data['categorie'],
                notes=piece_data.get('notes', ''),
                pertinence=piece_data.get('pertinence', 5)
            )
            st.session_state.pieces_selectionnees[piece_id] = piece
    
    # Importer les styles appris
    if 'learned_styles' in config_data:
        st.session_state.learned_styles = config_data['learned_styles']
    
    # Importer le papier en-tête
    if 'letterhead_template' in config_data and config_data['letterhead_template']:
        lt = config_data['letterhead_template']
        st.session_state.letterhead_template = LetterheadTemplate(
            name=lt['name'],
            header_content=lt['header_content'],
            footer_content=lt['footer_content'],
            font_family=lt.get('font_family', 'Arial'),
            font_size=lt.get('font_size', 11),
            line_spacing=lt.get('line_spacing', 1.5),
            margins=lt.get('margins', {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5})
        )