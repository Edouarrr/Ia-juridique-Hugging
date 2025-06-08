# modules/redaction_assistee.py
"""Page de rédaction assistée par IA avec analyse de style"""

import streamlit as st
import io
import asyncio
from datetime import datetime

from config.app_config import LLMProvider
from models.dataclasses import StylePattern
from managers.multi_llm_manager import MultiLLMManager
from managers.style_analyzer import StyleAnalyzer
from managers.dynamic_generators import generate_dynamic_templates
from utils.helpers import clean_key

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def show_page():
    """Affiche la page de rédaction assistée"""
    st.header("📝 Rédaction assistée par IA")
    
    # Initialiser l'analyseur de style
    if 'style_analyzer' not in st.session_state:
        st.session_state.style_analyzer = StyleAnalyzer()
    
    # Onglets
    tabs = st.tabs(["✍️ Rédaction", "📚 Apprentissage de style", "📚 Modèles"])
    
    with tabs[0]:
        show_redaction_tab()
    
    with tabs[1]:
        show_style_learning_tab()
    
    with tabs[2]:
        show_templates_tab()

def show_redaction_tab():
    """Onglet de rédaction"""
    st.markdown("### ✍️ Créer un nouvel acte")
    
    # Type d'acte
    col1, col2 = st.columns(2)
    
    with col1:
        type_acte = st.text_input(
            "Type d'acte",
            placeholder="Ex: Conclusions en défense, Plainte, Assignation...",
            key="type_acte_redaction"
        )
        
        client_nom = st.text_input(
            "Client",
            placeholder="Nom du client",
            key="client_nom_acte"
        )
        
        partie_adverse = st.text_input(
            "Partie adverse",
            placeholder="Nom de la partie adverse",
            key="partie_adverse_acte"
        )
    
    with col2:
        juridiction = st.text_input(
            "Juridiction",
            placeholder="Ex: Tribunal judiciaire de Paris",
            key="juridiction_acte"
        )
        
        numero_affaire = st.text_input(
            "N° RG / Référence",
            placeholder="Ex: 23/00123",
            key="numero_affaire_acte"
        )
        
        infraction = st.text_input(
            "Infraction/Objet",
            placeholder="Ex: Abus de biens sociaux",
            key="infraction_acte"
        )
    
    # Éléments à inclure
    st.markdown("### 📋 Éléments à développer")
    
    # Faits
    faits = st.text_area(
        "Résumé des faits",
        placeholder="Décrivez brièvement les faits principaux...",
        height=150,
        key="faits_acte"
    )
    
    # Arguments
    arguments = st.text_area(
        "Points clés à développer",
        placeholder="""Ex:
- Absence d'élément intentionnel
- Prescription acquise
- Nullité de la procédure
- Bonne foi du client""",
        height=150,
        key="arguments_acte"
    )
    
    # Pièces à citer
    if st.session_state.pieces_selectionnees:
        st.markdown("### 📎 Pièces sélectionnées à citer")
        pieces_a_citer = []
        
        for piece_id, piece in st.session_state.pieces_selectionnees.items():
            if st.checkbox(
                f"{piece.titre} ({piece.categorie})",
                value=True,
                key=f"cite_piece_{piece_id}"
            ):
                pieces_a_citer.append(piece)
        
        st.session_state.pieces_a_citer = pieces_a_citer
    
    # Style à appliquer
    st.markdown("### 🎨 Style de rédaction")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Styles appris disponibles
        styles_disponibles = ["Style standard"] + list(st.session_state.get('learned_styles', {}).keys())
        
        style_choisi = st.selectbox(
            "Style à appliquer",
            styles_disponibles,
            key="style_choisi_acte"
        )
        
        ton_redaction = st.select_slider(
            "Ton",
            options=["Très formel", "Formel", "Équilibré", "Direct", "Combatif"],
            value="Formel",
            key="ton_redaction_acte"
        )
    
    with col2:
        longueur_cible = st.select_slider(
            "Longueur",
            options=["Concis", "Standard", "Détaillé", "Très détaillé"],
            value="Standard",
            key="longueur_acte"
        )
        
        # Modèle à utiliser
        if st.checkbox("Utiliser un modèle", key="use_template_check"):
            modeles_disponibles = list(st.session_state.get('custom_templates', {}).keys())
            if 'template_to_use' in st.session_state:
                modeles_disponibles.append("Modèle importé")
            
            if modeles_disponibles:
                modele_choisi = st.selectbox(
                    "Modèle",
                    modeles_disponibles,
                    key="modele_choisi"
                )
            else:
                st.info("Aucun modèle disponible. Créez-en dans l'onglet Modèles.")
    
    # Documents de référence à analyser
    if st.session_state.azure_documents:
        with st.expander("📄 Documents de référence (optionnel)"):
            st.info("Sélectionnez des documents similaires pour guider le style")
            docs_reference = []
            
            for doc_id, doc in list(st.session_state.azure_documents.items())[:10]:
                if st.checkbox(
                    doc.title[:80],
                    key=f"ref_doc_{doc_id}"
                ):
                    docs_reference.append(doc_id)
            
            st.session_state.docs_reference_acte = docs_reference
    
    # Bouton de génération
    if st.button("🚀 Générer l'acte", type="primary", key="generer_acte"):
        generate_acte(type_acte, style_choisi)

def show_style_learning_tab():
    """Onglet d'apprentissage de style"""
    st.markdown("### 📚 Apprentissage de style")
    
    st.info("""
    Cette fonctionnalité analyse vos documents pour apprendre votre style de rédaction
    et l'appliquer automatiquement aux nouveaux actes.
    """)
    
    # Sélection des documents modèles
    st.markdown("#### 📄 Sélectionner les documents modèles")
    
    # Upload de documents
    uploaded_files = st.file_uploader(
        "Charger des documents modèles",
        type=['docx', 'txt', 'pdf'],
        accept_multiple_files=True,
        key="upload_style_docs"
    )
    
    # Ou sélectionner parmi les documents Azure
    if st.session_state.azure_documents:
        st.markdown("##### Ou sélectionner parmi vos documents")
        
        docs_for_style = []
        for doc_id, doc in list(st.session_state.azure_documents.items())[:20]:
            if st.checkbox(
                f"{doc.title[:80]}",
                key=f"style_doc_{doc_id}"
            ):
                docs_for_style.append(doc_id)
    
    # Type d'acte pour ce style
    style_name = st.text_input(
        "Nom du style",
        placeholder="Ex: Conclusions pénales, Plaintes commerciales...",
        key="style_name_input"
    )
    
    # Analyser le style
    if st.button("🔍 Analyser le style", key="analyze_style_button"):
        if style_name and (uploaded_files or docs_for_style):
            analyze_style(style_name, uploaded_files, docs_for_style)
    
    # Afficher les styles appris
    if st.session_state.get('learned_styles'):
        st.markdown("### 🎨 Styles appris")
        
        for style_name, style_data in st.session_state.learned_styles.items():
            with st.expander(f"📝 {style_name}"):
                st.json(style_data.get('summary', {}))
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Documents analysés", style_data.get('doc_count', 0))
                    st.metric("Longueur moyenne", f"{style_data.get('avg_length', 0):,.0f} mots")
                
                with col2:
                    if st.button(f"🗑️ Supprimer", key=f"delete_style_{clean_key(style_name)}"):
                        del st.session_state.learned_styles[style_name]
                        st.rerun()

def show_templates_tab():
    """Onglet des modèles"""
    st.markdown("### 📚 Bibliothèque de modèles")
    
    # Options pour générer des modèles dynamiques
    st.markdown("#### 🤖 Générer des modèles personnalisés")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        type_modele_generer = st.text_input(
            "Type d'acte pour lequel générer des modèles",
            placeholder="Ex: Plainte avec constitution de partie civile, Mémoire en défense...",
            key="type_modele_generer"
        )
    
    with col2:
        if st.button("🎯 Générer", key="generer_modeles_button"):
            if type_modele_generer:
                with st.spinner("Génération de modèles intelligents..."):
                    # Contexte optionnel
                    context = {
                        'client': st.session_state.get('client_nom_acte', ''),
                        'infraction': st.session_state.get('infraction_acte', ''),
                        'juridiction': st.session_state.get('juridiction_acte', '')
                    }
                    
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    
                    modeles = loop.run_until_complete(
                        generate_dynamic_templates(type_modele_generer, context)
                    )
                    
                    # Stocker dans le cache
                    cache_key = f"templates_{clean_key(type_modele_generer)}"
                    st.session_state.dynamic_templates[cache_key] = modeles
                    
                    st.success("✅ Modèles générés avec succès!")
    
    # Afficher les modèles générés dynamiquement
    if st.session_state.dynamic_templates:
        st.markdown("#### 🎨 Modèles générés par IA")
        
        for cache_key, modeles in st.session_state.dynamic_templates.items():
            type_clean = cache_key.replace("templates_", "").replace("_", " ").title()
            
            with st.expander(f"📁 Modèles pour : {type_clean}", expanded=True):
                for titre, contenu in modeles.items():
                    st.markdown(f"**{titre}**")
                    
                    # Aperçu du modèle
                    with st.container():
                        st.text_area(
                            "Aperçu",
                            value=contenu[:500] + "..." if len(contenu) > 500 else contenu,
                            height=200,
                            disabled=True,
                            key=f"preview_dyn_template_{clean_key(titre)}"
                        )
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if st.button(f"📋 Utiliser", key=f"use_dyn_template_{clean_key(titre)}"):
                            st.session_state.template_to_use = contenu
                            st.session_state.selected_tab = 0  # Retour à l'onglet rédaction
                            st.success("✅ Modèle copié. Retournez à l'onglet Rédaction.")
                    
                    with col2:
                        if st.button(f"👁️ Voir complet", key=f"view_dyn_template_{clean_key(titre)}"):
                            with st.expander("Modèle complet", expanded=True):
                                st.text_area(
                                    "Contenu",
                                    value=contenu,
                                    height=600,
                                    key=f"full_dyn_template_{clean_key(titre)}"
                                )
                    
                    with col3:
                        st.download_button(
                            "💾 Télécharger",
                            contenu,
                            f"{clean_key(titre)}.txt",
                            "text/plain",
                            key=f"download_dyn_template_{clean_key(titre)}"
                        )
                    
                    st.markdown("---")
    
    # Modèles personnalisés existants
    if st.session_state.get('custom_templates'):
        st.markdown("#### 📑 Modèles personnalisés")
        
        for template_name, template_content in st.session_state.custom_templates.items():
            with st.expander(f"📄 {template_name}"):
                st.text_area(
                    "Contenu",
                    value=template_content,
                    height=300,
                    key=f"view_template_{clean_key(template_name)}"
                )
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button(f"🗑️ Supprimer", key=f"delete_template_{clean_key(template_name)}"):
                        del st.session_state.custom_templates[template_name]
                        st.rerun()
                
                with col2:
                    st.download_button(
                        "💾 Télécharger",
                        template_content,
                        f"{clean_key(template_name)}.txt",
                        "text/plain",
                        key=f"download_template_{clean_key(template_name)}"
                    )
    
    # Créer un nouveau modèle
    st.markdown("#### ➕ Créer un nouveau modèle")
    
    new_template_name = st.text_input(
        "Nom du modèle",
        placeholder="Ex: Conclusions récapitulatives",
        key="new_template_name"
    )
    
    new_template_content = st.text_area(
        "Contenu du modèle",
        placeholder="Utilisez des balises comme [CLIENT], [PARTIE_ADVERSE], [FAITS], etc.",
        height=300,
        key="new_template_content"
    )
    
    if st.button("💾 Sauvegarder le modèle", key="save_new_template"):
        if new_template_name and new_template_content:
            if 'custom_templates' not in st.session_state:
                st.session_state.custom_templates = {}
            
            st.session_state.custom_templates[new_template_name] = new_template_content
            st.success(f"✅ Modèle '{new_template_name}' sauvegardé!")
            st.rerun()

def analyze_style(style_name: str, uploaded_files, doc_ids):
    """Analyse le style des documents sélectionnés"""
    with st.spinner("🔄 Analyse du style en cours..."):
        analyzer = st.session_state.style_analyzer
        documents = []
        
        # Récupérer le contenu des documents
        if uploaded_files:
            for file in uploaded_files:
                content = file.read()
                if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # Traiter les fichiers Word
                    if DOCX_AVAILABLE:
                        doc = DocxDocument(io.BytesIO(content))
                        text = "\n".join([p.text for p in doc.paragraphs])
                        documents.append({
                            'title': file.name,
                            'content': text
                        })
                else:
                    # Fichiers texte
                    documents.append({
                        'title': file.name,
                        'content': content.decode('utf-8', errors='ignore')
                    })
        
        # Documents Azure
        for doc_id in doc_ids:
            if doc_id in st.session_state.azure_documents:
                doc = st.session_state.azure_documents[doc_id]
                documents.append({
                    'title': doc.title,
                    'content': doc.content
                })
        
        if documents:
            # Analyser le style
            style_pattern = analyzer.analyze_documents(documents, style_name)
            
            # Sauvegarder le style
            if 'learned_styles' not in st.session_state:
                st.session_state.learned_styles = {}
            
            st.session_state.learned_styles[style_name] = {
                'pattern': style_pattern.__dict__,
                'doc_count': len(documents),
                'avg_length': sum(len(d['content'].split()) for d in documents) // len(documents),
                'summary': {
                    'structure': style_pattern.structure,
                    'formules_types': style_pattern.formules[:5],
                    'vocabulaire_freq': dict(list(style_pattern.vocabulaire.items())[:10])
                }
            }
            
            st.success(f"✅ Style '{style_name}' appris avec succès!")
            st.info(f"📊 {len(documents)} documents analysés")

def generate_acte(type_acte: str, style_choisi: str):
    """Génère l'acte avec le style choisi"""
    if not type_acte:
        st.error("❌ Veuillez spécifier le type d'acte")
        return
    
    # Construire le prompt
    prompt = build_generation_prompt(type_acte, style_choisi)
    
    # Générer avec l'IA
    llm_manager = MultiLLMManager()
    
    with st.spinner("🔄 Génération en cours..."):
        if llm_manager.clients:
            # Utiliser le premier provider disponible
            provider = list(llm_manager.clients.keys())[0]
            
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            response = loop.run_until_complete(
                llm_manager.query_single_llm(
                    provider,
                    prompt,
                    "Tu es un avocat expert en droit pénal des affaires, spécialisé dans la rédaction d'actes juridiques."
                )
            )
            
            if response['success']:
                # Afficher le résultat
                st.markdown("### 📄 Acte généré")
                
                # Zone d'édition
                acte_genere = st.text_area(
                    "Vous pouvez modifier l'acte généré",
                    value=response['response'],
                    height=600,
                    key="acte_genere_edit"
                )
                
                # Options d'export
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.download_button(
                        "💾 Télécharger (.txt)",
                        acte_genere,
                        f"{clean_key(type_acte)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        "text/plain",
                        key="download_txt_acte"
                    )
                
                with col2:
                    if DOCX_AVAILABLE:
                        # Créer un document Word
                        doc = DocxDocument()
                        doc.add_heading(type_acte, 0)
                        
                        # Ajouter le contenu
                        for paragraph in acte_genere.split('\n'):
                            if paragraph.strip():
                                doc.add_paragraph(paragraph)
                        
                        # Sauvegarder en mémoire
                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.download_button(
                            "💾 Télécharger (.docx)",
                            docx_buffer.getvalue(),
                            f"{clean_key(type_acte)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx_acte"
                        )
                
                with col3:
                    if st.button("📧 Préparer l'envoi", key="prepare_send"):
                        st.info("Fonctionnalité d'envoi à implémenter")
            else:
                st.error(f"❌ Erreur : {response['error']}")
        else:
            st.error("❌ Aucune IA disponible")

def build_generation_prompt(type_acte: str, style_choisi: str) -> str:
    """Construit le prompt pour la génération"""
    prompt = f"Rédige un(e) {type_acte} professionnel(le) en droit pénal des affaires.\n\n"
    
    # Ajouter les informations du formulaire
    if st.session_state.get('client_nom_acte'):
        prompt += f"Client : {st.session_state.client_nom_acte}\n"
    
    if st.session_state.get('partie_adverse_acte'):
        prompt += f"Partie adverse : {st.session_state.partie_adverse_acte}\n"
    
    if st.session_state.get('juridiction_acte'):
        prompt += f"Juridiction : {st.session_state.juridiction_acte}\n"
    
    if st.session_state.get('numero_affaire_acte'):
        prompt += f"Référence : {st.session_state.numero_affaire_acte}\n"
    
    if st.session_state.get('infraction_acte'):
        prompt += f"Infraction/Objet : {st.session_state.infraction_acte}\n"
    
    # Ajouter les faits
    if st.session_state.get('faits_acte'):
        prompt += f"\nFaits :\n{st.session_state.faits_acte}\n"
    
    # Ajouter les arguments
    if st.session_state.get('arguments_acte'):
        prompt += f"\nPoints clés à développer :\n{st.session_state.arguments_acte}\n"
    
    # Ajouter les références aux pièces
    if st.session_state.get('pieces_a_citer'):
        prompt += "\nPièces à citer :\n"
        for i, piece in enumerate(st.session_state.pieces_a_citer, 1):
            prompt += f"- Pièce n°{i} : {piece.titre}\n"
    
    # Ajouter les instructions de style
    if style_choisi != "Style standard" and style_choisi in st.session_state.get('learned_styles', {}):
        style_data = st.session_state.learned_styles[style_choisi]
        prompt += f"\nAppliquer le style '{style_choisi}' avec les caractéristiques suivantes :\n"
        prompt += f"- Structure type : {style_data['pattern']['structure']}\n"
        prompt += f"- Formules caractéristiques : {style_data['pattern']['formules'][:3]}\n"
    
    # Ajouter les paramètres de rédaction
    prompt += f"\nTon : {st.session_state.get('ton_redaction_acte', 'Formel')}\n"
    prompt += f"Longueur : {st.session_state.get('longueur_acte', 'Standard')}\n"
    
    # Ajouter le modèle si sélectionné
    if st.session_state.get('use_template_check') and st.session_state.get('template_to_use'):
        prompt += f"\nUtiliser le modèle suivant comme base :\n{st.session_state.template_to_use}\n"
    
    return prompt