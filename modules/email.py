# modules/email.py
"""Module de gestion des emails pour l'application juridique"""

import streamlit as st
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from typing import Dict, List, Any, Optional
import re
import json
import io
from dataclasses import dataclass, field

# Dataclasses intégrées
@dataclass
class EmailConfig:
    """Configuration d'un email"""
    to: List[str]
    subject: str
    body: str
    cc: List[str] = field(default_factory=list)
    bcc: List[str] = field(default_factory=list)
    attachments: List[Dict[str, Any]] = field(default_factory=list)
    priority: str = "normal"
    
    def add_attachment(self, filename: str, data: bytes, mimetype: str):
        """Ajoute une pièce jointe"""
        self.attachments.append({
            'filename': filename,
            'data': data,
            'mimetype': mimetype
        })

# Fonctions helper intégrées
def is_valid_email(email: str) -> bool:
    """Vérifie si une adresse email est valide"""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def format_file_size(size_bytes: int) -> str:
    """Formate la taille d'un fichier"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} TB"

# Types MIME pour les pièces jointes
ATTACHMENT_MIME_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'txt': 'text/plain',
    'html': 'text/html',
    'json': 'application/json',
    'csv': 'text/csv',
    'zip': 'application/zip'
}

def process_email_request(query: str, analysis: dict):
    """Traite une demande d'envoi par email"""
    
    st.markdown("### 📧 Envoi par email")
    
    # Initialiser la session state si nécessaire
    if 'email_config' not in st.session_state:
        st.session_state.email_config = None
    
    # Déterminer les destinataires
    recipients = extract_email_recipients(query, analysis)
    
    # Configuration de l'email
    if st.session_state.email_config is None:
        email_config = create_email_config(recipients, analysis)
        st.session_state.email_config = email_config
    else:
        email_config = st.session_state.email_config
    
    # Interface de configuration
    show_email_configuration(email_config)
    
    # Préparer le contenu et les pièces jointes
    prepare_email_content_and_attachments(email_config, analysis)
    
    # Aperçu
    show_email_preview(email_config)
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📧 Envoyer", type="primary", key="send_email_main", use_container_width=True):
            send_email_with_progress(email_config)
    
    with col2:
        if st.button("💾 Sauvegarder brouillon", key="save_draft", use_container_width=True):
            save_email_draft(email_config)
    
    with col3:
        if st.button("🔄 Réinitialiser", key="reset_email", use_container_width=True):
            st.session_state.email_config = None
            st.rerun()

def extract_email_recipients(query: str, analysis: dict) -> List[str]:
    """Extrait les destinataires depuis la requête"""
    
    recipients = []
    
    # Extraction des emails depuis la requête
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    found_emails = re.findall(email_pattern, query)
    recipients.extend(found_emails)
    
    # Extraction depuis l'analyse
    if analysis.get('entities', {}).get('emails'):
        recipients.extend(analysis['entities']['emails'])
    
    # Récupérer depuis l'historique si pas de destinataires
    if not recipients and st.session_state.get('last_email_recipients'):
        recipients = st.session_state.last_email_recipients
    
    # Récupérer depuis les contacts favoris
    if not recipients and st.session_state.get('favorite_contacts'):
        # Proposer de sélectionner parmi les favoris
        pass
    
    return list(set(recipients))  # Dédupliquer

def create_email_config(recipients: List[str], analysis: dict) -> EmailConfig:
    """Crée la configuration de l'email"""
    
    # Déterminer l'objet selon le contexte
    subject = determine_email_subject(analysis)
    
    # Corps initial avec signature
    body = generate_email_body(analysis)
    
    # Ajouter la signature par défaut
    if st.session_state.get('default_signature'):
        signatures = st.session_state.get('email_signatures', get_default_signatures())
        if st.session_state.default_signature in signatures:
            body += "\n\n" + signatures[st.session_state.default_signature]['content']
    
    config = EmailConfig(
        to=recipients,
        subject=subject,
        body=body,
        cc=[],
        bcc=[],
        attachments=[],
        priority="normal"
    )
    
    return config

def determine_email_subject(analysis: dict) -> str:
    """Détermine l'objet de l'email selon le contexte"""
    
    if st.session_state.get('redaction_result'):
        doc_type = st.session_state.redaction_result.get('type', 'document')
        reference = analysis.get('reference', '')
        
        subjects = {
            'conclusions': f"Conclusions en défense - {reference}",
            'plainte': f"Dépôt de plainte - {reference}",
            'constitution_pc': f"Constitution de partie civile - {reference}",
            'assignation': f"Assignation - {reference}",
            'mémoire': f"Mémoire - {reference}",
            'courrier': f"Courrier juridique - {reference}"
        }
        
        return subjects.get(doc_type, f"Document juridique - {reference}")
    
    elif st.session_state.get('current_bordereau'):
        return f"Bordereau de communication de pièces - {analysis.get('reference', 'Dossier')}"
    
    elif st.session_state.get('synthesis_result'):
        return f"Synthèse - {analysis.get('reference', 'Dossier')}"
    
    else:
        return f"Document - {datetime.now().strftime('%d/%m/%Y')}"

def generate_email_body(analysis: dict) -> str:
    """Génère le corps de l'email selon le contexte"""
    
    reference = analysis.get('reference', 'votre dossier')
    
    # Vérifier si un template est sélectionné
    if st.session_state.get('selected_email_template'):
        templates = st.session_state.get('email_templates', get_default_email_templates())
        template = templates.get(st.session_state.selected_email_template)
        if template:
            body = template['body']
            # Remplacer les variables
            body = body.replace('[référence]', reference)
            body = body.replace('[Référence]', reference)
            return body
    
    # Sinon, générer selon le contexte
    if st.session_state.get('redaction_result'):
        result = st.session_state.redaction_result
        doc_type_names = {
            'conclusions': 'les conclusions en défense',
            'plainte': 'la plainte',
            'constitution_pc': 'la constitution de partie civile',
            'assignation': "l'assignation",
            'mémoire': 'le mémoire',
            'courrier': 'le courrier'
        }
        
        doc_name = doc_type_names.get(result['type'], 'le document')
        
        body = f"""Maître,

Je vous prie de bien vouloir trouver ci-joint {doc_name} concernant {reference}.

Ce document a été préparé conformément à vos instructions et intègre l'ensemble des éléments discutés.

Je reste à votre disposition pour tout complément d'information.

Bien cordialement,"""

    elif st.session_state.get('current_bordereau'):
        bordereau = st.session_state.current_bordereau
        piece_count = bordereau['metadata']['piece_count']
        
        body = f"""Maître,

Veuillez trouver ci-joint le bordereau de communication de pièces pour {reference}.

Ce bordereau récapitule {piece_count} pièces communiquées, organisées par catégorie.

Les pièces physiques vous seront transmises selon les modalités convenues.

Cordialement,"""

    else:
        body = f"""Bonjour,

Veuillez trouver ci-joint le document demandé concernant {reference}.

N'hésitez pas à me contacter pour toute question.

Cordialement,"""
    
    return body

def show_email_configuration(email_config: EmailConfig):
    """Interface de configuration de l'email"""
    
    st.markdown("#### 📮 Configuration de l'email")
    
    # Quick actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Contacts favoris
        if st.session_state.get('favorite_contacts'):
            selected_contact = st.selectbox(
                "📇 Contacts favoris",
                [""] + list(st.session_state.favorite_contacts.keys()),
                key="quick_contact_select"
            )
            
            if selected_contact:
                contact_email = st.session_state.favorite_contacts[selected_contact]
                if contact_email not in email_config.to:
                    email_config.to.append(contact_email)
    
    with col2:
        # Templates
        templates = st.session_state.get('email_templates', get_default_email_templates())
        selected_template = st.selectbox(
            "📝 Template",
            [""] + list(templates.keys()),
            format_func=lambda x: templates[x]['name'] if x else "Choisir un template",
            key="template_select"
        )
        
        if selected_template:
            st.session_state.selected_email_template = selected_template
            template = templates[selected_template]
            email_config.subject = template['subject'].replace('[Référence]', 
                                                               st.session_state.get('current_reference', ''))
            email_config.body = generate_email_body({})
    
    with col3:
        # Groupes
        if st.button("👥 Groupes de diffusion", key="show_groups"):
            st.session_state.show_email_groups = not st.session_state.get('show_email_groups', False)
    
    # Gestion des groupes
    if st.session_state.get('show_email_groups'):
        show_email_groups_management(email_config)
    
    st.divider()
    
    # Destinataires
    col1, col2, col3 = st.columns(3)
    
    with col1:
        to_addresses = st.text_area(
            "À (To) *",
            value=", ".join(email_config.to),
            help="Séparez les adresses par des virgules",
            key="email_to_field",
            height=70
        )
        email_config.to = [email.strip() for email in to_addresses.split(',') if email.strip()]
    
    with col2:
        cc_addresses = st.text_area(
            "Cc",
            value=", ".join(email_config.cc),
            help="Copie carbone",
            key="email_cc_field",
            height=70
        )
        email_config.cc = [email.strip() for email in cc_addresses.split(',') if email.strip()]
    
    with col3:
        bcc_addresses = st.text_area(
            "Cci",
            value=", ".join(email_config.bcc),
            help="Copie carbone invisible",
            key="email_bcc_field",
            height=70
        )
        email_config.bcc = [email.strip() for email in bcc_addresses.split(',') if email.strip()]
    
    # Validation des emails
    all_emails = email_config.to + email_config.cc + email_config.bcc
    invalid_emails = [email for email in all_emails if email and not is_valid_email(email)]
    
    if invalid_emails:
        st.error(f"❌ Emails invalides : {', '.join(invalid_emails)}")
    
    # Objet
    email_config.subject = st.text_input(
        "Objet *",
        value=email_config.subject,
        key="email_subject_field"
    )
    
    # Options avancées
    with st.expander("⚙️ Options avancées"):
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            email_config.priority = st.selectbox(
                "Priorité",
                ["low", "normal", "high"],
                index=1,
                format_func=lambda x: {"low": "🔵 Faible", "normal": "⚪ Normale", "high": "🔴 Haute"}[x],
                key="email_priority"
            )
        
        with col2:
            request_receipt = st.checkbox(
                "Accusé de réception",
                key="email_receipt"
            )
        
        with col3:
            request_read = st.checkbox(
                "Confirmation de lecture",
                key="email_read_confirm"
            )
        
        with col4:
            encrypt_email = st.checkbox(
                "Chiffrer l'email",
                key="email_encrypt",
                help="Nécessite la configuration du chiffrement"
            )

def show_email_groups_management(email_config: EmailConfig):
    """Gestion des groupes de diffusion"""
    
    with st.container():
        st.markdown("##### 👥 Groupes de diffusion")
        
        # Initialiser les groupes s'ils n'existent pas
        if 'email_groups' not in st.session_state:
            st.session_state.email_groups = {
                'tribunal': {
                    'name': 'Tribunal',
                    'emails': ['greffe@tribunal.fr', 'president@tribunal.fr']
                },
                'adversaires': {
                    'name': 'Parties adverses',
                    'emails': []
                }
            }
        
        groups = st.session_state.email_groups
        
        # Sélection de groupe
        col1, col2 = st.columns([3, 1])
        
        with col1:
            selected_group = st.selectbox(
                "Ajouter un groupe",
                [""] + list(groups.keys()),
                format_func=lambda x: groups[x]['name'] if x else "Sélectionner un groupe",
                key="select_email_group"
            )
        
        with col2:
            if st.button("➕ Ajouter", key="add_group_to_email") and selected_group:
                group_emails = groups[selected_group]['emails']
                for email in group_emails:
                    if email not in email_config.to:
                        email_config.to.append(email)
                st.rerun()
        
        # Gestion des groupes
        if st.checkbox("Gérer les groupes", key="manage_groups_check"):
            manage_email_groups()

def manage_email_groups():
    """Interface de gestion des groupes"""
    
    groups = st.session_state.email_groups
    
    # Créer un nouveau groupe
    with st.form("new_group_form"):
        st.markdown("**Nouveau groupe**")
        col1, col2 = st.columns(2)
        
        with col1:
            new_group_key = st.text_input("Identifiant", placeholder="ex: confreres")
        
        with col2:
            new_group_name = st.text_input("Nom", placeholder="ex: Confrères")
        
        new_group_emails = st.text_area(
            "Emails (un par ligne)",
            placeholder="email1@example.com\nemail2@example.com"
        )
        
        if st.form_submit_button("Créer le groupe"):
            if new_group_key and new_group_name:
                emails = [e.strip() for e in new_group_emails.split('\n') if e.strip()]
                valid_emails = [e for e in emails if is_valid_email(e)]
                
                groups[new_group_key] = {
                    'name': new_group_name,
                    'emails': valid_emails
                }
                st.success(f"✅ Groupe '{new_group_name}' créé avec {len(valid_emails)} email(s)")
                st.rerun()
    
    # Modifier les groupes existants
    st.markdown("**Groupes existants**")
    
    for group_key, group_data in groups.items():
        with st.expander(f"{group_data['name']} ({len(group_data['emails'])} emails)"):
            # Éditer les emails
            updated_emails = st.text_area(
                "Emails",
                value="\n".join(group_data['emails']),
                key=f"edit_group_{group_key}"
            )
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button(f"💾 Sauvegarder", key=f"save_group_{group_key}"):
                    emails = [e.strip() for e in updated_emails.split('\n') if e.strip()]
                    valid_emails = [e for e in emails if is_valid_email(e)]
                    group_data['emails'] = valid_emails
                    st.success(f"✅ Groupe mis à jour ({len(valid_emails)} emails)")
                    st.rerun()
            
            with col2:
                if st.button(f"🗑️ Supprimer", key=f"delete_group_{group_key}"):
                    del groups[group_key]
                    st.rerun()

def prepare_email_content_and_attachments(email_config: EmailConfig, analysis: dict):
    """Prépare le contenu et les pièces jointes"""
    
    st.markdown("#### 📎 Pièces jointes")
    
    # Déterminer les pièces jointes potentielles
    available_attachments = get_available_attachments()
    
    if available_attachments:
        # Organiser par catégorie
        categories = {}
        for name, data in available_attachments.items():
            category = data.get('category', 'Autres')
            if category not in categories:
                categories[category] = []
            categories[category].append((name, data))
        
        # Afficher par catégorie
        selected_attachments = []
        
        for category, items in categories.items():
            st.markdown(f"**{category}**")
            
            for name, data in items:
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    if st.checkbox(name, key=f"attach_{name}", value=True):
                        selected_attachments.append(name)
                
                with col2:
                    st.caption(data.get('type', 'document'))
                
                with col3:
                    if data.get('size'):
                        st.caption(format_file_size(data['size']))
        
        # Format des pièces jointes
        if selected_attachments:
            st.divider()
            
            col1, col2 = st.columns(2)
            
            with col1:
                attachment_format = st.selectbox(
                    "Format des documents",
                    ["pdf", "docx", "txt", "html"],
                    key="attachment_format"
                )
            
            with col2:
                compress_attachments = st.checkbox(
                    "Compresser en ZIP",
                    key="compress_attachments",
                    help="Recommandé pour plus de 5 fichiers"
                )
            
            # Préparer les pièces jointes
            email_config.attachments = []  # Réinitialiser
            
            for attachment_name in selected_attachments:
                attachment_data = prepare_attachment(
                    attachment_name,
                    available_attachments[attachment_name],
                    attachment_format
                )
                
                if attachment_data:
                    email_config.add_attachment(
                        f"{attachment_name}.{attachment_format}",
                        attachment_data,
                        ATTACHMENT_MIME_TYPES.get(attachment_format, "application/octet-stream")
                    )
            
            # Compression si demandée
            if compress_attachments and len(email_config.attachments) > 1:
                compress_email_attachments(email_config)
        
        # Afficher la liste des pièces jointes
        if email_config.attachments:
            st.info(f"📎 {len(email_config.attachments)} pièce(s) jointe(s)")
            
            total_size = sum(len(att['data']) for att in email_config.attachments)
            
            if total_size > 25 * 1024 * 1024:  # 25 MB
                st.warning("⚠️ Taille totale importante. Certains serveurs peuvent rejeter l'email.")
            
            # Liste avec actions
            for i, att in enumerate(email_config.attachments):
                col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                
                with col1:
                    st.write(f"• {att['filename']}")
                
                with col2:
                    st.caption(format_file_size(len(att['data'])))
                
                with col3:
                    # Aperçu pour certains types
                    if att['filename'].endswith('.txt'):
                        if st.button("👁️", key=f"preview_att_{i}"):
                            st.text_area(
                                f"Aperçu de {att['filename']}", 
                                value=att['data'].decode('utf-8', errors='ignore')[:500] + "...",
                                height=200,
                                key=f"preview_content_{i}"
                            )
                
                with col4:
                    if st.button("❌", key=f"remove_att_{i}"):
                        email_config.attachments.pop(i)
                        st.rerun()
    
    # Upload manuel
    st.divider()
    st.markdown("**Ajouter des fichiers**")
    
    uploaded_files = st.file_uploader(
        "Glisser-déposer ou parcourir",
        accept_multiple_files=True,
        key="email_manual_attachments",
        type=['pdf', 'docx', 'xlsx', 'txt', 'csv', 'json', 'zip']
    )
    
    if uploaded_files:
        for file in uploaded_files:
            # Vérifier si le fichier n'est pas déjà ajouté
            existing_names = [att['filename'] for att in email_config.attachments]
            if file.name not in existing_names:
                email_config.add_attachment(
                    file.name,
                    file.read(),
                    file.type or "application/octet-stream"
                )
                st.success(f"✅ {file.name} ajouté")

def compress_email_attachments(email_config: EmailConfig):
    """Compresse les pièces jointes en un fichier ZIP"""
    
    try:
        import zipfile
        
        # Créer un buffer pour le ZIP
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for att in email_config.attachments:
                zip_file.writestr(att['filename'], att['data'])
        
        # Remplacer les pièces jointes par le ZIP
        zip_data = zip_buffer.getvalue()
        zip_filename = f"documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        email_config.attachments = [{
            'filename': zip_filename,
            'data': zip_data,
            'mimetype': 'application/zip'
        }]
        
        st.success(f"✅ Fichiers compressés dans {zip_filename}")
        
    except ImportError:
        st.warning("⚠️ Module zipfile non disponible")

def get_available_attachments() -> Dict[str, Any]:
    """Récupère les documents disponibles pour attachement"""
    
    attachments = {}
    
    # Document rédigé
    if st.session_state.get('redaction_result'):
        result = st.session_state.redaction_result
        attachments[f"{result['type']}_principal"] = {
            'type': 'redaction',
            'content': result['document'],
            'category': '📄 Documents rédigés',
            'size': len(result['document'].encode('utf-8'))
        }
    
    # Bordereau
    if st.session_state.get('current_bordereau'):
        bordereau = st.session_state.current_bordereau
        attachments['bordereau_pieces'] = {
            'type': 'bordereau',
            'content': bordereau,
            'category': '📋 Bordereaux',
            'size': len(str(bordereau).encode('utf-8'))
        }
    
    # Synthèse
    if st.session_state.get('synthesis_result'):
        content = st.session_state.synthesis_result.get('content', '')
        attachments['synthese'] = {
            'type': 'synthesis',
            'content': content,
            'category': '📊 Analyses',
            'size': len(content.encode('utf-8'))
        }
    
    # Analyse
    if st.session_state.get('ai_analysis_results'):
        content = st.session_state.ai_analysis_results.get('content', '')
        attachments['analyse'] = {
            'type': 'analysis',
            'content': content,
            'category': '📊 Analyses',
            'size': len(content.encode('utf-8'))
        }
    
    # Documents uploadés
    if st.session_state.get('uploaded_documents'):
        for i, doc in enumerate(st.session_state.uploaded_documents):
            attachments[f"document_upload_{i}"] = {
                'type': 'uploaded',
                'content': doc.get('content', ''),
                'category': '📤 Documents uploadés',
                'size': doc.get('size', 0)
            }
    
    return attachments

def prepare_attachment(name: str, data: Dict[str, Any], format: str) -> bytes:
    """Prépare une pièce jointe dans le format demandé"""
    
    content = data.get('content', '')
    
    if format == 'pdf':
        return create_pdf_attachment(content, data['type'])
    
    elif format == 'docx':
        return create_docx_attachment(content, data['type'])
    
    elif format == 'html':
        return create_html_attachment(content, data['type'])
    
    else:  # txt
        return content.encode('utf-8') if isinstance(content, str) else str(content).encode('utf-8')

def create_pdf_attachment(content: Any, doc_type: str) -> bytes:
    """Crée une pièce jointe PDF"""
    try:
        # Essayer d'utiliser reportlab s'il est disponible
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        styles = getSampleStyleSheet()
        
        # Titre
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='black',
            spaceAfter=30
        )
        
        story.append(Paragraph(doc_type.replace('_', ' ').title(), title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Contenu
        if isinstance(content, str):
            # Diviser le contenu en paragraphes
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.replace('\n', '<br/>'), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
        
        elif isinstance(content, dict):
            # Traiter les structures complexes
            for key, value in content.items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # Construire le PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError:
        # Fallback : créer un PDF simple en texte
        pdf_content = f"""%%PDF-1.4
1 0 obj

/Type /Catalog
/Pages 2 0 R
>>
endobj

2 0 obj

/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj

3 0 obj

/Type /Page
/Parent 2 0 R
/Resources 
/Font 
/F1 4 0 R
>>
>>
/MediaBox [0 0 612 792]
/Contents 5 0 R
>>
endobj

4 0 obj

/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
endobj

5 0 obj

/Length {len(str(content))}
>>
stream
BT
/F1 12 Tf
50 750 Td
({doc_type.title()}) Tj
0 -20 Td
({str(content)[:100]}...) Tj
ET
endstream
endobj

xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000262 00000 n
0000000341 00000 n
trailer

/Size 6
/Root 1 0 R
>>
startxref
{500 + len(str(content))}
%%EOF"""
        
        return pdf_content.encode('latin-1')

def create_docx_attachment(content: Any, doc_type: str) -> bytes:
    """Crée une pièce jointe DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Titre
        title = doc.add_heading(doc_type.replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph()
        
        # Contenu
        if isinstance(content, str):
            # Ajouter le contenu avec mise en forme
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Détecter les titres (lignes en majuscules)
                    if para.isupper() and len(para) < 100:
                        doc.add_heading(para, level=2)
                    else:
                        p = doc.add_paragraph(para)
                        # Mise en forme basique
                        if para.startswith('ARTICLE') or para.startswith('Article'):
                            p.runs[0].bold = True
        
        elif isinstance(content, dict):
            # Bordereau ou autre structure
            if 'header' in content:
                doc.add_paragraph(content['header'])
                doc.add_paragraph()
            
            if 'pieces' in content:
                doc.add_heading('Liste des pièces', level=1)
                
                # Créer un tableau pour les pièces
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Grid Accent 1'
                
                # En-têtes
                header_cells = table.rows[0].cells
                header_cells[0].text = 'N°'
                header_cells[1].text = 'Titre'
                header_cells[2].text = 'Description'
                
                # Pièces
                for piece in content.get('pieces', []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(piece.get('numero', ''))
                    row_cells[1].text = piece.get('titre', '')
                    row_cells[2].text = piece.get('description', '')
        
        # Sauvegarder
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError:
        # Fallback texte
        return create_txt_attachment(content, doc_type)

def create_html_attachment(content: Any, doc_type: str) -> bytes:
    """Crée une pièce jointe HTML"""
    
    html_template = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{doc_type.replace('_', ' ').title()}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 40px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
        }}
        .meta {{
            color: #7f8c8d;
            font-style: italic;
            margin-bottom: 30px;
        }}
        .piece {{
            background: #f8f9fa;
            padding: 15px;
            margin: 10px 0;
            border-left: 4px solid #3498db;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
    </style>
</head>
<body>
    <h1>{doc_type.replace('_', ' ').title()}</h1>
    <p class="meta">Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</p>
"""
    
    if isinstance(content, str):
        # Convertir le texte en HTML
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                if para.isupper() and len(para) < 100:
                    html_template += f"<h2>{para}</h2>\n"
                else:
                    html_template += f"<p>{para.replace(chr(10), '<br>')}</p>\n"
    
    elif isinstance(content, dict):
        # Traiter les structures complexes
        if 'pieces' in content:
            html_template += "<h2>Liste des pièces</h2>\n<table>\n<tr><th>N°</th><th>Titre</th><th>Description</th></tr>\n"
            
            for piece in content.get('pieces', []):
                html_template += f"<tr><td>{piece.get('numero', '')}</td>"
                html_template += f"<td>{piece.get('titre', '')}</td>"
                html_template += f"<td>{piece.get('description', '')}</td></tr>\n"
            
            html_template += "</table>\n"
    
    html_template += """
</body>
</html>"""
    
    return html_template.encode('utf-8')

def create_txt_attachment(content: Any, doc_type: str) -> bytes:
    """Crée une pièce jointe TXT basique"""
    
    txt_content = f"{doc_type.replace('_', ' ').upper()}\n"
    txt_content += f"{'=' * len(doc_type)}\n\n"
    txt_content += f"Date : {datetime.now().strftime('%d/%m/%Y')}\n\n"
    
    if isinstance(content, str):
        txt_content += content
    elif isinstance(content, dict):
        txt_content += json.dumps(content, indent=2, ensure_ascii=False)
    else:
        txt_content += str(content)
    
    return txt_content.encode('utf-8')

def show_email_preview(email_config: EmailConfig):
    """Affiche un aperçu de l'email"""
    
    with st.expander("👁️ Aperçu de l'email", expanded=True):
        # En-têtes
        st.markdown("**En-têtes**")
        
        # Affichage stylisé des en-têtes
        headers_html = f"""
        <div style="background: #f0f2f6; padding: 15px; border-radius: 5px; font-family: monospace;">
            <div><b>De :</b> {st.secrets.get('smtp_username', st.session_state.get('smtp_config', {}).get('username', '[Votre email]'))}</div>
            <div><b>À :</b> {', '.join(email_config.to) if email_config.to else '[Aucun destinataire]'}</div>
        """
        
        if email_config.cc:
            headers_html += f"<div><b>Cc :</b> {', '.join(email_config.cc)}</div>"
        
        if email_config.bcc:
            headers_html += f"<div><b>Cci :</b> {', '.join(email_config.bcc)}</div>"
        
        headers_html += f"""
            <div><b>Objet :</b> {email_config.subject}</div>
            <div><b>Date :</b> {datetime.now().strftime('%a %d %b %Y %H:%M:%S')}</div>
        """
        
        if email_config.priority != "normal":
            priority_icons = {"high": "🔴", "low": "🔵"}
            headers_html += f"<div><b>Priorité :</b> {priority_icons.get(email_config.priority, '')} {email_config.priority.title()}</div>"
        
        headers_html += "</div>"
        
        st.markdown(headers_html, unsafe_allow_html=True)
        
        st.divider()
        
        # Corps
        st.markdown("**Message**")
        
        # Mode édition ou aperçu
        if st.session_state.get('edit_email_mode'):
            email_config.body = st.text_area(
                "Corps du message",
                value=email_config.body,
                height=400,
                key="edit_email_body"
            )
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("✅ Valider les modifications", key="save_email_edits", use_container_width=True):
                    st.session_state.edit_email_mode = False
                    st.rerun()
            
            with col2:
                if st.button("❌ Annuler", key="cancel_email_edits", use_container_width=True):
                    st.session_state.edit_email_mode = False
                    st.rerun()
        else:
            # Aperçu stylisé du corps
            body_preview = email_config.body.replace('\n', '<br>')
            st.markdown(f"""
            <div style="background: white; padding: 20px; border: 1px solid #ddd; border-radius: 5px; min-height: 300px;">
                {body_preview}
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("✏️ Modifier le message", key="start_edit_mode"):
                st.session_state.edit_email_mode = True
                st.rerun()
        
        # Pièces jointes
        if email_config.attachments:
            st.divider()
            st.markdown(f"**Pièces jointes ({len(email_config.attachments)})**")
            
            total_size = sum(len(att['data']) for att in email_config.attachments)
            
            # Tableau des pièces jointes
            attachments_data = []
            for att in email_config.attachments:
                attachments_data.append({
                    'Fichier': att['filename'],
                    'Taille': format_file_size(len(att['data'])),
                    'Type': att.get('mimetype', 'Unknown').split('/')[-1].upper()
                })
            
            st.table(attachments_data)
            
            st.caption(f"💾 Taille totale : {format_file_size(total_size)}")

def send_email_with_progress(email_config: EmailConfig):
    """Envoie l'email avec barre de progression"""
    
    # Vérifications
    if not email_config.to:
        st.error("❌ Aucun destinataire spécifié")
        return
    
    if not email_config.subject:
        st.error("❌ L'objet est requis")
        return
    
    # Configuration SMTP
    smtp_config = get_smtp_configuration()
    
    if not smtp_config:
        st.error("❌ Configuration email manquante. Configurez les paramètres SMTP.")
        show_smtp_configuration_help()
        return
    
    # Container pour les messages
    progress_container = st.container()
    
    with progress_container:
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Étape 1 : Préparation
            status_text.text("📝 Préparation du message...")
            progress_bar.progress(10)
            
            msg = create_mime_message(email_config, smtp_config['username'])
            
            # Étape 2 : Connexion
            status_text.text("📡 Connexion au serveur...")
            progress_bar.progress(30)
            
            # Configuration du serveur selon le type
            if smtp_config['server'] == 'smtp.gmail.com':
                server = smtplib.SMTP(smtp_config['server'], smtp_config['port'])
                server.starttls()
            elif smtp_config['port'] == 465:
                server = smtplib.SMTP_SSL(smtp_config['server'], smtp_config['port'])
            else:
                server = smtplib.SMTP(smtp_config['server'], smtp_config['port'])
                if smtp_config.get('use_tls', True):
                    server.starttls()
            
            # Étape 3 : Authentification
            status_text.text("🔐 Authentification...")
            progress_bar.progress(50)
            
            server.login(smtp_config['username'], smtp_config['password'])
            
            # Étape 4 : Envoi
            status_text.text("📤 Envoi du message...")
            progress_bar.progress(70)
            
            # Tous les destinataires
            all_recipients = email_config.to + email_config.cc + email_config.bcc
            
            server.send_message(msg, to_addrs=all_recipients)
            
            # Étape 5 : Finalisation
            progress_bar.progress(90)
            status_text.text("✅ Finalisation...")
            
            server.quit()
            
            # Étape 6 : Confirmation
            progress_bar.progress(100)
            status_text.text("✅ Email envoyé avec succès !")
            
            # Sauvegarder l'historique
            save_email_history(email_config)
            
            # Notification de succès détaillée
            success_message = f"""
            ### ✅ Email envoyé avec succès !
            
            **Destinataires :** {len(all_recipients)} personne(s)  
            **Pièces jointes :** {len(email_config.attachments)} fichier(s)  
            **Heure d'envoi :** {datetime.now().strftime('%H:%M:%S')}  
            """
            
            if email_config.attachments:
                total_size = sum(len(att['data']) for att in email_config.attachments)
                success_message += f"**Taille totale :** {format_file_size(total_size)}"
            
            st.success(success_message)
            
            # Options post-envoi
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📧 Nouveau message", key="new_after_send"):
                    st.session_state.email_config = None
                    st.rerun()
            
            with col2:
                if st.button("📋 Copier les destinataires", key="copy_recipients"):
                    st.session_state.last_email_recipients = email_config.to
                    st.info("Destinataires copiés")
            
            with col3:
                if st.button("📜 Voir l'historique", key="view_history"):
                    st.session_state.show_email_history = True
            
            # Nettoyer l'état
            st.session_state.last_email_recipients = email_config.to
            
            # Logs
            log_email_sent(email_config)
            
        except smtplib.SMTPAuthenticationError:
            progress_bar.empty()
            status_text.empty()
            st.error("""
            ❌ **Erreur d'authentification**
            
            Vérifiez vos identifiants SMTP :
            - Email correct
            - Mot de passe d'application (pas le mot de passe normal)
            - Validation en 2 étapes activée (pour Gmail)
            """)
            
        except smtplib.SMTPRecipientsRefused as e:
            progress_bar.empty()
            status_text.empty()
            refused = list(e.recipients.keys())
            st.error(f"❌ Destinataires refusés : {', '.join(refused)}")
            
        except smtplib.SMTPException as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"❌ Erreur SMTP : {str(e)}")
            
        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"❌ Erreur inattendue : {str(e)}")
            
            # Afficher plus de détails en mode debug
            if st.checkbox("Afficher les détails techniques"):
                st.code(str(e))
        
        finally:
            # Nettoyer l'interface après un délai
            import time
            time.sleep(3)
            progress_bar.empty()
            status_text.empty()

def get_smtp_configuration() -> Optional[Dict[str, Any]]:
    """Récupère la configuration SMTP"""
    
    # D'abord vérifier les secrets Streamlit
    if hasattr(st, 'secrets'):
        if all(key in st.secrets for key in ['smtp_server', 'smtp_port', 'smtp_username', 'smtp_password']):
            return {
                'server': st.secrets['smtp_server'],
                'port': int(st.secrets['smtp_port']),
                'username': st.secrets['smtp_username'],
                'password': st.secrets['smtp_password'],
                'use_tls': st.secrets.get('smtp_use_tls', True)
            }
    
    # Sinon vérifier la configuration en session
    if st.session_state.get('smtp_config'):
        return st.session_state.smtp_config
    
    return None

def create_mime_message(email_config: EmailConfig, from_address: str) -> MIMEMultipart:
    """Crée le message MIME complet"""
    
    msg = MIMEMultipart('mixed')
    
    # En-têtes principaux
    msg['From'] = from_address
    msg['To'] = ', '.join(email_config.to)
    
    if email_config.cc:
        msg['Cc'] = ', '.join(email_config.cc)
    
    msg['Subject'] = email_config.subject
    msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
    
    # En-têtes additionnels
    msg['Message-ID'] = f"<{datetime.now().strftime('%Y%m%d%H%M%S')}.{hash(email_config.subject)}@juridique-app>"
    
    # Priorité
    if email_config.priority == "high":
        msg['X-Priority'] = '1'
        msg['Importance'] = 'high'
        msg['X-MSMail-Priority'] = 'High'
    elif email_config.priority == "low":
        msg['X-Priority'] = '5'
        msg['Importance'] = 'low'
        msg['X-MSMail-Priority'] = 'Low'
    
    # Demandes de confirmation
    if st.session_state.get('email_receipt'):
        msg['Return-Receipt-To'] = from_address
        msg['Disposition-Notification-To'] = from_address
    
    if st.session_state.get('email_read_confirm'):
        msg['X-Confirm-Reading-To'] = from_address
    
    # Corps du message
    body_part = MIMEMultipart('alternative')
    
    # Version texte
    text_part = MIMEText(email_config.body, 'plain', 'utf-8')
    body_part.attach(text_part)
    
    # Version HTML (optionnelle)
    html_body = f"""
    <html>
    <body style="font-family: Arial, sans-serif;">
        {email_config.body.replace(chr(10), '<br>')}
    </body>
    </html>
    """
    html_part = MIMEText(html_body, 'html', 'utf-8')
    body_part.attach(html_part)
    
    msg.attach(body_part)
    
    # Pièces jointes
    for attachment in email_config.attachments:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment['data'])
        encoders.encode_base64(part)
        
        # Nom de fichier sécurisé
        safe_filename = attachment['filename'].encode('utf-8').decode('ascii', 'ignore')
        
        part.add_header(
            'Content-Disposition',
            f'attachment; filename="{safe_filename}"'
        )
        
        # Type MIME spécifique si disponible
        if 'mimetype' in attachment:
            part.set_type(attachment['mimetype'])
        
        msg.attach(part)
    
    return msg

def show_smtp_configuration_help():
    """Affiche l'aide pour la configuration SMTP"""
    
    with st.expander("⚙️ Configuration SMTP requise", expanded=True):
        st.markdown("""
        ### Configuration des paramètres email
        
        Pour envoyer des emails, vous devez configurer les paramètres SMTP.
        """)
        
        # Onglets pour différents fournisseurs
        provider_tabs = st.tabs(["Gmail", "Outlook", "Autre", "Configuration manuelle"])
        
        with provider_tabs[0]:
            st.markdown("""
            #### Configuration Gmail
            
            1. **Activer la validation en 2 étapes** dans votre compte Google
            2. **Créer un mot de passe d'application** :
               - Allez dans [Paramètres Google](https://myaccount.google.com/security)
               - Sécurité → Validation en 2 étapes → Mots de passe d'application
               - Créer un nouveau mot de passe pour "Email"
            3. **Utiliser ces paramètres** :
               - Serveur : `smtp.gmail.com`
               - Port : `587`
               - Email : votre adresse Gmail
               - Mot de passe : le mot de passe d'application (pas votre mot de passe normal)
            """)
            
            if st.button("Configurer Gmail", key="quick_gmail"):
                st.session_state.smtp_config = {
                    'server': 'smtp.gmail.com',
                    'port': 587,
                    'username': '',
                    'password': '',
                    'use_tls': True
                }
        
        with provider_tabs[1]:
            st.markdown("""
            #### Configuration Outlook/Office 365
            
            **Paramètres Outlook.com** :
            - Serveur : `smtp-mail.outlook.com`
            - Port : `587`
            - Authentification : Oui
            - TLS : Oui
            
            **Paramètres Office 365** :
            - Serveur : `smtp.office365.com`
            - Port : `587`
            - Authentification : Oui
            - TLS : Oui
            """)
            
            if st.button("Configurer Outlook", key="quick_outlook"):
                st.session_state.smtp_config = {
                    'server': 'smtp-mail.outlook.com',
                    'port': 587,
                    'username': '',
                    'password': '',
                    'use_tls': True
                }
        
        with provider_tabs[2]:
            st.markdown("""
            #### Autres fournisseurs
            
            **Yahoo Mail** :
            - Serveur : `smtp.mail.yahoo.com`
            - Port : `587` ou `465`
            
            **Orange/Wanadoo** :
            - Serveur : `smtp.orange.fr`
            - Port : `465`
            
            **Free** :
            - Serveur : `smtp.free.fr`
            - Port : `465`
            
            **SFR** :
            - Serveur : `smtp.sfr.fr`
            - Port : `465`
            """)
        
        with provider_tabs[3]:
            # Formulaire de configuration manuelle
            with st.form("smtp_config_form"):
                st.markdown("#### Configuration manuelle")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    server = st.text_input("Serveur SMTP", placeholder="smtp.example.com")
                    username = st.text_input("Email/Nom d'utilisateur", placeholder="votre.email@example.com")
                
                with col2:
                    port = st.number_input("Port", value=587, min_value=1, max_value=65535)
                    password = st.text_input("Mot de passe", type="password")
                
                col3, col4 = st.columns(2)
                
                with col3:
                    use_tls = st.checkbox("Utiliser TLS/STARTTLS", value=True)
                
                with col4:
                    use_ssl = st.checkbox("Utiliser SSL", value=False)
                
                if st.form_submit_button("💾 Sauvegarder et tester", type="primary"):
                    if all([server, username, password]):
                        config = {
                            'server': server,
                            'port': port,
                            'username': username,
                            'password': password,
                            'use_tls': use_tls,
                            'use_ssl': use_ssl
                        }
                        
                        # Tester la configuration
                        if test_smtp_configuration(config):
                            st.session_state.smtp_config = config
                            st.success("✅ Configuration sauvegardée et testée avec succès !")
                            st.rerun()
                    else:
                        st.error("❌ Veuillez remplir tous les champs")
        
        # Option secrets.toml
        with st.expander("💡 Configuration via secrets.toml (recommandé)"):
            st.markdown("""
            Pour une configuration permanente, ajoutez dans `.streamlit/secrets.toml` :
            
            ```toml
            smtp_server = "smtp.gmail.com"
            smtp_port = 587
            smtp_username = "votre.email@gmail.com"
            smtp_password = "votre_mot_de_passe_app"
            smtp_use_tls = true
            ```
            
            Cette méthode est plus sécurisée car les identifiants ne sont pas stockés dans le code.
            """)

def save_email_draft(email_config: EmailConfig):
    """Sauvegarde un brouillon d'email"""
    
    if 'email_drafts' not in st.session_state:
        st.session_state.email_drafts = []
    
    # Vérifier si c'est une mise à jour d'un brouillon existant
    draft_id = st.session_state.get('current_draft_id')
    
    if draft_id:
        # Mettre à jour le brouillon existant
        for draft in st.session_state.email_drafts:
            if draft['id'] == draft_id:
                draft['updated_at'] = datetime.now()
                draft['config'] = email_config
                draft['subject'] = email_config.subject
                st.success("✅ Brouillon mis à jour")
                return
    
    # Créer un nouveau brouillon
    draft = {
        'id': f"draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{len(st.session_state.email_drafts)}",
        'created_at': datetime.now(),
        'updated_at': datetime.now(),
        'config': email_config,
        'subject': email_config.subject or "Sans objet",
        'preview': email_config.body[:100] + "..." if len(email_config.body) > 100 else email_config.body
    }
    
    st.session_state.email_drafts.append(draft)
    st.session_state.current_draft_id = draft['id']
    
    st.success("✅ Brouillon sauvegardé")
    
    # Notification avec options
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Continuer l'édition", key="continue_draft"):
            pass  # Reste sur la page actuelle
    
    with col2:
        if st.button("Voir les brouillons", key="view_drafts"):
            st.session_state.show_email_drafts = True

def save_email_history(email_config: EmailConfig):
    """Sauvegarde l'historique des emails envoyés"""
    
    if 'email_history' not in st.session_state:
        st.session_state.email_history = []
    
    history_entry = {
        'id': f"sent_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
        'sent_at': datetime.now(),
        'to': email_config.to.copy(),
        'cc': email_config.cc.copy(),
        'bcc': email_config.bcc.copy(),
        'subject': email_config.subject,
        'body_preview': email_config.body[:200] + "..." if len(email_config.body) > 200 else email_config.body,
        'attachments': [
            {
                'filename': att['filename'],
                'size': len(att['data'])
            }
            for att in email_config.attachments
        ],
        'attachments_count': len(email_config.attachments),
        'priority': email_config.priority,
        'total_size': sum(len(att['data']) for att in email_config.attachments)
    }
    
    st.session_state.email_history.append(history_entry)
    
    # Limiter l'historique à 100 entrées
    if len(st.session_state.email_history) > 100:
        st.session_state.email_history = st.session_state.email_history[-100:]
    
    # Supprimer le brouillon s'il existe
    if st.session_state.get('current_draft_id'):
        st.session_state.email_drafts = [
            d for d in st.session_state.get('email_drafts', [])
            if d['id'] != st.session_state.current_draft_id
        ]
        st.session_state.current_draft_id = None

def log_email_sent(email_config: EmailConfig):
    """Enregistre l'envoi dans les logs"""
    
    log_entry = {
        'timestamp': datetime.now().isoformat(),
        'action': 'email_sent',
        'recipients_count': len(email_config.to + email_config.cc + email_config.bcc),
        'attachments_count': len(email_config.attachments),
        'subject': email_config.subject[:50] + '...' if len(email_config.subject) > 50 else email_config.subject,
        'status': 'success'
    }
    
    # Ajouter aux statistiques
    if 'email_stats' not in st.session_state:
        st.session_state.email_stats = {
            'total_sent': 0,
            'total_recipients': 0,
            'total_attachments': 0,
            'last_sent': None
        }
    
    st.session_state.email_stats['total_sent'] += 1
    st.session_state.email_stats['total_recipients'] += log_entry['recipients_count']
    st.session_state.email_stats['total_attachments'] += log_entry['attachments_count']
    st.session_state.email_stats['last_sent'] = datetime.now()
    
    # Logger (pour le moment, juste print)
    print(f"[EMAIL LOG] {log_entry}")

def show_email_interface():
    """Interface principale de gestion des emails"""
    
    st.markdown("### 📧 Centre de messagerie")
    
    # Statistiques rapides
    if st.session_state.get('email_stats'):
        stats = st.session_state.email_stats
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Emails envoyés", stats['total_sent'])
        
        with col2:
            st.metric("Destinataires", stats['total_recipients'])
        
        with col3:
            st.metric("Pièces jointes", stats['total_attachments'])
        
        with col4:
            if stats['last_sent']:
                st.metric("Dernier envoi", stats['last_sent'].strftime('%H:%M'))
    
    # Tabs principaux
    tabs = st.tabs([
        "📤 Nouveau message",
        "📝 Brouillons",
        "📜 Historique",
        "📇 Contacts",
        "⚙️ Configuration"
    ])
    
    with tabs[0]:
        # Nouveau message
        process_email_request("", {})
    
    with tabs[1]:
        # Brouillons
        show_email_drafts()
    
    with tabs[2]:
        # Historique
        show_email_history()
    
    with tabs[3]:
        # Contacts
        show_contacts_management()
    
    with tabs[4]:
        # Configuration
        show_email_configuration_interface()

def show_email_drafts():
    """Affiche les brouillons d'emails"""
    
    drafts = st.session_state.get('email_drafts', [])
    
    if not drafts:
        st.info("📭 Aucun brouillon sauvegardé")
        
        if st.button("📝 Créer un nouveau message"):
            st.session_state.email_config = None
            st.rerun()
        
        return
    
    # Filtres et tri
    col1, col2, col3 = st.columns(3)
    
    with col1:
        search_term = st.text_input("🔍 Rechercher", key="search_drafts")
    
    with col2:
        sort_by = st.selectbox(
            "Trier par",
            ["Date (récent)", "Date (ancien)", "Objet"],
            key="sort_drafts"
        )
    
    with col3:
        if st.button("🗑️ Supprimer tous", key="delete_all_drafts"):
            if st.checkbox("Confirmer la suppression"):
                st.session_state.email_drafts = []
                st.rerun()
    
    # Filtrer les brouillons
    filtered_drafts = drafts
    
    if search_term:
        filtered_drafts = [
            d for d in filtered_drafts
            if search_term.lower() in d['subject'].lower() or
               search_term.lower() in d.get('preview', '').lower()
        ]
    
    # Trier
    if sort_by == "Date (récent)":
        filtered_drafts.sort(key=lambda x: x['updated_at'], reverse=True)
    elif sort_by == "Date (ancien)":
        filtered_drafts.sort(key=lambda x: x['updated_at'])
    else:  # Objet
        filtered_drafts.sort(key=lambda x: x['subject'])
    
    st.write(f"**{len(filtered_drafts)} brouillon(s)**")
    
    # Afficher les brouillons
    for draft in filtered_drafts:
        with st.container():
            col1, col2, col3, col4, col5 = st.columns([3, 2, 1, 1, 1])
            
            with col1:
                st.markdown(f"**{draft['subject']}**")
                st.caption(draft.get('preview', '')[:80] + "...")
            
            with col2:
                created = draft['created_at'].strftime('%d/%m %H:%M')
                updated = draft['updated_at'].strftime('%d/%m %H:%M')
                
                if created != updated:
                    st.caption(f"Modifié le {updated}")
                else:
                    st.caption(f"Créé le {created}")
                
                # Destinataires
                if draft['config'].to:
                    recipients = ", ".join(draft['config'].to[:2])
                    if len(draft['config'].to) > 2:
                        recipients += f" +{len(draft['config'].to) - 2}"
                    st.caption(f"À : {recipients}")
            
            with col3:
                att_count = len(draft['config'].attachments)
                if att_count > 0:
                    st.caption(f"📎 {att_count}")
            
            with col4:
                if st.button("✏️", key=f"edit_draft_{draft['id']}", help="Éditer"):
                    st.session_state.email_config = draft['config']
                    st.session_state.current_draft_id = draft['id']
                    st.session_state.current_tab = 0  # Retour à l'onglet nouveau message
                    st.rerun()
            
            with col5:
                if st.button("🗑️", key=f"delete_draft_{draft['id']}", help="Supprimer"):
                    st.session_state.email_drafts.remove(draft)
                    st.rerun()
        
        st.divider()

def show_email_history():
    """Affiche l'historique des emails envoyés"""
    
    history = st.session_state.get('email_history', [])
    
    if not history:
        st.info("📭 Aucun email envoyé")
        return
    
    # Statistiques de l'historique
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Total envoyé", len(history))
    
    with col2:
        total_recipients = sum(len(h['to'] + h['cc'] + h.get('bcc', [])) for h in history)
        st.metric("Total destinataires", total_recipients)
    
    with col3:
        total_attachments = sum(h['attachments_count'] for h in history)
        st.metric("Total pièces jointes", total_attachments)
    
    st.divider()
    
    # Filtres
    col1, col2, col3 = st.columns(3)
    
    with col1:
        date_range = st.date_input(
            "Période",
            value=[],
            key="email_history_date_range"
        )
    
    with col2:
        search_filter = st.text_input(
            "🔍 Rechercher",
            placeholder="Objet, destinataire...",
            key="email_history_search"
        )
    
    with col3:
        filter_priority = st.selectbox(
            "Priorité",
            ["Toutes", "Haute", "Normale", "Faible"],
            key="filter_priority"
        )
    
    # Appliquer les filtres
    filtered_history = history
    
    if date_range:
        if len(date_range) == 2:
            start_date, end_date = date_range
            filtered_history = [
                h for h in filtered_history
                if start_date <= h['sent_at'].date() <= end_date
            ]
    
    if search_filter:
        search_lower = search_filter.lower()
        filtered_history = [
            h for h in filtered_history
            if search_lower in h['subject'].lower() or
               any(search_lower in recipient.lower() for recipient in h['to'])
        ]
    
    if filter_priority != "Toutes":
        priority_map = {"Haute": "high", "Normale": "normal", "Faible": "low"}
        filtered_history = [
            h for h in filtered_history
            if h['priority'] == priority_map[filter_priority]
        ]
    
    # Pagination
    items_per_page = 20
    total_pages = max(1, (len(filtered_history) - 1) // items_per_page + 1)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        current_page = st.number_input(
            "Page",
            min_value=1,
            max_value=total_pages,
            value=1,
            key="history_page"
        )
    
    # Afficher l'historique paginé
    start_idx = (current_page - 1) * items_per_page
    end_idx = start_idx + items_per_page
    
    st.write(f"**{len(filtered_history)} email(s) - Page {current_page}/{total_pages}**")
    
    for entry in reversed(filtered_history[start_idx:end_idx]):
        with st.container():
            col1, col2, col3, col4, col5 = st.columns([3, 2, 1, 1, 1])
            
            with col1:
                # Icône de priorité
                priority_icons = {"high": "🔴", "low": "🔵", "normal": ""}
                icon = priority_icons.get(entry['priority'], "")
                
                st.markdown(f"{icon} **{entry['subject']}**")
                
                # Aperçu du corps
                if entry.get('body_preview'):
                    st.caption(entry['body_preview'][:100] + "...")
            
            with col2:
                st.caption(f"📅 {entry['sent_at'].strftime('%d/%m/%Y %H:%M')}")
                
                # Destinataires
                recipients_info = f"À : {len(entry['to'])}"
                if entry['cc']:
                    recipients_info += f", Cc : {len(entry['cc'])}"
                st.caption(recipients_info)
            
            with col3:
                if entry['attachments_count'] > 0:
                    st.caption(f"📎 {entry['attachments_count']}")
                    
                    if entry.get('total_size'):
                        st.caption(format_file_size(entry['total_size']))
            
            with col4:
                if st.button("👁️", key=f"view_sent_{entry['id']}", help="Voir détails"):
                    show_email_details(entry)
            
            with col5:
                if st.button("♻️", key=f"reuse_sent_{entry['id']}", help="Réutiliser"):
                    # Créer une nouvelle config basée sur cet historique
                    new_config = EmailConfig(
                        to=entry['to'].copy(),
                        cc=entry['cc'].copy(),
                        bcc=entry.get('bcc', []).copy(),
                        subject=f"Re: {entry['subject']}",
                        body="",  # Corps vide pour nouveau message
                        priority=entry['priority']
                    )
                    
                    st.session_state.email_config = new_config
                    st.session_state.current_tab = 0
                    st.rerun()
        
        st.divider()
    
    # Export de l'historique
    st.markdown("### 📥 Export")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📄 Export JSON", key="export_json"):
            export_email_history(filtered_history, format='json')
    
    with col2:
        if st.button("📊 Export CSV", key="export_csv"):
            export_email_history(filtered_history, format='csv')
    
    with col3:
        if st.button("📑 Export Excel", key="export_excel"):
            export_email_history(filtered_history, format='excel')

def show_email_details(entry: Dict[str, Any]):
    """Affiche les détails d'un email envoyé"""
    
    with st.expander(f"Détails : {entry['subject']}", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Informations d'envoi**")
            st.write(f"Date : {entry['sent_at'].strftime('%d/%m/%Y à %H:%M:%S')}")
            st.write(f"ID : {entry['id']}")
            
            priority_labels = {"high": "🔴 Haute", "normal": "⚪ Normale", "low": "🔵 Faible"}
            st.write(f"Priorité : {priority_labels.get(entry['priority'], entry['priority'])}")
        
        with col2:
            st.markdown("**Destinataires**")
            st.write(f"À : {', '.join(entry['to'])}")
            
            if entry['cc']:
                st.write(f"Cc : {', '.join(entry['cc'])}")
            
            if entry.get('bcc'):
                st.write(f"Cci : {', '.join(entry['bcc'])}")
        
        st.divider()
        
        # Pièces jointes
        if entry.get('attachments'):
            st.markdown("**Pièces jointes**")
            
            for att in entry['attachments']:
                st.write(f"• {att['filename']} ({format_file_size(att['size'])})")
        
        # Aperçu du message
        if entry.get('body_preview'):
            st.divider()
            st.markdown("**Aperçu du message**")
            st.text_area("", value=entry['body_preview'], height=200, disabled=True)

def show_contacts_management():
    """Gestion des contacts et carnets d'adresses"""
    
    st.markdown("### 📇 Gestion des contacts")
    
    # Initialiser les contacts
    if 'favorite_contacts' not in st.session_state:
        st.session_state.favorite_contacts = {
            'Greffe TJ': 'greffe@tribunal.fr',
            'Avocat adverse': 'contact@cabinet-adverse.fr'
        }
    
    contacts = st.session_state.favorite_contacts
    
    # Ajouter un contact
    with st.form("add_contact_form"):
        st.markdown("#### ➕ Ajouter un contact")
        
        col1, col2 = st.columns(2)
        
        with col1:
            new_name = st.text_input("Nom", placeholder="Ex: Maître Dupont")
        
        with col2:
            new_email = st.text_input("Email", placeholder="dupont@cabinet.fr")
        
        col3, col4 = st.columns(2)
        
        with col3:
            new_phone = st.text_input("Téléphone", placeholder="+33 1 23 45 67 89")
        
        with col4:
            new_notes = st.text_input("Notes", placeholder="Cabinet X, spécialiste...")
        
        if st.form_submit_button("Ajouter le contact", type="primary"):
            if new_name and new_email:
                if is_valid_email(new_email):
                    # Stocker avec plus d'infos
                    if 'contacts_detailed' not in st.session_state:
                        st.session_state.contacts_detailed = {}
                    
                    st.session_state.contacts_detailed[new_name] = {
                        'email': new_email,
                        'phone': new_phone,
                        'notes': new_notes,
                        'added_date': datetime.now()
                    }
                    
                    # Ajouter aussi aux favoris simples
                    contacts[new_name] = new_email
                    
                    st.success(f"✅ Contact '{new_name}' ajouté")
                    st.rerun()
                else:
                    st.error("❌ Email invalide")
            else:
                st.error("❌ Nom et email requis")
    
    st.divider()
    
    # Import/Export
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📥 Importer CSV", key="import_contacts"):
            show_import_contacts_dialog()
    
    with col2:
        if st.button("📤 Exporter", key="export_contacts"):
            export_contacts()
    
    with col3:
        search_contact = st.text_input("🔍 Rechercher", key="search_contacts")
    
    # Afficher les contacts
    st.markdown("#### 📋 Liste des contacts")
    
    # Filtrer si recherche
    filtered_contacts = contacts
    if search_contact:
        filtered_contacts = {
            name: email for name, email in contacts.items()
            if search_contact.lower() in name.lower() or
               search_contact.lower() in email.lower()
        }
    
    if not filtered_contacts:
        st.info("Aucun contact trouvé")
    else:
        # Affichage en tableau
        contacts_data = []
        
        for name, email in filtered_contacts.items():
            # Récupérer les infos détaillées si disponibles
            detailed = st.session_state.get('contacts_detailed', {}).get(name, {})
            
            contacts_data.append({
                'Nom': name,
                'Email': email,
                'Téléphone': detailed.get('phone', '-'),
                'Notes': detailed.get('notes', '-'),
                'Actions': name  # Pour les boutons
            })
        
        # Créer le tableau
        for contact in contacts_data:
            col1, col2, col3, col4, col5, col6 = st.columns([2, 3, 2, 2, 1, 1])
            
            with col1:
                st.write(contact['Nom'])
            
            with col2:
                st.write(contact['Email'])
            
            with col3:
                st.write(contact['Téléphone'])
            
            with col4:
                st.write(contact['Notes'])
            
            with col5:
                if st.button("✏️", key=f"edit_contact_{contact['Actions']}"):
                    st.session_state.editing_contact = contact['Actions']
            
            with col6:
                if st.button("🗑️", key=f"delete_contact_{contact['Actions']}"):
                    del contacts[contact['Actions']]
                    if 'contacts_detailed' in st.session_state:
                        if contact['Actions'] in st.session_state.contacts_detailed:
                            del st.session_state.contacts_detailed[contact['Actions']]
                    st.rerun()
            
            # Mode édition
            if st.session_state.get('editing_contact') == contact['Actions']:
                with st.form(f"edit_form_{contact['Actions']}"):
                    edited_email = st.text_input("Email", value=contact['Email'])
                    edited_phone = st.text_input("Téléphone", value=contact['Téléphone'])
                    edited_notes = st.text_input("Notes", value=contact['Notes'])
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        if st.form_submit_button("💾 Sauvegarder"):
                            contacts[contact['Actions']] = edited_email
                            
                            if 'contacts_detailed' not in st.session_state:
                                st.session_state.contacts_detailed = {}
                            
                            st.session_state.contacts_detailed[contact['Actions']] = {
                                'email': edited_email,
                                'phone': edited_phone,
                                'notes': edited_notes,
                                'added_date': datetime.now()
                            }
                            
                            st.session_state.editing_contact = None
                            st.rerun()
                    
                    with col2:
                        if st.form_submit_button("❌ Annuler"):
                            st.session_state.editing_contact = None
                            st.rerun()
            
            st.divider()

def show_import_contacts_dialog():
    """Dialog pour importer des contacts"""
    
    with st.expander("📥 Importer des contacts", expanded=True):
        st.markdown("""
        Format CSV attendu :
        ```
        Nom,Email,Téléphone,Notes
        Maître Dupont,dupont@cabinet.fr,+33123456789,Cabinet X
        ```
        """)
        
        uploaded_file = st.file_uploader(
            "Choisir un fichier CSV",
            type=['csv'],
            key="import_contacts_file"
        )
        
        if uploaded_file:
            try:
                import pandas as pd
                
                df = pd.read_csv(uploaded_file)
                
                # Vérifier les colonnes
                required_cols = ['Nom', 'Email']
                if not all(col in df.columns for col in required_cols):
                    st.error(f"❌ Colonnes requises : {', '.join(required_cols)}")
                    return
                
                # Aperçu
                st.write("Aperçu des contacts à importer :")
                st.dataframe(df.head(10))
                
                if st.button("✅ Importer", key="confirm_import"):
                    imported = 0
                    errors = 0
                    
                    for _, row in df.iterrows():
                        if is_valid_email(row['Email']):
                            st.session_state.favorite_contacts[row['Nom']] = row['Email']
                            
                            # Détails si disponibles
                            if 'contacts_detailed' not in st.session_state:
                                st.session_state.contacts_detailed = {}
                            
                            st.session_state.contacts_detailed[row['Nom']] = {
                                'email': row['Email'],
                                'phone': row.get('Téléphone', ''),
                                'notes': row.get('Notes', ''),
                                'added_date': datetime.now()
                            }
                            
                            imported += 1
                        else:
                            errors += 1
                    
                    st.success(f"✅ {imported} contacts importés")
                    if errors > 0:
                        st.warning(f"⚠️ {errors} contacts ignorés (email invalide)")
                    
                    st.rerun()
                    
            except Exception as e:
                st.error(f"❌ Erreur lors de l'import : {str(e)}")

def export_contacts():
    """Exporte les contacts"""
    
    contacts = st.session_state.favorite_contacts
    detailed = st.session_state.get('contacts_detailed', {})
    
    # Créer le CSV
    csv_lines = ["Nom,Email,Téléphone,Notes"]
    
    for name, email in contacts.items():
        details = detailed.get(name, {})
        phone = details.get('phone', '')
        notes = details.get('notes', '')
        
        csv_lines.append(f'"{name}","{email}","{phone}","{notes}"')
    
    csv_content = '\n'.join(csv_lines)
    
    # Télécharger
    st.download_button(
        "💾 Télécharger CSV",
        csv_content,
        f"contacts_{datetime.now().strftime('%Y%m%d')}.csv",
        "text/csv",
        key="download_contacts"
    )

def export_email_history(history: List[Dict[str, Any]], format: str = 'json'):
    """Exporte l'historique des emails dans différents formats"""
    
    if format == 'json':
        # Préparer les données pour JSON
        export_data = []
        
        for entry in history:
            export_entry = {
                'id': entry['id'],
                'date': entry['sent_at'].isoformat(),
                'destinataires_to': entry['to'],
                'destinataires_cc': entry['cc'],
                'destinataires_cci': entry.get('bcc', []),
                'objet': entry['subject'],
                'pieces_jointes': entry.get('attachments_count', 0),
                'priorite': entry.get('priority', 'normal'),
                'taille_totale': entry.get('total_size', 0)
            }
            export_data.append(export_entry)
        
        # JSON
        json_str = json.dumps(export_data, indent=2, ensure_ascii=False)
        
        st.download_button(
            "💾 Télécharger JSON",
            json_str,
            f"historique_emails_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "application/json",
            key="download_email_history_json"
        )
    
    elif format == 'csv':
        # CSV
        csv_lines = ["Date,Heure,Objet,Destinataires (À),CC,Pièces jointes,Priorité"]
        
        for entry in history:
            date = entry['sent_at'].strftime('%d/%m/%Y')
            time = entry['sent_at'].strftime('%H:%M:%S')
            subject = entry['subject'].replace('"', '""')  # Échapper les guillemets
            to = ';'.join(entry['to'])
            cc = ';'.join(entry['cc'])
            attachments = str(entry.get('attachments_count', 0))
            priority = entry.get('priority', 'normal')
            
            csv_lines.append(f'"{date}","{time}","{subject}","{to}","{cc}",{attachments},{priority}')
        
        csv_content = '\n'.join(csv_lines)
        
        st.download_button(
            "💾 Télécharger CSV",
            csv_content.encode('utf-8-sig'),  # BOM pour Excel
            f"historique_emails_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            "text/csv",
            key="download_email_history_csv"
        )
    
    elif format == 'excel':
        try:
            import pandas as pd
            import io
            
            # Préparer les données pour DataFrame
            data_rows = []
            
            for entry in history:
                data_rows.append({
                    'Date': entry['sent_at'].strftime('%d/%m/%Y'),
                    'Heure': entry['sent_at'].strftime('%H:%M:%S'),
                    'Objet': entry['subject'],
                    'Destinataires (À)': ', '.join(entry['to']),
                    'CC': ', '.join(entry['cc']),
                    'CCI': ', '.join(entry.get('bcc', [])),
                    'Pièces jointes': entry.get('attachments_count', 0),
                    'Taille totale (KB)': entry.get('total_size', 0) // 1024,
                    'Priorité': entry.get('priority', 'normal').title()
                })
            
            df = pd.DataFrame(data_rows)
            
            # Créer le fichier Excel
            output = io.BytesIO()
            
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df.to_excel(writer, sheet_name='Historique', index=False)
                
                # Formater
                workbook = writer.book
                worksheet = writer.sheets['Historique']
                
                # Format pour les en-têtes
                header_format = workbook.add_format({
                    'bold': True,
                    'bg_color': '#4CAF50',
                    'font_color': 'white'
                })
                
                # Appliquer le format aux en-têtes
                for col_num, value in enumerate(df.columns.values):
                    worksheet.write(0, col_num, value, header_format)
                
                # Ajuster la largeur des colonnes
                worksheet.set_column('A:B', 12)  # Date et Heure
                worksheet.set_column('C:C', 40)  # Objet
                worksheet.set_column('D:F', 30)  # Destinataires
                worksheet.set_column('G:I', 15)  # Stats
            
            output.seek(0)
            
            st.download_button(
                "💾 Télécharger Excel",
                output,
                f"historique_emails_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_email_history_excel"
            )
            
        except ImportError:
            st.error("❌ pandas/xlsxwriter requis pour l'export Excel")
            # Fallback vers CSV
            export_email_history(history, format='csv')

def show_email_configuration_interface():
    """Interface de configuration des emails"""
    
    st.markdown("#### ⚙️ Configuration email")
    
    # État de la configuration
    smtp_config = get_smtp_configuration()
    
    if smtp_config:
        col1, col2 = st.columns(2)
        
        with col1:
            st.success(f"✅ Configuré ({smtp_config['server']})")
        
        with col2:
            if st.button("🔌 Tester la connexion", key="test_smtp"):
                test_smtp_connection(smtp_config)
    else:
        st.warning("❌ Configuration SMTP non définie")
        show_smtp_configuration_help()
    
    st.divider()
    
    # Templates d'emails
    st.markdown("#### 📝 Templates d'emails")
    
    if 'email_templates' not in st.session_state:
        st.session_state.email_templates = get_default_email_templates()
    
    email_templates = st.session_state.email_templates
    
    # Ajouter un nouveau template
    if st.checkbox("➕ Créer un nouveau template", key="create_template_check"):
        with st.form("new_template_form"):
            col1, col2 = st.columns(2)
            
            with col1:
                template_key = st.text_input("Identifiant", placeholder="mon_template")
            
            with col2:
                template_name = st.text_input("Nom", placeholder="Mon template")
            
            template_subject = st.text_input("Objet", placeholder="[Type] - [Référence]")
            template_body = st.text_area("Corps", height=200, placeholder="Contenu du template...")
            
            if st.form_submit_button("Créer le template"):
                if template_key and template_name:
                    email_templates[template_key] = {
                        'name': template_name,
                        'subject': template_subject,
                        'body': template_body
                    }
                    st.success(f"✅ Template '{template_name}' créé")
                    st.rerun()
    
    # Modifier les templates existants
    for template_name, template in email_templates.items():
        with st.expander(f"📧 {template['name']}"):
            # Formulaire d'édition
            with st.form(f"edit_template_{template_name}"):
                new_subject = st.text_input(
                    "Objet",
                    value=template['subject'],
                    key=f"subject_{template_name}"
                )
                
                new_body = st.text_area(
                    "Corps",
                    value=template['body'],
                    height=200,
                    key=f"body_{template_name}"
                )
                
                st.caption("Variables disponibles : [référence], [Référence], [date], [destinataire]")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.form_submit_button("💾 Sauvegarder"):
                        template['subject'] = new_subject
                        template['body'] = new_body
                        st.success("✅ Template sauvegardé")
                        st.rerun()
                
                with col2:
                    if st.form_submit_button("🗑️ Supprimer"):
                        del email_templates[template_name]
                        st.rerun()
    
    st.divider()
    
    # Signatures
    st.markdown("#### ✍️ Signatures")
    
    if 'email_signatures' not in st.session_state:
        st.session_state.email_signatures = get_default_signatures()
    
    signatures = st.session_state.email_signatures
    
    # Sélection de la signature par défaut
    col1, col2 = st.columns(2)
    
    with col1:
        selected_signature = st.selectbox(
            "Signature par défaut",
            options=[''] + list(signatures.keys()),
            format_func=lambda x: "Aucune" if x == '' else signatures[x]['name'],
            key="default_signature_select"
        )
        
        if selected_signature:
            st.session_state.default_signature = selected_signature
    
    with col2:
        if st.button("➕ Nouvelle signature", key="add_signature"):
            st.session_state.creating_signature = True
    
    # Créer une nouvelle signature
    if st.session_state.get('creating_signature'):
        with st.form("new_signature_form"):
            col1, col2 = st.columns(2)
            
            with col1:
                sig_key = st.text_input("Identifiant", placeholder="ma_signature")
            
            with col2:
                sig_name = st.text_input("Nom", placeholder="Ma signature")
            
            sig_content = st.text_area(
                "Contenu",
                height=150,
                placeholder="[Votre nom]\n[Fonction]\n[Email]\n[Téléphone]"
            )
            
            col3, col4 = st.columns(2)
            
            with col3:
                if st.form_submit_button("Créer"):
                    if sig_key and sig_name:
                        signatures[sig_key] = {
                            'name': sig_name,
                            'content': sig_content
                        }
                        st.session_state.creating_signature = False
                        st.success(f"✅ Signature '{sig_name}' créée")
                        st.rerun()
            
            with col4:
                if st.form_submit_button("Annuler"):
                    st.session_state.creating_signature = False
                    st.rerun()
    
    # Modifier les signatures existantes
    for sig_key, signature in signatures.items():
        with st.expander(f"✍️ {signature['name']}"):
            with st.form(f"edit_sig_{sig_key}"):
                new_content = st.text_area(
                    "Contenu",
                    value=signature['content'],
                    height=150,
                    key=f"sig_content_{sig_key}"
                )
                
                # Aperçu
                st.markdown("**Aperçu :**")
                st.text(new_content)
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.form_submit_button("💾 Sauvegarder"):
                        signature['content'] = new_content
                        st.success("✅ Signature sauvegardée")
                        st.rerun()
                
                with col2:
                    if st.form_submit_button("🗑️ Supprimer"):
                        del signatures[sig_key]
                        # Si c'était la signature par défaut
                        if st.session_state.get('default_signature') == sig_key:
                            st.session_state.default_signature = ''
                        st.rerun()
    
    st.divider()
    
    # Paramètres avancés
    st.markdown("#### 🔧 Paramètres avancés")
    
    with st.expander("Options d'envoi"):
        col1, col2 = st.columns(2)
        
        with col1:
            always_request_receipt = st.checkbox(
                "Toujours demander un accusé de réception",
                value=st.session_state.get('always_request_receipt', False),
                key="always_receipt"
            )
            
            always_request_read = st.checkbox(
                "Toujours demander une confirmation de lecture",
                value=st.session_state.get('always_request_read', False),
                key="always_read"
            )
        
        with col2:
            default_priority = st.selectbox(
                "Priorité par défaut",
                ["normal", "high", "low"],
                format_func=lambda x: {"normal": "Normale", "high": "Haute", "low": "Faible"}[x],
                key="default_priority"
            )
            
            max_attachment_size = st.number_input(
                "Taille max des pièces jointes (MB)",
                min_value=1,
                max_value=100,
                value=25,
                key="max_attachment_size"
            )
        
        if st.button("💾 Sauvegarder les paramètres"):
            st.session_state.email_settings = {
                'always_request_receipt': always_request_receipt,
                'always_request_read': always_request_read,
                'default_priority': default_priority,
                'max_attachment_size': max_attachment_size
            }
            st.success("✅ Paramètres sauvegardés")

def test_smtp_connection(config: Dict[str, Any]) -> bool:
    """Teste la connexion SMTP"""
    
    with st.spinner("Test de connexion en cours..."):
        try:
            # Connexion selon le type de serveur
            if config.get('use_ssl') or config['port'] == 465:
                server = smtplib.SMTP_SSL(config['server'], config['port'], timeout=10)
            else:
                server = smtplib.SMTP(config['server'], config['port'], timeout=10)
                if config.get('use_tls', True):
                    server.starttls()
            
            # Authentification
            server.login(config['username'], config['password'])
            
            # Fermer la connexion
            server.quit()
            
            st.success(f"""
            ✅ **Connexion réussie !**
            
            Serveur : {config['server']}:{config['port']}  
            Utilisateur : {config['username']}  
            Sécurité : {'SSL' if config.get('use_ssl') else 'TLS' if config.get('use_tls', True) else 'Aucune'}
            """)
            
            return True
            
        except smtplib.SMTPAuthenticationError:
            st.error("""
            ❌ **Erreur d'authentification**
            
            Vérifiez :
            - Nom d'utilisateur et mot de passe
            - Mot de passe d'application (Gmail, Outlook)
            - Validation en 2 étapes activée
            """)
            
        except smtplib.SMTPServerDisconnected:
            st.error("""
            ❌ **Connexion impossible**
            
            Vérifiez :
            - Serveur et port corrects
            - Connexion internet
            - Pare-feu/antivirus
            """)
            
        except Exception as e:
            st.error(f"❌ **Erreur** : {str(e)}")
        
        return False

def get_default_email_templates() -> Dict[str, Dict[str, str]]:
    """Retourne les templates d'emails par défaut"""
    
    return {
        'document_juridique': {
            'name': 'Document juridique',
            'subject': 'Document juridique - [Référence]',
            'body': """Maître,

Je vous prie de bien vouloir trouver ci-joint le document demandé concernant [référence].

Ce document a été préparé conformément à vos instructions.

Je reste à votre disposition pour tout complément d'information.

Bien cordialement,"""
        },
        'bordereau_pieces': {
            'name': 'Bordereau de pièces',
            'subject': 'Bordereau de communication de pièces - [Référence]',
            'body': """Maître,

Veuillez trouver ci-joint le bordereau de communication de pièces pour [référence].

Les pièces physiques vous seront transmises selon les modalités convenues.

Cordialement,"""
        },
        'transmission_urgente': {
            'name': 'Transmission urgente',
            'subject': 'URGENT - [Objet] - [Référence]',
            'body': """Maître,

Suite à notre échange, je vous transmets en urgence [description].

Compte tenu des délais, je vous serais reconnaissant de bien vouloir me confirmer la bonne réception.

Bien cordialement,"""
        },
        'conclusions': {
            'name': 'Envoi de conclusions',
            'subject': 'Conclusions - [Référence] - [Date]',
            'body': """Maître,

Je vous prie de bien vouloir trouver ci-joint mes conclusions en [demande/défense] dans l'affaire [référence].

Ces conclusions ont été déposées ce jour au greffe du Tribunal.

Je vous prie de croire, Maître, à l'assurance de mes sentiments confraternels les meilleurs.

Bien cordialement,"""
        },
        'demande_pieces': {
            'name': 'Demande de pièces',
            'subject': 'Demande de communication de pièces - [Référence]',
            'body': """Maître,

Dans le cadre du dossier [référence], je vous serais reconnaissant de bien vouloir me communiquer les pièces suivantes :

- [Pièce 1]
- [Pièce 2]
- [Pièce 3]

Je vous remercie par avance de votre diligence.

Bien cordialement,"""
        }
    }

def get_default_signatures() -> Dict[str, Dict[str, str]]:
    """Retourne les signatures par défaut"""
    
    return {
        'formelle': {
            'name': 'Formelle',
            'content': """[Votre nom]
[Votre fonction]
[Cabinet/Société]
[Adresse]
[Code postal] [Ville]
Tél : [Téléphone]
Email : [Email]"""
        },
        'simple': {
            'name': 'Simple',
            'content': """Cordialement,

[Votre nom]
[Téléphone]"""
        },
        'complete': {
            'name': 'Complète avec mentions légales',
            'content': """[Votre nom]
[Votre fonction]
[Cabinet/Société]
[Adresse]
[Code postal] [Ville]
Tél : [Téléphone]
Mobile : [Mobile]
Email : [Email]
Site web : [Site]

Ce message et ses pièces jointes sont confidentiels et établis à l'intention exclusive de leur destinataire. Toute utilisation non conforme à sa destination, toute diffusion ou toute publication, totale ou partielle, est interdite, sauf autorisation expresse."""
        },
        'avocat': {
            'name': 'Avocat',
            'content': """Maître [Votre nom]
Avocat au Barreau de [Ville]
[Cabinet]
[Adresse]
[Code postal] [Ville]
Tél : [Téléphone]
Fax : [Fax]
Email : [Email]
Toque : [Numéro]

Membre d'une association agréée, le règlement par chèque est accepté."""
        }
    }

# Fonction pour intégration avec d'autres modules
def prepare_and_send_document(document_type: str, content: str, reference: str = ""):
    """Prépare et envoie un document par email"""
    
    # Créer la configuration
    email_config = EmailConfig(
        to=[],
        subject=f"{document_type.title()} - {reference}",
        body=generate_email_body({'reference': reference})
    )
    
    # Ajouter la signature par défaut
    if st.session_state.get('default_signature'):
        signatures = st.session_state.get('email_signatures', get_default_signatures())
        if st.session_state.default_signature in signatures:
            email_config.body += "\n\n" + signatures[st.session_state.default_signature]['content']
    
    # Ajouter le document en pièce jointe
    attachment_data = create_docx_attachment(content, document_type)
    email_config.add_attachment(
        f"{document_type}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        attachment_data,
        ATTACHMENT_MIME_TYPES['docx']
    )
    
    # Stocker la configuration
    st.session_state.email_config = email_config
    
    # Marquer pour afficher l'interface email
    st.session_state.show_email_interface = True
    st.rerun()

# Fonction helper pour l'intégration
def quick_send_email(to: List[str], subject: str, body: str, attachments: List[Dict[str, Any]] = None):
    """Envoi rapide d'un email sans interface"""
    
    config = EmailConfig(
        to=to,
        subject=subject,
        body=body,
        attachments=attachments or []
    )
    
    smtp_config = get_smtp_configuration()
    if not smtp_config:
        st.error("❌ Configuration SMTP requise")
        return False
    
    try:
        msg = create_mime_message(config, smtp_config['username'])
        
        with smtplib.SMTP(smtp_config['server'], smtp_config['port']) as server:
            server.starttls()
            server.login(smtp_config['username'], smtp_config['password'])
            server.send_message(msg)
        
        save_email_history(config)
        return True
        
    except Exception as e:
        st.error(f"❌ Erreur d'envoi : {str(e)}")
        return False