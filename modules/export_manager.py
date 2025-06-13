# modules/export_manager.py
"""Gestionnaire unifié pour tous les exports (Word, PDF, Excel, etc.)"""

import streamlit as st
from datetime import datetime
from typing import Dict, Any, List, Optional, Union, Tuple, Callable
from dataclasses import dataclass
import io
import base64
import re
from pathlib import Path
import json

# Imports optionnels avec gestion des dépendances
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    st.warning("⚠️ pandas non disponible - Fonctionnalités Excel limitées")

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx non disponible - Export Word désactivé")

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    st.warning("⚠️ reportlab non disponible - Export PDF désactivé")

# Configuration centralisée
from config.cahier_des_charges import STYLE_DOCUMENT, HIERARCHIE_NUMEROTATION

# ========================= TYPES ET CONFIGURATIONS =========================

@dataclass
class ExportConfig:
    """Configuration unifiée pour tous les exports"""
    format: str  # 'docx', 'pdf', 'xlsx', 'txt', 'json', 'html', 'csv', 'rtf'
    title: str
    content: Union[str, Dict, List, pd.DataFrame]
    metadata: Optional[Dict] = None
    style_options: Optional[Dict] = None
    include_metadata: bool = True
    include_toc: bool = False
    include_header_footer: bool = True
    watermark: Optional[str] = None
    password: Optional[str] = None
    compression: bool = False
    
    def __post_init__(self):
        """Validation et initialisation des valeurs par défaut"""
        if self.style_options is None:
            self.style_options = self._get_default_style_options()
        
        if self.metadata is None:
            self.metadata = {
                'created_at': datetime.now(),
                'format': self.format,
                'generator': 'Export Manager v2.0'
            }
    
    def _get_default_style_options(self) -> Dict:
        """Options de style par défaut selon le format"""
        defaults = {
            'docx': {
                'font': 'Times New Roman',
                'font_size': 12,
                'line_spacing': '1.5',
                'margins': 'Normales',
                'paper_size': 'A4',
                'orientation': 'Portrait',
                'numbering': True,
                'legal_formatting': True
            },
            'pdf': {
                'font': 'Times-Roman',
                'font_size': 12,
                'margins': 2.0,  # cm
                'paper_size': 'A4',
                'orientation': 'Portrait',
                'compress_images': True,
                'embed_fonts': True
            },
            'xlsx': {
                'freeze_headers': True,
                'autofilter': True,
                'conditional_formatting': False,
                'include_charts': False,
                'column_autowidth': True
            },
            'html': {
                'theme': 'Juridique',
                'responsive': True,
                'include_css': True,
                'include_js': False,
                'print_friendly': True
            }
        }
        
        return defaults.get(self.format, {})

# ========================= GESTIONNAIRE PRINCIPAL =========================

class ExportManager:
    """Gestionnaire central pour tous les types d'export"""
    
    def __init__(self):
        self.exporters = {
            'docx': WordExporter(),
            'pdf': PDFExporter(),
            'xlsx': ExcelExporter(),
            'txt': TextExporter(),
            'json': JSONExporter(),
            'html': HTMLExporter(),
            'csv': CSVExporter(),
            'rtf': RTFExporter()
        }
        
        self._init_styles()
    
    def _init_styles(self):
        """Initialise les styles et configurations"""
        self.juridique_styles = {
            'title': {'size': 16, 'bold': True, 'align': 'center'},
            'heading1': {'size': 14, 'bold': True, 'underline': True},
            'heading2': {'size': 12, 'bold': True},
            'normal': {'size': 12, 'align': 'justify'},
            'citation': {'size': 10, 'italic': True, 'indent': 0.5}
        }
    
    def export(self, config: ExportConfig) -> Tuple[bytes, str]:
        """
        Point d'entrée principal pour tous les exports
        
        Returns:
            Tuple[bytes, str]: (contenu exporté, extension de fichier)
        """
        
        # Validation
        if config.format not in self.exporters:
            raise ValueError(f"Format non supporté : {config.format}")
        
        # Vérifier la disponibilité
        exporter = self.exporters[config.format]
        if not exporter.is_available():
            # Fallback sur format texte
            st.warning(f"⚠️ Export {config.format.upper()} non disponible, export en TXT")
            config.format = 'txt'
            exporter = self.exporters['txt']
        
        # Effectuer l'export
        try:
            content = exporter.export(config)
            
            # Post-traitement si nécessaire
            if config.compression:
                content = self._compress_content(content, config)
            
            if config.password and config.format in ['pdf', 'docx']:
                content = self._encrypt_content(content, config)
            
            return content, exporter.get_extension()
            
        except Exception as e:
            st.error(f"❌ Erreur lors de l'export : {str(e)}")
            # Fallback sur texte simple
            fallback_config = ExportConfig(
                format='txt',
                title=config.title,
                content=str(config.content),
                metadata=config.metadata
            )
            return self.exporters['txt'].export(fallback_config), 'txt'
    
    def get_available_formats(self, content_type: str = None) -> List[str]:
        """Retourne les formats disponibles selon le type de contenu"""
        available = []
        
        for format_name, exporter in self.exporters.items():
            if exporter.is_available():
                # Filtrer selon le type de contenu
                if content_type == 'table' and format_name in ['xlsx', 'csv', 'html']:
                    available.insert(0, format_name)  # Priorité pour les tableaux
                elif content_type == 'document' and format_name in ['docx', 'pdf']:
                    available.insert(0, format_name)  # Priorité pour les documents
                else:
                    available.append(format_name)
        
        return available
    
    def show_export_interface(self, 
                            content: Any,
                            title: str,
                            content_type: str = 'document',
                            key_prefix: str = "export") -> Optional[Tuple[bytes, str]]:
        """Interface Streamlit unifiée pour l'export"""
        
        st.markdown("### 📤 Options d'export")
        
        # Formats disponibles
        formats = self.get_available_formats(content_type)
        
        if not formats:
            st.error("❌ Aucun format d'export disponible")
            return None
        
        col1, col2, col3 = st.columns([2, 2, 1])
        
        with col1:
            selected_format = st.selectbox(
                "Format",
                formats,
                format_func=lambda x: self._get_format_display_name(x),
                key=f"{key_prefix}_format"
            )
        
        with col2:
            filename = st.text_input(
                "Nom du fichier",
                value=self._generate_filename(title, selected_format),
                key=f"{key_prefix}_filename"
            )
        
        with col3:
            st.markdown("&nbsp;")  # Spacer
            export_clicked = st.button(
                "📥 Exporter",
                type="primary",
                key=f"{key_prefix}_button",
                use_container_width=True
            )
        
        # Options avancées
        with st.expander("⚙️ Options avancées", expanded=False):
            config = self._show_advanced_options(selected_format, key_prefix)
        else:
            config = ExportConfig(
                format=selected_format,
                title=title,
                content=content
            )
        
        # Exécuter l'export
        if export_clicked:
            with st.spinner(f"Génération du fichier {selected_format.upper()}..."):
                try:
                    # Configuration complète
                    config.format = selected_format
                    config.title = title
                    config.content = content
                    
                    # Export
                    exported_data, extension = self.export(config)
                    
                    # Téléchargement
                    st.download_button(
                        f"⬇️ Télécharger {selected_format.upper()}",
                        exported_data,
                        f"{filename}.{extension}",
                        mime=self._get_mime_type(selected_format),
                        key=f"{key_prefix}_download"
                    )
                    
                    st.success(f"✅ Export {selected_format.upper()} prêt !")
                    
                    # Statistiques
                    self._show_export_stats(exported_data, selected_format)
                    
                    return exported_data, extension
                    
                except Exception as e:
                    st.error(f"❌ Erreur : {str(e)}")
                    return None
        
        return None
    
    def _show_advanced_options(self, format: str, key_prefix: str) -> ExportConfig:
        """Affiche les options avancées selon le format"""
        
        config = ExportConfig(format=format, title="", content="")
        
        col1, col2 = st.columns(2)
        
        with col1:
            config.include_metadata = st.checkbox(
                "Inclure les métadonnées",
                value=True,
                key=f"{key_prefix}_metadata"
            )
            
            config.include_toc = st.checkbox(
                "Table des matières",
                value=format in ['docx', 'pdf', 'html'],
                key=f"{key_prefix}_toc",
                disabled=format not in ['docx', 'pdf', 'html']
            )
        
        with col2:
            config.watermark = st.text_input(
                "Filigrane",
                placeholder="Confidentiel",
                key=f"{key_prefix}_watermark",
                disabled=format not in ['docx', 'pdf']
            )
            
            config.compression = st.checkbox(
                "Compresser (ZIP)",
                value=False,
                key=f"{key_prefix}_compress"
            )
        
        # Options spécifiques au format
        if format == 'docx':
            st.markdown("**Options Word**")
            col1, col2 = st.columns(2)
            
            with col1:
                config.style_options['font'] = st.selectbox(
                    "Police",
                    ["Times New Roman", "Arial", "Calibri", "Garamond"],
                    key=f"{key_prefix}_font"
                )
                
                config.style_options['font_size'] = st.number_input(
                    "Taille (pt)",
                    8, 20, 12,
                    key=f"{key_prefix}_font_size"
                )
            
            with col2:
                config.style_options['line_spacing'] = st.selectbox(
                    "Interligne",
                    ["1.0", "1.15", "1.5", "2.0"],
                    index=2,
                    key=f"{key_prefix}_spacing"
                )
                
                config.style_options['paper_size'] = st.selectbox(
                    "Format papier",
                    ["A4", "Letter", "Legal"],
                    key=f"{key_prefix}_paper"
                )
        
        elif format == 'xlsx':
            st.markdown("**Options Excel**")
            config.style_options['include_charts'] = st.checkbox(
                "Générer des graphiques",
                value=False,
                key=f"{key_prefix}_charts"
            )
            
            config.style_options['conditional_formatting'] = st.checkbox(
                "Mise en forme conditionnelle",
                value=False,
                key=f"{key_prefix}_conditional"
            )
        
        return config
    
    def _get_format_display_name(self, format: str) -> str:
        """Nom d'affichage pour chaque format"""
        names = {
            'docx': '📄 Word (.docx)',
            'pdf': '📕 PDF',
            'xlsx': '📊 Excel (.xlsx)',
            'txt': '📝 Texte (.txt)',
            'json': '🔧 JSON',
            'html': '🌐 HTML',
            'csv': '📋 CSV',
            'rtf': '📃 RTF'
        }
        return names.get(format, format.upper())
    
    def _get_mime_type(self, format: str) -> str:
        """Type MIME pour chaque format"""
        mime_types = {
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'pdf': 'application/pdf',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'txt': 'text/plain',
            'json': 'application/json',
            'html': 'text/html',
            'csv': 'text/csv',
            'rtf': 'application/rtf'
        }
        return mime_types.get(format, 'application/octet-stream')
    
    def _generate_filename(self, title: str, format: str) -> str:
        """Génère un nom de fichier intelligent"""
        # Nettoyer le titre
        clean_title = re.sub(r'[^\w\s-]', '', title)
        clean_title = re.sub(r'[-\s]+', '_', clean_title)
        
        # Date
        date_str = datetime.now().strftime('%Y%m%d')
        
        return f"{clean_title}_{date_str}".lower()[:50]
    
    def _show_export_stats(self, data: bytes, format: str):
        """Affiche les statistiques de l'export"""
        size = len(data)
        
        if size < 1024:
            size_str = f"{size} octets"
        elif size < 1024 * 1024:
            size_str = f"{size / 1024:.1f} Ko"
        else:
            size_str = f"{size / (1024 * 1024):.1f} Mo"
        
        st.info(f"📊 Taille du fichier : {size_str}")
    
    def _compress_content(self, content: bytes, config: ExportConfig) -> bytes:
        """Compresse le contenu en ZIP"""
        import zipfile
        
        buffer = io.BytesIO()
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            filename = f"{config.title}.{self.exporters[config.format].get_extension()}"
            zip_file.writestr(filename, content)
            
            # Ajouter les métadonnées
            if config.metadata:
                metadata_str = json.dumps(config.metadata, indent=2, default=str)
                zip_file.writestr('metadata.json', metadata_str.encode('utf-8'))
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def _encrypt_content(self, content: bytes, config: ExportConfig) -> bytes:
        """Chiffre le contenu (placeholder - implémenter selon besoins)"""
        # Pour une vraie implémentation, utiliser cryptography ou similaire
        return content

# ========================= EXPORTERS SPÉCIFIQUES =========================

class BaseExporter:
    """Classe de base pour tous les exporters"""
    
    def is_available(self) -> bool:
        """Vérifie si l'exporter est disponible"""
        return True
    
    def get_extension(self) -> str:
        """Retourne l'extension du fichier"""
        return 'txt'
    
    def export(self, config: ExportConfig) -> bytes:
        """Méthode d'export à implémenter"""
        raise NotImplementedError

class WordExporter(BaseExporter):
    """Exporter pour documents Word"""
    
    def is_available(self) -> bool:
        return DOCX_AVAILABLE
    
    def get_extension(self) -> str:
        return 'docx'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers Word avec mise en forme avancée"""
        
        doc = docx.Document()
        
        # Configuration du document
        self._setup_document(doc, config)
        
        # Ajouter le contenu selon le type
        if isinstance(config.content, str):
            self._add_text_content(doc, config)
        elif isinstance(config.content, pd.DataFrame):
            self._add_dataframe_content(doc, config)
        elif isinstance(config.content, list):
            self._add_list_content(doc, config)
        elif isinstance(config.content, dict):
            self._add_dict_content(doc, config)
        
        # Sauvegarder
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _setup_document(self, doc, config: ExportConfig):
        """Configure le document Word"""
        
        # Marges
        sections = doc.sections
        for section in sections:
            margins = {
                'Normales': (1, 1, 1, 1),
                'Étroites': (0.5, 0.5, 0.5, 0.5),
                'Larges': (1.5, 1.5, 1.5, 1.5)
            }
            
            margin_values = margins.get(
                config.style_options.get('margins', 'Normales'),
                (1, 1, 1, 1)
            )
            
            section.top_margin = Inches(margin_values[0])
            section.bottom_margin = Inches(margin_values[1])
            section.left_margin = Inches(margin_values[2])
            section.right_margin = Inches(margin_values[3])
        
        # Styles
        self._setup_styles(doc, config)
    
    def _setup_styles(self, doc, config: ExportConfig):
        """Configure les styles du document"""
        
        # Style normal
        normal_style = doc.styles['Normal']
        normal_style.font.name = config.style_options.get('font', 'Times New Roman')
        normal_style.font.size = Pt(config.style_options.get('font_size', 12))
        
        # Interligne
        spacing_map = {
            '1.0': WD_LINE_SPACING.SINGLE,
            '1.15': WD_LINE_SPACING.ONE_POINT_FIVE,
            '1.5': WD_LINE_SPACING.ONE_POINT_FIVE,
            '2.0': WD_LINE_SPACING.DOUBLE
        }
        
        spacing = spacing_map.get(
            config.style_options.get('line_spacing', '1.5'),
            WD_LINE_SPACING.ONE_POINT_FIVE
        )
        normal_style.paragraph_format.line_spacing_rule = spacing
        
        # Styles juridiques si demandés
        if config.style_options.get('legal_formatting', True):
            self._create_legal_styles(doc, config)
    
    def _create_legal_styles(self, doc, config: ExportConfig):
        """Crée les styles juridiques"""
        
        styles = doc.styles
        
        # Style pour sections principales
        if 'JuridiqueSection' not in [s.name for s in styles]:
            section_style = styles.add_style('JuridiqueSection', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = config.style_options.get('font', 'Times New Roman')
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(12)
    
    def _add_text_content(self, doc, config: ExportConfig):
        """Ajoute du contenu texte"""
        
        # En-tête
        if config.include_header_footer:
            self._add_header(doc, config)
        
        # Titre
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(config.title.upper())
        run.font.size = Pt(16)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Métadonnées
        if config.include_metadata and config.metadata:
            self._add_metadata_section(doc, config)
        
        # Table des matières
        if config.include_toc:
            self._add_toc(doc, config)
        
        # Contenu principal
        self._process_text_content(doc, config.content, config)
        
        # Pied de page
        if config.include_header_footer:
            self._add_footer(doc, config)
    
    def _process_text_content(self, doc, content: str, config: ExportConfig):
        """Traite et ajoute le contenu texte avec détection de structure"""
        
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                doc.add_paragraph()
                continue
            
            # Détection de structure
            if re.match(r'^[IVX]+\.\s+', line):
                # Section principale
                para = doc.add_paragraph(line, style='JuridiqueSection' if 'JuridiqueSection' in [s.name for s in doc.styles] else 'Heading 1')
            elif re.match(r'^[A-Z]\.\s+', line):
                # Sous-section
                para = doc.add_paragraph(line, style='Heading 2')
                para.runs[0].font.bold = True
            elif re.match(r'^\d+\.\s+', line):
                # Liste numérotée
                para = doc.add_paragraph(line, style='List Number')
            elif line.strip().isupper() and len(line.strip()) < 100:
                # Titre en majuscules
                para = doc.add_paragraph(line.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].font.bold = True
            else:
                # Paragraphe normal
                para = doc.add_paragraph(line.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_dataframe_content(self, doc, config: ExportConfig):
        """Ajoute un DataFrame comme tableau"""
        
        df = config.content
        
        # Titre
        doc.add_heading(config.title, 1)
        
        # Créer le tableau
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
        table.style = 'Light List'
        
        # En-têtes
        for i, col in enumerate(df.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col)
            # Mettre en gras
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Données
        for row_idx, row in df.iterrows():
            for col_idx, value in enumerate(row):
                table.rows[row_idx + 1].cells[col_idx].text = str(value)
    
    def _add_header(self, doc, config: ExportConfig):
        """Ajoute un en-tête au document"""
        section = doc.sections[0]
        header = section.header
        
        # Texte de l'en-tête
        header_para = header.paragraphs[0]
        header_para.text = f"{config.title} - {datetime.now().strftime('%d/%m/%Y')}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _add_footer(self, doc, config: ExportConfig):
        """Ajoute un pied de page avec numérotation"""
        section = doc.sections[0]
        footer = section.footer
        
        # Numéro de page
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter le champ de numéro de page (complexe en python-docx)
        footer_para.text = "Page "

class PDFExporter(BaseExporter):
    """Exporter pour PDF"""
    
    def is_available(self) -> bool:
        return PDF_AVAILABLE
    
    def get_extension(self) -> str:
        return 'pdf'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers PDF"""
        
        buffer = io.BytesIO()
        
        # Configuration des marges
        margins = config.style_options.get('margins', 2.0)
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=margins * cm,
            leftMargin=margins * cm,
            topMargin=margins * cm,
            bottomMargin=margins * cm
        )
        
        # Styles
        styles = self._create_styles(config)
        
        # Construire le contenu
        story = []
        
        # Titre
        story.append(Paragraph(config.title, styles['Title']))
        story.append(Spacer(1, 12))
        
        # Métadonnées
        if config.include_metadata and config.metadata:
            story.extend(self._add_metadata_pdf(config, styles))
        
        # Contenu principal
        if isinstance(config.content, str):
            story.extend(self._process_text_pdf(config.content, styles))
        elif isinstance(config.content, pd.DataFrame):
            story.extend(self._add_dataframe_pdf(config.content, styles))
        
        # Générer le PDF
        doc.build(story)
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def _create_styles(self, config: ExportConfig) -> dict:
        """Crée les styles pour le PDF"""
        
        styles = getSampleStyleSheet()
        
        # Style personnalisé
        styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=styles['Normal'],
            fontName=config.style_options.get('font', 'Times-Roman'),
            fontSize=config.style_options.get('font_size', 12),
            leading=config.style_options.get('font_size', 12) * 1.5,
            alignment=4  # Justifié
        ))
        
        return styles
    
    def _process_text_pdf(self, text: str, styles: dict) -> List:
        """Convertit le texte en éléments PDF"""
        
        elements = []
        
        for line in text.split('\n'):
            if line.strip():
                elements.append(Paragraph(line, styles['CustomNormal']))
            else:
                elements.append(Spacer(1, 12))
        
        return elements

class ExcelExporter(BaseExporter):
    """Exporter pour Excel"""
    
    def is_available(self) -> bool:
        return PANDAS_AVAILABLE
    
    def get_extension(self) -> str:
        return 'xlsx'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers Excel"""
        
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            
            # Export selon le type de contenu
            if isinstance(config.content, pd.DataFrame):
                self._export_dataframe(writer, config)
            elif isinstance(config.content, list):
                self._export_list(writer, config)
            elif isinstance(config.content, dict):
                self._export_dict(writer, config)
            else:
                # Conversion en DataFrame simple
                df = pd.DataFrame({'Contenu': [str(config.content)]})
                df.to_excel(writer, sheet_name='Export', index=False)
            
            # Ajouter les métadonnées
            if config.include_metadata and config.metadata:
                self._add_metadata_sheet(writer, config)
            
            # Formatage
            self._apply_formatting(writer, config)
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def _export_dataframe(self, writer, config: ExportConfig):
        """Exporte un DataFrame"""
        
        df = config.content
        sheet_name = 'Données'
        
        df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        # Options de formatage
        if config.style_options.get('freeze_headers', True):
            writer.sheets[sheet_name].freeze_panes = 'A2'
        
        if config.style_options.get('autofilter', True):
            writer.sheets[sheet_name].auto_filter.ref = writer.sheets[sheet_name].dimensions
    
    def _apply_formatting(self, writer, config: ExportConfig):
        """Applique le formatage Excel"""
        
        if config.style_options.get('column_autowidth', True):
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width

class TextExporter(BaseExporter):
    """Exporter pour texte simple"""
    
    def get_extension(self) -> str:
        return 'txt'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers texte simple"""
        
        lines = []
        
        # En-tête
        lines.append("=" * 80)
        lines.append(config.title.upper().center(80))
        lines.append("=" * 80)
        lines.append("")
        
        # Métadonnées
        if config.include_metadata and config.metadata:
            lines.append("INFORMATIONS")
            lines.append("-" * 40)
            
            for key, value in config.metadata.items():
                if not isinstance(value, (dict, list)):
                    lines.append(f"{key}: {value}")
            
            lines.append("")
        
        # Contenu
        lines.append("CONTENU")
        lines.append("-" * 40)
        lines.append("")
        
        if isinstance(config.content, str):
            lines.append(config.content)
        else:
            lines.append(str(config.content))
        
        # Pied
        lines.append("")
        lines.append("=" * 80)
        lines.append(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}")
        
        return '\n'.join(lines).encode('utf-8')

class JSONExporter(BaseExporter):
    """Exporter pour JSON"""
    
    def get_extension(self) -> str:
        return 'json'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers JSON"""
        
        data = {
            'title': config.title,
            'metadata': config.metadata or {},
            'content': None,
            'export_date': datetime.now().isoformat()
        }
        
        # Conversion du contenu
        if isinstance(config.content, pd.DataFrame):
            data['content'] = config.content.to_dict('records')
        elif isinstance(config.content, (dict, list)):
            data['content'] = config.content
        else:
            data['content'] = str(config.content)
        
        return json.dumps(data, indent=2, ensure_ascii=False, default=str).encode('utf-8')

class HTMLExporter(BaseExporter):
    """Exporter pour HTML"""
    
    def get_extension(self) -> str:
        return 'html'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers HTML"""
        
        # Template HTML
        html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{config.title}</title>
"""
        
        # CSS
        if config.style_options.get('include_css', True):
            html += self._get_css(config)
        
        html += """</head>
<body>
    <div class="container">
"""
        
        # Titre
        html += f"<h1>{config.title}</h1>\n"
        
        # Métadonnées
        if config.include_metadata and config.metadata:
            html += self._add_metadata_html(config)
        
        # Contenu
        html += '<div class="content">\n'
        
        if isinstance(config.content, str):
            html += self._process_text_html(config.content)
        elif isinstance(config.content, pd.DataFrame):
            html += config.content.to_html(classes='table table-striped', index=False)
        else:
            html += f"<pre>{json.dumps(config.content, indent=2, default=str)}</pre>"
        
        html += """
        </div>
    </div>
</body>
</html>"""
        
        return html.encode('utf-8')
    
    def _get_css(self, config: ExportConfig) -> str:
        """Retourne le CSS selon le thème"""
        
        theme = config.style_options.get('theme', 'Juridique')
        
        return """
    <style>
        body {
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }
        
        h1 {
            color: #8b0000;
            text-align: center;
            border-bottom: 3px solid #8b0000;
            padding-bottom: 10px;
        }
        
        .metadata {
            background: #f5f5f5;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }
        
        .content {
            text-align: justify;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #8b0000;
            color: white;
        }
        
        @media print {
            body { margin: 0; }
            .no-print { display: none; }
        }
    </style>
"""

class CSVExporter(BaseExporter):
    """Exporter pour CSV"""
    
    def is_available(self) -> bool:
        return PANDAS_AVAILABLE
    
    def get_extension(self) -> str:
        return 'csv'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers CSV"""
        
        if isinstance(config.content, pd.DataFrame):
            return config.content.to_csv(index=False).encode('utf-8')
        elif isinstance(config.content, list):
            df = pd.DataFrame(config.content)
            return df.to_csv(index=False).encode('utf-8')
        else:
            # Fallback
            return str(config.content).encode('utf-8')

class RTFExporter(BaseExporter):
    """Exporter pour RTF"""
    
    def get_extension(self) -> str:
        return 'rtf'
    
    def export(self, config: ExportConfig) -> bytes:
        """Export vers RTF"""
        
        rtf = r'{\rtf1\ansi\deff0 {\fonttbl{\f0 Times New Roman;}}' + '\n'
        rtf += r'\f0\fs24 '  # Police 12pt
        
        # Titre
        rtf += r'{\b\fs32 ' + config.title + r'\par}' + '\n'
        rtf += r'\par' + '\n'
        
        # Contenu
        if isinstance(config.content, str):
            lines = config.content.split('\n')
            for line in lines:
                if line.strip():
                    # Échapper les caractères spéciaux
                    line = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                    rtf += line + r'\par '
                else:
                    rtf += r'\par '
        
        rtf += '}'
        
        return rtf.encode('utf-8')

# ========================= INSTANCE GLOBALE =========================

# Créer une instance globale du gestionnaire
export_manager = ExportManager()

# ========================= FONCTIONS UTILITAIRES =========================

def quick_export(content: Any, 
                title: str = "Export",
                format: str = "docx",
                **kwargs) -> Tuple[bytes, str]:
    """
    Fonction utilitaire pour un export rapide
    
    Example:
        data, ext = quick_export(df, "Rapport", "xlsx")
    """
    
    config = ExportConfig(
        format=format,
        title=title,
        content=content,
        **kwargs
    )
    
    return export_manager.export(config)

def show_export_button(content: Any,
                      title: str,
                      key: str = "export",
                      formats: List[str] = None) -> None:
    """
    Affiche un bouton d'export simple
    
    Example:
        show_export_button(document_content, "Mon Document", key="doc1")
    """
    
    if formats is None:
        formats = export_manager.get_available_formats()
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        selected_format = st.selectbox(
            "Format",
            formats,
            format_func=lambda x: export_manager._get_format_display_name(x),
            key=f"{key}_format_simple"
        )
    
    with col2:
        if st.button("📥 Exporter", key=f"{key}_btn_simple"):
            try:
                data, ext = quick_export(content, title, selected_format)
                
                st.download_button(
                    f"⬇️ Télécharger",
                    data,
                    f"{title}_{datetime.now().strftime('%Y%m%d')}.{ext}",
                    mime=export_manager._get_mime_type(selected_format),
                    key=f"{key}_download_simple"
                )
            except Exception as e:
                st.error(f"Erreur : {str(e)}")