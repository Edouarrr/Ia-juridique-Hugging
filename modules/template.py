"""Module de gestion des templates de documents juridiques - Version améliorée avec multi-IA"""

import json
import os
import re
import sys
import time
from collections import defaultdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import plotly.express as px
import plotly.graph_objects as go
import streamlit as st

# Ajouter le chemin parent pour importer utils
sys.path.append(str(Path(__file__).parent.parent))
from utils import clean_key, format_legal_date, truncate_text

# Icons pour les catégories
CATEGORY_ICONS = {
    'Procédure civile': '⚖️',
    'Pénal': '🚨',
    'Commercial': '💼',
    'Social': '👥',
    'Administratif': '🏛️',
    'Famille': '👨‍👩‍👧‍👦',
    'Immobilier': '🏠',
    'Propriété intellectuelle': '💡',
    'Fiscal': '💰',
    'International': '🌍',
    'Autre': '📄'
}

# Styles de templates
TEMPLATE_STYLES = {
    'formel': {'icon': '🎩', 'description': 'Style traditionnel et protocolaire'},
    'persuasif': {'icon': '💪', 'description': 'Argumentation convaincante'},
    'technique': {'icon': '🔬', 'description': 'Précision juridique maximale'},
    'synthétique': {'icon': '📊', 'description': 'Concis et structuré'},
    'pédagogique': {'icon': '🎓', 'description': 'Clair et explicatif'},
    'moderne': {'icon': '✨', 'description': 'Approche contemporaine'},
    'agressif': {'icon': '⚔️', 'description': 'Défense offensive'}
}

# Templates prédéfinis étendus
BUILTIN_TEMPLATES = {
    'conclusions_defense': {
        'id': 'conclusions_defense',
        'name': 'Conclusions en défense',
        'type': 'conclusions',
        'structure': [
            'POUR : [Partie défenderesse]',
            '',
            'CONTRE : [Partie demanderesse]',
            '',
            'PLAISE AU TRIBUNAL',
            '',
            'I. FAITS ET PROCÉDURE',
            '   A. Rappel des faits',
            '   B. Procédure',
            '',
            'II. DISCUSSION',
            '   A. Sur la recevabilité',
            '      1. [Point de recevabilité]',
            '   B. Sur le fond',
            '      1. Sur [premier moyen]',
            '      2. Sur [deuxième moyen]',
            '',
            'III. SUR LES DEMANDES',
            '   A. Sur les demandes principales',
            '   B. Sur les demandes accessoires',
            '',
            'PAR CES MOTIFS',
            '',
            'REJETER l\'ensemble des demandes',
            'CONDAMNER [partie adverse] aux entiers dépens'
        ],
        'style': 'formel',
        'category': 'Procédure civile',
        'description': 'Template standard pour des conclusions en défense',
        'tags': ['défense', 'civil', 'tribunal'],
        'difficulty': 'intermédiaire',
        'estimated_time': '30-45 min',
        'is_builtin': True
    },
    
    'plainte_simple': {
        'id': 'plainte_simple',
        'name': 'Plainte simple',
        'type': 'plainte',
        'structure': [
            'Madame, Monsieur le Procureur de la République',
            'Tribunal judiciaire de [Ville]',
            '',
            'PLAINTE',
            '',
            'Je soussigné(e) [Nom, Prénom]',
            'Né(e) le [Date] à [Lieu]',
            'Demeurant [Adresse]',
            'Profession : [Profession]',
            '',
            'Ai l\'honneur de porter plainte contre :',
            '[Identité du mis en cause ou X]',
            '',
            'POUR LES FAITS SUIVANTS :',
            '',
            '[Exposé détaillé des faits]',
            '',
            'Ces faits sont susceptibles de constituer :',
            '- [Qualification juridique 1]',
            '- [Qualification juridique 2]',
            '',
            'PRÉJUDICES SUBIS :',
            '[Description des préjudices]',
            '',
            'PIÈCES JOINTES :',
            '1. [Pièce 1]',
            '2. [Pièce 2]',
            '',
            'Je me tiens à votre disposition pour tout complément.',
            '',
            'Fait à [Ville], le [Date]',
            'Signature'
        ],
        'style': 'formel',
        'category': 'Pénal',
        'description': 'Template pour une plainte simple au Procureur',
        'tags': ['plainte', 'pénal', 'procureur'],
        'difficulty': 'facile',
        'estimated_time': '15-20 min',
        'is_builtin': True
    },
    
    'mise_en_demeure': {
        'id': 'mise_en_demeure',
        'name': 'Mise en demeure',
        'type': 'courrier',
        'structure': [
            '[Expéditeur]',
            '[Adresse expéditeur]',
            '',
            '[Destinataire]',
            '[Adresse destinataire]',
            '',
            'Lettre recommandée avec AR',
            'N° [Numéro AR]',
            '',
            'Objet : MISE EN DEMEURE',
            '',
            '[Ville], le [Date]',
            '',
            'Madame, Monsieur,',
            '',
            'Par [contrat/accord] en date du [Date contrat], vous vous êtes engagé(e) à [obligation].',
            '',
            'Or, à ce jour, [constat du manquement].',
            '',
            'Cette situation me cause un préjudice [description du préjudice].',
            '',
            'En conséquence, je vous mets en demeure de :',
            '- [Action demandée 1]',
            '- [Action demandée 2]',
            '',
            'Et ce, dans un délai de [Nombre] jours à compter de la réception de la présente.',
            '',
            'À défaut, je me réserve le droit de :',
            '- Saisir la justice',
            '- Réclamer des dommages et intérêts',
            '- Engager toute procédure utile',
            '',
            'Dans l\'attente de votre réponse, je vous prie d\'agréer, Madame, Monsieur, l\'expression de mes salutations distinguées.',
            '',
            '[Signature]'
        ],
        'style': 'formel',
        'category': 'Commercial',
        'description': 'Mise en demeure type pour obligation non respectée',
        'tags': ['mise en demeure', 'commercial', 'pré-contentieux'],
        'difficulty': 'facile',
        'estimated_time': '10-15 min',
        'is_builtin': True
    }
}

def run():
    """Fonction principale du module avec design amélioré"""
    
    # Configuration de la page
    st.markdown("""
        <style>
        /* Animations */
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        .template-card {
            animation: fadeIn 0.3s ease-out;
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            border-radius: 15px;
            padding: 20px;
            margin: 10px 0;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            transition: all 0.3s ease;
        }
        
        .template-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 8px 15px rgba(0,0,0,0.2);
        }
        
        .metric-card {
            background: white;
            border-radius: 10px;
            padding: 15px;
            text-align: center;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        
        .tag-pill {
            display: inline-block;
            padding: 4px 12px;
            margin: 2px;
            background: #e1e8f0;
            border-radius: 20px;
            font-size: 0.85em;
            color: #4a5568;
        }
        
        .ai-badge {
            display: inline-flex;
            align-items: center;
            gap: 5px;
            padding: 5px 15px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 20px;
            font-weight: 500;
        }
        
        .preview-box {
            background: #f8f9fa;
            border: 2px dashed #dee2e6;
            border-radius: 10px;
            padding: 20px;
            font-family: 'Courier New', monospace;
        }
        
        /* Variables highlighting */
        .variable-highlight {
            background: #fff3cd;
            padding: 2px 6px;
            border-radius: 4px;
            font-weight: bold;
            color: #856404;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Header avec animation
    col1, col2, col3 = st.columns([2, 1, 1])
    
    with col1:
        st.markdown("# 📄 Templates Juridiques")
        st.markdown("*Générateur intelligent de documents avec IA*")
    
    with col2:
        # Statistiques temps réel
        stats = get_template_stats()
        st.metric(
            "Templates actifs",
            stats['total'],
            f"+{stats['recent']} cette semaine"
        )
    
    with col3:
        # Mode IA
        ai_mode = st.selectbox(
            "🤖 Mode IA",
            ["🎯 Optimisé", "🔀 Fusion Multi-IA", "⚡ Rapide", "🔬 Précision Max"],
            key="template_ai_mode"
        )
    
    # Initialisation état
    init_template_state()
    
    # Navigation intelligente
    if st.session_state.template_state.get('quick_action'):
        handle_quick_action()
    else:
        show_main_interface()

def init_template_state():
    """Initialise l'état du module"""
    if 'template_state' not in st.session_state:
        st.session_state.template_state = {
            'view_mode': 'gallery',  # gallery, list, kanban
            'selected_template': None,
            'filter_tags': [],
            'ai_providers': {},
            'generation_history': [],
            'favorites': [],
            'recent_used': [],
            'quick_action': None
        }

def get_template_stats() -> Dict[str, int]:
    """Calcule les statistiques des templates"""
    all_templates = get_all_templates()
    
    stats = {
        'total': len(all_templates),
        'recent': 0,
        'custom': 0,
        'most_used': None
    }
    
    week_ago = datetime.now() - timedelta(days=7)
    
    for template in all_templates.values():
        if hasattr(template, 'created_at') and template.created_at > week_ago:
            stats['recent'] += 1
        
        if not template.get('is_builtin'):
            stats['custom'] += 1
    
    return stats

def show_main_interface():
    """Interface principale avec navigation fluide"""
    
    # Barre de recherche intelligente
    search_col1, search_col2, search_col3 = st.columns([3, 1, 1])
    
    with search_col1:
        search_query = st.text_input(
            "🔍 Recherche intelligente",
            placeholder="Ex: conclusions défense, plainte vol, mise en demeure...",
            key="template_smart_search"
        )
        
        # Suggestions en temps réel
        if search_query:
            suggestions = get_search_suggestions(search_query)
            if suggestions:
                selected = st.selectbox(
                    "Suggestions",
                    suggestions,
                    key="search_suggestions"
                )
                if st.button("Utiliser", key="use_suggestion"):
                    st.session_state.template_state['quick_action'] = ('apply', selected)
                    st.rerun()
    
    with search_col2:
        view_mode = st.selectbox(
            "📊 Vue",
            ["🎨 Galerie", "📋 Liste", "🗂️ Kanban"],
            key="template_view_mode"
        )
        st.session_state.template_state['view_mode'] = view_mode.split()[1].lower()
    
    with search_col3:
        if st.button("➕ Créer", use_container_width=True, type="primary"):
            st.session_state.template_state['quick_action'] = ('create', None)
            st.rerun()
    
    # Filtres avancés
    with st.expander("🎯 Filtres avancés", expanded=False):
        show_advanced_filters()
    
    # Quick actions
    show_quick_actions()
    
    # Affichage principal
    if st.session_state.template_state['view_mode'] == 'galerie':
        show_gallery_view(search_query)
    elif st.session_state.template_state['view_mode'] == 'liste':
        show_list_view(search_query)
    else:
        show_kanban_view(search_query)
    
    # Analytics
    if st.sidebar.checkbox("📊 Analytics", key="show_template_analytics"):
        show_template_analytics()

def show_quick_actions():
    """Actions rapides basées sur le contexte"""
    
    st.markdown("### ⚡ Actions rapides")
    
    cols = st.columns(5)
    
    quick_actions = [
        ("📝 Conclusions", "conclusions_defense", "⚖️"),
        ("🚨 Plainte", "plainte_simple", "🚨"),
        ("📨 Mise en demeure", "mise_en_demeure", "💼"),
        ("📄 Requête", "requete", "⚖️"),
        ("🎯 Sur mesure", None, "✨")
    ]
    
    for idx, (label, template_id, icon) in enumerate(quick_actions):
        with cols[idx]:
            if st.button(
                f"{icon}\n{label}",
                key=f"quick_{template_id or 'custom'}",
                use_container_width=True
            ):
                if template_id:
                    st.session_state.template_state['selected_template'] = template_id
                    st.session_state.template_state['quick_action'] = ('apply', template_id)
                else:
                    st.session_state.template_state['quick_action'] = ('create', None)
                st.rerun()

def show_gallery_view(search_query: str):
    """Vue galerie des templates avec design moderne"""
    
    templates = get_filtered_templates(search_query)
    
    if not templates:
        st.info("🔍 Aucun template trouvé. Essayez d'autres critères ou créez-en un nouveau !")
        return
    
    # Grouper par catégorie
    categories = defaultdict(list)
    for tid, template in templates.items():
        categories[template.get('category', 'Autre')].append((tid, template))
    
    # Afficher par catégorie
    for category, items in sorted(categories.items()):
        icon = CATEGORY_ICONS.get(category, '📄')
        st.markdown(f"### {icon} {category}")
        
        # Grille responsive
        cols = st.columns(3)
        
        for idx, (tid, template) in enumerate(items):
            with cols[idx % 3]:
                show_template_card_modern(tid, template)

def show_template_card_modern(template_id: str, template: Dict):
    """Carte template avec design moderne"""
    
    with st.container():
        # Card container avec style
        card_html = f"""
        <div class="template-card">
            <div style="display: flex; justify-content: space-between; align-items: start;">
                <div>
                    <h4 style="margin: 0; color: #2d3748;">
                        {TEMPLATE_STYLES.get(template.get('style', 'formel'), {}).get('icon', '📄')} 
                        {template['name']}
                    </h4>
                    <p style="color: #718096; font-size: 0.9em; margin: 5px 0;">
                        {template.get('description', 'Template juridique')}
                    </p>
                </div>
                <div class="ai-badge">
                    <span>✨ IA</span>
                </div>
            </div>
            
            <div style="margin: 15px 0;">
                <span class="tag-pill">📁 {template.get('type', 'document')}</span>
                <span class="tag-pill">⏱️ {template.get('estimated_time', '20 min')}</span>
                <span class="tag-pill">📊 {template.get('difficulty', 'moyen')}</span>
            </div>
            
            <div style="margin-top: 10px;">
        """
        
        # Tags
        if template.get('tags'):
            for tag in template['tags'][:3]:
                card_html += f'<span class="tag-pill">#{tag}</span>'
        
        card_html += "</div></div>"
        
        st.markdown(card_html, unsafe_allow_html=True)
        
        # Actions
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("👁️", key=f"preview_{template_id}", help="Aperçu"):
                show_template_preview_modern(template)
        
        with col2:
            if st.button("⚡", key=f"quick_use_{template_id}", help="Utilisation rapide"):
                quick_apply_template(template_id, template)
        
        with col3:
            is_favorite = template_id in st.session_state.template_state.get('favorites', [])
            if st.button(
                "⭐" if is_favorite else "☆",
                key=f"fav_{template_id}",
                help="Retirer des favoris" if is_favorite else "Ajouter aux favoris"
            ):
                toggle_favorite(template_id)

def show_template_preview_modern(template: Dict):
    """Aperçu moderne du template avec mise en forme"""
    
    with st.expander("👁️ Aperçu du template", expanded=True):
        # Header info
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Type", template.get('type', 'document'))
        with col2:
            st.metric("Style", TEMPLATE_STYLES.get(template.get('style', 'formel'), {}).get('description', ''))
        with col3:
            st.metric("Variables", len(extract_variables(template.get('structure', []))))
        
        # Structure avec highlighting
        st.markdown("#### 📋 Structure")
        
        preview_html = '<div class="preview-box">'
        
        for line in template.get('structure', []):
            # Highlight variables
            highlighted = re.sub(
                r'\[([^\]]+)\]',
                r'<span class="variable-highlight">[\1]</span>',
                line
            )
            preview_html += f"<div>{highlighted}</div>"
        
        preview_html += "</div>"
        
        st.markdown(preview_html, unsafe_allow_html=True)
        
        # Variables détectées
        variables = extract_variables(template.get('structure', []))
        if variables:
            st.markdown("#### 🔤 Variables détectées")
            var_cols = st.columns(4)
            for idx, var in enumerate(variables):
                with var_cols[idx % 4]:
                    st.info(f"[{var}]")

def quick_apply_template(template_id: str, template: Dict):
    """Application rapide d'un template avec IA"""
    
    st.session_state.template_state['selected_template'] = template_id
    st.session_state.template_state['quick_action'] = ('apply', template_id)
    st.rerun()

def handle_quick_action():
    """Gère les actions rapides"""
    
    action, data = st.session_state.template_state['quick_action']
    
    if action == 'create':
        show_create_wizard()
    elif action == 'apply':
        show_apply_wizard(data)
    
    # Bouton retour
    if st.button("⬅️ Retour", key="back_from_quick"):
        st.session_state.template_state['quick_action'] = None
        st.rerun()

def show_create_wizard():
    """Assistant de création de template avec IA"""
    
    st.markdown("## 🎨 Créer un nouveau template")
    
    # Étape 1: Type et inspiration
    st.markdown("### 1️⃣ Type de document")
    
    col1, col2 = st.columns(2)
    
    with col1:
        doc_type = st.selectbox(
            "Type de document",
            ["conclusions", "plainte", "assignation", "requête", "courrier", "contrat", "autre"],
            key="wizard_doc_type"
        )
        
        category = st.selectbox(
            "Catégorie",
            list(CATEGORY_ICONS.keys()),
            key="wizard_category"
        )
    
    with col2:
        style = st.selectbox(
            "Style souhaité",
            list(TEMPLATE_STYLES.keys()),
            format_func=lambda x: f"{TEMPLATE_STYLES[x]['icon']} {x.capitalize()} - {TEMPLATE_STYLES[x]['description']}",
            key="wizard_style"
        )
        
        ai_help = st.checkbox(
            "🤖 Assistance IA pour la création",
            value=True,
            key="wizard_ai_help"
        )
    
    # Étape 2: Contenu
    st.markdown("### 2️⃣ Contenu du template")
    
    if ai_help:
        # Génération assistée par IA
        st.info("💡 L'IA va vous aider à créer la structure optimale")
        
        context = st.text_area(
            "Décrivez votre besoin",
            placeholder="Ex: J'ai besoin d'un template pour des conclusions en défense dans un litige commercial, avec une section sur la prescription et les dommages-intérêts...",
            height=100,
            key="wizard_context"
        )
        
        if context and st.button("🎯 Générer la structure avec l'IA", type="primary"):
            with st.spinner("🤖 Génération en cours..."):
                structure = generate_template_structure_with_ai(
                    doc_type, category, style, context
                )
                st.session_state.wizard_generated_structure = structure
    
    # Structure (générée ou manuelle)
    initial_structure = st.session_state.get('wizard_generated_structure', '')
    
    structure = st.text_area(
        "Structure du template",
        value=initial_structure,
        height=400,
        placeholder="Utilisez [NomVariable] pour créer des variables...",
        key="wizard_structure"
    )
    
    # Étape 3: Métadonnées
    st.markdown("### 3️⃣ Informations")
    
    col1, col2 = st.columns(2)
    
    with col1:
        name = st.text_input(
            "Nom du template",
            placeholder="Ex: Conclusions défense commerciale",
            key="wizard_name"
        )
        
        estimated_time = st.select_slider(
            "Temps estimé",
            options=["5-10 min", "10-20 min", "20-30 min", "30-45 min", "45-60 min", "+1h"],
            value="20-30 min",
            key="wizard_time"
        )
    
    with col2:
        difficulty = st.select_slider(
            "Difficulté",
            options=["Débutant", "Facile", "Intermédiaire", "Avancé", "Expert"],
            value="Intermédiaire",
            key="wizard_difficulty"
        )
        
        tags = st.text_input(
            "Tags (séparés par des virgules)",
            placeholder="Ex: commercial, défense, prescription",
            key="wizard_tags"
        )
    
    description = st.text_area(
        "Description",
        placeholder="Décrivez l'usage et les particularités de ce template...",
        height=100,
        key="wizard_description"
    )
    
    # Actions finales
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("👁️ Prévisualiser", type="secondary"):
            if name and structure:
                preview_template = {
                    'name': name,
                    'type': doc_type,
                    'structure': structure.split('\n'),
                    'style': style,
                    'category': category,
                    'description': description,
                    'tags': [t.strip() for t in tags.split(',')] if tags else [],
                    'difficulty': difficulty,
                    'estimated_time': estimated_time
                }
                show_template_preview_modern(preview_template)
    
    with col2:
        if st.button("💾 Sauvegarder", type="primary"):
            if not name or not structure:
                st.error("❌ Nom et structure requis")
            else:
                save_new_template(
                    name, doc_type, structure, style, category,
                    description, tags, difficulty, estimated_time
                )
    
    with col3:
        if st.button("🎯 Sauvegarder et utiliser"):
            if name and structure:
                template_id = save_new_template(
                    name, doc_type, structure, style, category,
                    description, tags, difficulty, estimated_time
                )
                if template_id:
                    st.session_state.template_state['quick_action'] = ('apply', template_id)
                    st.rerun()

def generate_template_structure_with_ai(
    doc_type: str,
    category: str,
    style: str,
    context: str
) -> str:
    """Génère une structure de template avec l'IA"""
    
    # Sélection du modèle IA
    ai_provider = select_ai_provider_for_generation()
    
    prompt = f"""Génère la structure d'un template juridique avec ces caractéristiques :
Type de document : {doc_type}
Catégorie : {category}
Style : {style} ({TEMPLATE_STYLES[style]['description']})
Contexte : {context}

INSTRUCTIONS :
1. Crée une structure claire et professionnelle
2. Utilise [NomVariable] pour les éléments à personnaliser
3. Respecte les conventions du type de document
4. Adapte le ton au style demandé
5. Inclus toutes les sections nécessaires

Retourne UNIQUEMENT la structure du template, ligne par ligne."""
    
    try:
        # Multi-IA si activé
        if st.session_state.template_state.get('ai_mode') == '🔀 Fusion Multi-IA':
            structure = generate_with_multi_ai_fusion(prompt)
        else:
            from managers.multi_llm_manager import MultiLLMManager
            llm_manager = MultiLLMManager()
            
            response = llm_manager.query_single_llm(
                ai_provider,
                prompt,
                "Tu es un expert en rédaction de documents juridiques."
            )
            
            if response['success']:
                structure = response['response']
            else:
                structure = ""
        
        return structure
        
    except Exception as e:
        st.error(f"❌ Erreur génération IA : {str(e)}")
        return ""

def generate_with_multi_ai_fusion(prompt: str) -> str:
    """Génération avec fusion de plusieurs IA"""
    
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    if len(llm_manager.clients) < 2:
        # Fallback sur un seul modèle
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            prompt,
            "Tu es un expert en rédaction juridique."
        )
        return response['response'] if response['success'] else ""
    
    # Génération parallèle
    with st.spinner("🔀 Fusion Multi-IA en cours..."):
        responses = []
        
        # Progress bar
        progress = st.progress(0)
        
        for idx, provider in enumerate(llm_manager.clients.keys()):
            if idx >= 3:  # Limiter à 3 modèles
                break
            
            try:
                response = llm_manager.query_single_llm(
                    provider,
                    prompt + f"\n\nGénère une version {['structurée', 'détaillée', 'concise'][idx]}.",
                    "Tu es un expert en rédaction juridique."
                )
                
                if response['success']:
                    responses.append({
                        'provider': provider,
                        'content': response['response']
                    })
                
                progress.progress((idx + 1) / min(3, len(llm_manager.clients)))
                
            except Exception as e:
                st.warning(f"⚠️ {provider} non disponible")
        
        progress.empty()
    
    if not responses:
        return ""
    
    # Fusion des réponses
    return fuse_ai_responses(responses, prompt)

def fuse_ai_responses(responses: List[Dict], original_prompt: str) -> str:
    """Fusionne intelligemment les réponses de plusieurs IA"""
    
    if len(responses) == 1:
        return responses[0]['content']
    
    # Analyser et extraire les meilleures parties
    fusion_prompt = f"""Fusionne ces {len(responses)} versions de template en gardant le meilleur de chaque :

{chr(10).join([f"VERSION {i+1} ({r['provider']}):\n{r['content']}\n" for i, r in enumerate(responses)])}

INSTRUCTIONS DE FUSION :
1. Garde la structure la plus complète et logique
2. Intègre les meilleures formulations de chaque version
3. Assure la cohérence globale
4. Conserve toutes les variables [NomVariable]
5. Élimine les redondances

Retourne la version fusionnée optimale."""
    
    # Utiliser le premier modèle disponible pour la fusion
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    provider = list(llm_manager.clients.keys())[0]
    response = llm_manager.query_single_llm(
        provider,
        fusion_prompt,
        "Tu es un expert en fusion de contenus juridiques."
    )
    
    return response['response'] if response['success'] else responses[0]['content']

def show_apply_wizard(template_id: str):
    """Assistant d'application de template avec IA avancée"""
    
    templates = get_all_templates()
    template = templates.get(template_id)
    
    if not template:
        st.error("❌ Template introuvable")
        return
    
    st.markdown(f"## 🎯 Utiliser : {template['name']}")
    
    # Infos template
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Type", template.get('type', 'document'))
    with col2:
        st.metric("Catégorie", template.get('category', 'Autre'))
    with col3:
        st.metric("Temps estimé", template.get('estimated_time', '20 min'))
    with col4:
        st.metric("Variables", len(extract_variables(template.get('structure', []))))
    
    # Configuration IA
    st.markdown("### 🤖 Configuration de l'IA")
    
    col1, col2 = st.columns(2)
    
    with col1:
        ai_mode = st.selectbox(
            "Mode de génération",
            ["🎯 Mono-IA optimisée", "🔀 Fusion Multi-IA", "⚡ Génération rapide", "🔬 Analyse approfondie"],
            key="apply_ai_mode"
        )
        
        if ai_mode == "🎯 Mono-IA optimisée":
            selected_provider = select_ai_provider_for_generation()
    
    with col2:
        enrich_options = st.multiselect(
            "Enrichissements",
            ["📚 Jurisprudence", "📊 Analyse contextuelle", "💡 Suggestions", "🔍 Vérification cohérence"],
            default=["📊 Analyse contextuelle"],
            key="apply_enrich"
        )
    
    # Variables
    variables = extract_variables(template.get('structure', []))
    variable_values = {}
    
    if variables:
        st.markdown("### 📝 Remplir les variables")
        
        # Suggestion automatique des valeurs
        if st.button("🤖 Remplissage intelligent", key="smart_fill"):
            variable_values = smart_fill_variables(variables, template)
            st.session_state.prefilled_vars = variable_values
            st.success("✅ Variables pré-remplies avec l'IA")
            st.rerun()
        
        # Formulaire variables
        var_cols = st.columns(2)
        
        for idx, var in enumerate(variables):
            with var_cols[idx % 2]:
                default = st.session_state.get('prefilled_vars', {}).get(var, '')
                
                if var.lower() in ['date', 'date du jour']:
                    value = st.date_input(
                        f"📅 {var}",
                        value=datetime.now(),
                        key=f"var_{var}"
                    )
                    variable_values[var] = value.strftime('%d/%m/%Y')
                elif var.lower() in ['montant', 'somme', 'prix']:
                    value = st.number_input(
                        f"💰 {var}",
                        value=0.0,
                        format="%.2f",
                        key=f"var_{var}"
                    )
                    variable_values[var] = f"{value:,.2f} €"
                else:
                    value = st.text_input(
                        f"✏️ {var}",
                        value=default,
                        key=f"var_{var}"
                    )
                    variable_values[var] = value
    
    # Options avancées
    with st.expander("⚙️ Options avancées", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            tone_adjustment = st.slider(
                "Ajustement du ton",
                -2, 2, 0,
                format_func=lambda x: ["Très formel", "Formel", "Équilibré", "Direct", "Persuasif"][x+2],
                key="tone_adjust"
            )
            
            length_factor = st.slider(
                "Niveau de détail",
                0.5, 2.0, 1.0, 0.1,
                format_func=lambda x: f"{x:.1f}x",
                key="length_factor"
            )
        
        with col2:
            output_format = st.selectbox(
                "Format de sortie",
                ["📄 Word", "📋 PDF", "📝 Texte", "🌐 HTML"],
                key="output_format"
            )
            
            include_annotations = st.checkbox(
                "Inclure annotations",
                value=False,
                key="include_annotations"
            )
    
    # Génération
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("👁️ Aperçu", type="secondary"):
            show_generation_preview(template, variable_values)
    
    with col2:
        if st.button("🚀 Générer", type="primary", use_container_width=True):
            generate_document_advanced(
                template, variable_values, ai_mode,
                enrich_options, tone_adjustment,
                length_factor, output_format
            )
    
    with col3:
        if st.button("💾 Sauvegarder config"):
            save_generation_config(template_id, variable_values, ai_mode, enrich_options)

def smart_fill_variables(variables: List[str], template: Dict) -> Dict[str, str]:
    """Remplit intelligemment les variables avec l'IA"""
    
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    # Contexte depuis la session
    context = {
        'current_date': datetime.now().strftime('%d/%m/%Y'),
        'city': 'Paris',
        'user_info': st.session_state.get('user_info', {}),
        'case_info': st.session_state.get('current_case_info', {})
    }
    
    prompt = f"""Suggère des valeurs réalistes pour ces variables de template juridique :
Type de document : {template.get('type')}
Variables : {', '.join([f'[{v}]' for v in variables])}
Contexte disponible : {json.dumps(context, ensure_ascii=False)}

Pour chaque variable, donne une valeur appropriée et réaliste.
Format de réponse : JSON avec clé = nom de variable, valeur = suggestion"""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            prompt,
            "Tu es un assistant juridique qui aide à remplir des documents."
        )
        
        if response['success']:
            # Parser la réponse JSON
            import re
            json_match = re.search(r'\{[^}]+\}', response['response'], re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        
    except Exception as e:
        st.warning(f"⚠️ Remplissage auto partiel : {str(e)}")
    
    # Valeurs par défaut
    defaults = {}
    for var in variables:
        if 'date' in var.lower():
            defaults[var] = datetime.now().strftime('%d/%m/%Y')
        elif 'ville' in var.lower():
            defaults[var] = 'Paris'
        elif 'année' in var.lower():
            defaults[var] = str(datetime.now().year)
        else:
            defaults[var] = ''
    
    return defaults

def generate_document_advanced(
    template: Dict,
    variable_values: Dict[str, str],
    ai_mode: str,
    enrich_options: List[str],
    tone_adjustment: int,
    length_factor: float,
    output_format: str
):
    """Génération avancée du document avec options multiples"""
    
    with st.spinner("🚀 Génération en cours..."):
        # Progress tracking
        progress = st.progress(0)
        status = st.empty()
        
        # Étape 1: Remplacement des variables
        status.text("📝 Remplacement des variables...")
        content = '\n'.join(template.get('structure', []))
        
        for var, value in variable_values.items():
            content = content.replace(f"[{var}]", value)
        
        progress.progress(25)
        
        # Étape 2: Enrichissement IA
        if ai_mode != "⚡ Génération rapide":
            status.text("🤖 Enrichissement par IA...")
            
            if ai_mode == "🔀 Fusion Multi-IA":
                content = enrich_with_multi_ai(
                    content, template, tone_adjustment, length_factor
                )
            else:
                content = enrich_with_single_ai(
                    content, template, tone_adjustment, length_factor
                )
        
        progress.progress(50)
        
        # Étape 3: Enrichissements optionnels
        if "📚 Jurisprudence" in enrich_options:
            status.text("📚 Ajout de la jurisprudence...")
            content = add_jurisprudence_advanced(content, template)
            progress.progress(65)
        
        if "📊 Analyse contextuelle" in enrich_options:
            status.text("📊 Analyse contextuelle...")
            content = add_contextual_analysis(content, template)
            progress.progress(80)
        
        if "🔍 Vérification cohérence" in enrich_options:
            status.text("🔍 Vérification de cohérence...")
            coherence_report = check_document_coherence(content, template)
            progress.progress(90)
        
        # Étape 4: Formatage final
        status.text("📄 Formatage final...")
        formatted_content = format_final_document(content, output_format)
        
        progress.progress(100)
        status.text("✅ Génération terminée !")
        
        # Sauvegarder le résultat
        result = {
            'content': formatted_content,
            'template': template['name'],
            'generated_at': datetime.now(),
            'ai_mode': ai_mode,
            'enrichments': enrich_options,
            'format': output_format,
            'variables': variable_values
        }
        
        st.session_state.last_generation = result
        
        # Afficher le résultat
        show_generation_result(result)

def show_generation_result(result: Dict):
    """Affiche le résultat de la génération"""
    
    st.success("✅ Document généré avec succès !")
    
    # Métriques
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        word_count = len(result['content'].split())
        st.metric("Mots", word_count)
    
    with col2:
        st.metric("Temps", datetime.now().strftime('%H:%M'))
    
    with col3:
        st.metric("Mode IA", result['ai_mode'].split()[1])
    
    with col4:
        st.metric("Enrichissements", len(result['enrichments']))
    
    # Aperçu
    with st.expander("📄 Aperçu du document", expanded=True):
        st.text_area(
            "Contenu généré",
            value=result['content'],
            height=400,
            key="result_preview"
        )
    
    # Actions
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        format_icon = {
            "📄 Word": "📄",
            "📋 PDF": "📋",
            "📝 Texte": "📝",
            "🌐 HTML": "🌐"
        }
        
        if result['format'] == "📄 Word":
            file_data = export_to_word(result['content'])
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        elif result['format'] == "📋 PDF":
            file_data = export_to_pdf(result['content'])
            mime = "application/pdf"
            ext = "pdf"
        elif result['format'] == "🌐 HTML":
            file_data = export_to_html(result['content'])
            mime = "text/html"
            ext = "html"
        else:
            file_data = result['content'].encode()
            mime = "text/plain"
            ext = "txt"
        
        st.download_button(
            f"{format_icon[result['format']]} Télécharger",
            data=file_data,
            file_name=f"{result['template']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}",
            mime=mime,
            use_container_width=True
        )
    
    with col2:
        if st.button("✏️ Modifier", use_container_width=True):
            st.session_state.editing_result = result
            st.rerun()
    
    with col3:
        if st.button("📧 Envoyer", use_container_width=True):
            show_send_dialog(result)
    
    with col4:
        if st.button("🔄 Régénérer", use_container_width=True):
            st.session_state.template_state['quick_action'] = None
            st.rerun()
    
    # Historique
    if 'generation_history' not in st.session_state:
        st.session_state.generation_history = []
    
    st.session_state.generation_history.append(result)
    
    # Feedback
    st.markdown("---")
    st.markdown("### 📝 Feedback")
    
    rating = st.feedback("thumbs", key="generation_feedback")
    
    if rating is not None:
        save_feedback(result, rating)

def show_template_analytics():
    """Affiche les analytics des templates"""
    
    st.markdown("## 📊 Analytics Templates")
    
    # Métriques globales
    col1, col2, col3, col4 = st.columns(4)
    
    all_templates = get_all_templates()
    history = st.session_state.get('generation_history', [])
    
    with col1:
        st.metric(
            "Total templates",
            len(all_templates),
            f"+{len([t for t in all_templates.values() if not t.get('is_builtin', True)])} custom"
        )
    
    with col2:
        st.metric(
            "Générations",
            len(history),
            f"+{len([h for h in history if (datetime.now() - h['generated_at']).days < 7])} cette semaine"
        )
    
    with col3:
        if history:
            avg_time = sum([(h['generated_at'] - datetime.now()).seconds for h in history[-10:]]) / min(10, len(history))
            st.metric("Temps moyen", f"{avg_time:.0f}s")
        else:
            st.metric("Temps moyen", "N/A")
    
    with col4:
        favorites = st.session_state.template_state.get('favorites', [])
        st.metric("Favoris", len(favorites))
    
    # Graphiques
    if history:
        # Usage par type
        fig_types = px.pie(
            values=[1] * len(history),
            names=[h['template'] for h in history],
            title="Répartition par template"
        )
        st.plotly_chart(fig_types, use_container_width=True)
        
        # Timeline
        dates = [h['generated_at'].date() for h in history]
        date_counts = defaultdict(int)
        for d in dates:
            date_counts[d] += 1
        
        fig_timeline = go.Figure()
        fig_timeline.add_trace(go.Scatter(
            x=list(date_counts.keys()),
            y=list(date_counts.values()),
            mode='lines+markers',
            name='Générations',
            line=dict(color='#667eea', width=3),
            marker=dict(size=10)
        ))
        fig_timeline.update_layout(
            title="Évolution des générations",
            xaxis_title="Date",
            yaxis_title="Nombre de générations"
        )
        st.plotly_chart(fig_timeline, use_container_width=True)

# Fonctions utilitaires

def get_all_templates() -> Dict[str, Dict]:
    """Récupère tous les templates disponibles"""
    templates = BUILTIN_TEMPLATES.copy()
    
    # Ajouter les templates custom
    if 'custom_templates' in st.session_state:
        templates.update(st.session_state.custom_templates)
    
    return templates

def get_filtered_templates(search_query: str = "") -> Dict[str, Dict]:
    """Filtre les templates selon la recherche"""
    templates = get_all_templates()
    
    if not search_query:
        return templates
    
    search_lower = search_query.lower()
    filtered = {}
    
    for tid, template in templates.items():
        searchable = f"{template['name']} {template.get('type', '')} {template.get('category', '')} {' '.join(template.get('tags', []))}".lower()
        
        if search_lower in searchable:
            filtered[tid] = template
    
    return filtered

def extract_variables(structure: List[str]) -> List[str]:
    """Extrait les variables d'une structure"""
    variables = set()
    
    for line in structure:
        matches = re.findall(r'\[([^\]]+)\]', line)
        variables.update(matches)
    
    return sorted(list(variables))

def get_search_suggestions(query: str) -> List[str]:
    """Suggestions de recherche intelligentes"""
    suggestions = []
    
    templates = get_all_templates()
    query_lower = query.lower()
    
    # Suggestions basées sur le nom
    for template in templates.values():
        if query_lower in template['name'].lower():
            suggestions.append(template['name'])
    
    # Suggestions basées sur les tags
    for template in templates.values():
        for tag in template.get('tags', []):
            if query_lower in tag.lower() and template['name'] not in suggestions:
                suggestions.append(f"{template['name']} (tag: {tag})")
    
    return suggestions[:5]

def select_ai_provider_for_generation() -> str:
    """Sélectionne le provider IA pour la génération"""
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    if not llm_manager.clients:
        st.error("❌ Aucun modèle IA configuré")
        return None
    
    providers = list(llm_manager.clients.keys())
    
    # Sélection intelligente basée sur les préférences
    if 'preferred_provider' in st.session_state:
        if st.session_state.preferred_provider in providers:
            return st.session_state.preferred_provider
    
    # Par défaut, prendre le premier disponible
    return providers[0]

def toggle_favorite(template_id: str):
    """Ajoute/retire des favoris"""
    if 'favorites' not in st.session_state.template_state:
        st.session_state.template_state['favorites'] = []
    
    favorites = st.session_state.template_state['favorites']
    
    if template_id in favorites:
        favorites.remove(template_id)
    else:
        favorites.append(template_id)
    
    st.rerun()

def save_new_template(
    name: str, doc_type: str, structure: str,
    style: str, category: str, description: str,
    tags: str, difficulty: str, estimated_time: str
) -> str:
    """Sauvegarde un nouveau template"""
    
    template_id = f"custom_{clean_key(name)}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    template = {
        'id': template_id,
        'name': name,
        'type': doc_type,
        'structure': structure.split('\n'),
        'style': style,
        'category': category,
        'description': description,
        'tags': [t.strip() for t in tags.split(',')] if tags else [],
        'difficulty': difficulty,
        'estimated_time': estimated_time,
        'created_at': datetime.now(),
        'updated_at': datetime.now(),
        'usage_count': 0,
        'is_builtin': False
    }
    
    if 'custom_templates' not in st.session_state:
        st.session_state.custom_templates = {}
    
    st.session_state.custom_templates[template_id] = template
    
    st.success(f"✅ Template '{name}' créé avec succès !")
    
    # Ajouter aux favoris automatiquement
    if 'favorites' not in st.session_state.template_state:
        st.session_state.template_state['favorites'] = []
    st.session_state.template_state['favorites'].append(template_id)
    
    return template_id

def enrich_with_single_ai(
    content: str, template: Dict,
    tone_adjustment: int, length_factor: float
) -> str:
    """Enrichit le contenu avec un seul modèle IA"""
    
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    provider = select_ai_provider_for_generation()
    if not provider:
        return content
    
    tone_map = {
        -2: "très formel et protocolaire",
        -1: "formel",
        0: "équilibré",
        1: "direct et clair",
        2: "persuasif et engagé"
    }
    
    prompt = f"""Enrichis ce document juridique en respectant ces consignes :
Type : {template.get('type')}
Style : {template.get('style')}
Ton : {tone_map[tone_adjustment]}
Facteur de longueur : {length_factor}x (1.0 = normal, <1 = plus court, >1 = plus long)

DOCUMENT :
{content}

INSTRUCTIONS :
1. Conserve EXACTEMENT la structure existante
2. Complète les sections vides ou peu développées
3. Adapte le ton selon les instructions
4. Respecte le facteur de longueur
5. Utilise un vocabulaire juridique approprié
6. Assure la cohérence globale

Retourne le document enrichi."""
    
    try:
        response = llm_manager.query_single_llm(
            provider,
            prompt,
            f"Tu es un expert en rédaction juridique spécialisé en {template.get('category', 'droit')}."
        )
        
        if response['success']:
            return response['response']
            
    except Exception as e:
        st.error(f"❌ Erreur enrichissement : {str(e)}")
    
    return content

def enrich_with_multi_ai(
    content: str, template: Dict,
    tone_adjustment: int, length_factor: float
) -> str:
    """Enrichit avec fusion de plusieurs IA"""
    
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    if len(llm_manager.clients) < 2:
        return enrich_with_single_ai(content, template, tone_adjustment, length_factor)
    
    # Prompts spécialisés pour chaque IA
    specialized_prompts = [
        ("Structure et forme", "Focus sur l'organisation et la présentation"),
        ("Contenu juridique", "Focus sur la précision juridique et les références"),
        ("Argumentation", "Focus sur la force persuasive et la logique")
    ]
    
    responses = []
    
    for idx, (provider, client) in enumerate(list(llm_manager.clients.items())[:3]):
        if idx < len(specialized_prompts):
            focus_name, focus_desc = specialized_prompts[idx]
            
            prompt = f"""Enrichis ce document juridique avec un focus particulier :
FOCUS : {focus_name} - {focus_desc}
Type : {template.get('type')}
Ton : {'formel' if tone_adjustment < 0 else 'persuasif' if tone_adjustment > 0 else 'équilibré'}

DOCUMENT :
{content}

Retourne le document enrichi en respectant le focus demandé."""
            
            try:
                response = llm_manager.query_single_llm(
                    provider,
                    prompt,
                    f"Tu es un expert en {focus_name.lower()} pour documents juridiques."
                )
                
                if response['success']:
                    responses.append({
                        'provider': provider,
                        'focus': focus_name,
                        'content': response['response']
                    })
                    
            except Exception:
                continue
    
    if not responses:
        return content
    
    # Fusion intelligente
    return fuse_specialized_responses(responses, template)

def fuse_specialized_responses(responses: List[Dict], template: Dict) -> str:
    """Fusionne les réponses spécialisées"""
    
    if len(responses) == 1:
        return responses[0]['content']
    
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    fusion_prompt = f"""Fusionne ces {len(responses)} versions enrichies en gardant le meilleur de chaque approche :

{chr(10).join([f"{r['focus'].upper()} ({r['provider']}):\n{r['content']}\n" for r in responses])}

INSTRUCTIONS :
1. Combine les forces de chaque version
2. Garde la meilleure structure
3. Intègre le contenu juridique le plus précis
4. Conserve les arguments les plus convaincants
5. Assure une cohérence parfaite
6. Élimine toute redondance

Type de document : {template.get('type')}

Retourne la version fusionnée optimale."""
    
    provider = list(llm_manager.clients.keys())[0]
    response = llm_manager.query_single_llm(
        provider,
        fusion_prompt,
        "Tu es un expert en fusion de documents juridiques."
    )
    
    return response['response'] if response['success'] else responses[0]['content']

def add_jurisprudence_advanced(content: str, template: Dict) -> str:
    """Ajoute la jurisprudence de manière intelligente"""
    
    # Extraction des concepts juridiques
    legal_concepts = extract_legal_concepts(content)
    
    # Recherche de jurisprudence (simulée)
    jurisprudence = search_jurisprudence(legal_concepts, template.get('type'))
    
    if not jurisprudence:
        return content
    
    # Insertion intelligente selon le type de document
    if template.get('type') in ['conclusions', 'memoire']:
        # Insérer dans la section discussion
        insertion_point = content.find("DISCUSSION")
        if insertion_point > -1:
            # Trouver la fin de la section
            next_section = content.find("\n\n", insertion_point + 100)
            
            juris_section = "\n\n   C. Jurisprudence applicable\n\n"
            for j in jurisprudence[:5]:
                juris_section += f"      {j['citation']}\n"
                juris_section += f"      → {j['principe']}\n\n"
            
            content = content[:next_section] + juris_section + content[next_section:]
    
    return content

def add_contextual_analysis(content: str, template: Dict) -> str:
    """Ajoute une analyse contextuelle au document"""
    
    from managers.multi_llm_manager import MultiLLMManager
    llm_manager = MultiLLMManager()
    
    if not llm_manager.clients:
        return content
    
    prompt = f"""Analyse ce document juridique et suggère des améliorations contextuelles :

DOCUMENT :
{content[:2000]}...

ANALYSE DEMANDÉE :
1. Points forts du document
2. Points à renforcer
3. Arguments supplémentaires pertinents
4. Risques juridiques éventuels

Retourne une analyse concise en 5-10 lignes."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            prompt,
            "Tu es un analyste juridique expert."
        )
        
        if response['success']:
            # Ajouter l'analyse en note
            analysis = f"\n\n---\n📊 ANALYSE CONTEXTUELLE (IA)\n{response['response']}\n---\n"
            return content + analysis
            
    except Exception:
        pass
    
    return content

def check_document_coherence(content: str, template: Dict) -> Dict:
    """Vérifie la cohérence du document"""
    
    report = {
        'score': 100,
        'issues': [],
        'suggestions': []
    }
    
    # Vérifications basiques
    lines = content.split('\n')
    
    # Structure
    expected_sections = extract_main_sections(template.get('structure', []))
    found_sections = extract_main_sections(lines)
    
    missing = set(expected_sections) - set(found_sections)
    if missing:
        report['issues'].append(f"Sections manquantes : {', '.join(missing)}")
        report['score'] -= 10 * len(missing)
    
    # Variables non remplacées
    unresolved = re.findall(r'\[[^\]]+\]', content)
    if unresolved:
        report['issues'].append(f"Variables non remplies : {', '.join(set(unresolved))}")
        report['score'] -= 5 * len(set(unresolved))
    
    # Cohérence des dates
    dates = re.findall(r'\d{1,2}/\d{1,2}/\d{4}', content)
    if dates:
        try:
            parsed_dates = [datetime.strptime(d, '%d/%m/%Y') for d in dates]
            if max(parsed_dates) < min(parsed_dates):
                report['issues'].append("Incohérence chronologique détectée")
                report['score'] -= 15
        except:
            pass
    
    # Suggestions
    if report['score'] < 100:
        report['suggestions'].append("Vérifiez et corrigez les points signalés")
    
    return report

def extract_main_sections(lines: List[str]) -> List[str]:
    """Extrait les sections principales d'un document"""
    sections = []
    
    for line in lines:
        # Sections numérotées romaines
        if re.match(r'^[IVX]+\.\s+[A-Z]', line):
            sections.append(line.strip())
        # Sections en majuscules
        elif line.isupper() and len(line.split()) > 1:
            sections.append(line.strip())
    
    return sections

def extract_legal_concepts(content: str) -> List[str]:
    """Extrait les concepts juridiques du contenu"""
    
    concepts = []
    
    # Patterns de concepts juridiques
    patterns = [
        r'(?:responsabilité|obligation|contrat|préjudice|dommages)',
        r'(?:prescription|forclusion|délai|procédure)',
        r'(?:compétence|recevabilité|fondé|moyens)',
        r'(?:articles?\s+\d+|L\.\s*\d+|R\.\s*\d+)',
        r'(?:Cass\.|CE\s|CA\s|TGI\s|TC\s)'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE)
        concepts.extend(matches)
    
    return list(set(concepts))[:10]

def search_jurisprudence(concepts: List[str], doc_type: str) -> List[Dict]:
    """Recherche de jurisprudence (simulée)"""
    
    # Base de jurisprudence simulée
    jurisprudence_db = [
        {
            'citation': 'Cass. civ. 1ère, 13 novembre 2021, n° 20-18.234',
            'principe': 'La responsabilité contractuelle suppose la preuve d\'un manquement'
        },
        {
            'citation': 'Cass. com., 7 septembre 2021, n° 19-24.309',
            'principe': 'La mise en demeure doit être claire et précise'
        },
        {
            'citation': 'CE, 15 juin 2021, n° 439436',
            'principe': 'Le délai de recours contentieux est d\'ordre public'
        }
    ]
    
    # Filtrer selon les concepts
    relevant = []
    for j in jurisprudence_db:
        for concept in concepts:
            if concept.lower() in j['principe'].lower():
                relevant.append(j)
                break
    
    return relevant

def format_final_document(content: str, output_format: str) -> str:
    """Formate le document selon le format de sortie"""
    
    if output_format == "🌐 HTML":
        # Conversion basique en HTML
        html_content = "<html><head><style>"
        html_content += """
        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }
        h1, h2, h3 { color: #2c3e50; }
        .section { margin: 20px 0; }
        .variable { background: #fff3cd; padding: 2px 6px; border-radius: 4px; }
        </style></head><body>"
        """
        
        # Conversion du contenu
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                if line.isupper():
                    html_content += f"<h2>{line}</h2>\n"
                elif re.match(r'^[IVX]+\.', line):
                    html_content += f"<h3>{line}</h3>\n"
                else:
                    html_content += f"<p>{line}</p>\n"
            else:
                html_content += "<br>\n"
        
        html_content += "</body></html>"
        return html_content
    
    return content

def export_to_word(content: str) -> bytes:
    """Exporte en format Word"""
    from io import BytesIO

    from docx import Document

    doc = Document()
        
        # Styles
        for line in content.split('\n'):
            if line.strip():
                if line.isupper():
                    doc.add_heading(line, level=1)
                elif re.match(r'^[IVX]+\.', line):
                    doc.add_heading(line, level=2)
                else:
                    doc.add_paragraph(line)
        
        # Sauvegarder en mémoire
        doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io.read()

def export_to_pdf(content: str) -> bytes:
    """Exporte en format PDF"""
    # Nécessiterait une bibliothèque PDF
    # Pour l'instant, retourner le contenu texte
    return content.encode('utf-8')

def export_to_html(content: str) -> bytes:
    """Exporte en format HTML"""
    html = format_final_document(content, "🌐 HTML")
    return html.encode('utf-8')

def show_generation_preview(template: Dict, variable_values: Dict[str, str]):
    """Aperçu de la génération"""
    
    with st.expander("👁️ Aperçu", expanded=True):
        # Remplacement simple des variables
        content = '\n'.join(template.get('structure', []))
        
        for var, value in variable_values.items():
            if value:
                content = content.replace(f"[{var}]", f"**{value}**")
        
        st.markdown(content)

def save_generation_config(
    template_id: str,
    variable_values: Dict,
    ai_mode: str,
    enrich_options: List[str]
):
    """Sauvegarde une configuration de génération"""
    
    config = {
        'template_id': template_id,
        'variables': variable_values,
        'ai_mode': ai_mode,
        'enrichments': enrich_options,
        'saved_at': datetime.now()
    }
    
    if 'saved_configs' not in st.session_state:
        st.session_state.saved_configs = []
    
    st.session_state.saved_configs.append(config)
    st.success("✅ Configuration sauvegardée")

def show_send_dialog(result: Dict):
    """Dialog d'envoi du document"""
    
    with st.expander("📧 Envoyer le document", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            recipient = st.text_input("Destinataire", key="send_recipient")
            subject = st.text_input(
                "Objet",
                value=f"{result['template']} - {datetime.now().strftime('%d/%m/%Y')}",
                key="send_subject"
            )
        
        with col2:
            cc = st.text_input("CC", key="send_cc")
            message = st.text_area("Message", height=100, key="send_message")
        
        if st.button("📤 Envoyer", type="primary"):
            st.success("✅ Document envoyé (simulation)")

def save_feedback(result: Dict, rating: int):
    """Sauvegarde le feedback utilisateur"""
    
    feedback = {
        'template': result['template'],
        'rating': rating,
        'timestamp': datetime.now(),
        'ai_mode': result['ai_mode']
    }
    
    if 'feedbacks' not in st.session_state:
        st.session_state.feedbacks = []
    
    st.session_state.feedbacks.append(feedback)
    
    if rating == 0:  # Thumbs down
        st.warning("Merci pour votre retour. Nous améliorerons la génération.")
    else:
        st.success("Merci pour votre retour positif !")

def show_advanced_filters():
    """Filtres avancés pour les templates"""
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        categories = st.multiselect(
            "Catégories",
            list(CATEGORY_ICONS.keys()),
            key="filter_categories"
        )
    
    with col2:
        difficulties = st.multiselect(
            "Difficulté",
            ["Débutant", "Facile", "Intermédiaire", "Avancé", "Expert"],
            key="filter_difficulties"
        )
    
    with col3:
        time_ranges = st.multiselect(
            "Temps estimé",
            ["5-10 min", "10-20 min", "20-30 min", "30-45 min", "45-60 min", "+1h"],
            key="filter_times"
        )
    
    with col4:
        show_favorites = st.checkbox("⭐ Favoris uniquement", key="filter_favorites")
        show_recent = st.checkbox("🕐 Récents uniquement", key="filter_recent")

def show_list_view(search_query: str):
    """Vue liste des templates"""
    
    templates = get_filtered_templates(search_query)
    
    if not templates:
        st.info("🔍 Aucun template trouvé")
        return
    
    # Table des templates
    data = []
    for tid, template in templates.items():
        data.append({
            'Nom': template['name'],
            'Type': template.get('type', 'N/A'),
            'Catégorie': template.get('category', 'N/A'),
            'Difficulté': template.get('difficulty', 'N/A'),
            'Temps': template.get('estimated_time', 'N/A'),
            'Utilisations': template.get('usage_count', 0),
            'ID': tid
        })
    
    # Afficher en dataframe interactif
    import pandas as pd
    df = pd.DataFrame(data)
    
    # Sélection
    selected = st.dataframe(
        df,
        use_container_width=True,
        hide_index=True,
        selection_mode="single-row",
        on_select="rerun"
    )
    
    if selected and selected.selection.rows:
        selected_id = data[selected.selection.rows[0]]['ID']
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("👁️ Aperçu", key="list_preview"):
                show_template_preview_modern(templates[selected_id])
        with col2:
            if st.button("⚡ Utiliser", key="list_use"):
                quick_apply_template(selected_id, templates[selected_id])
        with col3:
            if st.button("✏️ Modifier", key="list_edit"):
                st.session_state.template_state['quick_action'] = ('edit', selected_id)
                st.rerun()

def show_kanban_view(search_query: str):
    """Vue Kanban des templates par catégorie"""
    
    templates = get_filtered_templates(search_query)
    
    if not templates:
        st.info("🔍 Aucun template trouvé")
        return
    
    # Grouper par catégorie
    categories = defaultdict(list)
    for tid, template in templates.items():
        categories[template.get('category', 'Autre')].append((tid, template))
    
    # Colonnes Kanban
    cols = st.columns(min(len(categories), 4))
    
    for idx, (category, items) in enumerate(sorted(categories.items())):
        with cols[idx % len(cols)]:
            st.markdown(f"### {CATEGORY_ICONS.get(category, '📄')} {category}")
            st.markdown(f"*{len(items)} templates*")
            
            # Cards minimalistes
            for tid, template in items:
                with st.container():
                    st.markdown(f"""
                    <div style="background: white; padding: 10px; margin: 5px 0; 
                              border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                        <strong>{template['name']}</strong><br>
                        <small>{template.get('type', 'N/A')} • {template.get('estimated_time', 'N/A')}</small>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if st.button("→", key=f"kanban_use_{tid}", help="Utiliser"):
                        quick_apply_template(tid, template)

# Point d'entrée principal
if __name__ == "__main__":
    run()