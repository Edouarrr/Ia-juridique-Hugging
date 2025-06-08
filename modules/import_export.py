# modules/import_export.py
"""Module unifié d'import/export des documents et résultats avec toutes les fonctionnalités"""

import streamlit as st
from datetime import datetime
from typing import Dict, Any, List, Optional, Union, Tuple
import json
import base64
import io
import re
from pathlib import Path
import zipfile
from collections import defaultdict

# Imports optionnels avec gestion des dépendances
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

from models.dataclasses import Document
from utils.helpers import clean_key

# ====================
# CONSTANTES GLOBALES
# ====================

# Types MIME supportés
MIME_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'doc': 'application/msword',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'xls': 'application/vnd.ms-excel',
    'txt': 'text/plain',
    'html': 'text/html',
    'json': 'application/json',
    'csv': 'text/csv',
    'rtf': 'application/rtf',
    'zip': 'application/zip'
}

# Extensions supportées
SUPPORTED_EXTENSIONS = {
    'documents': ['pdf', 'docx', 'doc', 'txt', 'rtf'],
    'data': ['csv', 'xlsx', 'xls', 'json'],
    'all': ['pdf', 'docx', 'doc', 'txt', 'rtf', 'csv', 'xlsx', 'xls', 'json', 'html']
}

# ====================
# FONCTIONS D'IMPORT
# ====================

def process_import_request(query: str, analysis: dict):
    """Traite une demande d'import avec interface complète"""
    
    st.markdown("### 📥 Import de documents")
    
    # Configuration de l'import
    config = get_import_config(analysis)
    
    # Interface d'upload avancée
    col1, col2 = st.columns([3, 1])
    
    with col1:
        uploaded_files = st.file_uploader(
            "Sélectionner les fichiers",
            type=config['file_types'],
            accept_multiple_files=True,
            key="import_files_uploader",
            help="Formats supportés: PDF, DOCX, TXT, CSV, XLSX, RTF"
        )
    
    with col2:
        st.markdown("**Limites**")
        st.caption(f"Max: {config['max_size_mb']} MB/fichier")
        st.caption(f"Types: {', '.join(config['file_types'][:3])}...")
    
    if uploaded_files:
        # Validation des fichiers
        valid_files, invalid_files = validate_files(uploaded_files, config)
        
        if invalid_files:
            st.warning(f"⚠️ {len(invalid_files)} fichiers invalides ignorés")
            with st.expander("Voir les détails"):
                for file, reason in invalid_files:
                    st.text(f"❌ {file.name}: {reason}")
        
        if valid_files:
            # Options d'import
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                import_mode = st.selectbox(
                    "Mode d'import",
                    ["Ajouter", "Remplacer", "Fusionner", "Versionner"],
                    key="import_mode",
                    help="Ajouter: conserve l'existant | Remplacer: écrase | Fusionner: combine | Versionner: garde historique"
                )
            
            with col2:
                destination = st.selectbox(
                    "Destination",
                    ["Documents locaux", "Azure Blob", "Dossier spécifique"],
                    key="import_destination"
                )
            
            with col3:
                auto_analyze = st.checkbox(
                    "Analyser auto",
                    value=True,
                    key="auto_analyze_import",
                    help="Lance une analyse IA après import"
                )
            
            with col4:
                auto_tag = st.checkbox(
                    "Tagger auto",
                    value=True,
                    key="auto_tag_import",
                    help="Ajoute des tags automatiques"
                )
            
            # Référence/dossier avec suggestions
            reference = st.text_input(
                "Référence du dossier",
                value=analysis.get('reference', ''),
                placeholder="Ex: affaire_martin_2024",
                key="import_reference",
                help="Laissez vide pour utiliser la date"
            )
            
            # Métadonnées additionnelles
            with st.expander("➕ Métadonnées additionnelles", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    tags = st.text_input(
                        "Tags (séparés par des virgules)",
                        placeholder="contrat, immobilier, 2024",
                        key="import_tags"
                    )
                
                with col2:
                    category = st.selectbox(
                        "Catégorie",
                        ["Non classé", "Contrats", "Correspondances", "Procédures", "Expertises", "Autre"],
                        key="import_category"
                    )
                
                notes = st.text_area(
                    "Notes",
                    placeholder="Informations complémentaires...",
                    key="import_notes",
                    height=100
                )
            
            # Bouton d'import avec confirmation
            col1, col2, col3 = st.columns([1, 2, 1])
            
            with col2:
                if st.button(
                    f"🚀 Importer {len(valid_files)} fichiers", 
                    type="primary", 
                    key="execute_import",
                    use_container_width=True
                ):
                    # Confirmation pour mode destructif
                    if import_mode in ["Remplacer"] and not st.session_state.get('import_confirmed'):
                        st.warning("⚠️ Cette action va remplacer les documents existants!")
                        if st.button("✅ Confirmer", key="confirm_replace"):
                            st.session_state.import_confirmed = True
                            st.rerun()
                    else:
                        # Exécuter l'import
                        import_results = import_files_advanced(
                            valid_files,
                            config,
                            import_mode,
                            destination,
                            reference,
                            {
                                'tags': tags.split(',') if tags else [],
                                'category': category,
                                'notes': notes,
                                'auto_tag': auto_tag
                            }
                        )
                        
                        # Réinitialiser la confirmation
                        if 'import_confirmed' in st.session_state:
                            del st.session_state.import_confirmed
                        
                        # Afficher les résultats
                        display_import_results(import_results)
                        
                        # Analyse automatique si demandée
                        if auto_analyze and import_results['success_count'] > 0:
                            st.info("🤖 Lancement de l'analyse automatique...")
                            analyze_imported_documents(import_results['imported_docs'])

def get_import_config(analysis: dict) -> dict:
    """Configuration avancée de l'import"""
    
    config = {
        'file_types': analysis.get('file_types', SUPPORTED_EXTENSIONS['all']),
        'max_size_mb': st.session_state.get('import_max_size', 50),
        'extract_text': True,
        'preserve_formatting': False,
        'ocr_enabled': False
    }
    
    # Options avancées
    with st.expander("⚙️ Options avancées d'import", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            config['extract_metadata'] = st.checkbox(
                "Extraire les métadonnées",
                value=True,
                key="extract_metadata",
                help="Extrait auteur, date création, etc."
            )
            
            config['split_large_docs'] = st.checkbox(
                "Diviser les gros documents",
                value=False,
                key="split_large_docs",
                help="Divise les PDF de plus de 50 pages"
            )
            
            config['detect_language'] = st.checkbox(
                "Détecter la langue",
                value=False,
                key="detect_language"
            )
        
        with col2:
            config['preserve_formatting'] = st.checkbox(
                "Conserver la mise en forme",
                value=False,
                key="preserve_formatting",
                help="Garde styles et formatage (plus lent)"
            )
            
            config['extract_images'] = st.checkbox(
                "Extraire les images",
                value=False,
                key="extract_images"
            )
            
            config['ocr_enabled'] = st.checkbox(
                "OCR pour PDF scannés",
                value=False,
                key="ocr_enabled",
                help="Nécessite tesseract"
            )
        
        if config['split_large_docs']:
            config['split_size'] = st.number_input(
                "Taille max par partie (pages)",
                min_value=10,
                max_value=100,
                value=50,
                key="split_size"
            )
        
        config['encoding'] = st.selectbox(
            "Encodage des fichiers texte",
            ["auto", "utf-8", "latin-1", "cp1252", "iso-8859-1"],
            key="text_encoding"
        )
        
        # Options de traitement
        st.markdown("**Options de traitement**")
        
        config['clean_text'] = st.checkbox(
            "Nettoyer le texte",
            value=True,
            key="clean_text",
            help="Supprime caractères spéciaux et espaces multiples"
        )
        
        config['extract_tables'] = st.checkbox(
            "Extraire les tableaux",
            value=True,
            key="extract_tables"
        )
    
    return config

def validate_files(files: List, config: dict) -> Tuple[List, List]:
    """Valide les fichiers avant import"""
    
    valid_files = []
    invalid_files = []
    
    for file in files:
        # Vérifier l'extension
        ext = file.name.lower().split('.')[-1]
        if ext not in config['file_types']:
            invalid_files.append((file, f"Type non supporté: {ext}"))
            continue
        
        # Vérifier la taille
        size_mb = file.size / (1024 * 1024)
        if size_mb > config['max_size_mb']:
            invalid_files.append((file, f"Trop volumineux: {size_mb:.1f} MB"))
            continue
        
        # Vérifier le nom de fichier
        if not re.match(r'^[\w\-. ]+$', file.name):
            invalid_files.append((file, "Nom de fichier invalide"))
            continue
        
        valid_files.append(file)
    
    return valid_files, invalid_files

def import_files_advanced(
    files: List, 
    config: dict, 
    mode: str, 
    destination: str, 
    reference: str,
    metadata: dict
) -> dict:
    """Import avancé avec toutes les options"""
    
    results = {
        'total_files': len(files),
        'success_count': 0,
        'error_count': 0,
        'imported_docs': [],
        'errors': [],
        'warnings': []
    }
    
    # Générer la référence si vide
    if not reference:
        reference = f"import_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
    # Créer une barre de progression
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Container pour les logs détaillés
    log_container = st.container()
    
    for i, file in enumerate(files):
        try:
            status_text.text(f"Import de {file.name}...")
            
            # Lire le contenu
            content = file.read()
            
            # Détecter et extraire selon le type
            file_ext = file.name.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                text_content, file_metadata = extract_pdf_content_advanced(content, config)
            elif file_ext in ['docx', 'doc']:
                text_content, file_metadata = extract_docx_content_advanced(content, config)
            elif file_ext in ['xlsx', 'xls']:
                text_content, file_metadata = extract_excel_content_advanced(content, config)
            elif file_ext == 'csv':
                text_content, file_metadata = extract_csv_content_advanced(content, config)
            elif file_ext == 'rtf':
                text_content, file_metadata = extract_rtf_content(content, config)
            else:
                text_content, file_metadata = extract_text_content_advanced(content, config)
            
            # Nettoyer le texte si demandé
            if config.get('clean_text'):
                text_content = clean_extracted_text(text_content)
            
            # Détecter la langue si demandé
            if config.get('detect_language'):
                file_metadata['language'] = detect_language(text_content[:1000])
            
            # Auto-tagging si activé
            if metadata.get('auto_tag'):
                auto_tags = generate_auto_tags(text_content, file.name)
                metadata['tags'].extend(auto_tags)
            
            # Créer le document
            doc_id = generate_doc_id(file.name, reference)
            
            document = Document(
                id=doc_id,
                title=file.name,
                content=text_content,
                source=f"Import {reference}",
                metadata={
                    **file_metadata,
                    'imported_at': datetime.now().isoformat(),
                    'original_filename': file.name,
                    'file_size': file.size,
                    'file_size_mb': round(file.size / (1024 * 1024), 2),
                    'reference': reference,
                    'import_mode': mode,
                    'category': metadata.get('category', 'Non classé'),
                    'tags': list(set(metadata.get('tags', []))),
                    'notes': metadata.get('notes', ''),
                    'import_config': {
                        'preserve_formatting': config.get('preserve_formatting'),
                        'extract_metadata': config.get('extract_metadata'),
                        'extract_tables': config.get('extract_tables')
                    }
                }
            )
            
            # Stocker selon la destination et le mode
            store_document(document, content, destination, mode, reference)
            
            results['imported_docs'].append(document)
            results['success_count'] += 1
            
            # Log de succès
            with log_container:
                st.success(f"✅ {file.name} importé ({file_metadata.get('char_count', 0):,} caractères)")
            
        except Exception as e:
            results['errors'].append({
                'file': file.name,
                'error': str(e),
                'traceback': str(e.__class__.__name__)
            })
            results['error_count'] += 1
            
            # Log d'erreur
            with log_container:
                st.error(f"❌ {file.name}: {str(e)}")
        
        # Mettre à jour la progression
        progress_bar.progress((i + 1) / len(files))
    
    # Nettoyer l'interface
    progress_bar.empty()
    status_text.empty()
    
    return results

def extract_pdf_content_advanced(content: bytes, config: dict) -> Tuple[str, dict]:
    """Extraction avancée du contenu PDF avec toutes les options"""
    
    text = ""
    metadata = {'type': 'pdf', 'extraction_method': 'pypdf2'}
    
    if PYPDF2_AVAILABLE:
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            
            # Métadonnées du document
            if config.get('extract_metadata') and pdf_reader.metadata:
                metadata['pdf_metadata'] = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'producer': pdf_reader.metadata.get('/Producer', ''),
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                    'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                }
            
            metadata['page_count'] = len(pdf_reader.pages)
            metadata['is_encrypted'] = pdf_reader.is_encrypted
            
            # Extraction du texte page par page
            extracted_pages = []
            tables_found = 0
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    
                    # Détecter les tableaux (heuristique simple)
                    if config.get('extract_tables') and '|' in page_text:
                        tables_found += 1
                    
                    # Diviser si document trop grand
                    if config.get('split_large_docs') and page_num > 0 and page_num % config.get('split_size', 50) == 0:
                        page_text = f"\n\n[--- PARTIE {page_num // config.get('split_size', 50) + 1} - Page {page_num + 1} ---]\n\n" + page_text
                    
                    extracted_pages.append(page_text)
                    
                except Exception as e:
                    extracted_pages.append(f"\n[Erreur extraction page {page_num + 1}: {str(e)}]\n")
            
            text = "\n".join(extracted_pages)
            metadata['tables_detected'] = tables_found
            
            # OCR si aucun texte extrait et OCR activé
            if not text.strip() and config.get('ocr_enabled'):
                metadata['extraction_method'] = 'ocr_needed'
                text = "[Document scanné - OCR requis pour extraction du texte]"
            
        except Exception as e:
            text = f"Erreur extraction PDF: {str(e)}"
            metadata['extraction_error'] = str(e)
    else:
        text = "PyPDF2 non installé - Extraction PDF non disponible"
        metadata['extraction_error'] = "PyPDF2 not available"
    
    # Statistiques du texte
    metadata.update(get_text_statistics(text))
    
    return text, metadata

def extract_docx_content_advanced(content: bytes, config: dict) -> Tuple[str, dict]:
    """Extraction avancée du contenu DOCX"""
    
    text = ""
    metadata = {'type': 'docx'}
    
    if DOCX_AVAILABLE:
        try:
            doc = docx.Document(io.BytesIO(content))
            
            # Métadonnées du document
            if config.get('extract_metadata'):
                core_props = doc.core_properties
                metadata['docx_metadata'] = {
                    'author': core_props.author or '',
                    'title': core_props.title or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'comments': core_props.comments or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or ''
                }
            
            # Extraction structurée
            content_parts = []
            
            # Paragraphes avec styles
            for para in doc.paragraphs:
                if para.text.strip():
                    if config.get('preserve_formatting'):
                        style_info = f"[{para.style.name}] " if para.style else ""
                        content_parts.append(style_info + para.text)
                    else:
                        content_parts.append(para.text)
            
            # Tables
            table_count = 0
            if config.get('extract_tables', True):
                for table_idx, table in enumerate(doc.tables):
                    table_count += 1
                    content_parts.append(f"\n[TABLEAU {table_count}]")
                    
                    # En-têtes si première ligne
                    headers = []
                    for cell in table.rows[0].cells:
                        headers.append(cell.text.strip())
                    
                    if headers:
                        content_parts.append(" | ".join(headers))
                        content_parts.append("-" * (len(" | ".join(headers))))
                    
                    # Données
                    for row_idx, row in enumerate(table.rows[1:]):
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        content_parts.append(" | ".join(row_data))
                    
                    content_parts.append("[FIN TABLEAU]\n")
            
            # Sections et en-têtes/pieds de page
            if config.get('preserve_formatting'):
                for section in doc.sections:
                    if section.header.paragraphs:
                        content_parts.insert(0, f"[EN-TÊTE] {section.header.paragraphs[0].text}")
                    if section.footer.paragraphs:
                        content_parts.append(f"[PIED DE PAGE] {section.footer.paragraphs[0].text}")
            
            text = "\n".join(content_parts)
            
            # Métadonnées supplémentaires
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = table_count
            metadata['section_count'] = len(doc.sections)
            
        except Exception as e:
            text = f"Erreur extraction DOCX: {str(e)}"
            metadata['extraction_error'] = str(e)
    else:
        text = "python-docx non installé - Extraction DOCX non disponible"
        metadata['extraction_error'] = "python-docx not available"
    
    metadata.update(get_text_statistics(text))
    
    return text, metadata

def extract_excel_content_advanced(content: bytes, config: dict) -> Tuple[str, dict]:
    """Extraction avancée du contenu Excel"""
    
    text = ""
    metadata = {'type': 'excel'}
    
    if PANDAS_AVAILABLE:
        try:
            excel_file = pd.ExcelFile(io.BytesIO(content))
            metadata['sheet_count'] = len(excel_file.sheet_names)
            metadata['sheets'] = excel_file.sheet_names
            
            all_sheets_content = []
            total_rows = 0
            total_cols = 0
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                total_rows += len(df)
                total_cols = max(total_cols, len(df.columns))
                
                # Contenu de la feuille
                sheet_content = f"\n[FEUILLE: {sheet_name}]\n"
                sheet_content += f"Dimensions: {len(df)} lignes × {len(df.columns)} colonnes\n\n"
                
                if config.get('preserve_formatting'):
                    # Export complet avec formatage
                    sheet_content += df.to_string(max_rows=None)
                else:
                    # Version condensée
                    sheet_content += f"Colonnes: {', '.join(map(str, df.columns))}\n\n"
                    
                    # Aperçu des données
                    preview_rows = min(100, len(df))
                    for idx, row in df.head(preview_rows).iterrows():
                        values = [str(val)[:50] for val in row.values]
                        sheet_content += f"{idx}: {' | '.join(values)}\n"
                    
                    if len(df) > preview_rows:
                        sheet_content += f"\n... et {len(df) - preview_rows} lignes supplémentaires\n"
                
                # Statistiques si données numériques
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0 and config.get('extract_tables'):
                    sheet_content += "\n[STATISTIQUES]\n"
                    stats = df[numeric_cols].describe()
                    sheet_content += stats.to_string()
                
                all_sheets_content.append(sheet_content)
            
            text = "\n\n".join(all_sheets_content)
            
            metadata['total_rows'] = total_rows
            metadata['total_columns'] = total_cols
            metadata['has_numeric_data'] = any(df.select_dtypes(include=['number']).columns.tolist() for df in [pd.read_excel(excel_file, sheet) for sheet in excel_file.sheet_names])
            
        except Exception as e:
            text = f"Erreur extraction Excel: {str(e)}"
            metadata['extraction_error'] = str(e)
    else:
        text = "pandas non installé - Extraction Excel non disponible"
        metadata['extraction_error'] = "pandas not available"
    
    metadata.update(get_text_statistics(text))
    
    return text, metadata

def extract_csv_content_advanced(content: bytes, config: dict) -> Tuple[str, dict]:
    """Extraction avancée du contenu CSV"""
    
    text = ""
    metadata = {'type': 'csv'}
    
    if PANDAS_AVAILABLE:
        try:
            # Détection automatique de l'encodage et du délimiteur
            if config.get('encoding') == 'auto':
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                delimiters = [',', ';', '\t', '|']
                
                df = None
                successful_config = {}
                
                for enc in encodings:
                    for delim in delimiters:
                        try:
                            decoded = content.decode(enc)
                            df = pd.read_csv(io.StringIO(decoded), delimiter=delim)
                            
                            # Vérifier que le parsing est correct
                            if len(df.columns) > 1 and not df.empty:
                                successful_config = {'encoding': enc, 'delimiter': delim}
                                break
                        except:
                            continue
                    
                    if df is not None and not df.empty:
                        break
                
                if df is None:
                    raise ValueError("Impossible de parser le CSV avec les encodages/délimiteurs testés")
                
                metadata['detected_config'] = successful_config
            else:
                # Utiliser la configuration spécifiée
                df = pd.read_csv(
                    io.StringIO(content.decode(config.get('encoding', 'utf-8'))),
                    delimiter=config.get('delimiter', ',')
                )
            
            # Métadonnées
            metadata['row_count'] = len(df)
            metadata['column_count'] = len(df.columns)
            metadata['columns'] = list(df.columns)
            metadata['memory_usage'] = df.memory_usage(deep=True).sum()
            
            # Types de données
            metadata['column_types'] = df.dtypes.astype(str).to_dict()
            
            # Contenu formaté
            text = f"CSV: {metadata['row_count']} lignes × {metadata['column_count']} colonnes\n"
            text += f"Colonnes: {', '.join(df.columns)}\n\n"
            
            # Aperçu des données
            if len(df) <= 200:
                text += df.to_string()
            else:
                text += "[DÉBUT DU FICHIER - 100 premières lignes]\n"
                text += df.head(100).to_string()
                text += "\n\n[...]\n\n"
                text += "[FIN DU FICHIER - 100 dernières lignes]\n"
                text += df.tail(100).to_string()
            
            # Statistiques détaillées
            if config.get('extract_tables'):
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    text += "\n\n[STATISTIQUES DÉTAILLÉES]\n"
                    stats = df[numeric_cols].describe(percentiles=[.1, .25, .5, .75, .9])
                    text += stats.to_string()
                    
                    # Valeurs manquantes
                    missing = df.isnull().sum()
                    if missing.any():
                        text += "\n\n[VALEURS MANQUANTES]\n"
                        text += missing[missing > 0].to_string()
            
        except Exception as e:
            # Fallback avec module csv standard
            try:
                import csv
                
                decoded = content.decode(config.get('encoding', 'utf-8'))
                reader = csv.reader(io.StringIO(decoded))
                rows = list(reader)
                
                if rows:
                    metadata['row_count'] = len(rows)
                    metadata['column_count'] = len(rows[0]) if rows else 0
                    
                    text = f"CSV: {len(rows)} lignes\n"
                    if rows:
                        text += f"En-têtes: {', '.join(rows[0])}\n\n"
                        
                        for i, row in enumerate(rows[:min(100, len(rows))]):
                            text += f"{i}: {' | '.join(row)}\n"
                        
                        if len(rows) > 100:
                            text += f"\n... et {len(rows) - 100} lignes supplémentaires"
                
            except Exception as e2:
                text = f"Erreur extraction CSV: {str(e)} / {str(e2)}"
                metadata['extraction_error'] = f"{str(e)} / {str(e2)}"
    else:
        text = "pandas non installé - Extraction CSV limitée"
        metadata['extraction_error'] = "pandas not available"
    
    metadata.update(get_text_statistics(text))
    
    return text, metadata

def extract_text_content_advanced(content: bytes, config: dict) -> Tuple[str, dict]:
    """Extraction avancée du contenu texte avec détection d'encodage"""
    
    text = ""
    metadata = {'type': 'text'}
    
    # Détection automatique de l'encodage
    if config.get('encoding') == 'auto':
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16', 'ascii']
        
        for enc in encodings:
            try:
                text = content.decode(enc)
                metadata['detected_encoding'] = enc
                metadata['encoding_confidence'] = 'high'
                break
            except UnicodeDecodeError:
                continue
        
        if not text:
            # Fallback avec remplacement des caractères invalides
            text = content.decode('utf-8', errors='replace')
            metadata['detected_encoding'] = 'utf-8 (with errors)'
            metadata['encoding_confidence'] = 'low'
            metadata['encoding_errors'] = True
    else:
        try:
            text = content.decode(config.get('encoding', 'utf-8'))
            metadata['encoding'] = config.get('encoding', 'utf-8')
        except UnicodeDecodeError:
            text = content.decode(config.get('encoding', 'utf-8'), errors='replace')
            metadata['encoding_errors'] = True
    
    # Détection du format du texte
    metadata['has_markdown'] = bool(re.search(r'#{1,6}\s+\w+|^\*{1,2}\w+\*{1,2}|^\-\s+\w+', text, re.MULTILINE))
    metadata['has_code'] = bool(re.search(r'```[\s\S]*?```|`[^`]+`', text))
    metadata['has_urls'] = bool(re.search(r'https?://\S+', text))
    
    # Statistiques avancées
    metadata.update(get_text_statistics(text))
    
    return text, metadata

def extract_rtf_content(content: bytes, config: dict) -> Tuple[str, dict]:
    """Extraction du contenu RTF"""
    
    text = ""
    metadata = {'type': 'rtf'}
    
    try:
        # Extraction basique RTF (sans dépendance externe)
        rtf_text = content.decode('utf-8', errors='ignore')
        
        # Supprimer les commandes RTF basiques
        text = re.sub(r'\\[a-z]+\d*\s?', ' ', rtf_text)
        text = re.sub(r'[{}]', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Nettoyer les caractères spéciaux RTF
        text = text.replace('\\par', '\n')
        text = text.replace('\\tab', '\t')
        
        metadata['extraction_method'] = 'basic_regex'
        
    except Exception as e:
        text = f"Erreur extraction RTF: {str(e)}"
        metadata['extraction_error'] = str(e)
    
    metadata.update(get_text_statistics(text))
    
    return text, metadata

def get_text_statistics(text: str) -> dict:
    """Calcule des statistiques détaillées sur le texte"""
    
    stats = {
        'char_count': len(text),
        'char_count_no_spaces': len(text.replace(' ', '').replace('\n', '').replace('\t', '')),
        'word_count': len(text.split()),
        'line_count': len(text.split('\n')),
        'paragraph_count': len(re.split(r'\n\s*\n', text.strip())),
        'sentence_count': len(re.split(r'[.!?]+', text))
    }
    
    # Longueur moyenne des mots
    words = text.split()
    if words:
        stats['avg_word_length'] = sum(len(word) for word in words) / len(words)
    
    # Caractères spéciaux
    stats['has_special_chars'] = bool(re.search(r'[^\w\s.,!?;:\'\"-]', text))
    
    return stats

def clean_extracted_text(text: str) -> str:
    """Nettoie le texte extrait"""
    
    # Supprimer les espaces multiples
    text = re.sub(r'\s+', ' ', text)
    
    # Supprimer les lignes vides multiples
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    
    # Nettoyer les caractères de contrôle
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
    
    # Normaliser les sauts de ligne
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    return text.strip()

def detect_language(text: str) -> str:
    """Détecte la langue du texte (heuristique simple)"""
    
    # Mots fréquents par langue
    language_words = {
        'fr': ['le', 'de', 'la', 'et', 'les', 'des', 'en', 'un', 'une', 'que', 'pour', 'dans'],
        'en': ['the', 'of', 'and', 'to', 'in', 'is', 'for', 'that', 'with', 'on', 'as', 'by'],
        'es': ['el', 'la', 'de', 'que', 'y', 'en', 'un', 'por', 'con', 'para', 'los', 'las'],
        'de': ['der', 'die', 'und', 'in', 'den', 'von', 'zu', 'das', 'mit', 'sich', 'auf', 'ist']
    }
    
    text_lower = text.lower()
    words = text_lower.split()
    
    scores = {}
    for lang, common_words in language_words.items():
        score = sum(1 for word in words if word in common_words)
        scores[lang] = score
    
    if scores:
        detected = max(scores, key=scores.get)
        if scores[detected] > 5:  # Seuil minimal
            return detected
    
    return 'unknown'

def generate_auto_tags(text: str, filename: str) -> List[str]:
    """Génère des tags automatiques basés sur le contenu"""
    
    tags = []
    
    # Tags basés sur l'extension
    ext = filename.lower().split('.')[-1]
    if ext in ['pdf', 'docx', 'doc']:
        tags.append('document')
    elif ext in ['xlsx', 'xls', 'csv']:
        tags.append('données')
    
    # Tags basés sur le contenu (patterns juridiques)
    patterns = {
        'contrat': r'\b(contrat|convention|accord|engagement)\b',
        'juridique': r'\b(tribunal|juge|avocat|procédure|jugement)\b',
        'immobilier': r'\b(immeuble|propriété|bail|location|loyer)\b',
        'financier': r'\b(euro|€|paiement|facture|montant|prix)\b',
        'correspondance': r'\b(lettre|courrier|email|courriel|destinataire)\b',
        'expertise': r'\b(expert|expertise|évaluation|estimation)\b'
    }
    
    text_lower = text.lower()
    for tag, pattern in patterns.items():
        if re.search(pattern, text_lower, re.IGNORECASE):
            tags.append(tag)
    
    # Tags basés sur les dates
    if re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', text):
        tags.append('daté')
    
    return list(set(tags))[:5]  # Limiter à 5 tags

def generate_doc_id(filename: str, reference: str) -> str:
    """Génère un ID unique pour le document"""
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:17]
    clean_name = clean_key(filename)[:30]
    
    if reference:
        return f"{clean_key(reference)}_{clean_name}_{timestamp}"
    else:
        return f"doc_{clean_name}_{timestamp}"

def store_document(document: Document, original_content: bytes, destination: str, mode: str, reference: str):
    """Stocke le document selon la destination et le mode"""
    
    if destination == "Azure Blob":
        store_in_azure_advanced(document, original_content, reference, mode)
    elif destination == "Dossier spécifique":
        store_in_folder(document, original_content, reference, mode)
    else:
        store_locally_advanced(document, mode)

def store_locally_advanced(document: Document, mode: str):
    """Stockage local avancé avec gestion des modes"""
    
    if 'azure_documents' not in st.session_state:
        st.session_state.azure_documents = {}
    
    if mode == "Remplacer":
        # Supprimer les documents avec la même référence
        if document.metadata.get('reference'):
            keys_to_remove = [
                key for key, doc in st.session_state.azure_documents.items()
                if doc.metadata.get('reference') == document.metadata['reference']
            ]
            for key in keys_to_remove:
                del st.session_state.azure_documents[key]
    
    elif mode == "Fusionner":
        # Chercher un document similaire
        for key, existing_doc in st.session_state.azure_documents.items():
            if (existing_doc.title == document.title and 
                existing_doc.metadata.get('reference') == document.metadata.get('reference')):
                # Fusionner les contenus
                existing_doc.content += f"\n\n[--- FUSION {datetime.now().strftime('%Y-%m-%d %H:%M')} ---]\n\n"
                existing_doc.content += document.content
                existing_doc.metadata['last_merge'] = datetime.now().isoformat()
                existing_doc.metadata['merge_count'] = existing_doc.metadata.get('merge_count', 0) + 1
                
                # Fusionner les tags
                existing_tags = set(existing_doc.metadata.get('tags', []))
                new_tags = set(document.metadata.get('tags', []))
                existing_doc.metadata['tags'] = list(existing_tags.union(new_tags))
                
                return
    
    elif mode == "Versionner":
        # Créer une version
        version_num = 1
        base_id = document.id.rsplit('_v', 1)[0]
        
        # Trouver la dernière version
        for key in st.session_state.azure_documents:
            if key.startswith(base_id + '_v'):
                try:
                    v_num = int(key.split('_v')[-1])
                    version_num = max(version_num, v_num + 1)
                except:
                    pass
        
        document.id = f"{base_id}_v{version_num}"
        document.metadata['version'] = version_num
    
    # Ajouter le document
    st.session_state.azure_documents[document.id] = document

def store_in_azure_advanced(document: Document, original_content: bytes, reference: str, mode: str):
    """Stockage Azure avancé avec options"""
    
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        try:
            container = reference if reference else 'imports'
            
            # Nom du blob avec organisation par date
            date_path = datetime.now().strftime('%Y/%m/%d')
            blob_name = f"{date_path}/{document.id}_{document.title}"
            
            # Uploader le fichier original
            blob_manager.upload_blob(container, blob_name, original_content)
            
            # Métadonnées Azure
            document.metadata['azure_blob_path'] = f"{container}/{blob_name}"
            document.metadata['azure_upload_date'] = datetime.now().isoformat()
            
            # Stocker aussi localement pour accès rapide
            store_locally_advanced(document, mode)
            
        except Exception as e:
            st.error(f"Erreur Azure: {str(e)}")
            # Fallback stockage local
            store_locally_advanced(document, mode)
    else:
        # Pas de connexion Azure
        store_locally_advanced(document, mode)

def store_in_folder(document: Document, original_content: bytes, reference: str, mode: str):
    """Stockage dans un dossier spécifique"""
    
    # Créer la structure de dossiers
    base_path = Path("imports") / (reference if reference else "general")
    date_path = base_path / datetime.now().strftime('%Y/%m/%d')
    
    try:
        date_path.mkdir(parents=True, exist_ok=True)
        
        # Sauvegarder le fichier
        file_path = date_path / f"{document.id}_{document.title}"
        
        with open(file_path, 'wb') as f:
            f.write(original_content)
        
        # Métadonnées
        document.metadata['file_path'] = str(file_path)
        
        # Stocker aussi en mémoire
        store_locally_advanced(document, mode)
        
    except Exception as e:
        st.error(f"Erreur sauvegarde fichier: {str(e)}")
        store_locally_advanced(document, mode)

def display_import_results(results: dict):
    """Affiche les résultats détaillés de l'import"""
    
    # Résumé
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("📥 Fichiers traités", results['total_files'])
    
    with col2:
        st.metric("✅ Succès", results['success_count'], 
                  delta=f"{results['success_count']/results['total_files']*100:.0f}%")
    
    with col3:
        if results['error_count'] > 0:
            st.metric("❌ Erreurs", results['error_count'], delta_color="inverse")
        else:
            st.metric("✨ Sans erreur", "0")
    
    # Détails des imports réussis
    if results['success_count'] > 0:
        with st.expander("📄 Documents importés", expanded=True):
            for doc in results['imported_docs']:
                col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                
                with col1:
                    st.write(f"**{doc.title}**")
                    if doc.metadata.get('tags'):
                        tags_str = " ".join([f"`{tag}`" for tag in doc.metadata['tags'][:3]])
                        st.caption(tags_str)
                
                with col2:
                    size_mb = doc.metadata.get('file_size_mb', 0)
                    st.caption(f"📊 {size_mb:.1f} MB")
                    chars = doc.metadata.get('char_count', 0)
                    st.caption(f"📝 {chars:,} car.")
                
                with col3:
                    doc_type = doc.metadata.get('type', 'inconnu').upper()
                    st.caption(f"📄 {doc_type}")
                    if doc.metadata.get('page_count'):
                        st.caption(f"📑 {doc.metadata['page_count']} p.")
                
                with col4:
                    if st.button("👁️", key=f"view_{doc.id}", help="Aperçu"):
                        st.text_area(
                            f"Aperçu de {doc.title}",
                            value=doc.content[:1000] + "...",
                            height=200,
                            key=f"preview_{doc.id}"
                        )
    
    # Erreurs détaillées
    if results['error_count'] > 0:
        with st.expander("⚠️ Erreurs détaillées", expanded=False):
            for error in results['errors']:
                st.error(f"**{error['file']}**")
                st.write(f"Erreur: {error['error']}")
                if 'traceback' in error:
                    st.caption(f"Type: {error['traceback']}")
    
    # Avertissements
    if results.get('warnings'):
        with st.expander("⚠️ Avertissements", expanded=False):
            for warning in results['warnings']:
                st.warning(warning)

def analyze_imported_documents(documents: List[Document]):
    """Lance une analyse IA automatique des documents importés"""
    
    if not documents:
        return
    
    # Grouper par catégorie
    by_category = defaultdict(list)
    for doc in documents:
        category = doc.metadata.get('category', 'Non classé')
        by_category[category].append(doc)
    
    # Créer un prompt d'analyse
    analysis_prompt = f"""Analyse approfondie de {len(documents)} documents importés.

DOCUMENTS PAR CATÉGORIE:
"""
    
    for category, docs in by_category.items():
        analysis_prompt += f"\n{category} ({len(docs)} documents):\n"
        for doc in docs[:3]:  # Limiter l'aperçu
            preview = doc.content[:300].replace('\n', ' ')
            analysis_prompt += f"- {doc.title}: {preview}...\n"
    
    analysis_prompt += """
ANALYSE DEMANDÉE:
1. Classification et typologie des documents
2. Identification des parties prenantes principales
3. Chronologie des événements clés
4. Points juridiques critiques identifiés
5. Risques et opportunités détectés
6. Recommandations pour l'organisation et l'analyse approfondie
7. Documents manquants potentiels à rechercher

Fournis une analyse structurée et détaillée."""
    
    # Stocker pour traitement
    st.session_state.import_analysis_request = {
        'prompt': analysis_prompt,
        'documents': documents,
        'timestamp': datetime.now()
    }
    
    # Notification
    st.success(f"✅ {len(documents)} documents prêts pour l'analyse")
    st.info("💡 Utilisez l'onglet 'Analyse IA' pour lancer l'analyse approfondie")

# ====================
# FONCTIONS D'EXPORT
# ====================

def process_export_request(query: str, analysis: dict):
    """Traite une demande d'export avec interface complète"""
    
    st.markdown("### 📤 Export de documents")
    
    # Déterminer le contenu exportable
    export_content = determine_export_content(analysis)
    
    if not export_content:
        st.warning("⚠️ Aucun contenu à exporter")
        st.info("💡 Générez d'abord des résultats : recherche, analyse, rédaction, chronologie...")
        
        # Suggestions
        st.markdown("**Contenus exportables :**")
        st.markdown("- 🔍 Résultats de recherche")
        st.markdown("- 🤖 Analyses IA")
        st.markdown("- 📝 Documents rédigés")
        st.markdown("- 📅 Chronologies")
        st.markdown("- 🗺️ Cartographies")
        st.markdown("- 🔄 Comparaisons")
        return
    
    # Configuration de l'export
    config = get_export_config_advanced(export_content['type'])
    
    # Aperçu du contenu
    with st.expander(f"👁️ Aperçu : {export_content['title']}", expanded=False):
        # Statistiques du contenu
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Type", export_content['type'].title())
        with col2:
            chars = len(export_content.get('content', ''))
            st.metric("Taille", f"{chars:,} car.")
        with col3:
            if export_content.get('metadata', {}).get('document_count'):
                st.metric("Documents", export_content['metadata']['document_count'])
        
        # Aperçu textuel
        st.text_area(
            "Contenu",
            value=export_content['preview'],
            height=200,
            disabled=True
        )
    
    # Options d'export principales
    col1, col2 = st.columns(2)
    
    with col1:
        export_format = st.selectbox(
            "📄 Format d'export",
            config['available_formats'],
            key="export_format_select",
            format_func=lambda x: config['format_names'].get(x, x.upper())
        )
    
    with col2:
        st.markdown("**Options générales**")
        include_metadata = st.checkbox(
            "Inclure métadonnées",
            value=True,
            key="include_metadata",
            help="Ajoute date, auteur, sources..."
        )
        
        include_toc = st.checkbox(
            "Table des matières",
            value=export_format in ['docx', 'pdf', 'html'],
            key="include_toc",
            disabled=export_format not in ['docx', 'pdf', 'html']
        )
    
    # Options spécifiques au format
    format_options = get_format_specific_options_advanced(export_format, export_content['type'])
    
    # Options d'export avancées
    with st.expander("⚙️ Options avancées", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            compress = st.checkbox(
                "📦 Compresser (ZIP)",
                value=False,
                key="compress_export",
                help="Créé une archive ZIP"
            )
            
            if compress:
                compression_level = st.slider(
                    "Niveau de compression",
                    0, 9, 6,
                    key="compression_level"
                )
            
            watermark = st.checkbox(
                "💧 Filigrane",
                value=False,
                key="add_watermark",
                disabled=export_format not in ['pdf', 'docx']
            )
        
        with col2:
            encrypt = st.checkbox(
                "🔒 Chiffrement",
                value=False,
                key="encrypt_export",
                disabled=export_format not in ['pdf', 'zip']
            )
            
            if encrypt:
                password = st.text_input(
                    "Mot de passe",
                    type="password",
                    key="export_password"
                )
            
            split_files = st.checkbox(
                "📑 Diviser en plusieurs fichiers",
                value=False,
                key="split_export",
                help="Pour les gros exports"
            )
            
            if split_files:
                split_size = st.number_input(
                    "Taille max par fichier (MB)",
                    1, 100, 10,
                    key="split_size_mb"
                )
    
    # Nom du fichier avec suggestions intelligentes
    default_filename = generate_export_filename_advanced(
        export_content['type'], 
        export_format,
        export_content.get('metadata', {})
    )
    
    filename = st.text_input(
        "📝 Nom du fichier",
        value=default_filename,
        key="export_filename",
        help="Sans extension (ajoutée automatiquement)"
    )
    
    # Validation du nom de fichier
    if not re.match(r'^[\w\-. ]+$', filename):
        st.error("❌ Nom de fichier invalide (utilisez uniquement lettres, chiffres, -, _, .)")
        return
    
    # Boutons d'action
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        if st.button(
            f"💾 Exporter en {config['format_names'].get(export_format, export_format.upper())}", 
            type="primary", 
            key="execute_export",
            use_container_width=True
        ):
            with st.spinner(f"Génération du fichier {export_format.upper()}..."):
                try:
                    # Effectuer l'export
                    exported_data = perform_export_advanced(
                        export_content,
                        export_format,
                        config,
                        format_options,
                        {
                            'include_metadata': include_metadata,
                            'include_toc': include_toc,
                            'watermark': watermark,
                            'encrypt': encrypt,
                            'password': password if encrypt else None
                        }
                    )
                    
                    if exported_data:
                        # Post-traitement
                        final_data = exported_data
                        final_filename = f"{filename}.{export_format}"
                        mime_type = MIME_TYPES.get(export_format, 'application/octet-stream')
                        
                        # Compression si demandée
                        if compress:
                            final_data = compress_data_advanced(
                                exported_data, 
                                final_filename,
                                compression_level if 'compression_level' in locals() else 6
                            )
                            final_filename = f"{filename}.zip"
                            mime_type = 'application/zip'
                        
                        # Division si demandée
                        if split_files and len(final_data) > split_size * 1024 * 1024:
                            parts = split_export_data(final_data, split_size * 1024 * 1024)
                            
                            st.success(f"✅ Export divisé en {len(parts)} parties")
                            
                            for i, part_data in enumerate(parts):
                                part_filename = f"{filename}_part{i+1}.{export_format}"
                                if compress:
                                    part_filename += ".zip"
                                
                                st.download_button(
                                    f"⬇️ Télécharger partie {i+1}/{len(parts)}",
                                    part_data,
                                    part_filename,
                                    mime_type,
                                    key=f"download_part_{i}"
                                )
                        else:
                            # Export simple
                            st.download_button(
                                f"⬇️ Télécharger {export_format.upper()}",
                                final_data,
                                final_filename,
                                mime_type,
                                key="download_export_file"
                            )
                            
                            st.success(f"✅ Export {export_format.upper()} prêt")
                        
                        # Statistiques et logs
                        display_export_statistics(final_data, export_format, export_content)
                        
                        # Log de l'export
                        log_export_action(export_content, export_format, filename)
                        
                except Exception as e:
                    st.error(f"❌ Erreur lors de l'export : {str(e)}")
                    st.exception(e)

def determine_export_content(analysis: dict) -> Optional[Dict[str, Any]]:
    """Détermine intelligemment le contenu à exporter"""
    
    # Ordre de priorité des contenus
    export_priorities = [
        ('redaction_result', 'redaction', format_redaction_content),
        ('timeline_result', 'timeline', format_timeline_content),
        ('mapping_result', 'mapping', format_mapping_content),
        ('comparison_result', 'comparison', format_comparison_content),
        ('ai_analysis_results', 'analysis', format_analysis_content),
        ('search_results', 'search', format_search_content),
        ('selected_pieces', 'pieces', format_pieces_content)
    ]
    
    for state_key, content_type, formatter in export_priorities:
        if st.session_state.get(state_key):
            data = st.session_state[state_key]
            formatted_content = formatter(data)
            
            return {
                'type': content_type,
                'title': formatted_content['title'],
                'content': formatted_content['content'],
                'metadata': formatted_content.get('metadata', {}),
                'preview': formatted_content['content'][:500] + '...',
                'source_data': data
            }
    
    return None

def get_export_config_advanced(content_type: str) -> dict:
    """Configuration avancée selon le type de contenu"""
    
    base_formats = ['docx', 'pdf', 'txt', 'html', 'json']
    
    config = {
        'available_formats': base_formats.copy(),
        'default_format': 'docx',
        'format_names': {
            'docx': '📄 Word',
            'pdf': '📕 PDF',
            'txt': '📝 Texte',
            'html': '🌐 HTML',
            'json': '🔧 JSON',
            'xlsx': '📊 Excel',
            'csv': '📋 CSV'
        }
    }
    
    # Formats spécifiques par type
    if content_type in ['timeline', 'mapping', 'comparison']:
        config['available_formats'].insert(2, 'xlsx')
        config['default_format'] = 'xlsx'
    
    elif content_type in ['search', 'pieces']:
        config['available_formats'].extend(['xlsx', 'csv'])
    
    elif content_type == 'analysis':
        config['available_formats'].insert(0, 'pdf')
        config['default_format'] = 'pdf'
    
    return config

def get_format_specific_options_advanced(format: str, content_type: str) -> dict:
    """Options avancées spécifiques au format et au type de contenu"""
    
    options = {}
    
    if format == 'docx':
        with st.expander("📄 Options Word", expanded=False):
            col1, col2 = st.columns(2)
            
            with col1:
                options['font'] = st.selectbox(
                    "Police",
                    ["Times New Roman", "Arial", "Calibri", "Garamond", "Verdana"],
                    key="docx_font"
                )
                
                options['font_size'] = st.number_input(
                    "Taille police (pt)",
                    8, 16, 12,
                    key="docx_font_size"
                )
                
                options['line_spacing'] = st.selectbox(
                    "Interligne",
                    ["1.0", "1.15", "1.5", "2.0"],
                    index=2,
                    key="docx_line_spacing"
                )
            
            with col2:
                options['margins'] = st.selectbox(
                    "Marges",
                    ["Normales", "Étroites", "Larges", "Miroir"],
                    key="docx_margins"
                )
                
                options['paper_size'] = st.selectbox(
                    "Format papier",
                    ["A4", "Letter", "Legal"],
                    key="docx_paper_size"
                )
                
                options['numbering'] = st.checkbox(
                    "Numérotation pages",
                    value=True,
                    key="docx_numbering"
                )
            
            # Options spécifiques au contenu juridique
            if content_type in ['redaction', 'analysis']:
                st.markdown("**Options juridiques**")
                
                options['legal_formatting'] = st.checkbox(
                    "Mise en forme juridique",
                    value=True,
                    key="docx_legal_format",
                    help="Applique les standards de mise en forme juridique"
                )
                
                options['section_numbering'] = st.checkbox(
                    "Numérotation des sections",
                    value=True,
                    key="docx_section_numbers"
                )
    
    elif format == 'pdf':
        with st.expander("📕 Options PDF", expanded=False):
            col1, col2 = st.columns(2)
            
            with col1:
                options['page_size'] = st.selectbox(
                    "Format",
                    ["A4", "Letter", "Legal", "A3"],
                    key="pdf_page_size"
                )
                
                options['orientation'] = st.radio(
                    "Orientation",
                    ["Portrait", "Paysage"],
                    key="pdf_orientation"
                )
                
                options['margins'] = st.slider(
                    "Marges (cm)",
                    0.5, 3.0, 2.0, 0.5,
                    key="pdf_margins"
                )
            
            with col2:
                options['compress_images'] = st.checkbox(
                    "Compresser images",
                    value=True,
                    key="pdf_compress_images"
                )
                
                options['embed_fonts'] = st.checkbox(
                    "Incorporer polices",
                    value=True,
                    key="pdf_embed_fonts"
                )
                
                options['security'] = st.selectbox(
                    "Sécurité",
                    ["Aucune", "Lecture seule", "Impression interdite"],
                    key="pdf_security"
                )
    
    elif format == 'html':
        with st.expander("🌐 Options HTML", expanded=False):
            col1, col2 = st.columns(2)
            
            with col1:
                options['theme'] = st.selectbox(
                    "Thème",
                    ["Clair", "Sombre", "Juridique", "Minimaliste", "Corporate"],
                    key="html_theme"
                )
                
                options['responsive'] = st.checkbox(
                    "Design responsive",
                    value=True,
                    key="html_responsive"
                )
                
                options['include_css'] = st.checkbox(
                    "CSS intégré",
                    value=True,
                    key="html_css"
                )
            
            with col2:
                options['include_js'] = st.checkbox(
                    "JavaScript interactif",
                    value=False,
                    key="html_js"
                )
                
                options['print_friendly'] = st.checkbox(
                    "Optimisé impression",
                    value=True,
                    key="html_print"
                )
                
                options['single_page'] = st.checkbox(
                    "Page unique",
                    value=True,
                    key="html_single"
                )
    
    elif format == 'xlsx':
        with st.expander("📊 Options Excel", expanded=False):
            col1, col2 = st.columns(2)
            
            with col1:
                options['include_charts'] = st.checkbox(
                    "Générer graphiques",
                    value=content_type in ['timeline', 'mapping'],
                    key="xlsx_charts"
                )
                
                options['freeze_headers'] = st.checkbox(
                    "Figer en-têtes",
                    value=True,
                    key="xlsx_freeze"
                )
                
                options['autofilter'] = st.checkbox(
                    "Filtres auto",
                    value=True,
                    key="xlsx_filter"
                )
            
            with col2:
                options['conditional_formatting'] = st.checkbox(
                    "Mise en forme conditionnelle",
                    value=False,
                    key="xlsx_conditional"
                )
                
                options['data_validation'] = st.checkbox(
                    "Validation données",
                    value=False,
                    key="xlsx_validation"
                )
                
                options['pivot_table'] = st.checkbox(
                    "Tableau croisé dynamique",
                    value=content_type == 'analysis',
                    key="xlsx_pivot"
                )
    
    return options

def generate_export_filename_advanced(content_type: str, format: str, metadata: dict) -> str:
    """Génère un nom de fichier intelligent basé sur le contexte"""
    
    # Base du nom selon le type
    type_names = {
        'redaction': 'document',
        'timeline': 'chronologie',
        'mapping': 'cartographie',
        'comparison': 'comparaison',
        'analysis': 'analyse',
        'search': 'recherche',
        'pieces': 'pieces'
    }
    
    base_name = type_names.get(content_type, 'export')
    
    # Ajouter la référence si disponible
    if metadata.get('reference'):
        base_name = f"{clean_key(metadata['reference'])}_{base_name}"
    
    # Ajouter le type de document pour les rédactions
    if content_type == 'redaction' and metadata.get('type'):
        base_name = f"{base_name}_{clean_key(metadata['type'])}"
    
    # Date
    date_str = datetime.now().strftime('%Y%m%d')
    
    return f"{base_name}_{date_str}"

def perform_export_advanced(
    content: dict,
    format: str,
    config: dict,
    format_options: dict,
    general_options: dict
) -> Optional[bytes]:
    """Effectue l'export avec toutes les options avancées"""
    
    try:
        # Préparer les options complètes
        export_options = {
            **format_options,
            **general_options,
            'config': config,
            'content_type': content['type']
        }
        
        # Appeler la fonction d'export appropriée
        if format == 'docx':
            return export_to_docx_advanced(content, export_options)
        
        elif format == 'pdf':
            return export_to_pdf_advanced(content, export_options)
        
        elif format == 'txt':
            return export_to_txt_advanced(content, export_options)
        
        elif format == 'html':
            return export_to_html_advanced(content, export_options)
        
        elif format == 'json':
            return export_to_json_advanced(content, export_options)
        
        elif format == 'xlsx':
            return export_to_xlsx_advanced(content, export_options)
        
        elif format == 'csv':
            return export_to_csv_advanced(content, export_options)
        
        else:
            st.error(f"Format {format} non supporté")
            return None
            
    except Exception as e:
        st.error(f"Erreur export {format}: {str(e)}")
        return None

# [Les fonctions d'export spécifiques suivent...]
# Je vais continuer avec les fonctions d'export détaillées dans la prochaine partie

def export_to_docx_advanced(content: dict, options: dict) -> bytes:
    """Export Word avancé avec toutes les options"""
    
    if not DOCX_AVAILABLE:
        return export_to_txt_advanced(content, options)
    
    try:
        doc = docx.Document()
        
        # Configuration du document
        setup_docx_document_advanced(doc, options)
        
        # En-tête si demandé
        if options.get('include_metadata', True):
            add_docx_header_advanced(doc, content, options)
        
        # Table des matières si demandée
        if options.get('include_toc', False):
            add_docx_toc_advanced(doc, content)
        
        # Titre principal avec style
        add_docx_title(doc, content['title'], options)
        
        # Métadonnées en début de document
        if options.get('include_metadata', True):
            add_docx_metadata_advanced(doc, content, options)
        
        # Contenu principal selon le type
        if content['type'] == 'redaction' and options.get('legal_formatting', True):
            add_docx_legal_content_advanced(doc, content, options)
        else:
            add_docx_content_advanced(doc, content, options)
        
        # Pied de page
        if options.get('numbering', True):
            add_docx_footer_advanced(doc, content, options)
        
        # Filigrane si demandé
        if options.get('watermark', False):
            add_docx_watermark(doc, options)
        
        # Sauvegarder
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur création DOCX: {e}")
        return export_to_txt_advanced(content, options)

def setup_docx_document_advanced(doc, options: dict):
    """Configuration avancée du document Word"""
    
    # Marges
    sections = doc.sections
    for section in sections:
        margins = {
            'Normales': (1, 1, 1, 1),
            'Étroites': (0.5, 0.5, 0.5, 0.5),
            'Larges': (1.5, 1.5, 1.5, 1.5),
            'Miroir': (1, 1.25, 1, 1.5)
        }
        
        margin_values = margins.get(options.get('margins', 'Normales'), (1, 1, 1, 1))
        section.top_margin = Inches(margin_values[0])
        section.bottom_margin = Inches(margin_values[1])
        section.left_margin = Inches(margin_values[2])
        section.right_margin = Inches(margin_values[3])
        
        # Format papier
        if options.get('paper_size') == 'Letter':
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif options.get('paper_size') == 'Legal':
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)
    
    # Styles de base
    styles = doc.styles
    
    # Style normal
    normal_style = styles['Normal']
    normal_style.font.name = options.get('font', 'Times New Roman')
    normal_style.font.size = Pt(options.get('font_size', 12))
    
    # Interligne
    line_spacing_map = {
        '1.0': WD_LINE_SPACING.SINGLE,
        '1.15': WD_LINE_SPACING.ONE_POINT_FIVE,
        '1.5': WD_LINE_SPACING.ONE_POINT_FIVE,
        '2.0': WD_LINE_SPACING.DOUBLE
    }
    
    spacing = line_spacing_map.get(options.get('line_spacing', '1.5'), WD_LINE_SPACING.ONE_POINT_FIVE)
    normal_style.paragraph_format.line_spacing_rule = spacing
    
    # Styles personnalisés pour documents juridiques
    if options.get('legal_formatting', False):
        create_legal_styles(doc, options)

def create_legal_styles(doc, options: dict):
    """Crée les styles spécifiques aux documents juridiques"""
    
    styles = doc.styles
    
    # Style pour sections principales
    if 'LegalSection' not in [s.name for s in styles]:
        section_style = styles.add_style('LegalSection', WD_STYLE_TYPE.PARAGRAPH)
        section_style.base_style = styles['Heading 1']
        section_style.font.name = options.get('font', 'Times New Roman')
        section_style.font.size = Pt(14)
        section_style.font.bold = True
        section_style.font.color.rgb = RGBColor(0, 0, 0)
        section_style.paragraph_format.space_before = Pt(18)
        section_style.paragraph_format.space_after = Pt(12)
        section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Style pour sous-sections
    if 'LegalSubSection' not in [s.name for s in styles]:
        subsection_style = styles.add_style('LegalSubSection', WD_STYLE_TYPE.PARAGRAPH)
        subsection_style.base_style = styles['Heading 2']
        subsection_style.font.name = options.get('font', 'Times New Roman')
        subsection_style.font.size = Pt(12)
        subsection_style.font.bold = True
        subsection_style.font.underline = True
        subsection_style.paragraph_format.space_before = Pt(12)
        subsection_style.paragraph_format.space_after = Pt(6)
    
    # Style pour citations
    if 'LegalCitation' not in [s.name for s in styles]:
        citation_style = styles.add_style('LegalCitation', WD_STYLE_TYPE.PARAGRAPH)
        citation_style.font.name = options.get('font', 'Times New Roman')
        citation_style.font.size = Pt(10)
        citation_style.font.italic = True
        citation_style.paragraph_format.left_indent = Inches(0.5)
        citation_style.paragraph_format.right_indent = Inches(0.5)
        citation_style.paragraph_format.space_before = Pt(6)
        citation_style.paragraph_format.space_after = Pt(6)

def add_docx_title(doc, title: str, options: dict):
    """Ajoute le titre principal avec mise en forme"""
    
    # Titre centré
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Style du titre
    run = title_para.add_run(title.upper())
    run.font.name = options.get('font', 'Times New Roman')
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Espace après le titre
    doc.add_paragraph()

def add_docx_legal_content_advanced(doc, content: dict, options: dict):
    """Ajoute du contenu juridique avec mise en forme spécialisée"""
    
    lines = content['content'].split('\n')
    current_section = None
    
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Sections principales (I., II., etc.)
        if re.match(r'^[IVX]+\.\s+', line):
            if options.get('section_numbering', True):
                para = doc.add_paragraph(line, style='LegalSection')
            else:
                para = doc.add_paragraph(line.split('.', 1)[1].strip(), style='LegalSection')
            current_section = 'main'
        
        # Sous-sections (A., B., etc.)
        elif re.match(r'^[A-Z]\.\s+', line):
            if current_section == 'main':
                para = doc.add_paragraph(line, style='LegalSubSection')
            else:
                para = doc.add_paragraph(line)
                para.style.font.bold = True
        
        # Numéros (1., 2., etc.)
        elif re.match(r'^\d+\.\s+', line):
            para = doc.add_paragraph(line, style='List Number')
        
        # PAR CES MOTIFS
        elif line.strip().upper().startswith('PAR CES MOTIFS'):
            doc.add_page_break()
            para = doc.add_paragraph(line.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.style.font.bold = True
            para.style.font.size = Pt(14)
        
        # Citations
        elif line.strip().startswith('"') and line.strip().endswith('"'):
            doc.add_paragraph(line.strip(), style='LegalCitation')
        
        # Texte normal
        else:
            para = doc.add_paragraph(line.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Inches(0.5)

def add_docx_content_advanced(doc, content: dict, options: dict):
    """Ajoute le contenu avec mise en forme avancée"""
    
    # Traitement selon le type de contenu
    if content['type'] == 'timeline':
        add_docx_timeline_content(doc, content, options)
    elif content['type'] == 'mapping':
        add_docx_mapping_content(doc, content, options)
    elif content['type'] == 'comparison':
        add_docx_comparison_content(doc, content, options)
    else:
        # Contenu générique
        add_docx_generic_content(doc, content, options)

def add_docx_timeline_content(doc, content: dict, options: dict):
    """Ajoute une chronologie formatée"""
    
    doc.add_heading('Chronologie des événements', 1)
    
    # Table pour la chronologie
    if 'source_data' in content and 'events' in content['source_data']:
        events = content['source_data']['events']
        
        # Créer une table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light List'
        
        # En-têtes
        headers = ['Date', 'Événement', 'Acteurs/Source']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Mettre en gras
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Données
        for event in events:
            row = table.add_row()
            row.cells[0].text = str(event.get('date', 'N/A'))
            row.cells[1].text = event.get('description', '')
            
            details = []
            if event.get('actors'):
                details.append(f"Acteurs: {', '.join(event['actors'])}")
            if event.get('source'):
                details.append(f"Source: {event['source']}")
            row.cells[2].text = '\n'.join(details)
        
        # Ajuster la largeur des colonnes
        for row in table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(3.5)
            row.cells[2].width = Inches(2.5)
    else:
        # Fallback texte simple
        lines = content['content'].split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())

def export_to_pdf_advanced(content: dict, options: dict) -> bytes:
    """Export PDF avancé"""
    
    # Pour un vrai PDF, il faudrait utiliser reportlab ou weasyprint
    # Ici on génère un HTML optimisé pour l'impression PDF
    
    html_options = {
        **options,
        'print_friendly': True,
        'theme': 'Juridique'
    }
    
    html_content = export_to_html_advanced(content, html_options)
    
    # Ajouter des meta-tags pour l'impression PDF
    pdf_meta = """
    <style>
        @page {
            size: """ + options.get('page_size', 'A4') + """ """ + options.get('orientation', 'portrait').lower() + """;
            margin: """ + str(options.get('margins', 2)) + """cm;
        }
        
        @media print {
            body { margin: 0; }
            .no-print { display: none; }
            .page-break { page-break-after: always; }
        }
    </style>
    """
    
    # Insérer les styles d'impression
    html_str = html_content.decode('utf-8')
    html_str = html_str.replace('</head>', pdf_meta + '</head>')
    
    return html_str.encode('utf-8')

def export_to_txt_advanced(content: dict, options: dict) -> bytes:
    """Export texte avancé avec options"""
    
    # En-tête
    text = "=" * 80 + "\n"
    text += content['title'].upper().center(80) + "\n"
    text += "=" * 80 + "\n\n"
    
    # Métadonnées si demandées
    if options.get('include_metadata', True):
        text += "INFORMATIONS DU DOCUMENT\n"
        text += "-" * 40 + "\n"
        text += f"Date d'export    : {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n"
        text += f"Type de document : {content.get('type', 'N/A').title()}\n"
        
        if content.get('metadata'):
            for key, value in content['metadata'].items():
                if key not in ['content', 'responses', 'visualization', 'source_data'] and not isinstance(value, (dict, list)):
                    text += f"{key.replace('_', ' ').title():<16} : {value}\n"
        
        text += "\n" + "=" * 80 + "\n\n"
    
    # Table des matières si demandée et pertinente
    if options.get('include_toc', False) and content['type'] in ['redaction', 'analysis']:
        text += "TABLE DES MATIÈRES\n"
        text += "-" * 40 + "\n"
        
        # Extraire les sections
        sections = extract_sections_from_content(content['content'])
        for i, section in enumerate(sections, 1):
            text += f"{i}. {section}\n"
        
        text += "\n" + "=" * 80 + "\n\n"
    
    # Contenu principal
    text += "CONTENU\n"
    text += "-" * 40 + "\n\n"
    text += content['content']
    
    # Pied de document
    text += "\n\n" + "=" * 80 + "\n"
    text += f"Document généré automatiquement le {datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}\n"
    
    # Encodage avec gestion des erreurs
    return text.encode('utf-8', errors='replace')

def export_to_html_advanced(content: dict, options: dict) -> bytes:
    """Export HTML avancé avec thèmes et options"""
    
    # Thèmes disponibles
    themes = {
        'Clair': {
            'bg': '#ffffff',
            'text': '#333333',
            'primary': '#2c3e50',
            'secondary': '#34495e',
            'accent': '#3498db',
            'border': '#ecf0f1'
        },
        'Sombre': {
            'bg': '#1a1a1a',
            'text': '#e0e0e0',
            'primary': '#3498db',
            'secondary': '#2980b9',
            'accent': '#e74c3c',
            'border': '#2c2c2c'
        },
        'Juridique': {
            'bg': '#fafafa',
            'text': '#2c3e50',
            'primary': '#8b0000',
            'secondary': '#a52a2a',
            'accent': '#d4af37',
            'border': '#ddd'
        },
        'Minimaliste': {
            'bg': '#ffffff',
            'text': '#000000',
            'primary': '#000000',
            'secondary': '#666666',
            'accent': '#0066cc',
            'border': '#000000'
        },
        'Corporate': {
            'bg': '#f5f5f5',
            'text': '#333333',
            'primary': '#003366',
            'secondary': '#336699',
            'accent': '#ff6600',
            'border': '#cccccc'
        }
    }
    
    theme = themes.get(options.get('theme', 'Juridique'), themes['Juridique'])
    
    # Construction du HTML
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="generator" content="Legal Document Generator">
    <meta name="created" content="{datetime.now().isoformat()}">
    <title>{content['title']}</title>
"""
    
    # CSS si demandé
    if options.get('include_css', True):
        html += generate_advanced_css(theme, options)
    
    # JavaScript si demandé
    if options.get('include_js', False):
        html += generate_interactive_js(options)
    
    html += """</head>
<body>
"""
    
    # Navigation si multi-pages
    if not options.get('single_page', True):
        html += generate_navigation_menu(content)
    
    # Container principal
    html += f"""
    <div class="container">
        <header>
            <h1>{content['title']}</h1>
"""
    
    # Métadonnées
    if options.get('include_metadata', True):
        html += generate_html_metadata(content, theme)
    
    html += """
        </header>
        
        <main>
"""
    
    # Table des matières
    if options.get('include_toc', False):
        html += generate_html_toc(content)
    
    # Contenu principal
    html += generate_html_content(content, options)
    
    html += """
        </main>
        
        <footer>
            <p>Document généré le {0}</p>
        </footer>
    </div>
</body>
</html>""".format(datetime.now().strftime('%d/%m/%Y à %H:%M'))
    
    return html.encode('utf-8')

def generate_advanced_css(theme: dict, options: dict) -> str:
    """Génère le CSS avancé selon le thème et les options"""
    
    css = f"""
    <style>
        /* Reset et base */
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        /* Variables CSS */
        :root {{
            --bg-color: {theme['bg']};
            --text-color: {theme['text']};
            --primary-color: {theme['primary']};
            --secondary-color: {theme['secondary']};
            --accent-color: {theme['accent']};
            --border-color: {theme['border']};
            --font-family: 'Times New Roman', Georgia, serif;
            --font-size-base: 16px;
            --line-height: 1.6;
        }}
        
        /* Corps du document */
        body {{
            font-family: var(--font-family);
            font-size: var(--font-size-base);
            line-height: var(--line-height);
            color: var(--text-color);
            background-color: var(--bg-color);
            margin: 0;
            padding: 0;
        }}
        
        /* Container principal */
        .container {{
            max-width: {1200 if options.get('responsive', True) else 800}px;
            margin: 0 auto;
            padding: 20px;
            background-color: var(--bg-color);
        }}
        
        /* En-tête */
        header {{
            border-bottom: 3px solid var(--primary-color);
            padding-bottom: 20px;
            margin-bottom: 40px;
        }}
        
        h1 {{
            color: var(--primary-color);
            text-align: center;
            font-size: 2.5em;
            margin-bottom: 20px;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        
        h2 {{
            color: var(--primary-color);
            font-size: 1.8em;
            margin-top: 30px;
            margin-bottom: 15px;
            border-bottom: 2px solid var(--border-color);
            padding-bottom: 10px;
        }}
        
        h3 {{
            color: var(--secondary-color);
            font-size: 1.4em;
            margin-top: 20px;
            margin-bottom: 10px;
        }}
        
        /* Paragraphes */
        p {{
            text-align: justify;
            margin-bottom: 15px;
            text-indent: 2em;
        }}
        
        /* Métadonnées */
        .metadata {{
            background-color: rgba(0,0,0,0.05);
            border: 1px solid var(--border-color);
            border-radius: 5px;
            padding: 20px;
            margin-bottom: 30px;
        }}
        
        .metadata table {{
            width: 100%;
            border-collapse: collapse;
        }}
        
        .metadata td {{
            padding: 8px 12px;
            border-bottom: 1px solid var(--border-color);
        }}
        
        .metadata td:first-child {{
            font-weight: bold;
            width: 30%;
            color: var(--secondary-color);
        }}
        
        /* Table des matières */
        .toc {{
            background-color: rgba(0,0,0,0.03);
            border: 1px solid var(--border-color);
            border-radius: 5px;
            padding: 20px;
            margin-bottom: 40px;
        }}
        
        .toc h2 {{
            margin-top: 0;
            color: var(--primary-color);
        }}
        
        .toc ul {{
            list-style: none;
            padding-left: 0;
        }}
        
        .toc li {{
            margin-bottom: 10px;
            padding-left: 20px;
            position: relative;
        }}
        
        .toc li:before {{
            content: "▸";
            position: absolute;
            left: 0;
            color: var(--accent-color);
        }}
        
        /* Tables */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        
        th, td {{
            border: 1px solid var(--border-color);
            padding: 12px;
            text-align: left;
        }}
        
        th {{
            background-color: var(--primary-color);
            color: white;
            font-weight: bold;
        }}
        
        tr:nth-child(even) {{
            background-color: rgba(0,0,0,0.02);
        }}
        
        /* Listes */
        ul, ol {{
            margin-left: 40px;
            margin-bottom: 15px;
        }}
        
        li {{
            margin-bottom: 5px;
        }}
        
        /* Pied de page */
        footer {{
            margin-top: 60px;
            padding-top: 20px;
            border-top: 2px solid var(--border-color);
            text-align: center;
            color: var(--secondary-color);
            font-size: 0.9em;
        }}
        
        /* Impression */
        @media print {{
            body {{
                background-color: white;
                color: black;
            }}
            
            .container {{
                max-width: 100%;
                padding: 0;
            }}
            
            .no-print {{
                display: none !important;
            }}
            
            h1, h2, h3 {{
                color: black;
            }}
            
            .page-break {{
                page-break-after: always;
            }}
            
            @page {{
                margin: 2cm;
            }}
        }}
        
        /* Responsive */
        @media (max-width: 768px) {{
            .container {{
                padding: 10px;
            }}
            
            h1 {{
                font-size: 1.8em;
            }}
            
            h2 {{
                font-size: 1.4em;
            }}
            
            table {{
                font-size: 0.9em;
            }}
        }}
    </style>
"""
    
    return css

def generate_html_content(content: dict, options: dict) -> str:
    """Génère le contenu HTML formaté"""
    
    html_content = '<div class="content">\n'
    
    # Traitement spécifique selon le type
    if content['type'] == 'timeline':
        html_content += generate_timeline_html(content)
    elif content['type'] == 'mapping':
        html_content += generate_mapping_html(content)
    elif content['type'] == 'comparison':
        html_content += generate_comparison_html(content)
    else:
        # Contenu générique
        lines = content['content'].split('\n')
        
        in_list = False
        
        for line in lines:
            if not line.strip():
                if in_list:
                    html_content += '</ul>\n'
                    in_list = False
                continue
            
            # Détection des structures
            if line.strip().isupper() and len(line.strip()) > 5:
                html_content += f'<h2>{line.strip()}</h2>\n'
            
            elif re.match(r'^[IVX]+\.\s+', line.strip()):
                html_content += f'<h2>{line.strip()}</h2>\n'
            
            elif re.match(r'^[A-Z]\.\s+', line.strip()):
                html_content += f'<h3>{line.strip()}</h3>\n'
            
            elif line.strip().startswith(('- ', '• ', '* ')):
                if not in_list:
                    html_content += '<ul>\n'
                    in_list = True
                html_content += f'<li>{line.strip()[2:]}</li>\n'
            
            else:
                if in_list:
                    html_content += '</ul>\n'
                    in_list = False
                html_content += f'<p>{line.strip()}</p>\n'
        
        if in_list:
            html_content += '</ul>\n'
    
    html_content += '</div>\n'
    
    return html_content

def export_to_json_advanced(content: dict, options: dict) -> bytes:
    """Export JSON avancé avec structure complète"""
    
    # Structure de données complète
    export_data = {
        'metadata': {
            'title': content['title'],
            'type': content['type'],
            'export_date': datetime.now().isoformat(),
            'export_version': '2.0',
            'generator': 'Legal Document System'
        },
        'content': {
            'raw': content['content'],
            'formatted': None,
            'sections': []
        }
    }
    
    # Ajouter les métadonnées si demandées
    if options.get('include_metadata', True) and content.get('metadata'):
        export_data['document_metadata'] = {}
        
        for key, value in content['metadata'].items():
            if isinstance(value, (str, int, float, bool, list)):
                export_data['document_metadata'][key] = value
            elif isinstance(value, dict):
                # Filtrer les dictionnaires
                export_data['document_metadata'][key] = {
                    k: v for k, v in value.items()
                    if isinstance(v, (str, int, float, bool, list))
                }
            elif hasattr(value, 'isoformat'):
                export_data['document_metadata'][key] = value.isoformat()
            else:
                export_data['document_metadata'][key] = str(value)
    
    # Traitement spécifique selon le type
    if content['type'] == 'timeline' and 'source_data' in content:
        export_data['timeline_data'] = content['source_data'].get('events', [])
    
    elif content['type'] == 'mapping' and 'source_data' in content:
        export_data['mapping_data'] = {
            'entities': content['source_data'].get('entities', []),
            'relationships': content['source_data'].get('relationships', []),
            'analysis': content['source_data'].get('analysis', {})
        }
    
    elif content['type'] == 'comparison' and 'source_data' in content:
        export_data['comparison_data'] = content['source_data'].get('comparison', {})
    
    # Extraction des sections si applicable
    if options.get('include_toc', False):
        sections = extract_sections_from_content(content['content'])
        export_data['content']['sections'] = sections
    
    # Sérialisation avec gestion des erreurs
    json_str = json.dumps(
        export_data,
        indent=2,
        ensure_ascii=False,
        sort_keys=True,
        default=str  # Convertir les objets non sérialisables en string
    )
    
    return json_str.encode('utf-8')

def export_to_xlsx_advanced(content: dict, options: dict) -> bytes:
    """Export Excel avancé avec formatage et graphiques"""
    
    if not PANDAS_AVAILABLE:
        st.error("pandas requis pour l'export Excel")
        return export_to_csv_advanced(content, options)
    
    try:
        buffer = io.BytesIO()
        
        # Créer le writer avec options
        with pd.ExcelWriter(
            buffer,
            engine='openpyxl',
            datetime_format='DD/MM/YYYY',
            date_format='DD/MM/YYYY'
        ) as writer:
            
            # Export selon le type
            if content['type'] == 'timeline':
                export_timeline_excel_advanced(writer, content, options)
            
            elif content['type'] == 'mapping':
                export_mapping_excel_advanced(writer, content, options)
            
            elif content['type'] == 'comparison':
                export_comparison_excel_advanced(writer, content, options)
            
            elif content['type'] == 'search':
                export_search_excel_advanced(writer, content, options)
            
            elif content['type'] == 'pieces':
                export_pieces_excel_advanced(writer, content, options)
            
            else:
                # Export générique
                export_generic_excel_advanced(writer, content, options)
            
            # Ajouter une feuille de métadonnées
            if options.get('include_metadata', True):
                add_metadata_sheet_advanced(writer, content, options)
            
            # Appliquer le formatage global
            apply_excel_formatting(writer, options)
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur export Excel: {e}")
        return export_to_csv_advanced(content, options)

def export_timeline_excel_advanced(writer, content: dict, options: dict):
    """Export avancé de chronologie vers Excel"""
    
    if 'source_data' in content and 'events' in content['source_data']:
        events = content['source_data']['events']
        
        # Créer le DataFrame
        df = pd.DataFrame(events)
        
        # Formater les colonnes
        column_mapping = {
            'date': 'Date',
            'description': 'Description',
            'actors': 'Acteurs',
            'source': 'Source',
            'type': 'Type'
        }
        
        df.rename(columns=column_mapping, inplace=True)
        
        # Convertir les listes en strings
        if 'Acteurs' in df.columns:
            df['Acteurs'] = df['Acteurs'].apply(lambda x: ', '.join(x) if isinstance(x, list) else x)
        
        # Trier par date si possible
        if 'Date' in df.columns:
            try:
                df['Date'] = pd.to_datetime(df['Date'], format='%d/%m/%Y', errors='coerce')
                df.sort_values('Date', inplace=True)
            except:
                pass
        
        # Écrire dans Excel
        df.to_excel(writer, sheet_name='Chronologie', index=False)
        
        # Formater la feuille
        worksheet = writer.sheets['Chronologie']
        
        # Largeurs de colonnes
        column_widths = {
            'A': 15,  # Date
            'B': 50,  # Description
            'C': 30,  # Acteurs
            'D': 20,  # Source
            'E': 15   # Type
        }
        
        for col, width in column_widths.items():
            if col <= chr(64 + len(df.columns)):
                worksheet.column_dimensions[col].width = width
        
        # Mise en forme conditionnelle si demandée
        if options.get('conditional_formatting', False):
            # Colorer les dates importantes
            from openpyxl.formatting.rule import CellIsRule
            from openpyxl.styles import PatternFill
            
            red_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
            
            # Ajouter des règles selon les mots-clés
            worksheet.conditional_formatting.add(
                f'B2:B{len(df)+1}',
                CellIsRule(
                    operator='containsText',
                    formula=['important'],
                    stopIfTrue=True,
                    fill=red_fill
                )
            )

def compress_data_advanced(data: bytes, filename: str, level: int = 6) -> bytes:
    """Compression avancée avec niveau configurable"""
    
    buffer = io.BytesIO()
    
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED, compresslevel=level) as zip_file:
        # Ajouter le fichier principal
        zip_file.writestr(filename, data)
        
        # Ajouter un fichier README
        readme_content = f"""
Archive générée automatiquement
==============================

Fichier : {filename}
Date    : {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
Taille  : {len(data):,} octets

Cette archive a été créée par le système d'export de documents juridiques.
"""
        
        zip_file.writestr('README.txt', readme_content.encode('utf-8'))
    
    buffer.seek(0)
    return buffer.getvalue()

def split_export_data(data: bytes, max_size: int) -> List[bytes]:
    """Divise les données en plusieurs parties"""
    
    parts = []
    
    for i in range(0, len(data), max_size):
        part = data[i:i + max_size]
        parts.append(part)
    
    return parts

def display_export_statistics(data: bytes, format: str, content: dict):
    """Affiche des statistiques détaillées sur l'export"""
    
    size_bytes = len(data)
    size_kb = size_bytes / 1024
    size_mb = size_kb / 1024
    
    # Métriques principales
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📄 Format", format.upper())
    
    with col2:
        if size_mb > 1:
            st.metric("📊 Taille", f"{size_mb:.2f} MB")
        else:
            st.metric("📊 Taille", f"{size_kb:.2f} KB")
    
    with col3:
        # Estimation selon le format
        if format == 'docx':
            pages = int(size_kb / 3)  # Approximation
            st.metric("📑 Pages", f"~{pages}")
        elif format == 'txt':
            chars = len(content.get('content', ''))
            st.metric("✏️ Caractères", f"{chars:,}")
        elif format == 'xlsx':
            st.metric("📊 Type", content['type'].title())
    
    with col4:
        compression_ratio = 1 - (size_bytes / len(content.get('content', 'x').encode('utf-8')))
        if compression_ratio > 0:
            st.metric("🗜️ Compression", f"{compression_ratio:.0%}")

def log_export_action(content: dict, format: str, filename: str):
    """Enregistre l'action d'export dans les logs"""
    
    if 'export_history' not in st.session_state:
        st.session_state.export_history = []
    
    st.session_state.export_history.append({
        'timestamp': datetime.now(),
        'content_type': content['type'],
        'format': format,
        'filename': filename,
        'size': len(content.get('content', ''))
    })

# Fonctions helper additionnelles

def extract_sections_from_content(content: str) -> List[str]:
    """Extrait les sections principales du contenu"""
    
    sections = []
    
    # Patterns pour détecter les sections
    patterns = [
        r'^[IVX]+\.\s+(.+)$',  # I. Section
        r'^#{1,3}\s+(.+)$',     # # Markdown
        r'^[A-Z][A-Z\s]+$'      # TITRE EN MAJUSCULES
    ]
    
    lines = content.split('\n')
    
    for line in lines:
        for pattern in patterns:
            match = re.match(pattern, line.strip())
            if match:
                section_title = match.group(1) if match.groups() else line.strip()
                sections.append(section_title.strip())
                break
    
    return sections

def clean_key(text: str) -> str:
    """Nettoie une chaîne pour utilisation comme clé/nom de fichier"""
    
    # Remplacer les caractères spéciaux
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[-\s]+', '_', text)
    
    return text.lower()

# Formatters pour différents types de contenu

def format_redaction_content(data: dict) -> dict:
    """Formate le contenu d'une rédaction"""
    
    return {
        'title': f"{data.get('type', 'Document').title()}",
        'content': data.get('document', ''),
        'metadata': {
            'type': data.get('type'),
            'timestamp': data.get('timestamp'),
            'reference': data.get('reference'),
            'quality_score': data.get('quality_score')
        }
    }

def format_timeline_content(data: dict) -> dict:
    """Formate le contenu d'une chronologie"""
    
    content = f"CHRONOLOGIE - {data.get('type', '').upper()}\n"
    content += f"Générée le {data.get('timestamp', datetime.now()).strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Basée sur {data.get('document_count', 0)} documents\n"
    content += "=" * 60 + "\n\n"
    
    for event in data.get('events', []):
        content += f"{event.get('date', 'N/A')} - {event.get('description', '')}\n"
        
        if event.get('actors'):
            content += f"   Acteurs: {', '.join(event['actors'])}\n"
        
        if event.get('source'):
            content += f"   Source: {event['source']}\n"
        
        content += "\n"
    
    return {
        'title': f"Chronologie {data.get('type', '')}",
        'content': content,
        'metadata': data
    }

def format_mapping_content(data: dict) -> dict:
    """Formate le contenu d'une cartographie"""
    
    content = f"CARTOGRAPHIE DES RELATIONS - {data.get('type', '').upper()}\n"
    content += f"Générée le {data.get('timestamp', datetime.now()).strftime('%d/%m/%Y à %H:%M')}\n"
    content += "=" * 60 + "\n\n"
    
    # Statistiques
    if 'analysis' in data:
        analysis = data['analysis']
        content += "STATISTIQUES:\n"
        content += f"- Nombre d'entités: {analysis.get('node_count', 0)}\n"
        content += f"- Nombre de relations: {analysis.get('edge_count', 0)}\n"
        content += f"- Densité du réseau: {analysis.get('density', 0):.2%}\n\n"
        
        if analysis.get('key_players'):
            content += "ACTEURS PRINCIPAUX:\n"
            for i, player in enumerate(analysis['key_players'], 1):
                content += f"{i}. {player}\n"
    
    return {
        'title': f"Cartographie {data.get('type', '')}",
        'content': content,
        'metadata': data
    }

def format_comparison_content(data: dict) -> dict:
    """Formate le contenu d'une comparaison"""
    
    content = f"COMPARAISON - {data.get('type', '').upper()}\n"
    content += f"Générée le {data.get('timestamp', datetime.now()).strftime('%d/%m/%Y à %H:%M')}\n"
    content += f"Documents comparés: {data.get('document_count', 0)}\n"
    content += "=" * 60 + "\n\n"
    
    comparison = data.get('comparison', {})
    
    # Convergences
    if comparison.get('convergences'):
        content += "POINTS DE CONVERGENCE:\n"
        content += "-" * 40 + "\n"
        for conv in comparison['convergences']:
            content += f"- {conv.get('point', '')}\n"
            content += f"  {conv.get('details', '')}\n\n"
    
    # Divergences
    if comparison.get('divergences'):
        content += "\nPOINTS DE DIVERGENCE:\n"
        content += "-" * 40 + "\n"
        for div in comparison['divergences']:
            content += f"- {div.get('point', '')}\n"
            content += f"  {div.get('details', '')}\n\n"
    
    return {
        'title': f"Comparaison {data.get('type', '')}",
        'content': content,
        'metadata': data
    }

def format_analysis_content(data: dict) -> dict:
    """Formate le contenu d'une analyse"""
    
    return {
        'title': "Analyse IA",
        'content': data.get('content', ''),
        'metadata': {
            'type': data.get('type', 'Générale'),
            'timestamp': data.get('timestamp', datetime.now()),
            'document_count': data.get('document_count', 0),
            'query': data.get('query')
        }
    }

def format_search_content(data: list) -> dict:
    """Formate les résultats de recherche"""
    
    content = f"RÉSULTATS DE RECHERCHE\n"
    content += f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Nombre de résultats: {len(data)}\n"
    content += "=" * 60 + "\n\n"
    
    for i, result in enumerate(data, 1):
        content += f"{i}. {result.get('title', 'Sans titre')}\n"
        content += f"   Source: {result.get('source', 'N/A')}\n"
        content += f"   Pertinence: {result.get('score', 0):.0%}\n"
        if result.get('content'):
            content += f"   Aperçu: {result['content'][:200]}...\n"
        content += "\n"
    
    return {
        'title': "Résultats de recherche",
        'content': content,
        'metadata': {
            'count': len(data),
            'timestamp': datetime.now()
        }
    }

def format_pieces_content(data: list) -> dict:
    """Formate une liste de pièces"""
    
    content = f"LISTE DES PIÈCES\n"
    content += f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Nombre de pièces: {len(data)}\n"
    content += "=" * 60 + "\n\n"
    
    # Grouper par catégorie
    by_category = defaultdict(list)
    for piece in data:
        by_category[piece.categorie].append(piece)
    
    for category, pieces in sorted(by_category.items()):
        content += f"\n{category.upper()} ({len(pieces)} pièces)\n"
        content += "-" * 40 + "\n"
        
        for piece in pieces:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            content += "\n"
    
    return {
        'title': "Liste des pièces",
        'content': content,
        'metadata': {
            'count': len(data),
            'timestamp': datetime.now()
        }
    }

# Point d'entrée principal
def show_import_export_interface():
    """Interface principale unifiée d'import/export"""
    
    st.markdown("### 📥📤 Import/Export de documents")
    
    tab1, tab2, tab3 = st.tabs(["📥 Import", "📤 Export", "📊 Historique"])
    
    with tab1:
        process_import_request("", {})
    
    with tab2:
        process_export_request("", {})
    
    with tab3:
        show_import_export_history()

def show_import_export_history():
    """Affiche l'historique des imports/exports"""
    
    st.markdown("### 📊 Historique des opérations")
    
    # Historique des exports
    if st.session_state.get('export_history'):
        st.markdown("#### 📤 Exports récents")
        
        df_exports = pd.DataFrame(st.session_state.export_history)
        df_exports['timestamp'] = pd.to_datetime(df_exports['timestamp'])
        df_exports = df_exports.sort_values('timestamp', ascending=False)
        
        # Afficher le tableau
        st.dataframe(
            df_exports[['timestamp', 'content_type', 'format', 'filename', 'size']].head(10),
            use_container_width=True
        )
    
    # Statistiques
    col1, col2 = st.columns(2)
    
    with col1:
        if st.session_state.get('azure_documents'):
            st.metric("📄 Documents importés", len(st.session_state.azure_documents))
    
    with col2:
        if st.session_state.get('export_history'):
            st.metric("📤 Exports effectués", len(st.session_state.export_history))