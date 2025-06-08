# modules/recherche.py
"""Module de recherche unifi√©e avec 100% des fonctionnalit√©s int√©gr√©es"""

import streamlit as st
import asyncio
from datetime import datetime, timedelta
import re
import json
from collections import defaultdict, Counter
from typing import List, Dict, Optional, Any, Tuple, Union
import io
import base64
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

# Imports optionnels avec gestion d'erreur
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("‚ö†Ô∏è pandas non install√© - Certaines fonctionnalit√©s seront limit√©es")

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False
    print("‚ö†Ô∏è plotly non install√© - Visualisations simplifi√©es")

try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False
    print("‚ö†Ô∏è networkx non install√© - Analyse r√©seau limit√©e")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("‚ö†Ô∏è python-docx non install√© - Export Word non disponible")

# Imports des modules internes
from config.app_config import SearchMode, ANALYSIS_PROMPTS_AFFAIRES, LLMProvider
from managers.azure_blob_manager import AzureBlobManager
from managers.azure_search_manager import AzureSearchManager
from managers.multi_llm_manager import MultiLLMManager
from managers.document_manager import display_import_interface
from managers.dynamic_generators import generate_dynamic_search_prompts
from managers.legal_search import LegalSearchManager, display_legal_search_interface
from managers.jurisprudence_verifier import JurisprudenceVerifier
from models.dataclasses import Document, PieceSelectionnee
from utils.helpers import clean_key

# Configuration des styles de r√©daction
REDACTION_STYLES = {
    'formel': {
        'name': 'Formel',
        'description': 'Style juridique classique et solennel',
        'tone': 'respectueux et distant',
        'vocabulary': 'technique et pr√©cis'
    },
    'persuasif': {
        'name': 'Persuasif',
        'description': 'Style argumentatif et convaincant',
        'tone': 'assertif et engag√©',
        'vocabulary': 'percutant et imag√©'
    },
    'technique': {
        'name': 'Technique',
        'description': 'Style factuel et d√©taill√©',
        'tone': 'neutre et objectif',
        'vocabulary': 'sp√©cialis√© et exhaustif'
    },
    'synth√©tique': {
        'name': 'Synth√©tique',
        'description': 'Style concis et efficace',
        'tone': 'direct et clair',
        'vocabulary': 'simple et pr√©cis'
    },
    'p√©dagogique': {
        'name': 'P√©dagogique',
        'description': 'Style explicatif et accessible',
        'tone': 'bienveillant et didactique',
        'vocabulary': 'vulgaris√© et illustr√©'
    }
}

# Templates pr√©d√©finis
DOCUMENT_TEMPLATES = {
    'conclusions_defense': {
        'name': 'Conclusions en d√©fense',
        'structure': [
            'I. FAITS ET PROC√âDURE',
            'II. DISCUSSION',
            '   A. Sur la recevabilit√©',
            '   B. Sur le fond',
            '      1. Sur l\'√©l√©ment mat√©riel',
            '      2. Sur l\'√©l√©ment intentionnel',
            '      3. Sur le pr√©judice',
            'III. PAR CES MOTIFS'
        ],
        'style': 'formel'
    },
    'plainte_simple': {
        'name': 'Plainte simple',
        'structure': [
            'Objet : Plainte',
            'EXPOS√â DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PR√âJUDICES SUBIS',
            'DEMANDES',
            'PI√àCES JOINTES'
        ],
        'style': 'formel'
    },
    'plainte_avec_cpc': {
        'name': 'Plainte avec constitution de partie civile',
        'structure': [
            'Objet : Plainte avec constitution de partie civile',
            'EXPOS√â DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PR√âJUDICES SUBIS',
            'CONSTITUTION DE PARTIE CIVILE',
            '√âVALUATION DU PR√âJUDICE',
            'DEMANDES',
            'PI√àCES JOINTES'
        ],
        'style': 'formel'
    },
    'mise_en_demeure': {
        'name': 'Mise en demeure',
        'structure': [
            'MISE EN DEMEURE',
            'Rappel des faits',
            'Obligations non respect√©es',
            'D√©lai accord√©',
            'Cons√©quences du d√©faut',
            'R√©serves'
        ],
        'style': 'persuasif'
    },
    'note_synthese': {
        'name': 'Note de synth√®se',
        'structure': [
            'SYNTH√àSE EX√âCUTIVE',
            'I. CONTEXTE',
            'II. ANALYSE',
            'III. RECOMMANDATIONS',
            'IV. PLAN D\'ACTION'
        ],
        'style': 'synth√©tique'
    }
}

# === FONCTIONS DE TRAITEMENT POUR IMPORT ===

def process_import_request(query: str, analysis: dict):
    """Traite une demande d'import"""
    
    st.info("üì• Interface d'import activ√©e")
    
    # Configuration de l'import depuis la config
    file_types = analysis['details'].get('file_types', ['pdf', 'docx', 'txt'])
    destination = st.session_state.get('import_destination', 'Documents locaux')
    
    # Si des fichiers sont d√©j√† upload√©s
    if st.session_state.get('import_files'):
        with st.spinner("üì• Import des fichiers en cours..."):
            imported_count = 0
            
            for file in st.session_state.import_files:
                try:
                    # Lire le contenu
                    content = file.read()
                    
                    # Traiter selon le type
                    if file.name.endswith('.pdf'):
                        text_content = extract_text_from_pdf(content)
                    elif file.name.endswith('.docx'):
                        text_content = extract_text_from_docx(content)
                    else:
                        text_content = content.decode('utf-8', errors='ignore')
                    
                    # Stocker selon la destination
                    if destination == 'Azure Blob':
                        store_in_azure_blob(file.name, content, analysis.get('reference'))
                    else:
                        store_locally(file.name, text_content, analysis.get('reference'))
                    
                    imported_count += 1
                    
                except Exception as e:
                    st.error(f"‚ùå Erreur import {file.name}: {str(e)}")
            
            st.success(f"‚úÖ {imported_count} fichiers import√©s avec succ√®s")
            
            # Analyse automatique si demand√©e
            if st.session_state.get('auto_analyze_import', True):
                st.info("ü§ñ Analyse automatique des documents import√©s...")
                analyze_imported_documents()
    
    else:
        st.warning("‚ö†Ô∏è Aucun fichier s√©lectionn√©. Utilisez l'interface d'import ci-dessous.")

def extract_text_from_pdf(content: bytes) -> str:
    """Extrait le texte d'un PDF"""
    try:
        import PyPDF2
        import io
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text
        
    except ImportError:
        return "PDF - Installation de PyPDF2 requise pour l'extraction"
    except Exception as e:
        return f"Erreur extraction PDF: {str(e)}"

def extract_text_from_docx(content: bytes) -> str:
    """Extrait le texte d'un DOCX"""
    if DOCX_AVAILABLE:
        try:
            import io
            doc = docx.Document(io.BytesIO(content))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"Erreur extraction DOCX: {str(e)}"
    else:
        return "DOCX - Installation de python-docx requise"

def store_locally(filename: str, content: str, reference: str = None):
    """Stocke un document localement"""
    doc_id = f"imported_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{clean_key(filename)}"
    
    if 'azure_documents' not in st.session_state:
        st.session_state.azure_documents = {}
    
    st.session_state.azure_documents[doc_id] = Document(
        id=doc_id,
        title=filename,
        content=content,
        source=f"Import {reference if reference else 'direct'}",
        metadata={
            'imported_at': datetime.now().isoformat(),
            'reference': reference,
            'original_filename': filename
        }
    )

def store_in_azure_blob(filename: str, content: bytes, reference: str = None):
    """Stocke un document dans Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        container = reference if reference else 'imports'
        blob_name = f"{datetime.now().strftime('%Y%m%d')}_{filename}"
        
        try:
            blob_manager.upload_blob(container, blob_name, content)
            st.success(f"‚úÖ {filename} upload√© dans Azure")
        except Exception as e:
            st.error(f"‚ùå Erreur Azure: {str(e)}")
    else:
        st.warning("‚ö†Ô∏è Azure Blob non connect√© - Stockage local")
        store_locally(filename, content.decode('utf-8', errors='ignore'), reference)

def analyze_imported_documents():
    """Analyse automatique des documents import√©s"""
    # R√©cup√©rer les derniers documents import√©s
    recent_docs = []
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if 'imported_at' in doc.metadata:
            import_time = datetime.fromisoformat(doc.metadata['imported_at'])
            if (datetime.now() - import_time).seconds < 300:  # Moins de 5 minutes
                recent_docs.append(doc)
    
    if recent_docs:
        # Lancer une analyse automatique
        analysis_prompt = f"""Analyse ces {len(recent_docs)} documents r√©cemment import√©s.
        
Identifie :
1. Type de document
2. Parties impliqu√©es
3. Dates importantes
4. Points juridiques cl√©s
5. Risques identifi√©s
Documents :
{chr(10).join([f"- {doc.title}: {doc.content[:500]}..." for doc in recent_docs[:5]])}"""
        
        llm_manager = MultiLLMManager()
        if llm_manager.clients:
            provider = list(llm_manager.clients.keys())[0]
            response = llm_manager.query_single_llm(
                provider,
                analysis_prompt,
                "Tu es un expert en analyse documentaire juridique."
            )
            
            if response['success']:
                st.session_state.import_analysis = response['response']
                st.success("‚úÖ Analyse des imports termin√©e")

# === FONCTIONS DE TRAITEMENT POUR EXPORT ===

def process_export_request(query: str, analysis: dict):
    """Traite une demande d'export"""
    
    format = analysis['details'].get('format', 'docx')
    
    # D√©terminer ce qu'il faut exporter
    content_to_export = None
    filename_base = "export"
    
    # Priorit√© : r√©daction > analyse > recherche > s√©lection
    if st.session_state.get('redaction_result'):
        content_to_export = st.session_state.redaction_result['document']
        filename_base = f"{st.session_state.redaction_result['type']}"
    
    elif st.session_state.get('timeline_result'):
        content_to_export = export_timeline_content(st.session_state.timeline_result)
        filename_base = "chronologie"
    
    elif st.session_state.get('mapping_result'):
        content_to_export = export_mapping_content(st.session_state.mapping_result)
        filename_base = "cartographie"
    
    elif st.session_state.get('comparison_result'):
        content_to_export = export_comparison_content(st.session_state.comparison_result)
        filename_base = "comparaison"
    
    elif st.session_state.get('ai_analysis_results'):
        content_to_export = format_ai_analysis_for_export(st.session_state.ai_analysis_results)
        filename_base = "analyse"
    
    elif st.session_state.get('search_results'):
        content_to_export = format_search_results_for_export(st.session_state.search_results)
        filename_base = "recherche"
    
    else:
        st.warning("‚ö†Ô∏è Aucun contenu √† exporter")
        return
    
    # Exporter selon le format
    export_functions = {
        'docx': export_to_docx,
        'pdf': export_to_pdf,
        'txt': export_to_txt,
        'html': export_to_html,
        'xlsx': export_to_xlsx
    }
    
    export_func = export_functions.get(format, export_to_txt)
    
    try:
        exported_data = export_func(content_to_export, analysis)
        
        # Cr√©er le bouton de t√©l√©chargement
        st.download_button(
            f"üíæ T√©l√©charger {format.upper()}",
            exported_data,
            f"{filename_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}",
            get_mime_type(format),
            key=f"download_export_{format}"
        )
        
        st.success(f"‚úÖ Export {format.upper()} pr√™t")
        
    except Exception as e:
        st.error(f"‚ùå Erreur export: {str(e)}")

def export_timeline_content(timeline_result: dict) -> str:
    """Exporte le contenu d'une chronologie"""
    content = f"CHRONOLOGIE - {timeline_result['type'].upper()}\n"
    content += f"G√©n√©r√©e le {timeline_result['timestamp'].strftime('%d/%m/%Y √† %H:%M')}\n"
    content += f"Bas√©e sur {timeline_result['document_count']} documents\n\n"
    
    for event in timeline_result['events']:
        content += f"{event.get('date', 'N/A')} - {event.get('description', '')}\n"
        if 'actors' in event:
            content += f"   Acteurs: {', '.join(event['actors'])}\n"
        if 'source' in event:
            content += f"   Source: {event['source']}\n"
        content += "\n"
    
    return content

def export_mapping_content(mapping_result: dict) -> str:
    """Exporte le contenu d'une cartographie"""
    content = f"CARTOGRAPHIE DES RELATIONS - {mapping_result['type'].upper()}\n"
    content += f"G√©n√©r√©e le {mapping_result['timestamp'].strftime('%d/%m/%Y √† %H:%M')}\n\n"
    
    content += f"STATISTIQUES:\n"
    content += f"- Entit√©s: {mapping_result['analysis']['node_count']}\n"
    content += f"- Relations: {mapping_result['analysis']['edge_count']}\n"
    content += f"- Densit√©: {mapping_result['analysis']['density']:.2%}\n\n"
    
    content += "ACTEURS PRINCIPAUX:\n"
    for i, player in enumerate(mapping_result['analysis']['key_players'], 1):
        content += f"{i}. {player}\n"
    
    content += "\nENTIT√âS:\n"
    for entity in mapping_result['entities']:
        content += f"- {entity['name']} ({entity.get('type', 'N/A')})\n"
    
    content += "\nRELATIONS:\n"
    for relation in mapping_result['relations']:
        content += f"- {relation['source']} ‚Üí {relation['target']}: {relation.get('type', 'N/A')}\n"
    
    return content

def export_comparison_content(comparison_result: dict) -> str:
    """Exporte le contenu d'une comparaison"""
    content = f"COMPARAISON - {comparison_result['type'].upper()}\n"
    content += f"G√©n√©r√©e le {comparison_result['timestamp'].strftime('%d/%m/%Y √† %H:%M')}\n"
    content += f"Documents compar√©s: {comparison_result['document_count']}\n\n"
    
    comparison = comparison_result['comparison']
    
    content += "CONVERGENCES:\n"
    for conv in comparison.get('convergences', []):
        content += f"- {conv['point']}\n  {conv['details']}\n\n"
    
    content += "DIVERGENCES:\n"
    for div in comparison.get('divergences', []):
        content += f"- {div['point']}\n  {div['details']}\n\n"
    
    if comparison.get('analysis'):
        content += f"ANALYSE:\n{comparison['analysis']}\n"
    
    return content

def export_to_docx(content: str, analysis: dict) -> bytes:
    """Exporte vers Word avec mise en forme"""
    if not DOCX_AVAILABLE:
        # Fallback en texte brut
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Styles personnalis√©s
        styles = doc.styles
        
        # En-t√™te
        if st.session_state.get('export_metadata', True):
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Document g√©n√©r√© le {datetime.now().strftime('%d/%m/%Y √† %H:%M')}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Titre principal
        title = doc.add_heading(analysis.get('document_type', 'Document').upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenu avec mise en forme
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            # D√©tecter les niveaux de titre
            if line.strip().isupper() and len(line.strip()) > 5:
                doc.add_heading(line.strip(), 1)
            elif line.strip().startswith('###'):
                doc.add_heading(line.strip().replace('#', '').strip(), 3)
            elif line.strip().startswith('##'):
                doc.add_heading(line.strip().replace('#', '').strip(), 2)
            elif line.strip().startswith('#'):
                doc.add_heading(line.strip().replace('#', '').strip(), 1)
            elif line.strip().startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
                p = doc.add_paragraph()
                p.add_run(line.strip()).bold = True
                p.style.font.size = Pt(14)
            elif line.strip().startswith(('A.', 'B.', 'C.', 'D.')):
                p = doc.add_paragraph(line.strip(), style='List Number')
                p.style.font.size = Pt(12)
            elif line.strip().startswith('-'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(line.strip())
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
        
        # Table des mati√®res si demand√©e
        if st.session_state.get('docx_toc', True):
            # Ins√©rer au d√©but
            doc.paragraphs[0].insert_paragraph_before("TABLE DES MATI√àRES", style='Heading 1')
            # Note: La vraie TOC n√©cessite des manipulations XML complexes
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur cr√©ation DOCX: {e}")
        return content.encode('utf-8')

def export_to_pdf(content: str, analysis: dict) -> bytes:
    """Exporte vers PDF"""
    # N√©cessiterait reportlab ou weasyprint
    st.warning("‚ö†Ô∏è Export PDF n√©cessite l'installation de biblioth√®ques suppl√©mentaires")
    
    # Pour l'instant, retourner le contenu texte
    return content.encode('utf-8')

def export_to_txt(content: str, analysis: dict) -> bytes:
    """Exporte en texte brut"""
    return content.encode('utf-8')

def export_to_html(content: str, analysis: dict) -> bytes:
    """Exporte vers HTML avec style"""
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <title>{analysis.get('document_type', 'Document').title()}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; margin: 40px; }}
        h1 {{ text-align: center; color: #2c3e50; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        p {{ text-align: justify; line-height: 1.6; }}
        .metadata {{ text-align: right; color: #95a5a6; font-size: 0.9em; }}
        .section {{ margin: 20px 0; }}
        ul, ol {{ margin-left: 20px; }}
    </style>
</head>
<body>
    <div class="metadata">G√©n√©r√© le {datetime.now().strftime('%d/%m/%Y √† %H:%M')}</div>
    <h1>{analysis.get('document_type', 'Document').upper()}</h1>
    <div class="content">
"""
    
    # Convertir le contenu en HTML
    lines = content.split('\n')
    in_list = False
    
    for line in lines:
        if not line.strip():
            if in_list:
                html += "</ul>"
                in_list = False
            continue
        
        if line.strip().isupper() and len(line.strip()) > 5:
            html += f"<h2>{line.strip()}</h2>\n"
        elif line.strip().startswith('-'):
            if not in_list:
                html += "<ul>\n"
                in_list = True
            html += f"<li>{line.strip()[1:].strip()}</li>\n"
        else:
            if in_list:
                html += "</ul>\n"
                in_list = False
            html += f"<p>{line.strip()}</p>\n"
    
    html += """
    </div>
</body>
</html>"""
    
    return html.encode('utf-8')

def export_to_xlsx(content: str, analysis: dict) -> bytes:
    """Exporte vers Excel"""
    if not PANDAS_AVAILABLE:
        st.error("‚ùå pandas requis pour l'export Excel")
        return content.encode('utf-8')
    
    try:
        import io
        
        # Cr√©er un DataFrame selon le type de contenu
        if 'timeline_result' in st.session_state:
            df = pd.DataFrame(st.session_state.timeline_result['events'])
        elif 'mapping_result' in st.session_state:
            df_entities = pd.DataFrame(st.session_state.mapping_result['entities'])
            df_relations = pd.DataFrame(st.session_state.mapping_result['relations'])
        else:
            # Contenu g√©n√©rique
            lines = content.split('\n')
            df = pd.DataFrame({'Contenu': lines})
        
        # Exporter vers Excel
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            if 'mapping_result' in st.session_state:
                df_entities.to_excel(writer, sheet_name='Entit√©s', index=False)
                df_relations.to_excel(writer, sheet_name='Relations', index=False)
            else:
                df.to_excel(writer, sheet_name='Donn√©es', index=False)
            
            # Ajouter une feuille de m√©tadonn√©es
            metadata = pd.DataFrame({
                'Propri√©t√©': ['Type', 'Date g√©n√©ration', 'Source'],
                'Valeur': [
                    analysis.get('document_type', 'Export'),
                    datetime.now().strftime('%d/%m/%Y %H:%M'),
                    analysis.get('reference', 'N/A')
                ]
            })
            metadata.to_excel(writer, sheet_name='M√©tadonn√©es', index=False)
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur export Excel: {e}")
        return content.encode('utf-8')

def get_mime_type(format: str) -> str:
    """Retourne le type MIME pour un format"""
    mime_types = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'html': 'text/html',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    }
    return mime_types.get(format, 'application/octet-stream')

# === FONCTIONS DE TRAITEMENT POUR EMAIL ===

def process_email_request(query: str, analysis: dict):
    """Traite une demande d'envoi par email"""
    
    recipients = analysis.get('recipients', [])
    
    if not recipients:
        st.error("‚ùå Aucun destinataire sp√©cifi√©")
        return
    
    # Pr√©parer le contenu
    content = prepare_email_content(analysis)
    
    if not content:
        st.warning("‚ö†Ô∏è Aucun contenu √† envoyer")
        return
    
    # Configuration email
    email_config = {
        'to': st.session_state.get('email_to', ', '.join(recipients)),
        'cc': st.session_state.get('email_cc', ''),
        'subject': st.session_state.get('email_subject', 'Document juridique'),
        'body': content['body'],
        'attachments': []
    }
    
    # Pr√©parer les pi√®ces jointes
    if st.session_state.get('email_attach_current', True):
        attachment_format = st.session_state.get('email_attachment_format', 'pdf')
        attachment_data = export_current_content(attachment_format)
        
        if attachment_data:
            email_config['attachments'].append({
                'filename': f"document.{attachment_format}",
                'data': attachment_data,
                'mime_type': get_mime_type(attachment_format)
            })
    
    # Afficher l'aper√ßu
    show_email_preview(email_config)
    
    # Bouton d'envoi
    if st.button("üìß Envoyer l'email", key="send_email_button"):
        if send_email(email_config):
            st.success("‚úÖ Email envoy√© avec succ√®s")
            st.session_state.email_sent = True
        else:
            st.error("‚ùå Erreur lors de l'envoi")

def prepare_email_content(analysis: dict) -> dict:
    """Pr√©pare le contenu de l'email"""
    content = {'body': '', 'attachments': []}
    
    # Corps de l'email selon le contexte
    if st.session_state.get('redaction_result'):
        result = st.session_state.redaction_result
        content['body'] = f"""Bonjour,
Veuillez trouver ci-joint le document {result['type']} demand√©.
Ce document a √©t√© g√©n√©r√© automatiquement √† partir de l'analyse de votre dossier.
Cordialement,
[Votre nom]"""
        
    else:
        content['body'] = """Bonjour,
Veuillez trouver ci-joint le document demand√©.
Cordialement,
[Votre nom]"""
    
    return content

def show_email_preview(email_config: dict):
    """Affiche un aper√ßu de l'email"""
    with st.expander("üìß Aper√ßu de l'email", expanded=True):
        st.text_input("√Ä:", value=email_config['to'], disabled=True)
        st.text_input("Cc:", value=email_config['cc'], disabled=True)
        st.text_input("Objet:", value=email_config['subject'], disabled=True)
        st.text_area("Corps:", value=email_config['body'], height=200, disabled=True)
        
        if email_config['attachments']:
            st.write("üìé Pi√®ces jointes:")
            for att in email_config['attachments']:
                st.write(f"- {att['filename']}")

def send_email(email_config: dict) -> bool:
    """Envoie l'email (fonction simplifi√©e)"""
    try:
        # Configuration SMTP (√† adapter)
        smtp_config = {
            'server': st.secrets.get('smtp_server', 'smtp.gmail.com'),
            'port': st.secrets.get('smtp_port', 587),
            'username': st.secrets.get('smtp_username', ''),
            'password': st.secrets.get('smtp_password', '')
        }
        
        if not smtp_config['username']:
            st.error("‚ùå Configuration email manquante")
            return False
        
        # Cr√©er le message
        msg = MIMEMultipart()
        msg['From'] = smtp_config['username']
        msg['To'] = email_config['to']
        msg['Cc'] = email_config['cc']
        msg['Subject'] = email_config['subject']
        
        # Corps
        msg.attach(MIMEText(email_config['body'], 'plain'))
        
        # Pi√®ces jointes
        for att in email_config['attachments']:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(att['data'])
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {att["filename"]}'
            )
            msg.attach(part)
        
        # Envoyer
        with smtplib.SMTP(smtp_config['server'], smtp_config['port']) as server:
            server.starttls()
            server.login(smtp_config['username'], smtp_config['password'])
            
            recipients = [email_config['to']]
            if email_config['cc']:
                recipients.extend(email_config['cc'].split(','))
            
            server.send_message(msg)
        
        return True
        
    except Exception as e:
        st.error(f"Erreur envoi email: {e}")
        return False

def export_current_content(format: str) -> bytes:
    """Exporte le contenu actuel dans le format demand√©"""
    # R√©utiliser les fonctions d'export
    analysis = st.session_state.get('current_analysis', {})
    
    if st.session_state.get('redaction_result'):
        content = st.session_state.redaction_result['document']
    else:
        content = "Aucun contenu disponible"
    
    export_functions = {
        'pdf': export_to_pdf,
        'docx': export_to_docx,
        'txt': export_to_txt
    }
    
    export_func = export_functions.get(format, export_to_txt)
    return export_func(content, analysis)

# === FONCTIONS DE TRAITEMENT POUR S√âLECTION DE PI√àCES ===

def process_piece_selection_request(query: str, analysis: dict):
    """Traite une demande de s√©lection de pi√®ces"""
    
    # Interface de s√©lection
    st.markdown("### üìã S√©lection de pi√®ces")
    
    # Collecter les documents disponibles
    available_docs = collect_available_documents(analysis)
    
    if not available_docs:
        st.warning("‚ö†Ô∏è Aucun document disponible")
        return
    
    # Grouper par cat√©gorie
    categories = group_documents_by_category(available_docs)
    
    # Interface de s√©lection par cat√©gorie
    selected_pieces = []
    
    for category, docs in categories.items():
        with st.expander(f"üìÅ {category} ({len(docs)} documents)", expanded=True):
            select_all = st.checkbox(f"Tout s√©lectionner - {category}", key=f"select_all_{category}")
            
            for doc in docs:
                is_selected = st.checkbox(
                    f"üìÑ {doc['title']}",
                    value=select_all,
                    key=f"select_doc_{doc['id']}",
                    help=f"Source: {doc.get('source', 'N/A')}"
                )
                
                if is_selected:
                    selected_pieces.append(PieceSelectionnee(
                        numero=len(selected_pieces) + 1,
                        titre=doc['title'],
                        description=doc.get('description', ''),
                        categorie=category,
                        date=doc.get('date'),
                        source=doc.get('source', ''),
                        pertinence=calculate_piece_relevance(doc, analysis)
                    ))
    
    # Sauvegarder la s√©lection
    st.session_state.selected_pieces = selected_pieces
    
    # Actions sur la s√©lection
    if selected_pieces:
        st.success(f"‚úÖ {len(selected_pieces)} pi√®ces s√©lectionn√©es")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("üìä Cr√©er bordereau", key="create_bordereau_from_selection"):
                process_bordereau_request(query, analysis)
        
        with col2:
            if st.button("üìÑ Synth√©tiser", key="synthesize_selection"):
                synthesize_selected_pieces(selected_pieces)
        
        with col3:
            if st.button("üì§ Exporter liste", key="export_piece_list"):
                export_piece_list(selected_pieces)

def collect_available_documents(analysis: dict) -> list:
    """Collecte tous les documents disponibles"""
    documents = []
    
    # Documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        documents.append({
            'id': doc_id,
            'title': doc.title,
            'content': doc.content,
            'source': doc.source,
            'metadata': doc.metadata
        })
    
    # Si r√©f√©rence sp√©cifique
    if analysis.get('reference'):
        ref_docs = search_by_reference(f"@{analysis['reference']}")
        documents.extend(ref_docs)
    
    return documents

def group_documents_by_category(documents: list) -> dict:
    """Groupe les documents par cat√©gorie"""
    categories = defaultdict(list)
    
    for doc in documents:
        # D√©terminer la cat√©gorie
        category = determine_document_category(doc)
        categories[category].append(doc)
    
    return dict(categories)

def determine_document_category(doc: dict) -> str:
    """D√©termine la cat√©gorie d'un document"""
    title_lower = doc.get('title', '').lower()
    content_lower = doc.get('content', '')[:500].lower()
    
    # Patterns de cat√©gories
    category_patterns = {
        'Proc√©dure': ['plainte', 'proc√®s-verbal', 'audition', 'perquisition', 'garde √† vue'],
        'Expertise': ['expertise', 'expert', 'rapport technique', 'analyse'],
        'Comptabilit√©': ['bilan', 'compte', 'comptable', 'facture', 'devis'],
        'Contrats': ['contrat', 'convention', 'accord', 'avenant'],
        'Correspondance': ['courrier', 'email', 'lettre', 'mail'],
        'Pi√®ces d\'identit√©': ['carte identit√©', 'passeport', 'kbis', 'statuts'],
        'Bancaire': ['relev√©', 'virement', 'compte bancaire', 'rib']
    }
    
    for category, keywords in category_patterns.items():
        if any(kw in title_lower or kw in content_lower for kw in keywords):
            return category
    
    return 'Autres'

def calculate_piece_relevance(doc: dict, analysis: dict) -> float:
    """Calcule la pertinence d'une pi√®ce"""
    score = 0.5
    
    # Si le document contient des mots-cl√©s de l'analyse
    if analysis.get('subject_matter'):
        if analysis['subject_matter'] in doc.get('content', '').lower():
            score += 0.3
    
    # Si r√©f√©rence dans le titre
    if analysis.get('reference'):
        if analysis['reference'] in doc.get('title', '').lower():
            score += 0.2
    
    return min(score, 1.0)

# === FONCTIONS DE TRAITEMENT POUR BORDEREAU ===

def process_bordereau_request(query: str, analysis: dict):
    """Traite une demande de cr√©ation de bordereau"""
    
    pieces = st.session_state.get('selected_pieces', [])
    
    if not pieces:
        st.warning("‚ö†Ô∏è Aucune pi√®ce s√©lectionn√©e pour le bordereau")
        return
    
    # Cr√©er le bordereau
    bordereau = create_bordereau(pieces, analysis)
    
    # Afficher le bordereau
    st.markdown("### üìä Bordereau de communication de pi√®ces")
    
    # En-t√™te
    st.text_area(
        "En-t√™te du bordereau",
        value=bordereau['header'],
        height=150,
        key="bordereau_header"
    )
    
    # Table des pi√®ces
    if PANDAS_AVAILABLE:
        df = pd.DataFrame([
            {
                'N¬∞': p.numero,
                'Titre': p.titre,
                'Description': p.description,
                'Cat√©gorie': p.categorie,
                'Date': p.date.strftime('%d/%m/%Y') if p.date else 'N/A'
            }
            for p in pieces
        ])
        
        st.dataframe(df, use_container_width=True)
    else:
        # Affichage sans pandas
        for piece in pieces:
            st.write(f"**{piece.numero}.** {piece.titre}")
            if piece.description:
                st.caption(piece.description)
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.download_button(
            "üíæ T√©l√©charger bordereau",
            create_bordereau_document(bordereau),
            f"bordereau_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_bordereau"
        ):
            st.success("‚úÖ Bordereau t√©l√©charg√©")
    
    with col2:
        if st.button("‚úèÔ∏è Modifier s√©lection", key="modify_bordereau_selection"):
            st.session_state.show_piece_selection = True
            st.rerun()
    
    with col3:
        if st.button("üìß Envoyer bordereau", key="send_bordereau"):
            st.session_state.universal_query = f"envoyer bordereau @{analysis.get('reference', 'dossier')}"
            st.rerun()
    
    # Stocker le bordereau
    st.session_state.current_bordereau = bordereau

def create_bordereau(pieces: list, analysis: dict) -> dict:
    """Cr√©e un bordereau structur√©"""
    
    bordereau = {
        'header': f"""BORDEREAU DE COMMUNICATION DE PI√àCES
AFFAIRE : {analysis.get('reference', 'N/A').upper()}
DATE : {datetime.now().strftime('%d/%m/%Y')}
NOMBRE DE PI√àCES : {len(pieces)}
POUR : [√Ä compl√©ter]
CONTRE : [√Ä compl√©ter]
PI√àCES COMMUNIQU√âES :""",
        'pieces': pieces,
        'footer': """Je certifie que les pi√®ces communiqu√©es sont conformes aux originaux en ma possession.
Fait √† [Ville], le {datetime.now().strftime('%d/%m/%Y')}
[Signature]""",
        'metadata': {
            'created_at': datetime.now(),
            'piece_count': len(pieces),
            'categories': list(set(p.categorie for p in pieces)),
            'reference': analysis.get('reference')
        }
    }
    
    return bordereau

def create_bordereau_document(bordereau: dict) -> bytes:
    """Cr√©e le document Word du bordereau"""
    if not DOCX_AVAILABLE:
        # Fallback texte
        content = bordereau['header'] + '\n\n'
        
        for piece in bordereau['pieces']:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            content += f"   Cat√©gorie: {piece.categorie}\n"
            if piece.date:
                content += f"   Date: {piece.date.strftime('%d/%m/%Y')}\n"
            content += "\n"
        
        content += bordereau['footer']
        
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # En-t√™te
        for line in bordereau['header'].split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                if 'BORDEREAU' in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(16)
        
        # Table des pi√®ces
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # En-t√™tes de colonnes
        headers = ['N¬∞', 'Titre', 'Description', 'Cat√©gorie', 'Date']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Lignes de pi√®ces
        for piece in bordereau['pieces']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(piece.numero)
            row_cells[1].text = piece.titre
            row_cells[2].text = piece.description or ''
            row_cells[3].text = piece.categorie
            row_cells[4].text = piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A'
        
        # Pied de page
        doc.add_paragraph()
        for line in bordereau['footer'].split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur cr√©ation bordereau: {e}")
        return create_bordereau_document.__wrapped__(bordereau)  # Fallback

# === FONCTIONS DE TRAITEMENT POUR SYNTH√àSE ===

def process_synthesis_request(query: str, analysis: dict):
    """Traite une demande de synth√®se"""
    
    # D√©terminer la source de la synth√®se
    if st.session_state.get('selected_pieces'):
        content_to_synthesize = synthesize_selected_pieces(st.session_state.selected_pieces)
    
    elif analysis.get('reference'):
        docs = search_by_reference(f"@{analysis['reference']}")
        content_to_synthesize = synthesize_documents(docs)
    
    else:
        st.warning("‚ö†Ô∏è Aucun contenu √† synth√©tiser")
        return
    
    # Stocker le r√©sultat
    st.session_state.synthesis_result = content_to_synthesize

def synthesize_selected_pieces(pieces: list) -> dict:
    """Synth√©tise les pi√®ces s√©lectionn√©es"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le contexte
    context = "PI√àCES √Ä SYNTH√âTISER:\n\n"
    
    for piece in pieces[:20]:  # Limiter
        context += f"Pi√®ce {piece.numero}: {piece.titre}\n"
        context += f"Cat√©gorie: {piece.categorie}\n"
        if piece.description:
            context += f"Description: {piece.description}\n"
        context += "\n"
    
    # Prompt de synth√®se
    synthesis_prompt = f"""{context}
Cr√©e une synth√®se structur√©e de ces pi√®ces.
La synth√®se doit inclure:
1. Vue d'ensemble des pi√®ces
2. Points cl√©s par cat√©gorie
3. Chronologie si applicable
4. Liens et relations entre pi√®ces
5. Points d'attention juridiques
6. Recommandations
Format professionnel avec titres et sous-sections."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            synthesis_prompt,
            "Tu es un expert en analyse et synth√®se de documents juridiques."
        )
        
        if response['success']:
            return {
                'content': response['response'],
                'piece_count': len(pieces),
                'categories': list(set(p.categorie for p in pieces)),
                'timestamp': datetime.now()
            }
        else:
            return {'error': '√âchec de la synth√®se'}
            
    except Exception as e:
        return {'error': f'Erreur synth√®se: {str(e)}'}

def synthesize_documents(documents: list) -> dict:
    """Synth√©tise une liste de documents"""
    # Convertir en pi√®ces pour r√©utiliser la fonction
    pieces = []
    
    for i, doc in enumerate(documents):
        pieces.append(PieceSelectionnee(
            numero=i + 1,
            titre=doc.get('title', 'Sans titre'),
            description=doc.get('content', '')[:200] + '...' if doc.get('content') else '',
            categorie=determine_document_category(doc),
            source=doc.get('source', '')
        ))
    
    return synthesize_selected_pieces(pieces)

# === FONCTIONS DE TRAITEMENT POUR TEMPLATES ===

def process_template_request(query: str, analysis: dict):
    """Traite une demande li√©e aux templates"""
    
    action = analysis['details'].get('action', 'apply')
    
    if action == 'create':
        create_new_template()
    
    elif action == 'edit':
        edit_existing_template()
    
    else:  # apply
        apply_template()

def create_new_template():
    """Cr√©e un nouveau template"""
    
    template_name = st.session_state.get('new_template_name', '')
    
    if not template_name:
        st.warning("‚ö†Ô∏è Nom du template requis")
        return
    
    # Base du template
    base_template = st.session_state.get('base_template', 'Vide')
    
    if base_template == 'Vide':
        template_content = {
            'name': template_name,
            'structure': [],
            'style': 'formel',
            'category': st.session_state.get('template_category', 'Autre')
        }
    else:
        # Copier depuis un template existant
        template_content = DOCUMENT_TEMPLATES.get(base_template, {}).copy()
        template_content['name'] = template_name
    
    # √âditeur de structure
    st.markdown("### üìù Structure du template")
    
    structure = st.text_area(
        "Sections (une par ligne)",
        value='\n'.join(template_content.get('structure', [])),
        height=300,
        key="template_structure_editor"
    )
    
    template_content['structure'] = [s.strip() for s in structure.split('\n') if s.strip()]
    
    # Sauvegarder
    if st.button("üíæ Sauvegarder le template", key="save_new_template"):
        if 'saved_templates' not in st.session_state:
            st.session_state.saved_templates = {}
        
        st.session_state.saved_templates[clean_key(template_name)] = template_content
        st.success(f"‚úÖ Template '{template_name}' sauvegard√©")
        
        # Optionnel : sauvegarder dans un fichier
        save_templates_to_file()

def edit_existing_template():
    """√âdite un template existant"""
    
    template_to_edit = st.session_state.get('template_to_edit')
    
    if not template_to_edit:
        st.warning("‚ö†Ô∏è S√©lectionnez un template √† modifier")
        return
    
    # Charger le template
    if template_to_edit in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[template_to_edit].copy()
        is_builtin = True
    else:
        template = st.session_state.saved_templates.get(template_to_edit, {})
        is_builtin = False
    
    if not template:
        st.error("‚ùå Template introuvable")
        return
    
    # √âditeur
    st.markdown(f"### ‚úèÔ∏è √âdition du template '{template.get('name', template_to_edit)}'")
    
    if is_builtin:
        st.info("‚ÑπÔ∏è Template int√©gr√© - Les modifications seront sauvegard√©es comme nouveau template")
    
    # Nom
    new_name = st.text_input(
        "Nom du template",
        value=template.get('name', ''),
        key="edit_template_name"
    )
    
    # Structure
    structure = st.text_area(
        "Structure",
        value='\n'.join(template.get('structure', [])),
        height=300,
        key="edit_template_structure"
    )
    
    # Style
    style = st.selectbox(
        "Style par d√©faut",
        list(REDACTION_STYLES.keys()),
        index=list(REDACTION_STYLES.keys()).index(template.get('style', 'formel')),
        format_func=lambda x: REDACTION_STYLES[x]['name'],
        key="edit_template_style"
    )
    
    # Sauvegarder les modifications
    if st.button("üíæ Sauvegarder les modifications", key="save_template_edits"):
        updated_template = {
            'name': new_name,
            'structure': [s.strip() for s in structure.split('\n') if s.strip()],
            'style': style,
            'category': template.get('category', 'Autre')
        }
        
        if is_builtin:
            # Sauvegarder comme nouveau
            st.session_state.saved_templates[clean_key(new_name)] = updated_template
            st.success(f"‚úÖ Nouveau template '{new_name}' cr√©√©")
        else:
            # Mettre √† jour l'existant
            st.session_state.saved_templates[template_to_edit] = updated_template
            st.success(f"‚úÖ Template '{new_name}' mis √† jour")

def apply_template():
    """Applique un template s√©lectionn√©"""
    
    selected_template = st.session_state.get('selected_template')
    
    if not selected_template:
        st.warning("‚ö†Ô∏è S√©lectionnez un template √† appliquer")
        return
    
    # Charger le template
    if selected_template in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[selected_template]
    else:
        template = st.session_state.saved_templates.get(selected_template, {})
    
    if not template:
        st.error("‚ùå Template introuvable")
        return
    
    # Cr√©er une requ√™te de r√©daction avec le template
    st.session_state.universal_query = f"r√©diger {template.get('name', 'document')} avec template {selected_template}"
    
    # D√©finir le style
    st.session_state.redaction_style = template.get('style', 'formel')
    
    # D√©clencher la r√©daction
    st.info(f"‚úÖ Template '{template.get('name')}' appliqu√© - Lancez la r√©daction")

def save_templates_to_file():
    """Sauvegarde les templates dans un fichier"""
    try:
        import json
        
        templates_data = {
            'builtin': DOCUMENT_TEMPLATES,
            'custom': st.session_state.get('saved_templates', {})
        }
        
        # Cr√©er un fichier t√©l√©chargeable
        json_str = json.dumps(templates_data, indent=2, ensure_ascii=False)
        
        st.download_button(
            "üíæ Exporter tous les templates",
            json_str,
            f"templates_{datetime.now().strftime('%Y%m%d')}.json",
            "application/json",
            key="export_templates"
        )
        
    except Exception as e:
        st.error(f"Erreur sauvegarde templates: {e}")

# === FONCTIONS DE TRAITEMENT POUR JURISPRUDENCE ===

def process_jurisprudence_request(query: str, analysis: dict):
    """Traite une demande de recherche de jurisprudence"""
    
    # Utiliser le gestionnaire de recherche juridique
    st.session_state.jurisprudence_search_active = True
    
    # Rediriger vers l'onglet jurisprudence
    st.info("‚öñÔ∏è Recherche de jurisprudence activ√©e - Voir l'onglet Jurisprudence")

# === FONCTIONS DE TRAITEMENT POUR ANALYSE ===

def process_analysis_request(query: str, analysis: dict):
    """Traite une demande d'analyse IA"""
    
    # Collecter les documents pertinents
    documents = []
    
    if analysis.get('reference'):
        documents = search_by_reference(f"@{analysis['reference']}")
    else:
        # Recherche g√©n√©rale
        documents = perform_search(query)
    
    if not documents:
        st.warning("‚ö†Ô∏è Aucun document trouv√© pour l'analyse")
        return
    
    # D√©terminer le type d'analyse
    analysis_type = detect_analysis_type(query)
    
    # Lancer l'analyse appropri√©e
    with st.spinner("ü§ñ Analyse en cours..."):
        if analysis_type == 'risks':
            results = analyze_legal_risks(documents, query)
        elif analysis_type == 'compliance':
            results = analyze_compliance(documents, query)
        elif analysis_type == 'strategy':
            results = analyze_strategy(documents, query)
        else:
            results = perform_general_analysis(documents, query)
    
    # Stocker les r√©sultats
    st.session_state.ai_analysis_results = results

def detect_analysis_type(query: str) -> str:
    """D√©tecte le type d'analyse demand√©"""
    query_lower = query.lower()
    
    if any(word in query_lower for word in ['risque', 'danger', 'exposition', 'vuln√©rabilit√©']):
        return 'risks'
    elif any(word in query_lower for word in ['conformit√©', 'compliance', 'r√©glementation']):
        return 'compliance'
    elif any(word in query_lower for word in ['strat√©gie', 'd√©fense', 'approche', 'tactique']):
        return 'strategy'
    else:
        return 'general'

def analyze_legal_risks(documents: list, query: str) -> dict:
    """Analyse les risques juridiques"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le prompt d'analyse
    risk_prompt = f"""Analyse les risques juridiques dans ces documents.
DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}
QUESTION: {query}
Identifie et √©value:
1. RISQUES P√âNAUX
   - Infractions potentielles
   - √âl√©ments constitutifs
   - Niveau de risque (faible/moyen/√©lev√©)
   - Sanctions encourues
2. RISQUES CIVILS
   - Responsabilit√©s contractuelles
   - Pr√©judices potentiels
   - Montants estim√©s
3. RISQUES R√âPUTATIONNELS
   - Impact m√©diatique
   - Cons√©quences business
4. RECOMMANDATIONS
   - Actions pr√©ventives
   - Strat√©gies de mitigation
   - Priorit√©s
Format structur√© avec √©valuation pr√©cise."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            risk_prompt,
            "Tu es un expert en analyse de risques juridiques et compliance."
        )
        
        if response['success']:
            return {
                'type': 'risk_analysis',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

def perform_general_analysis(documents: list, query: str) -> dict:
    """Analyse g√©n√©rale des documents"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Prompt g√©n√©ral
    general_prompt = f"""Analyse ces documents pour r√©pondre √† la question.
DOCUMENTS:
{chr(10).join([f"- {doc.get('title', 'Sans titre')}: {doc.get('content', '')[:500]}..." for doc in documents[:10]])}
QUESTION: {query}
Fournis une analyse compl√®te et structur√©e."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            general_prompt,
            "Tu es un expert en analyse juridique."
        )
        
        if response['success']:
            return {
                'type': 'general_analysis',
                'content': response['response'],
                'document_count': len(documents),
                'timestamp': datetime.now(),
                'query': query
            }
            
    except Exception as e:
        return {'error': f'Erreur analyse: {str(e)}'}

# === FONCTIONS DE TRAITEMENT POUR RECHERCHE ===

def process_search_request(query: str, analysis: dict):
    """Traite une demande de recherche simple"""
    
    # Recherche selon le contexte
    if analysis.get('reference'):
        results = search_by_reference(f"@{analysis['reference']}")
    else:
        results = perform_search(query)
    
    # Stocker les r√©sultats
    st.session_state.search_results = results
    
    if not results:
        st.warning("‚ö†Ô∏è Aucun r√©sultat trouv√©")

def search_by_reference(reference: str) -> list:
    """Recherche par r√©f√©rence @"""
    results = []
    
    # Nettoyer la r√©f√©rence
    ref_clean = reference.replace('@', '').strip().lower()
    
    # Recherche dans les documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if ref_clean in doc.title.lower() or ref_clean in doc.source.lower():
            results.append({
                'id': doc_id,
                'title': doc.title,
                'content': doc.content,
                'source': doc.source,
                'score': 1.0 if ref_clean == doc.title.lower() else 0.8
            })
    
    # Recherche Azure si disponible
    search_manager = st.session_state.get('azure_search_manager')
    if search_manager and search_manager.is_connected():
        try:
            azure_results = search_manager.search(reference)
            results.extend(azure_results)
        except:
            pass
    
    # Trier par score
    results.sort(key=lambda x: x.get('score', 0), reverse=True)
    
    return results

def perform_search(query: str) -> list:
    """Effectue une recherche g√©n√©rale"""
    results = []
    
    # Recherche locale basique
    query_lower = query.lower()
    query_words = query_lower.split()
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        # Score simple bas√© sur les mots trouv√©s
        score = 0
        content_lower = doc.content.lower()
        title_lower = doc.title.lower()
        
        for word in query_words:
            if word in title_lower:
                score += 2
            if word in content_lower:
                score += 1
        
        if score > 0:
            results.append({
                'id': doc_id,
                'title': doc.title,
                'content': doc.content,
                'source': doc.source,
                'score': score / len(query_words)
            })
    
    # Recherche Azure si disponible
    search_manager = st.session_state.get('azure_search_manager')
    if search_manager and search_manager.is_connected():
        try:
            azure_results = search_manager.search(query)
            results.extend(azure_results)
        except:
            pass
    
    # Trier par score
    results.sort(key=lambda x: x.get('score', 0), reverse=True)
    
    return results[:50]  # Limiter √† 50 r√©sultats

# === FONCTIONS D'AFFICHAGE DES R√âSULTATS ===

def show_unified_results_tab():
    """Affiche tous les types de r√©sultats dans un onglet unifi√©"""
    
    # V√©rifier quel type de r√©sultat afficher
    has_results = False
    
    # R√âSULTATS DE R√âDACTION (Priorit√© 1)
    if st.session_state.get('redaction_result'):
        show_redaction_results()
        has_results = True
    
    # R√âSULTATS DE TIMELINE (Priorit√© 2)
    elif st.session_state.get('timeline_result'):
        show_timeline_results()
        has_results = True
    
    # R√âSULTATS DE MAPPING (Priorit√© 3)
    elif st.session_state.get('mapping_result'):
        show_mapping_results()
        has_results = True
    
    # R√âSULTATS DE COMPARAISON (Priorit√© 4)
    elif st.session_state.get('comparison_result'):
        show_comparison_results()
        has_results = True
    
    # R√âSULTATS DE SYNTH√àSE (Priorit√© 5)
    elif st.session_state.get('synthesis_result'):
        show_synthesis_results()
        has_results = True
    
    # R√âSULTATS D'ANALYSE IA (Priorit√© 6)
    elif st.session_state.get('ai_analysis_results'):
        show_ai_analysis_results()
        has_results = True
    
    # R√âSULTATS DE RECHERCHE (Priorit√© 7)
    elif st.session_state.get('search_results'):
        show_search_results()
        has_results = True
    
    # Message si aucun r√©sultat
    if not has_results:
        st.info("üí° Utilisez la barre de recherche universelle pour commencer")
        
        # Suggestions d'utilisation
        st.markdown("""
        ### üöÄ Exemples de commandes
        
        **Recherche :**
        - `contrats soci√©t√© XYZ`
        - `@affaire_martin documents comptables`
        
        **Analyse :**
        - `analyser les risques @dossier_p√©nal`
        - `identifier les infractions @affaire_corruption`
        
        **R√©daction :**
        - `r√©diger conclusions d√©fense @affaire_martin abus biens sociaux`
        - `cr√©er plainte avec constitution partie civile escroquerie`
        
        **Visualisations :**
        - `chronologie des faits @affaire_martin`
        - `cartographie des soci√©t√©s @groupe_abc`
        - `comparer les auditions @t√©moins`
        
        **Gestion :**
        - `s√©lectionner pi√®ces @dossier cat√©gorie proc√©dure`
        - `cr√©er bordereau @pi√®ces_s√©lectionn√©es`
        - `exporter analyse format word`
        """)

def show_redaction_results():
    """Affiche les r√©sultats de r√©daction avec toutes les fonctionnalit√©s"""
    
    result = st.session_state.redaction_result
    
    st.markdown("### üìù Document juridique g√©n√©r√©")
    
    # M√©tadonn√©es
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        doc_icons = {
            'conclusions': '‚öñÔ∏è Conclusions',
            'plainte': 'üìã Plainte',
            'constitution_pc': 'üõ°Ô∏è Constitution PC',
            'courrier': '‚úâÔ∏è Courrier',
            'assignation': 'üìú Assignation',
            'm√©moire': 'üìö M√©moire',
            'requ√™te': 'üìÑ Requ√™te'
        }
        st.metric("Type", doc_icons.get(result['type'], 'üìÑ Document'))
    
    with col2:
        providers_count = len([r for r in result.get('responses', []) if r.get('success')])
        st.metric("IA utilis√©es", providers_count)
    
    with col3:
        word_count = len(result['document'].split())
        st.metric("Mots", f"{word_count:,}")
    
    with col4:
        char_count = len(result['document'])
        st.metric("Caract√®res", f"{char_count:,}")
    
    # Zone d'√©dition principale
    st.markdown("#### üìÑ Contenu du document")
    
    # Ajuster la hauteur selon le type de document
    # Les conclusions et plaintes sont g√©n√©ralement tr√®s longues
    height_by_type = {
        'conclusions': 800,  # Tr√®s long pour dossiers complexes
        'plainte': 700,      # Long aussi
        'constitution_pc': 700,
        'm√©moire': 800,
        'assignation': 600,
        'requ√™te': 600,
        'courrier': 400
    }
    
    text_height = height_by_type.get(result['type'], 600)
    
    edited_content = st.text_area(
        "Vous pouvez √©diter le document",
        value=result['document'],
        height=text_height,
        key="edit_redaction_main",
        help="Le document a √©t√© g√©n√©r√© pour √™tre complet et d√©taill√©, adapt√© aux dossiers complexes"
    )
    
    # Barre d'outils d'√©dition
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        if st.button("üîÑ R√©g√©n√©rer", key="regenerate_main"):
            if st.session_state.get('last_universal_query'):
                process_universal_query(st.session_state.last_universal_query)
                st.rerun()
    
    with col2:
        if st.button("‚ûï Enrichir", key="enrich_document"):
            enrich_current_document(edited_content)
    
    with col3:
        if st.button("‚úÇÔ∏è Synth√©tiser", key="summarize_document"):
            create_document_summary(edited_content)
    
    with col4:
        if st.button("üîç V√©rifier", key="verify_document"):
            verify_document_content(edited_content)
    
    with col5:
        if st.button("üìä Statistiques", key="document_stats"):
            show_document_statistics(edited_content)
    
    # Options d'export
    st.markdown("#### üíæ Export et partage")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # Export Word avec mise en forme
        docx_data = create_formatted_docx(edited_content, result['type'])
        if st.download_button(
            "üìÑ Word (.docx)",
            docx_data,
            f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx_main"
        ):
            st.success("‚úÖ Document Word t√©l√©charg√©")
    
    with col2:
        # Export PDF (si disponible)
        try:
            pdf_data = create_pdf_from_content(edited_content, result['type'])
            if st.download_button(
                "üìë PDF",
                pdf_data,
                f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "application/pdf",
                key="download_pdf_main"
            ):
                st.success("‚úÖ PDF t√©l√©charg√©")
        except:
            st.info("PDF n√©cessite des libs suppl√©mentaires")
    
    with col3:
        # Export texte brut
        if st.download_button(
            "üìù Texte (.txt)",
            edited_content.encode('utf-8'),
            f"{result['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_txt_main"
        ):
            st.success("‚úÖ Fichier texte t√©l√©charg√©")
    
    with col4:
        # Pr√©parer email
        if st.button("üìß Envoyer", key="prepare_email_main"):
            prepare_document_email(edited_content, result['type'])
    
    # Versions individuelles des IA (si plusieurs)
    if len(result.get('responses', [])) > 1:
        with st.expander("ü§ñ Versions par IA", expanded=False):
            for i, response in enumerate(result['responses']):
                if response.get('success'):
                    st.markdown(f"#### {response['provider']}")
                    
                    # M√©triques par version
                    col1, col2 = st.columns([4, 1])
                    
                    with col1:
                        words = len(response['response'].split())
                        st.caption(f"üìä {words:,} mots")
                    
                    with col2:
                        if st.button(f"Utiliser", key=f"use_version_{i}"):
                            st.session_state.edit_redaction_main = response['response']
                            st.rerun()
                    
                    # Contenu
                    st.text_area(
                        f"Version {response['provider']}",
                        value=response['response'],
                        height=400,
                        key=f"version_{response['provider']}_{i}",
                        disabled=True
                    )
    
    # Jurisprudence utilis√©e
    if result.get('jurisprudence_used') and result.get('jurisprudence_references'):
        with st.expander("‚öñÔ∏è Jurisprudence cit√©e", expanded=False):
            for ref in result['jurisprudence_references']:
                st.markdown(f"- [{ref['title']}]({ref.get('url', '#')})")
                if ref.get('summary'):
                    st.caption(ref['summary'])

def create_formatted_docx(content: str, doc_type: str) -> bytes:
    """Cr√©e un document Word avec mise en forme professionnelle compl√®te"""
    
    if not DOCX_AVAILABLE:
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = docx.Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
        
        # Styles personnalis√©s pour documents juridiques
        styles = doc.styles
        
        # Style pour le titre principal
        title_style = styles.add_style('JuridicalTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # Style pour les sections principales (I., II., III.)
        heading1_style = styles.add_style('JuridicalHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.name = 'Times New Roman'
        heading1_style.font.size = Pt(14)
        heading1_style.font.bold = True
        heading1_style.paragraph_format.space_before = Pt(18)
        heading1_style.paragraph_format.space_after = Pt(12)
        
        # Style pour les sous-sections (A., B., C.)
        heading2_style = styles.add_style('JuridicalHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.name = 'Times New Roman'
        heading2_style.font.size = Pt(13)
        heading2_style.font.bold = True
        heading2_style.font.underline = True
        heading2_style.paragraph_format.space_before = Pt(12)
        heading2_style.paragraph_format.space_after = Pt(6)
        
        # Style pour le corps du texte
        body_style = styles['Normal']
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.first_line_indent = Inches(0.5)
        
        # Analyser et formater le contenu
        lines = content.split('\n')
        
        # Variables pour tracker l'√©tat
        in_header = True
        current_section = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Ligne vide - ajouter un saut seulement si pas dans l'en-t√™te
                if not in_header:
                    doc.add_paragraph()
                continue
            
            # D√©tecter le type de ligne
            # Titre principal (tout en majuscules, court)
            if line.isupper() and len(line.split()) < 10 and in_header:
                p = doc.add_paragraph(line, style='JuridicalTitle')
                in_header = False
            
            # En-t√™te (POUR:, CONTRE:, etc.)
            elif line.startswith(('POUR', 'CONTRE', 'TRIBUNAL', 'AFFAIRE')) and ':' in line:
                p = doc.add_paragraph()
                parts = line.split(':', 1)
                run1 = p.add_run(parts[0] + ':')
                run1.bold = True
                if len(parts) > 1:
                    p.add_run(' ' + parts[1])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Sections principales (I., II., III., etc.)
            elif re.match(r'^[IVX]+\.\s+', line):
                p = doc.add_paragraph(line, style='JuridicalHeading1')
                current_section = line
                in_header = False
            
            # Sous-sections (A., B., C., etc.)
            elif re.match(r'^[A-Z]\.\s+', line):
                p = doc.add_paragraph(line, style='JuridicalHeading2')
            
            # Sous-sous-sections (1., 2., 3., etc.)
            elif re.match(r'^\d+\.\s+', line):
                p = doc.add_paragraph(line)
                p.style.font.bold = True
                p.style.paragraph_format.left_indent = Inches(0.5)
            
            # Points avec tirets
            elif line.startswith('-'):
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                p.style.paragraph_format.left_indent = Inches(0.5)
            
            # PAR CES MOTIFS
            elif line.startswith('PAR CES MOTIFS'):
                # Saut de page avant PAR CES MOTIFS
                doc.add_page_break()
                p = doc.add_paragraph(line, style='JuridicalHeading1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Formules de conclusion
            elif any(line.startswith(phrase) for phrase in [
                'PLAISE AU TRIBUNAL',
                'PLAISE √Ä LA COUR',
                'IL EST DEMAND√â',
                'SOUS TOUTES R√âSERVES'
            ]):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style.font.bold = True
            
            # Texte normal
            else:
                p = doc.add_paragraph(line, style='Normal')
                
                # Mettre en gras les r√©f√©rences juridiques
                for run in p.runs:
                    text = run.text
                    # Articles de loi
                    if re.search(r'article\s+[LR]?\s*\d+', text, re.IGNORECASE):
                        run.font.bold = True
                    # Jurisprudence
                    elif re.search(r'(Cass\.|CA|CE|Cons\.\s*const\.)', text):
                        run.font.italic = True
        
        # Ajouter les m√©tadonn√©es dans le pied de page
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Document g√©n√©r√© le {datetime.now().strftime('%d/%m/%Y √† %H:%M')} - {doc_type.title()}"
        footer_para.style.font.size = Pt(9)
        footer_para.style.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Num√©rotation des pages
        # Note: La num√©rotation automatique n√©cessite des manipulations XML complexes
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur cr√©ation document Word: {e}")
        # Fallback : retourner le texte brut
        return content.encode('utf-8')

def create_pdf_from_content(content: str, doc_type: str) -> bytes:
    """Cr√©e un PDF √† partir du contenu (n√©cessite reportlab ou weasyprint)"""
    # Pour l'instant, retourner une version texte
    # L'impl√©mentation compl√®te n√©cessiterait reportlab
    return content.encode('utf-8')

def enrich_current_document(content: str):
    """Enrichit le document actuel avec plus de d√©tails"""
    st.info("üîÑ Enrichissement en cours...")
    # Impl√©menter l'enrichissement via IA

def show_timeline_results():
    """Affiche les r√©sultats de chronologie"""
    result = st.session_state.timeline_result
    
    st.markdown(f"### ‚è±Ô∏è Chronologie des {result['type']}")
    
    # M√©tadonn√©es
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("√âv√©nements", len(result['events']))
    with col2:
        st.metric("Documents analys√©s", result['document_count'])
    with col3:
        st.metric("Type", result['type'].title())
    
    # Visualisation
    if result.get('visualization'):
        st.plotly_chart(result['visualization'], use_container_width=True)
    else:
        # Affichage alternatif
        for event in result['events']:
            with st.container():
                col1, col2 = st.columns([1, 4])
                with col1:
                    st.write(f"**{event.get('date', 'N/A')}**")
                with col2:
                    st.write(event.get('description', ''))
                    if event.get('actors'):
                        st.caption(f"üë• {', '.join(event['actors'])}")

def show_mapping_results():
    """Affiche les r√©sultats de cartographie"""
    result = st.session_state.mapping_result
    
    st.markdown(f"### üó∫Ô∏è Cartographie des relations - {result['type']}")
    
    # Statistiques
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Entit√©s", result['analysis']['node_count'])
    with col2:
        st.metric("Relations", result['analysis']['edge_count'])
    with col3:
        st.metric("Densit√©", f"{result['analysis']['density']:.2%}")
    with col4:
        st.metric("Composantes", len(result['analysis']['components']))
    
    # Acteurs cl√©s
    if result['analysis']['key_players']:
        st.markdown("#### üéØ Acteurs principaux")
        for i, player in enumerate(result['analysis']['key_players'], 1):
            st.write(f"{i}. **{player}**")
    
    # Visualisation
    if result.get('visualization'):
        st.plotly_chart(result['visualization'], use_container_width=True)

def show_comparison_results():
    """Affiche les r√©sultats de comparaison"""
    result = st.session_state.comparison_result
    
    st.markdown(f"### üîÑ Comparaison - {result['type']}")
    
    # Statistiques
    stats = result['comparison'].get('statistics', {})
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Documents", result['document_count'])
    with col2:
        st.metric("Convergences", len(result['comparison'].get('convergences', [])))
    with col3:
        st.metric("Divergences", len(result['comparison'].get('divergences', [])))
    with col4:
        reliability = stats.get('reliability_score', 0)
        color = "üü¢" if reliability > 0.7 else "üü°" if reliability > 0.4 else "üî¥"
        st.metric("Fiabilit√©", f"{color} {reliability:.0%}")
    
    # Visualisations
    if result.get('visualizations'):
        for viz_name, viz in result['visualizations'].items():
            if viz:
                st.plotly_chart(viz, use_container_width=True)
    
    # D√©tails textuels
    tabs = st.tabs(["Convergences", "Divergences", "√âvolutions", "Analyse"])
    
    with tabs[0]:
        for conv in result['comparison'].get('convergences', []):
            st.write(f"**{conv['point']}**")
            st.write(conv['details'])
            st.divider()
    
    with tabs[1]:
        for div in result['comparison'].get('divergences', []):
            st.write(f"**{div['point']}**")
            st.write(div['details'])
            st.divider()
    
    with tabs[2]:
        for evo in result['comparison'].get('evolutions', []):
            st.write(f"**{evo.get('point', '√âvolution')}**")
            st.write(evo.get('details', ''))
            st.divider()
    
    with tabs[3]:
        st.write(result['comparison'].get('analysis', 'Pas d\'analyse disponible'))

def show_synthesis_results():
    """Affiche les r√©sultats de synth√®se"""
    result = st.session_state.synthesis_result
    
    if 'error' in result:
        st.error(f"‚ùå {result['error']}")
        return
    
    st.markdown("### üìù Synth√®se des documents")
    
    # M√©tadonn√©es
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Pi√®ces analys√©es", result.get('piece_count', 0))
    with col2:
        st.metric("Cat√©gories", len(result.get('categories', [])))
    with col3:
        st.metric("G√©n√©r√© le", result.get('timestamp', datetime.now()).strftime('%H:%M'))
    
    # Contenu de la synth√®se
    st.markdown("#### üìÑ Synth√®se")
    
    synthesis_content = st.text_area(
        "Contenu de la synth√®se",
        value=result.get('content', ''),
        height=600,
        key="synthesis_content_display"
    )
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.download_button(
            "üíæ T√©l√©charger",
            synthesis_content.encode('utf-8'),
            f"synthese_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_synthesis"
        ):
            st.success("‚úÖ Synth√®se t√©l√©charg√©e")
    
    with col2:
        if st.button("üìä Voir les pi√®ces", key="view_synthesis_pieces"):
            st.session_state.show_piece_selection = True
    
    with col3:
        if st.button("üîÑ R√©g√©n√©rer", key="regenerate_synthesis"):
            if st.session_state.get('selected_pieces'):
                synthesize_selected_pieces(st.session_state.selected_pieces)

def show_ai_analysis_results():
    """Affiche les r√©sultats d'analyse IA"""
    results = st.session_state.ai_analysis_results
    
    if 'error' in results:
        st.error(f"‚ùå {results['error']}")
        return
    
    analysis_titles = {
        'risk_analysis': '‚ö†Ô∏è Analyse des risques',
        'compliance': '‚úÖ Analyse de conformit√©',
        'strategy': 'üéØ Analyse strat√©gique',
        'general_analysis': 'ü§ñ Analyse g√©n√©rale'
    }
    
    st.markdown(f"### {analysis_titles.get(results.get('type'), 'ü§ñ Analyse IA')}")
    
    # M√©tadonn√©es
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Documents analys√©s", results.get('document_count', 0))
    with col2:
        st.metric("Type", results.get('type', 'general').replace('_', ' ').title())
    with col3:
        st.metric("G√©n√©r√© le", results.get('timestamp', datetime.now()).strftime('%H:%M'))
    
    # Question originale
    if results.get('query'):
        st.info(f"**Question :** {results['query']}")
    
    # Contenu de l'analyse
    st.markdown("#### üìä R√©sultats de l'analyse")
    
    analysis_content = st.text_area(
        "Analyse d√©taill√©e",
        value=results.get('content', ''),
        height=600,
        key="ai_analysis_content"
    )
    
    # Actions
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.download_button(
            "üíæ T√©l√©charger",
            analysis_content.encode('utf-8'),
            f"analyse_{results.get('type', 'general')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            key="download_analysis"
        ):
            st.success("‚úÖ Analyse t√©l√©charg√©e")
    
    with col2:
        if st.button("üìù Convertir en rapport", key="convert_to_report"):
            convert_analysis_to_report(analysis_content, results)
    
    with col3:
        if st.button("üìä Visualiser", key="visualize_analysis"):
            create_analysis_visualization(results)
    
    with col4:
        if st.button("üîÑ Approfondir", key="deepen_analysis"):
            deepen_current_analysis(results)

def show_search_results():
    """Affiche les r√©sultats de recherche"""
    results = st.session_state.search_results
    
    st.markdown(f"### üîç R√©sultats de recherche ({len(results)} documents)")
    
    if not results:
        st.info("Aucun r√©sultat trouv√©")
        return
    
    # Options d'affichage
    col1, col2, col3 = st.columns(3)
    
    with col1:
        sort_by = st.selectbox(
            "Trier par",
            ["Pertinence", "Titre", "Date", "Source"],
            key="sort_search_results"
        )
    
    with col2:
        view_mode = st.radio(
            "Affichage",
            ["Compact", "D√©taill√©"],
            key="view_mode_search",
            horizontal=True
        )
    
    with col3:
        results_per_page = st.selectbox(
            "R√©sultats par page",
            [10, 20, 50, 100],
            key="results_per_page"
        )
    
    # Trier les r√©sultats
    sorted_results = sort_search_results(results, sort_by)
    
    # Pagination
    total_pages = (len(sorted_results) - 1) // results_per_page + 1
    current_page = st.number_input(
        "Page",
        min_value=1,
        max_value=total_pages,
        value=1,
        key="search_page"
    )
    
    start_idx = (current_page - 1) * results_per_page
    end_idx = min(start_idx + results_per_page, len(sorted_results))
    
    # Afficher les r√©sultats
    for i, result in enumerate(sorted_results[start_idx:end_idx], start=start_idx + 1):
        with st.container():
            if view_mode == "Compact":
                col1, col2, col3 = st.columns([4, 1, 1])
                
                with col1:
                    st.markdown(f"**{i}. {result.get('title', 'Sans titre')}**")
                    st.caption(f"Source: {result.get('source', 'N/A')}")
                
                with col2:
                    if result.get('score'):
                        st.metric("Score", f"{result['score']:.0%}")
                
                with col3:
                    if st.button("Voir", key=f"view_result_{i}"):
                        show_document_detail(result)
            
            else:  # D√©taill√©
                st.markdown(f"**{i}. {result.get('title', 'Sans titre')}**")
                st.caption(f"Source: {result.get('source', 'N/A')} | Score: {result.get('score', 0):.0%}")
                
                # Extrait du contenu
                content = result.get('content', '')
                if content:
                    st.text_area(
                        "Extrait",
                        value=content[:500] + '...' if len(content) > 500 else content,
                        height=150,
                        key=f"extract_{i}",
                        disabled=True
                    )
                
                # Actions
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if st.button("üìñ Lire", key=f"read_{i}"):
                        show_document_detail(result)
                
                with col2:
                    if st.button("ü§ñ Analyser", key=f"analyze_{i}"):
                        analyze_single_document(result)
                
                with col3:
                    if st.button("üìã S√©lectionner", key=f"select_{i}"):
                        add_to_selection(result)
            
            st.divider()
    
    # Navigation pages
    if total_pages > 1:
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col1:
            if current_page > 1:
                if st.button("‚¨ÖÔ∏è Pr√©c√©dent"):
                    st.session_state.search_page = current_page - 1
                    st.rerun()
        
        with col2:
            st.write(f"Page {current_page} / {total_pages}")
        
        with col3:
            if current_page < total_pages:
                if st.button("Suivant ‚û°Ô∏è"):
                    st.session_state.search_page = current_page + 1
                    st.rerun()

def sort_search_results(results: list, sort_by: str) -> list:
    """Trie les r√©sultats de recherche"""
    if sort_by == "Pertinence":
        return sorted(results, key=lambda x: x.get('score', 0), reverse=True)
    elif sort_by == "Titre":
        return sorted(results, key=lambda x: x.get('title', ''))
    elif sort_by == "Source":
        return sorted(results, key=lambda x: x.get('source', ''))
    else:  # Date
        # Essayer de parser les dates si disponibles
        return results

def show_document_detail(document: dict):
    """Affiche le d√©tail d'un document"""
    st.session_state.current_document = document
    st.session_state.show_document_detail = True

def analyze_single_document(document: dict):
    """Lance l'analyse d'un seul document"""
    st.session_state.universal_query = f"analyser @{document.get('id', document.get('title', ''))}"
    st.rerun()

def add_to_selection(document: dict):
    """Ajoute un document √† la s√©lection"""
    if 'selected_documents' not in st.session_state:
        st.session_state.selected_documents = []
    
    st.session_state.selected_documents.append(document)
    st.success("‚úÖ Document ajout√© √† la s√©lection")

# === AUTRES FONCTIONS HELPER ===

def show_pieces_management_tab():
    """Affiche l'onglet de gestion des pi√®ces"""
    st.markdown("### üìã Gestion des pi√®ces")
    
    if st.session_state.get('selected_pieces'):
        pieces = st.session_state.selected_pieces
        
        # Statistiques
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total pi√®ces", len(pieces))
        with col2:
            categories = list(set(p.categorie for p in pieces))
            st.metric("Cat√©gories", len(categories))
        with col3:
            st.metric("Pertinence moyenne", f"{sum(p.pertinence for p in pieces) / len(pieces):.0%}")
        
        # Actions group√©es
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("üìä Cr√©er bordereau", key="create_bordereau_tab"):
                process_bordereau_request("", {})
        
        with col2:
            if st.button("üìù Synth√©tiser", key="synthesize_pieces_tab"):
                synthesize_selected_pieces(pieces)
        
        with col3:
            if st.button("üì§ Exporter liste", key="export_pieces_tab"):
                export_piece_list(pieces)
        
        with col4:
            if st.button("üóëÔ∏è Vider s√©lection", key="clear_selection_tab"):
                st.session_state.selected_pieces = []
                st.rerun()
        
        # Affichage des pi√®ces
        for piece in pieces:
            with st.container():
                col1, col2, col3, col4 = st.columns([1, 3, 1, 1])
                
                with col1:
                    st.write(f"**{piece.numero}**")
                
                with col2:
                    st.write(piece.titre)
                    if piece.description:
                        st.caption(piece.description)
                
                with col3:
                    st.caption(piece.categorie)
                
                with col4:
                    if st.button("‚ùå", key=f"remove_piece_{piece.numero}"):
                        pieces.remove(piece)
                        st.rerun()
            
            st.divider()
    
    else:
        st.info("Aucune pi√®ce s√©lectionn√©e. Utilisez la commande `s√©lectionner pi√®ces` dans la barre universelle.")

def export_piece_list(pieces: list):
    """Exporte la liste des pi√®ces"""
    content = "LISTE DES PI√àCES S√âLECTIONN√âES\n"
    content += f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
    content += f"Nombre de pi√®ces : {len(pieces)}\n\n"
    
    # Grouper par cat√©gorie
    by_category = defaultdict(list)
    for piece in pieces:
        by_category[piece.categorie].append(piece)
    
    for category, cat_pieces in by_category.items():
        content += f"\n{category.upper()} ({len(cat_pieces)} pi√®ces)\n"
        content += "-" * 50 + "\n"
        
        for piece in cat_pieces:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            if piece.date:
                content += f"   Date : {piece.date.strftime('%d/%m/%Y')}\n"
            content += "\n"
    
    # T√©l√©charger
    st.download_button(
        "üíæ T√©l√©charger la liste",
        content.encode('utf-8'),
        f"liste_pieces_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
        "text/plain",
        key="download_piece_list_file"
    )

def show_jurisprudence_tab():
    """Affiche l'onglet jurisprudence"""
    # Utiliser le gestionnaire existant
    display_legal_search_interface()

def show_explorer_tab():
    """Affiche l'explorateur de fichiers"""
    st.markdown("### üìÅ Explorateur de documents")
    
    # Sources disponibles
    sources = []
    
    if st.session_state.get('azure_documents'):
        sources.append("Documents locaux")
    
    if st.session_state.get('azure_blob_manager'):
        sources.append("Azure Blob Storage")
    
    if not sources:
        st.info("Aucune source de documents disponible")
        return
    
    selected_source = st.selectbox(
        "Source",
        sources,
        key="explorer_source"
    )
    
    if selected_source == "Documents locaux":
        show_local_documents_explorer()
    
    elif selected_source == "Azure Blob Storage":
        show_azure_blob_explorer()

def show_local_documents_explorer():
    """Affiche l'explorateur de documents locaux"""
    documents = st.session_state.get('azure_documents', {})
    
    if not documents:
        st.info("Aucun document local")
        return
    
    # Filtres
    col1, col2 = st.columns(2)
    
    with col1:
        search_filter = st.text_input(
            "Filtrer par nom",
            key="explorer_filter_name"
        )
    
    with col2:
        source_filter = st.selectbox(
            "Filtrer par source",
            ["Toutes"] + list(set(doc.source for doc in documents.values())),
            key="explorer_filter_source"
        )
    
    # Appliquer les filtres
    filtered_docs = {}
    for doc_id, doc in documents.items():
        if search_filter and search_filter.lower() not in doc.title.lower():
            continue
        if source_filter != "Toutes" and doc.source != source_filter:
            continue
        filtered_docs[doc_id] = doc
    
    # Afficher les documents
    st.write(f"**{len(filtered_docs)} documents**")
    
    for doc_id, doc in filtered_docs.items():
        with st.expander(f"üìÑ {doc.title}"):
            st.write(f"**ID :** {doc_id}")
            st.write(f"**Source :** {doc.source}")
            
            if doc.metadata:
                st.write("**M√©tadonn√©es :**")
                for key, value in doc.metadata.items():
                    st.write(f"- {key}: {value}")
            
            # Actions
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("üìñ Lire", key=f"read_doc_{doc_id}"):
                    st.session_state.current_document = {
                        'id': doc_id,
                        'title': doc.title,
                        'content': doc.content,
                        'source': doc.source
                    }
            
            with col2:
                if st.button("ü§ñ Analyser", key=f"analyze_doc_{doc_id}"):
                    st.session_state.universal_query = f"analyser @{doc_id}"
                    st.rerun()
            
            with col3:
                if st.button("üóëÔ∏è Supprimer", key=f"delete_doc_{doc_id}"):
                    del st.session_state.azure_documents[doc_id]
                    st.rerun()

def show_azure_blob_explorer():
    """Affiche l'explorateur Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if not blob_manager or not blob_manager.is_connected():
        st.warning("Azure Blob Storage non connect√©")
        return
    
    try:
        containers = blob_manager.list_containers()
        
        if not containers:
            st.info("Aucun conteneur disponible")
            return
        
        selected_container = st.selectbox(
            "Conteneur",
            containers,
            key="blob_container"
        )
        
        # Explorer le conteneur
        current_path = st.session_state.get('blob_current_path', '')
        
        # Breadcrumb
        if current_path:
            path_parts = current_path.split('/')
            breadcrumb = "üìÅ " + " > ".join(path_parts)
            st.write(breadcrumb)
            
            if st.button("‚¨ÜÔ∏è Dossier parent"):
                st.session_state.blob_current_path = '/'.join(path_parts[:-1])
                st.rerun()
        
        # Lister le contenu
        items = blob_manager.list_folder_contents(selected_container, current_path)
        
        # S√©parer dossiers et fichiers
        folders = [item for item in items if item['type'] == 'folder']
        files = [item for item in items if item['type'] == 'file']
        
        # Afficher les dossiers
        if folders:
            st.write("**üìÅ Dossiers**")
            for folder in folders:
                col1, col2 = st.columns([4, 1])
                
                with col1:
                    st.write(f"üìÅ {folder['name']}")
                
                with col2:
                    if st.button("Ouvrir", key=f"open_folder_{folder['name']}"):
                        st.session_state.blob_current_path = folder['path']
                        st.rerun()
        
        # Afficher les fichiers
        if files:
            st.write("**üìÑ Fichiers**")
            for file in files:
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    st.write(f"üìÑ {file['name']}")
                    st.caption(f"Taille: {file['size']} octets")
                
                with col2:
                    if st.button("üíæ", key=f"download_blob_{file['name']}"):
                        content = blob_manager.download_blob(selected_container, file['path'])
                        st.download_button(
                            "T√©l√©charger",
                            content,
                            file['name'],
                            key=f"download_actual_{file['name']}"
                        )
                
                with col3:
                    if st.button("ü§ñ", key=f"analyze_blob_{file['name']}"):
                        # T√©l√©charger et analyser
                        content = blob_manager.download_blob(selected_container, file['path'])
                        # Stocker temporairement
                        store_locally(
                            file['name'],
                            content.decode('utf-8', errors='ignore'),
                            f"azure:{selected_container}/{file['path']}"
                        )
                        st.session_state.universal_query = f"analyser @{file['name']}"
                        st.rerun()
        
    except Exception as e:
        st.error(f"Erreur exploration Azure: {e}")

def show_configuration_tab():
    """Affiche l'onglet de configuration"""
    st.markdown("### ‚öôÔ∏è Configuration")
    
    tabs = st.tabs(["IA", "Connexions", "Templates", "Pr√©f√©rences"])
    
    with tabs[0]:
        show_ai_configuration()
    
    with tabs[1]:
        show_connections_configuration()
    
    with tabs[2]:
        show_templates_configuration()
    
    with tabs[3]:
        show_preferences_configuration()

def show_ai_configuration():
    """Configuration des IA"""
    st.markdown("#### ü§ñ Configuration des IA")
    
    llm_manager = MultiLLMManager()
    
    # √âtat des connexions
    st.write("**√âtat des connexions :**")
    
    for provider in LLMProvider:
        col1, col2 = st.columns([3, 1])
        
        with col1:
            is_connected = provider in llm_manager.clients
            status = "‚úÖ Connect√©" if is_connected else "‚ùå Non connect√©"
            st.write(f"{provider.value}: {status}")
        
        with col2:
            if is_connected:
                if st.button(f"Tester", key=f"test_{provider.value}"):
                    test_llm_connection(provider)
    
    # Pr√©f√©rences par d√©faut
    st.write("**Pr√©f√©rences par d√©faut :**")
    
    default_providers = st.multiselect(
        "IA par d√©faut pour la r√©daction",
        [p.value for p in LLMProvider if p in llm_manager.clients],
        default=[p.value for p in list(llm_manager.clients.keys())[:2]] if llm_manager.clients else [],
        key="default_redaction_providers"
    )
    
    fusion_mode = st.selectbox(
        "Mode de fusion par d√©faut",
        ["üéØ Fusion intelligente", "üìã Comparaison c√¥te √† c√¥te", "üîó Synth√®se enrichie", "‚ö° Meilleure version"],
        key="default_fusion_mode"
    )

def test_llm_connection(provider: LLMProvider):
    """Teste la connexion √† une IA"""
    llm_manager = MultiLLMManager()
    
    with st.spinner(f"Test de {provider.value}..."):
        try:
            response = llm_manager.query_single_llm(
                provider,
                "R√©ponds simplement 'OK' si tu me re√ßois.",
                "Assistant de test"
            )
            
            if response['success']:
                st.success(f"‚úÖ {provider.value} fonctionne correctement")
            else:
                st.error(f"‚ùå Erreur: {response.get('error', 'Inconnue')}")
                
        except Exception as e:
            st.error(f"‚ùå Erreur connexion: {str(e)}")

def show_connections_configuration():
    """Configuration des connexions externes"""
    st.markdown("#### üîå Connexions externes")
    
    # Azure Blob Storage
    st.write("**Azure Blob Storage**")
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        st.success("‚úÖ Connect√©")
        if st.button("üîÑ Rafra√Æchir la connexion"):
            blob_manager.connect()
    else:
        st.warning("‚ùå Non connect√©")
        st.info("Configurez les variables d'environnement Azure")
    
    # Azure Search
    st.write("**Azure Cognitive Search**")
    search_manager = st.session_state.get('azure_search_manager')
    
    if search_manager and search_manager.is_connected():
        st.success("‚úÖ Connect√©")
    else:
        st.warning("‚ùå Non connect√©")
    
    # Email SMTP
    st.write("**Configuration Email**")
    
    smtp_configured = bool(st.secrets.get('smtp_username'))
    if smtp_configured:
        st.success("‚úÖ Email configur√©")
    else:
        st.warning("‚ùå Email non configur√©")
        st.info("Ajoutez les param√®tres SMTP dans les secrets")

def show_templates_configuration():
    """Configuration des templates"""
    st.markdown("#### üìÑ Gestion des templates")
    
    # Templates int√©gr√©s
    st.write("**Templates int√©gr√©s :**")
    for key, template in DOCUMENT_TEMPLATES.items():
        st.write(f"- {template['name']} ({template.get('style', 'N/A')})")
    
    # Templates personnalis√©s
    custom_templates = st.session_state.get('saved_templates', {})
    
    if custom_templates:
        st.write(f"**Templates personnalis√©s ({len(custom_templates)}) :**")
        
        for key, template in custom_templates.items():
            col1, col2, col3 = st.columns([3, 1, 1])
            
            with col1:
                st.write(f"- {template.get('name', key)}")
            
            with col2:
                if st.button("‚úèÔ∏è", key=f"edit_template_{key}"):
                    st.session_state.template_to_edit = key
                    st.session_state.universal_query = "modifier template"
            
            with col3:
                if st.button("üóëÔ∏è", key=f"delete_template_{key}"):
                    del st.session_state.saved_templates[key]
                    st.rerun()
    
    # Import/Export
    st.divider()
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("üì§ Exporter tous les templates"):
            save_templates_to_file()
    
    with col2:
        uploaded_file = st.file_uploader(
            "üì• Importer des templates",
            type=['json'],
            key="import_templates"
        )
        
        if uploaded_file:
            try:
                import json
                templates_data = json.load(uploaded_file)
                
                if 'custom' in templates_data:
                    st.session_state.saved_templates.update(templates_data['custom'])
                    st.success(f"‚úÖ {len(templates_data['custom'])} templates import√©s")
                    st.rerun()
                    
            except Exception as e:
                st.error(f"Erreur import: {e}")

def show_preferences_configuration():
    """Configuration des pr√©f√©rences utilisateur"""
    st.markdown("#### üé® Pr√©f√©rences")
    
    # Pr√©f√©rences d'affichage
    st.write("**Affichage :**")
    
    results_per_page = st.selectbox(
        "R√©sultats par page par d√©faut",
        [10, 20, 50, 100],
        index=1,
        key="pref_results_per_page"
    )
    
    default_view = st.radio(
        "Vue par d√©faut",
        ["Compact", "D√©taill√©"],
        key="pref_default_view",
        horizontal=True
    )
    
    # Pr√©f√©rences de r√©daction
    st.write("**R√©daction :**")
    
    auto_jurisprudence = st.checkbox(
        "Recherche automatique de jurisprudence",
        value=True,
        key="pref_auto_juris"
    )
    
    create_hyperlinks = st.checkbox(
        "Cr√©er des liens hypertextes automatiques",
        value=True,
        key="pref_hyperlinks"
    )
    
    default_doc_length = st.select_slider(
        "Longueur par d√©faut des documents",
        options=["Concis", "Standard", "D√©taill√©", "Tr√®s d√©taill√©", "Exhaustif"],
        value="Tr√®s d√©taill√©",
        key="pref_doc_length"
    )
    
    # Sauvegarder les pr√©f√©rences
    if st.button("üíæ Sauvegarder les pr√©f√©rences"):
        preferences = {
            'results_per_page': results_per_page,
            'default_view': default_view,
            'auto_jurisprudence': auto_jurisprudence,
            'create_hyperlinks': create_hyperlinks,
            'default_doc_length': default_doc_length
        }
        
        st.session_state.user_preferences = preferences
        st.success("‚úÖ Pr√©f√©rences sauvegard√©es")

# Fonctions helper additionnelles

def clear_universal_state():
    """Efface l'√©tat de l'interface universelle"""
    keys_to_clear = [
        'universal_query', 'last_universal_query', 'current_analysis',
        'redaction_result', 'timeline_result', 'mapping_result',
        'comparison_result', 'synthesis_result', 'ai_analysis_results',
        'search_results', 'selected_pieces', 'import_files'
    ]
    
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    
    st.success("‚úÖ Interface r√©initialis√©e")
    st.rerun()

def save_current_work():
    """Sauvegarde le travail actuel"""
    work_data = {
        'timestamp': datetime.now().isoformat(),
        'query': st.session_state.get('universal_query', ''),
        'analysis': st.session_state.get('current_analysis', {}),
        'results': {}
    }
    
    # Collecter tous les r√©sultats
    result_keys = [
        'redaction_result', 'timeline_result', 'mapping_result',
        'comparison_result', 'synthesis_result', 'ai_analysis_results',
        'search_results'
    ]
    
    for key in result_keys:
        if key in st.session_state:
            work_data['results'][key] = st.session_state[key]
    
    # Sauvegarder
    import json
    
    json_str = json.dumps(work_data, indent=2, ensure_ascii=False, default=str)
    
    st.download_button(
        "üíæ T√©l√©charger la sauvegarde",
        json_str,
        f"sauvegarde_travail_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        "application/json",
        key="download_work_save"
    )

def share_current_work():
    """Partage le travail actuel"""
    # Cr√©er un lien de partage ou export
    st.info("üîó Fonctionnalit√© de partage √† impl√©menter")
    
    # Pour l'instant, proposer l'export
    save_current_work()

def prepare_document_email(content: str, doc_type: str):
    """Pr√©pare l'envoi d'un document par email"""
    st.session_state.email_document = {
        'content': content,
        'type': doc_type
    }
    
    st.session_state.universal_query = f"envoyer {doc_type} par email"
    st.rerun()

def verify_document_content(content: str):
    """V√©rifie le contenu d'un document"""
    # V√©rifications basiques
    issues = []
    
    # Longueur
    if len(content) < 1000:
        issues.append("‚ö†Ô∏è Document tr√®s court (< 1000 caract√®res)")
    
    # R√©f√©rences juridiques
    if not re.search(r'article\s+\d+', content, re.IGNORECASE):
        issues.append("‚ö†Ô∏è Aucune r√©f√©rence d'article de loi")
    
    # Structure
    if not any(marker in content for marker in ['I.', 'II.', 'A.', 'B.']):
        issues.append("‚ö†Ô∏è Structure peu claire (pas de sections num√©rot√©es)")
    
    # Afficher le r√©sultat
    if issues:
        st.warning("Points d'attention :")
        for issue in issues:
            st.write(issue)
    else:
        st.success("‚úÖ Document bien structur√©")

def show_document_statistics(content: str):
    """Affiche les statistiques d'un document"""
    
    # Calculs
    words = content.split()
    sentences = content.split('.')
    paragraphs = content.split('\n\n')
    
    # R√©f√©rences
    law_refs = len(re.findall(r'article\s+[LR]?\s*\d+', content, re.IGNORECASE))
    juris_refs = len(re.findall(r'(Cass\.|CA|CE|Cons\.\s*const\.)', content))
    
    # Affichage
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Mots", f"{len(words):,}")
        st.metric("Phrases", f"{len(sentences):,}")
    
    with col2:
        st.metric("Paragraphes", len(paragraphs))
        st.metric("Mots/phrase", f"{len(words) / max(len(sentences), 1):.1f}")
    
    with col3:
        st.metric("Articles cit√©s", law_refs)
        st.metric("Jurisprudences", juris_refs)
    
    with col4:
        # Complexit√© (basique)
        avg_word_length = sum(len(w) for w in words) / max(len(words), 1)
        complexity = "Simple" if avg_word_length < 5 else "Moyen" if avg_word_length < 7 else "Complexe"
        st.metric("Complexit√©", complexity)
        st.metric("Longueur moy.", f"{avg_word_length:.1f} car/mot")

def create_document_summary(content: str):
    """Cr√©e un r√©sum√© du document"""
    llm_manager = MultiLLMManager()
    
    if not llm_manager.clients:
        st.error("Aucune IA disponible pour le r√©sum√©")
        return
    
    with st.spinner("Cr√©ation du r√©sum√©..."):
        summary_prompt = f"""Cr√©e un r√©sum√© ex√©cutif de ce document juridique.
Le r√©sum√© doit contenir :
1. Type et objet du document
2. Points cl√©s (3-5 points)
3. Demandes principales
4. Enjeux juridiques
Document :
{content[:5000]}...
R√©sum√© en 200-300 mots maximum."""
        
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            summary_prompt,
            "Tu es un expert en synth√®se juridique."
        )
        
        if response['success']:
            st.markdown("### üìã R√©sum√© ex√©cutif")
            st.write(response['response'])
            
            # Proposer de sauvegarder
            if st.button("üíæ Sauvegarder le r√©sum√©"):
                st.session_state.document_summary = response['response']

def convert_analysis_to_report(analysis_content: str, analysis_metadata: dict):
    """Convertit une analyse en rapport formel"""
    st.info("üìù Conversion en rapport formel...")
    
    # Cr√©er une requ√™te de r√©daction de rapport
    report_query = f"r√©diger rapport formel bas√© sur analyse {analysis_metadata.get('type', 'g√©n√©rale')}"
    
    # Injecter le contenu de l'analyse
    st.session_state.analysis_to_report = analysis_content

# modules/recherche.py
"""Module de recherche unifi√©e avec 100% des fonctionnalit√©s int√©gr√©es"""

import streamlit as st
import asyncio
from datetime import datetime, timedelta
import re
import json
from collections import defaultdict, Counter
from typing import List, Dict, Optional, Any, Tuple, Union
import io
import base64
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

# Imports optionnels avec gestion d'erreur
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("‚ö†Ô∏è pandas non install√© - Certaines fonctionnalit√©s seront limit√©es")

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False
    print("‚ö†Ô∏è plotly non install√© - Visualisations simplifi√©es")

try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False
    print("‚ö†Ô∏è networkx non install√© - Analyse r√©seau limit√©e")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("‚ö†Ô∏è python-docx non install√© - Export Word non disponible")

# Imports des modules internes
from config.app_config import SearchMode, ANALYSIS_PROMPTS_AFFAIRES, LLMProvider
from managers.azure_blob_manager import AzureBlobManager
from managers.azure_search_manager import AzureSearchManager
from managers.multi_llm_manager import MultiLLMManager
from managers.document_manager import display_import_interface
from managers.dynamic_generators import generate_dynamic_search_prompts
from managers.legal_search import LegalSearchManager, display_legal_search_interface
from managers.jurisprudence_verifier import JurisprudenceVerifier
from models.dataclasses import Document, PieceSelectionnee
from utils.helpers import clean_key

# Configuration des styles de r√©daction
REDACTION_STYLES = {
    'formel': {
        'name': 'Formel',
        'description': 'Style juridique classique et solennel',
        'tone': 'respectueux et distant',
        'vocabulary': 'technique et pr√©cis'
    },
    'persuasif': {
        'name': 'Persuasif',
        'description': 'Style argumentatif et convaincant',
        'tone': 'assertif et engag√©',
        'vocabulary': 'percutant et imag√©'
    },
    'technique': {
        'name': 'Technique',
        'description': 'Style factuel et d√©taill√©',
        'tone': 'neutre et objectif',
        'vocabulary': 'sp√©cialis√© et exhaustif'
    },
    'synth√©tique': {
        'name': 'Synth√©tique',
        'description': 'Style concis et efficace',
        'tone': 'direct et clair',
        'vocabulary': 'simple et pr√©cis'
    },
    'p√©dagogique': {
        'name': 'P√©dagogique',
        'description': 'Style explicatif et accessible',
        'tone': 'bienveillant et didactique',
        'vocabulary': 'vulgaris√© et illustr√©'
    }
}

# Templates pr√©d√©finis
DOCUMENT_TEMPLATES = {
    'conclusions_defense': {
        'name': 'Conclusions en d√©fense',
        'structure': [
            'I. FAITS ET PROC√âDURE',
            'II. DISCUSSION',
            '   A. Sur la recevabilit√©',
            '   B. Sur le fond',
            '      1. Sur l\'√©l√©ment mat√©riel',
            '      2. Sur l\'√©l√©ment intentionnel',
            '      3. Sur le pr√©judice',
            'III. PAR CES MOTIFS'
        ],
        'style': 'formel'
    },
    'plainte_simple': {
        'name': 'Plainte simple',
        'structure': [
            'Objet : Plainte',
            'EXPOS√â DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PR√âJUDICES SUBIS',
            'DEMANDES',
            'PI√àCES JOINTES'
        ],
        'style': 'formel'
    },
    'plainte_avec_cpc': {
        'name': 'Plainte avec constitution de partie civile',
        'structure': [
            'Objet : Plainte avec constitution de partie civile',
            'EXPOS√â DES FAITS',
            'QUALIFICATION JURIDIQUE',
            'PR√âJUDICES SUBIS',
            'CONSTITUTION DE PARTIE CIVILE',
            '√âVALUATION DU PR√âJUDICE',
            'DEMANDES',
            'PI√àCES JOINTES'
        ],
        'style': 'formel'
    },
    'mise_en_demeure': {
        'name': 'Mise en demeure',
        'structure': [
            'MISE EN DEMEURE',
            'Rappel des faits',
            'Obligations non respect√©es',
            'D√©lai accord√©',
            'Cons√©quences du d√©faut',
            'R√©serves'
        ],
        'style': 'persuasif'
    },
    'note_synthese': {
        'name': 'Note de synth√®se',
        'structure': [
            'SYNTH√àSE EX√âCUTIVE',
            'I. CONTEXTE',
            'II. ANALYSE',
            'III. RECOMMANDATIONS',
            'IV. PLAN D\'ACTION'
        ],
        'style': 'synth√©tique'
    }
}

# === FONCTIONS DE TRAITEMENT POUR IMPORT ===

def process_import_request(query: str, analysis: dict):
    """Traite une demande d'import"""
    
    st.info("üì• Interface d'import activ√©e")
    
    # Configuration de l'import depuis la config
    file_types = analysis['details'].get('file_types', ['pdf', 'docx', 'txt'])
    destination = st.session_state.get('import_destination', 'Documents locaux')
    
    # Si des fichiers sont d√©j√† upload√©s
    if st.session_state.get('import_files'):
        with st.spinner("üì• Import des fichiers en cours..."):
            imported_count = 0
            
            for file in st.session_state.import_files:
                try:
                    # Lire le contenu
                    content = file.read()
                    
                    # Traiter selon le type
                    if file.name.endswith('.pdf'):
                        text_content = extract_text_from_pdf(content)
                    elif file.name.endswith('.docx'):
                        text_content = extract_text_from_docx(content)
                    else:
                        text_content = content.decode('utf-8', errors='ignore')
                    
                    # Stocker selon la destination
                    if destination == 'Azure Blob':
                        store_in_azure_blob(file.name, content, analysis.get('reference'))
                    else:
                        store_locally(file.name, text_content, analysis.get('reference'))
                    
                    imported_count += 1
                    
                except Exception as e:
                    st.error(f"‚ùå Erreur import {file.name}: {str(e)}")
            
            st.success(f"‚úÖ {imported_count} fichiers import√©s avec succ√®s")
            
            # Analyse automatique si demand√©e
            if st.session_state.get('auto_analyze_import', True):
                st.info("ü§ñ Analyse automatique des documents import√©s...")
                analyze_imported_documents()
    
    else:
        st.warning("‚ö†Ô∏è Aucun fichier s√©lectionn√©. Utilisez l'interface d'import ci-dessous.")

def extract_text_from_pdf(content: bytes) -> str:
    """Extrait le texte d'un PDF"""
    try:
        import PyPDF2
        import io
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text
        
    except ImportError:
        return "PDF - Installation de PyPDF2 requise pour l'extraction"
    except Exception as e:
        return f"Erreur extraction PDF: {str(e)}"

def extract_text_from_docx(content: bytes) -> str:
    """Extrait le texte d'un DOCX"""
    if DOCX_AVAILABLE:
        try:
            import io
            doc = docx.Document(io.BytesIO(content))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"Erreur extraction DOCX: {str(e)}"
    else:
        return "DOCX - Installation de python-docx requise"

def store_locally(filename: str, content: str, reference: str = None):
    """Stocke un document localement"""
    doc_id = f"imported_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{clean_key(filename)}"
    
    if 'azure_documents' not in st.session_state:
        st.session_state.azure_documents = {}
    
    st.session_state.azure_documents[doc_id] = Document(
        id=doc_id,
        title=filename,
        content=content,
        source=f"Import {reference if reference else 'direct'}",
        metadata={
            'imported_at': datetime.now().isoformat(),
            'reference': reference,
            'original_filename': filename
        }
    )

def store_in_azure_blob(filename: str, content: bytes, reference: str = None):
    """Stocke un document dans Azure Blob"""
    blob_manager = st.session_state.get('azure_blob_manager')
    
    if blob_manager and blob_manager.is_connected():
        container = reference if reference else 'imports'
        blob_name = f"{datetime.now().strftime('%Y%m%d')}_{filename}"
        
        try:
            blob_manager.upload_blob(container, blob_name, content)
            st.success(f"‚úÖ {filename} upload√© dans Azure")
        except Exception as e:
            st.error(f"‚ùå Erreur Azure: {str(e)}")
    else:
        st.warning("‚ö†Ô∏è Azure Blob non connect√© - Stockage local")
        store_locally(filename, content.decode('utf-8', errors='ignore'), reference)

def analyze_imported_documents():
    """Analyse automatique des documents import√©s"""
    # R√©cup√©rer les derniers documents import√©s
    recent_docs = []
    
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        if 'imported_at' in doc.metadata:
            import_time = datetime.fromisoformat(doc.metadata['imported_at'])
            if (datetime.now() - import_time).seconds < 300:  # Moins de 5 minutes
                recent_docs.append(doc)
    
    if recent_docs:
        # Lancer une analyse automatique
        analysis_prompt = f"""Analyse ces {len(recent_docs)} documents r√©cemment import√©s.
        
Identifie :
1. Type de document
2. Parties impliqu√©es
3. Dates importantes
4. Points juridiques cl√©s
5. Risques identifi√©s
Documents :
{chr(10).join([f"- {doc.title}: {doc.content[:500]}..." for doc in recent_docs[:5]])}"""
        
        llm_manager = MultiLLMManager()
        if llm_manager.clients:
            provider = list(llm_manager.clients.keys())[0]
            response = llm_manager.query_single_llm(
                provider,
                analysis_prompt,
                "Tu es un expert en analyse documentaire juridique."
            )
            
            if response['success']:
                st.session_state.import_analysis = response['response']
                st.success("‚úÖ Analyse des imports termin√©e")

# === FONCTIONS DE TRAITEMENT POUR EXPORT ===

def process_export_request(query: str, analysis: dict):
    """Traite une demande d'export"""
    
    format = analysis['details'].get('format', 'docx')
    
    # D√©terminer ce qu'il faut exporter
    content_to_export = None
    filename_base = "export"
    
    # Priorit√© : r√©daction > analyse > recherche > s√©lection
    if st.session_state.get('redaction_result'):
        content_to_export = st.session_state.redaction_result['document']
        filename_base = f"{st.session_state.redaction_result['type']}"
    
    elif st.session_state.get('timeline_result'):
        content_to_export = export_timeline_content(st.session_state.timeline_result)
        filename_base = "chronologie"
    
    elif st.session_state.get('mapping_result'):
        content_to_export = export_mapping_content(st.session_state.mapping_result)
        filename_base = "cartographie"
    
    elif st.session_state.get('comparison_result'):
        content_to_export = export_comparison_content(st.session_state.comparison_result)
        filename_base = "comparaison"
    
    elif st.session_state.get('ai_analysis_results'):
        content_to_export = format_ai_analysis_for_export(st.session_state.ai_analysis_results)
        filename_base = "analyse"
    
    elif st.session_state.get('search_results'):
        content_to_export = format_search_results_for_export(st.session_state.search_results)
        filename_base = "recherche"
    
    else:
        st.warning("‚ö†Ô∏è Aucun contenu √† exporter")
        return
    
    # Exporter selon le format
    export_functions = {
        'docx': export_to_docx,
        'pdf': export_to_pdf,
        'txt': export_to_txt,
        'html': export_to_html,
        'xlsx': export_to_xlsx
    }
    
    export_func = export_functions.get(format, export_to_txt)
    
    try:
        exported_data = export_func(content_to_export, analysis)
        
        # Cr√©er le bouton de t√©l√©chargement
        st.download_button(
            f"üíæ T√©l√©charger {format.upper()}",
            exported_data,
            f"{filename_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}",
            get_mime_type(format),
            key=f"download_export_{format}"
        )
        
        st.success(f"‚úÖ Export {format.upper()} pr√™t")
        
    except Exception as e:
        st.error(f"‚ùå Erreur export: {str(e)}")

def export_timeline_content(timeline_result: dict) -> str:
    """Exporte le contenu d'une chronologie"""
    content = f"CHRONOLOGIE - {timeline_result['type'].upper()}\n"
    content += f"G√©n√©r√©e le {timeline_result['timestamp'].strftime('%d/%m/%Y √† %H:%M')}\n"
    content += f"Bas√©e sur {timeline_result['document_count']} documents\n\n"
    
    for event in timeline_result['events']:
        content += f"{event.get('date', 'N/A')} - {event.get('description', '')}\n"
        if 'actors' in event:
            content += f"   Acteurs: {', '.join(event['actors'])}\n"
        if 'source' in event:
            content += f"   Source: {event['source']}\n"
        content += "\n"
    
    return content

def export_mapping_content(mapping_result: dict) -> str:
    """Exporte le contenu d'une cartographie"""
    content = f"CARTOGRAPHIE DES RELATIONS - {mapping_result['type'].upper()}\n"
    content += f"G√©n√©r√©e le {mapping_result['timestamp'].strftime('%d/%m/%Y √† %H:%M')}\n\n"
    
    content += f"STATISTIQUES:\n"
    content += f"- Entit√©s: {mapping_result['analysis']['node_count']}\n"
    content += f"- Relations: {mapping_result['analysis']['edge_count']}\n"
    content += f"- Densit√©: {mapping_result['analysis']['density']:.2%}\n\n"
    
    content += "ACTEURS PRINCIPAUX:\n"
    for i, player in enumerate(mapping_result['analysis']['key_players'], 1):
        content += f"{i}. {player}\n"
    
    content += "\nENTIT√âS:\n"
    for entity in mapping_result['entities']:
        content += f"- {entity['name']} ({entity.get('type', 'N/A')})\n"
    
    content += "\nRELATIONS:\n"
    for relation in mapping_result['relations']:
        content += f"- {relation['source']} ‚Üí {relation['target']}: {relation.get('type', 'N/A')}\n"
    
    return content

def export_comparison_content(comparison_result: dict) -> str:
    """Exporte le contenu d'une comparaison"""
    content = f"COMPARAISON - {comparison_result['type'].upper()}\n"
    content += f"G√©n√©r√©e le {comparison_result['timestamp'].strftime('%d/%m/%Y √† %H:%M')}\n"
    content += f"Documents compar√©s: {comparison_result['document_count']}\n\n"
    
    comparison = comparison_result['comparison']
    
    content += "CONVERGENCES:\n"
    for conv in comparison.get('convergences', []):
        content += f"- {conv['point']}\n  {conv['details']}\n\n"
    
    content += "DIVERGENCES:\n"
    for div in comparison.get('divergences', []):
        content += f"- {div['point']}\n  {div['details']}\n\n"
    
    if comparison.get('analysis'):
        content += f"ANALYSE:\n{comparison['analysis']}\n"
    
    return content

def export_to_docx(content: str, analysis: dict) -> bytes:
    """Exporte vers Word avec mise en forme"""
    if not DOCX_AVAILABLE:
        # Fallback en texte brut
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Styles personnalis√©s
        styles = doc.styles
        
        # En-t√™te
        if st.session_state.get('export_metadata', True):
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Document g√©n√©r√© le {datetime.now().strftime('%d/%m/%Y √† %H:%M')}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Titre principal
        title = doc.add_heading(analysis.get('document_type', 'Document').upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenu avec mise en forme
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            # D√©tecter les niveaux de titre
            if line.strip().isupper() and len(line.strip()) > 5:
                doc.add_heading(line.strip(), 1)
            elif line.strip().startswith('###'):
                doc.add_heading(line.strip().replace('#', '').strip(), 3)
            elif line.strip().startswith('##'):
                doc.add_heading(line.strip().replace('#', '').strip(), 2)
            elif line.strip().startswith('#'):
                doc.add_heading(line.strip().replace('#', '').strip(), 1)
            elif line.strip().startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
                p = doc.add_paragraph()
                p.add_run(line.strip()).bold = True
                p.style.font.size = Pt(14)
            elif line.strip().startswith(('A.', 'B.', 'C.', 'D.')):
                p = doc.add_paragraph(line.strip(), style='List Number')
                p.style.font.size = Pt(12)
            elif line.strip().startswith('-'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(line.strip())
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
        
        # Table des mati√®res si demand√©e
        if st.session_state.get('docx_toc', True):
            # Ins√©rer au d√©but
            doc.paragraphs[0].insert_paragraph_before("TABLE DES MATI√àRES", style='Heading 1')
            # Note: La vraie TOC n√©cessite des manipulations XML complexes
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur cr√©ation DOCX: {e}")
        return content.encode('utf-8')

def export_to_pdf(content: str, analysis: dict) -> bytes:
    """Exporte vers PDF"""
    # N√©cessiterait reportlab ou weasyprint
    st.warning("‚ö†Ô∏è Export PDF n√©cessite l'installation de biblioth√®ques suppl√©mentaires")
    
    # Pour l'instant, retourner le contenu texte
    return content.encode('utf-8')

def export_to_txt(content: str, analysis: dict) -> bytes:
    """Exporte en texte brut"""
    return content.encode('utf-8')

def export_to_html(content: str, analysis: dict) -> bytes:
    """Exporte vers HTML avec style"""
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <title>{analysis.get('document_type', 'Document').title()}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; margin: 40px; }}
        h1 {{ text-align: center; color: #2c3e50; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        p {{ text-align: justify; line-height: 1.6; }}
        .metadata {{ text-align: right; color: #95a5a6; font-size: 0.9em; }}
        .section {{ margin: 20px 0; }}
        ul, ol {{ margin-left: 20px; }}
    </style>
</head>
<body>
    <div class="metadata">G√©n√©r√© le {datetime.now().strftime('%d/%m/%Y √† %H:%M')}</div>
    <h1>{analysis.get('document_type', 'Document').upper()}</h1>
    <div class="content">
"""
    
    # Convertir le contenu en HTML
    lines = content.split('\n')
    in_list = False
    
    for line in lines:
        if not line.strip():
            if in_list:
                html += "</ul>"
                in_list = False
            continue
        
        if line.strip().isupper() and len(line.strip()) > 5:
            html += f"<h2>{line.strip()}</h2>\n"
        elif line.strip().startswith('-'):
            if not in_list:
                html += "<ul>\n"
                in_list = True
            html += f"<li>{line.strip()[1:].strip()}</li>\n"
        else:
            if in_list:
                html += "</ul>\n"
                in_list = False
            html += f"<p>{line.strip()}</p>\n"
    
    html += """
    </div>
</body>
</html>"""
    
    return html.encode('utf-8')

def export_to_xlsx(content: str, analysis: dict) -> bytes:
    """Exporte vers Excel"""
    if not PANDAS_AVAILABLE:
        st.error("‚ùå pandas requis pour l'export Excel")
        return content.encode('utf-8')
    
    try:
        import io
        
        # Cr√©er un DataFrame selon le type de contenu
        if 'timeline_result' in st.session_state:
            df = pd.DataFrame(st.session_state.timeline_result['events'])
        elif 'mapping_result' in st.session_state:
            df_entities = pd.DataFrame(st.session_state.mapping_result['entities'])
            df_relations = pd.DataFrame(st.session_state.mapping_result['relations'])
        else:
            # Contenu g√©n√©rique
            lines = content.split('\n')
            df = pd.DataFrame({'Contenu': lines})
        
        # Exporter vers Excel
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            if 'mapping_result' in st.session_state:
                df_entities.to_excel(writer, sheet_name='Entit√©s', index=False)
                df_relations.to_excel(writer, sheet_name='Relations', index=False)
            else:
                df.to_excel(writer, sheet_name='Donn√©es', index=False)
            
            # Ajouter une feuille de m√©tadonn√©es
            metadata = pd.DataFrame({
                'Propri√©t√©': ['Type', 'Date g√©n√©ration', 'Source'],
                'Valeur': [
                    analysis.get('document_type', 'Export'),
                    datetime.now().strftime('%d/%m/%Y %H:%M'),
                    analysis.get('reference', 'N/A')
                ]
            })
            metadata.to_excel(writer, sheet_name='M√©tadonn√©es', index=False)
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur export Excel: {e}")
        return content.encode('utf-8')

def get_mime_type(format: str) -> str:
    """Retourne le type MIME pour un format"""
    mime_types = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'html': 'text/html',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    }
    return mime_types.get(format, 'application/octet-stream')

# === FONCTIONS DE TRAITEMENT POUR EMAIL ===

def process_email_request(query: str, analysis: dict):
    """Traite une demande d'envoi par email"""
    
    recipients = analysis.get('recipients', [])
    
    if not recipients:
        st.error("‚ùå Aucun destinataire sp√©cifi√©")
        return
    
    # Pr√©parer le contenu
    content = prepare_email_content(analysis)
    
    if not content:
        st.warning("‚ö†Ô∏è Aucun contenu √† envoyer")
        return
    
    # Configuration email
    email_config = {
        'to': st.session_state.get('email_to', ', '.join(recipients)),
        'cc': st.session_state.get('email_cc', ''),
        'subject': st.session_state.get('email_subject', 'Document juridique'),
        'body': content['body'],
        'attachments': []
    }
    
    # Pr√©parer les pi√®ces jointes
    if st.session_state.get('email_attach_current', True):
        attachment_format = st.session_state.get('email_attachment_format', 'pdf')
        attachment_data = export_current_content(attachment_format)
        
        if attachment_data:
            email_config['attachments'].append({
                'filename': f"document.{attachment_format}",
                'data': attachment_data,
                'mime_type': get_mime_type(attachment_format)
            })
    
    # Afficher l'aper√ßu
    show_email_preview(email_config)
    
    # Bouton d'envoi
    if st.button("üìß Envoyer l'email", key="send_email_button"):
        if send_email(email_config):
            st.success("‚úÖ Email envoy√© avec succ√®s")
            st.session_state.email_sent = True
        else:
            st.error("‚ùå Erreur lors de l'envoi")

def prepare_email_content(analysis: dict) -> dict:
    """Pr√©pare le contenu de l'email"""
    content = {'body': '', 'attachments': []}
    
    # Corps de l'email selon le contexte
    if st.session_state.get('redaction_result'):
        result = st.session_state.redaction_result
        content['body'] = f"""Bonjour,
Veuillez trouver ci-joint le document {result['type']} demand√©.
Ce document a √©t√© g√©n√©r√© automatiquement √† partir de l'analyse de votre dossier.
Cordialement,
[Votre nom]"""
        
    else:
        content['body'] = """Bonjour,
Veuillez trouver ci-joint le document demand√©.
Cordialement,
[Votre nom]"""
    
    return content

def show_email_preview(email_config: dict):
    """Affiche un aper√ßu de l'email"""
    with st.expander("üìß Aper√ßu de l'email", expanded=True):
        st.text_input("√Ä:", value=email_config['to'], disabled=True)
        st.text_input("Cc:", value=email_config['cc'], disabled=True)
        st.text_input("Objet:", value=email_config['subject'], disabled=True)
        st.text_area("Corps:", value=email_config['body'], height=200, disabled=True)
        
        if email_config['attachments']:
            st.write("üìé Pi√®ces jointes:")
            for att in email_config['attachments']:
                st.write(f"- {att['filename']}")

def send_email(email_config: dict) -> bool:
    """Envoie l'email (fonction simplifi√©e)"""
    try:
        # Configuration SMTP (√† adapter)
        smtp_config = {
            'server': st.secrets.get('smtp_server', 'smtp.gmail.com'),
            'port': st.secrets.get('smtp_port', 587),
            'username': st.secrets.get('smtp_username', ''),
            'password': st.secrets.get('smtp_password', '')
        }
        
        if not smtp_config['username']:
            st.error("‚ùå Configuration email manquante")
            return False
        
        # Cr√©er le message
        msg = MIMEMultipart()
        msg['From'] = smtp_config['username']
        msg['To'] = email_config['to']
        msg['Cc'] = email_config['cc']
        msg['Subject'] = email_config['subject']
        
        # Corps
        msg.attach(MIMEText(email_config['body'], 'plain'))
        
        # Pi√®ces jointes
        for att in email_config['attachments']:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(att['data'])
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {att["filename"]}'
            )
            msg.attach(part)
        
        # Envoyer
        with smtplib.SMTP(smtp_config['server'], smtp_config['port']) as server:
            server.starttls()
            server.login(smtp_config['username'], smtp_config['password'])
            
            recipients = [email_config['to']]
            if email_config['cc']:
                recipients.extend(email_config['cc'].split(','))
            
            server.send_message(msg)
        
        return True
        
    except Exception as e:
        st.error(f"Erreur envoi email: {e}")
        return False

def export_current_content(format: str) -> bytes:
    """Exporte le contenu actuel dans le format demand√©"""
    # R√©utiliser les fonctions d'export
    analysis = st.session_state.get('current_analysis', {})
    
    if st.session_state.get('redaction_result'):
        content = st.session_state.redaction_result['document']
    else:
        content = "Aucun contenu disponible"
    
    export_functions = {
        'pdf': export_to_pdf,
        'docx': export_to_docx,
        'txt': export_to_txt
    }
    
    export_func = export_functions.get(format, export_to_txt)
    return export_func(content, analysis)

# === FONCTIONS DE TRAITEMENT POUR S√âLECTION DE PI√àCES ===

def process_piece_selection_request(query: str, analysis: dict):
    """Traite une demande de s√©lection de pi√®ces"""
    
    # Interface de s√©lection
    st.markdown("### üìã S√©lection de pi√®ces")
    
    # Collecter les documents disponibles
    available_docs = collect_available_documents(analysis)
    
    if not available_docs:
        st.warning("‚ö†Ô∏è Aucun document disponible")
        return
    
    # Grouper par cat√©gorie
    categories = group_documents_by_category(available_docs)
    
    # Interface de s√©lection par cat√©gorie
    selected_pieces = []
    
    for category, docs in categories.items():
        with st.expander(f"üìÅ {category} ({len(docs)} documents)", expanded=True):
            select_all = st.checkbox(f"Tout s√©lectionner - {category}", key=f"select_all_{category}")
            
            for doc in docs:
                is_selected = st.checkbox(
                    f"üìÑ {doc['title']}",
                    value=select_all,
                    key=f"select_doc_{doc['id']}",
                    help=f"Source: {doc.get('source', 'N/A')}"
                )
                
                if is_selected:
                    selected_pieces.append(PieceSelectionnee(
                        numero=len(selected_pieces) + 1,
                        titre=doc['title'],
                        description=doc.get('description', ''),
                        categorie=category,
                        date=doc.get('date'),
                        source=doc.get('source', ''),
                        pertinence=calculate_piece_relevance(doc, analysis)
                    ))
    
    # Sauvegarder la s√©lection
    st.session_state.selected_pieces = selected_pieces
    
    # Actions sur la s√©lection
    if selected_pieces:
        st.success(f"‚úÖ {len(selected_pieces)} pi√®ces s√©lectionn√©es")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("üìä Cr√©er bordereau", key="create_bordereau_from_selection"):
                process_bordereau_request(query, analysis)
        
        with col2:
            if st.button("üìÑ Synth√©tiser", key="synthesize_selection"):
                synthesize_selected_pieces(selected_pieces)
        
        with col3:
            if st.button("üì§ Exporter liste", key="export_piece_list"):
                export_piece_list(selected_pieces)

def collect_available_documents(analysis: dict) -> list:
    """Collecte tous les documents disponibles"""
    documents = []
    
    # Documents locaux
    for doc_id, doc in st.session_state.get('azure_documents', {}).items():
        documents.append({
            'id': doc_id,
            'title': doc.title,
            'content': doc.content,
            'source': doc.source,
            'metadata': doc.metadata
        })
    
    # Si r√©f√©rence sp√©cifique
    if analysis.get('reference'):
        ref_docs = search_by_reference(f"@{analysis['reference']}")
        documents.extend(ref_docs)
    
    return documents

def group_documents_by_category(documents: list) -> dict:
    """Groupe les documents par cat√©gorie"""
    categories = defaultdict(list)
    
    for doc in documents:
        # D√©terminer la cat√©gorie
        category = determine_document_category(doc)
        categories[category].append(doc)
    
    return dict(categories)

def determine_document_category(doc: dict) -> str:
    """D√©termine la cat√©gorie d'un document"""
    title_lower = doc.get('title', '').lower()
    content_lower = doc.get('content', '')[:500].lower()
    
    # Patterns de cat√©gories
    category_patterns = {
        'Proc√©dure': ['plainte', 'proc√®s-verbal', 'audition', 'perquisition', 'garde √† vue'],
        'Expertise': ['expertise', 'expert', 'rapport technique', 'analyse'],
        'Comptabilit√©': ['bilan', 'compte', 'comptable', 'facture', 'devis'],
        'Contrats': ['contrat', 'convention', 'accord', 'avenant'],
        'Correspondance': ['courrier', 'email', 'lettre', 'mail'],
        'Pi√®ces d\'identit√©': ['carte identit√©', 'passeport', 'kbis', 'statuts'],
        'Bancaire': ['relev√©', 'virement', 'compte bancaire', 'rib']
    }
    
    for category, keywords in category_patterns.items():
        if any(kw in title_lower or kw in content_lower for kw in keywords):
            return category
    
    return 'Autres'

def calculate_piece_relevance(doc: dict, analysis: dict) -> float:
    """Calcule la pertinence d'une pi√®ce"""
    score = 0.5
    
    # Si le document contient des mots-cl√©s de l'analyse
    if analysis.get('subject_matter'):
        if analysis['subject_matter'] in doc.get('content', '').lower():
            score += 0.3
    
    # Si r√©f√©rence dans le titre
    if analysis.get('reference'):
        if analysis['reference'] in doc.get('title', '').lower():
            score += 0.2
    
    return min(score, 1.0)

# === FONCTIONS DE TRAITEMENT POUR BORDEREAU ===

def process_bordereau_request(query: str, analysis: dict):
    """Traite une demande de cr√©ation de bordereau"""
    
    pieces = st.session_state.get('selected_pieces', [])
    
    if not pieces:
        st.warning("‚ö†Ô∏è Aucune pi√®ce s√©lectionn√©e pour le bordereau")
        return
    
    # Cr√©er le bordereau
    bordereau = create_bordereau(pieces, analysis)
    
    # Afficher le bordereau
    st.markdown("### üìä Bordereau de communication de pi√®ces")
    
    # En-t√™te
    st.text_area(
        "En-t√™te du bordereau",
        value=bordereau['header'],
        height=150,
        key="bordereau_header"
    )
    
    # Table des pi√®ces
    if PANDAS_AVAILABLE:
        df = pd.DataFrame([
            {
                'N¬∞': p.numero,
                'Titre': p.titre,
                'Description': p.description,
                'Cat√©gorie': p.categorie,
                'Date': p.date.strftime('%d/%m/%Y') if p.date else 'N/A'
            }
            for p in pieces
        ])
        
        st.dataframe(df, use_container_width=True)
    else:
        # Affichage sans pandas
        for piece in pieces:
            st.write(f"**{piece.numero}.** {piece.titre}")
            if piece.description:
                st.caption(piece.description)
    
    # Actions
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.download_button(
            "üíæ T√©l√©charger bordereau",
            create_bordereau_document(bordereau),
            f"bordereau_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_bordereau"
        ):
            st.success("‚úÖ Bordereau t√©l√©charg√©")
    
    with col2:
        if st.button("‚úèÔ∏è Modifier s√©lection", key="modify_bordereau_selection"):
            st.session_state.show_piece_selection = True
            st.rerun()
    
    with col3:
        if st.button("üìß Envoyer bordereau", key="send_bordereau"):
            st.session_state.universal_query = f"envoyer bordereau @{analysis.get('reference', 'dossier')}"
            st.rerun()
    
    # Stocker le bordereau
    st.session_state.current_bordereau = bordereau

def create_bordereau(pieces: list, analysis: dict) -> dict:
    """Cr√©e un bordereau structur√©"""
    
    bordereau = {
        'header': f"""BORDEREAU DE COMMUNICATION DE PI√àCES
AFFAIRE : {analysis.get('reference', 'N/A').upper()}
DATE : {datetime.now().strftime('%d/%m/%Y')}
NOMBRE DE PI√àCES : {len(pieces)}
POUR : [√Ä compl√©ter]
CONTRE : [√Ä compl√©ter]
PI√àCES COMMUNIQU√âES :""",
        'pieces': pieces,
        'footer': """Je certifie que les pi√®ces communiqu√©es sont conformes aux originaux en ma possession.
Fait √† [Ville], le {datetime.now().strftime('%d/%m/%Y')}
[Signature]""",
        'metadata': {
            'created_at': datetime.now(),
            'piece_count': len(pieces),
            'categories': list(set(p.categorie for p in pieces)),
            'reference': analysis.get('reference')
        }
    }
    
    return bordereau

def create_bordereau_document(bordereau: dict) -> bytes:
    """Cr√©e le document Word du bordereau"""
    if not DOCX_AVAILABLE:
        # Fallback texte
        content = bordereau['header'] + '\n\n'
        
        for piece in bordereau['pieces']:
            content += f"{piece.numero}. {piece.titre}\n"
            if piece.description:
                content += f"   {piece.description}\n"
            content += f"   Cat√©gorie: {piece.categorie}\n"
            if piece.date:
                content += f"   Date: {piece.date.strftime('%d/%m/%Y')}\n"
            content += "\n"
        
        content += bordereau['footer']
        
        return content.encode('utf-8')
    
    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # En-t√™te
        for line in bordereau['header'].split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                if 'BORDEREAU' in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(16)
        
        # Table des pi√®ces
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # En-t√™tes de colonnes
        headers = ['N¬∞', 'Titre', 'Description', 'Cat√©gorie', 'Date']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Lignes de pi√®ces
        for piece in bordereau['pieces']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(piece.numero)
            row_cells[1].text = piece.titre
            row_cells[2].text = piece.description or ''
            row_cells[3].text = piece.categorie
            row_cells[4].text = piece.date.strftime('%d/%m/%Y') if piece.date else 'N/A'
        
        # Pied de page
        doc.add_paragraph()
        for line in bordereau['footer'].split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Sauvegarder
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Erreur cr√©ation bordereau: {e}")
        return create_bordereau_document.__wrapped__(bordereau)  # Fallback

# === FONCTIONS DE TRAITEMENT POUR SYNTH√àSE ===

def process_synthesis_request(query: str, analysis: dict):
    """Traite une demande de synth√®se"""
    
    # D√©terminer la source de la synth√®se
    if st.session_state.get('selected_pieces'):
        content_to_synthesize = synthesize_selected_pieces(st.session_state.selected_pieces)
    
    elif analysis.get('reference'):
        docs = search_by_reference(f"@{analysis['reference']}")
        content_to_synthesize = synthesize_documents(docs)
    
    else:
        st.warning("‚ö†Ô∏è Aucun contenu √† synth√©tiser")
        return
    
    # Stocker le r√©sultat
    st.session_state.synthesis_result = content_to_synthesize

def synthesize_selected_pieces(pieces: list) -> dict:
    """Synth√©tise les pi√®ces s√©lectionn√©es"""
    
    llm_manager = MultiLLMManager()
    if not llm_manager.clients:
        return {'error': 'Aucune IA disponible'}
    
    # Construire le contexte
    context = "PI√àCES √Ä SYNTH√âTISER:\n\n"
    
    for piece in pieces[:20]:  # Limiter
        context += f"Pi√®ce {piece.numero}: {piece.titre}\n"
        context += f"Cat√©gorie: {piece.categorie}\n"
        if piece.description:
            context += f"Description: {piece.description}\n"
        context += "\n"
    
    # Prompt de synth√®se
    synthesis_prompt = f"""{context}
Cr√©e une synth√®se structur√©e de ces pi√®ces.
La synth√®se doit inclure:
1. Vue d'ensemble des pi√®ces
2. Points cl√©s par cat√©gorie
3. Chronologie si applicable
4. Liens et relations entre pi√®ces
5. Points d'attention juridiques
6. Recommandations
Format professionnel avec titres et sous-sections."""
    
    try:
        provider = list(llm_manager.clients.keys())[0]
        response = llm_manager.query_single_llm(
            provider,
            synthesis_prompt,
            "Tu es un expert en analyse et synth√®se de documents juridiques."
        )
        
        if response['success']:
            return {
                'content': response['response'],
                'piece_count': len(pieces),
                'categories': list(set(p.categorie for p in pieces)),
                'timestamp': datetime.now()
            }
        else:
            return {'error': '√âchec de la synth√®se'}
            
    except Exception as e:
        return {'error': f'Erreur synth√®se: {str(e)}'}

def synthesize_documents(documents: list) -> dict:
    """Synth√©tise une liste de documents"""
    # Convertir en pi√®ces pour r√©utiliser la fonction
    pieces = []
    
    for i, doc in enumerate(documents):
        pieces.append(PieceSelectionnee(
            numero=i + 1,
            titre=doc.get('title', 'Sans titre'),
            description=doc.get('content', '')[:200] + '...' if doc.get('content') else '',
            categorie=determine_document_category(doc),
            source=doc.get('source', '')
        ))
    
    return synthesize_selected_pieces(pieces)

# === FONCTIONS DE TRAITEMENT POUR TEMPLATES ===

def process_template_request(query: str, analysis: dict):
    """Traite une demande li√©e aux templates"""
    
    action = analysis['details'].get('action', 'apply')
    
    if action == 'create':
        create_new_template()
    
    elif action == 'edit':
        edit_existing_template()
    
    else:  # apply
        apply_template()

def create_new_template():
    """Cr√©e un nouveau template"""
    
    template_name = st.session_state.get('new_template_name', '')
    
    if not template_name:
        st.warning("‚ö†Ô∏è Nom du template requis")
        return
    
    # Base du template
    base_template = st.session_state.get('base_template', 'Vide')
    
    if base_template == 'Vide':
        template_content = {
            'name': template_name,
            'structure': [],
            'style': 'formel',
            'category': st.session_state.get('template_category', 'Autre')
        }
    else:
        # Copier depuis un template existant
        template_content = DOCUMENT_TEMPLATES.get(base_template, {}).copy()
        template_content['name'] = template_name
    
    # √âditeur de structure
    st.markdown("### üìù Structure du template")
    
    structure = st.text_area(
        "Sections (une par ligne)",
        value='\n'.join(template_content.get('structure', [])),
        height=300,
        key="template_structure_editor"
    )
    
    template_content['structure'] = [s.strip() for s in structure.split('\n') if s.strip()]
    
    # Sauvegarder
    if st.button("üíæ Sauvegarder le template", key="save_new_template"):
        if 'saved_templates' not in st.session_state:
            st.session_state.saved_templates = {}
        
        st.session_state.saved_templates[clean_key(template_name)] = template_content
        st.success(f"‚úÖ Template '{template_name}' sauvegard√©")
        
        # Optionnel : sauvegarder dans un fichier
        save_templates_to_file()

def edit_existing_template():
    """√âdite un template existant"""
    
    template_to_edit = st.session_state.get('template_to_edit')
    
    if not template_to_edit:
        st.warning("‚ö†Ô∏è S√©lectionnez un template √† modifier")
        return
    
    # Charger le template
    if template_to_edit in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[template_to_edit].copy()
        is_builtin = True
    else:
        template = st.session_state.saved_templates.get(template_to_edit, {})
        is_builtin = False
    
    if not template:
        st.error("‚ùå Template introuvable")
        return
    
    # √âditeur
    st.markdown(f"### ‚úèÔ∏è √âdition du template '{template.get('name', template_to_edit)}'")
    
    if is_builtin:
        st.info("‚ÑπÔ∏è Template int√©gr√© - Les modifications seront sauvegard√©es comme nouveau template")
    
    # Nom
    new_name = st.text_input(
        "Nom du template",
        value=template.get('name', ''),
        key="edit_template_name"
    )
    
    # Structure
    structure = st.text_area(
        "Structure",
        value='\n'.join(template.get('structure', [])),
        height=300,
        key="edit_template_structure"
    )
    
    # Style
    style = st.selectbox(
        "Style par d√©faut",
        list(REDACTION_STYLES.keys()),
        index=list(REDACTION_STYLES.keys()).index(template.get('style', 'formel')),
        format_func=lambda x: REDACTION_STYLES[x]['name'],
        key="edit_template_style"
    )
    
    # Sauvegarder les modifications
    if st.button("üíæ Sauvegarder les modifications", key="save_template_edits"):
        updated_template = {
            'name': new_name,
            'structure': [s.strip() for s in structure.split('\n') if s.strip()],
            'style': style,
            'category': template.get('category', 'Autre')
        }
        
        if is_builtin:
            # Sauvegarder comme nouveau
            st.session_state.saved_templates[clean_key(new_name)] = updated_template
            st.success(f"‚úÖ Nouveau template '{new_name}' cr√©√©")
        else:
            # Mettre √† jour l'existant
            st.session_state.saved_templates[template_to_edit] = updated_template
            st.success(f"‚úÖ Template '{new_name}' mis √† jour")

def apply_template():
    """Applique un template s√©lectionn√©"""
    
    selected_template = st.session_state.get('selected_template')
    
    if not selected_template:
        st.warning("‚ö†Ô∏è S√©lectionnez un template √† appliquer")
        return
    
    # Charger le template
    if selected_template in DOCUMENT_TEMPLATES:
        template = DOCUMENT_TEMPLATES[selected_template]
    else:
        template = st.session_state.saved_templates.get(selected_template, {})
    
    if not template:
        st.error("‚ùå Template introuvable")
        return
    
    # Cr√©er une requ√™te de r√©daction avec le template
    st.session_state.universal_query = f"r√©diger {template.get('name', 'document')} avec template {selected_template}"
    
    # D√©finir le style
    st.session_state.redaction_style = template.get('style', 'formel')
    
    # D√©clencher la r√©daction
    st.info(f"‚úÖ Template '{template.get('name')}' appliqu√© - Lancez la r√©daction")

def save_templates_to_file():
    """Sauvegarde les templates dans un fichier"""
    try:
        import json
        
        templates_data = {
            'builtin': DOCUMENT_TEMPLATES,
            'custom': st.session_state.get('saved_templates', {})
        }
        
        # Cr√©er un fichier t√©l√©chargeable
        json_str = json.dumps(templates_data, indent=2, ensure_ascii=False)
        
        st.download_button(
            "üíæ Exporter tous les templates",
            json_str,
            f"templates_{datetime.now().strftime('%Y%m%d')}.json",
            "application/json",
            key="export_templates"
        )
        
    except Exception as e:
        st.error(f"Erreur sauvegarde templates: {e}")

# === FONCTIONS DE TRAITEMENT POUR JURISPRUDENCE ===

def process_jurisprudence_request(query: str, analysis: dict):
    """Traite une demande de recherche de jurisprudence"""
    